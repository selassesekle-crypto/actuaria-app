"""
╔══════════════════════════════════════════════════════════════════════════════╗
║  ACTUARIA — AGENT RAPPORT SANTÉ : RAPPORT ACTUARIEL SANTÉ                  ║
║  Direction Santé-Prévoyance · Sous-direction Santé                         ║
╠══════════════════════════════════════════════════════════════════════════════╣
║                                                                              ║
║  RÔLE : Rapport actuariel consolidé Santé — livrable professionnel          ║
║         pour CA, ACPR, commissaire aux comptes, actuaire désigné.           ║
║                                                                              ║
║  PIPELINE M1 → M5 :                                                         ║
║  M1 — Ingestion & validation                                                ║
║       Agrège S1 (tarification), S2 (provisionnement), S3 (reporting)        ║
║       + SP-REG2 (IFRS 17) + SP-REG3 (ANI/100%S) si disponibles            ║
║                                                                              ║
║  M2 — Hypothèses actuarielles Santé                                         ║
║       DREES 2023, FNMF 2023, ANI 2013, EIOPA CoC 6%                        ║
║                                                                              ║
║  M3 — Calculs actuariels Santé                                              ║
║       LR par poste, PSAP/IBNR, BE, RA (CoC 6%), SCR NSLT, MCR              ║
║                                                                              ║
║  M4 — Synthèse & avis actuariel                                             ║
║       RAG consolidé, avis FAVORABLE / AVEC RÉSERVES / DÉFAVORABLE          ║
║                                                                              ║
║  M5 — Livrables professionnels                                              ║
║       · Rapport HTML premium Navy/Gold                                       ║
║       · Rapport PDF (WeasyPrint — couleurs forcées)                         ║
║       · Rapport Word (.docx)                                                 ║
║       · Classeur Excel 4 onglets                                             ║
║       · Narration Claude API (actuaire Santé senior)                        ║
║                                                                              ║
║  ENTRÉES (toutes optionnelles sauf S1 ou S2) :                             ║
║    result_s1   → S1 Léonie (tarification santé)                            ║
║    result_s2   → S2 Selma  (provisionnement santé)                         ║
║    result_s3   → S3 Binta  (reporting S2/QRT)                              ║
║    result_reg2 → SP-REG2   (IFRS 17)                                       ║
║    result_reg3 → SP-REG3   (ANI + 100% Santé)                              ║
║                                                                              ║
║  VERSION : 1.0 — Juillet 2026                                               ║
╚══════════════════════════════════════════════════════════════════════════════╝
"""

import base64
import hashlib
import io
import json
import logging
import re
import warnings
from datetime import datetime
from pathlib import Path
from typing import Dict, Optional, Tuple

from core import frontiere_llm
from core import traitement_ia

import numpy as np

warnings.filterwarnings("ignore")
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(name)s | %(levelname)s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)

try:
    import plotly.graph_objects as go
    PLOTLY_OK = True
except ImportError:
    PLOTLY_OK = False

# ── Palette ActuarIA ──────────────────────────────────────────────────────────
NAVY      = "#0B1E3D"
NAVY_L    = "#1E3A5F"
GOLD      = "#C9A84C"
ROUGE     = "#C0392B"
ORANGE    = "#E67E22"
VERT      = "#1E8449"
SLATE     = "#8A9BB0"
BLANC     = "#F0F4F8"
BLEU      = "#3498DB"
VIOLET    = "#9B59B6"

# ── Paramètres réglementaires ─────────────────────────────────────────────────
SCR_SIGMA_PREM   = 0.05    # σ primes santé NSLT — Art.148 RD 2015/35
SCR_SIGMA_RES    = 0.14    # σ réserves santé NSLT
MCR_COEFF_PREM   = 0.0453  # coefficient primes MCR santé
MCR_COEFF_RES    = 0.0351  # coefficient provisions MCR santé
MCR_PLANCHER     = 2_500_000.0  # plancher absolu Art.129 S2
COC_RATE         = 0.06    # EIOPA CoC rate — IFRS 17 §B91
SEUIL_SCR        = 100.0   # seuil minimal % — Art.129 S2
SEUIL_SCR_CIBLE  = 130.0   # seuil cible interne

# ── Logo SVG ──────────────────────────────────────────────────────────────────
LOGO_SVG = (
    '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 110" width="400" height="110">'
    '<rect width="400" height="110" fill="#0F2E52" rx="8"/>'
    '<g transform="translate(42,55)">'
    '<circle cx="0" cy="-27" r="5" fill="none" stroke="#4A8FD4" stroke-width="1.2"/>'
    '<circle cx="0" cy="-9"  r="5" fill="none" stroke="#4A8FD4" stroke-width="1.2"/>'
    '<circle cx="0" cy="9"   r="5" fill="none" stroke="#4A8FD4" stroke-width="1.2"/>'
    '<circle cx="0" cy="27"  r="5" fill="none" stroke="#4A8FD4" stroke-width="1.2"/>'
    '<circle cx="26" cy="-16" r="5" fill="none" stroke="#C9A84C" stroke-width="1.2"/>'
    '<circle cx="26" cy="4"   r="5" fill="none" stroke="#C9A84C" stroke-width="1.2"/>'
    '<circle cx="26" cy="24"  r="5" fill="none" stroke="#C9A84C" stroke-width="1.2"/>'
    '<circle cx="52" cy="4"   r="8" fill="none" stroke="#C9A84C" stroke-width="1.5"/>'
    '<text x="52" y="8" text-anchor="middle" font-family="Georgia,serif" font-size="8" font-weight="700" fill="#C9A84C">A</text>'
    '</g>'
    '<text x="122" y="50" font-family="Georgia,serif" font-size="38" font-weight="700">'
    '<tspan fill="#FFFFFF">Actuar</tspan><tspan fill="#C9A84C">IA</tspan>'
    '</text>'
    '<line x1="122" y1="58" x2="328" y2="58" stroke="#C9A84C" stroke-width="0.6" opacity="0.5"/>'
    '<text x="122" y="73" font-family="Georgia,serif" font-size="7.5" fill="#8A9AB0" letter-spacing="3">ACTUARIAL INTELLIGENCE PLATFORM</text>'
    '</svg>'
)
LOGO_URI = "data:image/svg+xml;base64," + base64.b64encode(LOGO_SVG.encode()).decode()


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _f(v, dec=0) -> str:
    if v is None:
        return "—"
    try:
        fv = float(v)
        if not np.isfinite(fv):
            return "—"
        sep = "\u202f"
        return f"{fv:,.0f}\u202f€".replace(",", sep) if dec == 0 else f"{fv:,.{dec}f}".replace(",", sep)
    except Exception:
        return "—"

def _pct(v, dec=1) -> str:
    if v is None:
        return "—"
    try:
        fv = float(v)
        return "—" if not np.isfinite(fv) else f"{fv:.{dec}f}\u202f%"
    except Exception:
        return "—"

def _clean(txt) -> str:
    if not txt:
        return ""
    txt = re.sub(r"[■□▪▸►═╔╗╚╝║─]+", "", str(txt))
    txt = re.sub(r"={4,}", "", txt)
    txt = re.sub(r"\n{3,}", "\n\n", txt)
    return txt.strip()

def _statut_col(s: str) -> str:
    return {"VERT": VERT, "AMBRE": ORANGE, "ROUGE": ROUGE}.get(s, SLATE)


# ══════════════════════════════════════════════════════════════════════════════
# SYSTEM PROMPT CLAUDE API — ACTUAIRE SANTÉ SENIOR
# ══════════════════════════════════════════════════════════════════════════════

SYSTEM_PROMPT_SANTE = """\
Tu es un actuaire Santé senior certifié par l'Institut des Actuaires (IA), \
expert en complémentaire santé collective et individuelle, avec 20 ans d'expérience \
auprès de mutuelles (Livre II Code de la Mutualité), institutions de prévoyance \
et assureurs mixtes.

Tu rédiges le commentaire actuariel d'un rapport de santé destiné à être présenté \
au Conseil d'Administration et à l'actuaire désigné. Ce rapport couvre la tarification, \
le provisionnement et le reporting réglementaire Solvabilité II + IFRS 17 \
pour un portefeuille de complémentaire santé.

RÈGLES ABSOLUES :
0. FORMAT : §N — TITRE pour les sections, ### pour les sous-titres, \
**gras** pour les termes importants, - tirets pour les listes. \
PAS de tableaux Markdown. PAS de blockquotes >. Sépare les sections par une ligne vide.
1. LANGUE : Français professionnel. Anglais uniquement pour les termes consacrés (BE, RA, SCR, MCR, IBNR).
2. RIGUEUR : Chaque affirmation est justifiée par les données fournies.
3. CHIFFRES : En euros avec séparateurs (ex : 2\u202f526\u202f597\u202f€). Pourcentages avec une décimale.
4. RÉFÉRENCES : DREES 2023, FNMF 2023, ANI 2013, Art.148 RD 2015/35 (SCR NSLT), \
IFRS 17 §B91 (CoC 6%), Art.252 RD 2015/35 (MCR), Art.129 S2 (seuils).
5. ALERTES : Ne jamais minimiser. Présenter avec l'implication réelle pour la solvabilité.
6. POSTURE : Assertif mais prudent. Recommandations claires avec justification.
7. INTERDIT : Phrases génériques sans données. Toujours lier chaque observation aux chiffres.

STRUCTURE OBLIGATOIRE EN 7 SECTIONS :
§1 — CONTEXTE ET QUALITÉ DU PORTEFEUILLE
§2 — SINISTRALITÉ PAR POSTE DE GARANTIE
§3 — PROVISIONNEMENT PSAP / IBNR
§4 — BEST ESTIMATE ET RISK ADJUSTMENT (IFRS 17 §B91)
§5 — SCR SANTÉ NSLT ET RATIO DE SOLVABILITÉ
§6 — CONFORMITÉ ANI 2013 ET 100% SANTÉ
§7 — CONCLUSION ET RECOMMANDATIONS POUR LE CONSEIL D'ADMINISTRATION\
"""


# ══════════════════════════════════════════════════════════════════════════════
# CONTEXTE CLAUDE API
# ══════════════════════════════════════════════════════════════════════════════

def _construire_contexte_sante(m1: Dict, m3: Dict, date_arrete: str) -> str:
    postes    = m3.get("postes", {})
    hyp       = m3.get("hypotheses_consolidees", [])

    lines = [
        f"DOSSIER SANTÉ — Arrêté {date_arrete}",
        "(Identité de l'organisme non transmise : ne jamais la nommer ni l'inventer.)",
        "",
        "=== PORTEFEUILLE ===",
        f"Assurés : {m1.get('nb_assures', '—'):,}",
        f"Âge moyen : {m1.get('age_moyen', '—')} ans",
        f"Type contrat : {m1.get('contrat', '—')}",
        f"Niveau garantie : {m1.get('garantie_niveau', '—')}",
        f"Source données : {m1.get('source_donnees', '—')}",
        "",
        "=== TARIFICATION (S1) ===",
        f"Prime pure unitaire : {_f(m3.get('prime_pure'))}",
        f"Prime commerciale : {_f(m3.get('prime_commerciale'))}",
        f"Primes acquises totales : {_f(m3.get('primes_acquises'))}",
        f"Loss Ratio attendu (tarification) : {_pct(m3.get('lr_attendu', 0) * 100)}",
        f"Référence marché FNMF 2023 : {m1.get('lr_fnmf_ref', '65%-75% garantie confort')}",
        "",
        "=== SINISTRALITÉ PAR POSTE ===",
    ]
    for p, v in postes.items():
        lines.append(
            f"  {p} : sin={_f(v.get('sinistre_annuel'))} | "
            f"IBNR={_pct(v.get('taux_ibnr', 0) * 100)} | "
            f"source={v.get('source', '—')}"
        )
    lines += [
        "",
        "=== PROVISIONNEMENT (S2) ===",
        f"PSAP dossiers connus : {_f(m3.get('psap_dossiers'))}",
        f"PSAP IBNR (cadence santé 1-3 mois) : {_f(m3.get('psap_ibnr'))}",
        f"PSAP Total : {_f(m3.get('psap_total'))}",
        f"PREC : {_f(m3.get('prec'))}",
        f"Provision totale : {_f(m3.get('provision_totale'))}",
        f"Loss Ratio observé : {_pct(m3.get('loss_ratio', 0) * 100)}",
        f"A/E ratio (sinistres réels/attendus) : {m3.get('ae_ratio', '—')}",
        "",
        "=== BEST ESTIMATE & RISK ADJUSTMENT (IFRS 17) ===",
        f"BE Santé : {_f(m3.get('be_sante'))}",
        f"Risk Adjustment (CoC 6% EIOPA §B91) : {_f(m3.get('risk_adjustment'))}",
        f"TP Santé (BE + RA) : {_f(m3.get('tp_sante'))}",
        f"Ratio TP/BE : {m3.get('ratio_tp_be', 0):.3f} (cible [1.00 — 1.50])",
        "",
        "=== SCR / MCR (SOLVABILITÉ II) ===",
        f"SCR Santé NSLT (Art.148 RD 2015/35) : {_f(m3.get('scr_sante'))}",
        f"  dont SCR primes (σ=5%) : {_f(m3.get('scr_prem'))}",
        f"  dont SCR réserves (σ=14%) : {_f(m3.get('scr_res'))}",
        f"  dont SCR catastrophe santé : {_f(m3.get('scr_cat'))}",
        f"MCR Santé (Art.252 RD 2015/35) : {_f(m3.get('mcr_sante'))}",
        f"Fonds propres éligibles : {_f(m3.get('fonds_propres'))}",
        f"Ratio SCR : {_pct(m3.get('ratio_scr_pct'))}",
        f"Ratio MCR : {_pct(m3.get('ratio_mcr_pct'))}",
        "",
        "=== CONFORMITÉ ANI 2013 ===",
        f"ANI 2013 conforme : {m3.get('ani_conforme', '—')}",
        f"Détail : {m3.get('ani_detail_str', '—')}",
        "",
        "=== HYPOTHÈSES ===",
    ]
    for h in hyp:
        lines.append(
            f"  {h['id']} — {h['hypothese']} : {h['statut']} | {h.get('valeur', '')[:120]}"
        )
    lines += [
        "",
        "Rédige le commentaire actuariel complet en 7 sections.",
    ]
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
# NARRATION CLAUDE API (3 niveaux)
# ══════════════════════════════════════════════════════════════════════════════

def _narration_claude_api(m1: Dict, m3: Dict, date_arrete: str) -> str:
    try:
        ctx = _construire_contexte_sante(m1, m3, date_arrete)
        resp = frontiere_llm.appeler(
            modele=frontiere_llm.MODELE_ETABLI,
            max_tokens=10000,
            systeme=SYSTEM_PROMPT_SANTE,
            messages=[{"role": "user", "content": ctx}],
            cle=frontiere_llm.cle_api_ou_secrets(),
        )
        return frontiere_llm.texte_des_blocs(resp)
    except Exception as e:
        logging.getLogger("actuaria.sp.rapport_sante").warning(f"Narration non générée : {e}")
        raise

def _narration_fallback(m3: Dict, commentaire: str) -> str:
    return _clean(commentaire) or _clean(m3.get("commentaire_consolide", ""))

def _generer_narration(m1, m3, date_arrete, commentaire) -> Tuple[str, str]:
    try:
        txt = _narration_claude_api(m1, m3, date_arrete)
        if txt:
            return txt, "claude_api"
    except Exception:
        pass
    txt = _narration_fallback(m3, commentaire)
    if txt:
        return txt, "fallback"
    return "", "aucune"


# ══════════════════════════════════════════════════════════════════════════════
# MARKDOWN → HTML
# ══════════════════════════════════════════════════════════════════════════════

def _md_to_html(texte: str) -> str:
    if not texte:
        return '<p style="color:#8A9BB0;font-style:italic;">Narration non disponible.</p>'
    txt = texte.strip()
    # §N — TITRE
    txt = re.sub(
        r"(§\d+\s*[—\-–]\s*[^\n]+)",
        lambda m: '<div class="comm-section-title">' + m.group(1).strip() + "</div>",
        txt,
    )
    # ### sous-titres
    txt = re.sub(
        r"^###\s+(.+)$",
        lambda m: '<div class="comm-h4">' + m.group(1).strip() + "</div>",
        txt, flags=re.MULTILINE,
    )
    # **gras**
    txt = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", txt, flags=re.DOTALL)
    # *italique*
    txt = re.sub(r"\*(.+?)\*", r"<em>\1</em>", txt)
    # Puces
    txt = re.sub(r"^[-•]\s+(.+)$", r"<li>\1</li>", txt, flags=re.MULTILINE)
    txt = re.sub(
        r"(<li>.*?</li>\n?)+",
        lambda m: '<ul class="comm-ul">' + m.group(0) + "</ul>",
        txt, flags=re.DOTALL,
    )
    # Paragraphes
    result = ""
    for bloc in re.split(r"\n{2,}", txt):
        bloc = bloc.strip()
        if not bloc:
            continue
        if bloc.startswith("<"):
            result += bloc + "\n"
        else:
            result += '<p class="comm-p">' + bloc.replace("\n", " ") + "</p>\n"
    return result or '<p class="comm-p">' + texte + "</p>"


# ══════════════════════════════════════════════════════════════════════════════
# CSS PREMIUM (identique Non-Vie + print-color-adjust pour PDF)
# ══════════════════════════════════════════════════════════════════════════════

def _css() -> str:
    return """
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,700;0,900;1,400&family=Inter:wght@300;400;500;600;700&display=swap');

/* ── FORCER LES COULEURS DE FOND EN PDF (WeasyPrint) ── */
* {
  -webkit-print-color-adjust: exact !important;
  print-color-adjust: exact !important;
  box-sizing: border-box;
  margin: 0;
  padding: 0;
}

:root {
  --navy:       #0B1E3D;
  --navy-mid:   #132844;
  --navy-light: #1E3A5F;
  --gold:       #C9A84C;
  --gold-light: #E2C97E;
  --gold-pale:  rgba(201,168,76,0.12);
  --rouge:      #C0392B;
  --rouge-pale: rgba(192,57,43,0.08);
  --orange:     #E67E22;
  --vert:       #1E8449;
  --slate:      #8A9BB0;
  --bg:         #F5F7FA;
  --white:      #FFFFFF;
  --text:       #1C2B3A;
  --text-mid:   #3D5166;
  --border:     #DDE4EE;
  --bleu:       #3498DB;
}

body {
  font-family: 'Inter', sans-serif;
  font-size: 10pt;
  color: var(--text);
  background: var(--bg);
  line-height: 1.7;
}

.rapport-container {
  max-width: 960px;
  margin: 0 auto;
  background: var(--white);
  box-shadow: 0 4px 40px rgba(11,30,61,0.14);
}

/* ── PAGE DE GARDE ── */
.page-garde {
  position: relative;
  min-height: 680px;
  display: flex;
  flex-direction: column;
  overflow: hidden;
  background: var(--navy) !important;
  page-break-after: always;
}
.garde-dots {
  position: absolute;
  top: 0; right: 0;
  width: 320px; height: 100%;
  background-image: radial-gradient(circle, rgba(201,168,76,0.25) 1px, transparent 1px);
  background-size: 22px 22px;
  -webkit-mask-image: linear-gradient(to left, rgba(0,0,0,0.7) 0%, transparent 100%);
  mask-image: linear-gradient(to left, rgba(0,0,0,0.7) 0%, transparent 100%);
}
.garde-diagonal {
  position: absolute;
  top:0; left:0; right:0; bottom:0;
  background: linear-gradient(135deg, transparent 0%, transparent 62%, rgba(201,168,76,0.06) 62%, rgba(201,168,76,0.06) 100%);
}
.garde-accent-line {
  position: absolute;
  left: 52px; top: 0; bottom: 0;
  width: 2px;
  background: linear-gradient(to bottom, transparent 0%, var(--gold) 20%, var(--gold) 80%, transparent 100%);
  opacity: 0.5;
}
.garde-inner {
  position: relative;
  z-index: 2;
  display: flex;
  flex-direction: column;
  height: 100%;
  padding: 0;
}
.garde-header {
  display: flex;
  justify-content: space-between;
  align-items: flex-start;
  padding: 36px 56px 0;
}
.garde-logo-wrap img { height: 52px; }
.badge-confidentiel {
  font-size: 7.5pt;
  font-weight: 700;
  color: var(--rouge);
  letter-spacing: 2.5px;
  border: 1.5px solid var(--rouge);
  padding: 5px 12px;
  text-transform: uppercase;
}
.garde-hero {
  flex: 1;
  display: flex;
  flex-direction: column;
  justify-content: center;
  padding: 48px 56px 0 80px;
}
.garde-eyebrow {
  font-size: 8pt;
  font-weight: 600;
  color: var(--gold);
  letter-spacing: 4px;
  text-transform: uppercase;
  margin-bottom: 20px;
}
.garde-titre {
  font-family: 'Playfair Display', serif;
  font-size: 48pt;
  font-weight: 900;
  color: var(--white);
  line-height: 1.05;
  letter-spacing: -1.5px;
  margin-bottom: 8px;
}
.garde-titre em { font-style: italic; color: var(--gold); }
.garde-subtitle {
  font-family: 'Playfair Display', serif;
  font-size: 16pt;
  color: var(--gold-light);
  margin-bottom: 40px;
}
.garde-meta {
  display: flex;
  flex-wrap: wrap;
  gap: 12px 32px;
  padding: 28px 80px;
  border-top: 1px solid rgba(201,168,76,0.2);
  margin-top: auto;
}
.garde-meta-item {
  display: flex;
  flex-direction: column;
  gap: 2px;
}
.garde-meta-label {
  font-size: 7pt;
  color: var(--slate);
  letter-spacing: 2px;
  text-transform: uppercase;
}
.garde-meta-val {
  font-size: 10pt;
  font-weight: 600;
  color: var(--white);
}
.garde-statut-badge {
  display: inline-flex;
  align-items: center;
  gap: 8px;
  padding: 8px 20px;
  border-radius: 3px;
  font-size: 9pt;
  font-weight: 700;
  letter-spacing: 1px;
  margin: 16px 80px;
}

/* ── SECTIONS ── */
.section {
  padding: 40px 56px;
  border-bottom: 1px solid var(--border);
}
.section-header {
  display: flex;
  align-items: center;
  gap: 16px;
  margin-bottom: 24px;
  padding-bottom: 12px;
  border-bottom: 2px solid var(--gold);
}
.section-num {
  width: 36px; height: 36px;
  background: var(--navy) !important;
  color: var(--gold);
  font-family: 'Playfair Display', serif;
  font-size: 14pt;
  font-weight: 700;
  display: flex;
  align-items: center;
  justify-content: center;
  flex-shrink: 0;
}
.section-title {
  font-family: 'Playfair Display', serif;
  font-size: 15pt;
  font-weight: 700;
  color: var(--navy);
}

/* ── TABLEAUX ── */
.kpi-grid {
  display: grid;
  grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
  gap: 16px;
  margin-bottom: 24px;
}
.kpi-card {
  background: var(--bg) !important;
  border: 1px solid var(--border);
  border-top: 3px solid var(--gold);
  padding: 16px 20px;
}
.kpi-label {
  font-size: 7.5pt;
  color: var(--slate);
  letter-spacing: 1.5px;
  text-transform: uppercase;
  margin-bottom: 6px;
}
.kpi-val {
  font-size: 18pt;
  font-weight: 700;
  color: var(--navy);
  font-family: 'Playfair Display', serif;
}
.kpi-sub { font-size: 8pt; color: var(--slate); margin-top: 4px; }

.data-table {
  width: 100%;
  border-collapse: collapse;
  font-size: 9pt;
  margin: 16px 0;
}
.data-table th {
  background: var(--navy) !important;
  color: var(--white);
  padding: 10px 14px;
  text-align: left;
  font-weight: 600;
  letter-spacing: 0.5px;
  font-size: 8.5pt;
}
.data-table td {
  padding: 9px 14px;
  border-bottom: 1px solid var(--border);
  color: var(--text);
}
.data-table tr:nth-child(even) td { background: #F8FAFB !important; }
.data-table tr:last-child td { font-weight: 600; background: var(--gold-pale) !important; }

/* ── HYPOTHÈSES ── */
.hyp-row {
  display: flex;
  align-items: flex-start;
  gap: 14px;
  padding: 12px 16px;
  border-left: 3px solid transparent;
  margin-bottom: 8px;
  background: var(--bg) !important;
}
.hyp-row.validee  { border-color: var(--vert);   background: rgba(30,132,73,0.05) !important; }
.hyp-row.justif   { border-color: var(--orange);  background: rgba(230,126,34,0.06) !important; }
.hyp-row.invalide { border-color: var(--rouge);   background: rgba(192,57,43,0.06) !important; }
.hyp-id { font-weight: 700; font-size: 9pt; min-width: 30px; color: var(--navy); }
.hyp-label { font-size: 9pt; font-weight: 600; color: var(--navy); }
.hyp-valeur { font-size: 8.5pt; color: var(--text-mid); margin-top: 2px; }

/* ── NARRATION ── */
.narration-block {
  background: var(--white) !important;
  border: 1px solid var(--border);
  border-left: 4px solid var(--gold);
  padding: 28px 32px;
  margin: 8px 0;
}
.comm-section-title {
  font-family: 'Playfair Display', serif;
  font-size: 12pt;
  font-weight: 700;
  color: var(--navy);
  margin: 20px 0 10px;
  padding-bottom: 6px;
  border-bottom: 1px solid var(--gold-pale);
}
.comm-h4 {
  font-size: 10pt;
  font-weight: 600;
  color: var(--navy);
  margin: 14px 0 6px;
}
.comm-p {
  font-size: 9.5pt;
  color: var(--text-mid);
  line-height: 1.75;
  margin-bottom: 10px;
}
.comm-ul {
  padding-left: 20px;
  margin: 8px 0 12px;
}
.comm-ul li {
  font-size: 9.5pt;
  color: var(--text-mid);
  margin-bottom: 5px;
}

/* ── AVIS ACTUARIEL ── */
.avis-block {
  padding: 20px 28px;
  margin: 24px 0;
  border-radius: 2px;
}
.avis-favorable  { background: rgba(30,132,73,0.08) !important; border-left: 4px solid var(--vert); }
.avis-reserves   { background: rgba(230,126,34,0.08) !important; border-left: 4px solid var(--orange); }
.avis-defavorable{ background: rgba(192,57,43,0.08) !important; border-left: 4px solid var(--rouge); }
.avis-titre { font-size: 11pt; font-weight: 700; margin-bottom: 6px; }

/* ── PIED DE PAGE ── */
.pied-de-page {
  background: var(--navy) !important;
  display: flex;
  align-items: center;
  gap: 24px;
  padding: 20px 56px;
}
.pied-de-page img { height: 32px; opacity: 0.85; }
.pied-meta {
  font-size: 7.5pt;
  color: var(--slate);
  line-height: 1.7;
  flex: 1;
}
.confidentiel-footer {
  font-weight: 700;
  color: var(--rouge);
  letter-spacing: 1px;
}

@media print {
  * { -webkit-print-color-adjust: exact !important; print-color-adjust: exact !important; }
  .page-garde { background: var(--navy) !important; }
  .section-num { background: var(--navy) !important; }
}
</style>
"""


# ══════════════════════════════════════════════════════════════════════════════
# EXPORT HTML
# ══════════════════════════════════════════════════════════════════════════════

# /!\ UN SEUL LIBELLE POUR LES DEUX FORMATS. L'HTML ne nommait sa source QU'EN
# CAS DE SUCCES et le Word ne la nommait JAMAIS -- on ne lui passait meme pas
# l'information. Un repli etait donc indiscernable d'une narration reussie
# dans les deux livrables signes. Meme correction que le lot T1 en tarification,
# ou le raisonnement est ecrit en toutes lettres.
def _badge_narration_sante(source: str) -> str:
    """Le libelle d'origine publie au lecteur, identique en HTML et en Word."""
    libelle = {
        "claude_api": "✦ Narration générée par ActuarIA Intelligence",
        "fallback": "📝 Mode standard",
    }.get(source, "")
    return traitement_ia.avec_engagement(libelle)


def _export_html_sante(m1: Dict, m3: Dict, m4: Dict,
                       narration: str, source_narration: str,
                       entite: str, date_arrete: str) -> str:
    logger = logging.getLogger("actuaria.sp.rapport_sante")
    try:
        dt       = datetime.now().strftime("%d/%m/%Y")
        rag      = m4.get("statut_rag", "AMBRE")
        avis     = m4.get("avis_actuariel", "AVEC RÉSERVES")
        col_rag  = _statut_col(rag)
        ic_rag   = "🟢" if rag == "VERT" else ("🟡" if rag == "AMBRE" else "🔴")

        # Avis CSS
        avis_cls = "avis-favorable" if avis == "FAVORABLE" else ("avis-defavorable" if avis == "DÉFAVORABLE" else "avis-reserves")

        # Hypothèses
        hyp_html = ""
        for h in m4.get("hypotheses", []):
            st = h.get("statut", "")
            cls = "validee" if st == "VALIDÉE" else ("justif" if st == "À JUSTIFIER" else "invalide")
            ic  = "✅" if st == "VALIDÉE" else ("⚠️" if st == "À JUSTIFIER" else "❌")
            hyp_html += (
                f'<div class="hyp-row {cls}">'
                f'<span class="hyp-id">{h["id"]}</span>'
                f'<div>'
                f'<div class="hyp-label">{ic} {h["hypothese"]}</div>'
                f'<div class="hyp-valeur">→ {h.get("valeur", "")} : <strong>{st}</strong></div>'
                f'</div></div>\n'
            )

        # Tableau sinistralité par poste
        postes = m3.get("postes", {})
        postes_rows = ""
        for p, v in postes.items():
            postes_rows += (
                f"<tr><td>{p.replace('_',' ').title()}</td>"
                f"<td>{_f(v.get('sinistre_annuel'))}</td>"
                f"<td>{_pct(v.get('taux_ibnr', 0) * 100)}</td>"
                f"<td>{_f(v.get('ibnr', 0))}</td>"
                f"<td>{v.get('source', '—')}</td></tr>\n"
            )

        # Narration
        narration_html = _md_to_html(narration) if narration else (
            '<p style="color:#8A9BB0;font-style:italic;">Narration non disponible — vérifier la clé API.</p>'
        )
        source_badge = _badge_narration_sante(source_narration)
        if source_badge:
            source_badge = ('<span style="font-size:7pt;color:#8A9BB0;'
                            'font-style:italic;">' + source_badge + '</span>')

        html = (
            "<!DOCTYPE html>\n<html lang='fr'>\n<head>\n"
            "<meta charset='UTF-8'/>\n"
            "<meta name='viewport' content='width=device-width, initial-scale=1.0'/>\n"
            f"<title>Rapport Santé — {entite} — {date_arrete}</title>\n"
            + _css() +
            "</head>\n<body>\n"
            '<div class="rapport-container">\n'

            # ── PAGE DE GARDE ──
            '<div class="page-garde">\n'
            '  <div class="garde-dots"></div>\n'
            '  <div class="garde-diagonal"></div>\n'
            '  <div class="garde-accent-line"></div>\n'
            '  <div class="garde-inner">\n'
            '    <div class="garde-header">\n'
            f'      <div class="garde-logo-wrap"><img src="{LOGO_URI}" alt="ActuarIA"/></div>\n'
            '      <div><span class="badge-confidentiel">Confidentiel</span></div>\n'
            '    </div>\n'
            '    <div class="garde-hero">\n'
            '      <div class="garde-eyebrow">Rapport Actuariel · Sous-direction Santé</div>\n'
            '      <div class="garde-titre">Rapport<br><em>Santé</em></div>\n'
            '      <div class="garde-subtitle">Complémentaire Santé · Analyse Actuarielle Complète</div>\n'
            '    </div>\n'
            f'    <div class="garde-statut-badge" style="background:{col_rag}22;border:1.5px solid {col_rag};color:{col_rag};">'
            f'      {ic_rag} Statut {rag} — {avis}'
            '    </div>\n'
            '    <div class="garde-meta">\n'
            f'      <div class="garde-meta-item"><span class="garde-meta-label">Entité</span><span class="garde-meta-val">{entite}</span></div>\n'
            f'      <div class="garde-meta-item"><span class="garde-meta-label">Arrêté</span><span class="garde-meta-val">{date_arrete}</span></div>\n'
            f'      <div class="garde-meta-item"><span class="garde-meta-label">Assurés</span><span class="garde-meta-val">{m1.get("nb_assures", "—"):,}</span></div>\n'
            f'      <div class="garde-meta-item"><span class="garde-meta-label">Primes acquises</span><span class="garde-meta-val">{_f(m3.get("primes_acquises"))}</span></div>\n'
            f'      <div class="garde-meta-item"><span class="garde-meta-label">Édition</span><span class="garde-meta-val">{dt}</span></div>\n'
            '    </div>\n'
            '  </div>\n'
            '</div>\n\n'

            # ── SECTION 1 : SYNTHÈSE ──
            '<div class="section">\n'
            '  <div class="section-header">'
            '    <div class="section-num">1</div>'
            '    <div class="section-title">Synthèse exécutive</div>'
            '  </div>\n'
            '  <div class="kpi-grid">\n'
            f'    <div class="kpi-card"><div class="kpi-label">Best Estimate Santé</div><div class="kpi-val">{_f(m3.get("be_sante"))}</div><div class="kpi-sub">PSAP total — IFRS 17 §33</div></div>\n'
            f'    <div class="kpi-card"><div class="kpi-label">Risk Adjustment</div><div class="kpi-val">{_f(m3.get("risk_adjustment"))}</div><div class="kpi-sub">CoC 6% EIOPA — §B91</div></div>\n'
            f'    <div class="kpi-card"><div class="kpi-label">SCR Santé NSLT</div><div class="kpi-val">{_f(m3.get("scr_sante"))}</div><div class="kpi-sub">Art.148 RD 2015/35</div></div>\n'
            f'    <div class="kpi-card"><div class="kpi-label">Ratio SCR</div><div class="kpi-val" style="color:{_statut_col(rag)};">{_pct(m3.get("ratio_scr_pct"))}</div><div class="kpi-sub">Seuil ≥ 100% Art.129 S2</div></div>\n'
            f'    <div class="kpi-card"><div class="kpi-label">Loss Ratio observé</div><div class="kpi-val">{_pct(m3.get("loss_ratio", 0) * 100)}</div><div class="kpi-sub">FNMF 2023 : 65-75%</div></div>\n'
            f'    <div class="kpi-card"><div class="kpi-label">ANI 2013</div><div class="kpi-val">{"✅ Conforme" if m3.get("ani_conforme") else "⚠️ Non conforme"}</div><div class="kpi-sub">Art. L911-7 CSS</div></div>\n'
            '  </div>\n'
            '</div>\n\n'

            # ── SECTION 2 : SINISTRALITÉ PAR POSTE ──
            '<div class="section">\n'
            '  <div class="section-header">'
            '    <div class="section-num">2</div>'
            '    <div class="section-title">Sinistralité par poste de garantie</div>'
            '  </div>\n'
            '  <table class="data-table">\n'
            '    <thead><tr><th>Poste</th><th>Sinistre/an</th><th>Taux IBNR</th><th>IBNR</th><th>Source</th></tr></thead>\n'
            f'    <tbody>{postes_rows}</tbody>\n'
            '  </table>\n'
            '</div>\n\n'

            # ── SECTION 3 : PROVISIONNEMENT ──
            '<div class="section">\n'
            '  <div class="section-header">'
            '    <div class="section-num">3</div>'
            '    <div class="section-title">Provisionnement PSAP / IBNR</div>'
            '  </div>\n'
            '  <table class="data-table">\n'
            '    <thead><tr><th>Composante</th><th>Montant</th><th>Référence</th></tr></thead>\n'
            '    <tbody>\n'
            f'      <tr><td>PSAP dossiers connus</td><td>{_f(m3.get("psap_dossiers"))}</td><td>Dossiers ouverts</td></tr>\n'
            f'      <tr><td>PSAP IBNR (cadence santé)</td><td>{_f(m3.get("psap_ibnr"))}</td><td>5-25% selon poste — FNMF 2023</td></tr>\n'
            f'      <tr><td>PSAP Total</td><td>{_f(m3.get("psap_total"))}</td><td>Dossiers + IBNR</td></tr>\n'
            f'      <tr><td>PREC</td><td>{_f(m3.get("prec"))}</td><td>Risques en cours</td></tr>\n'
            f'      <tr><td>Provision totale</td><td>{_f(m3.get("provision_totale"))}</td><td>PSAP + PREC</td></tr>\n'
            '    </tbody>\n'
            '  </table>\n'
            '</div>\n\n'

            # ── SECTION 4 : BE + RA ──
            '<div class="section">\n'
            '  <div class="section-header">'
            '    <div class="section-num">4</div>'
            '    <div class="section-title">Best Estimate et Risk Adjustment — IFRS 17</div>'
            '  </div>\n'
            '  <table class="data-table">\n'
            '    <thead><tr><th>Composante</th><th>Montant</th><th>Méthode / Référence</th></tr></thead>\n'
            '    <tbody>\n'
            f'      <tr><td>Best Estimate Santé (BE)</td><td>{_f(m3.get("be_sante"))}</td><td>PSAP total — IFRS 17 §33</td></tr>\n'
            f'      <tr><td>Risk Adjustment (RA)</td><td>{_f(m3.get("risk_adjustment"))}</td><td>CoC 6% EIOPA — IFRS 17 §B91</td></tr>\n'
            f'      <tr><td>TP Santé (BE + RA)</td><td>{_f(m3.get("tp_sante"))}</td><td>Provisions techniques S2</td></tr>\n'
            f'      <tr><td>Ratio TP/BE</td><td>{m3.get("ratio_tp_be", 0):.3f}</td><td>Cible [1.00 — 1.50]</td></tr>\n'
            '    </tbody>\n'
            '  </table>\n'
            '</div>\n\n'

            # ── SECTION 5 : SCR / MCR ──
            '<div class="section">\n'
            '  <div class="section-header">'
            '    <div class="section-num">5</div>'
            '    <div class="section-title">SCR Santé NSLT et ratio de solvabilité</div>'
            '  </div>\n'
            '  <table class="data-table">\n'
            '    <thead><tr><th>Composante</th><th>Montant</th><th>Référence</th></tr></thead>\n'
            '    <tbody>\n'
            f'      <tr><td>SCR primes (σ=5%)</td><td>{_f(m3.get("scr_prem"))}</td><td>Art.148 RD 2015/35</td></tr>\n'
            f'      <tr><td>SCR réserves (σ=14%)</td><td>{_f(m3.get("scr_res"))}</td><td>Art.148 RD 2015/35</td></tr>\n'
            f'      <tr><td>SCR catastrophe santé</td><td>{_f(m3.get("scr_cat"))}</td><td>Art.159 RD 2015/35</td></tr>\n'
            f'      <tr><td>SCR Santé Total</td><td>{_f(m3.get("scr_sante"))}</td><td>Formule standard EIOPA</td></tr>\n'
            f'      <tr><td>MCR Santé</td><td>{_f(m3.get("mcr_sante"))}</td><td>Art.252 RD 2015/35</td></tr>\n'
            f'      <tr><td>Fonds propres éligibles</td><td>{_f(m3.get("fonds_propres"))}</td><td>Bilan S2</td></tr>\n'
            f'      <tr><td>Ratio SCR</td><td style="color:{_statut_col(rag)};font-weight:700;">{_pct(m3.get("ratio_scr_pct"))} {"✅" if m3.get("ratio_scr_pct", 0) >= SEUIL_SCR else "❌"}</td><td>Seuil ≥ 100% — Art.129 S2</td></tr>\n'
            f'      <tr><td>Ratio MCR</td><td>{_pct(m3.get("ratio_mcr_pct"))}</td><td>Art.129 S2</td></tr>\n'
            '    </tbody>\n'
            '  </table>\n'
            '</div>\n\n'

            # ── SECTION 6 : HYPOTHÈSES ──
            '<div class="section">\n'
            '  <div class="section-header">'
            '    <div class="section-num">6</div>'
            '    <div class="section-title">Validation des hypothèses actuarielles</div>'
            '  </div>\n'
            + hyp_html +
            '</div>\n\n'

            # ── SECTION 7 : NARRATION ──
            '<div class="section">\n'
            '  <div class="section-header">'
            '    <div class="section-num">7</div>'
            '    <div class="section-title">Commentaire actuariel</div>'
            '  </div>\n'
            '  <div class="narration-block">\n'
            + narration_html +
            '  </div>\n'
            + source_badge +
            '\n</div>\n\n'

            # ── AVIS ACTUARIEL ──
            '<div class="section">\n'
            f'  <div class="avis-block {avis_cls}">\n'
            f'    <div class="avis-titre">{ic_rag} Avis actuariel : {avis}</div>\n'
            f'    <div style="font-size:9pt;">{m4.get("avis_detail", "")}</div>\n'
            '  </div>\n'
            '</div>\n\n'

            # ── PIED DE PAGE ──
            '<div class="pied-de-page">\n'
            f'  <img src="{LOGO_URI}" alt="ActuarIA"/>\n'
            '  <div class="pied-meta">\n'
            f'    {entite} · Rapport Santé · Arrêté {date_arrete} · Édition {dt}<br>\n'
            '    <span style="font-size:6.5pt;color:#8A9BB0;">Rapport établi conformément à l\'Art. 77 et 148 de la Directive Solvabilité II, '
            'à l\'IFRS 17 §B91 et au Guide Institut des Actuaires 2023</span><br>\n'
            '    <span class="confidentiel-footer">CONFIDENTIEL — USAGE STRICTEMENT ACTUARIEL</span>\n'
            '  </div>\n'
            '</div>\n\n'

            '</div>\n</body>\n</html>'
        )
        logger.info(f"HTML Santé : {len(html):,} chars — narration={source_narration}")
        return html

    except Exception as e:
        logging.getLogger("actuaria.sp.rapport_sante").error(f"export_html : {e}", exc_info=True)
        return f"<html><body><h1>Erreur : {e}</h1></body></html>"


# ══════════════════════════════════════════════════════════════════════════════
# EXPORT PDF
# ══════════════════════════════════════════════════════════════════════════════

def _export_pdf_sante(html: str) -> bytes:
    try:
        from weasyprint import HTML as WH
        pdf = WH(string=html).write_pdf()
        logging.getLogger("actuaria.sp.rapport_sante").info(f"PDF : {len(pdf):,} bytes")
        return pdf
    except ImportError:
        logging.getLogger("actuaria.sp.rapport_sante").warning("weasyprint non installé")
        return b""
    except Exception as e:
        logging.getLogger("actuaria.sp.rapport_sante").error(f"PDF : {e}")
        return b""


# ══════════════════════════════════════════════════════════════════════════════
# EXPORT WORD
# ══════════════════════════════════════════════════════════════════════════════

def _export_word_sante(m1: Dict, m3: Dict, m4: Dict,
                       narration: str, source_narration: str, entite: str, date_arrete: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        logging.getLogger("actuaria.sp.rapport_sante").error(f"python-docx absent : {e}")
        return b""

    try:
        def rgb(h):
            h = h.lstrip("#")
            return RGBColor(int(h[:2], 16), int(h[2:4], 16), int(h[4:], 16))

        NR = rgb(NAVY); GR = rgb(GOLD); WR = rgb("#FFFFFF")
        GrR = rgb(SLATE); RgR = rgb(ROUGE); VR = rgb(VERT); AR = rgb(ORANGE)

        dt     = datetime.now().strftime("%d/%m/%Y")
        rag    = m4.get("statut_rag", "AMBRE")
        avis   = m4.get("avis_actuariel", "AVEC RÉSERVES")
        s_col  = VR if rag == "VERT" else (RgR if rag == "ROUGE" else AR)

        doc = Document()
        for s in doc.sections:
            s.top_margin = Cm(2); s.bottom_margin = Cm(2)
            s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

        def _bg(cell, hex6):
            tc = cell._tc; tcp = tc.get_or_add_tcPr()
            sd = OxmlElement("w:shd")
            sd.set(qn("w:fill"), hex6.lstrip("#"))
            sd.set(qn("w:color"), "auto"); sd.set(qn("w:val"), "clear")
            tcp.append(sd)

        def _run(p, txt, bold=False, italic=False, sz=10, col=None):
            r = p.add_run(str(txt)); r.bold = bold; r.italic = italic
            r.font.size = Pt(sz)
            if col: r.font.color.rgb = col
            return r

        def _h(txt, lv=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            sz = {1: 16, 2: 12, 3: 10}.get(lv, 10)
            c = NR if lv == 1 else GR
            _run(p, txt, bold=True, sz=sz, col=c)

        def _sep():
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr")
            bo = OxmlElement("w:bottom")
            bo.set(qn("w:val"), "single"); bo.set(qn("w:sz"), "6")
            bo.set(qn("w:space"), "1"); bo.set(qn("w:color"), "C9A84C")
            pBdr.append(bo); pPr.append(pBdr)

        def _tbl(heads, rows, ws=None):
            t = doc.add_table(rows=1 + len(rows), cols=len(heads))
            t.style = "Table Grid"
            for i, hd in enumerate(heads):
                c = t.rows[0].cells[i]; _bg(c, "0B1E3D")
                pp = c.paragraphs[0]
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = pp.add_run(str(hd))
                r.bold = True; r.font.size = Pt(9); r.font.color.rgb = WR
            for ri, row in enumerate(rows):
                for ci, v in enumerate(row):
                    c = t.rows[ri + 1].cells[ci]
                    if ri % 2 == 1: _bg(c, "EEF2F7")
                    pp = c.paragraphs[0]
                    r = pp.add_run(str(v) if v is not None else "—")
                    r.font.size = Pt(9)
            if ws:
                for i, w in enumerate(ws):
                    for row in t.rows:
                        row.cells[i].width = Cm(w)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # Page de garde
        p = doc.add_paragraph()
        _run(p, "Actuar", bold=True, sz=28, col=NR)
        _run(p, "IA", bold=True, sz=28, col=GR)
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, "RAPPORT ACTUARIEL SANTÉ\n", bold=True, sz=18, col=NR)
        _run(p, "Complémentaire Santé — Analyse Actuarielle Complète", sz=12, col=GR)
        doc.add_paragraph()
        p = doc.add_paragraph()
        _run(p, "Statut : ", sz=11, col=NR)
        _run(p, f"{rag} — {avis}", bold=True, sz=11, col=s_col)
        doc.add_paragraph()
        _tbl(
            ["Entité", "Arrêté", "Assurés", "Primes acquises"],
            [[entite, date_arrete, f"{m1.get('nb_assures', '—'):,}", _f(m3.get("primes_acquises"))]],
            ws=[4.0, 3.0, 3.0, 4.0],
        )
        doc.add_page_break()

        # 1. Synthèse
        _h("1. Synthèse exécutive"); _sep()
        _tbl(
            ["Indicateur", "Valeur", "Référence"],
            [
                ["Best Estimate Santé", _f(m3.get("be_sante")), "IFRS 17 §33"],
                ["Risk Adjustment", _f(m3.get("risk_adjustment")), "CoC 6% — §B91"],
                ["TP Santé (BE+RA)", _f(m3.get("tp_sante")), "Provisions techniques S2"],
                ["SCR Santé NSLT", _f(m3.get("scr_sante")), "Art.148 RD 2015/35"],
                ["MCR Santé", _f(m3.get("mcr_sante")), "Art.252 RD 2015/35"],
                ["Ratio SCR", _pct(m3.get("ratio_scr_pct")), "≥ 100% Art.129 S2"],
                ["Loss Ratio observé", _pct(m3.get("loss_ratio", 0) * 100), "FNMF 2023 : 65-75%"],
                ["ANI 2013", "Conforme" if m3.get("ani_conforme") else "Non conforme", "Art. L911-7 CSS"],
            ],
            ws=[5.5, 4.0, 6.5],
        )
        doc.add_page_break()

        # 2. Sinistralité
        _h("2. Sinistralité par poste de garantie"); _sep()
        rows_p = []
        for p_name, v in m3.get("postes", {}).items():
            rows_p.append([
                p_name.replace("_", " ").title(),
                _f(v.get("sinistre_annuel")),
                _pct(v.get("taux_ibnr", 0) * 100),
                _f(v.get("ibnr", 0)),
                v.get("source", "—"),
            ])
        if rows_p:
            _tbl(["Poste", "Sinistre/an", "Taux IBNR", "IBNR", "Source"], rows_p,
                 ws=[3.5, 3.0, 2.5, 3.0, 4.0])
        doc.add_page_break()

        # 3. Provisionnement
        _h("3. Provisionnement PSAP / IBNR"); _sep()
        _tbl(
            ["Composante", "Montant", "Référence"],
            [
                ["PSAP dossiers connus", _f(m3.get("psap_dossiers")), "Dossiers ouverts"],
                ["PSAP IBNR", _f(m3.get("psap_ibnr")), "5-25% selon poste — FNMF 2023"],
                ["PSAP Total", _f(m3.get("psap_total")), "Dossiers + IBNR"],
                ["PREC", _f(m3.get("prec")), "Risques en cours"],
                ["Provision totale", _f(m3.get("provision_totale")), "PSAP + PREC"],
            ],
            ws=[5.0, 4.0, 7.0],
        )
        doc.add_page_break()

        # 4. BE + RA
        _h("4. Best Estimate et Risk Adjustment — IFRS 17"); _sep()
        _tbl(
            ["Composante", "Montant", "Méthode"],
            [
                ["Best Estimate (BE)", _f(m3.get("be_sante")), "PSAP total — §33 IFRS 17"],
                ["Risk Adjustment (RA)", _f(m3.get("risk_adjustment")), "CoC 6% EIOPA — §B91"],
                ["TP Santé (BE + RA)", _f(m3.get("tp_sante")), "Provisions techniques S2"],
                ["Ratio TP/BE", f"{m3.get('ratio_tp_be', 0):.3f}", "Cible [1.00 — 1.50]"],
            ],
            ws=[5.0, 4.0, 7.0],
        )
        doc.add_page_break()

        # 5. SCR/MCR
        _h("5. SCR Santé NSLT et ratio de solvabilité"); _sep()
        _tbl(
            ["Composante", "Montant", "Référence"],
            [
                ["SCR primes (σ=5%)", _f(m3.get("scr_prem")), "Art.148 RD 2015/35"],
                ["SCR réserves (σ=14%)", _f(m3.get("scr_res")), "Art.148 RD 2015/35"],
                ["SCR catastrophe santé", _f(m3.get("scr_cat")), "Art.159 RD 2015/35"],
                ["SCR Santé Total", _f(m3.get("scr_sante")), "Formule standard EIOPA"],
                ["MCR Santé", _f(m3.get("mcr_sante")), "Art.252 RD 2015/35"],
                ["Fonds propres éligibles", _f(m3.get("fonds_propres")), "Bilan S2"],
                ["Ratio SCR", _pct(m3.get("ratio_scr_pct")), "≥ 100% — Art.129 S2"],
            ],
            ws=[5.0, 4.0, 7.0],
        )
        doc.add_page_break()

        # 6. Hypothèses
        _h("6. Validation des hypothèses actuarielles"); _sep()
        rows_h = []
        for h in m4.get("hypotheses", []):
            st = h.get("statut", "")
            rows_h.append([h["id"], h["hypothese"][:60], st, str(h.get("valeur", ""))[:80]])
        if rows_h:
            _tbl(["ID", "Hypothèse", "Statut", "Valeur"], rows_h, ws=[1.5, 5.5, 2.5, 6.5])
        doc.add_page_break()

        # 7. Narration
        _h("7. Commentaire actuariel"); _sep()
        if narration:
            for sec in re.split(r"(?=§\d+\s*[—\-–])", _clean(narration)):
                sec = sec.strip()
                if not sec: continue
                ls = sec.split("\n", 1)
                if ls[0]: _h(ls[0], lv=2)
                if len(ls) > 1:
                    for ln in ls[1].split("\n"):
                        ln = ln.strip()
                        if ln:
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(3)
                            p.paragraph_format.left_indent = Cm(0.3)
                            _run(p, ln, sz=9, col=NR)
        else:
            p = doc.add_paragraph()
            _run(p, "Narration non disponible.", sz=9, italic=True)
        # /!\ LA MENTION D ORIGINE, DANS LE WORD AUSSI. Elle en etait
        # totalement absente : ce format ne recevait meme pas la source.
        _badge = _badge_narration_sante(source_narration)
        if _badge:
            p = doc.add_paragraph()
            _run(p, _badge, sz=7, italic=True, col=GrR)
        doc.add_paragraph()

        # Pied
        _sep()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"ActuarIA · {entite} · Rapport Santé · Arrêté {date_arrete} · {dt} · CONFIDENTIEL",
             sz=7, italic=True, col=GrR)

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        wb = buf.read()
        logging.getLogger("actuaria.sp.rapport_sante").info(f"Word : {len(wb):,} bytes")
        return wb

    except Exception as e:
        logging.getLogger("actuaria.sp.rapport_sante").error(f"Word : {e}", exc_info=True)
        return b""


# ══════════════════════════════════════════════════════════════════════════════
# EXPORT EXCEL
# ══════════════════════════════════════════════════════════════════════════════

def _export_excel_sante(m1: Dict, m3: Dict, m4: Dict,
                        entite: str, date_arrete: str) -> bytes:
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return b""

    try:
        wb = openpyxl.Workbook()
        NAVY_XL  = "0B1E3D"
        GOLD_XL  = "C9A84C"

        def _style_header(ws, row, cols):
            for col in range(1, cols + 1):
                c = ws.cell(row=row, column=col)
                c.fill = PatternFill("solid", fgColor=NAVY_XL)
                c.font = Font(color="FFFFFF", bold=True, size=10)
                c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        def _style_row(ws, row, cols, even=False):
            for col in range(1, cols + 1):
                c = ws.cell(row=row, column=col)
                if even:
                    c.fill = PatternFill("solid", fgColor="EEF2F7")
                c.font = Font(size=9)
                c.alignment = Alignment(vertical="center")

        def _title(ws, txt):
            ws.merge_cells("A1:D1")
            c = ws["A1"]
            c.value = txt
            c.font = Font(color=GOLD_XL, bold=True, size=13)
            c.fill = PatternFill("solid", fgColor=NAVY_XL)
            c.alignment = Alignment(horizontal="center", vertical="center")
            ws.row_dimensions[1].height = 30

        # ── Onglet 1 : Synthèse ──
        ws1 = wb.active
        ws1.title = "Synthèse"
        _title(ws1, f"RAPPORT SANTÉ — {entite} — Arrêté {date_arrete}")
        ws1.append(["Indicateur", "Valeur", "Référence", ""])
        _style_header(ws1, 2, 4)
        rows = [
            ["Best Estimate Santé", _f(m3.get("be_sante")), "IFRS 17 §33"],
            ["Risk Adjustment (CoC 6%)", _f(m3.get("risk_adjustment")), "IFRS 17 §B91"],
            ["TP Santé (BE + RA)", _f(m3.get("tp_sante")), "Provisions techniques S2"],
            ["SCR Santé NSLT", _f(m3.get("scr_sante")), "Art.148 RD 2015/35"],
            ["MCR Santé", _f(m3.get("mcr_sante")), "Art.252 RD 2015/35"],
            ["Fonds propres éligibles", _f(m3.get("fonds_propres")), "Bilan S2"],
            ["Ratio SCR", _pct(m3.get("ratio_scr_pct")), "≥ 100% Art.129 S2"],
            ["Loss Ratio observé", _pct(m3.get("loss_ratio", 0) * 100), "FNMF 2023 : 65-75%"],
            ["ANI 2013", "Conforme" if m3.get("ani_conforme") else "Non conforme", "Art. L911-7 CSS"],
            ["Statut RAG", m4.get("statut_rag", "—"), ""],
            ["Avis actuariel", m4.get("avis_actuariel", "—"), ""],
        ]
        for i, row in enumerate(rows):
            ws1.append(row + [""])
            _style_row(ws1, i + 3, 4, even=(i % 2 == 1))
        ws1.column_dimensions["A"].width = 32
        ws1.column_dimensions["B"].width = 20
        ws1.column_dimensions["C"].width = 28

        # ── Onglet 2 : Sinistralité par poste ──
        ws2 = wb.create_sheet("Sinistralité")
        _title(ws2, "Sinistralité par poste de garantie — DREES 2023")
        ws2.append(["Poste", "Sinistre/an (€)", "Taux IBNR", "IBNR (€)", "Source"])
        _style_header(ws2, 2, 5)
        for i, (p_name, v) in enumerate(m3.get("postes", {}).items()):
            ws2.append([
                p_name.replace("_", " ").title(),
                _f(v.get("sinistre_annuel")),
                _pct(v.get("taux_ibnr", 0) * 100),
                _f(v.get("ibnr", 0)),
                v.get("source", "—"),
            ])
            _style_row(ws2, i + 3, 5, even=(i % 2 == 1))
        for col in ["A", "B", "C", "D", "E"]:
            ws2.column_dimensions[col].width = 22

        # ── Onglet 3 : Provisionnement ──
        ws3 = wb.create_sheet("Provisionnement")
        _title(ws3, "Provisionnement PSAP / IBNR / PREC")
        ws3.append(["Composante", "Montant (€)", "Référence"])
        _style_header(ws3, 2, 3)
        prov_rows = [
            ["PSAP dossiers connus", _f(m3.get("psap_dossiers")), "Dossiers ouverts"],
            ["PSAP IBNR", _f(m3.get("psap_ibnr")), "5-25% selon poste — FNMF 2023"],
            ["PSAP Total", _f(m3.get("psap_total")), "Dossiers + IBNR"],
            ["PREC", _f(m3.get("prec")), "Risques en cours"],
            ["Provision totale", _f(m3.get("provision_totale")), "PSAP + PREC"],
        ]
        for i, row in enumerate(prov_rows):
            ws3.append(row)
            _style_row(ws3, i + 3, 3, even=(i % 2 == 1))
        ws3.column_dimensions["A"].width = 30
        ws3.column_dimensions["B"].width = 20
        ws3.column_dimensions["C"].width = 35

        # ── Onglet 4 : SCR / MCR ──
        ws4 = wb.create_sheet("SCR-MCR")
        _title(ws4, "SCR Santé NSLT et Solvabilité II")
        ws4.append(["Composante", "Montant (€)", "Référence"])
        _style_header(ws4, 2, 3)
        scr_rows = [
            ["SCR primes (σ=5%)", _f(m3.get("scr_prem")), "Art.148 RD 2015/35"],
            ["SCR réserves (σ=14%)", _f(m3.get("scr_res")), "Art.148 RD 2015/35"],
            ["SCR catastrophe santé", _f(m3.get("scr_cat")), "Art.159 RD 2015/35"],
            ["SCR Santé Total", _f(m3.get("scr_sante")), "Formule standard EIOPA"],
            ["MCR Santé", _f(m3.get("mcr_sante")), "Art.252 RD 2015/35"],
            ["Fonds propres éligibles", _f(m3.get("fonds_propres")), "Bilan S2"],
            ["Ratio SCR", _pct(m3.get("ratio_scr_pct")), "≥ 100% — Art.129 S2"],
            ["Ratio MCR", _pct(m3.get("ratio_mcr_pct")), "Art.129 S2"],
        ]
        for i, row in enumerate(scr_rows):
            ws4.append(row)
            _style_row(ws4, i + 3, 3, even=(i % 2 == 1))
        ws4.column_dimensions["A"].width = 30
        ws4.column_dimensions["B"].width = 20
        ws4.column_dimensions["C"].width = 30

        buf = io.BytesIO()
        wb.save(buf); buf.seek(0)
        xb = buf.read()
        logging.getLogger("actuaria.sp.rapport_sante").info(f"Excel : {len(xb):,} bytes")
        return xb

    except Exception as e:
        logging.getLogger("actuaria.sp.rapport_sante").error(f"Excel : {e}", exc_info=True)
        return b""


# ══════════════════════════════════════════════════════════════════════════════
# AGENT PRINCIPAL
# ══════════════════════════════════════════════════════════════════════════════

class AgentRapportSante:
    """
    Agent Rapport Santé — Rapport actuariel consolidé Sous-direction Santé.
    Pipeline M1→M5 : ingestion S1+S2+S3+REG → calculs → Claude API → 4 formats.
    """

    NOM     = "RapportSante"
    CODE    = "SP-RAPPORT-SANTE"
    VERSION = "1.0"
    MANAGER = "Amira (Directrice SP)"

    def __init__(self, audit_path: str = "audit", verbose: bool = True):
        self.audit_path = Path(audit_path)
        self.audit_path.mkdir(parents=True, exist_ok=True)
        self.logger = logging.getLogger("actuaria.sp.rapport_sante")
        self.verbose = verbose

    # ── run ───────────────────────────────────────────────────────────────────
    def run(self,
            result_s1=None,
            result_s2=None,
            result_s3=None,
            result_reg2=None,
            result_reg3=None,
            entite: str = "Mutuelle Santé",
            date_arrete: str = "",
            fonds_propres: float = 0.0,
            generer_graphiques: bool = True) -> Dict:

        t0  = datetime.now()
        aid = f"RS_{t0.strftime('%Y%m%d_%H%M%S')}"
        date_arrete = date_arrete or t0.strftime("%d/%m/%Y")

        try:
            # ── M1 : Ingestion ────────────────────────────────────────────────
            m1 = self._m1_ingestion(result_s1, result_s2, result_s3, fonds_propres)
            self.logger.info(
                f"[{aid}] M1 | assurés={m1.get('nb_assures', '—')} | "
                f"modules={m1.get('modules_disponibles', [])}"
            )

            # ── M2 : Hypothèses ───────────────────────────────────────────────
            m2 = self._m2_hypotheses(result_s1, result_s2, result_s3)

            # ── M3 : Calculs ──────────────────────────────────────────────────
            m3 = self._m3_calculs(result_s1, result_s2, result_s3, m1, fonds_propres)

            # ── M4 : Synthèse & avis ─────────────────────────────────────────
            hyp = self._hypotheses(m3, result_s1, result_s2)
            rag = self._rag(hyp, m3)
            avis, avis_detail = self._avis(rag, m3)
            m4 = {
                "statut_rag":     rag,
                "avis_actuariel": avis,
                "avis_detail":    avis_detail,
                "hypotheses":     hyp,
            }

            # Hash session
            session_hash = hashlib.sha256(
                f"{aid}{m3.get('be_sante', 0)}{m3.get('scr_sante', 0)}{rag}".encode()
            ).hexdigest()[:8].upper()

            # ── M5 : Livrables ───────────────────────────────────────────────
            commentaire = self._commentaire(rag, m3, m4, entite, date_arrete)
            narration, src_narration = _generer_narration(
                m1, m3, date_arrete, commentaire
            )
            html_bytes = _export_html_sante(m1, m3, m4, narration, src_narration, entite, date_arrete)
            pdf_bytes  = _export_pdf_sante(html_bytes)
            word_bytes = _export_word_sante(m1, m3, m4, narration, src_narration, entite, date_arrete)
            excel_bytes = _export_excel_sante(m1, m3, m4, entite, date_arrete)

            # Graphiques
            gph = {}
            if generer_graphiques and PLOTLY_OK:
                gph = self._graphiques(m3, m4)

            # Audit
            self._audit(aid, m3, rag, session_hash)
            if self.verbose:
                self._console(aid, rag, avis, m3, session_hash)

            duree = (datetime.now() - t0).total_seconds()

            return {
                "success":      True,
                "agent":        self.NOM,
                "version":      self.VERSION,
                "audit_id":     aid,
                "statut_rag":   rag,
                "session_hash": session_hash,
                # Données clés
                "be_sante":         round(m3.get("be_sante", 0), 2),
                "risk_adjustment":  round(m3.get("risk_adjustment", 0), 2),
                "tp_sante":         round(m3.get("tp_sante", 0), 2),
                "scr_sante":        round(m3.get("scr_sante", 0), 2),
                "mcr_sante":        round(m3.get("mcr_sante", 0), 2),
                "ratio_scr_pct":    round(m3.get("ratio_scr_pct", 0), 1),
                "loss_ratio":       round(m3.get("loss_ratio", 0), 4),
                "ani_conforme":     m3.get("ani_conforme", False),
                "avis_actuariel":   avis,
                "modules_disponibles": m1.get("modules_disponibles", []),
                # M1-M4
                "m1": m1, "m2": m2, "m3": m3, "m4": m4,
                # Livrables M5
                "html_bytes":   html_bytes.encode() if isinstance(html_bytes, str) else html_bytes,
                "pdf_bytes":    pdf_bytes,
                "word_bytes":   word_bytes,
                "excel_bytes":  excel_bytes,
                # Standard ActuarIA
                "hypotheses":   hyp,
                "commentaire":  commentaire,
                "graphiques":   gph,
                "duree_sec":    round(duree, 2),
                "erreur":       None,
            }

        except Exception as e:
            self.logger.error(f"[{aid}] ERREUR : {e}", exc_info=True)
            return self._erreur(str(e), aid)

    # ── M1 : Ingestion ────────────────────────────────────────────────────────
    def _m1_ingestion(self, s1, s2, s3, fonds_propres: float) -> Dict:
        modules = []
        if s1 and s1.get("success"): modules.append("S1")
        if s2 and s2.get("success"): modules.append("S2")
        if s3 and s3.get("success"): modules.append("S3")

        if not modules:
            raise ValueError("Au moins S1 ou S2 requis pour le rapport Santé")

        # Portefeuille depuis S1
        nb_assures     = int((s1 or {}).get("nb_assures", 1000))
        age_moyen      = float((s1 or {}).get("age_moyen", 42.0))
        contrat        = (s1 or {}).get("contrat", "collectif")
        garantie       = (s1 or {}).get("garantie_niveau", "confort")
        source_donnees = (s1 or {}).get("source_donnees", "—")
        ani_conforme   = bool((s1 or {}).get("ani_conforme", True))

        # Fonds propres
        fp = fonds_propres
        if fp <= 0 and s3 and s3.get("success"):
            fp = float(s3.get("fonds_propres", 0))

        return {
            "modules_disponibles": modules,
            "nb_assures":     nb_assures,
            "age_moyen":      age_moyen,
            "contrat":        contrat,
            "garantie_niveau": garantie,
            "source_donnees": source_donnees,
            "ani_conforme":   ani_conforme,
            "fonds_propres":  fp,
            "lr_fnmf_ref":    "65%-75% garantie confort — FNMF 2023",
        }

    # ── M2 : Hypothèses ───────────────────────────────────────────────────────
    def _m2_hypotheses(self, s1, s2, s3) -> Dict:
        """Collecte les hypothèses des agents sources."""
        hyp_s1 = (s1 or {}).get("hypotheses", []) if s1 else []
        hyp_s2 = (s2 or {}).get("hypotheses", []) if s2 else []
        hyp_s3 = (s3 or {}).get("hypotheses", []) if s3 else []
        return {
            "hypotheses_s1": hyp_s1,
            "hypotheses_s2": hyp_s2,
            "hypotheses_s3": hyp_s3,
            "referentiels": [
                "DREES 2023 — Comptes de la Santé",
                "FNMF 2023 — Sinistralité mutuelles",
                "ANI 2013 — Art. L911-7 CSS",
                "EIOPA CoC 6% — IFRS 17 §B91",
                "RD 2015/35 — Art.148 (SCR NSLT), Art.252 (MCR)",
            ],
        }

    # ── M3 : Calculs ─────────────────────────────────────────────────────────
    def _m3_calculs(self, s1, s2, s3, m1: Dict, fonds_propres: float) -> Dict:
        # ── Tarification (S1) ──
        prime_pure    = float((s1 or {}).get("prime_pure", 0))
        prime_comm    = float((s1 or {}).get("prime_commerciale", 0))
        primes_acq    = float((s1 or {}).get("primes_acquises", 0))
        lr_attendu    = float((s1 or {}).get("ratio_sp_attendu", 0))
        ani_conforme  = bool((s1 or {}).get("ani_conforme", True))
        ani_detail    = (s1 or {}).get("ani_detail", {})
        postes_s1     = (s1 or {}).get("postes", {})

        # ANI détail string
        ani_detail_str = " | ".join(
            f"{p}: {d.get('note', '—')}"
            for p, d in ani_detail.get("detail", {}).items()
            if d.get("seuil", 0) > 0
        ) or "N/A"

        # ── Provisionnement (S2) ──
        psap_dossiers   = float((s2 or {}).get("psap_dossiers", 0))
        psap_ibnr       = float((s2 or {}).get("psap_ibnr", 0))
        psap_total      = float((s2 or {}).get("psap_total", 0))
        prec            = float((s2 or {}).get("prec", 0))
        prov_totale     = float((s2 or {}).get("provision_totale", 0))
        loss_ratio      = float((s2 or {}).get("loss_ratio", 0))
        psap_par_poste  = (s2 or {}).get("psap_par_poste", {})
        ibnr_par_poste  = (s2 or {}).get("ibnr_par_poste", {})

        # Sinistres attendus (S2 via S1)
        sin_att = float((s2 or {}).get("sinistres_attendus", primes_acq * lr_attendu))
        sin_pay = float((s2 or {}).get("sinistres_payes", 0))
        ae_ratio = round(sin_pay / max(sin_att, 1), 3) if sin_att > 0 else "—"

        # ── Consolidation postes ──
        postes_consolides = {}
        for p_name, v in postes_s1.items():
            ibnr_v = ibnr_par_poste.get(p_name, {})
            postes_consolides[p_name] = {
                "sinistre_annuel": v.get("sinistre_annuel", 0),
                "taux_ibnr":       ibnr_v.get("taux_ibnr", 0),
                "ibnr":            ibnr_v.get("ibnr", 0),
                "source":          v.get("source", "DREES_2023"),
            }

        # ── Best Estimate & RA (S2 ou recalculer depuis S3) ──
        be_sante = float((s2 or {}).get("psap_total", psap_total))
        if s3 and s3.get("success"):
            be_sante     = float(s3.get("be_sante", be_sante))
            risk_adj     = float(s3.get("risk_adjustment", 0))
            tp_sante     = float(s3.get("tp_sante", be_sante + risk_adj))
            ratio_tp_be  = tp_sante / max(be_sante, 1)
            scr_sante    = float(s3.get("scr_sante", 0))
            scr_prem     = float(s3.get("scr_prem", 0))
            scr_res      = float(s3.get("scr_res", 0))
            scr_cat      = float(s3.get("scr_cat", 0))
            mcr_sante    = float(s3.get("mcr_sante", 0))
            fp           = float(s3.get("fonds_propres", fonds_propres))
            ratio_scr    = float(s3.get("ratio_scr_pct", 0))
            ratio_mcr    = float(s3.get("ratio_mcr_pct", 0))
        else:
            # Recalcul depuis formule standard EIOPA si S3 absent
            scr_prem  = 3.0 * SCR_SIGMA_PREM * primes_acq
            scr_res   = 3.0 * SCR_SIGMA_RES  * be_sante
            scr_sous  = float(np.sqrt(scr_prem**2 + 2 * 0.5 * scr_prem * scr_res + scr_res**2))
            scr_cat   = primes_acq * 0.01
            scr_sante = float(np.sqrt(scr_sous**2 + scr_cat**2))
            mcr_lin   = MCR_COEFF_PREM * primes_acq + MCR_COEFF_RES * be_sante
            plancher  = max(0.25 * scr_sante, MCR_PLANCHER)
            plafond   = 0.45 * scr_sante
            mcr_sante = max(min(mcr_lin, plafond), plancher)
            _scr_p    = SCR_SIGMA_PREM * primes_acq * 3
            risk_adj  = max(_scr_p * COC_RATE * 0.5, be_sante * 0.01)
            tp_sante  = be_sante + risk_adj
            ratio_tp_be = tp_sante / max(be_sante, 1)
            fp        = m1.get("fonds_propres", fonds_propres)
            ratio_scr = fp / max(scr_sante, 1) * 100
            ratio_mcr = fp / max(mcr_sante, 1) * 100

        if primes_acq <= 0 and s3:
            primes_acq = float(s3.get("primes_acquises", primes_acq))

        return {
            # Tarification
            "prime_pure":         prime_pure,
            "prime_commerciale":  prime_comm,
            "primes_acquises":    primes_acq,
            "lr_attendu":         lr_attendu,
            "loss_ratio":         loss_ratio,
            "ae_ratio":           ae_ratio,
            "ani_conforme":       ani_conforme,
            "ani_detail_str":     ani_detail_str,
            # Postes
            "postes":             postes_consolides,
            # Provisionnement
            "psap_dossiers":      psap_dossiers,
            "psap_ibnr":          psap_ibnr,
            "psap_total":         psap_total,
            "prec":               prec,
            "provision_totale":   prov_totale,
            # BE + RA
            "be_sante":           be_sante,
            "risk_adjustment":    risk_adj,
            "tp_sante":           tp_sante,
            "ratio_tp_be":        ratio_tp_be,
            # SCR / MCR
            "scr_sante":          scr_sante,
            "scr_prem":           scr_prem,
            "scr_res":            scr_res,
            "scr_cat":            scr_cat,
            "mcr_sante":          mcr_sante,
            "fonds_propres":      fp,
            "ratio_scr_pct":      ratio_scr,
            "ratio_mcr_pct":      ratio_mcr,
            # Hypothèses consolidées (pour narration)
            "hypotheses_consolidees": [],
        }

    # ── Hypothèses ───────────────────────────────────────────────────────────
    def _hypotheses(self, m3: Dict, s1, s2) -> list:
        hyp = []

        # H1 — Ratio SCR ≥ 100%
        ok1 = m3.get("ratio_scr_pct", 0) >= SEUIL_SCR
        hyp.append({
            "id": "H1",
            "hypothese": f"Ratio SCR ≥ {SEUIL_SCR:.0f}% — Art.129 Directive S2",
            "valeur": f"Ratio SCR = {m3.get('ratio_scr_pct', 0):.1f}% {'≥' if ok1 else '<'} {SEUIL_SCR:.0f}%",
            "statut": "VALIDÉE" if ok1 else "NON VALIDÉE",
            "critique": True,
        })

        # H2 — Loss Ratio ≤ 85%
        lr = m3.get("loss_ratio", 0)
        if lr <= 0.85:
            h2_s, h2_m = "VALIDÉE", f"LR = {lr*100:.1f}% ≤ 85% ✅"
        elif lr <= 0.95:
            h2_s, h2_m = "À JUSTIFIER", f"LR = {lr*100:.1f}% ∈ [85%,95%] — surveiller"
        else:
            h2_s, h2_m = "NON VALIDÉE", f"LR = {lr*100:.1f}% > 95% — contrat déficitaire"
        hyp.append({"id": "H2", "hypothese": "Loss Ratio ≤ 85% — sinistralité maîtrisée",
                    "valeur": h2_m, "statut": h2_s, "critique": True})

        # H3 — Ratio TP/BE ∈ [1.0, 1.5]
        rtp = m3.get("ratio_tp_be", 0)
        ok3 = 1.0 <= rtp <= 1.5
        hyp.append({
            "id": "H3",
            "hypothese": "Ratio TP/BE ∈ [1.0, 1.5] — CoC 6% EIOPA §B91",
            "valeur": f"TP/BE = {rtp:.3f} {'∈ [1.0,1.5] ✅' if ok3 else '— hors plage à justifier'}",
            "statut": "VALIDÉE" if ok3 else "À JUSTIFIER",
            "critique": False,
        })

        # H4 — ANI 2013
        ok4 = bool(m3.get("ani_conforme", True))
        hyp.append({
            "id": "H4",
            "hypothese": "Conformité ANI 2013 — Art. L911-7 CSS",
            "valeur": "Panier ANI conforme ✅" if ok4 else f"Non conforme : {m3.get('ani_detail_str', '—')}",
            "statut": "VALIDÉE" if ok4 else "NON VALIDÉE",
            "critique": True,
        })

        # H5 — Ratio SCR cible ≥ 130%
        ok5 = m3.get("ratio_scr_pct", 0) >= SEUIL_SCR_CIBLE
        hyp.append({
            "id": "H5",
            "hypothese": f"Ratio SCR ≥ {SEUIL_SCR_CIBLE:.0f}% (cible interne)",
            "valeur": f"Ratio SCR = {m3.get('ratio_scr_pct', 0):.1f}% {'≥' if ok5 else '<'} {SEUIL_SCR_CIBLE:.0f}%",
            "statut": "VALIDÉE" if ok5 else "À JUSTIFIER",
            "critique": False,
        })

        # Injecter dans m3 pour la narration
        m3["hypotheses_consolidees"] = hyp
        return hyp

    def _rag(self, hyp: list, m3: Dict) -> str:
        if m3.get("ratio_scr_pct", 0) < SEUIL_SCR or m3.get("ratio_mcr_pct", 0) < 100:
            return "ROUGE"
        non_val = [h for h in hyp if h["statut"] == "NON VALIDÉE" and h["critique"]]
        if non_val:
            return "ROUGE"
        a_just = [h for h in hyp if h["statut"] == "À JUSTIFIER"]
        return "AMBRE" if a_just else "VERT"

    def _avis(self, rag: str, m3: Dict) -> Tuple[str, str]:
        if rag == "VERT":
            return "FAVORABLE", (
                f"Le portefeuille santé présente un profil de risque maîtrisé. "
                f"Ratio SCR = {m3.get('ratio_scr_pct', 0):.1f}% — solvabilité conforme Art.129 S2. "
                f"Loss Ratio = {m3.get('loss_ratio', 0)*100:.1f}% — sinistralité dans la norme FNMF 2023."
            )
        elif rag == "AMBRE":
            return "AVEC RÉSERVES", (
                f"Le portefeuille santé nécessite une attention sur les points signalés. "
                f"Ratio SCR = {m3.get('ratio_scr_pct', 0):.1f}%. "
                "Des ajustements tarifaires ou de provisionnement peuvent être nécessaires."
            )
        else:
            return "DÉFAVORABLE", (
                f"Situation nécessitant une action corrective immédiate. "
                f"Ratio SCR = {m3.get('ratio_scr_pct', 0):.1f}% — insuffisant (seuil 100% Art.129 S2). "
                "Plan de rétablissement à soumettre à l'ACPR."
            )

    # ── Commentaire fallback ──────────────────────────────────────────────────
    def _commentaire(self, rag: str, m3: Dict, m4: Dict,
                     entite: str, date_arrete: str) -> str:
        ic = "🟢" if rag == "VERT" else ("🟡" if rag == "AMBRE" else "🔴")
        avis = m4.get("avis_actuariel", "—")
        L = [
            "=" * 70,
            f"  RAPPORT ACTUARIEL SANTÉ — {self.VERSION}",
            f"  {ic} STATUT : {rag} | {entite} | Arrêté {date_arrete}",
            f"  AVIS : {avis}",
            "=" * 70, "",
            "📊 BILAN SANTÉ", "─" * 50,
            f"  Best Estimate     : {_f(m3.get('be_sante'))}",
            f"  Risk Adjustment   : {_f(m3.get('risk_adjustment'))} (CoC 6% EIOPA §B91)",
            f"  TP Santé          : {_f(m3.get('tp_sante'))}",
            f"  SCR Santé NSLT    : {_f(m3.get('scr_sante'))}",
            f"  MCR Santé         : {_f(m3.get('mcr_sante'))}",
            f"  Ratio SCR         : {_pct(m3.get('ratio_scr_pct'))}",
            f"  Loss Ratio observé: {_pct(m3.get('loss_ratio', 0) * 100)}",
            f"  ANI 2013          : {'✅ Conforme' if m3.get('ani_conforme') else '❌ Non conforme'}",
            "", "📋 HYPOTHÈSES", "─" * 50,
        ]
        for h in m4.get("hypotheses", []):
            ic_h = "✅" if h["statut"] == "VALIDÉE" else ("⚠️" if h["statut"] == "À JUSTIFIER" else "❌")
            L += [f"  {ic_h} [{h['id']}] {h['hypothese']}",
                  f"       → {h['valeur']} : {h['statut']}"]
        return "\n".join(L)

    # ── Graphiques ───────────────────────────────────────────────────────────
    def _graphiques(self, m3: Dict, m4: Dict) -> Dict:
        gph = {}
        rag = m4.get("statut_rag", "AMBRE")
        col_rag = {"VERT": VERT, "AMBRE": ORANGE, "ROUGE": ROUGE}.get(rag, SLATE)
        LAYOUT = dict(
            paper_bgcolor=NAVY, plot_bgcolor=NAVY_L,
            font=dict(family="Inter, Arial", color=BLANC, size=11),
            margin=dict(l=16, r=16, t=60, b=60), height=300,
            hoverlabel=dict(bgcolor="#243F6A", bordercolor=GOLD, font_size=12, font_color=BLANC),
        )

        # G1 — BE / RA / TP
        try:
            fig = go.Figure(go.Bar(
                x=["Best Estimate", "Risk Adjustment", "TP Santé"],
                y=[m3.get("be_sante", 0), m3.get("risk_adjustment", 0), m3.get("tp_sante", 0)],
                marker_color=[BLEU, VIOLET, GOLD],
                text=[_f(v) for v in [m3.get("be_sante", 0), m3.get("risk_adjustment", 0), m3.get("tp_sante", 0)]],
                textposition="outside",
                textfont=dict(color=BLANC, size=10),
            ))
            fig.update_layout(**LAYOUT,
                title=dict(text="G1 — BE / Risk Adjustment / TP Santé (IFRS 17 §B91)", font=dict(color=GOLD, size=12), x=0.01),
                showlegend=False,
                xaxis=dict(tickfont=dict(color=BLANC), showgrid=False),
                yaxis=dict(visible=False),
            )
            gph["be_ra_tp"] = fig
        except Exception as e:
            self.logger.warning(f"G1 : {e}")

        # G2 — Jauge Ratio SCR
        try:
            ratio = m3.get("ratio_scr_pct", 0)
            c_g = VERT if ratio >= SEUIL_SCR_CIBLE else (ORANGE if ratio >= SEUIL_SCR else ROUGE)
            fig = go.Figure(go.Indicator(
                mode="gauge+number", value=ratio,
                number=dict(suffix=" %", font=dict(color=c_g, size=28), valueformat=".1f"),
                title=dict(text="Ratio SCR Santé", font=dict(color=c_g, size=12)),
                gauge=dict(
                    axis=dict(range=[0, max(300, ratio * 1.2)], tickcolor=SLATE),
                    bar=dict(color=c_g, thickness=0.25),
                    bgcolor=NAVY_L,
                    steps=[
                        dict(range=[0, 100],  color="rgba(192,57,43,0.15)"),
                        dict(range=[100, 130], color="rgba(230,126,34,0.15)"),
                        dict(range=[130, max(300, ratio * 1.2)], color="rgba(30,132,73,0.12)"),
                    ],
                    threshold=dict(line=dict(color=VERT, width=3), thickness=0.8, value=SEUIL_SCR_CIBLE),
                ),
            ))
            fig.update_layout(**LAYOUT,
                title=dict(text="G2 — Solvabilité S2 — Santé NSLT", font=dict(color=GOLD, size=12)),
            )
            gph["jauge_scr"] = fig
        except Exception as e:
            self.logger.warning(f"G2 : {e}")

        # G3 — Décomposition SCR
        try:
            fig = go.Figure(go.Bar(
                x=["SCR primes", "SCR réserves", "SCR catastrophe", "SCR Total"],
                y=[m3.get("scr_prem", 0), m3.get("scr_res", 0), m3.get("scr_cat", 0), m3.get("scr_sante", 0)],
                marker_color=[BLEU, VIOLET, ORANGE, GOLD],
                text=[_f(v) for v in [m3.get("scr_prem", 0), m3.get("scr_res", 0), m3.get("scr_cat", 0), m3.get("scr_sante", 0)]],
                textposition="outside",
                textfont=dict(color=BLANC, size=10),
            ))
            fig.update_layout(**LAYOUT,
                title=dict(text="G3 — Décomposition SCR Santé NSLT (Art.148 RD 2015/35)", font=dict(color=GOLD, size=12), x=0.01),
                showlegend=False,
                xaxis=dict(tickfont=dict(color=BLANC), showgrid=False),
                yaxis=dict(visible=False),
            )
            gph["scr_decomposition"] = fig
        except Exception as e:
            self.logger.warning(f"G3 : {e}")

        # G4 — Scorecard hypothèses
        try:
            fig = go.Figure()
            for h in m4.get("hypotheses", []):
                c = VERT if h["statut"] == "VALIDÉE" else (ORANGE if h["statut"] == "À JUSTIFIER" else ROUGE)
                ic_h = "✅" if h["statut"] == "VALIDÉE" else ("⚠️" if h["statut"] == "À JUSTIFIER" else "❌")
                s = 1.0 if h["statut"] == "VALIDÉE" else (0.5 if h["statut"] == "À JUSTIFIER" else 0.0)
                fig.add_trace(go.Bar(
                    x=[s], y=[h["hypothese"][:45]], orientation="h",
                    marker_color=c, width=0.5, opacity=0.85,
                    text=f"{ic_h} {h['statut']}", textposition="outside",
                    textfont=dict(color=c, size=10), showlegend=False,
                    hovertemplate=f"<b>{h['hypothese']}</b><br>{h.get('valeur', '')}<extra></extra>",
                ))
            fig.update_layout(**LAYOUT,
                title=dict(text="G4 — Scorecard Rapport Santé", font=dict(color=col_rag, size=12), x=0.01),
                xaxis=dict(range=[0, 1.6], visible=False),
                yaxis=dict(tickfont=dict(color=BLANC, size=9), showgrid=False),
                barmode="overlay", height=280,
            )
            gph["scorecard"] = fig
        except Exception as e:
            self.logger.warning(f"G4 : {e}")

        return gph

    # ── Audit ─────────────────────────────────────────────────────────────────
    def _audit(self, aid: str, m3: Dict, rag: str, session_hash: str):
        try:
            entry = {
                "audit_id":    aid,
                "timestamp":   datetime.now().isoformat(),
                "statut_rag":  rag,
                "session_hash": session_hash,
                "be_sante":    round(m3.get("be_sante", 0), 2),
                "scr_sante":   round(m3.get("scr_sante", 0), 2),
                "ratio_scr":   round(m3.get("ratio_scr_pct", 0), 1),
                "loss_ratio":  round(m3.get("loss_ratio", 0), 4),
            }
            log = self.audit_path / "rapport_sante_audit.jsonl"
            with open(log, "a", encoding="utf-8") as f:
                f.write(json.dumps(entry, ensure_ascii=False) + "\n")
        except Exception:
            pass

    def _console(self, aid: str, rag: str, avis: str, m3: Dict, session_hash: str):
        ic = "🟢" if rag == "VERT" else ("🟡" if rag == "AMBRE" else "🔴")
        self.logger.info(
            f"[{aid}] {ic} {rag} | {avis} | "
            f"BE={m3.get('be_sante', 0):,.0f}€ | "
            f"SCR={m3.get('scr_sante', 0):,.0f}€ | "
            f"Ratio={m3.get('ratio_scr_pct', 0):.1f}% | "
            f"LR={m3.get('loss_ratio', 0)*100:.1f}% | "
            f"Hash={session_hash}"
        )

    # ── Erreur ────────────────────────────────────────────────────────────────
    def _erreur(self, msg: str, aid: str = "") -> Dict:
        return {
            "success":      False,
            "agent":        self.NOM,
            "version":      self.VERSION,
            "audit_id":     aid,
            "statut_rag":   "ROUGE",
            "session_hash": "",
            "be_sante":     0.0,
            "risk_adjustment": 0.0,
            "tp_sante":     0.0,
            "scr_sante":    0.0,
            "mcr_sante":    0.0,
            "ratio_scr_pct": 0.0,
            "loss_ratio":   0.0,
            "ani_conforme": False,
            "avis_actuariel": "DÉFAVORABLE",
            "modules_disponibles": [],
            "m1": {}, "m2": {}, "m3": {}, "m4": {},
            "html_bytes":   b"",
            "pdf_bytes":    b"",
            "word_bytes":   b"",
            "excel_bytes":  b"",
            "hypotheses":   [],
            "commentaire":  "",
            "graphiques":   {},
            "duree_sec":    0.0,
            "erreur":       msg,
        }
