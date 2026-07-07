# =============================================================================
#  ActuarIA — Agent Rapport Santé SP v1.0
#  sp_rapport_sante.py
#
#  3 formats : HTML · PDF (weasyprint) · Word (.docx)
#  Narration via Claude API (claude-sonnet-4-6) — System Prompt actuaire Santé
#
#  Inputs  : result_s1 (S1 Léonie) · result_s2 (S2 Selma) · result_s3 (S3 Binta)
#             result_reg2 (IFRS17) · result_reg3 (ANI/100%Santé) [optionnels]
#  Outputs : str HTML · bytes PDF · bytes Word
# =============================================================================

from __future__ import annotations
import base64, io, logging, os, re
from datetime import datetime
from typing import Dict, Optional

import numpy as np

logger = logging.getLogger("actuaria.sp.rapport_sante")

# ── Réimporter logo et CSS depuis n5_rapport pour cohérence visuelle ──────────
try:
    from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
        LOGO_SVG, LOGO_URI, _css,
        _f, _pct, _s, _clean, _statut_col, _statut_label,
        _md_to_html, _nettoyer_narration,
    )
    _N5_IMPORT = True
except ImportError:
    _N5_IMPORT = False
    # Fallback minimal si n5 non accessible
    NAVY = "#0B1E3D"; GOLD = "#C9A84C"; BLANC = "#FFFFFF"
    ROUGE = "#C0392B"; VERT = "#1E8449"; ORANGE = "#E67E22"; SLATE = "#8A9BB0"
    LOGO_URI = ""

    def _f(v, dec=0):
        if v is None: return "—"
        try:
            fv = float(v)
            return f"{fv:,.{dec}f} €".replace(",", "\u202f") if dec > 0 else f"{fv:,.0f} €".replace(",", "\u202f")
        except: return "—"

    def _pct(v, dec=1):
        if v is None: return "—"
        try: return f"{float(v):.{dec}f} %"
        except: return "—"

    def _s(v): return str(v).strip() if v else "—"
    def _clean(txt): return re.sub(r"\n{3,}", "\n\n", str(txt or "")).strip()
    def _statut_col(s): return {"VERT":"#1E8449","AMBRE":"#E67E22","ROUGE":"#C0392B"}.get(s,"#8A9BB0")
    def _statut_label(s):
        return {"VERT":"Statut RAG : Vert — Conforme","AMBRE":"Statut RAG : Ambre — Vigilance",
                "ROUGE":"Statut RAG : Rouge — Surveillance renforcée"}.get(s,"Statut : " + s)
    def _nettoyer_narration(t): return t or ""
    def _md_to_html(t): return f'<p style="font-size:9.5pt;">{t}</p>' if t else ""
    def _css(): return "<style>body{font-family:Inter,sans-serif;}</style>"


# =============================================================================
#  SYSTEM PROMPT CLAUDE API — ACTUAIRE SANTÉ SENIOR
# =============================================================================

SYSTEM_PROMPT_SANTE = (
    "Tu es un actuaire Santé senior certifié par l'Institut des Actuaires (IA), "
    "spécialiste en complémentaire santé individuelle et collective, avec 20 ans "
    "d'expérience auprès de mutuelles, institutions de prévoyance et organismes "
    "complémentaires d'assurance maladie (OCAM).\n\n"
    "Tu rédiges le commentaire actuariel d'un rapport de tarification et "
    "provisionnement Santé destiné au Conseil d'Administration et à l'actuaire désigné. "
    "Ce rapport sera soumis à l'ACPR dans le cadre du reporting Solvabilité 2 "
    "(module NSLT) et de la communication IFRS 17.\n\n"
    "CONTEXTE RÉGLEMENTAIRE ET TECHNIQUE :\n"
    "- Solvabilité 2 : module Santé Non-SLT (NSLT), Art. 148-162 RD 2015/35, "
    "  facteur sigma NSLT, formule standard SCR Santé, QRT S.13.01\n"
    "- IFRS 17 : BE santé, Risk Adjustment CoC 6%, CSM, PAA vs GMM\n"
    "- ANI 2013 (Art. L911-7 CSS) : panier minimum complémentaire collective, "
    "  médecine/hospit/dentaire/optique — applicable uniquement aux contrats "
    "  collectifs obligatoires\n"
    "- 100% Santé (réforme 2019) : paniers A/B/C, optique/dentaire/audiologie, "
    "  reste à charge zéro\n"
    "- Tables de référence : DREES Comptes de la Santé 2023, FNMF Rapport "
    "  sinistralité mutuelles 2023, CCAM, NGAP\n"
    "- Cadence de règlement santé : 97% des sinistres liquidés en 3 mois "
    "  (vs 18-36 mois IARD, 60+ mois prévoyance)\n"
    "- Ratio S/P normatif : [65%-85%] mutuelles santé\n"
    "- LR marché FNMF 2023 : eco=55-65%, confort=65-75%, premium/luxe=75-85%\n\n"
    "RÈGLES ABSOLUES :\n"
    "0. FORMAT : §N — TITRE pour les sections, ### sous-titres, **gras** termes clés, "
    "   - tirets pour listes. PAS de tableaux Markdown. PAS de blockquotes.\n"
    "1. LANGUE : Français professionnel. Anglais pour termes consacrés seulement.\n"
    "2. RIGUEUR : Chaque affirmation justifiée par les données. Citer les articles réglementaires.\n"
    "3. CHIFFRES : En euros avec séparateurs. Pourcentages avec une décimale.\n"
    "4. RÉFÉRENCES : Art. 148-162 S2 (NSLT), Art. L911-7 CSS (ANI), DREES 2023, FNMF 2023.\n"
    "5. ALERTES : Ne jamais minimiser. Implication réelle pour le bilan S2.\n"
    "6. CAUSALITÉ : LR > 85% -> sinistralité hors norme -> impact BE et SCR NSLT.\n"
    "7. INCERTITUDE : Quantifier via A/E ratio, CV, fourchettes FNMF.\n"
    "8. POSTURE : Assertif mais prudent. Recommandations claires et actionnables.\n"
    "9. INTERDIT : Phrases génériques sans données.\n"
    "10. SPÉCIFIQUE SANTÉ : Distinguer PSAP vs IBNR (faible en santé : 5-25%) vs PREC. "
    "    Ne pas confondre LR tarifaire (S1) et LR observé (S2).\n\n"
    "STRUCTURE OBLIGATOIRE EN 7 SECTIONS :\n"
    "§1 — CONTEXTE ET PORTEFEUILLE SANTÉ\n"
    "§2 — TARIFICATION ET COHÉRENCE DU LOSS RATIO\n"
    "§3 — PROVISIONNEMENT ET CADENCE DE RÈGLEMENT\n"
    "§4 — CONFORMITÉ RÉGLEMENTAIRE S2/IFRS 17/ANI\n"
    "§5 — HYPOTHÈSES ACTUARIELLES ET BENCHMARKS MARCHÉ\n"
    "§6 — RISQUES IDENTIFIÉS ET POINTS DE VIGILANCE\n"
    "§7 — CONCLUSION ET RECOMMANDATIONS POUR LE CONSEIL D'ADMINISTRATION"
)


# =============================================================================
#  CONSTRUCTION DU CONTEXTE CLAUDE
# =============================================================================

def _construire_contexte(s1, s2, s3, reg2, reg3, arrete):
    pa   = float(s1.get("primes_acquises",  0) or 0)
    lr_t = float(s1.get("ratio_sp_attendu", 0) or 0)
    be   = float(s2.get("be_sante",         0) or 0)
    psap = float(s2.get("psap_dossiers",    0) or 0)
    ibnr = float(s2.get("psap_ibnr",        0) or 0)
    prec = float(s2.get("prec",             0) or 0)
    ra   = float(s2.get("risk_adjustment",  0) or 0)
    tp   = float(s2.get("tp_sante",         0) or 0)
    lr_o = float(s2.get("loss_ratio",       0) or 0)
    scr  = float(s3.get("scr_sante",        0) or 0)
    mcr  = float(s3.get("mcr",              0) or 0)
    r_scr= float(s3.get("ratio_scr_pct",    0) or 0)

    lines = [
        f"DOSSIER SANTÉ — Arrêté {arrete}",
        f"Contrat : {s1.get('contrat','—')} | Garantie : {s1.get('garantie_niveau','—')} "
        f"| {s1.get('nb_assures','—')} assurés | Âge moyen : {s1.get('age_moyen','—')} ans",
        "",
        "=== TARIFICATION (S1 Léonie — DREES 2023) ===",
        f"Primes acquises : {_f(pa)}",
        f"Loss Ratio tarifaire (attendu) : {_pct(lr_t*100)} | Norme : [65%-85%]",
        "",
        "=== PROVISIONNEMENT (S2 Selma) ===",
        f"PSAP dossiers : {_f(psap)}",
        f"IBNR santé (5-25% — liquidation <3 mois) : {_f(ibnr)} | {_pct(ibnr/psap*100 if psap else 0)} du PSAP",
        f"PREC : {_f(prec)}",
        f"Best Estimate Santé : {_f(be)} | {_pct(be/pa*100 if pa else 0)} des primes",
        f"Risk Adjustment (CoC 8%) : {_f(ra)}",
        f"TP Santé : {_f(tp)}",
        f"Loss Ratio observé : {_pct(lr_o*100)} | Norme ≤ 85%",
        "",
        "=== RÉGLEMENTATION S2 (S3 Binta) ===",
        f"SCR Santé NSLT (Art. 148-162 S2) : {_f(scr)} | {_pct(scr/be*100 if be else 0)} BE",
        f"MCR Santé : {_f(mcr)}",
        f"Ratio SCR : {_pct(r_scr)} | Cible >= 130%",
        "",
        "=== HYPOTHÈSES S1 ===",
    ]
    for h in s1.get("hypotheses", []):
        lines.append(f"  {h.get('id','?')} [{h.get('statut','?')}] : {h.get('valeur','')}")
    lines += ["", "=== HYPOTHÈSES S2 ==="]
    for h in s2.get("hypotheses", []):
        lines.append(f"  {h.get('id','?')} [{h.get('statut','?')}] : {h.get('valeur','')}")
    lines += ["", "=== HYPOTHÈSES S3 ==="]
    for h in s3.get("hypotheses", []):
        lines.append(f"  {h.get('id','?')} [{h.get('statut','?')}] : {h.get('valeur','')}")
    ani = s1.get("ani_detail", {})
    lines += ["", "=== CONFORMITÉ ANI 2013 ===",
              f"Note globale : {ani.get('note_globale','—')}"]
    for poste, d in ani.get("detail", {}).items():
        lines.append(f"  {poste} : {d.get('note','—')}")
    lines.append("")
    lines.append("Rédige le commentaire actuariel complet en 7 sections.")
    return "\n".join(lines)


# =============================================================================
#  NARRATION
# =============================================================================

def _narration_claude_api(contexte):
    try:
        import anthropic
        api_key = os.environ.get("ANTHROPIC_API_KEY")
        if not api_key:
            try:
                import streamlit as st
                api_key = st.secrets.get("ANTHROPIC_API_KEY")
            except Exception:
                pass
        if not api_key:
            raise ValueError("ANTHROPIC_API_KEY non définie")
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=10000,
            system=SYSTEM_PROMPT_SANTE,
            messages=[{"role": "user", "content": contexte}],
        )
        return resp.content[0].text
    except Exception as e:
        logger.warning(f"Claude API Santé indisponible : {e}")
        raise

def _generer_narration(s1, s2, s3, reg2, reg3, arrete):
    try:
        ctx = _construire_contexte(s1, s2, s3, reg2, reg3, arrete)
        txt = _narration_claude_api(ctx)
        if txt:
            return txt, "claude_api"
    except Exception:
        pass
    for src in [s2, s1, s3]:
        com = _clean(src.get("commentaire", ""))
        if com:
            return com, "templates"
    return "", "aucune"


# =============================================================================
#  EXPORT HTML
# =============================================================================

def export_html(result_s1, result_s2, result_s3,
                result_reg2=None, result_reg3=None,
                ref_client="", arrete="", audit_id="") -> str:
    try:
        s1 = result_s1 or {}; s2 = result_s2 or {}; s3 = result_s3 or {}
        reg2 = result_reg2 or {}; reg3 = result_reg3 or {}

        dt  = datetime.now().strftime("%d/%m/%Y")
        arr = arrete or dt
        cli = ref_client or "À renseigner"

        rags = [s1.get("statut_rag","VERT"), s2.get("statut_rag","VERT"), s3.get("statut_rag","VERT")]
        rag  = "ROUGE" if "ROUGE" in rags else ("AMBRE" if "AMBRE" in rags else "VERT")
        s_label = _statut_label(rag)
        s_cls   = rag.lower()

        pa   = float(s1.get("primes_acquises",  0) or 0)
        lr_t = float(s1.get("ratio_sp_attendu", 0) or 0)
        be   = float(s2.get("be_sante",         0) or 0)
        psap = float(s2.get("psap_dossiers",    0) or 0)
        ibnr = float(s2.get("psap_ibnr",        0) or 0)
        prec = float(s2.get("prec",             0) or 0)
        ra   = float(s2.get("risk_adjustment",  0) or 0)
        tp   = float(s2.get("tp_sante",         0) or 0)
        lr_o = float(s2.get("loss_ratio",       0) or 0)
        scr  = float(s3.get("scr_sante",        0) or 0)
        mcr  = float(s3.get("mcr",              0) or 0)
        fpp  = float(s3.get("fonds_propres",    0) or 0)
        r_scr= float(s3.get("ratio_scr_pct",    0) or 0)
        r_mcr= float(s3.get("ratio_mcr_pct",    0) or 0)

        narration, source = _generer_narration(s1, s2, s3, reg2, reg3, arr)
        narration_html    = _md_to_html(narration)
        source_badge = {"claude_api": "✦ ActuarIA Intelligence",
                        "templates": "📝 Mode standard", "aucune": ""}.get(source, "")

        # KPIs page de garde
        logo_src = LOGO_URI if _N5_IMPORT else ""
        kpis_html = (
            f'<div class="garde-kpi"><div class="kpi-label">Client</div>'            f'<div class="kpi-value">{cli}</div></div>'            f'<div class="garde-kpi"><div class="kpi-label">Arrêté</div>'            f'<div class="kpi-value">{arr}</div></div>'            f'<div class="garde-kpi"><div class="kpi-label">Primes acquises</div>'            f'<div class="kpi-value highlight">{_f(pa)}</div></div>'            f'<div class="garde-kpi"><div class="kpi-label">LR tarifaire S1</div>'            f'<div class="kpi-value">{_pct(lr_t*100)}</div>'            f'<div class="kpi-sub">Norme [65%-85%]</div></div>'            f'<div class="garde-kpi"><div class="kpi-label">BE Santé | Ratio SCR</div>'            f'<div class="kpi-value">{_f(be)}</div>'            f'<div class="kpi-sub">{_pct(r_scr)} (cible ≥130%)</div></div>'
        )

        # Hypothèses combinées S1+S2+S3
        all_hyp = ([(h,"S1") for h in s1.get("hypotheses",[])] +
                   [(h,"S2") for h in s2.get("hypotheses",[])] +
                   [(h,"S3") for h in s3.get("hypotheses",[])])
        hyp_cards = ""
        for h, agent in all_hyp:
            statut = h.get("statut","")
            cls = ("hyp-ok" if statut=="VALIDÉE" else
                   "hyp-err" if statut=="NON VALIDÉE" else "hyp-warn")
            lbl = {"VALIDÉE":"✓ VALIDÉE","NON VALIDÉE":"✗ NON VALIDÉE"}.get(statut,"⚠ "+statut)
            hyp_cards += (
                f'<div class="hyp-card {cls}">'                f'<div class="hyp-label">'                f'<div class="hyp-code">[{agent}] {h.get("id","?")} — {h.get("hypothese","")[:40]}</div>'                f'<div class="hyp-code" style="margin-top:4px;">{lbl}</div></div>'                f'<div class="hyp-text">{_s(h.get("valeur",""))}</div></div>'
            )

        # Tableau provisions
        tbl_prov = (
            '<table class="premium"><thead><tr>'            '<th>Composante</th><th class="right">Montant</th>'            '<th class="right">% Primes</th><th>Référence</th></tr></thead><tbody>'
            f'<tr><td class="label">PSAP Dossiers</td>'            f'<td class="right"><span class="mono">{_f(psap)}</span></td>'            f'<td class="right">{_pct(psap/pa*100 if pa else 0)}</td>'            f'<td>Sinistres déclarés en instance</td></tr>'            f'<tr><td class="label">IBNR Santé (5-25% — liquidation &lt;3 mois)</td>'            f'<td class="right"><span class="mono">{_f(ibnr)}</span></td>'            f'<td class="right">{_pct(ibnr/pa*100 if pa else 0)}</td>'            f'<td>Art. 77 S2 — cadence rapide santé</td></tr>'            f'<tr><td class="label">PREC</td>'            f'<td class="right"><span class="mono">{_f(prec)}</span></td>'            f'<td class="right">{_pct(prec/pa*100 if pa else 0)}</td>'            f'<td>Art. 78 S2</td></tr>'            f'<tr><td class="label">Best Estimate Santé</td>'            f'<td class="right"><span class="mono">{_f(be)}</span></td>'            f'<td class="right">{_pct(be/pa*100 if pa else 0)}</td>'            f'<td>Art. 77 §1 Directive S2</td></tr>'            f'<tr><td class="label">Risk Adjustment (CoC 8% — floor 3%)</td>'            f'<td class="right"><span class="mono">{_f(ra)}</span></td>'            f'<td class="right">{_pct(ra/pa*100 if pa else 0)}</td>'            f'<td>IFRS 17 §B119</td></tr>'            f'<tr class="highlight-gold"><td class="label">TP Santé (BE + RA)</td>'            f'<td class="right"><span class="mono">{_f(tp)}</span></td>'            f'<td class="right">{_pct(tp/pa*100 if pa else 0)}</td>'            f'<td>Art. 77 §1 S2 — bilan prudentiel</td></tr>'            '</tbody></table>'
        )

        # Tableau SCR
        lr_stat  = "✅ CONFORME" if lr_o<=0.85 else ("⚠️ VIGILANCE" if lr_o<=0.95 else "❌ ALERTE")
        scr_stat = "✅ CONFORME" if r_scr>=130 else ("⚠️ VIGILANCE" if r_scr>=100 else "❌ INSUFFISANT")
        tbl_scr = (
            '<table class="premium"><thead><tr>'            '<th>Composante</th><th class="center">Valeur</th>'            '<th class="center">Statut</th><th>Référence</th></tr></thead><tbody>'
            f'<tr><td class="label">Loss Ratio observé</td>'            f'<td class="center"><span class="mono">{_pct(lr_o*100)}</span></td>'            f'<td class="center">{lr_stat}</td><td>Norme [65%-85%] mutuelles</td></tr>'            f'<tr><td class="label">SCR Santé NSLT</td>'            f'<td class="center"><span class="mono">{_f(scr)}</span></td>'            f'<td class="center">{_pct(scr/be*100 if be else 0)} BE</td>'            f'<td>Art. 148-162 RD 2015/35 — QRT S.13.01</td></tr>'            f'<tr><td class="label">MCR Santé</td>'            f'<td class="center"><span class="mono">{_f(mcr)}</span></td>'            f'<td class="center">—</td><td>Art. 252 RD 2015/35</td></tr>'            f'<tr><td class="label">Fonds Propres</td>'            f'<td class="center"><span class="mono">{_f(fpp)}</span></td>'            f'<td class="center">—</td><td>Art. 87 S2</td></tr>'            f'<tr class="highlight-gold"><td class="label">Ratio SCR</td>'            f'<td class="center" style="font-weight:700;">{_pct(r_scr)}</td>'            f'<td class="center">{scr_stat}</td><td>Cible pratique ≥ 130%</td></tr>'            f'<tr class="highlight-gold"><td class="label">Ratio MCR</td>'            f'<td class="center" style="font-weight:700;">{_pct(r_mcr)}</td>'            f'<td class="center">{("✅" if r_mcr>=100 else "❌")}</td>'            f'<td>Art. 139 S2 — seuil déclencheur</td></tr>'            '</tbody></table>'
        )

        # ANI
        ani = s1.get("ani_detail", {})
        ani_note = ani.get("note_globale", "—")
        ani_ok = "conforme" in ani_note.lower()
        ani_col = "#1E8449" if ani_ok else "#C0392B"
        ani_block = (
            f'<div style="background:{ani_col};color:#fff;padding:12px 18px;'            f'border-radius:6px;font-size:9pt;font-weight:600;margin-bottom:16px;">'            f'{ani_note}</div>'
        )
        ani_tbl = (
            '<table class="premium"><thead><tr>'            '<th>Poste</th><th class="center">Seuil ANI (€)</th>'            '<th class="center">Charge mutuelle (€)</th>'            '<th class="center">Conforme</th><th>Note</th>'            '</tr></thead><tbody>'
        )
        for i, (poste, d) in enumerate(ani.get("detail", {}).items()):
            ok = d.get("ok", True)
            ani_tbl += (
                f'<tr><td class="label">{poste}</td>'                f'<td class="center"><span class="mono">{_f(d.get("seuil",0))}</span></td>'                f'<td class="center"><span class="mono">{_f(d.get("charge",0))}</span></td>'                f'<td class="center">{("✅" if ok else "❌")}</td>'                f'<td>{_s(d.get("note",""))}</td></tr>'
            )
        ani_tbl += "</tbody></table>"

        # Tableau tarification postes
        postes = s1.get("postes", {})
        tbl_postes = (
            '<table class="premium"><thead><tr>'            '<th>Poste</th><th class="right">Fréq./an</th>'            '<th class="right">Coût moyen</th><th class="right">Charge mutuelle</th>'            '<th class="right">Sinistre annuel</th><th class="center">Source</th>'            '</tr></thead><tbody>'
        )
        total_sin = 0.0
        for i, (poste, v) in enumerate(postes.items()):
            sin = float(v.get("sinistre_annuel",0) or 0)
            total_sin += sin
            tbl_postes += (
                f'<tr><td class="label">{poste}</td>'                f'<td class="right">{v.get("frequence_an","—")}</td>'                f'<td class="right"><span class="mono">{_f(v.get("cout_moyen"))}</span></td>'                f'<td class="right"><span class="mono">{_f(v.get("charge_mutuelle"))}</span></td>'                f'<td class="right"><span class="mono">{_f(sin)}</span></td>'                f'<td class="center">{v.get("source","—")}</td></tr>'
            )
        tbl_postes += (
            f'<tr class="highlight-gold"><td class="label">TOTAL</td>'            f'<td class="right">—</td><td class="right">—</td><td class="right">—</td>'            f'<td class="right"><span class="mono">{_f(total_sin)}</span></td>'            f'<td class="center">—</td></tr></tbody></table>'
        )

        # Alertes et recommandations
        alertes = ""
        for src_dict in [s1, s2, s3]:
            for h in src_dict.get("hypotheses", []):
                if h.get("statut") in ("NON VALIDÉE", "À JUSTIFIER"):
                    txt = _clean(h.get("valeur", ""))
                    if txt:
                        cls = "hyp-err" if h.get("statut")=="NON VALIDÉE" else "hyp-warn"
                        alertes += (
                            f'<div class="hyp-card {cls}" style="margin-bottom:10px;">'                            f'<div class="hyp-label"><div class="hyp-code">Point de vigilance</div></div>'                            f'<div class="hyp-text">{txt}</div></div>'
                        )
        if not alertes:
            alertes = '<p style="color:#1E8449;font-size:9pt;">✅ Aucun point de vigilance majeur identifié.</p>'

        # Assemblage HTML final
        css_str = _css() if callable(_css) else ""
        html = (
            '<!DOCTYPE html>\n<html lang="fr">\n<head>\n'            '<meta charset="UTF-8">\n'            f'<title>Rapport Santé — {cli} — {arr}</title>\n'            '<script src="https://cdn.jsdelivr.net/npm/plotly.js-dist@2.26.0/plotly.min.js"></script>\n'            + css_str +
            '\n</head>\n<body>\n<div class="rapport-container">\n'
            # PAGE DE GARDE
            '<div class="page-garde">\n'            '<div class="garde-bg"><div class="garde-dots"></div>'            '<div class="garde-diagonal"></div><div class="garde-accent-line"></div></div>\n'            '<div class="garde-inner">\n'            '<div class="garde-header">\n'            f'<div class="garde-logo-wrap"><img src="{logo_src}" alt="ActuarIA"/></div>\n'            '<div class="garde-badges"><span class="badge-confidentiel">⬛ Confidentiel</span></div>\n'            '</div>\n'            '<div class="garde-hero">\n'            '<div class="garde-eyebrow">Rapport Actuariel Santé</div>\n'            '<div class="garde-titre">Complémentaire<br><em>Santé</em></div>\n'            f'<div class="garde-subtitle">Arrêté au {arr}</div>\n'            '<div class="garde-sep"><div class="garde-sep-line"></div>'            '<div class="garde-sep-diamond"></div><div class="garde-sep-line"></div></div>\n'            f'<div class="garde-statut garde-statut-{s_cls}">'            f'<div class="statut-dot statut-dot-{s_cls}"></div>'            f'<div class="statut-label-{s_cls}">{s_label}</div></div>\n'            '</div>\n'            f'<div class="garde-footer"><div class="garde-kpis">{kpis_html}</div></div>\n'            '</div>\n</div>\n'
            # CORPS
            '<div class="rapport-body">\n'
            # S1 — Synthèse
            '<div class="section-header"><span class="section-num">01</span>'            '<span class="section-titre">Synthèse exécutive</span></div>\n'            '<div class="section-body">\n'            '<div class="kpi-grid">\n'            f'<div class="kpi-card"><div class="kpi-card-label">Primes acquises</div>'            f'<div class="kpi-card-value">{_f(pa)}</div>'            f'<div class="kpi-card-sub">{s1.get("nb_assures","—")} assurés</div></div>\n'            f'<div class="kpi-card"><div class="kpi-card-label">LR tarifaire (S1)</div>'            f'<div class="kpi-card-value">{_pct(lr_t*100)}</div>'            f'<div class="kpi-card-sub">Norme [65%-85%]</div></div>\n'            f'<div class="kpi-card kpi-card-rouge"><div class="kpi-card-label">Best Estimate Santé</div>'            f'<div class="kpi-card-value kpi-card-value-rouge">{_f(be)}</div>'            f'<div class="kpi-card-sub">LR observé : {_pct(lr_o*100)}</div></div>\n'            f'<div class="kpi-card"><div class="kpi-card-label">Ratio SCR</div>'            f'<div class="kpi-card-value">{_pct(r_scr)}</div>'            f'<div class="kpi-card-sub">Cible ≥ 130%</div></div>\n'            '</div>\n</div>\n<div class="section-divider"></div>\n'
            # S2 — Tarification
            '<div class="section-header"><span class="section-num">02</span>'            '<span class="section-titre">Tarification — S1 Léonie (DREES 2023)</span></div>\n'            '<div class="section-body">\n'            '<div class="table-section-title">Sinistralité par poste</div>\n'            + tbl_postes +
            '\n</div>\n<div class="section-divider"></div>\n'
            # S3 — Provisionnement
            '<div class="section-header"><span class="section-num">03</span>'            '<span class="section-titre">Provisionnement — S2 Selma</span></div>\n'            '<div class="section-body">\n'            + tbl_prov +
            '\n</div>\n<div class="section-divider"></div>\n'
            # S4 — Réglementation
            '<div class="section-header"><span class="section-num">04</span>'            '<span class="section-titre">Conformité réglementaire — S2 · IFRS 17 · ANI 2013</span></div>\n'            '<div class="section-body">\n'            + tbl_scr +
            '<div class="table-section-title" style="margin-top:24px;">'            'Conformité ANI 2013 — Panier minimum (Art. L911-7 CSS)</div>\n'            + ani_block + ani_tbl +
            '\n</div>\n<div class="section-divider"></div>\n'
            # S5 — Hypothèses
            '<div class="section-header"><span class="section-num">05</span>'            '<span class="section-titre">Hypothèses actuarielles</span></div>\n'            f'<div class="section-body"><div class="hyp-grid">{hyp_cards}</div></div>\n'            '<div class="section-divider"></div>\n'
            # S6 — Commentaire Claude
            '<div class="section-header"><span class="section-num">06</span>'            '<span class="section-titre">Commentaire actuariel</span></div>\n'            '<div class="section-body">\n'            '<div class="commentaire-wrap">\n'            '<div class="commentaire-header">'            f'<span class="commentaire-header-title">Rapport Santé · Arrêté {arr}</span>'            f'<span class="commentaire-ai-badge">{source_badge}</span></div>\n'            '<div class="commentaire-body">\n'            '<p style="font-size:8pt;color:var(--slate);margin-bottom:20px;font-style:italic;">'            "Commentaire à destination du Conseil d'Administration et de l'Actuaire Désigné.<br>"
            "Document soumis à l'ACPR dans le cadre du reporting Solvabilité 2.</p>\n"
            + narration_html +
            f'\n<div class="comm-footer">✦ Narration générée par ActuarIA Intelligence · SP-Santé v1.0 · {dt}</div>\n'            '</div></div>\n</div>\n<div class="section-divider"></div>\n'
            # S7 — Jugement
            '<div class="section-header"><span class="section-num">07</span>'            '<span class="section-titre">Jugement actuariel &amp; Recommandations</span></div>\n'            f'<div class="section-body"><div class="hyp-grid">{alertes}</div></div>\n'
            '</div>\n'
            # PIED
            f'<div class="pied-de-page">'            f'<div class="pied-logo"><img src="{logo_src}" alt="ActuarIA"/></div>'            f'<div class="pied-meta">{cli} · Santé · Arrêté {arr} · {dt}<br>'            '<span style="font-size:6.5pt;color:#8A9AB0;">'            "Rapport établi conformément aux Art. 77, 105 et 148-162 de la Directive Solvabilité II "
            "et au Guide Institut des Actuaires 2023</span><br>"
            '<span class="confidentiel-footer">CONFIDENTIEL — USAGE STRICTEMENT ACTUARIEL</span>'            '</div></div>\n'            '</div>\n</body>\n</html>'
        )

        logger.info(f"HTML Santé : {len(html):,} chars — narration={source}")
        return html

    except Exception as e:
        logger.error(f"export_html Santé : {e}", exc_info=True)
        return f"<html><body><h1>Erreur : {e}</h1></body></html>"


# =============================================================================
#  EXPORT PDF
# =============================================================================

def export_pdf(result_s1, result_s2, result_s3,
               result_reg2=None, result_reg3=None,
               ref_client="", arrete="", audit_id="") -> bytes:
    try:
        from weasyprint import HTML as WH
        html = export_html(result_s1, result_s2, result_s3,
                           result_reg2, result_reg3, ref_client, arrete, audit_id)
        pdf = WH(string=html).write_pdf()
        logger.info(f"PDF Santé : {len(pdf):,} bytes")
        return pdf
    except ImportError:
        logger.error("weasyprint non installé")
        return b""
    except Exception as e:
        logger.error(f"export_pdf Santé : {e}", exc_info=True)
        return b""


# =============================================================================
#  EXPORT WORD
# =============================================================================

def export_word(result_s1, result_s2, result_s3,
                result_reg2=None, result_reg3=None,
                ref_client="", arrete="", audit_id="") -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        logger.error(f"python-docx absent : {e}")
        return b""

    try:
        s1 = result_s1 or {}; s2 = result_s2 or {}; s3 = result_s3 or {}

        def rgb(h):
            h = h.lstrip("#")
            return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))

        NR=rgb("#0B1E3D"); GR=rgb("#C9A84C"); BR=rgb("#FFFFFF")
        GrR=rgb("#8A9BB0"); RgR=rgb("#C0392B"); VR=rgb("#1E8449"); AR=rgb("#E67E22")

        dt  = datetime.now().strftime("%d/%m/%Y")
        arr = arrete or dt
        cli = ref_client or "À renseigner"

        rags = [s1.get("statut_rag","VERT"), s2.get("statut_rag","VERT"), s3.get("statut_rag","VERT")]
        rag  = "ROUGE" if "ROUGE" in rags else ("AMBRE" if "AMBRE" in rags else "VERT")

        narration, source = _generer_narration(s1, s2, s3,
                                               result_reg2 or {}, result_reg3 or {}, arr)

        pa   = float(s1.get("primes_acquises",  0) or 0)
        lr_t = float(s1.get("ratio_sp_attendu", 0) or 0)
        be   = float(s2.get("be_sante",         0) or 0)
        psap = float(s2.get("psap_dossiers",    0) or 0)
        ibnr = float(s2.get("psap_ibnr",        0) or 0)
        prec = float(s2.get("prec",             0) or 0)
        ra   = float(s2.get("risk_adjustment",  0) or 0)
        tp   = float(s2.get("tp_sante",         0) or 0)
        lr_o = float(s2.get("loss_ratio",       0) or 0)
        scr  = float(s3.get("scr_sante",        0) or 0)
        mcr  = float(s3.get("mcr",              0) or 0)
        fpp  = float(s3.get("fonds_propres",    0) or 0)
        r_scr= float(s3.get("ratio_scr_pct",    0) or 0)
        r_mcr= float(s3.get("ratio_mcr_pct",    0) or 0)

        doc = Document()
        for s in doc.sections:
            s.top_margin=Cm(2); s.bottom_margin=Cm(2)
            s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

        def _bg(cell, hex6):
            tc=cell._tc; tcp=tc.get_or_add_tcPr()
            sd=OxmlElement("w:shd")
            sd.set(qn("w:fill"), hex6.lstrip("#"))
            sd.set(qn("w:color"), "auto"); sd.set(qn("w:val"), "clear")
            tcp.append(sd)

        def _run(p, txt, bold=False, italic=False, sz=10, col=None):
            r=p.add_run(str(txt)); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
            if col: r.font.color.rgb=col
            return r

        def _h(txt, lv=1, col=None):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
            _run(p, txt, bold=True, sz={1:16,2:12,3:10}.get(lv,10),
                 col=col or (NR if lv==1 else GR))

        def _sep():
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
            bo=OxmlElement("w:bottom")
            bo.set(qn("w:val"),"single"); bo.set(qn("w:sz"),"6")
            bo.set(qn("w:space"),"1"); bo.set(qn("w:color"),"C9A84C")
            pBdr.append(bo); pPr.append(pBdr)

        def _tbl(heads, rows, ws=None):
            t=doc.add_table(rows=1+len(rows), cols=len(heads)); t.style="Table Grid"
            for i,hd in enumerate(heads):
                c=t.rows[0].cells[i]; _bg(c,"0B1E3D")
                pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=pp.add_run(str(hd)); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=BR
            for ri,row in enumerate(rows):
                for ci,v in enumerate(row):
                    c=t.rows[ri+1].cells[ci]
                    if ri%2==1: _bg(c,"EEF2F7")
                    pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r=pp.add_run(str(v) if v is not None else "—"); r.font.size=Pt(9)
            if ws:
                for i,w in enumerate(ws):
                    for row in t.rows: row.cells[i].width=Cm(w)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)

        # Page de garde
        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p,"RAPPORT ACTUARIEL SANTÉ\n",bold=True,sz=20,col=NR)
        _run(p,"Complémentaire Santé",bold=True,sz=14,col=GR)
        doc.add_paragraph()
        s_col_r = VR if rag=="VERT" else AR if rag=="AMBRE" else RgR
        p=doc.add_paragraph()
        _run(p,"Statut : ",sz=11,col=NR); _run(p,rag,bold=True,sz=11,col=s_col_r)
        doc.add_paragraph()
        _tbl(["Client","Arrêté","Primes acquises","LR tarifaire"],
             [[cli, arr, _f(pa), _pct(lr_t*100)]],ws=[3.5,2.5,4.0,3.0])
        doc.add_page_break()

        _h("1. Synthèse exécutive"); _sep()
        _tbl(["Indicateur","Valeur","Indicateur","Valeur"],
             [["Primes acquises",_f(pa),"Best Estimate Santé",_f(be)],
              ["LR tarifaire (S1)",_pct(lr_t*100),"LR observé (S2)",_pct(lr_o*100)],
              ["SCR Santé NSLT",_f(scr),"Ratio SCR",_pct(r_scr)],
              ["MCR Santé",_f(mcr),"Ratio MCR",_pct(r_mcr)]],ws=[4.0,3.5,4.0,3.5])
        doc.add_page_break()

        _h("2. Tarification — S1 Léonie (DREES 2023)"); _sep()
        postes = s1.get("postes", {})
        if postes:
            _tbl(["Poste","Fréq./an","Coût moyen","Charge mut.","Sinistre/an","Source"],
                 [[p, v.get("frequence_an","—"), _f(v.get("cout_moyen")),
                   _f(v.get("charge_mutuelle")), _f(v.get("sinistre_annuel")),
                   v.get("source","—")] for p,v in postes.items()],
                 ws=[3.5,2.0,3.0,3.0,3.0,3.0])
        doc.add_page_break()

        _h("3. Provisionnement — S2 Selma"); _sep()
        _tbl(["Composante","Montant","% Primes","Référence"],
             [["PSAP Dossiers",_f(psap),_pct(psap/pa*100 if pa else 0),"Sinistres connus"],
              ["IBNR Santé (5-25%)",_f(ibnr),_pct(ibnr/pa*100 if pa else 0),"Art. 77 S2"],
              ["PREC",_f(prec),_pct(prec/pa*100 if pa else 0),"Art. 78 S2"],
              ["Best Estimate",_f(be),_pct(be/pa*100 if pa else 0),"Art. 77 §1 S2"],
              ["Risk Adjustment",_f(ra),_pct(ra/pa*100 if pa else 0),"IFRS 17 §B119"],
              ["TP Santé",_f(tp),_pct(tp/pa*100 if pa else 0),"Art. 77 §1 S2"]],
             ws=[4.5,3.0,2.5,5.0])
        doc.add_page_break()

        _h("4. Conformité réglementaire — S2 · IFRS 17 · ANI 2013"); _sep()
        _tbl(["Composante","Valeur","Statut","Référence"],
             [["Loss Ratio observé",_pct(lr_o*100),
               "✅" if lr_o<=0.85 else "⚠️","Norme [65%-85%]"],
              ["SCR Santé NSLT",_f(scr),_pct(scr/be*100 if be else 0),"Art. 148-162 S2"],
              ["MCR Santé",_f(mcr),"—","Art. 252 RD 2015/35"],
              ["Fonds Propres",_f(fpp),"—","Art. 87 S2"],
              ["Ratio SCR",_pct(r_scr),"✅" if r_scr>=130 else "⚠️","≥ 130% cible"],
              ["Ratio MCR",_pct(r_mcr),"✅" if r_mcr>=100 else "❌","≥ 100% requis"]],
             ws=[4.5,3.0,2.0,5.5])
        doc.add_page_break()

        _h("5. Hypothèses actuarielles"); _sep()
        all_hyp = ([(h,"S1") for h in s1.get("hypotheses",[])] +
                   [(h,"S2") for h in s2.get("hypotheses",[])] +
                   [(h,"S3") for h in s3.get("hypotheses",[])])
        if all_hyp:
            _tbl(["Src","ID","Hypothèse","Résultat","Statut"],
                 [[ag,h.get("id",""),h.get("hypothese","")[:55],
                   h.get("valeur","")[:75],h.get("statut","")]
                  for h,ag in all_hyp],
                 ws=[1.0,1.5,5.0,5.5,2.5])
        doc.add_page_break()

        _h("6. Commentaire actuariel"); _sep()
        if narration:
            sections_n = re.split(r"(?=§\d+\s*[—\-–])", _clean(narration))
            for sec in sections_n:
                sec = sec.strip()
                if not sec: continue
                ls = sec.split("\n", 1)
                if ls[0]: _h(ls[0], lv=2)
                if len(ls) > 1:
                    for ln in ls[1].split("\n"):
                        ln = ln.strip()
                        if ln:
                            p=doc.add_paragraph()
                            p.paragraph_format.space_after=Pt(3)
                            p.paragraph_format.left_indent=Cm(0.3)
                            _run(p, ln, sz=9, col=NR)
            if source == "claude_api":
                p=doc.add_paragraph()
                _run(p,"✦ Narration générée par ActuarIA Intelligence",sz=7,italic=True,col=GrR)
        else:
            p=doc.add_paragraph()
            _run(p,"Narration non disponible.",sz=9,italic=True)
        doc.add_page_break()

        _h("7. Jugement actuariel & Recommandations"); _sep()
        for src_dict in [s1, s2, s3]:
            for h in src_dict.get("hypotheses", []):
                if h.get("statut") in ("NON VALIDÉE", "À JUSTIFIER"):
                    txt = _clean(h.get("valeur",""))
                    if txt:
                        p=doc.add_paragraph()
                        p.paragraph_format.left_indent=Cm(0.4)
                        ic = "❌" if h.get("statut")=="NON VALIDÉE" else "⚠️"
                        _run(p, f"{ic}  ", sz=10, col=RgR if ic=="❌" else AR)
                        _run(p, txt, sz=9, col=NR)

        _sep()
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"ActuarIA · {cli} · Santé · Arrêté {arr} · {dt} · CONFIDENTIEL",
             sz=7, italic=True, col=GrR)

        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        wb=buf.read()
        logger.info(f"Word Santé : {len(wb):,} bytes")
        return wb

    except Exception as e:
        logger.error(f"export_word Santé : {e}", exc_info=True)
        return b""
