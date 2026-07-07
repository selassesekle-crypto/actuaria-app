# =============================================================================
#  ActuarIA — Agent Rapport Prévoyance SP v2.0
#  sp_rapport_prevoyance.py
#
#  3 formats : HTML · PDF (weasyprint) · Word (.docx)
#  Narration via Claude API (claude-sonnet-4-6) — System Prompt actuaire Prévoyance
#  Style : reproduction exacte du rapport premium n5_rapport.py
#
#  Inputs  : result_p1 · result_p2 · result_p3 · result_p4
#             result_alm · result_reg1 [optionnels]
#  Outputs : str HTML · bytes PDF · bytes Word
# =============================================================================

from __future__ import annotations
import base64, io, logging, os, re
from datetime import datetime
from typing import Dict, Optional, Tuple

import numpy as np

logger = logging.getLogger('actuaria.sp.rapport_prevoyance')

# ── Réimporter logo et CSS depuis n5_rapport pour cohérence visuelle ──────────
try:
    from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
        LOGO_SVG, LOGO_URI, _css,
        _f, _pct, _s, _clean, _statut_col, _statut_label,
        _md_to_html, _nettoyer_narration,
    )
    _N5_IMPORT = True
except ImportError:
    _N5_IMPORT = False
    LOGO_URI = ''

    def _f(v, dec=0):
        if v is None: return '—'
        try:
            fv = float(v)
            sep = '\u202f'
            return f'{fv:,.0f}\u202f\u20ac'.replace(',', sep) if dec == 0 else f'{fv:,.{dec}f}'.replace(',', sep)
        except: return '—'

    def _pct(v, dec=1):
        if v is None: return '—'
        try: return f'{float(v):.{dec}f}\u202f%'
        except: return '—'

    def _s(v):
        if v is None: return '—'
        return re.sub(r'\s+', ' ', str(v)).strip() or '—'

    def _clean(txt):
        if not txt: return ''
        txt = re.sub(r'[■□▪▸►═╔╗╚╝║─]+', '', str(txt))
        txt = re.sub(r'={4,}', '', txt)
        txt = re.sub(r'\n{3,}', '\n\n', txt)
        return txt.strip()

    def _statut_col(s): return {'VERT':'#1E8449','AMBRE':'#E67E22','ROUGE':'#C0392B'}.get(s,'#8A9BB0')
    def _statut_label(s):
        return {'VERT':'Statut RAG : Vert — Conforme','AMBRE':'Statut RAG : Ambre — Vigilance',
                'ROUGE':'Statut RAG : Rouge — Surveillance renforcée'}.get(s,'Statut : ' + s)
    def _nettoyer_narration(t): return t or ''
    def _md_to_html(t):
        if not t: return '<p class="comm-p" style="color:#8A9BB0;font-style:italic;">Narration non disponible.</p>'
        return '<p class="comm-p">' + t + '</p>'
    def _css(): return '<style>body{font-family:Inter,sans-serif;}</style>'

# =============================================================================
#  SYSTEM PROMPT CLAUDE API — ACTUAIRE PRÉVOYANCE SENIOR
# =============================================================================

SYSTEM_PROMPT_PREVOYANCE = """\
Tu es un actuaire Prévoyance senior certifié par l'Institut des Actuaires (IA), \
spécialiste en prévoyance collective (incapacité/invalidité/décès), avec 20 ans \
d'expérience auprès d'institutions de prévoyance, mutuelles et groupes de \
protection sociale. Expert en tables BCAC, modèles de Markov multi-états \
et provisionnement triangle ITT.

Tu rédiges le commentaire actuariel d'un rapport de tarification, \
provisionnement et reporting réglementaire Prévoyance destiné au Conseil \
d'Administration et à l'actuaire désigné. \
Ce rapport sera soumis à l'ACPR dans le cadre du reporting Solvabilité 2 \
(module SLT Invalidité) et de la communication IFRS 17.

CONTEXTE RÉGLEMENTAIRE ET TECHNIQUE :
- Solvabilité 2 : module Santé SLT (invalidité-morbidité), Art. 145 RD 2015/35, \
  chocs morbidité +35% (incidence) et cessation -20% (guérison), QRT S.14.01, \
  SCR Invalidité agrégé avec corrélation ρ=0.25 EIOPA Annexe IV
- MCR Prévoyance : Art. 252 RD 2015/35, coefficients α=0.0338/β=0.0191, \
  plancher absolu 3,7 M€ (vs 2,5 M€ santé), bornes 25%/45% SCR
- IFRS 17 : BE prévoyance = PM Rentes IP + PSAP ITT, Risk Adjustment CoC 6%, \
  floor 3%, CSM prévoyance
- Tables de référence incapacité : BCAC 2019 (CTIP/INSEE), \
  probabilités annuelles q_AI/q_IA/q_IP par âge et CSP
- Tables de mortalité : TH0002/TF0002, TD 88-90 (décès prévoyance)
- Durées moyennes ITT : 8-26 mois selon âge (CTIP/BCAC), carence 90 jours
- Probabilité maintien ITT 6 mois : 20-40% selon âge (BCAC 2019)
- Passage ITT → IP : q_IP|ITT augmente avec l'âge : 4% à 25 ans → 35% à 60 ans
- Ratio BE/PA attendu : ≤ 50% prévoyance collective bien tarifée
- SCR/BE prévoyance : [0.20-1.50] plage normale. Choc dominant = morbidité +35%
- LCR actif/passif : ≥ 100% recommandé. Duration gap : ≤ 2 ans idéal
- Triangle ITT : CL, Mack, BF (LR CTIP 2023), Bootstrap ODP. \
  Horizon liquidation : 60-84 mois (vs 36 mois IARD, 3 mois santé)
- PM Rentes IP : actuariat vie, mortalité TH/TF0002, taux technique réglementaire

RÈGLES ABSOLUES :
0. FORMAT : §N — TITRE pour les sections, ### pour les sous-titres, \
**gras** pour les termes importants, - tirets pour les listes. \
PAS de tableaux Markdown. PAS de blockquotes >. Sépare les sections par une ligne vide.
1. LANGUE : Français professionnel. Anglais uniquement pour les termes consacrés.
2. RIGUEUR : Chaque affirmation est justifiée par les données fournies. \
   Citer les articles réglementaires exacts.
3. CHIFFRES : En euros avec séparateurs (ex : 2 526 597 €). Pourcentages avec une décimale.
4. RÉFÉRENCES : Art. 145 S2 (SLT), Art. 252 S2 (MCR), BCAC 2019, \
   CTIP 2023, IFRS 17 §B96, TD 88-90, TH0002.
5. ALERTES : Ne jamais minimiser. Implication réelle pour le ratio SCR et la solvabilité.
6. CAUSALITÉ : q_AI CSP ouvrier=1.35×BCAC → incidence ITT élevée → hausse BE → \
   SCR morbidité +35% amplifié → ratio SCR dégradé → révision tarifaire P1 recommandée.
7. INCERTITUDE : Toujours quantifier via CV triangle, A/E ratio BCAC, \
   fourchettes CTIP, p99.5 bootstrap.
8. POSTURE : Assertif mais prudent. Recommandations actionnables avec chiffrage de l'impact.
9. INTERDIT : Phrases génériques sans données. Ne jamais conclure favorablement \
   si ratio SCR < 130% ou maintien ITT > 50% à 6 mois.
10. SPÉCIFIQUE PRÉVOYANCE : Distinguer toujours ITT (court terme, triangle) \
    vs IP (long terme, PM rentes, actuariat vie). \
    Ne pas confondre q_AI (incidence annuelle) et q_IP|ITT (consolidation conditionnelle).

STRUCTURE OBLIGATOIRE EN 8 SECTIONS :
§1 — CONTEXTE ET PORTEFEUILLE PRÉVOYANCE
§2 — TARIFICATION ET COHÉRENCE DES HYPOTHÈSES DE RISQUE
§3 — TABLES DE MORBIDITÉ ET CHAÎNE DE MARKOV
§4 — PROVISIONNEMENT ITT ET RENTES IP
§5 — REPORTING RÉGLEMENTAIRE S2 (SCR/MCR/QRT S.14.01)
§6 — GESTION ACTIF-PASSIF (ALM)
§7 — RISQUES IDENTIFIÉS ET POINTS DE VIGILANCE
§8 — CONCLUSION ET RECOMMANDATIONS POUR LE CONSEIL D'ADMINISTRATION\
"""


# =============================================================================
#  CONSTRUCTION DU CONTEXTE CLAUDE
# =============================================================================

def _construire_contexte(p1, p2, p3, p4, alm, reg1, arrete):
    pa   = float(p1.get('primes_acquises',    0) or 0)
    age  = p1.get('age',             '—')
    csp  = p1.get('categorie',       '—')
    fran = p1.get('franchise_jours', 90)
    be   = float(p3.get('be_prevoyance',  0) or 0)
    be_i = float(p3.get('be_itt',         0) or 0)
    pm_r = float(p3.get('pm_rentes_ip',   0) or 0)
    psap = float(p3.get('psap_total',     0) or 0)
    lr_o = float(p3.get('loss_ratio',     0) or 0)
    scr  = float(p4.get('scr_invalidite', 0) or 0)
    scr_m= float(p4.get('scr_morbidite',  0) or 0)
    mcr  = float(p4.get('mcr',            0) or 0)
    fpp  = float(p4.get('fonds_propres',  0) or 0)
    r_scr= float(p4.get('ratio_scr_pct',  0) or 0)
    r_mcr= float(p4.get('ratio_mcr_pct',  0) or 0)
    trans = p2.get('transitions', {})
    dur   = p2.get('esperances',   {})
    dur_a = float(alm.get('duration_actif',  0) or 0)
    dur_p = float(alm.get('duration_passif', 0) or 0)
    lcr   = float(alm.get('lcr',             0) or 0)
    cv    = float(p3.get('cv_inter',         0) or 0)
    lines = [
        f'DOSSIER PRÉVOYANCE — Arrêté {arrete}',
        f'Assuré : {age} ans | CSP : {csp} | Franchise : {fran} jours',
        '',
        '=== TARIFICATION (P1 Axel — BCAC 2019) ===',
        f'Primes acquises : {_f(pa)}',
        f'Taux cotisation : {_pct(p1.get("taux_cotisation_pct", 0))}',
        f'Taux ITT BCAC : {_pct(float(p1.get("taux_itt", 0) or 0)*100, 3)} /an',
        f'Taux IP annuel : {_pct(float(p1.get("taux_ip", 0) or 0)*100, 4)} /an',
        '',
        '=== TABLES MORBIDITÉ (P2 Rayan — BCAC 2019) ===',
        f'q_AI (Actif→ITT) : {_pct(float(trans.get("q_AI", 0) or 0)*100, 4)} /an',
        f'q_IA (ITT→Actif guérison) : {_pct(float(trans.get("q_IA", 0) or 0)*100)} /an',
        f'q_IP|ITT (consolidation) : {_pct(float(trans.get("q_IP_cond", 0) or 0)*100)}',
        f'Durée moyenne ITT (>franchise) : {dur.get("duree_moy_itt_mois","—")} mois',
        f'Espérance durée IP : {dur.get("esperance_duree_ip_ans","—")} ans',
        '',
        '=== PROVISIONNEMENT (P3 Élodie — CL/Mack/BF/Bootstrap) ===',
        f'BE ITT (triangle) : {_f(be_i)} | {_pct(be_i/pa*100 if pa else 0)} des primes',
        f'PM Rentes IP (actuariat vie) : {_f(pm_r)}',
        f'PSAP ITT : {_f(psap)}',
        f'BE Prévoyance total : {_f(be)} | {_pct(be/pa*100 if pa else 0)} des primes | Norme ≤ 50%',
        f'Loss Ratio CTIP : {_pct(lr_o*100)} | Norme ≤ 85%',
        f'CV inter-méthodes : {_pct(cv)} | Norme < 20%',
        '',
        '=== RÉGLEMENTATION S2 (P4 Valentin — Art. 145 RD 2015/35) ===',
        f'SCR Morbidité (+35% incidence) : {_f(scr_m)} — choc dominant',
        f'SCR Invalidité agrégé (EIOPA ρ=0.25) : {_f(scr)}',
        f'MCR Prévoyance : {_f(mcr)} (plancher 3,7 M€)',
        f'Fonds Propres : {_f(fpp)}',
        f'Ratio SCR : {_pct(r_scr)} | Cible ≥ 130%',
        f'Ratio MCR : {_pct(r_mcr)} | Requis ≥ 100%',
        f'Ratio SCR/BE : {_pct(scr/be*100 if be else 0)} | Plage normale [20%-150%]',
        '',
        '=== ALM ===',
        f'Duration actif : {dur_a:.2f} ans | Duration passif : {dur_p:.2f} ans',
        f'Gap duration : {dur_a-dur_p:+.2f} ans | Alerte si > 3 ans',
        f'LCR : {_pct(lcr*100)} | Requis ≥ 100%',
        '',
        '=== HYPOTHÈSES P1 ===',
    ]
    for h in p1.get('hypotheses', []):
        lines.append(f'  {h.get("id","?")} [{h.get("statut","?")}] : {h.get("valeur","")}')
    lines += ['', '=== HYPOTHÈSES P2 ===']
    for h in p2.get('hypotheses', []):
        lines.append(f'  {h.get("id","?")} [{h.get("statut","?")}] : {h.get("valeur","")}')
    lines += ['', '=== HYPOTHÈSES P3 ===']
    for h in p3.get('hypotheses', []):
        lines.append(f'  {h.get("id","?")} [{h.get("statut","?")}] : {h.get("valeur","")}')
    lines += ['', '=== HYPOTHÈSES P4 ===']
    for h in p4.get('hypotheses', []):
        lines.append(f'  {h.get("id","?")} [{h.get("statut","?")}] : {h.get("valeur","")}')
    lines.append('')
    lines.append('Rédige le commentaire actuariel complet en 8 sections.')
    return '\n'.join(lines)


# =============================================================================
#  NARRATION CLAUDE API
# =============================================================================

def _narration_claude_api(contexte):
    try:
        import anthropic
        api_key = os.environ.get('ANTHROPIC_API_KEY')
        if not api_key:
            try:
                import streamlit as st
                api_key = st.secrets.get('ANTHROPIC_API_KEY')
            except Exception:
                pass
        if not api_key:
            raise ValueError('ANTHROPIC_API_KEY non définie')
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model='claude-sonnet-4-6',
            max_tokens=10000,
            system=SYSTEM_PROMPT_PREVOYANCE,
            messages=[{'role': 'user', 'content': contexte}],
        )
        return resp.content[0].text
    except Exception as e:
        logger.warning(f'Claude API Prévoyance indisponible : {e}')
        raise

def _generer_narration(p1, p2, p3, p4, alm, reg1, arrete) -> Tuple[str, str]:
    try:
        ctx = _construire_contexte(p1, p2, p3, p4, alm, reg1, arrete)
        txt = _narration_claude_api(ctx)
        if txt:
            return txt, 'claude_api'
    except Exception:
        pass
    for src in [p4, p3, p1]:
        com = _clean(src.get('commentaire', ''))
        if com:
            return com, 'templates'
    return '', 'aucune'


# =============================================================================
#  CSS ADDITIONNEL (hyp-err absent de n5_rapport)
# =============================================================================

_CSS_ADDON = """
<style>
.hyp-err .hyp-label { background: rgba(192,57,43,0.10); color: var(--rouge); border-right: 3px solid var(--rouge); }
</style>
"""


# =============================================================================
#  EXPORT HTML
# =============================================================================

def export_html(result_p1, result_p2, result_p3, result_p4,
                result_alm=None, result_reg1=None,
                ref_client='', arrete='', audit_id='') -> str:
    try:
        p1 = result_p1 or {}; p2 = result_p2 or {}
        p3 = result_p3 or {}; p4 = result_p4 or {}
        alm = result_alm or {}; reg1 = result_reg1 or {}

        dt  = datetime.now().strftime('%d/%m/%Y')
        arr = arrete or dt
        cli = ref_client or 'À renseigner'

        rags = [p1.get('statut_rag','VERT'), p2.get('statut_rag','VERT'),
                p3.get('statut_rag','VERT'), p4.get('statut_rag','VERT')]
        rag  = 'ROUGE' if 'ROUGE' in rags else ('AMBRE' if 'AMBRE' in rags else 'VERT')
        s_label = _statut_label(rag)
        s_cls   = rag.lower()

        pa   = float(p1.get('primes_acquises',   0) or 0)
        be   = float(p3.get('be_prevoyance',     0) or 0)
        be_i = float(p3.get('be_itt',            0) or 0)
        pm_r = float(p3.get('pm_rentes_ip',      0) or 0)
        psap = float(p3.get('psap_total',        0) or 0)
        prec = float(p3.get('prec',              0) or 0)
        ra   = float(p3.get('risk_adjustment',   0) or 0)
        tp   = float(p3.get('tp_prevoyance',     0) or 0)
        lr_o = float(p3.get('loss_ratio',        0) or 0)
        scr  = float(p4.get('scr_invalidite',    0) or 0)
        scr_m= float(p4.get('scr_morbidite',     0) or 0)
        scr_c= float(p4.get('scr_cessation',     0) or 0)
        scr_l= float(p4.get('scr_longevite',     0) or 0)
        mcr  = float(p4.get('mcr',               0) or 0)
        fpp  = float(p4.get('fonds_propres',     0) or 0)
        r_scr= float(p4.get('ratio_scr_pct',     0) or 0)
        r_mcr= float(p4.get('ratio_mcr_pct',     0) or 0)
        dur_a= float(alm.get('duration_actif',   0) or 0)
        dur_p= float(alm.get('duration_passif',  0) or 0)
        lcr  = float(alm.get('lcr',              0) or 0)

        narration, source = _generer_narration(p1, p2, p3, p4, alm, reg1, arr)
        narration_html = _md_to_html(narration)
        source_badge = {
            'claude_api': '✦ ActuarIA Intelligence',
            'templates':  '📝 Mode standard',
            'aucune':     '',
        }.get(source, '')

        logo_src = LOGO_URI if _N5_IMPORT else ''

        # KPIs page de garde
        kpis_html = (
            '<div class="garde-kpi"><div class="kpi-label">Client</div>'
            f'<div class="kpi-value">{cli}</div></div>'
            '<div class="garde-kpi"><div class="kpi-label">Arrêté</div>'
            f'<div class="kpi-value">{arr}</div></div>'
            '<div class="garde-kpi"><div class="kpi-label">Primes acquises</div>'
            f'<div class="kpi-value highlight">{_f(pa)}</div></div>'
            '<div class="garde-kpi"><div class="kpi-label">BE Prévoyance</div>'
            f'<div class="kpi-value">{_f(be)}</div>'
            f'<div class="kpi-sub">LR CTIP : {_pct(lr_o*100)}</div></div>'
            '<div class="garde-kpi"><div class="kpi-label">SCR | Ratio SCR</div>'
            f'<div class="kpi-value">{_f(scr)}</div>'
            f'<div class="kpi-sub">{_pct(r_scr)} (cible ≥130%)</div></div>'
        )

        # Hypothèses P1+P2+P3+P4
        all_hyp = (
            [(h,'P1') for h in p1.get('hypotheses',[])] +
            [(h,'P2') for h in p2.get('hypotheses',[])] +
            [(h,'P3') for h in p3.get('hypotheses',[])] +
            [(h,'P4') for h in p4.get('hypotheses',[])]
        )
        hyp_cards = ''
        for h, agent in all_hyp:
            statut = h.get('statut', '')
            cls  = 'hyp-ok'   if statut == 'VALIDÉE'     else (
                   'hyp-err'  if statut == 'NON VALIDÉE' else 'hyp-warn')
            lbl  = {'VALIDÉE': '✓ VALIDÉE', 'NON VALIDÉE': '✗ NON VALIDÉE'}.get(
                   statut, '⚠ ' + statut)
            hyp_cards += (
                f'<div class="hyp-card {cls}">'
                f'<div class="hyp-label">'
                f'<div class="hyp-code">[{agent}] {h.get("id","?")} — {h.get("hypothese","")[:40]}</div>'
                f'<div class="hyp-code" style="margin-top:4px;">{lbl}</div></div>'
                f'<div class="hyp-text">{_s(h.get("valeur",""))}</div></div>'
            )

        # Tableau transitions P2
        trans = p2.get('transitions', {})
        tbl_trans = (
            '<table class="premium"><thead><tr>'
            '<th>Transition</th><th class="right">Probabilité</th>'
            '<th>Source</th></tr></thead><tbody>'
            f'<tr><td class="label">q_AI : Actif → ITT (incidence)</td>'
            f'<td class="right">{_pct(float(trans.get("q_AI",0) or 0)*100, 4)} /an</td>'
            f'<td>BCAC 2019</td></tr>'
            f'<tr><td class="label">q_IA : ITT → Actif (guérison)</td>'
            f'<td class="right">{_pct(float(trans.get("q_IA",0) or 0)*100)} /an</td>'
            f'<td>BCAC 2019</td></tr>'
            f'<tr><td class="label">q_IP|ITT : Consolidation (invalidité)</td>'
            f'<td class="right">{_pct(float(trans.get("q_IP_cond",0) or 0)*100)}</td>'
            f'<td>BCAC 2019</td></tr>'
            '<tr class="highlight-gold">'
            f'<td class="label">q_IP annuel : Actif → IP (composé)</td>'
            f'<td class="right">{_pct(float(trans.get("q_IP_annuel",0) or 0)*100, 4)}</td>'
            f'<td>q_AI × q_IP|ITT</td></tr>'
            '</tbody></table>'
        )

        # Tableau provisions
        tbl_prov = (
            '<table class="premium"><thead><tr>'
            '<th>Composante</th><th class="right">Montant</th>'
            '<th class="right">% Primes</th><th>Référence</th></tr></thead><tbody>'
            f'<tr><td class="label">BE ITT (triangle CL/Mack/BF)</td>'
            f'<td class="right"><span class="mono">{_f(be_i)}</span></td>'
            f'<td class="right">{_pct(be_i/pa*100 if pa else 0)}</td>'
            f'<td>Horizon liquidation 60-84 mois</td></tr>'
            f'<tr><td class="label">PM Rentes IP (actuariat vie)</td>'
            f'<td class="right"><span class="mono">{_f(pm_r)}</span></td>'
            f'<td class="right">{_pct(pm_r/pa*100 if pa else 0)}</td>'
            f'<td>TH0002/TF0002 — TD 88-90</td></tr>'
            f'<tr><td class="label">PSAP ITT (dossiers en cours)</td>'
            f'<td class="right"><span class="mono">{_f(psap)}</span></td>'
            f'<td class="right">{_pct(psap/pa*100 if pa else 0)}</td>'
            f'<td>Sinistres déclarés en instance</td></tr>'
            f'<tr><td class="label">PREC (provision risques en cours)</td>'
            f'<td class="right"><span class="mono">{_f(prec)}</span></td>'
            f'<td class="right">{_pct(prec/pa*100 if pa else 0)}</td>'
            f'<td>Art. 78 S2</td></tr>'
            f'<tr><td class="label">Best Estimate Prévoyance</td>'
            f'<td class="right"><span class="mono">{_f(be)}</span></td>'
            f'<td class="right">{_pct(be/pa*100 if pa else 0)}</td>'
            f'<td>Art. 77 §1 S2 — norme ≤ 50% PA</td></tr>'
            f'<tr><td class="label">Risk Adjustment (CoC 6% — floor 3%)</td>'
            f'<td class="right"><span class="mono">{_f(ra)}</span></td>'
            f'<td class="right">{_pct(ra/pa*100 if pa else 0)}</td>'
            f'<td>IFRS 17 §B119</td></tr>'
            '<tr class="highlight-gold">'
            f'<td class="label">TP Prévoyance (BE + RA)</td>'
            f'<td class="right"><span class="mono">{_f(tp)}</span></td>'
            f'<td class="right">{_pct(tp/pa*100 if pa else 0)}</td>'
            f'<td>Art. 77 §1 S2 — bilan prudentiel</td></tr>'
            '</tbody></table>'
        )

        # Tableau SCR/MCR
        scr_stat = '✅ CONFORME' if r_scr>=130 else ('⚠️ VIGILANCE' if r_scr>=100 else '❌ INSUFFISANT')
        scr_be   = scr/be*100 if be else 0
        tbl_scr = (
            '<table class="premium"><thead><tr>'
            '<th>Composante</th><th class="center">Montant</th>'
            '<th class="center">Ratio/Info</th><th>Référence</th></tr></thead><tbody>'
            f'<tr><td class="label">SCR Morbidité (+35% incidence)</td>'
            f'<td class="center"><span class="mono">{_f(scr_m)}</span></td>'
            f'<td class="center">Choc dominant</td><td>Art. 145 §2(a) RD 2015/35</td></tr>'
            f'<tr><td class="label">SCR Cessation (-20% guérison)</td>'
            f'<td class="center"><span class="mono">{_f(scr_c)}</span></td>'
            f'<td class="center">ρ=0.25</td><td>Art. 145 §2(b) RD 2015/35</td></tr>'
            f'<tr><td class="label">SCR Longévité (-20% mortalité IP)</td>'
            f'<td class="center"><span class="mono">{_f(scr_l)}</span></td>'
            f'<td class="center">ρ=0.25</td><td>Art. 145 §2(c) RD 2015/35</td></tr>'
            f'<tr><td class="label">SCR Invalidité (agrégé EIOPA)</td>'
            f'<td class="center"><span class="mono">{_f(scr)}</span></td>'
            f'<td class="center">{_pct(scr_be)} BE</td><td>EIOPA Annexe IV</td></tr>'
            f'<tr><td class="label">MCR Prévoyance</td>'
            f'<td class="center"><span class="mono">{_f(mcr)}</span></td>'
            f'<td class="center">Plancher 3,7 M€</td><td>Art. 252 RD 2015/35</td></tr>'
            f'<tr><td class="label">Fonds Propres</td>'
            f'<td class="center"><span class="mono">{_f(fpp)}</span></td>'
            f'<td class="center">—</td><td>Art. 87 S2</td></tr>'
            '<tr class="highlight-gold">'
            f'<td class="label">Ratio SCR</td>'
            f'<td class="center" style="font-weight:700;">{_pct(r_scr)}</td>'
            f'<td class="center">{scr_stat}</td><td>Cible pratique ≥ 130%</td></tr>'
            '<tr class="highlight-gold">'
            f'<td class="label">Ratio MCR</td>'
            f'<td class="center" style="font-weight:700;">{_pct(r_mcr)}</td>'
            f'<td class="center">{("✅" if r_mcr>=100 else "❌")}</td>'
            f'<td>Art. 139 S2 — seuil déclencheur</td></tr>'
            '</tbody></table>'
        )

        # Tableau ALM
        gap_dur = dur_a - dur_p
        gap_ok  = abs(gap_dur) <= 2
        lcr_ok  = lcr >= 1.0
        tbl_alm = (
            '<table class="premium"><thead><tr>'
            '<th>Indicateur ALM</th><th class="center">Valeur</th>'
            '<th class="center">Norme</th><th class="center">Statut</th>'
            '</tr></thead><tbody>'
            f'<tr><td class="label">Duration actif</td>'
            f'<td class="center">{dur_a:.2f} ans</td><td class="center">—</td>'
            f'<td class="center">—</td></tr>'
            f'<tr><td class="label">Duration passif</td>'
            f'<td class="center">{dur_p:.2f} ans</td><td class="center">—</td>'
            f'<td class="center">—</td></tr>'
            f'<tr><td class="label">Gap de duration</td>'
            f'<td class="center">{gap_dur:+.2f} ans</td><td class="center">≤ 2 ans</td>'
            f'<td class="center">{("✅" if gap_ok else "⚠️")}</td></tr>'
            f'<tr><td class="label">LCR (liquidité)</td>'
            f'<td class="center">{_pct(lcr*100)}</td><td class="center">≥ 100%</td>'
            f'<td class="center">{("✅" if lcr_ok else "❌")}</td></tr>'
            f'<tr><td class="label">Immunisation Redington</td>'
            f'<td class="center">{alm.get("redington_ok","—")}</td>'
            f'<td class="center">Convexité ≥ 0</td><td class="center">—</td></tr>'
            '<tr class="highlight-gold">'
            f'<td class="label">BV01 stress +100bp</td>'
            f'<td class="center">{_f(alm.get("bv01_stress_100",0))}</td>'
            f'<td class="center">Impact taux</td><td class="center">—</td></tr>'
            '</tbody></table>'
        )

        # Alertes section 09
        alertes = ''
        for src_dict in [p1, p2, p3, p4]:
            for h in src_dict.get('hypotheses', []):
                if h.get('statut') in ('NON VALIDÉE', 'À JUSTIFIER'):
                    txt = _clean(h.get('valeur', ''))
                    if txt:
                        cls = 'hyp-err' if h.get('statut') == 'NON VALIDÉE' else 'hyp-warn'
                        alertes += (
                            f'<div class="hyp-card {cls}" style="margin-bottom:10px;">'
                            f'<div class="hyp-label"><div class="hyp-code">Point de vigilance</div></div>'
                            f'<div class="hyp-text">{txt}</div></div>'
                        )
        if not alertes:
            alertes = '<p style="color:#1E8449;font-size:9pt;">✅ Aucun point de vigilance majeur identifié.</p>'

        css_str = _css() if callable(_css) else ''

        # Assemblage HTML
        html = (
            '<!DOCTYPE html>\n<html lang="fr">\n<head>\n'
            '<meta charset="UTF-8">\n'
            f'<title>Rapport Prévoyance — {cli} — {arr}</title>\n'
            '<script src="https://cdn.jsdelivr.net/npm/plotly.js-dist@2.26.0/plotly.min.js"></script>\n'
            + css_str + _CSS_ADDON +
            '\n</head>\n<body>\n<div class="rapport-container">\n'

            # PAGE DE GARDE
            '<div class="page-garde">\n'
            '<div class="garde-bg">'
            '<div class="garde-dots"></div>'
            '<div class="garde-diagonal"></div>'
            '<div class="garde-accent-line"></div></div>\n'
            '<div class="garde-inner">\n'
            '<div class="garde-header">\n'
            f'<div class="garde-logo-wrap"><img src="{logo_src}" alt="ActuarIA"/></div>\n'
            '<div class="garde-badges"><span class="badge-confidentiel">⬛ Confidentiel</span></div>\n'
            '</div>\n'
            '<div class="garde-hero">\n'
            '<div class="garde-eyebrow">Rapport Actuariel Prévoyance</div>\n'
            '<div class="garde-titre">Prévoyance<br><em>Collective</em></div>\n'
            f'<div class="garde-subtitle">Arrêté au {arr}</div>\n'
            '<div class="garde-sep">'
            '<div class="garde-sep-line"></div>'
            '<div class="garde-sep-diamond"></div>'
            '<div class="garde-sep-line"></div></div>\n'
            f'<div class="garde-statut garde-statut-{s_cls}">\n'
            f'  <div class="statut-dot statut-dot-{s_cls}"></div>\n'
            f'  <div class="statut-label-{s_cls}">{s_label}</div>\n'
            '</div>\n'
            '</div>\n'
            f'<div class="garde-footer"><div class="garde-kpis">{kpis_html}</div></div>\n'
            '</div>\n</div>\n\n'

            # CORPS
            '<div class="rapport-body">\n\n'

            # 01 Synthèse
            '<div class="section-header"><span class="section-num">01</span>'
            '<span class="section-titre">Synthèse exécutive</span></div>\n'
            '<div class="section-body">\n'
            '<div class="kpi-grid">\n'
            f'<div class="kpi-card"><div class="kpi-card-label">Primes acquises</div>'
            f'<div class="kpi-card-value">{_f(pa)}</div>'
            f'<div class="kpi-card-sub">{p1.get("nb_assures","—")} assuré(s)</div></div>\n'
            f'<div class="kpi-card"><div class="kpi-card-label">BE Prévoyance</div>'
            f'<div class="kpi-card-value">{_f(be)}</div>'
            f'<div class="kpi-card-sub">LR CTIP : {_pct(lr_o*100)}</div></div>\n'
            f'<div class="kpi-card kpi-card-rouge"><div class="kpi-card-label">SCR Invalidité</div>'
            f'<div class="kpi-card-value kpi-card-value-rouge">{_f(scr)}</div>'
            f'<div class="kpi-card-sub">Art. 145 RD 2015/35</div></div>\n'
            f'<div class="kpi-card"><div class="kpi-card-label">Ratio SCR</div>'
            f'<div class="kpi-card-value">{_pct(r_scr)}</div>'
            f'<div class="kpi-card-sub">Cible ≥ 130%</div></div>\n'
            '</div>\n</div>\n<div class="section-divider"></div>\n\n'

            # 02 Tarification
            '<div class="section-header"><span class="section-num">02</span>'
            '<span class="section-titre">Tarification — P1 Axel (BCAC 2019)</span></div>\n'
            '<div class="section-body">\n'
            '<div class="table-section-title">Paramètres de tarification</div>\n'
            + _build_table_tarifP1(p1) +
            '\n</div>\n<div class="section-divider"></div>\n\n'

            # 03 Tables morbidité
            '<div class="section-header"><span class="section-num">03</span>'
            '<span class="section-titre">Tables de morbidité — P2 Rayan (BCAC 2019)</span></div>\n'
            '<div class="section-body">\n'
            '<div class="table-section-title">Probabilités de transition — Chaîne de Markov</div>\n'
            + tbl_trans +
            '\n</div>\n<div class="section-divider"></div>\n\n'

            # 04 Provisionnement
            '<div class="section-header"><span class="section-num">04</span>'
            '<span class="section-titre">Provisionnement — P3 Élodie</span></div>\n'
            '<div class="section-body">\n'
            + tbl_prov +
            '\n</div>\n<div class="section-divider"></div>\n\n'

            # 05 Reporting S2
            '<div class="section-header"><span class="section-num">05</span>'
            '<span class="section-titre">Reporting S2 — P4 Valentin · QRT S.14.01</span></div>\n'
            '<div class="section-body">\n'
            + tbl_scr +
            '\n</div>\n<div class="section-divider"></div>\n\n'

            # 06 ALM
            '<div class="section-header"><span class="section-num">06</span>'
            '<span class="section-titre">Gestion Actif-Passif (ALM)</span></div>\n'
            '<div class="section-body">\n'
            + tbl_alm +
            '\n</div>\n<div class="section-divider"></div>\n\n'

            # 07 Hypothèses
            '<div class="section-header"><span class="section-num">07</span>'
            '<span class="section-titre">Hypothèses actuarielles</span></div>\n'
            '<div class="section-body">\n'
            f'<div class="hyp-grid">{hyp_cards}</div>\n'
            '</div>\n<div class="section-divider"></div>\n\n'

            # 08 Commentaire Claude
            '<div class="section-header"><span class="section-num">08</span>'
            '<span class="section-titre">Commentaire actuariel</span></div>\n'
            '<div class="section-body">\n'
            '<div class="commentaire-wrap">\n'
            '<div class="commentaire-header">\n'
            f'  <span class="commentaire-header-title">Rapport Prévoyance · Arrêté {arr}</span>\n'
            f'  <span class="commentaire-ai-badge">{source_badge}</span>\n'
            '</div>\n'
            '<div class="commentaire-body">\n'
            '<p style="font-size:8pt;color:var(--slate);margin-bottom:20px;font-style:italic;">'
            "Commentaire à destination du Conseil d'Administration et de l'Actuaire Désigné.<br>"
            "Document soumis à l'ACPR dans le cadre du reporting Solvabilité 2.</p>\n"
            + narration_html +
            f'\n<div class="comm-footer">✦ Narration générée par ActuarIA Intelligence · Agent SP-Prévoyance v2.0 · {dt}</div>\n'
            '</div>\n</div>\n'
            '</div>\n<div class="section-divider"></div>\n\n'

            # 09 Jugement
            '<div class="section-header"><span class="section-num">09</span>'
            '<span class="section-titre">Jugement actuariel &amp; Recommandations</span></div>\n'
            '<div class="section-body">\n'
            f'<div class="hyp-grid">{alertes}</div>\n'
            '</div>\n\n'

            '</div>\n\n'

            # PIED DE PAGE
            '<div class="pied-de-page">\n'
            f'  <div class="pied-logo"><img src="{logo_src}" alt="ActuarIA"/></div>\n'
            '  <div class="pied-meta">'
            + f'{cli} · Prévoyance · Arrêté {arr} · {dt}<br>'
            '<span style="font-size:6.5pt;color:#8A9AB0;">'
            "Rapport établi conformément aux Art. 77, 105 et 145 de la Directive Solvabilité II "
            "et au Guide Institut des Actuaires 2023</span><br>"
            '<span class="confidentiel-footer">CONFIDENTIEL — USAGE STRICTEMENT ACTUARIEL</span>\n'
            '  </div>\n'
            '</div>\n\n'
            '</div>\n</body>\n</html>'
        )

        logger.info(f'HTML Prévoyance v2 : {len(html):,} chars — narration={source}')
        return html

    except Exception as e:
        logger.error(f'export_html Prévoyance : {e}', exc_info=True)
        return f'<html><body><h1>Erreur : {e}</h1></body></html>'


def _build_table_tarifP1(p1) -> str:
    kpis = [
        ('Prime commerciale',      _f(p1.get('prime_commerciale'))),
        ('Taux de cotisation',     _pct(p1.get('taux_cotisation_pct'))),
        ('Salaire brut',           _f(p1.get('salaire_brut'))),
        ('Taux rente IPP',         _pct(float(p1.get('taux_rente_ipp', 0) or 0)*100)),
        ('Franchise ITT',          f"{p1.get('franchise_jours', 90)} jours"),
        ('Durée contrat',          f"{p1.get('duree_contrat', 20)} ans"),
        ('Taux ITT BCAC',          _pct(float(p1.get('taux_itt', 0) or 0)*100, 3)),
        ('Taux IP annuel BCAC',    _pct(float(p1.get('taux_ip',  0) or 0)*100, 4)),
    ]
    tbl = (
        '<table class="premium"><thead><tr>'
        '<th>Indicateur</th><th class="right">Valeur</th>'
        '</tr></thead><tbody>'
    )
    for i, (k, v) in enumerate(kpis):
        row_cls = '' if i % 2 == 0 else ''
        tbl += f'<tr><td class="label">{k}</td><td class="right"><span class="mono">{v}</span></td></tr>'
    tbl += '</tbody></table>'
    return tbl


# =============================================================================
#  EXPORT PDF
# =============================================================================

def export_pdf(result_p1, result_p2, result_p3, result_p4,
               result_alm=None, result_reg1=None,
               ref_client='', arrete='', audit_id='') -> bytes:
    try:
        from weasyprint import HTML as WH
        html = export_html(result_p1, result_p2, result_p3, result_p4,
                           result_alm, result_reg1, ref_client, arrete, audit_id)
        pdf = WH(string=html).write_pdf()
        logger.info(f'PDF Prévoyance : {len(pdf):,} bytes')
        return pdf
    except ImportError:
        logger.error('weasyprint non installé')
        return b''
    except Exception as e:
        logger.error(f'export_pdf Prévoyance : {e}', exc_info=True)
        return b''


# =============================================================================
#  EXPORT WORD
# =============================================================================

def export_word(result_p1, result_p2, result_p3, result_p4,
                result_alm=None, result_reg1=None,
                ref_client='', arrete='', audit_id='') -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        logger.error(f'python-docx absent : {e}')
        return b''

    try:
        p1 = result_p1 or {}; p2 = result_p2 or {}
        p3 = result_p3 or {}; p4 = result_p4 or {}
        alm = result_alm or {}

        def rgb(h):
            h = h.lstrip('#')
            return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))

        NR=rgb('#0B1E3D'); GR=rgb('#C9A84C'); BR=rgb('#FFFFFF')
        GrR=rgb('#8A9BB0'); RgR=rgb('#C0392B'); VR=rgb('#1E8449'); AR=rgb('#E67E22')

        dt  = datetime.now().strftime('%d/%m/%Y')
        arr = arrete or dt
        cli = ref_client or 'À renseigner'

        rags = [p1.get('statut_rag','VERT'), p2.get('statut_rag','VERT'),
                p3.get('statut_rag','VERT'), p4.get('statut_rag','VERT')]
        rag  = 'ROUGE' if 'ROUGE' in rags else ('AMBRE' if 'AMBRE' in rags else 'VERT')

        narration, source = _generer_narration(p1, p2, p3, p4, alm, result_reg1 or {}, arr)

        pa   = float(p1.get('primes_acquises',   0) or 0)
        be   = float(p3.get('be_prevoyance',     0) or 0)
        be_i = float(p3.get('be_itt',            0) or 0)
        pm_r = float(p3.get('pm_rentes_ip',      0) or 0)
        psap = float(p3.get('psap_total',        0) or 0)
        prec = float(p3.get('prec',              0) or 0)
        ra   = float(p3.get('risk_adjustment',   0) or 0)
        tp   = float(p3.get('tp_prevoyance',     0) or 0)
        lr_o = float(p3.get('loss_ratio',        0) or 0)
        scr  = float(p4.get('scr_invalidite',    0) or 0)
        mcr  = float(p4.get('mcr',               0) or 0)
        fpp  = float(p4.get('fonds_propres',     0) or 0)
        r_scr= float(p4.get('ratio_scr_pct',     0) or 0)
        r_mcr= float(p4.get('ratio_mcr_pct',     0) or 0)
        dur_a= float(alm.get('duration_actif',   0) or 0)
        dur_p= float(alm.get('duration_passif',  0) or 0)
        lcr  = float(alm.get('lcr',              0) or 0)

        doc = Document()
        for s in doc.sections:
            s.top_margin=Cm(2); s.bottom_margin=Cm(2)
            s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

        def _bg(cell, hex6):
            tc=cell._tc; tcp=tc.get_or_add_tcPr()
            sd=OxmlElement('w:shd')
            sd.set(qn('w:fill'), hex6.lstrip('#'))
            sd.set(qn('w:color'), 'auto'); sd.set(qn('w:val'), 'clear')
            tcp.append(sd)

        def _run(p, txt, bold=False, italic=False, sz=10, col=None):
            r=p.add_run(str(txt)); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
            if col: r.font.color.rgb=col
            return r

        def _h(txt, lv=1, col=None):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
            _run(p, txt, bold=True, sz={1:16,2:12,3:10}.get(lv,10),
                 col=col or (NR if lv==1 else GR))

        def _sep():
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
            bo=OxmlElement('w:bottom')
            bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'6')
            bo.set(qn('w:space'),'1'); bo.set(qn('w:color'),'C9A84C')
            pBdr.append(bo); pPr.append(pBdr)

        def _tbl(heads, rows, ws=None):
            t=doc.add_table(rows=1+len(rows), cols=len(heads)); t.style='Table Grid'
            for i,hd in enumerate(heads):
                c=t.rows[0].cells[i]; _bg(c,'0B1E3D')
                pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=pp.add_run(str(hd)); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=BR
            for ri,row in enumerate(rows):
                for ci,v in enumerate(row):
                    c=t.rows[ri+1].cells[ci]
                    if ri%2==1: _bg(c,'EEF2F7')
                    pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r=pp.add_run(str(v) if v is not None else '—'); r.font.size=Pt(9)
            if ws:
                for i,w in enumerate(ws):
                    for row in t.rows: row.cells[i].width=Cm(w)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)

        # Page de garde
        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p,'RAPPORT ACTUARIEL PRÉVOYANCE\n',bold=True,sz=20,col=NR)
        _run(p,'Prévoyance Collective',bold=True,sz=14,col=GR)
        doc.add_paragraph()
        s_col_r = VR if rag=='VERT' else AR if rag=='AMBRE' else RgR
        p=doc.add_paragraph()
        _run(p,'Statut : ',sz=11,col=NR); _run(p,rag,bold=True,sz=11,col=s_col_r)
        doc.add_paragraph()
        _tbl(['Client','Arrêté','Primes acquises','BE Prévoyance'],
             [[cli, arr, _f(pa), _f(be)]],ws=[3.5,2.5,4.0,4.0])
        doc.add_page_break()

        _h('1. Synthèse exécutive'); _sep()
        _tbl(['Indicateur','Valeur','Indicateur','Valeur'],
             [['Primes acquises',_f(pa),'BE Prévoyance',_f(be)],
              ['BE ITT (triangle)',_f(be_i),'PM Rentes IP',_f(pm_r)],
              ['SCR Invalidité',_f(scr),'Ratio SCR',_pct(r_scr)],
              ['MCR Prévoyance',_f(mcr),'Ratio MCR',_pct(r_mcr)],
              ['LR CTIP',_pct(lr_o*100),'LCR',_pct(lcr*100)]],ws=[4.0,3.5,4.0,3.5])
        doc.add_page_break()

        _h('2. Tarification — P1 Axel (BCAC 2019)'); _sep()
        _tbl(['Indicateur','Valeur','Indicateur','Valeur'],
             [['Prime commerciale',_f(p1.get('prime_commerciale')),
               'Taux cotisation',_pct(p1.get('taux_cotisation_pct'))],
              ['Salaire brut',_f(p1.get('salaire_brut')),
               'Franchise ITT',f"{p1.get('franchise_jours',90)} jours"],
              ['Taux ITT BCAC',_pct(float(p1.get('taux_itt',0) or 0)*100, 3),
               'Taux IP annuel',_pct(float(p1.get('taux_ip',0) or 0)*100, 4)]],
             ws=[4.0,3.5,4.0,3.5])
        doc.add_page_break()

        _h('3. Tables de morbidité — P2 Rayan (BCAC 2019)'); _sep()
        trans = p2.get('transitions', {})
        _tbl(['Transition','Probabilité','Source'],
             [['q_AI : Actif → ITT',_pct(float(trans.get('q_AI',0) or 0)*100,4),'BCAC 2019'],
              ['q_IA : ITT → Actif',_pct(float(trans.get('q_IA',0) or 0)*100),'BCAC 2019'],
              ['q_IP|ITT : Consolidation',_pct(float(trans.get('q_IP_cond',0) or 0)*100),'BCAC 2019'],
              ['q_IP : Actif → IP (annuel)',_pct(float(trans.get('q_IP_annuel',0) or 0)*100,4),'q_AI × q_IP|ITT']],
             ws=[6.0,3.5,4.5])
        doc.add_page_break()

        _h('4. Provisionnement — P3 Élodie'); _sep()
        _tbl(['Composante','Montant','% Primes','Référence'],
             [['BE ITT (triangle)',_f(be_i),_pct(be_i/pa*100 if pa else 0),'CL/Mack/BF/Bootstrap'],
              ['PM Rentes IP',_f(pm_r),_pct(pm_r/pa*100 if pa else 0),'TH0002/TD88-90'],
              ['PSAP ITT',_f(psap),_pct(psap/pa*100 if pa else 0),'Sinistres connus'],
              ['PREC',_f(prec),_pct(prec/pa*100 if pa else 0),'Art. 78 S2'],
              ['Best Estimate',_f(be),_pct(be/pa*100 if pa else 0),'Art. 77 §1 S2'],
              ['Risk Adjustment',_f(ra),_pct(ra/pa*100 if pa else 0),'IFRS 17 §B119'],
              ['TP Prévoyance',_f(tp),_pct(tp/pa*100 if pa else 0),'Art. 77 §1 S2']],
             ws=[4.5,3.0,2.5,4.0])
        doc.add_page_break()

        _h('5. Reporting S2 — P4 Valentin · QRT S.14.01'); _sep()
        _tbl(['Composante','Valeur','Info','Référence'],
             [['SCR Morbidité +35%',_f(p4.get('scr_morbidite',0)),'Choc dominant','Art. 145 §2(a)'],
              ['SCR Cessation -20%',_f(p4.get('scr_cessation',0)),'ρ=0.25','Art. 145 §2(b)'],
              ['SCR Longévité -20%',_f(p4.get('scr_longevite',0)),'ρ=0.25','Art. 145 §2(c)'],
              ['SCR Invalidité agrégé',_f(scr),_pct(scr/be*100 if be else 0),'EIOPA Annexe IV'],
              ['MCR Prévoyance',_f(mcr),'Plancher 3,7 M€','Art. 252 RD 2015/35'],
              ['Fonds Propres',_f(fpp),'—','Art. 87 S2'],
              ['Ratio SCR',_pct(r_scr),'✅' if r_scr>=130 else '⚠️','Cible ≥ 130%'],
              ['Ratio MCR',_pct(r_mcr),'✅' if r_mcr>=100 else '❌','Art. 139 S2']],
             ws=[4.5,2.5,2.5,4.5])
        doc.add_page_break()

        _h('6. Gestion Actif-Passif (ALM)'); _sep()
        _tbl(['Indicateur','Valeur','Norme','Statut'],
             [['Duration actif',f'{dur_a:.2f} ans','—','—'],
              ['Duration passif',f'{dur_p:.2f} ans','—','—'],
              ['Gap duration',f'{dur_a-dur_p:+.2f} ans','≤ 2 ans','✅' if abs(dur_a-dur_p)<=2 else '⚠️'],
              ['LCR',_pct(lcr*100),'≥ 100%','✅' if lcr>=1 else '❌'],
              ['Redington',alm.get('redington_ok','—'),'Convexité ≥ 0','—']],
             ws=[4.5,3.0,2.5,4.0])
        doc.add_page_break()

        _h('7. Hypothèses actuarielles'); _sep()
        all_hyp = (
            [(h,'P1') for h in p1.get('hypotheses',[])] +
            [(h,'P2') for h in p2.get('hypotheses',[])] +
            [(h,'P3') for h in p3.get('hypotheses',[])] +
            [(h,'P4') for h in p4.get('hypotheses',[])]
        )
        if all_hyp:
            _tbl(['Src','ID','Hypothèse','Résultat','Statut'],
                 [[ag,h.get('id',''),h.get('hypothese','')[:55],
                   h.get('valeur','')[:75],h.get('statut','')]
                  for h,ag in all_hyp],
                 ws=[1.0,1.5,5.0,5.5,2.5])
        doc.add_page_break()

        _h('8. Commentaire actuariel'); _sep()
        if narration:
            sections_n = re.split(r'(?=§\d+\s*[—\-–])', _clean(narration))
            for sec in sections_n:
                sec = sec.strip()
                if not sec: continue
                ls = sec.split('\n', 1)
                if ls[0]: _h(ls[0], lv=2)
                if len(ls) > 1:
                    for ln in ls[1].split('\n'):
                        ln = ln.strip()
                        if ln:
                            p=doc.add_paragraph()
                            p.paragraph_format.space_after=Pt(3)
                            p.paragraph_format.left_indent=Cm(0.3)
                            _run(p, ln, sz=9, col=NR)
            if source == 'claude_api':
                p=doc.add_paragraph()
                _run(p,'✦ Narration générée par ActuarIA Intelligence',sz=7,italic=True,col=GrR)
        else:
            p=doc.add_paragraph()
            _run(p,'Narration non disponible.',sz=9,italic=True)
        doc.add_page_break()

        _h('9. Jugement actuariel & Recommandations'); _sep()
        for src_dict in [p1, p2, p3, p4]:
            for h in src_dict.get('hypotheses', []):
                if h.get('statut') in ('NON VALIDÉE', 'À JUSTIFIER'):
                    txt = _clean(h.get('valeur',''))
                    if txt:
                        p=doc.add_paragraph()
                        p.paragraph_format.left_indent=Cm(0.4)
                        ic = '❌' if h.get('statut')=='NON VALIDÉE' else '⚠️'
                        _run(p, f'{ic}  ', sz=10, col=RgR if ic=='❌' else AR)
                        _run(p, txt, sz=9, col=NR)

        _sep()
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f'ActuarIA · {cli} · Prévoyance · Arrêté {arr} · {dt} · CONFIDENTIEL',
             sz=7, italic=True, col=GrR)

        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        wb=buf.read()
        logger.info(f'Word Prévoyance : {len(wb):,} bytes')
        return wb

    except Exception as e:
        logger.error(f'export_word Prévoyance : {e}', exc_info=True)
        return b''
