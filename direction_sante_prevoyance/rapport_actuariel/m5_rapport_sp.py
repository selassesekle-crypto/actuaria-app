"""
m5_rapport_sp.py — Rapport actuariel Word + PDF Santé-Prévoyance
Direction Santé-Prévoyance

Produit :
  - Rapport Word (.docx) professionnel — sections 1-8
  - Rapport PDF via weasyprint (fallback HTML si indisponible)

Structure du rapport :
  Page de garde + résumé exécutif
  1. Présentation du portefeuille
  2. Sinistralité santé (LR par poste)
  3. Morbidité prévoyance (ITT/IP/décès)
  4. Provisions techniques (PSAP + PM + BE + RA)
  5. Solvabilité 2 (SCR/MCR consolidé)
  6. IFRS 17 (BE/RA/CSM/LC)
  7. Réglementation (ANI / 100%S / Contrat responsable)
  8. Avis actuariel + hypothèses
"""

from __future__ import annotations
import base64, io, logging, os
from datetime import datetime
from typing import Dict, List, Optional

import numpy as np

logger = logging.getLogger("actuaria.sp.rapport.word")

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY   = "#0F2E52"
OR     = "#C9A84C"
BLANC  = "#F0F4F8"
VERT   = "#2ECC71"
ROUGE  = "#E74C3C"
AMBRE  = "#F39C12"

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from weasyprint import HTML as WP_HTML
    WEASYPRINT_OK = True
except ImportError:
    WEASYPRINT_OK = False


def _rgb(hex_color: str):
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def _para_style(para, bold=False, size=11, color=None, align="left"):
    run = para.runs[0] if para.runs else para.add_run()
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = _rgb(color)
    para.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)

def _cell_bg(cell, hex_color: str):
    """Colorier le fond d'une cellule."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)

def _add_section_title(doc, numero: int, titre: str):
    p = doc.add_heading(f"{numero}. {titre}", level=1)
    p.runs[0].font.color.rgb = _rgb(OR)
    p.runs[0].font.size = Pt(13)

def _add_kpi_table(doc, headers, rows, col_widths=None):
    """Tableau KPI avec entête Navy/Gold."""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # En-tête
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _cell_bg(cell, NAVY.lstrip("#"))
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = _rgb(OR)
        run.font.size = Pt(9)
    # Données
    for r_idx, row in enumerate(rows, 1):
        bg = "1B3A5C" if r_idx % 2 == 0 else "243F6A"
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = str(val) if val is not None else "—"
            _cell_bg(cell, bg)
            # cell.text crée déjà un run — on l'utilise ou on en ajoute un
            para = cell.paragraphs[0]
            run  = para.runs[0] if para.runs else para.add_run(str(val) if val else "")
            run.font.color.rgb = _rgb(BLANC)
            run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def export_word_sp(data: Dict) -> bytes:
    """
    Génère le rapport actuariel SP au format Word (.docx).

    Parameters
    ----------
    data : dict — données consolidées du rapport SP

    Returns
    -------
    bytes — contenu du fichier .docx
    """
    if not DOCX_OK:
        logger.warning("python-docx non disponible — export Word SP désactivé")
        return b""

    doc = Document()

    # ── Page de garde ─────────────────────────────────────────────────────────
    doc.add_paragraph()
    titre = doc.add_paragraph()
    run_t = titre.add_run("RAPPORT ACTUARIEL SANTÉ-PRÉVOYANCE")
    run_t.bold = True; run_t.font.size = Pt(20)
    run_t.font.color.rgb = _rgb(NAVY)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sous_titre = doc.add_paragraph()
    run_s = sous_titre.add_run(data.get("entite", "ActuarIA"))
    run_s.font.size = Pt(14); run_s.font.color.rgb = _rgb(OR)
    sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_p = doc.add_paragraph()
    run_d = date_p.add_run(f"Arrêté au {data.get('date_arrete','')} | Généré le {datetime.now().strftime('%d/%m/%Y')}")
    run_d.font.size = Pt(11)
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    rag = data.get("statut_rag","?")
    rag_p = doc.add_paragraph()
    run_r = rag_p.add_run(f"Avis actuariel : {'FAVORABLE' if rag=='VERT' else 'AVEC RÉSERVES' if rag=='AMBRE' else 'DÉFAVORABLE'}")
    run_r.bold = True; run_r.font.size = Pt(13)
    run_r.font.color.rgb = _rgb(VERT if rag=="VERT" else AMBRE if rag=="AMBRE" else ROUGE)
    rag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ── Résumé exécutif ───────────────────────────────────────────────────────
    doc.add_heading("Résumé exécutif", level=1).runs[0].font.color.rgb = _rgb(NAVY)
    doc.add_paragraph(data.get("commentaire_resume", "Voir rapport complet ci-dessous."))

    doc.add_page_break()

    # ── Section 1 : Portefeuille ───────────────────────────────────────────────
    _add_section_title(doc, 1, "Présentation du portefeuille")
    _add_kpi_table(doc,
        ["Indicateur", "Santé", "Prévoyance", "Consolidé"],
        [
            ("Assurés",        f"{data.get('nb_assures_sante',0):,}",   f"{data.get('nb_assures_prev',0):,}", "—"),
            ("Primes (€)",     f"{data.get('pa_sante',0):,.0f}",        f"{data.get('pa_prev',0):,.0f}",    f"{data.get('pa_sante',0)+data.get('pa_prev',0):,.0f}"),
            ("Age moyen",      f"{data.get('age_moyen_sante',0):.1f}",  f"{data.get('age_moyen_prev',0):.1f}", "—"),
            ("Loss Ratio",     f"{data.get('lr_sante',0):.1%}",         f"{data.get('lr_prev',0):.1%}",     "—"),
        ]
    )

    # ── Section 2 : Sinistralité santé ────────────────────────────────────────
    _add_section_title(doc, 2, "Sinistralité santé — LR par poste (DREES 2023)")
    postes = data.get("sinistralite_par_poste", {})
    rows_sin = []
    for poste, infos in postes.items():
        if isinstance(infos, dict):
            rows_sin.append((
                poste.capitalize(),
                f"{infos.get('frequence_an',0):.2f}",
                f"{infos.get('charge_mutuelle',0):,.0f}€",
                "DREES 2023",
            ))
    if rows_sin:
        _add_kpi_table(doc, ["Poste","Fréquence/an","Charge mutuelle","Source"], rows_sin)
    else:
        doc.add_paragraph("Données sinistralité non disponibles.")

    # ── Section 3 : Morbidité prévoyance ──────────────────────────────────────
    _add_section_title(doc, 3, "Morbidité prévoyance — ITT/IP/Décès (BCAC 2019)")
    morb = data.get("morbidite", {})
    doc.add_paragraph(
        f"Taux ITT (employe, 42 ans) : {morb.get('taux_itt',0):.2%} | "
        f"Taux IP : {morb.get('taux_ip',0):.3%} | "
        f"P(maintien 6m) : {morb.get('maintien_6m',0):.3f}"
    )

    # ── Section 4 : Provisions ────────────────────────────────────────────────
    _add_section_title(doc, 4, "Provisions techniques (PSAP + PM + BE + RA)")
    _add_kpi_table(doc,
        ["Poste","Santé (€)","Prévoyance (€)","Source"],
        [
            ("PSAP",              f"{data.get('psap_sante',0):,.0f}",    f"{data.get('psap_prev',0):,.0f}",    "Art.27 C.ass."),
            ("PM Rentes IP",      "—",                                    f"{data.get('pm_rentes_ip',0):,.0f}", "BCAC 2019 + EIOPA RFR"),
            ("Best Estimate",     f"{data.get('be_sante',0):,.0f}",      f"{data.get('be_prev',0):,.0f}",      "IFRS 17 §33"),
            ("Risk Adjustment",   f"{data.get('ra_sante',0):,.0f}",      f"{data.get('ra_prev',0):,.0f}",      "§B91 CoC 6%"),
        ]
    )

    # ── Section 5 : Solvabilité S2 ────────────────────────────────────────────
    _add_section_title(doc, 5, "Solvabilité 2 — SCR/MCR consolidé (ρ=0.25 EIOPA)")
    _add_kpi_table(doc,
        ["Module","Montant (€)","Source"],
        [
            ("SCR Santé NSLT",    f"{data.get('scr_sante',0):,.0f}",    "Art.148 RD 2015/35"),
            ("SCR Invalidité SLT",f"{data.get('scr_prev',0):,.0f}",     "Art.145 RD 2015/35"),
            ("Diversification",   f"-{data.get('diversification',0):,.0f}", "Annexe IV"),
            ("SCR Consolidé",     f"{data.get('scr_consolide',0):,.0f}","EIOPA"),
            ("Fonds Propres",     f"{data.get('fonds_propres',0):,.0f}","Bilan S2"),
            (f"Ratio SCR = {data.get('ratio_scr',0):.1f}%", "","FP/SCR seuil 100%"),
        ]
    )

    # ── Section 6 : IFRS 17 ───────────────────────────────────────────────────
    _add_section_title(doc, 6, "IFRS 17 — BE / RA / FCF / CSM / LC")
    i17 = data.get("ifrs17", {})
    _add_kpi_table(doc,
        ["Composante","Santé (PAA)","Prévoyance (GMM)","Réf."],
        [
            ("BE",  f"{i17.get('be_sante',0):,.0f}",  f"{i17.get('be_prev',0):,.0f}",  "§33"),
            ("RA",  f"{i17.get('ra_sante',0):,.0f}",  f"{i17.get('ra_prev',0):,.0f}",  "§B91"),
            ("FCF", f"{i17.get('fcf_sante',i17.get('be_sante',0)+i17.get('ra_sante',0)):,.0f}",
                    f"{i17.get('fcf_prev',i17.get('be_prev',0)+i17.get('ra_prev',0)):,.0f}", "§33"),
            ("CSM", f"{i17.get('csm_sante',0):,.0f}", f"{i17.get('csm_prev',0):,.0f}", "§38"),
        ]
    )

    # ── Section 7 : Réglementation ────────────────────────────────────────────
    _add_section_title(doc, 7, "Réglementation — ANI 2013 / 100% Santé / Contrat responsable")
    reg = data.get("reglementation", {})
    doc.add_paragraph(
        f"ANI 2013 : {'✅ Conforme' if reg.get('ani_conforme') else '❌ Non conforme'} — "
        f"100% Santé : {reg.get('sante_100_note','Non vérifié')} — "
        f"Contrat responsable : {reg.get('contrat_resp_note','Non vérifié')}"
    )

    # ── Section 8 : Avis actuariel ────────────────────────────────────────────
    _add_section_title(doc, 8, "Avis actuariel et hypothèses")
    avis = "FAVORABLE" if rag=="VERT" else "AVEC RÉSERVES" if rag=="AMBRE" else "DÉFAVORABLE"
    p_avis = doc.add_paragraph()
    run_av = p_avis.add_run(f"Avis : {avis}")
    run_av.bold = True
    run_av.font.color.rgb = _rgb(VERT if rag=="VERT" else AMBRE if rag=="AMBRE" else ROUGE)

    hyps = data.get("hypotheses_rapport", [])
    for h in hyps:
        ic = "✅" if h.get("statut")=="VALIDÉE" else ("⚠️" if h.get("statut")=="À JUSTIFIER" else "❌")
        doc.add_paragraph(f"{ic} {h.get('hypothese','')} → {h.get('valeur','')} : {h.get('statut','')}")

    doc.add_paragraph()
    p_sig = doc.add_paragraph()
    run_sig = p_sig.add_run(f"Généré par ActuarIA — {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run_sig.font.color.rgb = _rgb(OR); run_sig.font.size = Pt(9)
    p_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf_sp(data: Dict) -> bytes:
    """
    Génère le rapport SP au format PDF via WeasyPrint.
    Fallback HTML si WeasyPrint indisponible.

    Returns
    -------
    bytes — PDF ou HTML encodé en UTF-8
    """
    html = _build_html_sp(data)
    if WEASYPRINT_OK:
        try:
            return WP_HTML(string=html).write_pdf()
        except Exception as e:
            logger.warning(f"WeasyPrint PDF SP échoué : {e} — fallback HTML")
    return html.encode("utf-8")


def _build_html_sp(data: Dict) -> str:
    """Génère le HTML du rapport SP."""
    rag = data.get("statut_rag","?")
    avis = "FAVORABLE" if rag=="VERT" else "AVEC RÉSERVES" if rag=="AMBRE" else "DÉFAVORABLE"
    rag_color = VERT if rag=="VERT" else AMBRE if rag=="AMBRE" else ROUGE

    def fmt(v, pct=False, eur=False):
        if v is None or v == "": return "—"
        if pct: return f"{float(v):.1%}"
        if eur: return f"{float(v):,.0f} €"
        return str(v)

    postes_html = ""
    for poste, infos in data.get("sinistralite_par_poste",{}).items():
        if isinstance(infos, dict):
            postes_html += f"""<tr>
              <td>{poste.capitalize()}</td>
              <td>{infos.get('frequence_an',0):.2f}</td>
              <td>{infos.get('charge_mutuelle',0):,.2f} €</td>
              <td>DREES 2023</td>
            </tr>"""

    hyps_html = ""
    for h in data.get("hypotheses_rapport",[]):
        ic = "✅" if h.get("statut")=="VALIDÉE" else ("⚠️" if h.get("statut")=="À JUSTIFIER" else "❌")
        hyps_html += f"<li>{ic} <strong>{h.get('hypothese','')}</strong> — {h.get('valeur','')} : <em>{h.get('statut','')}</em></li>"

    return f"""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="utf-8">
<style>
  body {{ font-family: Arial, sans-serif; background:#F5F7FA; color:#1C2B3A; margin:0; padding:20px; }}
  .page-garde {{ background:{NAVY}; color:{BLANC}; padding:40px; text-align:center; margin-bottom:30px; }}
  .page-garde h1 {{ color:{OR}; font-size:24px; margin:0 0 10px 0; }}
  .page-garde h2 {{ font-size:16px; margin:0 0 8px 0; }}
  .avis {{ font-size:18px; font-weight:bold; color:{rag_color}; }}
  h2.section {{ color:{NAVY}; border-bottom:2px solid {OR}; padding-bottom:4px; }}
  table {{ width:100%; border-collapse:collapse; margin-bottom:20px; }}
  thead th {{ background:{NAVY}; color:{OR}; padding:8px; text-align:left; font-size:11px; }}
  tbody tr:nth-child(even) {{ background:#E8EDF3; }}
  tbody tr:nth-child(odd)  {{ background:#F5F7FA; }}
  tbody td {{ padding:6px 8px; font-size:10px; border-bottom:1px solid #DDE4EE; }}
  .footer {{ color:#8A9AB0; font-size:9px; text-align:right; margin-top:20px; }}
</style>
</head>
<body>

<div class="page-garde">
  <h1>RAPPORT ACTUARIEL SANTÉ-PRÉVOYANCE</h1>
  <h2>{data.get('entite','ActuarIA')}</h2>
  <p>Arrêté au {data.get('date_arrete','')} | Généré le {datetime.now().strftime('%d/%m/%Y')}</p>
  <p class="avis">Avis : {avis}</p>
</div>

<h2 class="section">1. Portefeuille</h2>
<table>
  <thead><tr><th>Indicateur</th><th>Santé</th><th>Prévoyance</th><th>Consolidé</th></tr></thead>
  <tbody>
    <tr><td>Primes acquises</td><td>{fmt(data.get('pa_sante'),eur=True)}</td><td>{fmt(data.get('pa_prev'),eur=True)}</td><td>{fmt(data.get('pa_sante',0)+data.get('pa_prev',0),eur=True)}</td></tr>
    <tr><td>Best Estimate</td><td>{fmt(data.get('be_sante'),eur=True)}</td><td>{fmt(data.get('be_prev'),eur=True)}</td><td>{fmt(data.get('be_sante',0)+data.get('be_prev',0),eur=True)}</td></tr>
    <tr><td>Loss Ratio</td><td>{fmt(data.get('lr_sante'),pct=True)}</td><td>{fmt(data.get('lr_prev'),pct=True)}</td><td>—</td></tr>
  </tbody>
</table>

<h2 class="section">2. Sinistralité santé</h2>
<table>
  <thead><tr><th>Poste</th><th>Fréquence/an</th><th>Charge mutuelle</th><th>Source</th></tr></thead>
  <tbody>{postes_html}</tbody>
</table>

<h2 class="section">3. Provisions techniques</h2>
<table>
  <thead><tr><th>Poste</th><th>Santé</th><th>Prévoyance</th><th>Référence</th></tr></thead>
  <tbody>
    <tr><td>PSAP</td><td>{fmt(data.get('psap_sante'),eur=True)}</td><td>{fmt(data.get('psap_prev'),eur=True)}</td><td>Art.27 C.ass.</td></tr>
    <tr><td>PM Rentes IP</td><td>—</td><td>{fmt(data.get('pm_rentes_ip'),eur=True)}</td><td>BCAC 2019 + EIOPA RFR</td></tr>
    <tr><td>Best Estimate</td><td>{fmt(data.get('be_sante'),eur=True)}</td><td>{fmt(data.get('be_prev'),eur=True)}</td><td>IFRS 17 §33</td></tr>
    <tr><td>Risk Adjustment</td><td>{fmt(data.get('ra_sante'),eur=True)}</td><td>{fmt(data.get('ra_prev'),eur=True)}</td><td>§B91 CoC 6%</td></tr>
  </tbody>
</table>

<h2 class="section">4. Solvabilité 2</h2>
<table>
  <thead><tr><th>Module</th><th>Montant</th><th>Référence</th></tr></thead>
  <tbody>
    <tr><td>SCR Santé NSLT</td><td>{fmt(data.get('scr_sante'),eur=True)}</td><td>Art.148 RD 2015/35</td></tr>
    <tr><td>SCR Invalidité SLT</td><td>{fmt(data.get('scr_prev'),eur=True)}</td><td>Art.145 RD 2015/35</td></tr>
    <tr><td>Bénéfice diversification</td><td>-{fmt(data.get('diversification'),eur=True)}</td><td>Annexe IV ρ=0.25</td></tr>
    <tr><td><strong>SCR Consolidé</strong></td><td><strong>{fmt(data.get('scr_consolide'),eur=True)}</strong></td><td>EIOPA</td></tr>
    <tr><td>Fonds Propres</td><td>{fmt(data.get('fonds_propres'),eur=True)}</td><td>Bilan S2</td></tr>
    <tr><td><strong>Ratio SCR</strong></td><td><strong>{data.get('ratio_scr',0):.1f}%</strong></td><td>Seuil 100% Art.129</td></tr>
  </tbody>
</table>

<h2 class="section">5. Hypothèses actuarielles</h2>
<ul>{hyps_html}</ul>

<p class="footer">Généré par ActuarIA · Direction Santé-Prévoyance · {datetime.now().strftime('%d/%m/%Y %H:%M')}</p>

</body></html>"""
