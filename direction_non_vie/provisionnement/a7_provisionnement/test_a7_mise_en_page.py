# -*- coding: utf-8 -*-
"""
=============================================================================
 A7 Ibrahim — ce que deviennent les couleurs et les pages à l'impression
=============================================================================

 ⚠️ DEUX DÉFAUTS DISTINCTS, ET UN SEUL RESSEMBLAIT À UN PROBLÈME DE
 CONVERSION.

 1. LE HTML. Un navigateur SUPPRIME par défaut les fonds colorés à
    l'impression — la case « Graphiques d'arrière-plan » du dialogue Ctrl+P
    est DÉCOCHÉE. Le rapport pose plus de 180 fonds (69 `background`, 5
    dégradés, 117 `rgba`) et le signal RAG en dépend : sans
    `print-color-adjust: exact`, un statut ROUGE s'imprime sur fond blanc,
    exactement comme un VERT.
    Et le bloc `@media print` qui existait FAISAIT PIRE QUE RIEN : son
    `@page { margin: 0 }` supprimait toute marge, si bien que le contenu
    courait jusqu'au bord physique de la feuille — zone que la plupart des
    imprimantes et convertisseurs rognent.

 2. LE WORD. Ce n'est PAS un problème de conversion : c'est du code de
    génération séparé, et il n'a jamais porté ces couleurs. Mesuré sur un
    document produit — 99 fonds de cellule, mais DEUX teintes seulement, le
    gris de zébrage et le navy des en-têtes. Les dix-sept lignes du tableau
    des hypothèses affichaient « VALIDÉE » et « NON VALIDÉE » dans la même
    encre noire. Depuis la décision B, le Word est le MAÎTRE D'IMPRESSION du
    PDF : ses couleurs portent le signal, elles ne le décorent pas.

 ⚠️ CE QUE CE FILET NE PEUT PAS VÉRIFIER, ET IL FAUT LE DIRE. Chrome en mode
 « headless » PRÉSERVE les fonds par défaut — vérifié par une expérience
 minimale : un fond `#C0392B` sans aucune règle d'impression ressort bien
 dans le PDF produit. Le symptôme ne se reproduit donc QUE dans le dialogue
 d'impression interactif. Ces verrous contrôlent que les règles sont émises
 et bien formées ; seul un rendu manuel dans Chrome ou Edge peut confirmer
 le résultat visuel.
=============================================================================
"""

import io
import re
import unittest
import zipfile

import numpy as np

from direction_non_vie.provisionnement.a7_provisionnement.agent import (
    AgentA7Provisionnement)
from direction_non_vie.provisionnement.a7_provisionnement.test_a7_ibrahim import (
    GENINS, _TRI_RECOURS_FORT)
from direction_non_vie.provisionnement.a7_provisionnement.test_a7_graphiques import (
    kaleido_declare, rendeur_substitue)

#: Le vocabulaire des verdicts, RELEVÉ sur huit exécutions — quatre triangles
#: × avec et sans exposition. Il est fini : la correspondance peut donc être
#: EXACTE, et une correspondance exacte ne produit aucun faux positif.
VERDICTS = {
    'VALIDÉE', 'À JUSTIFIER', 'NON VALIDÉE', 'REJETÉE', 'NON TESTABLE',
    'Non disponible', 'non calculée', 'VERT', 'AMBRE', 'ROUGE',
    '✓ Inclus', '⊘ Exclu', '✅ OUI', '❌ NON',
}
#: Les trois fonds du signal — mêmes teintes que le HTML et que l'Excel.
FONDS_RAG = {'EAF3DE': 'vert', 'FAEEDA': 'ambre', 'FCEBEB': 'rouge'}
#: Les trois encres correspondantes.
ENCRES_RAG = {'1E8449': 'vert', 'E67E22': 'ambre', 'C0392B': 'rouge'}

_CACHE = {}


def _run(triangle=GENINS, avec_expo=True, cle=None):
    if cle is not None and cle in _CACHE:
        return _CACHE[cle]
    C = np.asarray(triangle, dtype=float)
    kw = {}
    if avec_expo:
        kw['primes'] = np.full(C.shape[0], float(np.nanmean(C[:, 0])) * 8.0)
    # ⚠️ CES VERROUS PORTENT SUR LES MARGES ET LES COULEURS, pas sur les
    # pixels d'une figure. Ils faisaient rasteriser 26 figures — 143 s
    # mesurées — pour lire une taille de page et des fonds de cellule.
    with kaleido_declare(True), rendeur_substitue():
        r = AgentA7Provisionnement(verbose=False).run(
            source=C, mode_declare='cumule', generer_graphiques=True,
            generer_word=True, n_sim_bootstrap=60, seed=42, **kw)
    if cle is not None:
        _CACHE[cle] = r
    return r


def _xml_word(octets):
    return zipfile.ZipFile(io.BytesIO(octets)).read(
        'word/document.xml').decode('utf-8')


# =============================================================================
#  T1 — LE HTML SAIT S'IMPRIMER
# =============================================================================

class T1_Le_Html_Sait_S_Imprimer(unittest.TestCase):

    def _bloc_print(self, html):
        """Le contenu du `@media print`, accolades équilibrées."""
        i = html.find('@media print')
        self.assertGreater(i, 0, 'aucun bloc @media print dans le rapport')
        prof, fin = 0, None
        for k, ch in enumerate(html[i:]):
            if ch == '{':
                prof += 1
            elif ch == '}':
                prof -= 1
                if prof == 0:
                    fin = k
                    break
        self.assertIsNotNone(fin, 'bloc @media print non refermé')
        return html[i:i + fin + 1]

    def _bloc_print_sans_commentaires(self, html):
        """Le même bloc, commentaires CSS retirés.

        ⚠️ NÉCESSAIRE, ET C'EST UNE MESURE QUI L'A APPRIS : le bloc documente
        l'ancienne règle fautive en la CITANT — « `@page { margin: 0 }` était
        pire que pas de règle ». Une recherche naïve retrouve donc la citation
        et croit la règle toujours en place. Un commentaire n'est pas une
        règle : on le retire avant de juger.
        """
        return re.sub(r'/\*.*?\*/', '', self._bloc_print(html), flags=re.S)

    def test_les_fonds_survivent_a_l_impression(self):
        """⚠️ LA RÈGLE QUI RÉPARE LE SYMPTÔME SIGNALÉ.

        Les deux orthographes sont exigées : Chrome et Edge — qui font la
        conversion PDF — ne lisent encore que le préfixe `-webkit-` dans
        certaines versions.
        """
        bloc = self._bloc_print(_run(cle='genins').get('html') or '')
        self.assertIn('print-color-adjust: exact', bloc)
        self.assertIn('-webkit-print-color-adjust: exact', bloc)
        self.assertIn('!important', bloc,
                      'sans !important la règle est écrasée par les styles '
                      'des éléments')
        print('    OK C4-1 print-color-adjust: exact, avec le préfixe webkit '
              'et !important')

    def test_la_page_a_des_marges_et_le_bon_format(self):
        """⚠️ `@page { margin: 0 }` ÉTAIT PIRE QUE PAS DE RÈGLE.

        Et le rapport s'imprimait en US LETTER — mesuré sur un PDF produit
        par Chrome : MediaBox 612 × 792 pt avant, 595 × 842 pt (A4) après.
        Pour un rapport réglementaire français, ce n'est pas un détail.
        """
        bloc = self._bloc_print_sans_commentaires(
            _run(cle='genins').get('html') or '')
        self.assertIn('size: A4', bloc, 'le format de page n\'est pas fixé — '
                                        'le rendu retombe sur US Letter')
        m = re.search(r'@page\s*\{[^}]*margin:\s*([^;}]+)', bloc)
        self.assertIsNotNone(m, 'aucune marge déclarée pour @page')
        self.assertNotEqual(m.group(1).strip(), '0',
                            'la marge est nulle : le contenu court jusqu\'au '
                            'bord de la feuille et se fait rogner')
        self.assertIn('@page :first', bloc,
                      'la page de garde perd son fond perdu, ou impose sa '
                      'marge nulle au reste du document')
        print('    OK C4-2 @page A4, marges « %s », page de garde à part'
              % m.group(1).strip())

    def test_rien_d_indivisible_n_est_coupe_entre_deux_pages(self):
        """Un tableau scindé perd son en-tête, une figure scindée ne dit
        plus rien, et un titre seul en bas de page est la faute la plus
        visible d'un rapport."""
        bloc = self._bloc_print(_run(cle='genins').get('html') or '')
        for regle in ('break-inside: avoid', 'page-break-inside: avoid',
                      'break-after: avoid', 'display: table-header-group',
                      'orphans', 'widows'):
            self.assertIn(regle, bloc, 'règle absente : %s' % regle)
        for classe in ('table.premium', '.hyp-card', '.plotly-graph-div'):
            self.assertIn(classe, bloc,
                          '%s peut encore être coupé en deux' % classe)
        print('    OK C4-3 tableaux, cartes et figures indivisibles ; '
              'en-têtes répétés ; ni orpheline ni veuve')

    def test_les_classes_visees_existent_vraiment_dans_le_rapport(self):
        """⚠️ UNE RÈGLE QUI VISE UNE CLASSE ABSENTE NE PROTÈGE RIEN.

        C'est le défaut silencieux typique d'une feuille de style : elle a
        l'air complète et ne s'applique à personne.
        """
        html = _run(cle='genins').get('html') or ''
        bloc = self._bloc_print(html)
        classes = {c.lstrip('.') for c in
                   re.findall(r'(?:^|\s|,)(\.[a-z][a-z0-9-]+)', bloc)}
        self.assertTrue(classes, 'aucune classe ciblée')
        absentes = [c for c in classes
                    if not re.search(r'class="[^"]*\b%s\b' % re.escape(c),
                                     html)]
        self.assertEqual(absentes, [],
                         'règles d\'impression sur des classes qui '
                         'n\'existent pas : %s' % absentes)
        print('    OK C4-4 les %d classes visées existent toutes dans le '
              'rapport généré' % len(classes))


# =============================================================================
#  T2 — LE WORD PORTE LE SIGNAL, PAS SEULEMENT LE MOT
# =============================================================================

class T2_Le_Word_Porte_Le_Signal_Rag(unittest.TestCase):

    def test_les_verdicts_recoivent_un_fond_et_une_encre(self):
        r = _run(cle='genins')
        octets = r.get('word_bytes') or b''
        if not octets:
            self.skipTest('python-docx absent')
        xml = _xml_word(octets)
        fonds = re.findall(r'w:shd[^>]*w:fill="(%s)"' % '|'.join(FONDS_RAG), xml)
        encres = re.findall(r'w:color w:val="(%s)"' % '|'.join(ENCRES_RAG), xml)
        self.assertGreater(len(fonds), 0,
                           'aucun fond de statut dans le Word : les tableaux '
                           'ne portent plus le signal RAG')
        self.assertGreater(len(encres), 5,
                           'le RAG n\'existe que sur quelques runs isolés')
        print('    OK C4-5 Word : %d fonds de statut et %d encres RAG'
              % (len(fonds), len(encres)))

    def test_seules_les_cellules_de_verdict_sont_colorees(self):
        """⚠️ LA CORRESPONDANCE EST EXACTE, ET C'EST CE QUI LA REND SÛRE.

        Un seuil de longueur ou une recherche de sous-chaîne repeindrait la
        colonne « Message », qui contient « BFCC-H4 VALIDÉE — loss ratio… ».
        """
        import docx
        octets = _run(cle='genins').get('word_bytes') or b''
        if not octets:
            self.skipTest('python-docx absent')
        doc = docx.Document(io.BytesIO(octets))
        fautes = []
        for tbl in doc.tables:
            for ligne in tbl.rows[1:]:
                for cel in ligne.cells:
                    colore = any(
                        run.font.color is not None
                        and run.font.color.rgb is not None
                        and str(run.font.color.rgb) in ENCRES_RAG
                        for p in cel.paragraphs for run in p.runs)
                    if colore and cel.text.strip() not in VERDICTS:
                        fautes.append(cel.text.strip()[:50])
        self.assertEqual(fautes, [],
                         'cellules colorées hors vocabulaire : %s' % fautes)
        print('    OK C4-6 aucune cellule colorée hors du vocabulaire des '
              '%d verdicts' % len(VERDICTS))

    def test_les_trois_couleurs_du_rag_sont_reellement_atteignables(self):
        """⚠️ SANS CE VERROU, UN CHEMIN MORT PASSERAIT INAPERÇU.

        GenIns valide tout : le rouge et l'ambre n'y apparaissent presque
        pas. Il faut un triangle où des hypothèses tombent pour prouver que
        les trois teintes sortent.
        """
        octets = _run(_TRI_RECOURS_FORT, True, cle='recfort').get(
            'word_bytes') or b''
        if not octets:
            self.skipTest('python-docx absent')
        xml = _xml_word(octets)
        vus = {lbl for code, lbl in FONDS_RAG.items()
               if 'w:fill="%s"' % code in xml}
        self.assertEqual(vus, set(FONDS_RAG.values()),
                         'teintes manquantes : %s'
                         % (set(FONDS_RAG.values()) - vus))
        print('    OK C4-7 les trois fonds RAG sortent sur un triangle à '
              'hypothèses en défaut : %s' % ', '.join(sorted(vus)))

    def test_la_couleur_n_est_jamais_seule_porteuse_du_sens(self):
        """Imprimé en noir et blanc, ou lu par quelqu'un qui distingue mal le
        rouge du vert, le document doit encore dire ce qu'il dit."""
        import docx
        octets = _run(_TRI_RECOURS_FORT, True, cle='recfort').get(
            'word_bytes') or b''
        if not octets:
            self.skipTest('python-docx absent')
        doc = docx.Document(io.BytesIO(octets))
        muettes = []
        for tbl in doc.tables:
            for ligne in tbl.rows[1:]:
                for cel in ligne.cells:
                    for p in cel.paragraphs:
                        for run in p.runs:
                            c = run.font.color
                            if (c is not None and c.rgb is not None
                                    and str(c.rgb) in ENCRES_RAG
                                    and not run.text.strip()):
                                muettes.append(cel.text[:40])
        self.assertEqual(muettes, [],
                         'une cellule ne porte QUE de la couleur : %s'
                         % muettes)
        print('    OK C4-8 chaque cellule colorée porte aussi son mot')


if __name__ == '__main__':
    unittest.main(verbosity=2)
