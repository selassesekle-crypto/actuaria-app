# =============================================================================
#  ActuarIA — Agent A7 Ibrahim v5.0
#  n5_rapport.py  —  Rapport actuariel professionnel v5.0
#
#  Style : reproduction exacte du rapport premium ActuarIA
#  3 formats : HTML · Word (.docx) · Excel — le PDF par CONVERSION
#  Narration : Claude API → Templates → Données seules (3 niveaux)
#
#  Corrections vs premium :
#    · "Version A7-IBRAHIM · v4.0.0" supprimé de la page de garde
#    · Back-testing : tableau détaillé si données dispo, message sinon
#    · Section 7 : narration complète sans coupure
#    · LOB codes → libellés français (rc_medicale → RC Médicale)
#    · BF poids garanti ≥ 50% si méthode recommandée
#
#  Auteur  : ActuarIA v5.0
#  Version : 5.0.0
# =============================================================================

from __future__ import annotations
import base64, io, logging, re
from datetime import datetime
from typing import Dict, List, NamedTuple, Tuple

from core import frontiere_llm
from core import traitement_ia
import numpy as np

# Source UNIQUE d'affichage des hypothèses de BF et Cape Cod : une hypothèse non
# évaluée y ressort NON TESTABLE, jamais en valeur par défaut.
from .methodes_be import (ORDRE_AFFICHAGE, libelle, motif_exclusion,
                          reserve)
from .n5_graphiques import TITRES_FIGURES
from .n2_hypotheses_bfcc import lignes_hypotheses_bfcc
from .n2_hypotheses_clm import lignes_hypotheses_clm
from .n2_hypotheses_bootstrap import lignes_hypotheses_bootstrap
from .n2_hypotheses_munich import lignes_hypotheses_munich
# Source UNIQUE d'affichage de Munich CL, partagée par HTML, Word et Excel.
from .n3.munich_cl import lignes_munich_rapport
# Source UNIQUE d'affichage de Benktander, partagee par HTML, Word et Excel.
from .n3.benktander import lignes_benktander_rapport
# ⚠️ IMPORT MANQUANT, BUG DE PRODUCTION (lot F2, trouvé par ruff F821).
# `_construire_contexte` appelle `libelle_loss_ratio` ligne 237 sans jamais
# l'importer : l'appel levait `NameError`. Son unique appelant l'enveloppe
# dans un `except Exception` qui journalise « Claude API indisponible » — le
# commentaire actuariel par Claude n'a donc JAMAIS pu être produit, et le
# message accusait l'API d'une panne qui était dans le code.
from .n3.bf_cape_cod import libelle_loss_ratio
# Source UNIQUE du NOM de l'approche publiée dans `reserve_p*` — la même que
# l'Excel et le commentaire. Ces libellés étaient écrits en dur dans les deux
# formats de ce fichier, et « (retenue) » y était cloué sur le composé.
from .n4_best_estimate import (CLE_BOOT, CLE_COMPOSE, CLE_MACK,
                               libelle_percentiles, marque_retenue)

logger = logging.getLogger('actuaria.a7.rapport')

#: Marqueur en tête du repli d'`export_html`. Un commentaire HTML : invisible à
#: l'écran, détectable par le code. Il permet de distinguer une page d'erreur
#: d'un rapport, ce que ni la taille ni la validité du HTML ne permettent.
MARQUEUR_ECHEC_RAPPORT = '<!--ACTUARIA-RAPPORT-ECHEC-->'

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY      = '#0B1E3D'
NAVY_MID  = '#132844'
NAVY_L    = '#1E3A5F'
GOLD      = '#C9A84C'
GOLD_L    = '#E2C97E'
GOLD_PALE = 'rgba(201,168,76,0.12)'
CYAN      = '#4A8FD4'
ROUGE     = '#C0392B'
ORANGE    = '#E67E22'
VERT      = '#1E8449'
SLATE     = '#8A9BB0'
SLATE_L   = '#B8C5D3'
BG        = '#F5F7FA'
WHITE     = '#FFFFFF'
TEXT      = '#1C2B3A'
TEXT_MID  = '#3D5166'
BORDER    = '#DDE4EE'

# ── Mapping LOB ───────────────────────────────────────────────────────────────
LOB_LABELS = {
    'mrh':             'MRH — Multirisques Habitation',
    'rc_auto':         'RC Automobile',
    'rc_generale':     'RC Générale',
    'rc_medicale':     'RC Médicale',
    'construction':    'Construction (DO / RC Décennale)',
    'transport':       'Transport Maritime & Terrestre',
    'credit_caution':  'Crédit & Caution',
    'catastrophe_nat': 'Catastrophes Naturelles',
    'autre':           'Autre branche Non-Vie',
    'generique':       'Branche Non-Vie',
}

def _lob(code: str) -> str:
    if not code:
        return 'Branche Non-Vie'
    c = str(code).lower().strip().replace(' ', '_')
    return LOB_LABELS.get(c, code.replace('_', ' ').title())

# =============================================================================
#  LOGO SVG
# =============================================================================

LOGO_SVG = (
    '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 400 110" width="400" height="110">'
    '<rect width="400" height="110" fill="#0F2E52" rx="8"/>'
    '<g transform="translate(42,55)">'
    '<circle cx="0" cy="-27" r="5" fill="none" stroke="#4A8FD4" stroke-width="1.2"/>'
    '<circle cx="0" cy="-9"  r="5" fill="none" stroke="#4A8FD4" stroke-width="1.2"/>'
    '<circle cx="0" cy="9"   r="5" fill="none" stroke="#4A8FD4" stroke-width="1.2"/>'
    '<circle cx="0" cy="27"  r="5" fill="none" stroke="#4A8FD4" stroke-width="1.2"/>'
    '<circle cx="26" cy="-16" r="5" fill="none" stroke="#C9A84C" stroke-width="1.2"/>'
    '<circle cx="26" cy="4"   r="5" fill="none" stroke="#C9A84C" stroke-width="1.2"/>'
    '<circle cx="26" cy="24"  r="5" fill="none" stroke="#C9A84C" stroke-width="1.2"/>'
    '<circle cx="52" cy="4"   r="8" fill="none" stroke="#C9A84C" stroke-width="1.5"/>'
    '<text x="52" y="8" text-anchor="middle" font-family="Georgia,serif" font-size="8" font-weight="700" fill="#C9A84C">A</text>'
    '<line x1="5" y1="-27" x2="21" y2="-16" stroke="#4A8FD4" stroke-width="0.6" opacity="0.6"/>'
    '<line x1="5" y1="-9"  x2="21" y2="-16" stroke="#4A8FD4" stroke-width="0.6" opacity="0.6"/>'
    '<line x1="5" y1="-9"  x2="21" y2="4"   stroke="#4A8FD4" stroke-width="0.6" opacity="0.6"/>'
    '<line x1="5" y1="9"   x2="21" y2="4"   stroke="#4A8FD4" stroke-width="0.6" opacity="0.6"/>'
    '<line x1="5" y1="9"   x2="21" y2="24"  stroke="#4A8FD4" stroke-width="0.6" opacity="0.6"/>'
    '<line x1="5" y1="27"  x2="21" y2="24"  stroke="#4A8FD4" stroke-width="0.6" opacity="0.6"/>'
    '<line x1="31" y1="-16" x2="44" y2="1"  stroke="#C9A84C" stroke-width="0.8" opacity="0.7"/>'
    '<line x1="31" y1="4"   x2="44" y2="4"  stroke="#C9A84C" stroke-width="0.8" opacity="0.7"/>'
    '<line x1="31" y1="24"  x2="44" y2="8"  stroke="#C9A84C" stroke-width="0.8" opacity="0.7"/>'
    '</g>'
    '<text x="122" y="50" font-family="Georgia,serif" font-size="38" font-weight="700">'
    '<tspan fill="#FFFFFF">Actuar</tspan><tspan fill="#C9A84C">IA</tspan>'
    '</text>'
    '<line x1="122" y1="58" x2="328" y2="58" stroke="#C9A84C" stroke-width="0.6" opacity="0.5"/>'
    '<text x="122" y="73" font-family="Georgia,serif" font-size="7.5" fill="#8A9BB0" letter-spacing="3">ACTUARIAL INTELLIGENCE PLATFORM</text>'
    '</svg>'
)
LOGO_URI = 'data:image/svg+xml;base64,' + base64.b64encode(LOGO_SVG.encode()).decode()


# =============================================================================
#  HELPERS
# =============================================================================

def _f(v, dec=0) -> str:
    if v is None:
        return '—'
    try:
        fv = float(v)
        if not np.isfinite(fv):
            return '—'
        sep = '\u202f'
        return f"{fv:,.0f}\u202f\u20ac".replace(',', sep) if dec == 0 else f"{fv:,.{dec}f}".replace(',', sep)
    except Exception:
        return '—'

def _pct(v, dec=1) -> str:
    if v is None:
        return '—'
    try:
        fv = float(v)
        return '—' if not np.isfinite(fv) else f"{fv:.{dec}f}\u202f%"
    except Exception:
        return '—'

def _s(v) -> str:
    if v is None:
        return '—'
    return re.sub(r'\s+', ' ', str(v)).strip() or '—'


def _scr_publiable(sc):
    """Le SCR Provisions a publier, ou `None` quand il n'est pas calculable.

    ⚠️⚠️ CE QUE CETTE FONCTION REMPLACE PORTAIT DEUX DEFAUTS DANS LA MEME
    EXPRESSION, ET ELLE VIVAIT EN TROIS EXEMPLAIRES (contexte du LLM, HTML,
    Word) :

        SCP = float(sc.get('scr_provisions',
                           sc.get('scr_prov', BE * 0.30)) if sc else BE * 0.30)

    · SCR absent  -> un chiffre FABRIQUE a 30 % du Best Estimate, publie sans
      la moindre marque. Mesure : BE de 1 000 -> << SCR 300 EUR, ratio 30,0 % >>.
      Atteignable : `actuaria_app` passe `r_raw.get("n4", {})`, donc un `n4`
      partiel arrive bien jusqu'ici.
    · `scr_provisions` PRESENT A None -> `.get` ne prend JAMAIS son repli,
      `float(None)` leve, et le rapport ne se genere pas du tout.

    ⚠️ ET C'EST LE SECOND CAS QUI COMPTE, PARCE QUE SA CAUSE EST L'INVERSE D'UN
    DEFAUT. `None` est pose par `garde_fou_be_negatif` quand le Best Estimate
    est NEGATIF : N4 y refuse d'inventer, marque l'absence et passe en ROUGE --
    << marqueurs None + statut ROUGE, jamais de plancher silencieux >>.
    Le renderer DEFAISAIT ce garde-fou en plantant dessus. La couche de calcul
    faisait ce qu'il fallait ; la couche de rendu ne savait pas le rendre.

    Rendre `None` suffit a tout reparer : `_f` et `_pct` publient deja un tiret
    sur `None`. L'instrument existait, il etait court-circuite en amont.
    """
    if not sc:
        return None
    valeur = sc.get('scr_provisions')
    if valeur is None:
        valeur = sc.get('scr_prov')
    if valeur is None:
        return None
    try:
        v = float(valeur)
    except (TypeError, ValueError):                    # pragma: no cover
        return None
    return v if np.isfinite(v) else None


def _ratio_scr(scp, be):
    """Le ratio SCR/BE en %, ou `None` des qu'un des deux manque."""
    if scp is None or not be:
        return None
    return scp / float(be) * 100.0


#: ⚠️ CE QUE LE RAPPORT AFFICHE QUAND AUCUN ARRETE N'EST COMMUNIQUE. Jusqu'a ce
#: lot, HTML et Word ecrivaient `arr = arrete or dt` : la DATE DU JOUR se
#: faisait passer pour l'arrete, sous une etiquette << Arrete >> (sous-titre de
#: garde et KPI cote HTML, colonne de table cote Word) et jusque dans le prompt
#: du LLM via `_generer_narration`. C'est le defaut exact qu'A4a a ferme dans
#: `n5_commentaire` (`arrete or date_arrete or 'NON COMMUNIQUE'`), sans que la
#: correction atteigne les deux exports. Le libelle nomme desormais l'absence
#: au lieu de la combler -- meme mot que le livrable frere, pour qu'un lecteur
#: retrouve la meme chose partout.
ARRETE_ABSENT = 'non communiqué'


#: Ce que chaque rendu en texte simple (prompt LLM, Word) dit de l'etat du test
#: calendaire. Le HTML garde son propre balisage, mais lit le MEME etat.
_PHRASE_CALENDAIRE = {
    'indisponible': "Test calendaire indisponible : le GLM Poisson "
                    "age-cohorte ne s'est pas ajuste sur ce triangle. "
                    "Absence de test, pas absence d'effet",
    'diffus':       "Effet calendaire global reparti (aucune diagonale "
                    "isolee dominante)",
    'aucun':        "Aucun effet calendaire significatif",
}

#: Couleur de l'encadre d'avis actuariel, cote HTML, PAR STATUT RAG.
#:
#: ⚠️ MATCH EXACT SUR UN VOCABULAIRE FINI, jamais une recherche de mot dans le
#: texte de l'avis -- meme principe que `_RAG_CELLULE` cote Word. Le statut
#: inconnu tombe sur l'orange (a surveiller), jamais sur le vert : une couleur
#: par defaut ne doit pas rassurer.
_COULEUR_AVIS_HTML = {
    'VERT':  'var(--vert)',
    'AMBRE': 'var(--orange)',
    'ROUGE': 'var(--rouge)',
}

#: Ce que le prompt du LLM dit de l'etat calendaire. `indisponible` porte une
#: consigne explicite : ne pas conclure a une absence d'effet.
_CAL_PROMPT = {
    'present':      'SIGNIFICATIF',
    'diffus':       'significatif mais diffus (aucune diagonale dominante)',
    'aucun':        'non significatif',
    'indisponible': 'INDISPONIBLE (GLM non ajuste -- NE PAS conclure a une '
                    'absence d effet)',
}


def etat_calendaire(bz):
    """L'etat du test calendaire GLM Poisson APC, en UN mot -- SOURCE UNIQUE
    des trois rendus (prompt LLM, HTML, Word).

    ⚠️ POURQUOI CETTE FONCTION EXISTE. Le HTML distinguait deja
    << indisponible >> de << aucun effet >> (le FIX 2 de la section 6). Le
    Word et le prompt LLM, eux, deduisaient << Aucun effet calendaire
    significatif >> / << non significatif >> du SEUL compteur d'effets, SANS
    regarder si le test avait tourne. Sur un triangle ou le GLM ne s'ajuste
    pas, ils affirmaient une absence d'EFFET la ou il n'y avait qu'absence de
    TEST -- le faux de F3 (un << OK >> vert sur zero test), transpose au
    calendaire, dans le document SIGNE (Word) et dans le prompt du MODELE.

    ⚠️ LA PRECEDENCE REPRODUIT EXACTEMENT CELLE DU HTML pour que sa sortie
    reste identique au caractere pres : un effet compte prime ; sinon
    l'indisponibilite ; sinon le diffus ; sinon l'absence. Verrouille par un
    test de non-regression sur le HTML.

    Rend 'present' | 'indisponible' | 'diffus' | 'aucun'.
    """
    if int(bz.get('n_effets_significatifs', 0) or 0) > 0:
        return 'present'
    if not bz.get('glm_disponible', False):
        return 'indisponible'
    if bz.get('cal_significatif', False):
        return 'diffus'
    return 'aucun'


def _clean(txt) -> str:
    if not txt:
        return ''
    txt = re.sub(r'[■□▪▸►═╔╗╚╝║─]+', '', str(txt))
    txt = re.sub(r'={4,}', '', txt)
    txt = re.sub(r'\n{3,}', '\n\n', txt)
    return txt.strip()

def _statut_col(s: str) -> str:
    return {'VERT': VERT, 'AMBRE': ORANGE, 'ROUGE': ROUGE}.get(s, SLATE)

def _statut_label(s: str) -> str:
    return {
        'VERT':  'Statut RAG : Vert — Conforme',
        'AMBRE': 'Statut RAG : Ambre — Vigilance',
        'ROUGE': 'Statut RAG : Rouge — Surveillance renforcée',
    }.get(s, 'Statut : ' + s)


# =============================================================================
#  SYSTEM PROMPT CLAUDE API
# =============================================================================

SYSTEM_PROMPT = """\
Tu es un actuaire Non-Vie senior certifié par l'Institut des Actuaires (IA), \
expert en provisionnement Solvabilité 2 et IFRS 17, avec 20 ans d'expérience \
auprès de mutuelles, institutions de prévoyance et grands groupes d'assurance.

Tu rédiges le commentaire actuariel d'un rapport de provisionnement Non-Vie \
destiné à être présenté au Conseil d'Administration et à l'actuaire désigné. \
Ce rapport sera soumis à l'ACPR dans le cadre du reporting Solvabilité 2.

RÈGLES ABSOLUES :
0. FORMAT : §N — TITRE pour les sections, ### pour les sous-titres, \
**gras** pour les termes importants, - tirets pour les listes. \
PAS de tableaux Markdown. PAS de blockquotes >. Sépare les sections par une ligne vide.
1. LANGUE : Français professionnel. Anglais uniquement pour les termes consacrés.
2. RIGUEUR : Chaque affirmation est justifiée par les données fournies.
3. CHIFFRES : En euros avec séparateurs (ex : 2\u202f526\u202f597\u202f€). Pourcentages avec une décimale.
4. RÉFÉRENCES : Art. 77 S2, Art. 115 S2, Guide IA 2023, Mack 1993, Clark 2003.
5. ALERTES : Ne jamais minimiser. Présenter avec l'implication réelle pour le bilan S2.
6. CAUSALITÉ : H1 rejetée (corr=0.52) → CL biaisé → BF retenu → impact X\u202f€ sur BE.
7. INCERTITUDE : Toujours quantifier via CV ou intervalles de confiance.
8. POSTURE : Assertif mais prudent. Recommandations claires avec justification.
9. INTERDIT : Phrases génériques sans données. FAVORABLE si H1 rejetée ET BT ROUGE.

STRUCTURE OBLIGATOIRE EN 7 SECTIONS :
§1 — CONTEXTE ET QUALITÉ DES DONNÉES
§2 — VALIDATION DES HYPOTHÈSES ACTUARIELLES
§3 — RÉSULTATS PAR MÉTHODE ET CONVERGENCE
§4 — INCERTITUDE STOCHASTIQUE ET PROVISIONS DE PRÉCAUTION
§5 — BACK-TESTING ET QUALITÉ DU PROVISIONNEMENT HISTORIQUE
§6 — EFFETS CALENDAIRE ET RISQUES IDENTIFIÉS
§7 — CONCLUSION ET RECOMMANDATIONS POUR LE CONSEIL D'ADMINISTRATION\
"""


# =============================================================================
#  CONTEXTE CLAUDE API
# =============================================================================

def _construire_contexte(n2: Dict, n3: Dict, n4: Dict, lob_label: str, arrete: str) -> str:
    mk    = n3.get('mack', {})
    clark = n3.get('clark', {});         bz  = n3.get('glm_apc', {})
    bt    = n3.get('backtesting', {});   sc  = n4.get('scr', {})
    h1    = n2.get('h1_independance', {}); h2 = n2.get('h2_stabilite', {})
    pw    = n4.get('poids', {})
    BE  = float(n4.get('best_estimate', 0) or 0)
    SIG = float(mk.get('sigma_total', 0) or 0)
    CV  = float(n4.get('cv_inter_methodes', 0) or 0)
    P75 = float(n4.get('reserve_p75', 0) or 0)
    P90 = float(n4.get('reserve_p90', 0) or 0)
    P99 = float(n4.get('reserve_p99_5', 0) or 0)
    SCP = _scr_publiable(sc)
    SCR = _ratio_scr(SCP, BE)
    lines = [
        f"DOSSIER DE PROVISIONNEMENT — {lob_label.upper()} — Arrêté {arrete}",
        "",
        "=== TRIANGLE ===",
        f"Dimensions : {n2.get('dimensions', '—')} (n={n2.get('n_lignes', '?')} années × m={n2.get('n_colonnes', '?')} périodes)",
        f"Années de survenance : {n2.get('annee_debut', '?')} à {n2.get('annee_fin', '?')}",
        f"Méthode retenue : {n4.get('methode_facteurs', n2.get('methode_recommandee', '—'))}",
        "",
        "=== HYPOTHÈSES ===",
        # ⚠️ SYMETRIQUE DE H2 CI-DESSOUS. `ok` vaut True par defaut sur les DEUX
        # chemins non testables de `_h1` (scipy absent, triangle < 7 ans) : le
        # LLM recevait « VALIDÉE | corr_moy=0 | score=80/100 » et redigeait
        # dessus.
        f"H1 Indépendance : "
        f"{h1.get('statut', 'VALIDÉE' if h1.get('ok') else 'REJETÉE')} | "
        + (f"corr_moy={h1.get('corr_moy', '—')} | score={h1.get('score', '—')}/100"
           if h1.get('statut') != 'NON TESTABLE' else
           "aucune paire de colonnes testable"),
        f"  Message : {str(h1.get('message', ''))[:200]}",
        # ⚠️ LE STATUT VIENT DU MODULE, PLUS DU BOOLÉEN DE GATING. `ok` vaut True
        # par défaut quand aucune colonne n'est testable : cette ligne publiait
        # alors « VALIDÉE | CV=0 | score=80/100 ».
        f"H2 Stabilité : "
        f"{h2.get('statut', 'VALIDÉE' if h2.get('ok') else 'REJETÉE')} | "
        + (f"CV={h2.get('cv_moy', '—')} | score={h2.get('score', '—')}/100"
           if h2.get('statut') != 'NON TESTABLE' else
           "aucune periode testable"),
        # ⚠️ H1 publiait son message, pas H2 : l'asymétrie masquait le seul
        # endroit où le module disait « non testable ».
        f"  Message : {str(h2.get('message', ''))[:200]}",
        f"Loss ratio a priori : {libelle_loss_ratio(n3.get('bf', {}))} (source : {n3.get('bf', {}).get('source_lr', '—')}) — UNIQUE, produit par N3.",
        *[f"{l['libelle']} : {l['statut']} | {l['message'][:160]}"
          for l in lignes_hypotheses_bfcc(n2)],
        *[f"{l['libelle']} : {l['statut']} | {l['message'][:160]}"
          for l in lignes_hypotheses_bootstrap(n2)],
        "",
        "=== RÉSULTATS ===",
        # ⚠️ LE FAUX ZÉRO, TROISIÈME FOYER (lot C3a). Ces quatre lignes
        # lisaient `bf.get('reserve_totale')` sans garde : une méthode non
        # calculable ressortait à « 0 € ». Le référentiel rend None, et le
        # motif vient de N4.
        *[f"{libelle(_m)} : "
          f"{_f(reserve(n3, _m)) if reserve(n3, _m) is not None else motif_exclusion(n4, _m)}"
          f" | poids={_pct(pw.get(_m, 0)*100)}"
          for _m in ORDRE_AFFICHAGE],
        f"Clark LDF : {_f(clark.get('reserve_be_clark'))} | courbe={clark.get('courbe_choisie', '—')} | AIC={clark.get('aic_optimal', '—')}"
        + ('' if (clark.get('structure_monotone') or {}).get('compatible', True)
           else ' | NON PUBLIEE : structure incompatible (facteur de developpement < 1)'),
        f"BEST ESTIMATE S2 : {_f(BE)} | CV={_pct(CV)}",
        "",
        "=== INCERTITUDE ===",
        f"σ Mack={_f(SIG)} | P75={_f(P75)} | P90={_f(P90)} | P99.5={_f(P99)}",
        "",
        "=== SCR ===",
        f"SCR={_f(SCP)} | Ratio SCR/BE={_pct(SCR)}",
        "",
        "=== BACK-TESTING ===",
        f"Statut={bt.get('statut', '—')} | Score={bt.get('score_qualite', '—')}/100",
        f"N-1: {bt.get('n_rouge_n1', 0)} rouge / {bt.get('n_ambre_n1', 0)} ambre | N-2: {bt.get('n_rouge_n2', 0)} rouge / {bt.get('n_ambre_n2', 0)} ambre",
        f"Message: {str(bt.get('message', ''))[:300]}",
        "",
        "=== EFFETS CALENDAIRE (GLM POISSON APC) ===",
        # ⚠️ LE PROMPT NE DIT PLUS << non significatif >> QUAND LE TEST N'A PAS
        # TOURNE. `cal_significatif` faux couvrait DEUX cas -- test negatif et
        # test absent -- et le modele recevait la meme phrase pour les deux.
        f"Statut={bz.get('statut', '—')} | Sig.={bz.get('n_effets_significatifs', 0)}/{bz.get('n_diagonales_evaluees', 0)}",
        f"Réserve GLM APC={_f(bz.get('reserve_apc', 0))} (= chain-ladder par MLE) | "
        f"test calendaire {_CAL_PROMPT[etat_calendaire(bz)]} "
        f"(F quasi-Poisson, p={bz.get('p_calendaire', '—')})",
        f"Diagonales anormales: {', '.join(bz.get('diagonales_anormales', [])) or 'Aucune'}",
        f"Recommandation: {bz.get('recommandation', '—')}",
        "",
        "Rédige le commentaire actuariel complet en 7 sections.",
    ]
    return '\n'.join(lines)


# =============================================================================
#  GÉNÉRATION NARRATION (3 NIVEAUX)
# =============================================================================

def _narration_claude_api(n2, n3, n4, lob_label, arrete) -> str:
    try:
        ctx = _construire_contexte(n2, n3, n4, lob_label, arrete)
        resp = frontiere_llm.appeler(
            modele=frontiere_llm.MODELE_ETABLI, max_tokens=10000,
            systeme=SYSTEM_PROMPT,
            messages=[{'role': 'user', 'content': ctx}],
            cle=frontiere_llm.cle_api_ou_secrets(),
        )
        return frontiere_llm.texte_des_blocs(resp)
    except Exception as e:
        logger.warning(f'Narration non generee : {e}')
        raise

def _narration_templates(n4, commentaire) -> str:
    return _clean(commentaire) or _clean(n4.get('jugement', ''))


# =============================================================================
#  VERROU C2 — AUCUN NOMBRE PUBLIE QUI NE SOIT DANS LA CHARGE UTILE
# =============================================================================
#
#  ⚠️⚠️ CE VERROU PROUVE LA PROVENANCE, JAMAIS LA JUSTESSE. Il verifie qu'un
#  nombre imprime dans la narration EXISTE dans ce qui a ete transmis -- rien
#  de plus. Un modele qui INVERSERAIT deux valeurs justes passerait sans un
#  mot : les deux nombres sont dans la charge utile. Il attrape ce qui est
#  INVENTE ou PERIME, pas ce qui est mal attribue. Le dire ici plutot que de
#  laisser croire a une garantie qu'il n'offre pas.
#
#  ⚠️ LA ZONE FRANCHE NE COUVRE QUE DES FORMES DE REFERENCE, JAMAIS UN SEUIL
#  NUMERIQUE. Exempter << 5 annees >> ou << 3 observations >> rouvrirait
#  exactement ce que ce chantier a passe une semaine a fermer : ce sont des
#  affirmations sur le portefeuille. Seuls sortent les identifiants qui ne
#  designent aucune grandeur -- section, hypothese, annee de publication.
#
#  ⚠️ APPARIEMENT EXACT, ET LES ARRONDIS REMONTENT COMME ORPHELINS. Un poids
#  publie << 17 % >> pour une charge utile a 17,3 % est une perte de precision
#  SILENCIEUSE. Mieux vaut que la charge porte la valeur affichee.
#
#  ⚠️ TAUX MESURE : 16 % d'orphelins (45 nombres, 7 orphelins) sur la narration
#  DETERMINISTE (`n4['jugement']`) -- PAS sur du LLM. Aucune cle API n'etait
#  disponible et aucune narration Claude n'est archivee dans le depot. Ce
#  chiffre vaut comme ORDRE DE GRANDEUR ; personne ne doit le lire comme une
#  mesure du chemin LLM. Sur ces 7 : 1 seul faux positif reel, 2 arrondis,
#  4 VRAIS defauts -- dont les trois provisions du jugement (`fcfb3d3`).
#
#  ⚠️ IL JOURNALISE, IL NE LEVE PAS. Faire tomber un rapport pour un defaut de
#  narration reproduirait le defaut que le lot A a ferme : un renderer qui
#  plantait sur un garde-fou bien pose. Le refus vit dans les TESTS.

#: Formes qui ne designent AUCUNE grandeur -- elles sortent du controle.
#: ⚠️ MESUREES sur la narration reelle, pas devinees : section (<< §5 >>),
#: hypothese (<< H2 >>, << CLM-H1 >>, << BFCC-H4 >>), annee de publication
#: d'une methode (<< Mack 1993 >>, << Clark 2003 >>), article et reglement.
_REFERENCES_NON_CHIFFREES = re.compile(
    r'§\s*\d+'                                  # section
    r'|\b(?:CLM|BFCC|BOOT|MCL)-H\d+\b'          # hypotheses nommees
    r'|\bH\d\b'                                 # H1..H4
    # ⚠️ LA PARENTHESE EST OBLIGATOIRE DANS CE MOTIF, ET C'EST MESURE : la
    # narration ecrit AUSSI BIEN << Mack 1993 >> que << Mack (1993) >>. La
    # premiere version ne couvrait que la forme sans parenthese, et 1993
    # ressortait orphelin -- le SEUL faux positif reel des sept mesures.
    r'|\b(?:Mack|Clark|Renshaw|Verrall|Quarg|Benktander|B&Z|Bornhuetter)'
    r'\s+(?:&\s+\w+\s+)?\(?(?:19|20)\d\d\)?'    # methode + annee, () optionnelles
    r'|Art(?:icle)?\.?\s*\d+'                   # Art. 77, Art. 115
    r'|\b(?:19|20)\d\d/\d+\b'                   # reglement 2015/35
    r'|\bQIS\s*\d+\b|\bTP\.\d[\d.]*',           # QIS5 TP.5.26
    re.IGNORECASE)

#: Un nombre : chiffres, separateurs de milliers (virgule, point, espace fine
#: ou insecable), decimale optionnelle.
#: ⚠️ CE MOTIF EST LA PIECE CRITIQUE, ET IL M'A DEJA TROMPE. Une premiere
#: version ne reconnaissait pas la virgule comme separateur : elle scindait
#: << 8,057,830 >> en TROIS nombres, tous orphelins, et annoncait 42 % de faux
#: positifs. La narration n'inventait rien -- le detecteur fabriquait des
#: orphelins. Il est teste POUR LUI-MEME (classe dediee) avant d'etre cru.
_NOMBRE = re.compile(r'\d[\d.,    ]*\d|\d')


def _cle_nombre(brut: str) -> str:
    """La forme canonique d'un nombre, pour comparer sans le formatage.

    ⚠️ LE POINT DELICAT EST LE POINT : separateur de milliers dans
    << 1.234.567 >>, separateur DECIMAL dans << 11.0 >>. On ne tranche pas au
    jugé -- un point suivi d'exactement trois chiffres ET d'une frontiere est
    un separateur de milliers ; sinon c'est une decimale. Meme regle pour la
    virgule, dans l'autre sens.
    """
    x = re.sub(r'[    ]', '', brut.strip())
    x = re.sub(r'(?<=\d)[.,](?=\d{3}(?:\D|$))', '', x)   # milliers
    x = x.replace(',', '.').rstrip('.')                  # decimale unifiee
    if '.' in x:
        entier, dec = x.split('.', 1)
        dec = dec.rstrip('0')
        x = f'{entier.lstrip("0") or "0"}.{dec}' if dec else (entier.lstrip('0') or '0')
        return x
    return x.lstrip('0') or '0'


def nombres_publies(texte: str) -> list:
    """Les nombres d'un texte, hors formes de reference. Ordre preserve."""
    if not texte:
        return []
    return _NOMBRE.findall(_REFERENCES_NON_CHIFFREES.sub(' ', texte))


def orphelins_narration(narration: str, charge_utile: str) -> list:
    """Les nombres PUBLIES qui ne figurent pas dans la CHARGE UTILE.

    ⚠️ RIEN N'EST APPARIE PAR IDENTIFIANT : on compare des VALEURS. Un
    identifiant de correspondance obligerait le modele a le reproduire, et
    ferait echouer le controle sur sa mise en forme plutot que sur son fond.

    Rend la liste des formes brutes orphelines, dans l'ordre du texte.
    """
    connus = {_cle_nombre(n) for n in nombres_publies(charge_utile)}
    return [n for n in nombres_publies(narration)
            if _cle_nombre(n) not in connus]

def _generer_narration(n2, n3, n4, commentaire, lob_label, arrete) -> Tuple[str, str]:
    try:
        txt = _narration_claude_api(n2, n3, n4, lob_label, arrete)
        if txt:
            return txt, 'claude_api'
    except Exception:
        pass
    try:
        txt = _narration_templates(n4, commentaire)
        if txt:
            return txt, 'templates'
    except Exception:
        pass
    return '', 'aucune'


#: Le libellé d'origine de la narration, par source. UNE SEULE TABLE.
#:
#: ⚠️ IL Y EN AVAIT DEUX, ET ELLES DIVERGEAIENT SUR DEUX POINTS. L'HTML
#: écrivait « ✦ ActuarIA Intelligence », le Word « ✦ Narration générée par
#: ActuarIA Intelligence » — divergence de forme. Et surtout, le Word était
#: gardé par `if source == 'claude_api'` : le chemin `templates`, LE SEUL qui
#: s'exécute sans clé API, n'y publiait ni origine ni engagement.
_LIBELLE_SOURCE = {
    'claude_api': '✦ ActuarIA Intelligence',
    'templates':  '📝 Mode standard',
    'aucune':     '',
}


def badge_narration(source: str) -> str:
    """Le libellé d'origine suivi de la phrase d'engagement — ou rien.

    ⚠️ CE QUE PORTE L'HTML, LE WORD LE PORTE. C'est le format Word que
    l'actuaire envoie au commissaire aux comptes ; un document qui n'y dit pas
    ce qui a produit sa section 7, ni ce qu'elle engage, n'est pas opposable.

    ⚠️ RIEN QUAND IL N'Y A RIEN : `avec_engagement('')` rend `''`, donc une
    source inconnue ou `aucune` ne fait apparaître aucun badge — un rapport
    sans narration ne doit engager personne.
    """
    return traitement_ia.avec_engagement(_LIBELLE_SOURCE.get(source, ''))


# =============================================================================
#  RENDU MARKDOWN → HTML (style premium)
# =============================================================================

#: Les six formes du FAUX-TITRE que le modele ajoute avant le §1 malgre la
#: consigne. ⚠️ ELLES NE CHERCHENT QU'UNE SEULE CHOSE -- l'en-tete de document.
#: C'est pourquoi elles ne s'appliquent QUE dans le preambule (cf. plus bas) :
#: affinees ou non, elles restent des motifs de DEBUT DE LIGNE, et un motif de
#: debut de ligne applique au corps du rapport juge une phrase entiere sur ses
#: premiers mots. Mesure : << Ce rapport sera transmis a l'ACPR avec une
#: reserve de 12 M EUR >>, ecrite en §5, etait SUPPRIMEE -- le chiffre avec.
_FAUX_TITRE = (
    r'^#+\s+RAPPORT',
    r'^Arrêté\s+(au|Q\d)',
    r'^Document\s+destin',
    r'^Commentaire\s+destin',
    r'^Reporting\s+Solvabilit',
    r'^Ce\s+rapport.*ACPR',
)

#: ⚠️⚠️ CES DEUX MOTIFS SONT REPRIS DE `core.narration` (lot T6) -- ET A7 DOIT
#: L'Y REJOINDRE. Les transposer ici cree la ONZIEME implementation de markdown
#: d'un depot qui en comptait dix, et `core.narration` existe PRECISEMENT pour
#: supprimer cette divergence : il couvre neuf formes la ou A7 en couvre six,
#: et il analyse UNE fois pour le HTML et le Word.
#:
#: Le choix de transposer plutot que de migrer est DELIBERE et arbitre : ce lot
#: ferme une PERTE DE CONTENU mesuree, il ne migre pas trois directions (A7 +
#: deux services Sante-Prevoyance consomment ces fonctions). Une migration
#: merite sa propre mesure d'impact -- le rendu HTML des trois changerait
#: (neuf formes au lieu de six, <h5>/<blockquote>/<hr> nouveaux).
#:
#: ⚠️ QUI LIT CECI ET TOUCHE A LA CONVERSION : va voir `core.narration` AVANT
#: d'ajouter une regle ici. La bonne fin de cette histoire est qu'A7 appelle
#: `core.narration` et que ce bloc disparaisse.
_TABLEAU_MD = re.compile(r'^\s*\|.*\|\s*$')
_SEPARATRICE_MD = re.compile(r'^\s*\|?[\s:|-]{5,}\|?\s*$')   # |---|---|


def _sans_faux_titre(preambule: str) -> str:
    """Le preambule debarrasse de ses lignes de faux-titre -- le RESTE SURVIT.

    ⚠️ AVANT CE LOT, LE PREAMBULE ENTIER ETAIT JETE (`txt[premier_s.start():]`).
    Une ouverture legitime -- un paragraphe de contexte redige avant le §1 --
    disparaissait avec le faux-titre. On FILTRE desormais, on ne tronque plus.
    """
    gardees = [
        ligne for ligne in preambule.split(chr(10))
        if not any(re.match(p, ligne.strip(), re.IGNORECASE)
                   for p in _FAUX_TITRE)
    ]
    return chr(10).join(gardees)


def _tableaux_en_texte(txt: str) -> str:
    """Les lignes de tableau markdown, CONTENU CONSERVE, pipes retires.

    ⚠️ ELLES ETAIENT PUREMENT ET SIMPLEMENT EFFACEES. Mesure : une narration de
    81 caracteres tombait a 25, et << 7 746 000 EUR >> -- le Best Estimate --
    disparaissait du commentaire remis au CAC. Le prompt interdit les tableaux
    (regle 0), mais un modele en produit malgre tout : la consigne se relache,
    la conversion non.

    Regle reprise de `core.narration` : les cellules deviennent
    << cellule - cellule >>, et seule la ligne SEPARATRICE (|---|---|) part,
    car elle ne porte aucun contenu.
    """
    sortie = []
    for ligne in txt.split(chr(10)):
        if not _TABLEAU_MD.match(ligne):
            sortie.append(ligne)
            continue
        if _SEPARATRICE_MD.match(ligne):
            continue
        cellules = [c.strip() for c in ligne.strip().strip('|').split('|')]
        sortie.append(' · '.join(c for c in cellules if c))
    return chr(10).join(sortie)


def _nettoyer_narration(texte: str) -> str:
    """Retire le faux-titre du preambule, sans rien detruire du corps.

    ⚠️ CETTE FONCTION EFFACAIT DU VRAI, ET C'ETAIT MESURE. Trois pertes,
    toutes fermees ici :
      · les lignes de tableau markdown, effacees avec leurs chiffres ;
      · les six motifs de faux-titre appliques a TOUT le texte, qui
        supprimaient des phrases legitimes du corps du rapport ;
      · le preambule tronque en bloc, ouverture legitime comprise.

    ⚠️ ELLE N'EST PAS PRIVEE A A7, malgre le tiret bas : deux services
    Sante-Prevoyance l'importent (`sp_rapport_prevoyance`, `sp_rapport_sante`).
    Toute modification porte sur TROIS livrables -- d'ou la gate Vie+SP.
    """
    if not texte:
        return ''
    txt = texte.strip()

    # ── Le preambule : filtre, jamais tronque ────────────────────────────────
    # ⚠️ LA PORTEE EST LE REMEDE, PAS LA PRECISION DES MOTIFS. Le faux-titre ne
    # peut exister QU'AVANT le premier §. Au-dela, c'est du corps de rapport et
    # cela ne se filtre pas.
    # ⚠️ C'EST LA PRESENCE D'UN §, PAS SA POSITION, QUI DELIMITE LE PREAMBULE.
    # L'ancien code exigeait `start() > 10` -- un seuil herite dont l'effet
    # etait l'inverse du but : sur une narration BIEN FORMEE commencant par
    # << §1 >> (position 0), la condition etait fausse et le texte ENTIER
    # partait dans la branche sans §, donc filtre de bout en bout. C'est
    # exactement ce qui detruisait la phrase du §5. Un preambule vide est un
    # preambule : on le filtre, il rend '', et le corps passe intact.
    premier_s = re.search(r'§\s*\d+', txt)
    if premier_s:
        preambule = _sans_faux_titre(txt[:premier_s.start()]).strip()
        corps = txt[premier_s.start():]
        txt = (preambule + chr(10) * 2 + corps) if preambule else corps
    else:
        # Aucun § : rien ne delimite le corps, le texte entier est un
        # preambule potentiel -- seul cas ou le filtre porte partout.
        txt = _sans_faux_titre(txt)

    txt = _tableaux_en_texte(txt)

    # Normaliser les sauts de ligne
    txt = re.sub(r'\n{3,}', '\n\n', txt)

    return txt.strip()


def _md_to_html(texte: str) -> str:
    """Convertit le Markdown de Claude en HTML avec les classes premium."""
    if not texte:
        return '<p class="comm-p" style="color:#8A9BB0;font-style:italic;">Narration non disponible.</p>'

    # Nettoyer le header générique avant conversion
    texte = _nettoyer_narration(texte)
    if not texte:
        return '<p class="comm-p" style="color:#8A9BB0;font-style:italic;">Narration non disponible.</p>'

    txt = texte.strip()

    # §N — TITRE → comm-section-title
    def _section(m):
        t = m.group(1).strip()
        return '<div class="comm-section-title">' + t + '</div>'
    txt = re.sub(r'(§\d+\s*[—\-–]\s*[^\n]+)', _section, txt)

    # ### Sous-titre → comm-h4
    txt = re.sub(r'^###\s+(.+)$',
        lambda m: '<div class="comm-h4">' + m.group(1).strip() + '</div>',
        txt, flags=re.MULTILINE)

    # ## Titre → comm-h4
    txt = re.sub(r'^##\s+(.+)$',
        lambda m: '<div class="comm-h4">' + m.group(1).strip() + '</div>',
        txt, flags=re.MULTILINE)

    # --- → comm-divider
    txt = re.sub(r'^---+$', '<hr class="comm-divider">', txt, flags=re.MULTILINE)

    # **gras** → strong (y compris multi-lignes courts)
    txt = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', txt, flags=re.DOTALL)

    # *italique* → em
    txt = re.sub(r'\*(.+?)\*', r'<em>\1</em>', txt)

    # Puces - → comm-ul li
    txt = re.sub(r'^[-•]\s+(.+)$', r'<li>\1</li>', txt, flags=re.MULTILINE)
    txt = re.sub(
        r'(<li>.*?</li>\n?)+',
        lambda m: '<ul class="comm-ul">' + m.group(0) + '</ul>',
        txt, flags=re.DOTALL
    )

    # Paragraphes
    blocs = re.split(r'\n{2,}', txt)
    result = ''
    for bloc in blocs:
        bloc = bloc.strip()
        if not bloc:
            continue
        if bloc.startswith('<'):
            result += bloc + '\n'
        else:
            clean = bloc.replace('\n', ' ').strip()
            result += '<p class="comm-p">' + clean + '</p>\n'
    return result or '<p class="comm-p">' + texte + '</p>'


# =============================================================================
#  CSS PREMIUM
# =============================================================================

def _css() -> str:
    return """
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,700;0,900;1,400&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');

:root {
  --navy:        #0B1E3D;
  --navy-mid:    #132844;
  --navy-light:  #1E3A5F;
  --gold:        #C9A84C;
  --gold-light:  #E2C97E;
  --gold-pale:   rgba(201,168,76,0.12);
  --cyan:        #4A8FD4;
  --cyan-pale:   rgba(74,143,212,0.10);
  --rouge:       #C0392B;
  --rouge-pale:  rgba(192,57,43,0.08);
  --orange:      #E67E22;
  --vert:        #1E8449;
  --slate:       #8A9BB0;
  --slate-light: #B8C5D3;
  --bg:          #F5F7FA;
  --white:       #FFFFFF;
  --text:        #1C2B3A;
  --text-mid:    #3D5166;
  --border:      #DDE4EE;
}

* { box-sizing: border-box; margin: 0; padding: 0; }

body {
  font-family: 'Inter', sans-serif;
  font-size: 10pt;
  color: var(--text);
  background: var(--bg);
  line-height: 1.7;
  -webkit-font-smoothing: antialiased;
}

.rapport-container {
  max-width: 960px;
  margin: 0 auto;
  background: var(--white);
  box-shadow: 0 4px 40px rgba(11,30,61,0.14);
}

/* ── PAGE DE GARDE ── */
.page-garde {
  position: relative;
  min-height: 680px;
  display: flex;
  flex-direction: column;
  overflow: hidden;
  background: var(--navy);
  page-break-after: always;
}

.garde-bg { position: absolute; inset: 0; overflow: hidden; pointer-events: none; }

.garde-bg::before {
  content: '';
  position: absolute;
  top: -120px; right: -80px;
  width: 520px; height: 520px;
  border: 1px solid rgba(201,168,76,0.12);
  border-radius: 50%;
  transform: rotate(15deg);
}

.garde-bg::after {
  content: '';
  position: absolute;
  top: -60px; right: -20px;
  width: 360px; height: 360px;
  border: 1px solid rgba(201,168,76,0.20);
  border-radius: 50%;
}

.garde-dots {
  position: absolute;
  top: 0; right: 0;
  width: 320px; height: 100%;
  background-image: radial-gradient(circle, rgba(201,168,76,0.25) 1px, transparent 1px);
  background-size: 22px 22px;
  mask-image: linear-gradient(to left, rgba(0,0,0,0.7) 0%, transparent 100%);
  -webkit-mask-image: linear-gradient(to left, rgba(0,0,0,0.7) 0%, transparent 100%);
}

.garde-diagonal {
  position: absolute;
  top: 0; left: 0; right: 0; bottom: 0;
  background: linear-gradient(135deg, transparent 0%, transparent 62%, rgba(201,168,76,0.06) 62%, rgba(201,168,76,0.06) 100%);
}

.garde-accent-line {
  position: absolute;
  left: 52px; top: 0; bottom: 0;
  width: 2px;
  background: linear-gradient(to bottom, transparent 0%, var(--gold) 20%, var(--gold) 80%, transparent 100%);
  opacity: 0.5;
}

.garde-inner {
  position: relative;
  z-index: 2;
  display: flex;
  flex-direction: column;
  height: 100%;
  padding: 0;
}

.garde-header {
  display: flex;
  justify-content: space-between;
  align-items: flex-start;
  padding: 36px 56px 0;
}

.garde-logo-wrap img { height: 52px; }

.garde-badges {
  display: flex;
  flex-direction: column;
  align-items: flex-end;
  gap: 8px;
}

.badge-confidentiel {
  font-size: 7.5pt;
  font-weight: 700;
  color: var(--rouge);
  letter-spacing: 2.5px;
  border: 1.5px solid var(--rouge);
  padding: 5px 12px;
  text-transform: uppercase;
}

.garde-hero {
  flex: 1;
  display: flex;
  flex-direction: column;
  justify-content: center;
  padding: 48px 56px 0 80px;
}

.garde-eyebrow {
  font-size: 8pt;
  font-weight: 600;
  color: var(--gold);
  letter-spacing: 4px;
  text-transform: uppercase;
  margin-bottom: 20px;
}

.garde-titre {
  font-family: 'Playfair Display', serif;
  font-size: 52pt;
  font-weight: 900;
  color: var(--white);
  line-height: 1.05;
  letter-spacing: -1.5px;
  margin-bottom: 8px;
}

.garde-titre em {
  font-style: italic;
  color: var(--gold);
}

.garde-subtitle {
  font-family: 'Playfair Display', serif;
  font-size: 18pt;
  font-weight: 400;
  font-style: italic;
  color: var(--slate-light);
  margin-bottom: 32px;
}

.garde-sep {
  display: flex;
  align-items: center;
  gap: 16px;
  margin-bottom: 32px;
}

.garde-sep-line { height: 1px; width: 80px; background: var(--gold); opacity: 0.6; }

.garde-sep-diamond {
  width: 8px; height: 8px;
  background: var(--gold);
  transform: rotate(45deg);
  flex-shrink: 0;
}

.garde-statut {
  display: inline-flex;
  align-items: center;
  gap: 12px;
  padding: 10px 24px;
  margin-bottom: 0;
}

.garde-statut-rouge { background: rgba(192,57,43,0.15); border: 1px solid rgba(192,57,43,0.5); }
.garde-statut-ambre { background: rgba(230,126,34,0.15); border: 1px solid rgba(230,126,34,0.5); }
.garde-statut-vert  { background: rgba(30,132,73,0.15);  border: 1px solid rgba(30,132,73,0.5); }

.statut-dot {
  width: 10px; height: 10px;
  border-radius: 50%;
  box-shadow: 0 0 8px rgba(192,57,43,0.7);
  animation: pulse 2s ease-in-out infinite;
}

.statut-dot-rouge { background: var(--rouge); box-shadow: 0 0 8px rgba(192,57,43,0.7); }
.statut-dot-ambre { background: var(--orange); box-shadow: 0 0 8px rgba(230,126,34,0.7); }
.statut-dot-vert  { background: var(--vert);  box-shadow: 0 0 8px rgba(30,132,73,0.7); }

@keyframes pulse {
  0%, 100% { opacity: 1; transform: scale(1); }
  50% { opacity: 0.7; transform: scale(1.2); }
}

.statut-label-rouge { font-size: 8pt; font-weight: 700; color: var(--rouge); letter-spacing: 3px; text-transform: uppercase; }
.statut-label-ambre { font-size: 8pt; font-weight: 700; color: var(--orange); letter-spacing: 3px; text-transform: uppercase; }
.statut-label-vert  { font-size: 8pt; font-weight: 700; color: var(--vert);  letter-spacing: 3px; text-transform: uppercase; }

.garde-footer {
  border-top: 1px solid rgba(201,168,76,0.25);
  margin: 40px 56px 0;
  padding: 28px 0 40px;
}

.garde-kpis { display: grid; grid-template-columns: repeat(6, 1fr); gap: 0; }

.garde-kpi {
  padding: 0 20px 0 0;
  border-right: 1px solid rgba(255,255,255,0.08);
}

.garde-kpi:last-child { border-right: none; }

.kpi-label {
  font-size: 6.5pt;
  font-weight: 600;
  color: var(--slate);
  text-transform: uppercase;
  letter-spacing: 1.5px;
  margin-bottom: 6px;
}

.kpi-value {
  font-family: 'Playfair Display', serif;
  font-size: 12pt;
  font-weight: 700;
  color: var(--white);
  line-height: 1.2;
}

.kpi-value.highlight { color: var(--gold); font-size: 14pt; }
.kpi-sub { font-size: 7pt; color: var(--slate); margin-top: 3px; }

/* ── CORPS ── */
.rapport-body { padding: 0; }

.section-header {
  background: linear-gradient(to right, var(--navy) 0%, var(--navy-light) 100%);
  padding: 22px 56px;
  display: flex;
  align-items: center;
  gap: 20px;
}

.section-num {
  font-family: 'JetBrains Mono', monospace;
  font-size: 9pt;
  font-weight: 500;
  color: var(--gold);
  opacity: 0.8;
  min-width: 28px;
}

.section-titre {
  font-family: 'Playfair Display', serif;
  font-size: 16pt;
  font-weight: 700;
  color: var(--white);
}

.section-body { padding: 36px 56px 48px; }

.section-divider { height: 1px; background: var(--border); }

/* ── KPI CARDS ── */
.kpi-grid {
  display: grid;
  grid-template-columns: repeat(4, 1fr);
  gap: 12px;
  margin-bottom: 32px;
}

.kpi-card {
  background: var(--navy);
  border-radius: 6px;
  padding: 20px 18px;
  position: relative;
  overflow: hidden;
}

.kpi-card::before {
  content: '';
  position: absolute;
  top: 0; left: 0; right: 0;
  height: 3px;
  background: var(--gold);
}

.kpi-card-rouge::before { background: var(--rouge); }

.kpi-card-label { font-size: 7pt; font-weight: 600; color: var(--slate); text-transform: uppercase; letter-spacing: 1.5px; margin-bottom: 8px; }
.kpi-card-value { font-family: 'Playfair Display', serif; font-size: 18pt; font-weight: 700; color: var(--white); line-height: 1.1; margin-bottom: 6px; }
.kpi-card-value-rouge { color: var(--rouge) !important; }
.kpi-card-sub { font-size: 7.5pt; color: var(--gold); }

/* ── TABLEAUX PREMIUM ── */
.table-section-title {
  font-family: 'Playfair Display', serif;
  font-size: 10pt;
  font-weight: 600;
  color: var(--gold);
  margin: 28px 0 12px;
  padding-bottom: 6px;
  border-bottom: 1px solid var(--gold-pale);
}

table.premium { width: 100%; border-collapse: collapse; font-size: 9pt; }
table.premium thead tr { background: var(--navy); }
table.premium thead th { padding: 11px 14px; color: var(--white); font-weight: 600; font-size: 8pt; letter-spacing: 0.5px; text-align: left; }
table.premium thead th.right { text-align: right; }
table.premium thead th.center { text-align: center; }
table.premium tbody tr:nth-child(even) { background: #F8FAFB; }
table.premium tbody tr.highlight-gold { background: var(--gold-pale); font-weight: 700; }
table.premium tbody td { padding: 9px 14px; border-bottom: 1px solid var(--border); color: var(--text); }
table.premium tbody td.right { text-align: right; }
table.premium tbody td.center { text-align: center; }
table.premium tbody td.label { font-weight: 600; color: var(--navy); }
table.premium tbody td .mono { font-family: 'JetBrains Mono', monospace; font-size: 8.5pt; }

/* ── BADGES ── */
.badge { display: inline-flex; align-items: center; gap: 4px; font-size: 7.5pt; font-weight: 700; padding: 2px 8px; border-radius: 3px; letter-spacing: 0.5px; }
.badge-ok   { background: rgba(30,132,73,0.12);  color: var(--vert); }
.badge-warn { background: rgba(230,126,34,0.12); color: var(--orange); }
.badge-excl { background: rgba(138,155,176,0.15); color: var(--slate); }

/* ── HYPOTHÈSES ── */
.hyp-grid { display: flex; flex-direction: column; gap: 10px; }

.hyp-card {
  display: grid;
  grid-template-columns: 240px 1fr;
  gap: 0;
  border-radius: 6px;
  overflow: hidden;
  border: 1px solid var(--border);
}

.hyp-label { padding: 14px 18px; font-size: 8pt; font-weight: 700; }
.hyp-label .hyp-code { font-family: 'JetBrains Mono', monospace; font-size: 7.5pt; opacity: 0.9; }
.hyp-label .hyp-score { font-weight: 400; opacity: 0.8; margin-top: 4px; font-size: 7.5pt; }

.hyp-ok   .hyp-label { background: rgba(30,132,73,0.10);  color: var(--vert);   border-right: 3px solid var(--vert); }
.hyp-warn .hyp-label { background: rgba(230,126,34,0.10); color: var(--orange); border-right: 3px solid var(--orange); }

.hyp-text { padding: 14px 18px; font-size: 9pt; color: var(--text-mid); line-height: 1.6; background: #FAFBFC; }
/* La PUISSANCE se lit avec le verdict : même carte, registre visuel distinct
   pour qu'on ne la confonde pas avec le verdict lui-même. */
.hyp-puissance { padding: 10px 18px 14px 18px; font-size: 8.5pt; color: var(--text-mid);
                 line-height: 1.55; background: #FAFBFC; border-top: 1px dashed #D8DEE6;
                 font-style: italic; }

/* ── RECOMMANDATION BOX ── */
.recommandation-box {
  background: var(--navy);
  padding: 18px 24px;
  border-radius: 6px;
  margin-top: 20px;
  display: flex;
  align-items: flex-start;
  gap: 16px;
}

.reco-icon { font-size: 18pt; line-height: 1; flex-shrink: 0; color: var(--gold); }
.reco-label { font-size: 7pt; font-weight: 700; color: var(--gold); text-transform: uppercase; letter-spacing: 2px; margin-bottom: 4px; }
.reco-value { font-family: 'Playfair Display', serif; font-size: 13pt; font-weight: 700; color: var(--white); margin-bottom: 4px; }
.reco-text { font-size: 8.5pt; color: var(--slate); line-height: 1.5; }

/* ── EFFETS CALENDAIRE ── */
.bz-grid { display: flex; flex-direction: column; gap: 6px; }

.bz-item { display: flex; align-items: center; gap: 14px; padding: 10px 16px; border-radius: 4px; font-size: 9pt; }
.bz-warn { background: rgba(230,126,34,0.08); border-left: 3px solid var(--orange); }
.bz-info { background: var(--cyan-pale); border-left: 3px solid var(--cyan); }
.bz-ok   { background: rgba(30,132,73,0.06); border-left: 3px solid var(--vert); }

.bz-year { font-family: 'JetBrains Mono', monospace; font-size: 8.5pt; font-weight: 700; min-width: 42px; }
.bz-bar { flex: 1; height: 4px; background: var(--border); border-radius: 2px; overflow: hidden; }
.bz-fill { height: 100%; border-radius: 2px; }
.bz-pct { font-size: 8.5pt; font-weight: 600; min-width: 120px; text-align: right; }

/* ── COMMENTAIRE ACTUARIEL ── */
.commentaire-wrap { border: 1px solid var(--border); border-radius: 6px; overflow: hidden; }

.commentaire-header {
  background: var(--navy);
  padding: 14px 20px;
  display: flex;
  align-items: center;
  justify-content: space-between;
}

.commentaire-header-title { font-size: 8pt; font-weight: 600; color: var(--gold); letter-spacing: 2px; text-transform: uppercase; }
.commentaire-ai-badge { font-size: 7pt; color: var(--slate); font-family: 'JetBrains Mono', monospace; }
.commentaire-body { padding: 28px 32px; background: #FAFBFC; }

.comm-section-title {
  font-family: 'Playfair Display', serif;
  font-size: 11pt;
  font-weight: 700;
  color: var(--navy);
  border-bottom: 1.5px solid var(--gold);
  padding-bottom: 6px;
  margin: 24px 0 12px;
}

.comm-section-title:first-child { margin-top: 0; }
.comm-h4 { font-size: 9.5pt; font-weight: 700; color: var(--navy); margin: 16px 0 8px; }
.comm-p { font-size: 9.5pt; line-height: 1.85; color: var(--text-mid); margin-bottom: 12px; }
.comm-p strong { color: var(--navy); font-weight: 700; }

.comm-ul { margin: 10px 0 12px 20px; list-style: none; }
.comm-ul li { position: relative; padding-left: 16px; font-size: 9.5pt; line-height: 1.7; color: var(--text-mid); margin-bottom: 6px; }
.comm-ul li::before { content: '—'; position: absolute; left: 0; color: var(--gold); font-weight: 700; }
.comm-ul li strong { color: var(--navy); }

.comm-divider { border: none; border-top: 1px solid var(--border); margin: 20px 0; }
.comm-footer { font-size: 7.5pt; color: var(--slate); font-style: italic; text-align: right; border-top: 1px solid var(--border); padding-top: 10px; margin-top: 8px; }

/* ── BACK-TESTING ── */
.bt-summary { display: flex; gap: 12px; margin-bottom: 20px; flex-wrap: wrap; }
.bt-card { padding: 14px 20px; border-radius: 6px; text-align: center; min-width: 120px; }
.bt-card-rouge { background: var(--rouge); color: var(--white); }
.bt-card-ambre { background: var(--orange); color: var(--white); }
.bt-card-vert  { background: var(--vert);  color: var(--white); }
.bt-card-navy  { background: var(--navy);  color: var(--white); }
.bt-card-label { font-size: 7.5pt; opacity: 0.8; margin-bottom: 4px; }
.bt-card-value { font-family: 'Playfair Display', serif; font-size: 15pt; font-weight: 700; }
.bt-card-sub   { font-size: 7.5pt; color: var(--gold); margin-top: 2px; }

.bt-note {
  background: var(--rouge-pale);
  border-left: 3px solid var(--rouge);
  padding: 12px 16px;
  border-radius: 0 4px 4px 0;
  font-size: 9pt;
  color: var(--text-mid);
  line-height: 1.6;
  margin-top: 12px;
}

/* ── PIED DE PAGE ── */
.pied-de-page {
  background: var(--navy);
  padding: 20px 56px;
  display: flex;
  justify-content: space-between;
  align-items: center;
  border-top: 2px solid var(--gold);
}

.pied-logo img { height: 26px; }
.pied-meta { text-align: right; font-size: 7.5pt; color: var(--slate); line-height: 1.7; }
.pied-meta .confidentiel-footer { color: var(--rouge); font-weight: 700; letter-spacing: 1px; }

@media print {
  /* ⚠️ LA RÈGLE QUI RÉPARE LES COULEURS QUI DISPARAISSAIENT.
     Un navigateur SUPPRIME par défaut les fonds à l'impression — c'est une
     économie d'encre héritée du papier. Le rapport pose plus de 180 fonds
     colorés (69 `background`, 5 dégradés, 117 `rgba`) et le signal RAG en
     dépend : sans cette déclaration, un statut ROUGE s'imprime sur fond
     blanc, exactement comme un VERT. `print-color-adjust: exact` demande au
     moteur de rendre les fonds tels quels. Le préfixe `-webkit-` reste
     nécessaire : Chrome et Edge, qui font la conversion PDF, ne lisent
     encore que lui dans certaines versions. */
  *, *::before, *::after {
    -webkit-print-color-adjust: exact !important;
    print-color-adjust: exact !important;
  }

  /* ⚠️ `@page { margin: 0 }` ÉTAIT PIRE QUE PAS DE RÈGLE DU TOUT. Sans marge,
     le contenu court jusqu'au bord physique de la feuille : la plupart des
     imprimantes et des convertisseurs rognent cette zone, et un tableau y
     perd sa dernière colonne. La page de garde, elle, est conçue pour saigner
     — d'où `:first` qui lui rend sa marge nulle sans l'imposer au reste. */
  @page { size: A4; margin: 14mm 12mm 16mm 12mm; }
  @page :first { margin: 0; }

  body { background: white; }
  .rapport-container { box-shadow: none; max-width: none; margin: 0; }
  .page-garde { page-break-after: always; break-after: page; }

  /* CE QUI NE DOIT PAS ÊTRE COUPÉ EN DEUX PAR UNE FIN DE PAGE.
     Un tableau scindé perd son en-tête, une figure scindée ne veut plus rien
     dire, et une carte d'hypothèse coupée sépare le verdict de sa
     justification. `break-inside` est la propriété moderne, `page-break-inside`
     son ancêtre : les deux sont posées parce que les moteurs d'impression
     n'ont pas tous migré. */
  table.premium, .hyp-card, .kpi-grid, .hyp-grid, .bz-grid, .garde-kpis,
  .plotly-graph-div {
    break-inside: avoid;
    page-break-inside: avoid;
  }
  tr { break-inside: avoid; page-break-inside: avoid; }

  /* Un tableau qui franchit une page réimprime son en-tête. Sans cela, la
     suite d'un tableau de dix-sept hypothèses arrive sans ses colonnes. */
  thead { display: table-header-group; }

  /* UN TITRE NE RESTE JAMAIS SEUL EN BAS DE PAGE, séparé de ce qu'il
     annonce — c'est la faute de mise en page la plus visible d'un rapport. */
  .section-header, .table-section-title, .comm-section-title {
    break-after: avoid;
    page-break-after: avoid;
  }
  .section-divider { break-before: avoid; page-break-before: avoid; }

  /* Pas de ligne isolée en haut ou en bas de page dans la prose. */
  p, .comm-p, .hyp-text { orphans: 3; widows: 3; }

  /* Une figure Plotly a une largeur en pixels fixée par le navigateur ; sur
     une feuille A4 elle déborderait de la zone imprimable. */
  .plotly-graph-div { max-width: 100% !important; }
}
</style>"""


# =============================================================================
#  CONSTRUCTION DES BLOCS HTML
# =============================================================================

#: Quantile normal unilatéral à 90 % — même approximation normale que l'IC 95 %
#: de Clark (ultime ± 1,96·se). Publier le P90 sur une autre base que l'IC
#: reviendrait à faire cohabiter deux lois dans le même bloc.
Z90 = 1.281552


def _bloc_clark_incertitude(n2: Dict, n3: Dict) -> str:
    """Intervalle de prédiction de Clark (2003), décomposé année par année.

    Ce bloc existe parce que le correctif de couverture — l'IC passait de 30,8 %
    à 97,0 % une fois le facteur σ² appliqué à la Hessienne — produisait des
    clés (`ic_95`, `sigma2`, `se_parametre`, `se_processus`, `se_totale`,
    `se_reserve_totale`, `cv_reserve`) que RIEN n'affichait. Un intervalle
    réparé mais invisible ne vaut pas mieux qu'un intervalle faux.
    """
    clark = (n3 or {}).get('clark') or {}
    if not clark.get('disponible'):
        return ''

    struct = clark.get('structure_monotone') or {}
    if not struct.get('compatible', True):
        # La réserve n'est pas publiée : un intervalle autour d'elle ne le sera
        # pas davantage. On dit pourquoi plutôt que de laisser un vide.
        return (
            '<p style="font-size:8.5pt;color:var(--slate);font-style:italic;">'
            'Intervalle de prédiction non publié — ' + _s(struct.get('message'))
            + '</p>')

    ic     = clark.get('ic_95') or []
    se_par = clark.get('se_parametre') or []
    se_pro = clark.get('se_processus') or []
    se_tot = clark.get('se_totale') or []
    ibnr   = clark.get('ibnr_brut_par_annee') or []
    if not ic or ic[0][0] is None or not se_tot or se_tot[0] is None:
        return (
            '<p style="font-size:8.5pt;color:var(--slate);font-style:italic;">'
            'Intervalle de prédiction non calculé — la Hessienne de la '
            'log-vraisemblance n\'a pas pu être inversée sur ce triangle '
            '(convergence du MLE insuffisante).</p>')

    an0 = n2.get('annee_debut') if isinstance(n2, dict) else None
    try:
        an0 = int(an0)
    except (TypeError, ValueError):
        an0 = None

    lignes = ''
    for i in range(min(len(ic), len(se_tot), len(ibnr))):
        lo, hi = ic[i]
        etiq = str(an0 + i) if an0 is not None else f'Année {i + 1}'
        lignes += (
            '<tr><td class="label">' + etiq + '</td>'
            '<td class="right"><span class="mono">' + _f(ibnr[i]) + '</span></td>'
            '<td class="right"><span class="mono">' + _f(se_par[i] if i < len(se_par) else None) + '</span></td>'
            '<td class="right"><span class="mono">' + _f(se_pro[i] if i < len(se_pro) else None) + '</span></td>'
            '<td class="right"><span class="mono">' + _f(se_tot[i]) + '</span></td>'
            '<td class="center" style="font-size:8pt;"><span class="mono">'
            + _f(lo) + ' — ' + _f(hi) + '</span></td></tr>'
        )

    se_r = clark.get('se_reserve_totale')
    cv   = clark.get('cv_reserve')
    lignes += (
        '<tr class="highlight-gold"><td class="label" style="color:var(--navy);">'
        'TOTAL (réserve)</td>'
        '<td class="right" style="color:var(--navy);"><span class="mono">'
        + _f(clark.get('reserve_brute')) + '</span></td>'
        '<td class="center" colspan="2" style="font-size:8pt;color:var(--navy-light);">'
        'covariances comprises</td>'
        '<td class="right" style="color:var(--navy);"><span class="mono">'
        + _f(se_r) + '</span></td>'
        '<td class="center" style="color:var(--navy);">CV = '
        + (_pct(cv) if cv is not None else '—') + '</td></tr>'
    )

    s2  = clark.get('sigma2')
    df  = clark.get('df')
    phi = ((n3 or {}).get('bootstrap') or {}).get('phi')
    recoupement = ''
    if s2 and phi:
        rap = max(s2, phi) / min(s2, phi)
        recoupement = (
            ' Le Bootstrap ODP estime la même sur-dispersion à '
            + f'{phi:,.4f}'.replace(',', ' ')
            + f" — rapport ×{rap:.2f}. Deux ajustements différents du "
              "même paramètre : ils diffèrent légitimement, un écart d'un "
              "ordre de grandeur signalerait en revanche que l'un des deux "
              "ne mesure pas ce qu'il annonce."
        )

    return (
        '<table class="premium"><thead><tr>'
        '<th>Survenance</th><th class="right">IBNR</th>'
        '<th class="right">se paramètre</th><th class="right">se processus</th>'
        '<th class="right">se totale</th><th class="center">IC 95 %</th>'
        '</tr></thead><tbody>' + lignes + '</tbody></table>'
        '<p style="font-size:8pt;color:var(--slate);font-style:italic;margin-top:6px;">'
        'Clark (2003) décompose Var(R) en variance de PROCESSUS (σ²·R, aléa des '
        'paiements futurs) et variance de PARAMÈTRE (σ²·H⁻¹, incertitude sur la '
        'courbe et les ultimes). Le total agrégé fait intervenir les covariances : '
        'les ultimes partagent ω et θ, leurs erreurs sont corrélées. '
        + (f'σ² = {s2:,.4f}'.replace(',', ' ') if s2 else 'σ² non estimé')
        + (f' sur {df} degrés de liberté.' if df else '.')
        + recoupement +
        '</p>'
    )


def lignes_clark_rapport(n2: Dict, n3: Dict) -> List[Dict[str, str]]:
    """Les hypothèses de Clark, PRÉSENTÉES — pas gouvernées.

    ⚠️ AUCUN VERDICT N'EST CRÉÉ ICI, et c'est le point. Une investigation
    dédiée a mesuré qu'un test d'adéquation de la courbe de Clark — statistique
    de restriction emboîtée, calibrée sur vérité connue, parfaitement juste
    (5,0 % de fausses alarmes au seuil 5 %) et écrasamment puissante (statistique
    ×112 entre une vraie courbe de Clark et une cadence bimodale) — n'a AUCUNE
    relation avec l'erreur de réserve : ρ de Spearman = −0,001 sur 135 triangles
    (p = 0,993), erreur médiane 5,01 % contre 5,41 %. Publier un statut
    VALIDÉE / NON VALIDÉE là-dessus aurait recréé l'ancienne H4 : une alarme
    juste sur une grandeur qui ne commande rien.

    Ce que ces lignes rendent, ce sont donc des FAITS déjà calculés ailleurs :
    la précondition de forme (`structure_monotone`, n3), ce que le masque du MLE
    a réellement retiré (`increments_exclus_mle`, n3), et le verdict
    d'indépendance déjà rendu par CLM-H1 (n2), republié tel quel.
    """
    clark = (n3 or {}).get('clark') or {}
    if not clark.get('disponible'):
        return []

    lignes: List[Dict[str, str]] = []

    struct = clark.get('structure_monotone') or {}
    if struct.get('testable'):
        compat = bool(struct.get('compatible', True))
        lignes.append({
            'libelle': 'Clark — Représentabilité par une courbe monotone',
            'etat':    'Compatible' if compat else 'Incompatible',
            'ok':      'oui' if compat else 'non',
            'texte':   _s(struct.get('message')),
        })

    exc = clark.get('increments_exclus_mle') or {}
    if 'n_exclues' in exc:
        n_ex = int(exc.get('n_exclues', 0))
        n_ze = int(exc.get('zeros_conserves', 0))
        txt = (
            f"L'ajustement porte sur {clark.get('n_obs', '—')} cellules. "
            + (f"{n_ex} incrément(s) strictement négatif(s) retiré(s) "
               f"({exc.get('frac_exclue', 0):.1%} des cellules connues) : le "
               f"modèle de Poisson sur-dispersé n'est pas défini pour une "
               f"espérance négative."
               if n_ex else
               "Aucun incrément négatif : le modèle de Poisson sur-dispersé "
               "s'applique tel quel.")
            + (f" {n_ze} incrément(s) nul(s) CONSERVÉ(s) — le terme de "
               f"vraisemblance vaut alors −μ, parfaitement défini."
               if n_ze else "")
        )
        lignes.append({
            'libelle': 'Clark — Positivité des incréments ajustés',
            'etat':    'Aucune exclusion' if not n_ex else f'{n_ex} exclue(s)',
            'ok':      'oui' if not n_ex else 'non',
            'texte':   txt,
        })

    h1 = ((n2 or {}).get('clm') or {}).get('hypotheses', {}).get('CLM-H1') or {}
    if h1.get('statut'):
        lignes.append({
            'libelle': 'Clark — Indépendance des incréments (repris de CLM-H1)',
            'etat':    _s(h1.get('statut')),
            'ok':      'oui' if h1.get('statut') == 'VALIDÉE' else 'non',
            'texte':   ("Clark (2003) suppose les incréments indépendants, "
                        "exactement comme Chain Ladder. Verdict repris sans "
                        "recalcul : " + _s(h1.get('message'))),
        })

    return lignes


# ── LA RELECTURE ACTUARIELLE ────────────────────────────────────────────────
#
#  ⚠️ DEUX ETATS, ET PAS TROIS. Le rapport de tarification en porte trois
#  (validee / non enregistree / hors production), le troisieme derivant du
#  parametre `environnement` d'A6. A7 N'A AUCUNE NOTION D'ENVIRONNEMENT —
#  releve : les deux occurrences du mot dans l'agent et ici sont des
#  commentaires sur l'environnement de test et sur weasyprint. Transposer le
#  troisieme etat serait de la symetrie sans source : un etat qu'aucune
#  donnee ne peut produire.
#
#  ⚠️ ET L'ETAT NEGATIF EST ACTIF, C'EST TOUT LE POINT. A7 ne distinguait que
#  << signe >> et << rien >> : sans nom, le pied ne disait RIEN, et le lecteur
#  ne pouvait pas separer << relu par personne >> de << le champ n'a pas ete
#  transmis >>. C'est le silence que ce chantier ferme depuis le debut.

#  ⚠️ ET DEUX ETATS SE PORTENT PAR UN BOOLEEN, PAS PAR UN CODE. Le rapport de
#  tarification expose un `etat` textuel parce qu'il en a TROIS a distinguer.
#  Ici `alerte` porte a lui seul la distinction : ajouter une chaine
#  redondante aurait ete un champ que personne ne lit — le motif que ce
#  chantier ferme depuis le debut.


class TraceRelecture(NamedTuple):
    """L'état de la relecture : sa phrase, et s'il alerte."""
    texte: str
    alerte: bool


def trace_relecture(nom, numero_ia=None) -> TraceRelecture:
    """Les deux états de la relecture actuarielle, en toutes lettres.

    ⚠️ AUCUNE DATE N'EST FABRIQUÉE ICI. Une date posée au moment du rendu
    dirait quand le document a été produit, pas quand il a été relu — et le
    dépôt n'enregistre pas la seconde.
    """
    if nom and str(nom).strip():
        texte = 'Relu et validé par ' + str(nom).strip()
        if numero_ia and str(numero_ia).strip():
            texte += ' — N° IA ' + str(numero_ia).strip()
        return TraceRelecture(texte, False)
    return TraceRelecture(
        'Relecture actuarielle non enregistrée — ce rapport n\'a pas été relu '
        'par un actuaire identifié.', True)


def _build_blocks(n2, n3, n4, narration, source_narration, lob, cli, arr, dt, audit_id, methode, statut, graphiques_html, actuaire_nom='', actuaire_numero_ia='') -> Dict:
    mk    = n3.get('mack', {})
    clark = n3.get('clark', {});         bz  = n3.get('glm_apc', {})
    bt    = n3.get('backtesting', {});   sc  = n4.get('scr', {})
    h1    = n2.get('h1_independance', {}); h2 = n2.get('h2_stabilite', {})
    pw    = n4.get('poids', {})

    BE  = float(n4.get('best_estimate', 0) or 0)
    SIG = float(mk.get('sigma_total', 0) or 0)
    CV  = float(n4.get('cv_inter_methodes', 0) or 0)
    P90 = float(n4.get('reserve_p90', 0) or 0)
    P99 = float(n4.get('reserve_p99_5', 0) or 0)
    # Le nom de la grandeur que P90/P99 portent RÉELLEMENT — les cartes KPI
    # affichaient « (composé) » en dur, y compris quand ce n'en est pas un.
    _appr_h = libelle_percentiles(n4)
    # Percentiles Mack et Bootstrap séparés
    P90_mack = float(n4.get('reserve_p90_mack', P90) or P90)
    P75_boot = n4.get('reserve_p75_boot')
    P90_boot = n4.get('reserve_p90_boot')
    _boot_dispo = P75_boot is not None and float(P75_boot or 0) > 0
    SCP = _scr_publiable(sc)
    SCR = _ratio_scr(SCP, BE)

    s_label = _statut_label(statut)
    s_cls   = statut.lower()

    b = {}

    # ── PAGE DE GARDE ─────────────────────────────────────────────────────────
    # Titre branche : première partie en blanc, deuxième en or italic si " — "
    lob_parts = lob.split(' — ')
    if len(lob_parts) == 2:
        garde_titre = lob_parts[0] + '<br><em>' + lob_parts[1] + '</em>'
    else:
        # RC Médicale → RC / Médicale (style premium)
        words = lob.split()
        if len(words) >= 2:
            garde_titre = ' '.join(words[:-1]) + '<br><em>' + words[-1] + '</em>'
        else:
            garde_titre = '<em>' + lob + '</em>'

    b['garde_titre']   = garde_titre
    b['garde_statut_cls'] = 'garde-statut-' + s_cls
    b['garde_dot_cls']    = 'statut-dot-' + s_cls
    b['garde_label_cls']  = 'statut-label-' + s_cls
    b['garde_label']      = s_label

    # KPIs page de garde
    scr_style = 'color:var(--rouge);' if statut == 'ROUGE' else ''
    b['garde_kpis'] = (
        '<div class="garde-kpi"><div class="kpi-label">Client</div>'
        '<div class="kpi-value">' + cli + '</div></div>'

        '<div class="garde-kpi"><div class="kpi-label">Arrêté</div>'
        '<div class="kpi-value">' + arr + '</div></div>'

        '<div class="garde-kpi"><div class="kpi-label">Méthode principale</div>'
        '<div class="kpi-value" style="font-size:10pt;">'
        + methode.replace('_', ' ').replace('bornhuetter ferguson', 'Bornhuetter-Ferguson').title()
        + '</div></div>'

        '<div class="garde-kpi"><div class="kpi-label">Best Estimate (brut)</div>'
        '<div class="kpi-value highlight">' + _f(BE) + '</div></div>'

        '<div class="garde-kpi"><div class="kpi-label">SCR Provisions</div>'
        '<div class="kpi-value" style="font-size:11pt;' + scr_style + '">' + _f(SCP) + '</div>'
        '<div class="kpi-sub">Ratio SCR/BE : ' + _pct(SCR) + '</div></div>'

        '<div class="garde-kpi"><div class="kpi-label">Audit ID</div>'
        '<div class="kpi-value" style="font-family:\'JetBrains Mono\',monospace;font-size:8pt;color:var(--slate-light);">'
        + (audit_id or '—') + '</div></div>'
    )

    # ── SECTION 1 : KPI GRID ─────────────────────────────────────────────────
    # 4ème KPI : PT S2 si RM disponible, sinon P99,5
    _kpi4 = (
        '<div class="kpi-card" style="border-left:3px solid var(--gold);">'
        '<div class="kpi-card-label">Provisions Tech. S2</div>'
        '<div class="kpi-card-value">' + _f(n4.get('provisions_techniques_s2', 0)) + '</div>'
        '<div class="kpi-card-sub">BE + RM — Art. 77 §1</div>'
        '</div>'
        # ⚠️ MEME GARDE-FOU, MEME CHUTE, DEUX LIGNES PLUS LOIN.
        # `garde_fou_be_negatif` met AUSSI `risk_margin` a None ; `.get(k, 0)`
        # rend alors None — la cle EXISTE — et `None > 0` leve. Le `or 0`
        # traite les deux etats (absent, present a None) de la meme facon.
        if (n4.get('risk_margin') or 0) > 0
        else '<div class="kpi-card">'
        '<div class="kpi-card-label">Provision P99,5 (' + _appr_h + ')</div>'
        '<div class="kpi-card-value">' + _f(P99) + '</div>'
        '<div class="kpi-card-sub">P90 (' + _appr_h + ') : ' + _f(P90) + '</div>'
        '</div>'
    )
    b['kpi_grid'] = (
        '<div class="kpi-grid">'
        '<div class="kpi-card">'
        '<div class="kpi-card-label">Best Estimate (brut)</div>'
        '<div class="kpi-card-value">' + _f(BE) + '</div>'
        '<div class="kpi-card-sub">Réserve avant actualisation — S2 par A10</div>'
        '</div>'
        '<div class="kpi-card">'
        '<div class="kpi-card-label">σ Mack total</div>'
        '<div class="kpi-card-value">' + _f(SIG) + '</div>'
        '<div class="kpi-card-sub">CV inter-méthodes : ' + _pct(CV) + '</div>'
        '</div>'
        '<div class="kpi-card kpi-card-rouge">'
        '<div class="kpi-card-label">SCR Provisions</div>'
        '<div class="kpi-card-value kpi-card-value-rouge">' + _f(SCP) + '</div>'
        '<div class="kpi-card-sub">Ratio SCR/BE : ' + _pct(SCR) + '</div>'
        '</div>'
        + _kpi4
        + '</div>'
    )

    # ── SECTION 2 : MÉTHODES ─────────────────────────────────────────────────
    def _badge_statut(m_name, pw_val):
        if 'Mack' in m_name:   # Mack = volatilité (σ), non pondérée dans le BE (point = CL)
            return ('<span class="badge badge-excl" style="background:rgba(74,144,226,0.15);'
                    'color:#4A90E2;">σ (volatilité)</span>')
        if pw_val > 0:
            return '<span class="badge badge-ok">✓ Inclus</span>'
        return '<span class="badge badge-excl">⊘ Exclu</span>'

    # ⚠️ LE FAUX ZÉRO, DANS LE LIVRABLE QUE C1 A PROMU (lot C3a). Le HTML
    # publiait « Bornhuetter-Ferguson | 0 € | 0.0 % » sans exposition.
    # ⚠️ LA CLÉ VOYAGE AVEC LE LIBELLÉ, ET C'EST LA CORRECTION. Ce tableau
    # cherchait le motif d'exclusion par LIBELLÉ dans un dict indexé par CLÉ :
    # intersection vide, le repli « non calculée » était pris à tous les
    # coups. L'Excel publiait le motif détaillé, le HTML jamais — deux formats,
    # deux informations, sur le même fait.
    rows_m = [(_m, libelle(_m), reserve(n3, _m), pw.get(_m, 0), '—')
              for _m in ORDRE_AFFICHAGE]
    tbl = (
        '<table class="premium"><thead><tr>'
        '<th>Méthode</th>'
        '<th class="right">Réserve IBNR</th>'
        '<th class="center">Poids BE</th>'
        '<th class="center">Score</th>'
        '<th class="center">Statut</th>'
        '</tr></thead><tbody>'
    )
    for _cle_m, nom, res, pds, score in rows_m:
        s_txt = str(score) + ' / 100' if score != '—' else '—'
        tbl += (
            '<tr><td class="label">' + nom + '</td>'
            '<td class="right"><span class="mono">'
            + (_f(res) if res is not None
               else motif_exclusion(n4, _cle_m)) + '</span></td>'
            '<td class="center">' + _pct(pds * 100) + '</td>'
            '<td class="center">' + s_txt + '</td>'
            '<td class="center">' + _badge_statut(nom, pds) + '</td></tr>'
        )
    # Benktander — INFORMATIF, poids nul par construction : il est déjà un
    # mélange de Chain Ladder et de BF, l'inclure au BE les compterait deux
    # fois (Chain Ladder pèserait 44,2 % pour 25 % affichés).
    _gb = n3.get('benktander', {})
    if _gb.get('disponible'):
        tbl += (
            '<tr><td class="label">Benktander (1976)</td>'
            '<td class="right"><span class="mono">'
            + _f(_gb.get('reserve_totale')) + '</span></td>'
            '<td class="center">—</td>'
            '<td class="center">α = ' + _s(f"{_gb.get('alpha_moyen', 0):.4f}")
            + '</td>'
            '<td class="center"><span class="badge badge-excl">'
            'ⓘ Informatif</span></td></tr>'
        )
    # Clark
    if clark.get('disponible'):
        aic_val = str(clark.get('aic_optimal', '—'))
        if aic_val != '—':
            try:
                aic_f = float(aic_val)
                aic_disp = f'{aic_f/1e6:.0f}M' if abs(aic_f) > 1e6 else aic_val
            except Exception:
                aic_disp = aic_val
        else:
            aic_disp = '—'
        # Réserve retenue quand la courbe croissante de Clark ne peut pas
        # représenter le triangle : la colonne de détail porte le motif, sinon
        # un tiret se lirait comme une donnée manquante.
        _ck_struct = clark.get('structure_monotone') or {}
        _ck_detail = (
            'Structure incompatible — facteur(s) de développement sous 1'
            if not _ck_struct.get('compatible', True) else 'AIC = ' + aic_disp
        )
        tbl += (
            '<tr><td class="label">Clark LDF (' + _s(clark.get('courbe_choisie')).title() + ')</td>'
            '<td class="right"><span class="mono">' + _f(clark.get('reserve_be_clark')) + '</span></td>'
            '<td class="center">—</td>'
            '<td class="center" style="font-size:7.5pt;color:var(--slate);">' + _ck_detail + '</td>'
            '<td class="center"><span class="badge badge-excl">⊘ Exclu</span></td></tr>'
        )
    if bz.get('glm_disponible') and bz.get('reserve_apc'):
        _p_bz = bz.get('p_calendaire', 1) or 1
        _p_bz_txt = 'p\u202f&lt;\u202f0,0001' if _p_bz < 0.0001 else f'p\u202f=\u202f{_p_bz:.4f}'
        tbl += (
            '<tr><td class="label">GLM Poisson APC (= CL)</td>'
            '<td class="right"><span class="mono">' + _f(bz.get('reserve_apc')) + '</span></td>'
            '<td class="center">—</td>'
            '<td class="center" style="font-size:7.5pt;color:var(--slate);">' + _p_bz_txt + '</td>'
            '<td class="center"><span class="badge badge-excl" style="background:rgba(74,144,226,0.15);color:#4A90E2;">ℹ️ Informatif</span></td></tr>'
        )
    tbl += (
        '<tr class="highlight-gold">'
        '<td class="label" style="color:var(--navy);">⭐ Best Estimate (brut)</td>'
        '<td class="right" style="color:var(--navy);"><span class="mono">' + _f(BE) + '</span></td>'
        '<td class="center">100\u202f%</td>'
        '<td class="center">—</td>'
        '<td class="center" style="color:var(--navy-light);font-size:8.5pt;">→ Bilan S2</td>'
        '</tr></tbody></table>'
    )
    b['tableau_methodes'] = tbl

    # Table diagnostic — décomposition de l'incertitude (4 lignes, colonne Centre)
    _mk = n3.get('mack', {}); _bo = n3.get('bootstrap', {})
    P90_natif      = float(_mk.get('reserve_p90', 0) or 0)
    P90_COMPOSE    = float(n4.get('reserve_p90_compose', 0) or 0)
    SIG_COMPOSE    = float(n4.get('sigma_total_compose', SIG) or SIG)
    SIG_MACK       = float(n4.get('sigma_mack', SIG) or SIG)
    SIG_MACK_NATIF = float(_mk.get('sigma_total', SIG) or SIG)
    STD_BOOT       = float(_bo.get('std_bootstrap', 0) or 0)
    _bp90 = ('<span class="mono">' + _f(P90_boot) + '</span>') if _boot_dispo else '<span style="color:var(--slate)">—</span>'
    _bsig = ('<span class="mono">' + _f(STD_BOOT) + '</span>')  if _boot_dispo else '<span style="color:var(--slate)">—</span>'
    # ⚠️ « (retenue) », LE SURLIGNAGE OR ET LA PHRASE DE PIED étaient CLOUÉS sur
    # la ligne du composé. Les trois suivent désormais l'arbitrage de N4 : la
    # ligne marquée est celle qui porte réellement `reserve_p90`. « Mack natif »
    # ne peut jamais l'être — il est centré sur la réserve de Mack, pas sur le
    # BE publié — d'où sa clé vide, qu'aucun arbitrage ne peut désigner.
    _mo = '<span class="mono">'
    _lignes_i = [
        (CLE_COMPOSE, 'Incertitude composée',
         _mo + _f(P90_COMPOSE) + '</span>', _mo + _f(SIG_COMPOSE) + '</span>',
         'BE pondéré'),
        (CLE_MACK, 'Mack recentré',
         _mo + _f(P90_mack) + '</span>', _mo + _f(SIG_MACK) + '</span>',
         'BE pondéré'),
        ('', 'Mack natif',
         _mo + _f(P90_natif) + '</span>', _mo + _f(SIG_MACK_NATIF) + '</span>',
         'réserve Mack'),
        (CLE_BOOT, 'Bootstrap ODP', _bp90, _bsig, 'réserve Bootstrap'),
    ]
    _corps_i = ''
    for _cle, _base, _vp90, _vsig, _centre in _lignes_i:
        _cls = (' class="highlight-gold"'
                if _cle and n4.get('cle_percentiles') == _cle else '')
        _corps_i += (
            '<tr' + _cls + '><td class="label">'
            + marque_retenue(n4, _cle, _base) + '</td>'
            '<td class="right">' + _vp90 + '</td>'
            '<td class="right">' + _vsig + '</td>'
            '<td>' + _centre + '</td></tr>')
    tbl_i = (
        '<table class="premium"><thead><tr>'
        '<th>Approche</th><th class="right">P90</th><th class="right">σ</th><th>Centre</th>'
        '</tr></thead><tbody>' + _corps_i + '</tbody></table>'
        '<p style="font-size:8pt;color:var(--slate);font-style:italic;margin-top:6px;">'
        'Ces valeurs diffèrent par le σ (Mack seul / composé / bootstrap) et/ou le point de '
        'centrage (colonne Centre). Le livrable retient le P90 '
        + libelle_percentiles(n4).lower() + '.</p>'
    )
    # Clark porte lui aussi une incertitude, construite autrement : variance de
    # PROCESSUS + variance de PARAMÈTRE (Clark 2003), là où Mack agrège des σ_j
    # et le Bootstrap rééchantillonne. Elle manquait à ce tableau alors qu'elle
    # est calculée depuis le correctif de couverture (30,8 % → 97,0 %).
    # ⚠️ On s'appuie sur la réserve PUBLIÉE, pas sur `reserve_brute` : quand la
    # structure du triangle est incompatible avec une courbe monotone, Clark
    # retient sa réserve — faire apparaître son incertitude ici la republierait
    # par la bande, au milieu de Mack et du Bootstrap.
    _ck_se  = clark.get('se_reserve_totale')
    _ck_res = clark.get('reserve_be_clark')
    if clark.get('disponible') and _ck_se and _ck_res is not None:
        _ck_p90 = float(_ck_res) + Z90 * float(_ck_se)
        tbl_i = tbl_i.replace(
            '</tbody></table>',
            '<tr><td class="label">Clark LDF (paramètre + processus)</td>'
            '<td class="right"><span class="mono">' + _f(_ck_p90) + '</span></td>'
            '<td class="right"><span class="mono">' + _f(_ck_se) + '</span></td>'
            '<td>réserve Clark</td></tr>'
            '</tbody></table>', 1)
    b['tableau_incertitude'] = tbl_i
    b['tableau_clark_ic']    = _bloc_clark_incertitude(n2, n3)

    # Munich CL — n'apparaissait dans AUCUN des trois formats de livrable
    # (ni ce module pour HTML/Word, ni n5_excel). Seule méthode à exploiter
    # les deux triangles, et la seule totalement invisible.
    _lig_mcl = lignes_munich_rapport(n3)
    # Benktander — MEME SOURCE que le Word et l'Excel. Informatif : il ne
    # figure pas au Best Estimate, il en mesure la sensibilite a la ponderation.
    _lig_gb = lignes_benktander_rapport(n3, n4)
    b['tableau_benktander'] = ('' if not _lig_gb else (
        '<table class="premium"><thead><tr>'
        '<th>Élément</th><th>Valeur</th>'
        '</tr></thead><tbody>'
        + ''.join('<tr><td class="label">' + _s(a) + '</td>'
                  '<td>' + _s(v) + '</td></tr>' for a, v in _lig_gb)
        + '</tbody></table>'))

    b['tableau_munich'] = ('' if not _lig_mcl else (
        '<table class="premium"><thead><tr>'
        '<th>Indicateur</th><th class="right">Valeur</th><th>Lecture</th>'
        '</tr></thead><tbody>'
        + ''.join('<tr><td class="label">' + _s(a) + '</td>'
                  '<td class="right"><span class="mono">' + _s(v) + '</span></td>'
                  '<td style="font-size:8pt;color:var(--slate);">' + _s(c) + '</td></tr>'
                  for a, v, c in _lig_mcl)
        + '</tbody></table>'))

    # ── SECTION 3 : HYPOTHÈSES ────────────────────────────────────────────────
    hyp_cards = ''
    for label, h, code in [
        ('H1 — Indépendance des facteurs', h1, 'H1'),
        ('H2 — Stabilité des facteurs', h2, 'H2'),
    ]:
        if not h:
            continue
        # ⚠️⚠️ LA CARTE SE CONTREDISAIT ELLE-MEME, ET C'EST LE DEFAUT QUI A
        # OUVERT CE LOT. Le badge se deduisait de `ok` — True par defaut quand
        # rien n'est testable — pendant que le texte juste en dessous portait
        # le message honnete du module. Mesure sur un triangle 3x3 sain :
        #
        #     BADGE   : ✓ VALIDÉE · Score 80 / 100
        #     TEXTE   : H2 NON TESTABLE — aucune période ne porte 3 facteurs…
        #
        # Une pastille VERTE au-dessus de son propre dementi. Les deux
        # hypotheses etaient touchees.
        # ⚠️ UNE SEULE DERIVATION DEPUIS `ok`, ET ELLE EST LE REPLI DU `.get`.
        # La premiere version de ce correctif rebranchait sur `ok` pour choisir
        # le libelle : le controle du filet l'a REFUSEE, à juste titre. Le
        # badge se lit desormais du seul `statut`, par table — ce qui couvre
        # aussi tout statut futur sans nouvelle branche.
        # ⚠️⚠️ NOM LOCAL `statut_hyp`, ET C'EST UNE CORRECTION DE BUG. Le lot C
        # avait ecrit `statut` ici : cette boucle ECRASAIT le parametre
        # `statut` de la fonction -- le statut RAG du dossier. Apres elle, tout
        # code lisant `statut` recevait le verdict de la DERNIERE hypothese
        # rendue (<< VALIDÉE >>, << NON TESTABLE >>...), plus le RAG. Le defaut
        # est reste INVISIBLE tant que rien n'utilisait `statut` en aval ; il a
        # ete revele par le lot avis-couleur, qui l'a fait pour la premiere
        # fois. Verrouille par un controle AST : aucun parametre de
        # `_build_blocks` ne doit etre reaffecte dans son corps.
        statut_hyp = str(h.get('statut',
                               'VALIDÉE' if h.get('ok', True) else 'REJETÉE'))
        non_testable = statut_hyp == 'NON TESTABLE'
        cls   = 'hyp-ok' if statut_hyp == 'VALIDÉE' else 'hyp-warn'
        lbl   = {'VALIDÉE': '✓ VALIDÉE',
                 'REJETÉE': '⚠ REJETÉE'}.get(statut_hyp, '⚠ ' + statut_hyp)
        # ⚠️ UN SCORE NE SE PUBLIE PAS SOUS UN STATUT QUI DIT QU'IL N'Y A PAS EU
        # DE MESURE. « Score 80 / 100 » etait la valeur par defaut du module.
        score = '' if non_testable else ' · Score ' + _s(h.get('score')) + ' / 100'
        msg   = _s(h.get('message'))
        hyp_cards += (
            '<div class="hyp-card ' + cls + '">'
            '<div class="hyp-label">'
            '<div class="hyp-code">' + label + '</div>'
            '<div class="hyp-score">' + lbl + score + '</div>'
            '</div>'
            '<div class="hyp-text">' + msg + '</div>'
            '</div>'
        )
    # BFCC-H1..H6 — statut motivé, SANS score : ces hypothèses n'en produisent
    # aucun. Afficher « Score — / 100 » serait moins clair que ne rien afficher.
    # BOOT-H1..H4 — idem. La carte « H4 — Homoscédasticité Bootstrap ODP »
    # affichait un score sur 100 calculé sur le CV des variances des
    # facteurs de développement, rouge sur les trois triangles de référence.
    # MCL-H1..H5 rejoignent BFCC et BOOT : mêmes cartes, même convention de
    # statut. MCL-H4 y figure comme MENTION (NON TESTABLE assumé, cf. son
    # message) — elle ne peut pas être verte, on ne sait pas la tester.
    # ⚠️ CLM EN TÊTE, ET C'EST UN MANQUE QUI EST RÉPARÉ ICI. BFCC, Bootstrap et
    # Munich avaient chacun leur affichage ; CLM-H1..H4 — les hypothèses des
    # MÉTHODES PRINCIPALES, dont deux sont gatantes — n'atteignaient aucun
    # livrable. Elles viennent en premier parce que Chain Ladder et Mack
    # portent le Best Estimate.
    for ligne in (list(lignes_hypotheses_clm(n2))
                  + list(lignes_hypotheses_bfcc(n2))
                  + list(lignes_hypotheses_bootstrap(n2))
                  + list(lignes_hypotheses_munich(n2))):
        cls = ('hyp-ok' if ligne['statut'] == 'VALIDÉE'
               else 'hyp-warn' if ligne['statut'] != 'NON VALIDÉE' else 'hyp-warn')
        # La PUISSANCE se lit avec le verdict, jamais à sa place : elle dit ce
        # que le test POUVAIT détecter sur ce triangle-ci.
        # ⚠️ ON TESTE LA VALEUR BRUTE, PAS `_s(...)` : `_s('')` rend « — », qui
        # est vrai en Python. Le bloc serait alors posé sur les quinze
        # hypothèses qui n'ont pas de puissance, avec un tiret pour tout
        # contenu — le faux zéro que ce dépôt traque depuis plusieurs lots.
        brut = ligne.get('puissance_phrase')
        bloc_p = ('<div class="hyp-puissance">' + _s(brut) + '</div>'
                  if brut else '')
        hyp_cards += (
            '<div class="hyp-card ' + cls + '">'
            '<div class="hyp-label">'
            '<div class="hyp-code">' + _s(ligne['libelle']) + '</div>'
            '<div class="hyp-score">' + _s(ligne['statut']) + '</div>'
            '</div>'
            '<div class="hyp-text">' + _s(ligne['message']) + '</div>'
            + bloc_p +
            '</div>'
        )
    # Clark — PRÉSENTATION, pas gouvernance. Ces cartes n'affichent aucun
    # statut VALIDÉE / NON VALIDÉE : elles rendent des faits déjà calculés
    # (cf. `lignes_clark_rapport`, qui documente pourquoi aucun verdict
    # d'adéquation de courbe n'est produit).
    for ligne in lignes_clark_rapport(n2, n3):
        hyp_cards += (
            '<div class="hyp-card ' + ('hyp-ok' if ligne['ok'] == 'oui'
                                       else 'hyp-warn') + '">'
            '<div class="hyp-label">'
            '<div class="hyp-code">' + _s(ligne['libelle']) + '</div>'
            '<div class="hyp-score">' + _s(ligne['etat']) + '</div>'
            '</div>'
            '<div class="hyp-text">' + _s(ligne['texte']) + '</div>'
            '</div>'
        )
    b['hyp_cards'] = hyp_cards

    raison = _s(n2.get('raison_recommandation'))[:200]
    b['methode_rec_box'] = (
        '<div class="recommandation-box">'
        '<div class="reco-icon">→</div>'
        '<div>'
        '<div class="reco-label">Méthode recommandée par ActuarIA</div>'
        '<div class="reco-value">' + methode.replace('_', ' ').title() + '</div>'
        '<div class="reco-text">' + raison + '</div>'
        '</div>'
        '</div>'
    )

    # ── SECTION 4 : SCR ───────────────────────────────────────────────────────
    # En pourcentage et à la décimale : deux segments ont un σ à trois chiffres
    # significatifs (protection juridique 5,5 %, crédit 17,2 %) que l'affichage
    # brut du flottant rendait illisible.
    sigma_eiopa = _pct(float(sc.get('sigma_eiopa') or 0) * 100) if sc else '—'
    b['tableau_scr'] = (
        '<table class="premium"><thead><tr>'
        '<th>Composante</th><th class="center">Valeur</th><th>Référence réglementaire</th>'
        '</tr></thead><tbody>'
        '<tr><td class="label">Best Estimate (brut)</td>'
        '<td class="center"><span class="mono">' + _f(BE) + '</span></td>'
        '<td>Réserve brute — actualisation S2 en aval (A10)</td></tr>'
        '<tr><td class="label">Facteur σ EIOPA</td>'
        '<td class="center"><span class="mono">' + sigma_eiopa + '</span></td>'
        '<td>' + lob + ' — ' + _s(sc.get('reference_s2', 'Annexes II / XIV, Règlement 2015/35') if sc else 'Annexes II / XIV, Règlement 2015/35') + '</td></tr>'
        '<tr><td class="label">SCR Provisions</td>'
        '<td class="center"><span class="mono" style="color:var(--rouge);font-weight:700;">' + _f(SCP) + '</span></td>'
        '<td>SCR = 3 × σ(LoB) × BE — Art. 115</td></tr>'
        '<tr class="highlight-gold"><td class="label">Ratio SCR / BE</td>'
        '<td class="center" style="color:var(--rouge);font-weight:700;"><span class="mono">' + _pct(SCR) + '</span></td>'
        '<td>Cible pratique marché : &lt; 35\u202f%</td></tr>'
        # Risk Margin S2
        + ('<tr style="border-top:2px solid var(--navy);"><td class="label" style="color:var(--navy);font-weight:600;">Risk Margin S2</td>'
        '<td class="center"><span class="mono" style="color:var(--navy);font-weight:700;">'
        + _f(n4.get('risk_margin', 0)) + '</span></td>'
        # ⚠️ << Courbe EIOPA >> ÉTAIT ÉCRIT EN DUR. Avec un taux assumé par
        # l'actuaire, le rapport affichait « Courbe EIOPA Arrêté courant » —
        # il attribuait à EIOPA un taux que l'actuaire avait saisi lui-même.
        # C'est la PROVENANCE publiée qui parle désormais.
        '<td>Art. 77 §5 — CoC 6% · Méthode proportionnelle · '
        + _s(n4.get('source_courbe_rfr', '—')) + ' ('
        + _s(n4.get('date_courbe_rfr', '—')) + ')</td></tr>'
        '<tr class="highlight-gold"><td class="label">Provisions Techniques S2</td>'
        '<td class="center"><span class="mono" style="font-weight:700;">'
        + _f(n4.get('provisions_techniques_s2', 0)) + '</span></td>'
        '<td>PT S2 = BE + RM — Art. 77 §1</td></tr>')
        + '</tbody></table>'
    )

    # ── SECTION 5 : BACK-TESTING ──────────────────────────────────────────────
    bt_statut = _s(bt.get('statut', 'AMBRE'))
    bt_col_cls = 'bt-card-' + bt_statut.lower() if bt_statut in ('ROUGE', 'AMBRE', 'VERT') else 'bt-card-navy'
    bt_score  = _s(bt.get('score_qualite', '—'))
    # Recalculer depuis le tableau réel pour cohérence
    _bt_tab = bt.get('tableau', [])
    _SR, _SA = 15.0, 8.0
    _nr1 = _na1 = _nr2 = _na2 = 0
    for _r in _bt_tab:
        if not isinstance(_r, dict) or not _r.get('mature', True): continue
        for _hor in ['n1', 'n2']:
            try:
                _ep = abs(float(_r.get('ecart_pct_' + _hor, 0) or 0))
                if _ep >= _SR:
                    if _hor == 'n1': _nr1 += 1
                    else: _nr2 += 1
                elif _ep >= _SA:
                    if _hor == 'n1': _na1 += 1
                    else: _na2 += 1
            except Exception: pass
    n_rouge_n1_disp = _nr1 if _bt_tab else int(bt.get('n_rouge_n1', 0))
    n_ambre_n1_disp = _na1 if _bt_tab else int(bt.get('n_ambre_n1', 0))
    n_rouge_n2_disp = _nr2 if _bt_tab else int(bt.get('n_rouge_n2', 0))
    n_ambre_n2_disp = _na2 if _bt_tab else int(bt.get('n_ambre_n2', 0))

    b['bt_summary'] = (
        '<div class="bt-summary">'
        '<div class="bt-card ' + bt_col_cls + '">'
        '<div class="bt-card-label">Statut RAG</div>'
        '<div class="bt-card-value">' + bt_statut + '</div>'
        '</div>'
        '<div class="bt-card bt-card-navy">'
        '<div class="bt-card-label">Score qualité</div>'
        '<div class="bt-card-value">' + bt_score + '</div>'
        '<div class="bt-card-sub">/ 100</div>'
        '</div>'
        '<div class="bt-card bt-card-navy">'
        '<div class="bt-card-label">Alertes N-1</div>'
        '<div class="bt-card-value">' + str(n_rouge_n1_disp) + '🔴  ' + str(n_ambre_n1_disp) + '🟡</div>'
        '</div>'
        '<div class="bt-card bt-card-navy">'
        '<div class="bt-card-label">Alertes N-2</div>'
        '<div class="bt-card-value">' + str(n_rouge_n2_disp) + '🔴  ' + str(n_ambre_n2_disp) + '🟡</div>'
        '</div>'
        '</div>'
    )

    # Tableaux boni/mali détaillés
    b['tableau_bt_n1'] = _build_bt_table(bt, 'n1')
    b['tableau_bt_n2'] = _build_bt_table(bt, 'n2')

    msg_bt = _s(bt.get('message'))
    _bt_note_txt = msg_bt if msg_bt and msg_bt != '—' else ''
    b['bt_note'] = '<div class="bt-note">' + _bt_note_txt + '</div>' if _bt_note_txt else ''

    # ── SECTION 6 : EFFETS CALENDAIRE ────────────────────────────────────────
    # ⚠️ L'ETAT VIENT DE `etat_calendaire`, SOURCE PARTAGEE AVEC LE WORD ET LE
    # PROMPT. Les trois branches ci-dessous sont le rendu HTML des trois etats
    # sans effet ; le balisage est INCHANGE (verrouille par non-regression).
    _etat_cal = etat_calendaire(bz)
    bz_items = ''
    if _etat_cal != 'present':
        if _etat_cal == 'indisponible':
            # FIX 2 : test indisponible -> ne jamais afficher "aucun effet".
            bz_items = (
                '<div class="bz-item bz-info">'
                '<span class="bz-year" style="color:var(--cyan);">n/d</span>'
                '<div class="bz-bar"><div class="bz-fill" style="width:0%;background:var(--cyan);"></div></div>'
                '<span class="bz-pct" style="color:var(--cyan);">Test calendaire indisponible</span>'
                '</div>'
            )
        elif _etat_cal == 'diffus':
            # FIX 1 : F global significatif mais aucune diagonale isolee dominante (effet diffus).
            bz_items = (
                '<div class="bz-item bz-warn">'
                '<span class="bz-year" style="color:var(--orange);">Diffus</span>'
                '<div class="bz-bar"><div class="bz-fill" style="width:60%;background:var(--orange);"></div></div>'
                '<span class="bz-pct" style="color:var(--orange);">Effet calendaire global réparti (aucune diagonale isolée dominante)</span>'
                '</div>'
            )
        else:
            bz_items = (
                '<div class="bz-item bz-ok">'
                '<span class="bz-year" style="color:var(--vert);">&#10003; OK</span>'
                '<div class="bz-bar"><div class="bz-fill" style="width:0%;background:var(--vert);"></div></div>'
                '<span class="bz-pct" style="color:var(--vert);">Aucun effet significatif</span>'
                '</div>'
            )
    else:
        forts   = [e for e in bz.get('effets_calendaire', []) if e.get('significatif') and e.get('niveau') in ('FORT', 'MODÉRÉ')]
        faibles = [e for e in bz.get('effets_calendaire', []) if e.get('significatif') and e.get('niveau') == 'FAIBLE']
        for e in forts:
            amp  = float(e.get('amplitude_pct', 0))
            w    = min(abs(amp) * 5, 100)
            col  = 'var(--rouge)' if e.get('niveau') == 'FORT' else 'var(--orange)'
            niv  = e.get('niveau', '')
            sens = '↑' if e.get('sens') == 'hausse' else '↓'
            bz_items += (
                '<div class="bz-item bz-warn">'
                '<span class="bz-year">' + _s(e.get('annee_label')) + '</span>'
                '<div class="bz-bar"><div class="bz-fill" style="width:' + str(w) + '%;background:' + col + ';"></div></div>'
                '<span class="bz-pct" style="color:' + col + ';">' + f'{amp:+.1f}\u202f%' + ' ' + sens + ' ' + niv + '</span>'
                '</div>'
            )
        if faibles:
            n_f = len(faibles)
            bz_items += (
                '<div class="bz-item bz-info">'
                '<span class="bz-year" style="color:var(--cyan);">+' + str(n_f) + ' diag.</span>'
                '<div class="bz-bar"><div class="bz-fill" style="width:20%;background:var(--cyan);"></div></div>'
                '<span class="bz-pct" style="color:var(--cyan);">' + str(n_f) + ' diag. &lt; 8\u202f%</span>'
                '</div>'
            )
    b['bz_items'] = bz_items

    # -- Resultats GLM Poisson APC (test calendaire F quasi-Poisson) --
    glm_dispo  = bz.get('glm_disponible', False)
    p_cal      = bz.get('p_calendaire')
    cal_sig    = bz.get('cal_significatif')
    F_cal      = bz.get('F_calendaire')
    phi_disp   = bz.get('phi_dispersion')
    ddl_cal    = bz.get('ddl_calendaire')
    res_apc    = bz.get('reserve_apc')

    if glm_dispo and p_cal is not None:
        _cal_col  = 'var(--orange)' if cal_sig else 'var(--vert)'
        _p_fmt    = 'p &lt; 0,0001' if p_cal < 0.0001 else ('p = ' + f'{p_cal:.4f}'.replace('.', ','))
        _cal_txt  = ('Effet calendaire significatif au seuil 5 % (' + _p_fmt + ')'
                     if cal_sig else 'Effet calendaire non significatif (' + _p_fmt + ')')
        _stat_fmt = (('F = ' + f'{F_cal:.2f}'.replace('.', ',') + ' (ddl = ' + str(ddl_cal) + ')')
                     if F_cal is not None else '-')
        _phi_fmt  = (f'{phi_disp:,.0f}'.replace(',', ' ') if phi_disp else '-')
        b['bz_glm'] = (
            '<div class="section-body" style="margin-top:24px;padding:20px 28px;'
            'background:var(--navy-mid, #0D2137);border-radius:8px;'
            'border:1px solid rgba(212,175,55,0.25);">'

            '<div style="font-size:7.5pt;font-weight:700;color:var(--gold);'
            'text-transform:uppercase;letter-spacing:1.5px;margin-bottom:16px;">'
            '&#9670; GLM Poisson APC &mdash; Test F quasi-Poisson (effet calendaire)</div>'

            '<table style="width:100%;border-collapse:collapse;font-size:8.5pt;"><tbody>'
            '<tr style="border-bottom:1px solid rgba(255,255,255,0.08);">'
            '<td style="padding:8px 0;color:var(--slate);width:55%;">Conclusion du test F</td>'
            '<td style="padding:8px 0;font-weight:700;color:' + _cal_col + ';">' + _cal_txt + '</td>'
            '</tr>'
            '<tr style="border-bottom:1px solid rgba(255,255,255,0.08);">'
            '<td style="padding:8px 0;color:var(--slate);">Statistique de test</td>'
            '<td style="padding:8px 0;"><span class="mono" style="color:var(--white);">' + _stat_fmt + '</span>'
            '<span style="color:var(--slate);font-size:7.5pt;"> &mdash; &#966; = ' + _phi_fmt + ' (sur-dispersion Poisson)</span></td>'
            '</tr>'
            + ('<tr><td style="padding:8px 0;color:var(--slate);">Réserve GLM APC (âge-cohorte)</td>'
               '<td style="padding:8px 0;"><span class="mono" style="color:var(--gold);font-weight:700;">'
               + _f(res_apc) +
               '</span><span style="color:var(--slate);font-size:7.5pt;"> &mdash; = chain-ladder par MLE '
               '(Renshaw-Verrall 1998) ; le calendaire n’est pas projeté</span></td></tr>'
               if res_apc else '') +
            '</tbody></table>'

            '<div style="margin-top:12px;font-size:7.5pt;color:var(--slate);'
            'border-top:1px solid rgba(255,255,255,0.06);padding-top:10px;">'
            'GLM Poisson cross-classifié âge-période-cohorte (Renshaw &amp; Verrall 1998). '
            'H&#8320; : effets calendaires nuls &mdash; test F quasi-Poisson (LR / &#966;), rejeté si p &lt; 0,05. '
            'Le test capte la courbure (chocs, changements de régime) ; une inflation '
            'constante est non identifiable sur le triangle seul (VERT ne l’exclut pas).</div>'
            '</div>'
        )
    else:
        b['bz_glm'] = ''

    reco_bz = _s(bz.get('recommandation'))
    bz_reco_val = reco_bz.split(' — ')[0] if ' — ' in reco_bz else reco_bz
    bz_reco_txt = reco_bz.split(' — ')[1] if ' — ' in reco_bz else ''
    b['bz_reco_box'] = (
        '<div class="recommandation-box" style="margin-top:20px;">'
        '<div class="reco-icon" style="color:var(--cyan);">↗</div>'
        '<div>'
        '<div class="reco-label">Recommandation GLM Poisson APC</div>'
        '<div class="reco-value">' + bz_reco_val + '</div>'
        '<div class="reco-text">' + bz_reco_txt + '</div>'
        '</div>'
        '</div>'
    )

    # -- Sous-section : Barnett-Zehnwirth PTF (detection tendances/ruptures, etape 2a) --
    ptf = n3.get('bz_ptf', {})
    if ptf and ptf.get('disponible'):
        _p_rc   = ptf.get('ruptures_calendaires_testees', [])
        _p_scan = ptf.get('ruptures_candidates_scan', [])
        _p_sig  = ptf.get('calendaire_significatif')
        _p_col  = 'var(--orange)' if _p_sig else 'var(--vert)'
        if _p_rc:
            _p_rows = ''.join(
                '<tr><td style="padding:6px 0;color:var(--slate);">' + _s(r.get('annee_label')) + '</td>'
                '<td style="padding:6px 0;text-align:right;"><span class="mono" style="color:'
                + ('var(--orange)' if r.get('significatif') else 'var(--slate)') + ';">'
                + ('%+.1f %%/an' % float(r.get('delta_pct_par_an', 0)))
                + ((' <span style="font-weight:400;color:var(--slate);">[%+.1f %% ; %+.1f %%]</span>'
                    % (r['ic95_pct_par_an'][0], r['ic95_pct_par_an'][1])) if r.get('ic95_pct_par_an') else '')
                + '</span></td>'
                '<td style="padding:6px 0 6px 16px;color:var(--slate);font-size:7.5pt;">p = '
                + ('%.4f' % float(r.get('p_value', 1))) + (' (significatif)' if r.get('significatif') else ' (non sig.)')
                + '</td></tr>'
                for r in _p_rc)
            _p_rcbloc = ('<table style="width:100%;border-collapse:collapse;font-size:8.5pt;margin-top:8px;">'
                         '<tbody>' + _p_rows + '</tbody></table>')
        else:
            _p_rcbloc = ('<div style="font-size:8pt;color:var(--slate);margin-top:6px;">'
                         'Aucune rupture declaree par l\'actuaire (mode semi-manuel) ; '
                         'le scan ci-dessous propose des candidats a examiner.</div>')
        _p_scantxt = ', '.join(_s(c.get('annee_label')) + ' (t = ' + ('%.1f' % float(c.get('t_stat', 0))) + ')'
                               for c in _p_scan) or 'aucun'
        _p_sigb = ptf.get('sigma_par_bande')
        _p_sigtxt = (('sigma par bande dev = ' + ', '.join(str(v) for v in _p_sigb.values()))
                     if _p_sigb else ('sigma constante = ' + str(ptf.get('sigma_constante'))))
        b['ptf_block'] = (
            '<div class="section-body" style="margin-top:20px;padding:20px 28px;'
            'background:var(--navy-mid, #0D2137);border-radius:8px;border:1px solid rgba(212,175,55,0.25);">'
            '<div style="font-size:7.5pt;font-weight:700;color:var(--gold);text-transform:uppercase;'
            'letter-spacing:1.5px;margin-bottom:12px;">'
            '&#9670; Barnett-Zehnwirth PTF (log-normal) &mdash; Tendances &amp; ruptures</div>'
            '<div style="font-size:8.5pt;font-weight:700;color:' + _p_col + ';">'
            + ('Rupture calendaire significative detectee.' if _p_sig
               else 'Aucune rupture calendaire significative confirmee.') + '</div>'
            + _p_rcbloc +
            '<div style="font-size:8pt;color:var(--slate);margin-top:10px;">'
            'Candidats de rupture (scan diagnostic, <b>non</b> testes ni retenus) : ' + _p_scantxt + '. '
            + _p_sigtxt + '.</div>'
            '<div style="margin-top:10px;font-size:7.5pt;color:var(--slate);'
            'border-top:1px solid rgba(255,255,255,0.06);padding-top:8px;">'
            'Detection log-normale (B&amp;Z 2000) &mdash; complete le test F du GLM APC '
            '(plus puissante sur tendances moderees, localise les ruptures). Ruptures semi-manuelles. '
            '&#9888; B&amp;Z detecte les <b>changements</b> de tendance, pas une inflation <b>constante</b> '
            '(non identifiable sur le triangle seul). Aucune reserve produite (etape 2a).</div>'
            '</div>'
        )
    elif ptf and (not ptf.get('disponible')) and ptf.get('message'):
        b['ptf_block'] = (
            '<div class="section-body" style="margin-top:20px;padding:14px 20px;'
            'background:var(--cyan-pale);border-left:3px solid var(--cyan);border-radius:4px;">'
            '<div style="font-size:8pt;color:var(--slate);">&#9670; Barnett-Zehnwirth PTF : '
            + _s(ptf.get('message')) + '</div></div>'
        )
    else:
        b['ptf_block'] = ''

    # ── SECTION 7 : COMMENTAIRE ACTUARIEL ────────────────────────────────────
    try:
        narration_html = _md_to_html(narration)
    except Exception as e_nar:
        logger.warning(f'_md_to_html échoué : {e_nar}')
        narration_html = '<p class="comm-p">' + _clean(narration) + '</p>' if narration else \
            '<p class="comm-p" style="color:#8A9BB0;font-style:italic;">Narration non disponible.</p>'

    source_badge = badge_narration(source_narration)

    b['narration_html']  = narration_html
    b['commentaire_badge'] = source_badge
    b['commentaire_header_title'] = 'Rapport de Provisionnement — ' + lob + ' · Arrêté ' + arr

    # ── SECTION 8 : JUGEMENT ─────────────────────────────────────────────────
    alertes = ''
    for a in n4.get('alertes', n2.get('alertes', [])):
        at = _clean(str(a))
        if at:
            alertes += (
                '<div class="hyp-card hyp-warn" style="margin-bottom:10px;">'
                '<div class="hyp-label"><div class="hyp-code">Point de vigilance</div></div>'
                '<div class="hyp-text">' + at + '</div>'
                '</div>'
            )
    b['alertes'] = alertes

    recs = ''
    for r in n4.get('recommandations', []):
        rt = _clean(str(r))
        if rt:
            recs += (
                '<div class="hyp-card hyp-ok" style="margin-bottom:10px;">'
                '<div class="hyp-label"><div class="hyp-code">Recommandation</div></div>'
                '<div class="hyp-text">' + rt + '</div>'
                '</div>'
            )
    b['recommandations'] = recs

    avis_txt = _clean(n4.get('avis_actuariel', ''))
    if avis_txt:
        # ⚠️⚠️ LA COULEUR ETAIT TOUJOURS VERTE, PAS << verte sauf si >>. Elle
        # cherchait le mot << DEFAVORABLE >> dans `avis_actuariel` -- un mot que
        # ce champ NE PRODUIT JAMAIS : ses trois valeurs (n4_best_estimate,
        # l. 1315-1322) commencent TOUTES par << FAVORABLE >>. Le test etait
        # donc toujours faux, et l'avis d'un dossier ROUGE
        # (<< FAVORABLE SOUS RESERVE -- revisions requises avant bilan S2 >>)
        # s'affichait EN VERT. Le controle lisait le vocabulaire d'un AUTRE
        # champ : c'est `jugement` (l. 1919) qui ecrit << AVIS DEFAVORABLE >>.
        #
        # LA COULEUR S'ADOSSE DESORMAIS AU STATUT RAG, PAR MATCH EXACT -- meme
        # principe que `_RAG_CELLULE` cote Word : une table sur un vocabulaire
        # FINI (VERT / AMBRE / ROUGE), jamais une recherche de mot dans du
        # texte libre. Le statut est deja calcule et deja transmis ici.
        #
        # ⚠️ ET LA COULEUR N'EST JAMAIS SEULE PORTEUSE : le texte de l'avis
        # reste imprime tel quel. Le libelle lui-meme -- << FAVORABLE SOUS
        # RESERVE >> sur un dossier que `jugement` declare << a ne pas inscrire
        # au bilan >> -- est une INCOHERENCE DU PRODUCTEUR, nommee a l'ardoise
        # et NON traitee ici : elle porte sur le verdict, pas sur son rendu.
        avis_col = _COULEUR_AVIS_HTML.get(statut, 'var(--orange)')
        b['avis'] = (
            '<div style="background:' + avis_col + ';color:#fff;padding:16px 24px;'
            'border-radius:6px;margin-top:20px;font-size:10pt;font-weight:600;">'
            + avis_txt + '</div>'
        )
    else:
        b['avis'] = ''

    # ── PIED ─────────────────────────────────────────────────────────────────
    # Signature actuaire — DEUX ETATS, cf. `trace_relecture`.
    act_nom = actuaire_nom or ''
    act_ia  = actuaire_numero_ia or ''
    _tr = trace_relecture(act_nom, act_ia)
    b['pied_info'] = cli + ' · ' + lob + ' · Arrêté ' + arr + ' · Audit ID : ' + (audit_id or '—') + ' · ' + dt
    b['relecture'] = _tr
    # ⚠️ L'ALERTE SE VOIT DANS LES DEUX FORMATS. Le Word colore la mention en
    # ambre quand la relecture n'est pas enregistree ; sans cette regle, l'HTML
    # l'aurait rendue au gris du pied, et l'etat negatif aurait ete ACTIF dans
    # un format et discret dans l'autre — la divergence que ce lot ferme.
    b['signature_actuaire'] = (
        ('<span style="color:' + ORANGE + ';">' + _tr.texte + '</span>'
         if _tr.alerte else _tr.texte) + ' · ')
    b['actuaire_nom'] = act_nom
    b['actuaire_ia']  = act_ia

    return b


def _build_bt_table(bt: Dict, horizon: str) -> str:
    """
    Construit le tableau boni/mali pour un horizon donné.
    Affiche TOUTES les années :
    - Matures (pct_dev >= 75%) : évaluées avec statut ROUGE/AMBRE/VERT
    - Immatures : grisées, mention "Immature", non comptabilisées dans les alertes
    """
    tableau = bt.get('tableau', [])
    if not tableau:
        resultats = bt.get('resultats', {})
        hor_key = 'horizon_1' if horizon == 'n1' else 'horizon_2'
        tableau = resultats.get(hor_key, {}).get('annees', [])

    data = [r for r in tableau if isinstance(r, dict) and r.get('observe_n', 0) > 0]

    if not data:
        return (
            '<p style="font-size:9pt;color:var(--slate);font-style:italic;">'
            'Données de back-testing non disponibles pour cet arrêté.</p>'
        )

    hor_label  = 'N-1' if horizon == 'n1' else 'N-2'
    seuil_r, seuil_a = 15.0, 8.0

    html = (
        '<table class="premium"><thead><tr>'
        '<th>Année</th>'
        '<th class="right">Observé N (€)</th>'
        '<th class="right">Ultimate ' + hor_label + ' (€)</th>'
        '<th class="right">Boni/Mali (€)</th>'
        '<th class="center">Écart (%)</th>'
        '<th class="center">Développé</th>'
        '<th class="center">Statut</th>'
        '</tr></thead><tbody>'
    )

    for idx, row in enumerate(data):
        annee   = _s(row.get('annee_label', row.get('annee', '—')))
        obs     = row.get('observe_n', 0)
        ult     = row.get('ultimate_' + horizon)
        bm      = row.get('boni_mali_' + horizon)
        ecart   = row.get('ecart_pct_' + horizon)
        mature  = bool(row.get('mature', True))
        pct_dev = float(row.get('pct_developpe', 0))

        if not mature:
            # Immature — grisée et non évaluée
            html += (
                '<tr style="opacity:0.55;background:#F5F7FA;">'
                '<td class="label" style="color:var(--slate);">' + annee + '</td>'
                '<td class="right" style="color:var(--slate);"><span class="mono">' + _f(obs) + '</span></td>'
                '<td class="right" style="color:var(--slate);">—</td>'
                '<td class="right" style="color:var(--slate);">—</td>'
                '<td class="center" style="color:var(--slate);">—</td>'
                '<td class="center" style="font-size:8pt;color:var(--slate);">' + _pct(pct_dev) + '</td>'
                '<td class="center" style="font-size:8pt;color:var(--slate);font-style:italic;">Immature</td>'
                '</tr>'
            )
            continue

        # Mature — évaluée
        try:
            ecart_f = float(ecart) if ecart is not None else 0.0
            if abs(ecart_f) >= seuil_r:
                icon = '🔴'; row_bg = 'background:rgba(192,57,43,0.05);'
            elif abs(ecart_f) >= seuil_a:
                icon = '🟡'; row_bg = 'background:rgba(230,126,34,0.04);'
            else:
                icon = '✅'; row_bg = 'background:#fff;' if idx % 2 == 0 else 'background:#F8FAFB;'
        except Exception:
            icon = '—'; row_bg = ''

        html += (
            '<tr style="' + row_bg + '">'
            '<td class="label">' + annee + '</td>'
            '<td class="right"><span class="mono">' + _f(obs) + '</span></td>'
            '<td class="right"><span class="mono">' + _f(ult) + '</span></td>'
            '<td class="right"><span class="mono">' + _f(bm) + '</span></td>'
            '<td class="center">' + _pct(ecart) + '</td>'
            '<td class="center" style="font-size:8pt;color:var(--vert);">' + _pct(pct_dev) + '</td>'
            '<td class="center">' + icon + '</td>'
            '</tr>'
        )

    n_mat = sum(1 for r in data if r.get('mature', True))
    n_tot = len(data)
    n_imm = n_tot - n_mat
    note  = (' — ' + str(n_imm) + ' année(s) immature(s) non évaluée(s)') if n_imm > 0 else ''
    html += (
        '</tbody>'
        '<tfoot><tr style="background:var(--gold-pale);">'
        '<td colspan="7" style="padding:8px 14px;font-size:8pt;color:var(--text-mid);">'
        + str(n_mat) + ' année(s) mature(s) (≥ 75 %) évaluée(s) sur '
        + str(n_tot) + ' au total' + note
        + '</td></tr></tfoot>'
        '</table>'
    )
    return html
# =============================================================================
#  LA NUMÉROTATION DES FIGURES  (lot C3d)
# =============================================================================
#
#  ⚠️ LES IDENTIFIANTS TECHNIQUES NE SORTENT PAS DU CODE. `g1`, `g4`, `g9`,
#  `g15`… sautent : g7, g8, g11 et g13 ont été retirés au lot C3b. Un lecteur
#  de rapport n'a pas à subir l'histoire du dépôt, et « voir figure 7 » doit
#  pouvoir s'écrire dans le texte.
#
#  LE NUMÉRO EST POSITIONNEL, ET C'EST CE QUI GARANTIT LA CONTINUITÉ. Le
#  compteur n'avance QUE lorsqu'une figure est réellement rendue. Une figure
#  absente — sans exposition, `g10_h3` et `g15_exposition` ne se produisent
#  pas — ne laisse donc aucun trou : la suivante prend simplement son numéro.
#
#  ET L'ORDRE N'EST PAS TENU EN PARALLÈLE DU DOCUMENT, il EST le document :
#  chaque appel numérote à l'endroit exact où la figure s'insère. Une liste
#  d'ordre séparée dériverait au premier déplacement de section ; un compteur
#  posé dans le flux ne le peut pas.


def kaleido_disponible() -> bool:
    """Le module qui convertit une figure Plotly en image est-il là ?

    ⚠️ CE N'EST PAS UNE DÉPENDANCE COMME LES AUTRES. Un `.docx` ne sait porter
    que du raster : sans `kaleido`, plotly refuse `to_image` et AUCUNE figure
    ne peut entrer dans le Word. Trois issues étaient possibles ; celle
    retenue est la troisième :
      · l'exiger — l'export entier échouerait là où il manque. C'est le motif
        `weasyprint` que le lot C1 vient précisément de retirer ;
      · s'en passer — un rapport de provisionnement sans une seule figure,
        dans le format qui part chez un commissaire aux comptes ;
      · le rendre OPTIONNEL et NOMMER la dégradation. Le document porte alors
        la légende de chaque figure et la raison de son absence, et `run()`
        remonte `figures_word` dans `livrables_erreurs`.

    ⚠️ ET L'ÉPINGLE DE `requirements.txt` EST DATÉE, PAS CASSÉE — vérifié
    dans la source de plotly 6.5.2 : `kaleido` 0.2.1 fonctionne toujours,
    par le chemin « v0 », mais plotly annonce sa suppression (variable
    `ENGINE_SUPPORT_TIMELINE`, « September 2025 »). La version 1.0 qui le
    remplace exige Google Chrome sur la machine. Le choix entre les deux
    n'appartient pas à ce lot ; ne rien exiger, si.
    """
    import importlib.util
    return importlib.util.find_spec('kaleido') is not None


def rendre_image(figure) -> bytes:
    """Rasterise une figure Plotly. LE SEUL point du module qui le fasse.

    ⚠️ CETTE FONCTION EXISTE POUR ÊTRE SUBSTITUÉE PAR LES TESTS, et c'est un
    besoin de JUSTESSE, pas de vitesse. Mesuré avant ce lot : six tests
    faisaient rasteriser 94 figures — 562 s — parce que `kaleido` se trouvait
    installé sur la machine. Sans lui, les mêmes tests passaient en empruntant
    un AUTRE chemin. Une suite de tests qui n'exerce pas le même code selon la
    machine ne prouve pas la même chose selon la machine.

    Cinq de ces six tests ne parlent pas de figures — ils inspectent le
    vocabulaire publié, les marges de la page, les couleurs, le classeur
    Excel. Ils ont besoin qu'une image ENTRE dans le document, pas qu'elle
    soit belle : un rendeur substitué leur suffit, et le chemin réel reste
    vérifié par les deux tests qui, eux, portent sur les figures.
    """
    return figure.to_image(format='png', width=1100, height=520, scale=2)


class _NumeroteurFigures:
    """Numérote « Figure N » les figures présentes, dans l'ordre du document."""

    def __init__(self, graphiques_html):
        self._g = graphiques_html or {}
        self.n = 0
        self.rendues = []          # [(numero, cle)] — pour la trace et les tests

    def html(self, cle: str) -> str:
        """Le bloc HTML numéroté d'une figure, ou '' si elle n'existe pas."""
        contenu = self._g.get(cle) or ''
        if not contenu:
            return ''
        self.n += 1
        self.rendues.append((self.n, cle))
        titre = TITRES_FIGURES.get(cle, cle)
        return (
            '<div class="table-section-title">Figure ' + str(self.n)
            + ' — ' + titre + '</div>'
            '<div style="margin:12px 0;border:1px solid var(--border);'
            'border-radius:6px;overflow:hidden;">' + contenu + '</div>'
        )

    def legende(self, cle: str) -> str:
        """« Figure N — titre », sans le contenu : pour le Word."""
        self.n += 1
        self.rendues.append((self.n, cle))
        return 'Figure %d — %s' % (self.n, TITRES_FIGURES.get(cle, cle))


# =============================================================================
#  EXPORT HTML
# =============================================================================

def export_html(
    n1, n2, n3, n4,
    commentaire='', ref_client='', arrete='',
    audit_id='', lob_label='', graphiques=None,
    actuaire_nom='', actuaire_numero_ia='',
    narration=None, source_narration='',
) -> str:
    try:
        n1=n1 or {}; n2=n2 or {}; n3=n3 or {}; n4=n4 or {}

        dt      = datetime.now().strftime('%d/%m/%Y')
        # ⚠️ PLUS DE `arrete or dt` : la date du jour ne se fait plus
        # passer pour l'arrete. Voir ARRETE_ABSENT.
        arr     = arrete or ARRETE_ABSENT
        cli     = ref_client or 'À renseigner'
        lob     = _lob(lob_label or n2.get('lob_label', '') or n2.get('lob', ''))
        methode = n4.get('methode_facteurs', n2.get('methode_recommandee', '—'))
        statut  = n4.get('statut', 'AMBRE')

        # ⚠️⚠️ LA NARRATION EST RECUE SI ON LA DONNE, GENEREE SINON.
        # Mesure : un seul `run()` appelait `_generer_narration` DEUX fois --
        # une fois ici, une fois dans l'autre format. Tant que le texte
        # deterministe passait devant, les deux etaient identiques par
        # construction (12 699 caracteres, mesures a l'identique) et le
        # doublon etait inerte. Depuis que la narration LLM passe devant, ce
        # sont DEUX APPELS INDEPENDANTS : le HTML et le Word du meme dossier
        # porteraient deux textes differents, et le cout d'appel double.
        # Le repli `is None` garde intacts les appels qui ne la fournissent
        # pas -- l'application et une trentaine de tests.
        if narration is None:
            narration, source = _generer_narration(n2, n3, n4, commentaire,
                                                   lob, arr)
        else:
            source = source_narration

        # Graphiques Plotly — accepte go.Figure ou HTML pré-converti
        graphiques_html = {}
        if graphiques:
            for nom, fig_or_html in graphiques.items():
                try:
                    if isinstance(fig_or_html, str):
                        # Déjà converti en HTML (stockage session_state optimisé)
                        graphiques_html[nom] = fig_or_html
                    else:
                        # Objet go.Figure → convertir en HTML
                        import plotly.io as pio
                        graphiques_html[nom] = pio.to_html(
                            fig_or_html, full_html=False, include_plotlyjs=False,
                            config={'displayModeBar': False},
                        )
                except Exception as _eg:
                    logger.debug(f'Graphique {nom} ignoré : {_eg}')

        # Construire tous les blocs
        b = _build_blocks(n2, n3, n4, narration, source, lob, cli, arr, dt, audit_id, methode, statut, graphiques_html, actuaire_nom=actuaire_nom, actuaire_numero_ia=actuaire_numero_ia)

        # ⚠️ LE CATALOGUE ET LE RAPPORT COÏNCIDENT DEPUIS LE LOT C3d. Le HTML
        # portait CINQ figures sur quatorze, et les neuf absentes étaient
        # celles qui JUSTIFIENT la méthode — dont g17 et g18, les deux seuls
        # graphiques que le guide nomme. Un rapport qui montre la réponse et
        # cache le raisonnement est celui qu'on ne peut pas défendre.
        fig = _NumeroteurFigures(graphiques_html)

        # Assembler
        html = (
            '<!DOCTYPE html>\n<html lang="fr">\n<head>\n'
            '<meta charset="UTF-8">\n'
            '<title>Rapport Actuariel — ' + cli + ' — ' + arr + '</title>\n'
            '<script src="https://cdn.jsdelivr.net/npm/plotly.js-dist@2.26.0/plotly.min.js"></script>\n'
            + _css() +
            '\n</head>\n<body>\n<div class="rapport-container">\n\n'

            # ── PAGE DE GARDE ──────────────────────────────────────────────
            '<div class="page-garde">\n'
            '  <div class="garde-bg">\n'
            '    <div class="garde-dots"></div>\n'
            '    <div class="garde-diagonal"></div>\n'
            '    <div class="garde-accent-line"></div>\n'
            '  </div>\n'
            '  <div class="garde-inner">\n'
            '    <div class="garde-header">\n'
            '      <div class="garde-logo-wrap"><img src="' + LOGO_URI + '" alt="ActuarIA"/></div>\n'
            '      <div class="garde-badges">\n'
            '        <span class="badge-confidentiel">⬛ Confidentiel</span>\n'
            '      </div>\n'
            '    </div>\n'
            '    <div class="garde-hero">\n'
            '      <div class="garde-eyebrow">Rapport de Provisionnement Non-Vie</div>\n'
            '      <div class="garde-titre">' + b['garde_titre'] + '</div>\n'
            # ⚠️ « au » N'A DE SENS QUE DEVANT UNE DATE. Sans arrete, `arr` vaut
            # ARRETE_ABSENT : « Arrêté au non communiqué » serait fautif. Le
            # gabarit s'adapte au lieu d'imposer la preposition.
            '      <div class="garde-subtitle">Arrêté '
            + ('au ' + arr if arrete else arr) + '</div>\n'
            '      <div class="garde-sep">\n'
            '        <div class="garde-sep-line"></div>\n'
            '        <div class="garde-sep-diamond"></div>\n'
            '        <div class="garde-sep-line"></div>\n'
            '      </div>\n'
            '      <div class="garde-statut ' + b['garde_statut_cls'] + '">\n'
            '        <div class="statut-dot ' + b['garde_dot_cls'] + '"></div>\n'
            '        <div class="' + b['garde_label_cls'] + '">' + b['garde_label'] + '</div>\n'
            '      </div>\n'
            '    </div>\n'
            '    <div class="garde-footer">\n'
            '      <div class="garde-kpis">' + b['garde_kpis'] + '</div>\n'
            '    </div>\n'
            '  </div>\n'
            '</div>\n\n'

            # ── CORPS ──────────────────────────────────────────────────────
            '<div class="rapport-body">\n\n'

            # Section 1
            '<div class="section-header"><span class="section-num">01</span><span class="section-titre">Synthèse exécutive</span></div>\n'
            '<div class="section-body">\n'
            + b['kpi_grid']
            + fig.html('g5_convergence')
            + '\n</div>\n<div class="section-divider"></div>\n\n'

            # Section 2
            '<div class="section-header"><span class="section-num">02</span><span class="section-titre">Résultats par méthode actuarielle</span></div>\n'
            '<div class="section-body">\n'
            + b['tableau_methodes']
            + '<div class="table-section-title">Diagnostic — décomposition de l\'incertitude (outil analytique interne, non destiné au bilan)</div>\n'
            + b['tableau_incertitude']
            + (('<div class="table-section-title">Clark LDF (2003) — intervalle '
                'de prédiction, décomposé processus / paramètre</div>\n'
                + b['tableau_clark_ic']) if b.get('tableau_clark_ic') else '')
            + (('<div class="table-section-title">Munich Chain Ladder '
                '(Quarg &amp; Mack 2004) — rapprochement payé / engagé</div>\n'
                + b['tableau_munich']) if b.get('tableau_munich') else '')
            + (('<div class="table-section-title">Benktander (1976) — mélange '
                'par crédibilité, INFORMATIF (hors Best Estimate)</div>\n'
                + b['tableau_benktander']) if b.get('tableau_benktander') else '')
            + fig.html('g1_heatmap')
            + fig.html('g16_increments')
            + fig.html('g15_exposition')
            + fig.html('g2_cadences')
            + fig.html('g3_facteurs_cl')
            + fig.html('g4_reserve_annee')
            + fig.html('g6_bootstrap')
            + '\n</div>\n<div class="section-divider"></div>\n\n'

            # Section 3
            '<div class="section-header"><span class="section-num">03</span><span class="section-titre">Validation des hypothèses actuarielles</span></div>\n'
            '<div class="section-body">\n'
            '<div class="hyp-grid">' + b['hyp_cards'] + '</div>\n'
            # ⚠️ CETTE SECTION S'APPELLE « VALIDATION DES HYPOTHÈSES » ET NE
            # PORTAIT AUCUNE FIGURE. Les deux que le guide de l'Institut des
            # Actuaires nomme — §9.d.ii et §9.d.iii — sont ici, à leur place.
            + fig.html('g17_linearite')
            + fig.html('g18_residus')
            + fig.html('g9_h2')
            + fig.html('g10_h3')
            + b['methode_rec_box']
            + '\n</div>\n<div class="section-divider"></div>\n\n'

            # Section 4
            '<div class="section-header"><span class="section-num">04</span><span class="section-titre">SCR Provisions — Art. 115 Règlement délégué 2015/35</span></div>\n'
            '<div class="section-body">\n'
            + b['tableau_scr']
            + '\n</div>\n<div class="section-divider"></div>\n\n'

            # Section 5
            '<div class="section-header"><span class="section-num">05</span><span class="section-titre">Back-testing — Boni / Mali de liquidation</span></div>\n'
            '<div class="section-body">\n'
            + b['bt_summary']
            + '<div class="table-section-title">Tableau N-1 — Horizon un an</div>\n'
            + b['tableau_bt_n1']
            + '<div class="table-section-title" style="margin-top:20px;">Tableau N-2 — Horizon deux ans</div>\n'
            + b['tableau_bt_n2']
            + fig.html('g14_backtesting')
            + b['bt_note']
            + '\n</div>\n<div class="section-divider"></div>\n\n'

            # Section 6
            '<div class="section-header"><span class="section-num">06</span><span class="section-titre">Effets calendaires — GLM Poisson APC (Renshaw-Verrall 1998)</span></div>\n'
            '<div class="section-body">\n'
            '<div class="bz-grid">' + b['bz_items'] + '</div>\n'
            + b.get('bz_glm', '') + '\n'
            + b['bz_reco_box']
            + b.get('ptf_block', '')
            + '\n</div>\n<div class="section-divider"></div>\n\n'

            # Section 7 — COMMENTAIRE COMPLET
            '<div class="section-header"><span class="section-num">07</span><span class="section-titre">Commentaire actuariel</span></div>\n'
            '<div class="section-body">\n'
            '<div class="commentaire-wrap">\n'
            '  <div class="commentaire-header">\n'
            '    <span class="commentaire-header-title">' + b['commentaire_header_title'] + '</span>\n'
            '    <span class="commentaire-ai-badge">' + b['commentaire_badge'] + '</span>\n'
            '  </div>\n'
            '  <div class="commentaire-body">\n'
            '    <p style="font-size:8pt;color:var(--slate);margin-bottom:20px;font-style:italic;">'
            'Commentaire à destination du Conseil d\'Administration et de l\'Actuaire Désigné.<br>'
            'Document soumis à l\'ACPR dans le cadre du reporting Solvabilité 2.</p>\n'
            + b['narration_html']
            + '\n    <div class="comm-footer">✦ Narration générée par ActuarIA Intelligence · Agent A7 Ibrahim v5.0 · ' + dt + '</div>\n'
            '  </div>\n'
            '</div>\n'
            '</div>\n<div class="section-divider"></div>\n\n'

            # Section 8
            '<div class="section-header"><span class="section-num">08</span><span class="section-titre">Jugement actuariel &amp; Recommandations</span></div>\n'
            '<div class="section-body">\n'
            '<div class="hyp-grid">' + b['alertes'] + b['recommandations'] + '</div>\n'
            # Le guide exige que l'impact des retraitements soit étudié « dans
            # le rapport actuariel [...] via des tests de sensibilité ». Sa
            # place est donc auprès du jugement, pas dans les résultats.
            + fig.html('g12_sensibilites')
            + b['avis']
            + '\n</div>\n\n'

            '</div>\n\n'

            # ── PIED DE PAGE ───────────────────────────────────────────────
            '<div class="pied-de-page">\n'
            '  <div class="pied-logo"><img src="' + LOGO_URI + '" alt="ActuarIA"/></div>\n'
            '  <div class="pied-meta">'
            + b['signature_actuaire']
            + b['pied_info'] + '<br>'
            '    <span style="font-size:6.5pt;color:#8A9AB0;">'
            'Rapport établi conformément à l\'Art. 77 et 105 de la Directive Solvabilité II '
            'et au Guide Institut des Actuaires 2023</span><br>'
            '    <span class="confidentiel-footer">CONFIDENTIEL — USAGE STRICTEMENT ACTUARIEL</span>\n'
            '  </div>\n'
            '</div>\n\n'
            '</div>\n</body>\n</html>'
        )

        logger.info(f'HTML v5 : {len(html):,} chars — narration={source}')
        return html

    except Exception as e:
        logger.error(f'export_html échoué : {e}', exc_info=True)
        # Le repli porte un MARQUEUR reconnaissable. Sans lui, cette page d'erreur
        # est un HTML valide qu'un appelant ne distingue pas d'un rapport : c'est
        # ainsi qu'un import manquant a produit 88 octets pendant que la gate
        # entière passait au vert. `export_word`, qui appelle cette fonction,
        # refuse désormais de mettre en page un repli — il en aurait fait un PDF
        # volumineux et parfaitement valide de la page d'erreur.
        return (MARQUEUR_ECHEC_RAPPORT + '<html><body><h1>Erreur : '
                + str(e) + '</h1></body></html>')


# =============================================================================
#  EXPORT PDF
# =============================================================================

# ⚠️ `export_pdf` A ETE RETIREE (lot C1, decision B). Le PDF n'est plus
# GENERE : il s'obtient par CONVERSION du Word ou du HTML, ce qui supprime
# la dependance a weasyprint — absente de bien des environnements, et qui
# rendait `pdf_bytes` a zero octet sans que rien ne le dise. Le HTML est
# desormais une sortie a part entiere de `run()` : c'est lui le maitre
# d'impression, avec le Word.
# =============================================================================
#  EXPORT WORD
# =============================================================================

def export_word(n1, n2, n3, n4,
                commentaire='', ref_client='', arrete='',
                audit_id='', lob_label='', graphiques=None,
                actuaire_nom='', actuaire_numero_ia='',
                narration=None, source_narration='') -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        logger.error(f'python-docx absent : {e}'); return b''

    try:
        n1=n1 or {}; n2=n2 or {}; n3=n3 or {}; n4=n4 or {}

        def rgb(h):
            h = h.lstrip('#')
            return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))

        NR=rgb(NAVY); GR=rgb(GOLD); BR=rgb('#FFFFFF')
        GrR=rgb(SLATE); RgR=rgb(ROUGE); VR=rgb(VERT); AR=rgb(ORANGE)

        dt      = datetime.now().strftime('%d/%m/%Y')
        # ⚠️ PLUS DE `arrete or dt` : la date du jour ne se fait plus
        # passer pour l'arrete. Voir ARRETE_ABSENT.
        arr     = arrete or ARRETE_ABSENT
        cli     = ref_client or 'À renseigner'
        lob     = _lob(lob_label or n2.get('lob_label','') or n2.get('lob',''))
        methode = n4.get('methode_facteurs', n2.get('methode_recommandee','—'))
        statut  = n4.get('statut','AMBRE')
        # ⚠️⚠️ LA NARRATION EST RECUE SI ON LA DONNE, GENEREE SINON.
        # Mesure : un seul `run()` appelait `_generer_narration` DEUX fois --
        # une fois ici, une fois dans l'autre format. Tant que le texte
        # deterministe passait devant, les deux etaient identiques par
        # construction (12 699 caracteres, mesures a l'identique) et le
        # doublon etait inerte. Depuis que la narration LLM passe devant, ce
        # sont DEUX APPELS INDEPENDANTS : le HTML et le Word du meme dossier
        # porteraient deux textes differents, et le cout d'appel double.
        # Le repli `is None` garde intacts les appels qui ne la fournissent
        # pas -- l'application et une trentaine de tests.
        if narration is None:
            narration, source = _generer_narration(n2, n3, n4, commentaire,
                                                   lob, arr)
        else:
            source = source_narration

        mk  = n3.get('mack',{})
        sc  = n4.get('scr',{});           pw = n4.get('poids',{})
        bt  = n3.get('backtesting',{});   bz = n3.get('glm_apc',{})
        h1  = n2.get('h1_independance',{}); h2 = n2.get('h2_stabilite',{})

        BE  = float(n4.get('best_estimate',0) or 0)
        SIG = float(mk.get('sigma_total',0) or 0)
        P75 = float(n4.get('reserve_p75',0) or 0)
        P90 = float(n4.get('reserve_p90',0) or 0)
        P99 = float(n4.get('reserve_p99_5',0) or 0)
        # Le nom de la grandeur publiée, jamais « composé » en dur : la
        # bascule de N4 peut avoir retenu Mack recentré ou le Bootstrap.
        _appr_w = libelle_percentiles(n4)
        CV  = float(n4.get('cv_inter_methodes',0) or 0)
        SCP = _scr_publiable(sc)
        SCR = _ratio_scr(SCP, BE)
        # En pourcentage et à la décimale : deux segments ont un σ à trois
        # chiffres significatifs (protection juridique 5,5 %, crédit 17,2 %).
        SIG_EIOPA = _pct(float(sc.get('sigma_eiopa') or 0) * 100) if sc else '—'
        REF_S2 = _s(sc.get('reference_s2', 'Annexes II / XIV, Rgt 2015/35')
                    if sc else 'Annexes II / XIV, Rgt 2015/35')

        doc = Document()
        for s in doc.sections:
            s.top_margin=Cm(2); s.bottom_margin=Cm(2)
            s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

        def _bg(cell, hex6):
            tc=cell._tc; tcp=tc.get_or_add_tcPr()
            sd=OxmlElement('w:shd')
            sd.set(qn('w:fill'),hex6.lstrip('#'))
            sd.set(qn('w:color'),'auto'); sd.set(qn('w:val'),'clear')
            tcp.append(sd)

        def _run(p, txt, bold=False, italic=False, sz=10, col=None):
            r=p.add_run(str(txt)); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
            if col: r.font.color.rgb=col
            return r

        def _h(txt, lv=1, col=None):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
            sz={1:16,2:12,3:10}.get(lv,10)
            c=col or (NR if lv==1 else GR)
            _run(p,txt,bold=True,sz=sz,col=c)

        def _sep():
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
            bo=OxmlElement('w:bottom')
            bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'6')
            bo.set(qn('w:space'),'1'); bo.set(qn('w:color'),'C9A84C')
            pBdr.append(bo); pPr.append(pBdr)

        # ⚠️ LE SIGNAL RAG N'EXISTAIT PAS DANS LES TABLEAUX DU WORD, ET CE
        # SONT EUX QUI PORTENT LES VERDICTS. Mesuré sur un document généré :
        # 99 fonds de cellule, mais DEUX teintes seulement — le gris de
        # zébrage et le navy des en-têtes. Les dix-sept lignes du tableau des
        # hypothèses affichaient « VALIDÉE » et « NON VALIDÉE » dans
        # exactement la même encre noire. Depuis la décision B, le Word est le
        # MAÎTRE D'IMPRESSION du PDF : ce n'est plus de la cosmétique.
        #
        # LA CORRESPONDANCE EST EXACTE, PAS HEURISTIQUE. Le vocabulaire des
        # verdicts a été RELEVÉ sur huit exécutions — quatre triangles × avec
        # et sans exposition — et il est fini. Comparer le texte ENTIER de la
        # cellule à cette table ne peut pas produire de faux positif, là où un
        # seuil de longueur ou une recherche de sous-chaîne en produirait :
        # la colonne « Message » contient « BFCC-H4 VALIDÉE — loss ratio… »,
        # qui ne doit pas repeindre toute la cellule.
        #
        # ET LA COULEUR N'EST JAMAIS SEULE PORTEUSE. Le mot reste écrit :
        # imprimé en noir et blanc, ou lu par quelqu'un qui distingue mal le
        # rouge du vert, le document dit toujours « NON VALIDÉE ». La couleur
        # accélère la lecture, elle ne la remplace pas.
        _RAG_CELLULE = {
            'VALIDÉE':        (VR,  'EAF3DE'),
            'À JUSTIFIER':    (AR,  'FAEEDA'),
            'NON VALIDÉE':    (RgR, 'FCEBEB'),
            'REJETÉE':        (RgR, 'FCEBEB'),
            'NON TESTABLE':   (GrR, None),
            'Non disponible': (GrR, None),
            'non calculée':   (GrR, None),
            'VERT':           (VR,  'EAF3DE'),
            'AMBRE':          (AR,  'FAEEDA'),
            'ROUGE':          (RgR, 'FCEBEB'),
            '✓ Inclus':       (VR,  'EAF3DE'),
            '⊘ Exclu':        (RgR, 'FCEBEB'),
            '✅ OUI':          (VR,  'EAF3DE'),
            '❌ NON':          (RgR, 'FCEBEB'),
        }

        def _tbl(heads, rows, ws=None):
            t=doc.add_table(rows=1+len(rows), cols=len(heads)); t.style='Table Grid'
            for i,hd in enumerate(heads):
                c=t.rows[0].cells[i]; _bg(c,'0B1E3D')
                pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=pp.add_run(str(hd)); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=BR
            for ri,row in enumerate(rows):
                for ci,v in enumerate(row):
                    c=t.rows[ri+1].cells[ci]
                    txt = str(v) if v is not None else '—'
                    rag = _RAG_CELLULE.get(txt.strip())
                    if rag and rag[1]:
                        _bg(c, rag[1])
                    elif ri%2==1: _bg(c,'EEF2F7')
                    pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r=pp.add_run(txt); r.font.size=Pt(9)
                    if rag:
                        r.font.color.rgb = rag[0]; r.bold = True
            if ws:
                for i,w in enumerate(ws):
                    for row in t.rows: row.cells[i].width=Cm(w)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)

        # Page de garde
        try:
            import cairosvg
            logo_png = cairosvg.svg2png(bytestring=LOGO_SVG.encode(), output_width=280)
            doc.add_picture(io.BytesIO(logo_png), width=Cm(7))
        except Exception:
            p=doc.add_paragraph()
            _run(p,'Actuar',bold=True,sz=26,col=NR); _run(p,'IA',bold=True,sz=26,col=GR)

        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p,'RAPPORT DE PROVISIONNEMENT NON-VIE\n',bold=True,sz=20,col=NR)
        _run(p,lob,bold=True,sz=14,col=GR)
        doc.add_paragraph()
        s_col_r = VR if statut=='VERT' else AR if statut=='AMBRE' else RgR
        p=doc.add_paragraph()
        _run(p,'Statut : ',sz=11,col=NR); _run(p,statut,bold=True,sz=11,col=s_col_r)
        doc.add_paragraph()
        _tbl(['Client','Branche','Méthode','Arrêté'],
             [[cli,lob,methode.replace('_',' ').title(),arr]],ws=[3.5,4.0,4.0,2.5])
        doc.add_page_break()

        # ⚠️ LE WORD NE PORTAIT AUCUNE FIGURE (lot C3d). Zéro image mesurée
        # sur 45 040 octets, alors que le HTML en portait cinq et
        # l'application quatorze : le livrable qui VOYAGE était le plus
        # pauvre des trois.
        fig = _NumeroteurFigures(graphiques or {})
        _kaleido = kaleido_disponible()

        def _figure(cle):
            """Insère une figure numérotée, ou nomme la raison de son absence."""
            objet = (graphiques or {}).get(cle)
            if objet is None:
                return
            legende = fig.legende(cle)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            _run(p, legende, bold=True, sz=9, col=GR)
            if not _kaleido:
                q = doc.add_paragraph()
                q.paragraph_format.space_after = Pt(6)
                _run(q, "Figure non rendue : le module « kaleido » n'est pas "
                        "installé sur cette machine. Le rapport HTML la "
                        "porte.", sz=8, col=GR)
                return
            try:
                png = rendre_image(objet)
                doc.add_picture(io.BytesIO(png), width=Cm(16.5))
            except Exception as _ef:
                q = doc.add_paragraph()
                q.paragraph_format.space_after = Pt(6)
                _run(q, "Figure non rendue : %s." % type(_ef).__name__,
                     sz=8, col=GR)

        _h('1. Synthèse exécutive'); _sep()
        _tbl(['Indicateur','Valeur','Indicateur','Valeur'],
             [['Best Estimate (brut)',_f(BE),'σ Mack total',_f(SIG)],
              [f'P75 ({_appr_w})',_f(P75),'CV inter-méthodes',_pct(CV)],
              [f'P90 ({_appr_w})',_f(P90),'SCR Provisions',_f(SCP)],
              [f'P99.5 ({_appr_w})',_f(P99),'Ratio SCR/BE',_pct(SCR)]],ws=[4.5,3.5,4.5,3.5])

        _h("Diagnostic — décomposition de l'incertitude (outil analytique interne, non destiné au bilan)"); _sep()
        _tbl(['Approche','P90','σ','Centre'],
             [[marque_retenue(n4,CLE_COMPOSE,'Incertitude composée'),_f(n4.get('reserve_p90_compose',0) or 0),_f(n4.get('sigma_total_compose',SIG)),'BE pondéré'],
              [marque_retenue(n4,CLE_MACK,'Mack recentré'),_f(n4.get('reserve_p90_mack',P90)),_f(n4.get('sigma_mack',SIG)),'BE pondéré'],
              ['Mack natif',_f(mk.get('reserve_p90',0)),_f(mk.get('sigma_total',SIG)),'réserve Mack'],
              [marque_retenue(n4,CLE_BOOT,'Bootstrap ODP'),_f(n3.get('bootstrap',{}).get('p90') or 0),_f(n3.get('bootstrap',{}).get('std_bootstrap') or 0),'réserve Bootstrap']],ws=[5.0,3.0,3.0,3.0])
        # ⚠️ L'ORDRE DES FIGURES EST LE MÊME QUE DANS LE HTML, ET C'EST
        # VOLONTAIRE : « Figure 7 » doit désigner la même chose dans les deux
        # formats. Le compteur est positionnel dans chacun d'eux.
        _figure('g5_convergence')
        doc.add_page_break()

        _h('2. Résultats par méthode actuarielle'); _sep()
        _tbl(['Méthode','Réserve IBNR','Poids BE','Score','Statut'],
             # ⚠️ DEUX MENSONGES ICI, PAS UN (lot C3a). Le « 0 € » d'une
             # méthode non calculable, et surtout un « ✓ Inclus » ÉCRIT EN
             # DUR : sans exposition, le Word affirmait que
             # Bornhuetter-Ferguson était incluse au Best Estimate alors
             # qu'elle en était exclue. Le statut se déduit du poids, comme
             # dans le HTML.
             [[libelle(_m),
               (_f(reserve(n3, _m)) if reserve(n3, _m) is not None
                else motif_exclusion(n4, _m)),
               _pct(pw.get(_m, 0)*100), '—',
               ('σ (volatilité)' if _m == 'mack'
                else '✓ Inclus' if pw.get(_m, 0) > 0 else '⊘ Exclu')]
              for _m in ORDRE_AFFICHAGE] + [
              ['BEST ESTIMATE (brut)',_f(BE),'100 %','—','→ A10 (actualisation)']],ws=[4.5,3.5,2.5,2.5,3.0])

        # Benktander — MÊME SOURCE que le HTML et l'Excel, comme Munich.
        _lg_gb = lignes_benktander_rapport(n3, n4)
        if _lg_gb:
            _h("Benktander (1976) — mélange par crédibilité, INFORMATIF")
            _tbl(['Élément', 'Valeur'],
                 [[_s(k), _s(v)] for k, v in _lg_gb], ws=[5.0, 11.0])

        # Munich CL — MÊME SOURCE que le HTML et l'Excel. Le lot Clark avait
        # commencé par ne câbler que le HTML ; Munich n'était dans aucun des
        # trois. Les trois lisent désormais `lignes_munich_rapport`.
        _lignes_mcl = lignes_munich_rapport(n3)
        if _lignes_mcl:
            _h('Munich Chain Ladder (Quarg & Mack 2004) — rapprochement payé / engagé',
               lv=2)
            _tbl(['Indicateur', 'Valeur', 'Lecture'],
                 [[a, v, c] for a, v, c in _lignes_mcl], ws=[5.0, 3.5, 7.5])

        for _cle in ('g1_heatmap', 'g16_increments', 'g15_exposition',
                     'g2_cadences', 'g3_facteurs_cl', 'g4_reserve_annee',
                     'g6_bootstrap'):
            _figure(_cle)

        doc.add_page_break()

        _h('3. Validation des hypothèses actuarielles'); _sep()
        rows_h=[]
        for lbl,h in [('H1 Indépendance',h1),('H2 Stabilité',h2)]:
            if not h: continue
            # ⚠️ TROISIEME LIVRABLE, MEME DEFAUT. Le Word remis au CAC affichait
            # « VALIDÉE | 80/100 » sur une hypothese que le module declarait non
            # testable. Le statut vient desormais du module.
            # ⚠️ La troncature du message a 80 caracteres reste EN PLACE et
            # NOMMEE : c'est la classe (3) du releve B2, non arbitree. Le bloc
            # juste en dessous a deja retire la sienne pour ce motif.
            _st = str(h.get('statut',
                            'VALIDÉE' if h.get('ok', True) else 'REJETÉE'))
            _sc = '—' if _st == 'NON TESTABLE' else str(h.get('score', '—')) + '/100'
            rows_h.append([lbl, _st, _sc, str(h.get('message', ''))[:80]])
        # ⚠️ PLUS DE TRONCATURE. Les messages étaient coupés à 80 caractères et
        # les libellés à 38 : sur les 14 lignes d'un dossier avec Munich, 14
        # messages et 13 libellés étaient rognés en pleine phrase (médiane 170
        # caractères, maximum 532). Un rapport réglementaire peut se permettre
        # une ligne haute ; il ne peut pas se permettre une justification
        # coupée à « Cette hypothèse n'es ». La colonne Score, toujours « — »
        # pour ces lignes, cède sa largeur au message.
        # CLM en tête, comme dans le HTML : ces quatre-là gouvernent les
        # méthodes principales et n'atteignaient aucun livrable. La puissance
        # suit le message — le verdict seul ne dit pas ce qu'on pouvait voir.
        for ligne in (list(lignes_hypotheses_clm(n2))
                      + list(lignes_hypotheses_bfcc(n2))
                      + list(lignes_hypotheses_bootstrap(n2))
                      + list(lignes_hypotheses_munich(n2))):
            texte = ligne['message']
            if ligne.get('puissance_phrase'):
                texte = texte + ' ' + ligne['puissance_phrase']
            rows_h.append([ligne['libelle'], ligne['statut'], '—', texte])
        if rows_h: _tbl(['Hypothèse','Résultat','Score','Message'],rows_h,ws=[4.5,2.5,1.2,7.8])
        # Les deux graphiques que le guide nomme (§9.d.ii et §9.d.iii) sont
        # dans la section qui s'appelle « Validation des hypothèses ».
        _figure('g17_linearite')
        _figure('g18_residus')
        _figure('g9_h2')
        _figure('g10_h3')
        doc.add_page_break()

        _h('4. SCR Provisions — Art. 115 Rgt délégué (UE) 2015/35'); _sep()
        _tbl(['Composante','Valeur','Référence'],
             [['Best Estimate (brut)',_f(BE),'Art. 77 — avant actualisation'],
              ['Facteur σ — risque de réserve',SIG_EIOPA,REF_S2],
              ['SCR Provisions',_f(SCP),'3 × σ × BE — Art. 115'],
              ['Ratio SCR/BE',_pct(SCR),'< 35 %']],ws=[4.5,3.5,8.0])
        doc.add_page_break()

        _h('5. Back-testing — Boni / Mali de liquidation'); _sep()
        bt_s = str(bt.get('statut','—')); bt_sc = str(bt.get('score_qualite','—'))
        p=doc.add_paragraph()
        s_bt = VR if bt_s=='VERT' else AR if bt_s=='AMBRE' else RgR
        _run(p,'Statut : ',sz=9,col=NR); _run(p,bt_s,bold=True,sz=9,col=s_bt)
        _run(p,' | Score : '+bt_sc+'/100',sz=9,col=NR)
        _run(p,' | N-1 : '+str(bt.get('n_rouge_n1',0))+' rouge / '+str(bt.get('n_ambre_n1',0))+' ambre',sz=9,col=NR)
        _run(p,' | N-2 : '+str(bt.get('n_rouge_n2',0))+' rouge / '+str(bt.get('n_ambre_n2',0))+' ambre',sz=9,col=NR)
        _figure('g14_backtesting')
        doc.add_page_break()

        _h('6. Effets calendaire — GLM Poisson APC (Renshaw-Verrall 1998)'); _sep()
        # ⚠️⚠️ LE FAUX DE F3, TRANSPOSE AU WORD. Cette section affichait
        # << ✅ Aucun effet calendaire significatif >>, EN VERT, des que le
        # compteur valait zero -- SANS regarder si le test avait tourne. Sur
        # un triangle ou le GLM ne s'ajuste pas, le document remis au CAC
        # affirmait une absence d'effet que rien n'avait testee. Le HTML
        # distinguait deja les deux (FIX 2) ; le Word non. L'etat vient
        # desormais de `etat_calendaire`, source partagee.
        _etat_cal_w = etat_calendaire(bz)
        if _etat_cal_w == 'indisponible':
            p=doc.add_paragraph()
            _run(p, _PHRASE_CALENDAIRE['indisponible']+'.', sz=9, col=GrR)
        elif _etat_cal_w == 'diffus':
            p=doc.add_paragraph()
            _run(p, '⚠️ '+_PHRASE_CALENDAIRE['diffus']+'.', sz=9, col=AR)
        elif _etat_cal_w == 'aucun':
            p=doc.add_paragraph()
            _run(p, '✅ '+_PHRASE_CALENDAIRE['aucun']+'.', sz=9, col=VR)
        else:
            for e in bz.get('effets_calendaire',[]):
                if e.get('significatif'):
                    e_l=str(e.get('annee_label','—')); e_a=round(float(e.get('amplitude_pct',0)),1); e_n=str(e.get('niveau','—'))
                    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
                    _run(p,f'⚠️ {e_l} : {e_a:+.1f}% ({e_n})',sz=9,col=AR)
        doc.add_page_break()

        _h('7. Commentaire actuariel'); _sep()
        if narration:
            sections_n = re.split(r'(?=§\d+\s*[—\-–])', _clean(narration))
            for sec in sections_n:
                sec=sec.strip()
                if not sec: continue
                ls=sec.split('\n',1)
                if ls[0]: _h(ls[0],lv=2)
                if len(ls)>1:
                    for ln in ls[1].split('\n'):
                        ln=ln.strip()
                        if ln:
                            p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(3)
                            p.paragraph_format.left_indent=Cm(0.3)
                            _run(p,ln,sz=9,col=NR)
            # ⚠️ CE `if source=='claude_api'` LAISSAIT LE WORD MUET. Le chemin
            # `templates` — le SEUL qui s'exécute sans clé API — ne publiait
            # NI origine NI phrase d'engagement. Mesuré : « Cette narration
            # engage l'actuaire signataire. » comptait 1 en HTML et 0 en Word,
            # dans le format que l'actuaire envoie au commissaire aux comptes.
            # Le badge vient désormais de `badge_narration`, comme l'HTML.
            _badge = badge_narration(source)
            if _badge:
                p=doc.add_paragraph()
                _run(p,_badge,sz=7,italic=True,col=GrR)
        else:
            p=doc.add_paragraph(); _run(p,'Narration non disponible.',sz=9,italic=True)
        doc.add_page_break()

        _h('8. Jugement actuariel & Recommandations'); _sep()
        for a in n4.get('alertes',n2.get('alertes',[])):
            at=_clean(str(a))
            if at:
                p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
                _run(p,'⚠️  ',sz=10,col=AR); _run(p,at,sz=9,col=NR)
        for r in n4.get('recommandations',[]):
            rt=_clean(str(r))
            if rt:
                p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
                _run(p,'→  ',bold=True,sz=10,col=GR); _run(p,rt,sz=9,col=NR)
        avis_w=_clean(n4.get('avis_actuariel',''))
        if avis_w:
            doc.add_paragraph()
            p=doc.add_paragraph()
            # ⚠️ MEME DEFAUT QU'EN HTML, MEME REMEDE. `avis_actuariel` ne dit
            # JAMAIS << DEFAVORABLE >> : l'avis d'un dossier ROUGE sortait en
            # VERT dans le document signe. La couleur vient du STATUT, par la
            # meme table que le bandeau de garde -- `s_col_r`, deja calcule
            # plus haut sur ce meme statut.
            _run(p,avis_w,sz=10,bold=True,col=s_col_r)

        _figure('g12_sensibilites')

        _sep()
        # ⚠️ CE PIED NE PORTAIT NI SIGNATURE NI IDENTIFIANT D'AUDIT. Les trois
        # parametres `actuaire_nom`, `actuaire_numero_ia` et `audit_id`
        # etaient DECLARES dans la signature de `export_word` et lus nulle
        # part : 0 occurrence de << actuaire >> sur ses 375 lignes. L'HTML,
        # lui, imprimait les deux. Le document qui VOYAGE et qu'on SIGNE
        # etait le seul des deux a ne rien porter.
        # ⚠️ UN SEUL PARAGRAPHE, DEUX PASSAGES : la mention de relecture porte
        # sa propre couleur (ambre quand elle alerte) sans qu'un second
        # paragraphe centre soit necessaire.
        _tr = trace_relecture(actuaire_nom, actuaire_numero_ia)
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p, _tr.texte, sz=7, italic=True,
             col=AR if _tr.alerte else GrR).add_break()
        _run(p,'ActuarIA · '+cli+' · '+lob+' · Arrêté '+arr
             +' · Audit ID : '+(audit_id or '—')+' · '+dt+' · CONFIDENTIEL',
             sz=7,italic=True,col=GrR)

        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        wb=buf.read()
        logger.info(f'Word : {len(wb):,} bytes')
        return wb

    except Exception as e:
        logger.error(f'export_word : {e}', exc_info=True)
        return b''
