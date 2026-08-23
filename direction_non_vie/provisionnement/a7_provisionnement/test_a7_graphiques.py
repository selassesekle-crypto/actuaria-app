# -*- coding: utf-8 -*-
"""
=============================================================================
 A7 Ibrahim — le filet des graphiques (C2 · triage C3b · ajouts C3c)
=============================================================================

 ⚠️ POURQUOI CE FICHIER EXISTE.

 Les 14 graphiques d'A7 n'étaient épinglés par RIEN. Aucun test ne vérifiait
 qu'une courbe trace bien la grandeur que son étiquette annonce. Un graphique
 faux ne casse aucun test : il s'affiche, il est beau, et il ment.

 L'audit du lot C2 a recalculé chaque série INDÉPENDAMMENT depuis `C`, `n2`,
 `n3`, `n4`, puis l'a comparée à ce qui est réellement tracé. Ce fichier fige
 ce travail pour qu'il ne se reperde jamais : chaque verrou ci-dessous
 RECALCULE la série attendue depuis la source, il ne relit pas la figure.

 CE FILET N'EST PAS UN CONSTAT DE CONFORMITÉ GÉNÉRALE. L'audit a trouvé trois
 défauts, écrits ici comme des assertions de ce qui DEVRAIT être vrai et
 marquées `@unittest.expectedFailure`. Le marqueur se démonte tout seul quand
 le défaut est CORRIGÉ : le test réussit alors qu'on attendait un échec, et
 unittest fait ÉCHOUER la campagne (vérifié : `wasSuccessful()` rend False sur
 un `unexpectedSuccess`).

 ⚠️ ET IL NE SE DÉMONTE PAS QUAND LE SUJET EST SUPPRIMÉ — mesuré : un
 `KeyError` passe pour l'échec attendu, en silence. Les trois défauts ont
 disparu de trois façons différentes, et c'est instructif :
   · g5 et g11 CORRIGÉS au lot C3a — les marqueurs sont tombés d'eux-mêmes ;
   · g7 SUPPRIMÉ au lot C3b — c'est le verrou de CATALOGUE qui a échoué et
     forcé le retrait du test devenu orphelin.
 Un marqueur ne suffit donc pas : il lui faut un verrou de périmètre à côté.
=============================================================================
"""

import io
import re
import struct
import unittest
import zlib

import numpy as np

from direction_non_vie.provisionnement.a7_provisionnement.agent import (
    AgentA7Provisionnement)
from direction_non_vie.provisionnement.a7_provisionnement import (
    n5_rapport as _RAP)
from direction_non_vie.provisionnement.a7_provisionnement.n5_graphiques import (
    TITRES_FIGURES)
from direction_non_vie.provisionnement.a7_provisionnement.test_a7_ibrahim import (
    GENINS, _TRI_RECOURS_FORT)


# =============================================================================
#  LE RENDEUR SUBSTITUÉ — pour que la gate ne dépende pas d'un paquet
# =============================================================================
#
#  ⚠️ MESURÉ AVANT CE LOT : six tests faisaient rasteriser 94 figures, 562 s,
#  parce que `kaleido` se trouvait installé sur la machine. Sans lui, les
#  MÊMES tests passaient — en empruntant un autre chemin. Une suite qui
#  n'exerce pas le même code selon le poste ne prouve pas la même chose selon
#  le poste ; c'est un défaut de justesse, la lenteur n'en est que le symptôme
#  visible.
#
#  ⚠️ ET LE RENDEUR NE PEUT DÉPENDRE D'AUCUN PAQUET OPTIONNEL, sinon on
#  remplace une dépendance à `kaleido` par une dépendance à `Pillow` — le
#  patron d'origine en avait une, et il SAUTAIT là où Pillow manque. Ce PNG
#  est écrit avec `zlib` et `struct` de la bibliothèque standard : 263 octets,
#  aucune importation à faire.

def png_de_test(largeur: int, hauteur: int, rgb=(11, 30, 61)) -> bytes:
    """Un PNG valide, uni, sans aucune dépendance."""
    def _bloc(typ, donnees):
        return (struct.pack('>I', len(donnees)) + typ + donnees
                + struct.pack('>I', zlib.crc32(typ + donnees) & 0xFFFFFFFF))

    entete = struct.pack('>IIBBBBB', largeur, hauteur, 8, 2, 0, 0, 0)
    ligne = b'\x00' + bytes(rgb) * largeur        # filtre 0, puis les pixels
    return (b'\x89PNG\r\n\x1a\n' + _bloc(b'IHDR', entete)
            + _bloc(b'IDAT', zlib.compress(ligne * hauteur, 9))
            + _bloc(b'IEND', b''))


class rendeur_substitue:
    """Remplace le rasteriseur de `n5_rapport` le temps d'un bloc `with`.

    ⚠️ LES IMAGES DOIVENT ÊTRE DISTINCTES. OOXML déduplique les médias
    identiques : mesuré, quatorze PNG identiques ne laissent qu'UN fichier
    dans `word/media/`. Un test qui compte les médias diagnostiquerait alors
    un bug qui n'existe pas. Chaque appel rend donc une image d'une taille
    différente.

    `appels` conserve les figures reçues : un test peut vérifier QUE le
    rendeur a été sollicité, et combien de fois.
    """

    def __init__(self):
        self.appels = []

    def __enter__(self):
        self._vrai = _RAP.rendre_image
        def _faux(figure):
            self.appels.append(figure)
            n = len(self.appels)
            return png_de_test(200 + 3 * n, 100 + n)
        _RAP.rendre_image = _faux
        return self

    def __exit__(self, *_):
        _RAP.rendre_image = self._vrai
        return False


class kaleido_declare:
    """Force la réponse du prédicat `kaleido_disponible`, dans les DEUX sens.

    ⚠️ LE TEST DÉCIDE DU CHEMIN QU'IL EXERCE, PAS LA MACHINE. Avant ce lot,
    le chemin dégradé n'était atteint que sur un poste sans `kaleido` et le
    chemin nominal que sur un poste avec : chaque machine n'en voyait qu'un.
    Les deux sont désormais exercés partout.
    """

    def __init__(self, present: bool):
        self._present = present

    def __enter__(self):
        self._vrai = _RAP.kaleido_disponible
        _RAP.kaleido_disponible = lambda: self._present
        return self

    def __exit__(self, *_):
        _RAP.kaleido_disponible = self._vrai
        return False


#: Le catalogue, dans l'ordre de `generer_graphiques`.
#:
#: ⚠️ ILS ÉTAIENT QUATORZE JUSQU'AU LOT C3b, LE TRIAGE LES A RAMENÉS À DIX.
#: g7 (donut SCR) et g13 (le triangle une troisième fois) sont retirés, g8
#: devient un tableau, et g11 est fondu dans g4. C'est ce verrou-ci qui a
#: forcé la mise à jour du fichier : `expectedFailure` avale une SUPPRESSION
#: en silence (mesuré — un `KeyError` y passe sans bruit), donc le marqueur
#: de g7 ne pouvait pas se démonter tout seul. Le catalogue, lui, tombe.
#: ⚠️ LE LOT C3c EN AJOUTE QUATRE, ET LE VERROU EST TOMBÉ AU PREMIER — c'est
#: exactement ce qu'on lui demande. L'ordre suit la LECTURE et non la
#: numérotation : les numéros portent l'histoire du dépôt, pas le plan.
CATALOGUE = (
    # données
    'g1_heatmap', 'g16_increments', 'g15_exposition',
    # méthode
    'g2_cadences', 'g3_facteurs_cl',
    # validation des hypothèses — guide IA 2023, §9.d
    'g17_linearite', 'g18_residus', 'g9_h2', 'g10_h3',
    # résultat, incertitude, sensibilités, back-testing
    'g4_reserve_annee', 'g5_convergence', 'g6_bootstrap',
    'g12_sensibilites', 'g14_backtesting',
)

_CACHE = {}


#: L'exposition de référence des verrous — croissante, donc discriminante :
#: une exposition plate ne dirait rien d'une chronique.
EXPOSITION = np.arange(1, 11) * 4e5 + 3e6


def _run(avec_primes):
    """Un run complet, mis en cache — l'agent est cher, on ne le lance que 2×."""
    if avec_primes in _CACHE:
        return _CACHE[avec_primes]
    C = np.asarray(GENINS, dtype=float)
    kw = {}
    if avec_primes:
        kw['primes'] = EXPOSITION[:C.shape[0]]
    _CACHE[avec_primes] = AgentA7Provisionnement(verbose=False).run(
        source=C, mode_declare='cumule', generer_graphiques=True,
        generer_word=False, n_sim_bootstrap=60, seed=42, **kw)
    return _CACHE[avec_primes]


def _run_triangle(triangle, **kw):
    """Un run sur un AUTRE triangle — GenIns ne porte aucune reprise."""
    src = np.asarray(triangle, dtype=float)
    return AgentA7Provisionnement(verbose=False).run(
        source=src, mode_declare='cumule', generer_graphiques=True,
        generer_word=False, n_sim_bootstrap=60, seed=42, **kw)


def _figures(avec_primes=True):
    r = _run(avec_primes)
    g = r.get('graphiques') or {}
    if not g:
        raise unittest.SkipTest('plotly absent — aucun graphique produit')
    return g, np.asarray(GENINS, dtype=float), r['n2'], r['n3'], r['n4']


def _tr(fig, nom):
    """La trace nommée `nom` — exacte d'abord, sinon par préfixe.

    Le préfixe est nécessaire parce que g5 suffixe ses noms du poids retenu
    (« Chain Ladder (53%) ») ; l'exact passe en premier pour que « 1 » ne
    puisse pas attraper « 10 ».
    """
    for t in fig.data:
        if str(t.name or '') == nom:
            return t
    for t in fig.data:
        if str(t.name or '').startswith(nom):
            return t
    raise AssertionError(
        'trace %r absente ; présentes : %s'
        % (nom, [str(t.name) for t in fig.data]))


def _f(seq):
    """Une suite de flottants. `is None` et non `or ()` : un tableau numpy
    vide ou à plusieurs éléments n'a pas de valeur de vérité."""
    return [] if seq is None else [float(v) for v in seq]


# =============================================================================
#  T1 — LE CATALOGUE EST CLOS
# =============================================================================

class T1_Le_Catalogue_Est_Clos(unittest.TestCase):
    """Aucun graphique ne peut apparaître ni disparaître sans passer ici.

    Le verrou a déjà servi deux fois : au lot C3b quand le triage a fait
    tomber quatre graphiques, et au lot C3c quand quatre autres sont arrivés.
    """

    def test_le_catalogue_est_produit_en_entier_et_rien_d_autre(self):
        g, *_ = _figures()
        self.assertEqual(
            tuple(sorted(g)), tuple(sorted(CATALOGUE)),
            'catalogue modifié — épingle le nouveau graphique à sa source '
            'avant de l\'ajouter')
        print('    OK C2-0 les %d graphiques, et seulement eux'
              % len(CATALOGUE))

    def test_seuls_les_graphiques_d_exposition_peuvent_manquer(self):
        """⚠️ LE BON COMPORTEMENT FACE À UNE DONNÉE ABSENTE.

        Sans exposition, Bornhuetter-Ferguson n'existe pas et la chronique
        n'a rien à tracer. `g10_h3` et `g15_exposition` ne se rabattent PAS
        sur zéro : ils ne se produisent pas du tout, et l'orchestrateur les
        journalise en « pas de données ». C'était le contre-exemple de g5 et
        g11 avant le lot C3a ; c'est désormais le motif partagé,
        `methodes_be.disponible`.
        """
        g, _, _, n3, _ = _figures(avec_primes=False)
        self.assertFalse(n3['bf']['disponible'])
        self.assertEqual(sorted(set(CATALOGUE) - set(g)),
                         ['g10_h3', 'g15_exposition'],
                         'la liste des graphiques absents sans exposition a '
                         'changé')
        print('    OK C2-0b sans exposition : %d/%d — g10 et g15 s\'effacent '
              'au lieu de tracer zéro' % (len(g), len(CATALOGUE)))


# =============================================================================
#  T2 — CHAQUE GRAPHIQUE EST ÉPINGLÉ À SA SOURCE
# =============================================================================

class T2_Chaque_Graphique_Est_Epingle(unittest.TestCase):

    def test_g1_la_heatmap_est_le_triangle(self):
        g, C, _, _, _ = _figures()
        z = np.asarray(g['g1_heatmap'].data[0].z, dtype=float)
        self.assertEqual(z.shape, C.shape)
        n, m = C.shape
        connu = np.zeros(C.shape, dtype=bool)
        for i in range(n):
            connu[i, :min(n - i - 1, m - 1) + 1] = True
        np.testing.assert_allclose(z[connu], C[connu], rtol=1e-9)
        self.assertTrue(
            np.isnan(z[~connu]).all(),
            'une case FUTURE porte une valeur : le triangle brut la stocke à '
            'zéro, la heatmap doit la masquer et non peindre « 0 € payé »')
        print('    OK C2-1 g1 : %d cases connues == le triangle, %d cases '
              'futures masquées' % (int(connu.sum()), int((~connu).sum())))

    def test_g2_les_cadences_sont_la_fraction_developpee(self):
        g, C, _, n3, _ = _figures()
        pct = _f(n3['chain_ladder']['pct_developpe'])
        n, m = C.shape
        verifiees = 0
        for i in range(n):
            k = min(n - i - 1, m - 1)
            if C[i, k] <= 0:
                continue
            att = [round(C[i, j] / C[i, k] * pct[i], 4)
                   for j in range(k + 1) if C[i, j] > 0]
            obs = _f(_tr(g['g2_cadences'], 'An. %d' % i).y)
            self.assertEqual([round(v, 4) for v in obs], att,
                             'année %d : la courbe n\'est pas C[i,j]/C[i,k] '
                             '× pct_developpe[i]' % i)
            verifiees += 1
        self.assertGreaterEqual(verifiees, 9)
        y = _f(_tr(g['g2_cadences'], 'An. 0').y)
        self.assertLessEqual(max(y), 1.0001,
                             'une fraction développée dépasse 100 %')
        print('    OK C2-2 g2 : %d courbes == C[i,j]/C[i,k] × pct_developpe[i]'
              % verifiees)

    def test_g3_les_barres_sont_les_facteurs_chain_ladder(self):
        g, _, _, n3, _ = _figures()
        att = _f(n3['chain_ladder']['facteurs'])
        obs = _f(_tr(g['g3_facteurs_cl'], 'Facteurs CL').y)
        np.testing.assert_allclose(obs, att, rtol=1e-9)
        print('    OK C2-3 g3 : %d barres == n3.chain_ladder.facteurs'
              % len(att))

    def test_g4_la_reserve_par_annee_avec_son_sigma(self):
        """g4 fusionne l'ancien g4 et l'ancien g11, et porte σ (lot C3b).

        Trois choses vérifiées ici, dont deux qui n'existaient pas avant :
        l'IBNR est SIGNÉ (l'ancien traçait `max(v, 0)`, une année en reprise
        s'affichait à zéro), σ par année est présent, et les ultimes de chaque
        méthode disponible sont exacts.
        """
        g, _, _, n3, _ = _figures()
        fig = g['g4_reserve_annee']
        ibnr = _f(n3['chain_ladder']['ibnr_par_annee'])
        barres = _tr(fig, 'IBNR (Chain Ladder)')
        np.testing.assert_allclose(_f(barres.y), ibnr, rtol=1e-9)
        np.testing.assert_allclose(
            _f(barres.error_y.array),
            _f(n3['mack']['sigma_par_annee']), rtol=1e-9,
            err_msg='la barre d\'erreur n\'est pas σ par année de Mack')
        diag = _f(n3['chain_ladder']['last_diagonale'])
        np.testing.assert_allclose(_f(_tr(fig, 'Payé à date').y), diag,
                                   rtol=1e-9)
        for lbl, cle in (('Ultime Chain Ladder', 'chain_ladder'),
                         ('Ultime BF', 'bf'), ('Ultime Cape Cod', 'cape_cod')):
            ult = _f(n3.get(cle, {}).get('ultimates'))
            np.testing.assert_allclose(_f(_tr(fig, lbl).y), ult, rtol=1e-9,
                                       err_msg='série %r' % lbl)
        np.testing.assert_allclose(
            [d + i for d, i in zip(diag, ibnr)],
            _f(n3['chain_ladder']['ultimates']), rtol=1e-9,
            err_msg='payé à date + IBNR ne fait plus l\'ultime')
        print('    OK C2-4 g4 : IBNR signé ± σ Mack, payé à date et les '
              'ultimes des %d méthodes disponibles'
              % sum(1 for c in ('chain_ladder', 'bf', 'cape_cod')
                    if n3.get(c, {}).get('ultimates')))

    def test_g4_une_annee_en_reprise_descend_sous_l_axe(self):
        """⚠️ LE DÉFAUT QUE L'ANCIEN g4 REPORTAIT PAR ÉCRIT.

        Son code portait : « GARDE ANTI-PLANTAGE, PAS LA VRAIE CORRECTION […]
        la vraie correction est prévue au chantier rapport ». Le `max(v, 0)`
        affichait une année de recours à ZÉRO, et le cumul tracé dépassait la
        réserve — 1 483 contre 1 076 mesurés. Le chantier rapport, c'est ici.
        """
        r = _run_triangle(_TRI_RECOURS_FORT)
        g = r.get('graphiques') or {}
        if not g:
            self.skipTest('plotly absent')
        ibnr = _f(r['n3']['chain_ladder']['ibnr_par_annee'])
        self.assertTrue(any(v < 0 for v in ibnr),
                        'ce triangle ne porte plus de reprise — le verrou '
                        'ne prouverait plus rien')
        obs = _f(_tr(g['g4_reserve_annee'], 'IBNR (Chain Ladder)').y)
        np.testing.assert_allclose(obs, ibnr, rtol=1e-9)
        self.assertEqual(sum(1 for v in obs if v < 0),
                         sum(1 for v in ibnr if v < 0))
        self.assertAlmostEqual(
            sum(obs), float(r['n3']['chain_ladder']['reserve_totale']),
            places=2,
            msg='la somme des barres ne fait plus la réserve : un plancher '
                'est revenu')
        print('    OK C2-4b g4 : %d années en reprise tracées NÉGATIVES, et '
              'la somme des barres == la réserve' % sum(1 for v in obs if v < 0))

    def test_g5_chaque_barre_est_la_reserve_de_sa_methode(self):
        g, _, _, n3, n4 = _figures()
        att = {
            'Chain Ladder':         n3['chain_ladder']['reserve_totale'],
            'Bornhuetter-Ferguson': n3['bf']['reserve_totale'],
            'Cape Cod':             n3['cape_cod']['reserve_totale'],
            'BE S2':                n4['best_estimate'],
        }
        for lbl, v in att.items():
            obs = _f(_tr(g['g5_convergence'], lbl).y)
            self.assertAlmostEqual(obs[0], float(v), places=2,
                                   msg='barre %r' % lbl)
        be = float(n4['best_estimate'])
        bornes = [float(att['Bornhuetter-Ferguson']),
                  float(att['Chain Ladder'])]
        self.assertTrue(min(bornes) <= be <= max(bornes),
                        'le BE tracé sort de l\'enveloppe des méthodes')
        print('    OK C2-5 g5 : 4 barres == les réserves publiées, BE dans '
              'l\'enveloppe')

    def test_g6_l_histogramme_est_la_distribution_bootstrap(self):
        g, _, _, n3, _ = _figures()
        att = _f((n3.get('bootstrap') or {}).get('distribution'))
        obs = _f(g['g6_bootstrap'].data[0].x)
        self.assertEqual(len(obs), len(att))
        np.testing.assert_allclose(sorted(obs), sorted(att), rtol=1e-9)
        print('    OK C2-6 g6 : %d tirages == bootstrap.distribution'
              % len(att))

    def test_g9_la_heatmap_est_l_ecart_au_facteur_agrege(self):
        g, C, _, n3, _ = _figures()
        f_cl = _f(n3['chain_ladder']['facteurs'])
        z = np.asarray(g['g9_h2'].data[0].z, dtype=float)
        n, m = C.shape
        compares = 0
        for i in range(n):
            for j in range(min(n - i - 1, len(f_cl))):
                if not (C[i, j] > 0 and f_cl[j]):
                    continue
                att = (C[i, j + 1] / C[i, j] / f_cl[j] - 1.0) * 100.0
                self.assertLess(
                    abs(float(z[i][j]) - att), 0.051,
                    'case (%d,%d) : %s tracé, %.2f attendu'
                    % (i, j, z[i][j], att))
                compares += 1
        self.assertGreaterEqual(compares, 40)
        print('    OK C2-9 g9 : %d cases == (f_individuel / f_CL − 1) × 100'
              % compares)

    def test_g10_la_ligne_de_reference_est_le_lr_a_priori(self):
        g, _, _, n3, _ = _figures()
        lr = float(n3['bf']['lr_apriori'])
        lignes = [float(s.y0) for s in (g['g10_h3'].layout.shapes or ())
                  if s.type == 'line']
        self.assertTrue(
            any(abs(v - lr * 100.0) < 1e-6 for v in lignes),
            'la ligne de référence ne vaut pas lr_apriori × 100 (%.4f) ; '
            'lignes tracées : %s' % (lr * 100.0, lignes))
        print('    OK C2-10 g10 : ligne de référence == lr_apriori × 100 '
              '= %.2f' % (lr * 100.0))

    def test_g12_le_tornado_est_horizontal_et_mesure_les_ecarts_au_be(self):
        g, _, _, _, n4 = _figures()
        be = float(n4['best_estimate'])
        att = sorted(round(float(v) - be, 2)
                     for v in (n4.get('sensibilites') or {}).values())
        obs = []
        for t in g['g12_sensibilites'].data:
            self.assertEqual(t.orientation, 'h',
                             'un tornado tracé en vertical n\'est pas un '
                             'tornado')
            obs += [round(v, 2) for v in _f(t.x) if v]
        self.assertEqual(sorted(obs), [v for v in att if v],
                         'les barres ne sont pas les écarts au BE')
        print('    OK C2-12 g12 : %d écarts horizontaux == sensibilites − BE'
              % len(obs))

    def test_g14_le_backtesting_est_l_ecart_projete_observe(self):
        g, _, _, n3, _ = _figures()
        res = (n3.get('backtesting') or {}).get('resultats') or {}
        for nom, cle in (('Horizon N-1', 'horizon_1'),
                         ('Horizon N-2', 'horizon_2')):
            annees = (res.get(cle) or {}).get('annees') or []
            t = _tr(g['g14_backtesting'], nom)
            np.testing.assert_allclose(
                _f(t.y), [float(a['ecart_pct']) for a in annees], rtol=1e-9,
                err_msg='horizon %r' % nom)
            self.assertEqual([str(v) for v in t.x],
                             [a['annee_label'] for a in annees],
                             'les points de %r ne sont pas alignés sur les '
                             'années qu\'ils datent' % nom)
        print('    OK C2-14 g14 : les 2 horizons == ecart_pct par année, '
              'étiquettes alignées')


    # ── AJOUTS DU LOT C3c ──────────────────────────────────────────────────

    def test_g15_l_exposition_et_le_loss_ratio_qu_elle_implique(self):
        """Guide IA 2023, Figure 11 — la chronique des primes acquises.

        ⚠️ ET LA MESURE QUI JUSTIFIE L'AJOUT, PLUTÔT QUE MON ARGUMENT INITIAL.
        J'avais dit que ce graphique aurait rendu visible l'erreur ×60 du lot
        F1. C'est vrai, mais INCOMPLET : depuis F1, BFCC-H6 attrape déjà une
        exposition globalement ×60 (loss ratio poolé 1,8 %, hors plage
        plausible → NON VALIDÉE). La vraie valeur ajoutée est ailleurs, et
        elle est mesurée dans le test suivant : UNE SEULE année mal saisie
        passe la garde poolée.
        """
        g, _, _, n3, _ = _figures()
        expo = _f(EXPOSITION[:len(_f(n3['chain_ladder']['ultimates']))])
        obs = _f(_tr(g['g15_exposition'], 'Exposition').y)
        np.testing.assert_allclose(obs, expo, rtol=1e-9)
        ult = _f(n3['chain_ladder']['ultimates'])
        att = [u / e * 100.0 for u, e in zip(ult, expo)]
        np.testing.assert_allclose(
            _f(_tr(g['g15_exposition'], 'Loss ratio implicite').y), att,
            rtol=1e-9,
            err_msg="la courbe n'est pas ultime / exposition")
        print('    OK C3c-1 g15 : barres == exposition, courbe == ultime / '
              'exposition (%.0f %% à %.0f %%)' % (min(att), max(att)))

    def test_g15_voit_une_annee_que_la_garde_poolee_ne_voit_pas(self):
        """⚠️ LA PREUVE DE LA VALEUR MARGINALE, ET ELLE EST MESURÉE.

        Une SEULE année dont l'exposition est saisie soixante fois trop grande
        ne fait pas tomber BFCC-H6 : le loss ratio POOLÉ reste dans la plage
        plausible. Le graphique, lui, la montre à 1,6 % contre 109 % de
        médiane. Une garde agrégée ne peut pas localiser ; une chronique si.
        """
        expo = np.array(EXPOSITION, dtype=float)
        expo[4] *= 60
        r = _run_triangle(GENINS, primes=expo)
        g = r.get('graphiques') or {}
        if not g:
            self.skipTest('plotly absent')
        h6 = (r['n2'].get('bfcc', {}).get('hypotheses', {})
              .get('BFCC-H6', {}))
        self.assertEqual(str(h6.get('statut')), 'VALIDÉE',
                         "BFCC-H6 attrape désormais le cas — ce test ne "
                         "prouverait plus la valeur marginale de g15")
        lr = _f(_tr(g['g15_exposition'], 'Loss ratio implicite').y)
        autres = sorted(v for i, v in enumerate(lr) if i != 4)
        mediane = autres[len(autres) // 2]
        self.assertLess(lr[4] * 20, mediane,
                        "l'année aberrante ne ressort plus du lot")
        print('    OK C3c-2 g15 : BFCC-H6 VALIDÉE mais l\'année 4 ressort à '
              '%.1f %% contre %.0f %% de médiane' % (lr[4], mediane))

    def test_g16_les_increments_et_leurs_negatifs(self):
        """Le triangle que g1 cache par construction."""
        g, C, _, _, _ = _figures()
        z = np.asarray(g['g16_increments'].data[0].z, dtype=float)
        n, m = C.shape
        att = np.full(C.shape, np.nan)
        for i in range(n):
            for j in range(min(n - i - 1, m - 1) + 1):
                att[i, j] = C[i, j] if j == 0 else C[i, j] - C[i, j - 1]
        connu = ~np.isnan(att)
        np.testing.assert_allclose(z[connu], att[connu], rtol=1e-9)
        self.assertTrue(np.isnan(z[~connu]).all(),
                        'une case future porte un incrément')
        hm = g['g16_increments'].data[0]
        self.assertEqual(hm.zmid, 0, "l'échelle n'est plus centrée sur zéro")
        self.assertAlmostEqual(abs(float(hm.zmin)), abs(float(hm.zmax)),
                               places=6,
                               msg='échelle divergente asymétrique : un '
                                   'négatif se noierait')
        print('    OK C3c-3 g16 : %d incréments == C[i,j] − C[i,j−1], échelle '
              'centrée sur zéro' % int(connu.sum()))

    def test_g16_montre_les_reprises_que_g1_ne_montre_pas(self):
        """⚠️ LE SUJET MÊME DE L'AJOUT, VÉRIFIÉ SUR UN TRIANGLE À RECOURS."""
        r = _run_triangle(_TRI_RECOURS_FORT)
        g = r.get('graphiques') or {}
        if not g:
            self.skipTest('plotly absent')
        inc = np.asarray(g['g16_increments'].data[0].z, dtype=float)
        cum = np.asarray(g['g1_heatmap'].data[0].z, dtype=float)
        n_inc = int(np.nansum(inc < 0))
        self.assertGreater(n_inc, 0,
                           'ce triangle ne porte plus de reprise — le verrou '
                           'ne prouverait plus rien')
        self.assertEqual(int(np.nansum(cum < 0)), 0,
                         'les cumulés portent un négatif : le contraste qui '
                         'justifie g16 a disparu')
        print('    OK C3c-4 g16 : %d incréments négatifs visibles, %d sur les '
              'cumulés de g1' % (n_inc, int(np.nansum(cum < 0))))

    def test_g17_le_nuage_de_linearite_et_sa_droite_par_l_origine(self):
        """Guide IA 2023, §9.d.ii — l'hypothèse sur l'espérance."""
        g, C, _, n3, _ = _figures()
        fig = g['g17_linearite']
        n, m = C.shape
        f = _f(n3['chain_ladder']['facteurs'])
        nuages = [t for t in fig.data if t.mode == 'markers']
        droites = [t for t in fig.data if t.mode == 'lines']
        # Une transition n'est traçable qu'à partir de DEUX points : un nuage
        # d'un seul point ne montre aucune linéarité. La dernière transition
        # n'en a qu'un, elle est écartée — et c'est voulu.
        traçables = sum(1 for j in range(min(m - 1, len(f)))
                        if sum(1 for i in range(n)
                               if j + 1 <= min(n - i - 1, m - 1)
                               and C[i, j] > 0) >= 2)
        self.assertEqual(len(nuages), traçables)
        self.assertEqual(len(droites), traçables,
                         'chaque nuage doit porter SA droite par l\'origine')
        for j, (nuage, droite) in enumerate(zip(nuages, droites)):
            att_x = [float(C[i, j]) for i in range(n)
                     if j + 1 <= min(n - i - 1, m - 1) and C[i, j] > 0]
            att_y = [float(C[i, j + 1]) for i in range(n)
                     if j + 1 <= min(n - i - 1, m - 1) and C[i, j] > 0]
            np.testing.assert_allclose(_f(nuage.x), att_x, rtol=1e-9)
            np.testing.assert_allclose(_f(nuage.y), att_y, rtol=1e-9)
            self.assertEqual((float(droite.x[0]), float(droite.y[0])), (0.0, 0.0),
                             'la droite ne part pas de l\'origine')
            self.assertAlmostEqual(float(droite.y[1]) / float(droite.x[1]),
                                   f[j], places=9,
                                   msg='la pente n\'est pas le facteur f_%d' % j)
        print('    OK C3c-5 g17 : %d nuages == (C[i,j], C[i,j+1]), %d droites '
              'par l\'origine de pente f_j' % (len(nuages), len(droites)))

    def test_g18_les_residus_standardises(self):
        """Guide IA 2023, §9.d.iii — l'hypothèse sur la variance."""
        g, C, _, n3, _ = _figures()
        n, m = C.shape
        f = _f(n3['chain_ladder']['facteurs'])
        s2 = _f(n3['mack']['sigma2_par_colonne'])
        att = []
        for j in range(min(m - 1, len(f))):
            for i in range(n):
                if j + 1 > min(n - i - 1, m - 1) or C[i, j] <= 0:
                    continue
                brut = (C[i, j + 1] - f[j] * C[i, j]) / C[i, j] ** 0.5
                att.append((float(C[i, j]),
                            brut / s2[j] ** 0.5 if s2[j] > 0 else brut))
        obs = [(float(x), float(y)) for t in g['g18_residus'].data
               for x, y in zip(t.x or [], t.y or [])]
        self.assertEqual(len(obs), len(att))
        for a, b in zip(sorted(obs), sorted(att)):
            self.assertAlmostEqual(a[0], b[0], places=6)
            self.assertAlmostEqual(a[1], b[1], places=9)
        ys = [p[1] for p in obs]
        self.assertLess(max(abs(v) for v in ys), 10,
                        'la standardisation par σ_j ne tient plus : sans elle '
                        'la plage brute est [−617 ; 600] et le nuage est '
                        'illisible')
        print('    OK C3c-6 g18 : %d résidus == (C[i,j+1] − f_j·C[i,j])/√C[i,j] '
              'rapportés à σ_j, plage [%.2f ; %.2f]'
              % (len(att), min(ys), max(ys)))

    def test_g6_marque_le_be_retenu_et_ecrit_sa_reserve(self):
        """⚠️ LE CHIFFRE INCONFORTABLE EST PUBLIÉ, AVEC SA RÉSERVE.

        Le guide assigne à la mesure d'incertitude un usage précis :
        « objectiver la notion parfois floue de prudence dans les provisions
        [...] dans le cadre des échanges avec les CAC et/ou l'ACPR ». g6
        repérait P50, P75, P90 et P99.5 — et pas le montant qui part en
        comptabilité.

        LA RÉSERVE EST VÉRIFIÉE, PAS SEULEMENT LE REPÈRE : la distribution
        bootstrap est bâtie sur le SEUL Chain Ladder alors que le BE mêle
        trois méthodes. Publier un percentile sans dire contre quelle
        distribution remplacerait une ambiguïté par une autre.
        """
        g, _, _, n3, n4 = _figures()
        be = float(n4['best_estimate'])
        lignes = [float(s.x0) for s in (g['g6_bootstrap'].layout.shapes or ())
                  if s.type == 'line']
        self.assertTrue(any(abs(x - be) < 1e-6 for x in lignes),
                        'aucun repère sur le Best Estimate retenu')
        textes = [str(a.text) for a in
                  (g['g6_bootstrap'].layout.annotations or ())]
        marque = [t for t in textes if 'BE retenu' in t]
        self.assertEqual(len(marque), 1, 'le repère du BE n\'est pas nommé')
        self.assertIn('Chain Ladder', marque[0],
                      'le percentile est publié SANS dire contre quelle '
                      'distribution il est mesuré')
        dist = sorted(_f(n3['bootstrap']['distribution']))
        rang = sum(1 for v in dist if v <= be) / len(dist) * 100.0
        self.assertIn('P%.0f' % rang, marque[0],
                      'le percentile affiché n\'est pas celui de la '
                      'distribution publiée')
        print('    OK C3c-7 g6 : BE retenu marqué au P%.0f de la distribution '
              'Chain Ladder, réserve écrite sur le graphique' % rang)

# =============================================================================
#  T3 — LES DÉFAUTS CONSTATÉS EN C2, NON CORRIGÉS ICI
# =============================================================================

class T3_Les_Trois_Defauts_De_L_Audit_C2(unittest.TestCase):
    """⚠️ LES TROIS DÉFAUTS DE L'AUDIT C2, ET LEURS TROIS SORTS.

    Ils étaient écrits comme des assertions de ce qui DEVRAIT être vrai,
    marquées `@unittest.expectedFailure`. Aucun marqueur ne subsiste, et les
    trois chemins pour y arriver n'ont pas été les mêmes :

      · g5 et g11, CORRIGÉS au lot C3a. Le dispositif a fonctionné seul : la
        campagne est passée au rouge avec « unexpected successes=2 » et les
        marqueurs ont dû être retirés. Les deux tests ci-dessous sont
        désormais des verrous ordinaires.

      · g7, SUPPRIMÉ au lot C3b — et là le dispositif N'AURAIT PAS SUFFI.
        `expectedFailure` avale une erreur autant qu'un échec : le `KeyError`
        d'un graphique disparu serait passé pour l'échec attendu, en silence.
        C'est le verrou de CATALOGUE de T1 qui a échoué au passage de 14 à 10
        et forcé le retrait du test orphelin. Un marqueur a besoin d'un verrou
        de périmètre à côté de lui.
    """

    def test_g5_une_methode_indisponible_ne_devrait_pas_valoir_zero_euro(self):
        """CORRIGÉ AU LOT C3a — le marqueur `expectedFailure` a été retiré.

        `res_map.get(m, 0)` transformait « non calculable faute d'exposition »
        en « réserve de zéro euro », étiquette « 0€ » imprimée sur la barre.
        Le scénario sans primes est celui de l'oracle GenIns : ce n'était pas
        un cas de bord.

        g5 lit désormais `methodes_be.reserve()`, qui rend **None** et non
        zéro : une méthode exclue mais CALCULÉE reste tracée en grisé — c'est
        une information — et une méthode absente n'a plus de barre.
        """
        g, _, _, n3, _ = _figures(avec_primes=False)
        self.assertFalse(n3['bf']['disponible'])
        for lbl in ('Bornhuetter-Ferguson', 'Cape Cod'):
            try:
                t = _tr(g['g5_convergence'], lbl)
            except AssertionError:
                continue          # absente : c'est la bonne réponse
            self.fail('%r tracé à %s € alors que la méthode est indisponible'
                      % (lbl, _f(t.y)))

    def test_une_methode_indisponible_ne_devrait_pas_etre_une_ligne_a_zero(self):
        """CORRIGÉ AU LOT C3a, ET LE SUJET A SUIVI LA FUSION AU LOT C3b.

        Le défaut était dans g11 : deux séries plates à zéro affirmaient des
        ultimes INFÉRIEURS aux montants déjà payés. `n3['bf']['ultimates']`
        n'est pas VIDE quand BF est indisponible — elle est pleine de zéros,
        donc le `if ult:` la laissait passer. La garde porte sur
        `disponible()`, pas sur la longueur.

        g11 a été fondu dans g4 au lot C3b : le verrou vise le graphique
        fusionné, parce que c'est le sujet qui compte, pas le nom du support.
        """
        g, _, _, n3, _ = _figures(avec_primes=False)
        self.assertFalse(n3['cape_cod']['disponible'])
        for lbl in ('Ultime BF', 'Ultime Cape Cod'):
            try:
                t = _tr(g['g4_reserve_annee'], lbl)
            except AssertionError:
                continue          # absente : c'est la bonne réponse
            self.assertTrue(any(v for v in _f(t.y)),
                            '%r est une ligne plate à zéro' % lbl)



# =============================================================================
#  T4 — LE CATALOGUE ET LE RAPPORT COÏNCIDENT  (lot C3d)
# =============================================================================

class T4_Le_Catalogue_Atteint_Les_Livrables(unittest.TestCase):
    """⚠️ PRODUIRE UNE FIGURE ET NE PAS LA PUBLIER, C'EST NE PAS L'AVOIR FAITE.

    Avant le lot C3d, le HTML portait CINQ figures sur quatorze et le Word
    ZÉRO. Les neuf absentes du HTML étaient celles qui JUSTIFIENT la méthode —
    cadences, facteurs, hypothèses, sensibilités — pendant que celles qui
    donnent le RÉSULTAT passaient. C'est l'inverse de ce dont un commissaire
    aux comptes a besoin : le résultat, il le lit dans un tableau.
    """

    _NUM = re.compile(r'Figure (\d+) — ')

    def _numeros_html(self, html):
        return [int(n) for n in re.findall(r'Figure (\d+) — ', html)]

    def _numeros_word(self, octets):
        import docx
        doc = docx.Document(io.BytesIO(octets))
        return [int(m.group(1)) for m in
                (self._NUM.match(p.text) for p in doc.paragraphs) if m]

    def test_chaque_figure_produite_atteint_le_html(self):
        r = _run(True)
        html = r.get('html') or ''
        if not html:
            self.skipTest('HTML non produit')
        manquantes = [cle for cle in (r.get('graphiques') or {})
                      if TITRES_FIGURES[cle] not in html]
        self.assertEqual(manquantes, [],
                         'figures produites mais absentes du rapport : %s'
                         % manquantes)
        self.assertEqual(len(self._numeros_html(html)),
                         len(r.get('graphiques') or {}))
        print('    OK C3d-1 les %d figures produites sont TOUTES dans le HTML'
              % len(r.get('graphiques') or {}))

    def test_chaque_titre_du_catalogue_existe(self):
        """Un graphique sans titre sortirait avec sa clé technique."""
        self.assertEqual(sorted(TITRES_FIGURES), sorted(CATALOGUE),
                         'le registre des titres et le catalogue divergent')
        print('    OK C3d-2 les %d graphiques ont un titre de lecture'
              % len(CATALOGUE))

    def test_la_numerotation_reste_continue_quand_une_figure_manque(self):
        """⚠️ LE POINT QUI JUSTIFIE UN COMPTEUR POSITIONNEL.

        Sans exposition, `g10_h3` et `g15_exposition` ne se produisent pas.
        Une numérotation précalculée laisserait deux trous — « Figure 4 » puis
        « Figure 6 ». Le compteur n'avance que sur une figure RÉELLEMENT
        rendue : la suivante prend simplement son numéro.
        """
        for avec in (True, False):
            r = _run(avec)
            html = r.get('html') or ''
            if not html:
                self.skipTest('HTML non produit')
            nums = self._numeros_html(html)
            self.assertEqual(nums, list(range(1, len(nums) + 1)),
                             'numérotation à trous (exposition=%s) : %s'
                             % (avec, nums))
            self.assertEqual(len(nums), len(r.get('graphiques') or {}))
        sans = len(_run(False).get('graphiques') or {})
        avec = len(_run(True).get('graphiques') or {})
        self.assertLess(sans, avec,
                        "le scénario sans exposition ne retire plus aucune "
                        "figure — le test ne prouverait plus rien")
        print('    OK C3d-3 numérotation continue 1..%d et 1..%d, sans trou '
              'là où deux figures s\'effacent' % (avec, sans))

    def test_le_word_et_le_html_numerotent_pareil(self):
        """« Figure 7 » doit désigner la même chose dans les deux formats.

        ⚠️ CE TEST PORTE SUR DES NUMÉROS, PAS SUR DES PIXELS. Il rasterisait
        quatorze figures — 82 s mesurées — pour lire quatorze légendes. Le
        rendeur substitué lui donne exactement ce dont il a besoin : des
        images qui entrent dans le document.
        """
        C = np.asarray(GENINS, dtype=float)
        with kaleido_declare(True), rendeur_substitue():
            r = AgentA7Provisionnement(verbose=False).run(
                source=C, mode_declare='cumule', generer_graphiques=True,
                generer_word=True, n_sim_bootstrap=60, seed=42,
                primes=EXPOSITION[:C.shape[0]])
        html, mot = r.get('html') or '', r.get('word_bytes') or b''
        if not html or not mot:
            self.skipTest('HTML ou Word non produit')
        titres_html = re.findall(r'Figure \d+ — ([^<]{4,120})', html)
        import docx
        titres_word = [p.text.split(' — ', 1)[1]
                       for p in docx.Document(io.BytesIO(mot)).paragraphs
                       if self._NUM.match(p.text)]
        self.assertEqual(titres_html, titres_word,
                         "l'ordre des figures diverge entre HTML et Word")
        print('    OK C3d-4 les %d figures portent le même numéro dans le '
              'HTML et dans le Word' % len(titres_html))

    def test_le_word_nomme_l_absence_de_kaleido_au_lieu_de_la_taire(self):
        """⚠️ NI INSTALLATION FORCÉE, NI SILENCE.

        Un `.docx` ne porte que du raster : sans `kaleido`, plotly refuse
        `to_image`. Le Word se produit quand même, chaque légende porte la
        raison, et `run()` remonte `figures_word` — un livrable dégradé est
        DÉCLARÉ, jamais deviné. C'est la discipline du lot F2, étendue à une
        dépendance qui vit À L'INTÉRIEUR d'un export.
        """
        C = np.asarray(GENINS, dtype=float)
        with kaleido_declare(False):
            r = AgentA7Provisionnement(verbose=False).run(
                source=C, mode_declare='cumule', generer_graphiques=True,
                generer_word=True, n_sim_bootstrap=60, seed=42,
                primes=EXPOSITION[:C.shape[0]])
        mot = r.get('word_bytes') or b''
        if not mot:
            self.skipTest('Word non produit')
        import docx
        doc = docx.Document(io.BytesIO(mot))
        notes = [p.text for p in doc.paragraphs
                 if p.text.startswith('Figure non rendue')]
        erreur = (r.get('livrables_erreurs') or {}).get('figures_word')
        self.assertEqual(len(notes), len(self._numeros_word(mot)),
                         'une figure manque sans dire pourquoi')
        self.assertEqual(erreur, 'dependance_absente: kaleido',
                         'la dégradation n\'est pas remontée à l\'appelant')
        # ⚠️ ET LA DÉGRADATION DOIT ÊTRE LUE PAR TOUTE LA CHAÎNE : l'export
        # produit les mentions, l'agent remonte `figures_word`. Ces deux-là
        # lisaient le prédicat à deux endroits — l'un par référence figée.
        print('    OK C3d-5 dégradation FORCÉE : %d légendes, %d raisons '
              'nommées, et `figures_word` remonté' % (len(notes), len(notes)))

    def test_le_rasteriseur_demande_bien_un_png_a_la_bonne_taille(self):
        """⚠️ LE RENDEUR SUBSTITUÉ NE PROUVE RIEN SUR LES ARGUMENTS RÉELS.

        En remplaçant `rendre_image`, les tests cessent de voir ce que la
        vraie fonction demande à Plotly — et un `format='svg'` glissé là
        passerait tous les tests en produisant un `.docx` que Word refuse.
        Cette vérification-ci porte donc sur `rendre_image` elle-même, avec
        une figure de substitution qui note ce qu'on lui demande.
        """
        vus = {}

        class _Figure:
            def to_image(self, **kw):
                vus.update(kw)
                return b'png'

        self.assertEqual(_RAP.rendre_image(_Figure()), b'png')
        self.assertEqual(vus.get('format'), 'png')
        self.assertEqual((vus.get('width'), vus.get('height')), (1100, 520))
        # ⚠️⚠️ `scale` EST PASSE DE 2 A 1.5, ET CE TEST A FAIT SON TRAVAIL :
        # il a signale le changement, comme il doit. Il n'est mis a jour que
        # parce que le changement est INTENTIONNEL, MESURE et ARBITRE.
        #
        # RAISON : depuis que le HTML porte ses figures en image, C3 archive
        # ce document A CHAQUE RUN. Le poids devient un cout recurrent.
        #   scale=1    32 Ko/figure   plus petit texte  9 px  -- ECARTE
        #   scale=1.5  62 Ko/figure                    14 px  -- retenu
        #   scale=2    98 Ko/figure                    18 px
        # Effet mesure sur le dossier archive : 3 957 254 -> 2 619 797 octets,
        # soit 1,34 Mo PAR CLOTURE (le Word en profite aussi : `rendre_image`
        # sert les deux formats).
        #
        # ⚠️ CE QUE CE TEST PROTEGE RESTE ENTIER : le `format='png'` et les
        # dimensions. Un `format='svg'` glisse ici produirait toujours un
        # `.docx` que Word refuse, et il tomberait toujours.
        self.assertEqual(vus.get('scale'), 1.5)
        print('    OK C3d-7 le rasteriseur demande un PNG 1100x520 en x1.5')

    def test_une_figure_n_est_rasterisee_qu_une_fois(self):
        """⚠️⚠️ CHAQUE FIGURE ETAIT RENDUE DEUX FOIS — MESURE SUR UN RUN REEL.

        24 appels a `rendre_image` pour DOUZE figures distinctes : les douze,
        sans exception, une fois pour l'HTML et une fois pour le Word. 126,1 s
        cumulees, 5,3 s par appel. La MOITIE etait de la pure duplication, et
        un arrete reel la payait a chaque cloture.

        ⚠️ MEMOISER NE PEUT RIEN CHANGER A CE QUI EST PUBLIE : `rendre_image`
        ne prend AUCUN parametre de taille. Prouve par ailleurs octet par
        octet (les 12 images HTML et les 12 Word sont identiques avant et
        apres). Ce test-ci verrouille le MECANISME.
        """
        rendus = []

        class _Figure:
            def to_image(self, **kw):
                rendus.append(kw)
                return b'PNG-' + bytes([len(rendus)])

        f = _Figure()
        premier = _RAP.rendre_image(f)
        second = _RAP.rendre_image(f)
        self.assertEqual(len(rendus), 1,
                         'la meme figure est rasterisee plusieurs fois')
        self.assertEqual(premier, second, 'deux octets differents pour une '
                                          'seule et meme figure')
        # ⚠️ CONTRE-EPREUVE : une AUTRE figure ne doit pas recevoir l'image
        # de la premiere. Sans elle, un cache qui rendrait toujours la meme
        # chose passerait ce test au vert.
        autre = _RAP.rendre_image(_Figure())
        self.assertEqual(len(rendus), 2, 'une figure neuve n a pas ete rendue')
        self.assertNotEqual(premier, autre, 'une figure recoit l image d une '
                                            'autre')
        print('    OK C3d-8 une figure rendue une fois, une autre distinguee')

    def test_la_session_de_rasterisation_est_reentrante_et_ferme_toujours(self):
        """⚠️⚠️ UN NAVIGATEUR PAR FIGURE — MESURE, 3,8 s CONTRE 0,6 s.

        Le profil d'un run reel montrait `kaleido.close` a 134,4 s : un
        navigateur ouvert ET FERME a chaque image. Une session unique ramene
        les douze figures de 45,8 s a 7,4 s, et LES OCTETS SONT IDENTIQUES --
        c'etait la condition posee avant de coder.

        ⚠️ CE TEST NE FAIT TOURNER AUCUN AGENT : il substitue le demarrage et
        l'arret, qui coutent 3,65 s a la fermeture. Un test de mecanisme n'a
        pas a payer un navigateur.
        """
        # ⚠️ UNE PREMIERE VERSION LISAIT `sys.modules` : elle SAUTAIT
        # toujours, parce que `session_rasterisation` importe kaleido
        # PARESSEUSEMENT, a l'interieur. Un test saute ne prouve rien -- c'est
        # un controle qui atteste sans surveiller.
        if not _RAP.kaleido_disponible():                 # pragma: no cover
            self.skipTest('kaleido absent : la session est un no-op')
        import kaleido as kal
        demarrages, arrets = [], []
        vrai_start, vrai_stop = kal.start_sync_server, kal.stop_sync_server
        kal.start_sync_server = lambda **k: demarrages.append(1)
        kal.stop_sync_server = lambda **k: arrets.append(1)
        try:
            class _Fig:
                def to_image(self, **kw):
                    return b'PNG'

            self.assertEqual(_RAP._PROFONDEUR_SESSION[0], 0,
                             'une session est restee ouverte avant ce test')

            # ⚠️⚠️ UNE SESSION SANS RENDU N'OUVRE AUCUN NAVIGATEUR, et c'est
            # LE correctif de ce lot. Une premiere version demarrait a
            # l'entree : la gate a fait 888 CYCLES et elle est passee de
            # 2 886 s a 3 356 s -- PLUS LENTE. La plupart des tests
            # substituent `rendre_image` et ne rasterisent rien.
            with _RAP.session_rasterisation():
                pass
            self.assertEqual(demarrages, [],
                             'un navigateur ouvert pour RIEN')

            with _RAP.session_rasterisation():
                self.assertEqual(_RAP._PROFONDEUR_SESSION[0], 1)
                _RAP.rendre_image(_Fig())
                self.assertEqual(len(demarrages), 1, 'aucun navigateur ouvert')
                with _RAP.session_rasterisation():
                    # ⚠️ IMBRIQUEE : la profondeur monte, mais AUCUN second
                    # navigateur. C'est ce qui permet d'englober plusieurs
                    # livrables sans redemarrer.
                    self.assertEqual(_RAP._PROFONDEUR_SESSION[0], 2)
                    _RAP.rendre_image(_Fig())
                    self.assertEqual(len(demarrages), 1,
                                     'un second navigateur a ete ouvert')
                self.assertEqual(len(arrets), 0,
                                 'la session interne a ferme le navigateur')
            self.assertEqual(_RAP._PROFONDEUR_SESSION[0], 0)
            self.assertEqual((len(demarrages), len(arrets)), (1, 1),
                             'un demarrage, un arret, et pas davantage')

            # ⚠️⚠️ LE `finally` EST LE POINT CRITIQUE : un serveur laisse
            # ouvert retiendrait le processus apres la fin des tests --
            # exactement le defaut de gate observe une fois et jamais
            # explique. On ne l'introduit pas.
            with self.assertRaises(ValueError), \
                    _RAP.session_rasterisation():
                _RAP.rendre_image(_Fig())      # il y a donc bien a fermer
                raise ValueError('incident pendant le rendu')
            self.assertEqual(_RAP._PROFONDEUR_SESSION[0], 0,
                             'une exception laisse la session ouverte')
            self.assertEqual(len(arrets), 2,
                             'le navigateur survit a une exception')
        finally:
            kal.start_sync_server, kal.stop_sync_server = vrai_start, vrai_stop
        print('    OK C3d-10 session reentrante, et fermee meme sur exception')

    def test_les_livrables_sont_produits_SOUS_session(self):
        """⚠️ SANS CE VERROU, LE GAIN SE PERDRAIT EN SILENCE.

        Retirer la session ne casserait RIEN : les rapports sortiraient
        pareils, en rouvrant simplement un navigateur par figure. Un cout qui
        revient sans bruit est celui qu'on ne voit jamais.

        ⚠️ ON MESURE LA PROFONDEUR DEPUIS L'INTERIEUR DE LA RASTERISATION,
        c'est-a-dire le FAIT d'etre sous session -- pas la presence d'un mot
        dans le source. Le rendeur est substitue : aucune image reelle n'est
        produite, le test coute quelques secondes et non une minute.
        """
        profondeurs = []
        vrai = _RAP.rendre_image

        def _sonde(figure):
            profondeurs.append(_RAP._PROFONDEUR_SESSION[0])
            return png_de_test(200 + len(profondeurs), 100)

        _RAP.rendre_image = _sonde
        try:
            AgentA7Provisionnement(verbose=False).run(
                source=np.asarray(GENINS, dtype=float), mode_declare='cumule',
                generer_graphiques=True, n_sim_bootstrap=20, seed=42)
        finally:
            _RAP.rendre_image = vrai
        self.assertTrue(profondeurs, 'aucune figure rasterisee : mesure vide')
        self.assertTrue(all(p >= 1 for p in profondeurs),
                        f'des figures rendues HORS session : {profondeurs}')
        # ⚠️ ET LA SESSION SE REFERME : une session laissee ouverte
        # retiendrait le processus apres la gate.
        self.assertEqual(_RAP._PROFONDEUR_SESSION[0], 0,
                         'une session est restee ouverte apres le run')
        print(f'    OK C3d-11 les {len(profondeurs)} figures sont rendues '
              f'sous session, refermee ensuite')

    def test_le_cache_d_images_ne_survit_pas_a_sa_figure(self):
        """⚠️ LA CLE EST `id()`, ET UN `id` SE RECYCLE.

        Une figure Plotly n'est pas hashable : la cle ne peut pas etre la
        figure elle-meme. `id()` est donc employe -- mais un `id` libere est
        REATTRIBUE par CPython, et une entree survivante rendrait alors
        l'image d'une figure morte a une figure neuve. Le finaliseur weakref
        l'interdit ; ce test verifie qu'il est bien pose.
        """
        import gc

        class _Figure:
            def to_image(self, **kw):
                return b'PNG'

        # ⚠️⚠️ ON SUIT NOTRE PROPRE CLE, PAS LA TAILLE DU CACHE. Une premiere
        # version comparait `len(...)` avant et apres : elle passait ISOLEE et
        # ECHOUAIT dans la gate complete -- `gc.collect()` libere aussi les
        # figures des AUTRES tests, si bien que la taille finale etait
        # INFERIEURE a la taille initiale. Un test qui mesure un compteur
        # partage mesure le voisinage autant que lui-meme.
        f = _Figure()
        _RAP.rendre_image(f)
        cle = id(f)
        self.assertIn(cle, _RAP._IMAGES_RASTERISEES, 'rien n a ete memoise')
        del f
        gc.collect()
        self.assertNotIn(cle, _RAP._IMAGES_RASTERISEES,
                         'le cache survit a sa figure : un id recycle '
                         'rendrait une image etrangere')
        print('    OK C3d-9 le cache meurt avec sa figure, aucun id recycle')

    # ⚠️⚠️ UN TEST A ETE ECRIT ICI, PUIS RETIRE, ET C'EST UN ARBITRAGE.
    # Il comparait les images du HTML a celles du Word pour verifier que la
    # memoisation ne les desynchronise pas. Deux raisons de ne pas le garder :
    #
    #   · le fixture partage `_run` produit le HTML SANS le Word
    #     (`generer_word=False`) : le test aurait exige un run DE PLUS,
    #     mesure a ~55 s -- dans un lot dont l'objet est justement de rendre
    #     la suite moins chere ;
    #   · l'invariant qu'il verifiait DECOULE de C3d-8 : une meme figure rend
    #     les memes octets, donc les deux formats recoivent les memes images
    #     des lors qu'ils recoivent les memes figures -- ce que
    #     `test_le_word_et_le_html_numerotent_pareil` couvre deja.
    #
    # ⚠️ L'IDENTITE A ETE VERIFIEE, HORS SUITE, SUR LES PRODUITS REELS :
    # 12 images HTML et 12 images Word, empreintes SHA-256 identiques AVANT
    # et APRES la memoisation, et identiques entre les deux formats. Mesure
    # ponctuelle et non verrouillee -- c'est dit plutot que sous-entendu.

    def test_le_chemin_nominal_du_word_insere_bien_les_images(self):
        """⚠️ LE CHEMIN AVEC kaleido, EXERCÉ QUE LA MACHINE L'AIT OU NON.

        Avant ce lot, ce chemin n'était emprunté que sur un poste où
        `kaleido` se trouvait installé. Il est désormais FORCÉ, et le
        rasteriseur substitué vérifie que les images entrent bien dans le
        document — ce que la vraie fonction demande à Plotly étant vérifié
        juste au-dessus.

        ⚠️ LES PNG SONT TOUS DIFFÉRENTS, ET C'EST NÉCESSAIRE : OOXML
        déduplique les média identiques. Avec quatorze images identiques le
        `.docx` n'en contient qu'UNE, ce qui ferait croire à un bug qui
        n'existe pas — mesuré.
        """
        import zipfile
        r = _run(True)
        g = r.get('graphiques') or {}
        if not g:
            self.skipTest('plotly absent')
        with kaleido_declare(True), rendeur_substitue() as rendeur:
            mot = _RAP.export_word(r.get('n1'), r['n2'], r['n3'], r['n4'],
                                   commentaire=r.get('commentaire', ''),
                                   graphiques=g, lob_label='Test')
        media = [n for n in zipfile.ZipFile(io.BytesIO(mot)).namelist()
                 if n.startswith('word/media/')]
        self.assertEqual(len(media), len(g),
                         'toutes les figures ne sont pas entrées dans le .docx')
        self.assertEqual(len(rendeur.appels), len(g),
                         'le rasteriseur n\'a pas été appelé pour chaque figure')
        import docx
        notes = [p.text for p in docx.Document(io.BytesIO(mot)).paragraphs
                 if p.text.startswith('Figure non rendue')]
        self.assertEqual(notes, [],
                         'une dégradation est annoncée alors que le rendeur '
                         'répond')
        print('    OK C3d-6 chemin nominal FORCÉ : %d images dans le .docx, '
              'le rasteriseur appelé %d fois, aucune dégradation'
              % (len(media), len(rendeur.appels)))

if __name__ == '__main__':
    unittest.main(verbosity=2)
