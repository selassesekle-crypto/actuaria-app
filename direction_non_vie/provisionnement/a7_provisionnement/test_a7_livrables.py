# =============================================================================
#  Tests — lot « avant Bootstrap » : livrables, courbe des taux, affichages
#
#  CE FICHIER MESURE DES TAILLES ET DES CONTENUS, JAMAIS L'ABSENCE D'EXCEPTION.
#  C'est tout son propos. Trois fois de suite, un livrable A7 a été cassé sans
#  qu'aucun test ne bronche, parce que le code rattrapait son propre échec et
#  rendait un repli :
#    · `generer_graphiques` masqué par un paramètre → 0 graphique (lot T20) ;
#    · `export_excel` → 0 octet sur un None (lot 1 de l'harmonisation) ;
#    · `export_html` → 88 octets sur un NameError, LA GATE ENTIÈRE AU VERT
#      (lot BFCC — 596 tests verts avec le rapport HTML cassé).
#  Un test qui se contente d'appeler la fonction sans lever ne voit rien de tout
#  cela. Il faut regarder ce qui SORT.
# =============================================================================

import inspect
import io
import re
import os
import unittest
from datetime import date, timedelta

import numpy as np

from direction_non_vie.provisionnement.a7_provisionnement import (
    n5_commentaire as COMM_MOD)
from direction_non_vie.provisionnement.a7_provisionnement import (
    n5_excel as XL_MOD)
from direction_non_vie.provisionnement.a7_provisionnement import (
    n5_rapport as RAPP_MOD)
from direction_non_vie.provisionnement.a7_provisionnement.methodes_be import (
    disponible,
    libelle,
    motif_exclusion,
)
from direction_non_vie.provisionnement.a7_provisionnement.n5_commentaire import (
    _LOB_HANDLERS,
    _narration_lob,
    generer_commentaire,
)
from direction_non_vie.provisionnement.a7_provisionnement.config.lob_config import (
    LOB_CONFIG,
    get_lob_config,
)

from direction_non_vie.provisionnement.a7_provisionnement.agent import (
    _TAILLE_MIN_LIVRABLE,
    PORTEE_ARCHIVE,
    AgentA7Provisionnement,
    _dependance_absente,
    _produire_livrable,
    etiquette_methode_grands,
    verifier_archive,
)
from direction_non_vie.provisionnement.a7_provisionnement.config.lob_config import (
    LOB_CONFIG,
    get_lob_config,
)
from direction_non_vie.provisionnement.a7_provisionnement.config.rfr_eiopa import (
    DATE_COURBE,
    MOIS_ALERTE_PEREMPTION,
    MOIS_ROUGE_PEREMPTION,
    age_courbe_mois,
    diagnostic_peremption,
    get_courbe_embarquee,
)
from direction_non_vie.provisionnement.a7_provisionnement.n4_best_estimate import (
    CLE_BOOT,
    CLE_COMPOSE,
    CLE_MACK,
    LIBELLE_APPROCHE,
    libelle_percentiles,
    marque_retenue,
)
from direction_non_vie.provisionnement.a7_provisionnement.n5_excel import (
    export_excel,
)
from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
    ARRETE_ABSENT,
    MARQUEUR_ECHEC_RAPPORT,
    _build_blocks,
    _construire_contexte,
    _md_to_html,
    _nettoyer_narration,
    etat_calendaire,
    export_html,
    export_word,
)
from direction_non_vie.provisionnement.a7_provisionnement.test_a7_graphiques import (
    kaleido_declare,
    rendeur_substitue,
)
from direction_non_vie.provisionnement.a7_provisionnement.test_a7_ibrahim import (
    GENINS,
)


def _exposition(triangle, loss_ratio=0.70):
    C = np.array(triangle, dtype=float)
    return C.max(axis=1) * 1.6 / loss_ratio


# =============================================================================
#  1. LES LIVRABLES DÉCLARENT LEUR ÉCHEC — ILS NE LE TAISENT PLUS
# =============================================================================

class T1_Livrables_Declarent(unittest.TestCase):
    """Un livrable dégradé est DÉCLARÉ dans le résultat, jamais deviné."""

    @classmethod
    def setUpClass(cls):
        # ⚠️ CE VERROU PORTE SUR LA DÉCLARATION DES LIVRABLES, pas sur les
        # pixels. Il faisait rasteriser 14 figures — 103 s mesurées — pour
        # inspecter un classeur Excel.
        with kaleido_declare(True), rendeur_substitue():
            cls.r = AgentA7Provisionnement(verbose=False).run(
                source=np.array(GENINS, dtype=float), mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=300, seed=42,
                generer_graphiques=True, generer_word=True)

    def test_le_resultat_porte_un_bilan_des_livrables(self):
        self.assertTrue(self.r.get('success'), self.r.get('erreur'))
        self.assertIn('livrables_erreurs', self.r,
                      "le résultat doit porter le bilan des livrables")
        self.assertIsInstance(self.r['livrables_erreurs'], dict)
        print(f"    OK LIV-1 bilan présent : {self.r['livrables_erreurs'] or 'aucune erreur'}")

    def test_excel_est_un_vrai_classeur_pas_un_repli(self):
        """Le seul export dont la dépendance est toujours là : il doit sortir."""
        octets = self.r.get('excel_bytes', b'')
        self.assertGreater(len(octets), _TAILLE_MIN_LIVRABLE,
                           f"export_excel a rendu {len(octets)} octets — repli ?")
        # Un .xlsx est un ZIP : les deux premiers octets valent 'PK'.
        self.assertEqual(octets[:2], b'PK',
                         "ce ne sont pas les octets d'un classeur Excel")
        self.assertNotIn('excel', self.r['livrables_erreurs'])
        print(f"    OK LIV-2 Excel : {len(octets):,} octets, en-tête ZIP conforme")

    def test_un_livrable_vide_est_toujours_explique(self):
        """Aucun octet manquant sans motif — c'était exactement le trou."""
        for cle, nom in (('excel_bytes', 'excel'), ('word_bytes', 'word'),
                         ('html', 'html')):
            octets = self.r.get(cle, b'')
            if len(octets) >= _TAILLE_MIN_LIVRABLE:
                continue
            self.assertIn(nom, self.r['livrables_erreurs'],
                          f"{nom} rend {len(octets)} octets SANS motif déclaré")
            motif = self.r['livrables_erreurs'][nom]
            self.assertTrue(
                motif.startswith(('dependance_absente:', 'echec:', 'vide:')),
                f"motif non classé pour {nom} : {motif!r}")
            print(f"    OK LIV-3 {nom} vide, et le motif le dit : {motif}")

    def test_une_dependance_absente_nest_pas_un_echec_de_code(self):
        """La distinction qui manquait : bibliothèque absente ≠ bug."""
        for nom in ('word', 'pdf'):
            motif = self.r['livrables_erreurs'].get(nom)
            if motif is None:
                continue
            if _dependance_absente(nom):
                self.assertTrue(motif.startswith('dependance_absente:'),
                                f"{nom} : dépendance absente mais motif {motif!r}")
            else:
                self.assertFalse(
                    motif.startswith('dependance_absente:'),
                    f"{nom} : dépendance présente, le motif ne peut pas "
                    f"l'incriminer — {motif!r}")
        print("    OK LIV-4 dépendance absente et échec de code sont distingués")

    def test_un_echec_reel_est_classe_comme_tel(self):
        """Chemin d'échec exercé pour de vrai, pas seulement décrit."""
        def fabrique_qui_casse(**kw):
            raise TypeError('boum')
        octets, err = _produire_livrable('excel', fabrique_qui_casse)
        self.assertEqual(octets, b'')
        self.assertTrue(err.startswith('echec: TypeError'), err)

        octets, err = _produire_livrable('excel', lambda **kw: b'PK\x03\x04')
        self.assertTrue(err.startswith('vide:'),
                        f"4 octets doivent être classés 'vide:' — {err!r}")

        octets, err = _produire_livrable(
            'excel', lambda **kw: b'PK' + b'\x00' * _TAILLE_MIN_LIVRABLE)
        self.assertIsNone(err, "un livrable de taille normale n'a pas de motif")
        print("    OK LIV-5 les trois classements exercés : echec / vide / OK")


# =============================================================================
#  1bis. LE RAPPORT HTML — ET LE PDF QU'IL ALIMENTE
# =============================================================================

class T2_Rapport_HTML_Et_PDF(unittest.TestCase):
    """`export_html` rattrape ses échecs et rend une page d'erreur VALIDE."""

    def test_le_rapport_reel_ne_porte_pas_le_marqueur_dechec(self):
        r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=300, seed=42,
            generer_graphiques=False)
        html = export_html(r['n1'], r['n2'], r['n3'], r['n4'], {},
                           ref_client='test')
        self.assertFalse(html.startswith(MARQUEUR_ECHEC_RAPPORT),
                         "export_html est retombé sur son repli d'échec")
        self.assertGreater(len(html), 20_000, f"{len(html)} octets — repli ?")
        print(f"    OK LIV-6 rapport HTML : {len(html):,} octets, sans marqueur")

    def test_le_repli_porte_un_marqueur_detectable(self):
        """Sans marqueur, la page d'erreur est un HTML que rien ne distingue."""
        # n3 dépourvu des clés attendues → _build_blocks lève → repli.
        html = export_html({}, {}, {}, {}, {}, ref_client='casse')
        if html.startswith(MARQUEUR_ECHEC_RAPPORT):
            print(f"    OK LIV-7 repli marqué et détectable ({len(html)} octets)")
        else:
            # Le rapport a tenu sur des dicts vides : c'est légitime, mais alors
            # il doit être un VRAI rapport, pas un entre-deux silencieux.
            self.assertGreater(len(html), 20_000,
                               f"ni marqué comme repli, ni un rapport complet "
                               f"({len(html)} octets)")
            print(f"    OK LIV-7 aucun repli déclenché ({len(html):,} octets)")

    def test_un_repli_de_rapport_reste_trop_petit_pour_passer(self):
        """⚠️ CE TEST VISAIT `export_pdf`, RETIRE AU LOT C1 (decision B).

        Sa preoccupation survit et elle est intacte : `export_pdf` GONFLAIT le
        repli d'echec en un PDF volumineux et parfaitement valide, ce qui
        defaisait le controle de taille de `_produire_livrable`. Le PDF n'etant
        plus genere, le repli reste une chaine courte -- et la garde le
        rattrape. C'est ce que ce test verifie desormais : la protection tient
        PAR CONSTRUCTION, non par un controle supplementaire.
        """
        import direction_non_vie.provisionnement.a7_provisionnement.agent as ag
        repli = MARQUEUR_ECHEC_RAPPORT + '<html><body><h1>Erreur</h1></body></html>'
        self.assertLess(
            len(repli), ag._TAILLE_MIN_LIVRABLE,
            "le repli d'echec doit rester sous le seuil qui le disqualifie")
        octets, erreur = ag._produire_livrable('html', lambda **k: repli)
        self.assertTrue(erreur, "un repli doit etre declare comme anomalie")
        self.assertIn('vide', erreur)
        print(f"    OK LIV-8 le repli d'echec ({len(repli)} car.) reste sous "
              f"le seuil ({ag._TAILLE_MIN_LIVRABLE}) et est declare : {erreur}")


# =============================================================================
#  2. COURBE DES TAUX — LA PÉREMPTION EST MESURÉE, PAS SUPPOSÉE
# =============================================================================

class T3_Peremption_Courbe(unittest.TestCase):
    """La courbe embarquée ne peut plus se déclarer sans reproche."""

    def test_les_trois_statuts_sont_atteignables(self):
        courbe = date.fromisoformat(DATE_COURBE)
        cas = [
            (courbe + timedelta(days=15),  'VERT'),
            (courbe + timedelta(days=int(MOIS_ALERTE_PEREMPTION * 31)), 'AMBRE'),
            (courbe + timedelta(days=int(MOIS_ROUGE_PEREMPTION * 31)),  'ROUGE'),
        ]
        for quand, attendu in cas:
            d = diagnostic_peremption(quand)
            self.assertEqual(d['statut'], attendu,
                             f"{quand} → {d['statut']} ({d['age_mois']} mois)")
        print("    OK LIV-9 péremption : VERT / AMBRE / ROUGE tous atteignables")

    def test_un_arrete_passe_est_juge_a_sa_date(self):
        """Un arrêté contemporain de la courbe ne doit pas la déclarer périmée."""
        d = diagnostic_peremption(DATE_COURBE)
        self.assertEqual(d['statut'], 'VERT')
        self.assertLess(d['age_mois'], 1.0)
        print(f"    OK LIV-10 arrêté du {DATE_COURBE} : VERT (0 mois)")

    def test_la_courbe_embarquee_signale_son_age_aujourdhui(self):
        """Le cœur du point : plus de 'erreur': None inconditionnel.

        ⚠️ CE TEST N'ASSERTAIT PLUS RIEN APRÈS LA BASCULE DU LOT R2, ET IL
        PASSAIT QUAND MÊME. Son unique assertion vivait sous
        `if mois >= MOIS_ALERTE_PEREMPTION` ; la courbe embarquée étant
        passée du 31/03/2025 au 31/07/2026, la branche n'était plus prise et
        le test devenait creux — vert, silencieux, et sans valeur. Il vérifie
        désormais LES DEUX SENS, donc il ne peut plus se vider tout seul.
        """
        c = get_courbe_embarquee()
        self.assertIn('peremption', c)
        mois = age_courbe_mois()
        if mois >= MOIS_ALERTE_PEREMPTION:
            self.assertIsNotNone(
                c['erreur'],
                "la courbe est périmée et ne le déclare pas — c'était le bug")
        else:
            self.assertIsNone(
                c['erreur'],
                "la courbe est fraîche et se déclare pourtant en défaut — "
                "un reproche sans motif use le signal")
        print(f"    OK LIV-11 courbe du {DATE_COURBE} : {mois:.0f} mois, "
              f"statut {c['peremption']['statut']}, erreur "
              f"{'absente' if c['erreur'] is None else 'présente'}")

    def test_la_peremption_atteint_le_resultat_et_les_alertes(self):
        r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            n_sim_bootstrap=300, seed=42, generer_graphiques=False)
        diag = (r['n4'].get('peremption_courbe')
                or (r['n4'].get('risk_margin_data') or {}).get('peremption_courbe'))
        self.assertIsNotNone(diag, "le diagnostic doit remonter jusqu'à N4")
        alertes = ' '.join(str(a) for a in r['n4'].get('alertes', []))
        # ⚠️ MÊME DÉFAUT QUE LIV-11, ET MÊME CORRECTIF. L'assertion vivait
        # sous `if diag['statut'] in ('AMBRE','ROUGE')` ; la courbe embarquée
        # étant devenue fraîche au lot R2, la branche n'était plus prise et le
        # test passait sans rien vérifier. Les deux sens sont désormais
        # exigés : une courbe en défaut doit alerter, une courbe à jour ne
        # doit PAS — une alerte sans motif use le signal aussi sûrement qu'une
        # alerte manquante.
        if diag['statut'] in ('AMBRE', 'ROUGE'):
            self.assertIn('courbe', alertes.lower(),
                          "la péremption doit apparaître dans les alertes N4")
        else:
            self.assertNotIn('périmée', alertes.lower(),
                             "la courbe est à jour et N4 alerte quand même")
        print(f"    OK LIV-12 N4 porte le diagnostic : {diag['statut']} "
              f"({diag['age_mois']} mois)")


# =============================================================================
#  3 & 4. LES DEUX AFFICHAGES FAUX DE L'APPLICATION
# =============================================================================

class T4_Affichages_Application(unittest.TestCase):

    def test_munich_expose_bien_ses_deux_reserves(self):
        """`be_munich` n'existe pas — l'app affichait « 0 € ✅ ».

        MUNICH EST RÉELLEMENT EXERCÉE ICI. Il y faut DEUX conditions, et rater
        l'une des deux rend le test creux : un triangle d'engagés, et une LoB où
        `munich_cl_disponible` est vrai — sur 'generique' la méthode est
        désactivée, et le premier jet de ce test ne prouvait donc rien.
        """
        from direction_non_vie.provisionnement.a7_provisionnement.test_a7_ibrahim import (
            _MCL_ENG_SAIN,
            _MCL_PAYE,
        )
        r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(_MCL_PAYE, dtype=float), mode_declare='cumule',
            triangle_engage=np.array(_MCL_ENG_SAIN, dtype=float),
            lob='rc_auto_corporels',
            n_sim_bootstrap=300, seed=42, generer_graphiques=False)
        m = r['n3']['munich_cl']
        self.assertTrue(m.get('disponible'),
                        "le scénario doit RÉELLEMENT exercer Munich")
        self.assertNotIn('be_munich', m,
                         "si cette clé apparaît un jour, l'app doit la lire")
        for cle in ('be_munich_paye', 'be_munich_engage'):
            self.assertIn(cle, m, f"clé {cle} attendue par l'application")
            self.assertNotEqual(m[cle], 0.0,
                                f"{cle} nul : le scénario n'exerce rien")
        print(f"    OK LIV-13 Munich exercée : payé={m['be_munich_paye']:,.2f} "
              f"engagé={m['be_munich_engage']:,.2f} (l'app affichait 0 €)")

    def test_letiquette_grands_sinistres_suit_le_resultat(self):
        """Elle annonçait « BF auto » là où seul Chain Ladder était retenu."""
        f = etiquette_methode_grands
        code, lbl = f(['chain_ladder'])
        self.assertEqual(code, 'cl_separe')
        self.assertIn('Chain Ladder', lbl)
        self.assertNotIn('Bornhuetter', lbl)

        code, lbl = f(['bornhuetter_ferguson', 'chain_ladder'])
        self.assertEqual(code, 'bf_auto')
        self.assertIn('Bornhuetter-Ferguson', lbl)

        code, lbl = f([])
        self.assertEqual(code, 'manuel')
        print("    OK LIV-14 étiquette grands sinistres déduite du résultat")

    def test_un_triangle_sans_exposition_ne_produit_pas_de_bf(self):
        """La prémisse du correctif, vérifiée et non supposée."""
        r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            n_sim_bootstrap=300, seed=42, generer_graphiques=False)
        inc = r['n4']['methodes_incluses']
        self.assertNotIn('bornhuetter_ferguson', inc)
        code, _ = etiquette_methode_grands(inc)
        self.assertEqual(code, 'cl_separe',
                         "sans exposition, l'étiquette ne peut pas dire BF")
        print(f"    OK LIV-15 sans exposition : incluses={inc} → cl_separe")

    def test_plus_aucune_citation_du_guide_inexistante(self):
        """« §3.2 » ne correspond à aucune section du guide (numérotée en lettres)."""
        # `actuaria_app` n'est PAS importable dans la gate (streamlit absent) :
        # on lit son source comme un fichier. C'est le seul moyen de couvrir ce
        # fichier, et c'est aussi pourquoi il faut y laisser le moins de logique
        # possible.
        import io as _io
        from pathlib import Path
        racine = Path(__file__).resolve().parents[3]
        # ⚠️ CE BALAYAGE PORTAIT SUR DEUX FICHIERS, ET LE SECOND N'EXISTE
        # PLUS : `nv_triangle_builder.py` a été supprimé avec l'ancien chemin
        # de préparation des données. ⚠️ Il était ouvert PAR SON CHEMIN, pas
        # importé — sa disparition ne se serait donc pas vue à l'import mais
        # par un `FileNotFoundError` à l'exécution, et aucun relevé par AST
        # ne l'aurait signalée. C'est un relevé par CHAÎNE qui l'a trouvé.
        for chemin in (racine / 'actuaria_app.py',):
            src = _io.open(chemin, encoding='utf-8').read()
            self.assertNotIn(
                'Guide IA 2023 §3.', src,
                f"{chemin.name} cite une section inexistante du guide — les "
                f"sinistres graves sont traités en §4.c p36-37")
        print("    OK LIV-16 citations du guide corrigées (§4.c p36-37)")


# =============================================================================
#  5. LES DEUX LoB PRESQUE HOMONYMES
# =============================================================================

class T5_LoB_Distinguables(unittest.TestCase):
    """Elles partagent la LoB S2 et le σ, mais pas le régime de queue.

    Le renommage a été ÉCARTÉ après mesure du périmètre : `dommage_corporel_
    individuel` est une clé d'A10 (table σ et table de correspondance), figure
    dans le snapshot T7 du filet historique, et les deux noms sont des termes
    professionnels corrects. Renommer aurait désynchronisé A7 et A10 sans
    traiter le vrai risque — un choix fait sans voir la conséquence. Ce que ces
    tests verrouillent, c'est que la conséquence soit VISIBLE.
    """

    PAIRE = ('accidents_corporels', 'dommage_corporel_individuel')

    def test_elles_diffèrent_bien_sur_le_regime_de_queue(self):
        a, b = (get_lob_config(k) for k in self.PAIRE)
        self.assertNotEqual(a['risque_long'], b['risque_long'])
        # `lob_eiopa` (texte libre) a fusionné avec `sigma_eiopa` dans un champ
        # unique `segment_s2` au lot B10-a : les deux se contredisaient sur
        # trois LoB sur quinze. L'intention du test est inchangée, elle est
        # seulement exprimée sur le segment officiel plutôt que sur un libellé.
        self.assertEqual(a['segment_s2'], b['segment_s2'],
                         "même segment Solvabilité II — la différence est la queue")
        self.assertEqual(a['sigma_eiopa'], b['sigma_eiopa'])
        print(f"    OK LIV-17 même LoB S2 et même σ, risque_long "
              f"{a['risque_long']} contre {b['risque_long']}")

    def test_leurs_libelles_disent_lequel_est_long(self):
        for cle in self.PAIRE:
            cfg = get_lob_config(cle)
            libelle = cfg['label'].upper()
            attendu = 'LONGUE' if cfg['risque_long'] else 'COURTE'
            self.assertIn(attendu, libelle,
                          f"{cle} : le libellé doit dire son régime de queue")
        self.assertNotEqual(get_lob_config(self.PAIRE[0])['label'],
                            get_lob_config(self.PAIRE[1])['label'])
        print("    OK LIV-18 les libellés portent COURTE / LONGUE")

    def test_la_distinction_remonte_jusqua_lactuaire(self):
        """Une configuration que personne ne lit ne protège de rien."""
        for cle in self.PAIRE:
            r = AgentA7Provisionnement(verbose=False).run(
                source=np.array(GENINS, dtype=float), mode_declare='cumule',
                lob=cle, n_sim_bootstrap=300, seed=42, generer_graphiques=False)
            infos = ' '.join(str(i) for i in r['n2'].get('infos', []))
            self.assertIn('LoB retenue', infos,
                          f"{cle} : la distinction n'atteint pas les infos N2")
            self.assertIn('risque_long', infos)
        print("    OK LIV-19 la distinction apparaît dans les infos N2")

    def test_toute_lob_portant_distinction_la_voit_remonter(self):
        """Invariant général, pas un cas particulier codé en dur."""
        avec = [k for k, v in LOB_CONFIG.items() if v.get('distinction')]
        self.assertGreaterEqual(len(avec), 2)
        for cle in avec:
            self.assertIn(cle, LOB_CONFIG)
            self.assertTrue(LOB_CONFIG[cle]['distinction'].strip())
        print(f"    OK LIV-20 {len(avec)} LoB portent une distinction explicite")


# =============================================================================
#  LOT B (releve B2) — UN ARRETE NON COMMUNIQUE NE DEVIENT PAS LA DATE DU JOUR
# =============================================================================
#
#  ⚠️ LE FILET COUVRE LES TROIS LIVRABLES, PAS LE MODULE, ET C'EST LA LECON DU
#  LOT. A4a avait ferme ce defaut dans `n5_commentaire` et cru le lot complet ;
#  la correction n'avait atteint ni le HTML ni le Word, qui ecrivaient encore
#  `arr = arrete or dt` -- la date du jour sous une etiquette << Arrete >>.
#  Un filet borne au module l'aurait laisse revenir par un autre export.
#
#  ⚠️ ET LA TRACE PAR CONSOMMATEUR A CORRIGE MON PROPRE RELEVE. J'avais compte
#  l'Excel comme troisieme site du meme faux. Il ne l'est PAS : son `date_str`
#  est etiquete << Date rapport >> (la date de generation, pour laquelle
#  aujourd'hui est JUSTE quand l'arrete manque), pas << Arrete >>. Le variable
#  matchait le grep ; le consommateur dit autre chose. L'Excel porte un autre
#  defaut, plus doux -- il confond arrete et date de generation dans un seul
#  champ quand l'arrete EST fourni -- nomme a l'ardoise, hors de ce lot.
#  Le filet le VERIFIE quand meme : il s'assure que l'Excel ne presente jamais
#  la date du jour SOUS une etiquette d'arrete.
# =============================================================================

def _today_fr():
    # ⚠️ `datetime.now()` NAIF, ET C'EST DELIBERE (noqa DTZ005). Ce helper doit
    # reproduire A L'IDENTIQUE la date que la production calcule dans
    # `n5_rapport` (`datetime.now().strftime('%d/%m/%Y')`, sans fuseau) : la
    # comparaison << la date du jour n'apparait pas comme arrete >> n'a de sens
    # que si les deux chaines sont produites de la meme facon. Un tz les ferait
    # diverger aux frontieres de journee.
    # ⚠️ `noqa` MAL CIBLE AU LOT B, ET LE CONSTAT VAUT D'ETRE ECRIT : la
    # directive portait `DTZ005` seul, que `ruff` nu n'active pas -- il la
    # declarait donc INUTILE (RUF100) pendant que `scripts/proprete.py`, lui,
    # active DTZ005 et l'exigeait. Deux outils, deux configurations, une seule
    # ligne : il faut les deux codes pour satisfaire les deux.
    from datetime import datetime
    return datetime.now().strftime('%d/%m/%Y')  # noqa: DTZ005, RUF100


class T_Arrete_Les_Trois_Livrables_Ne_Fabriquent_Pas_La_Date(unittest.TestCase):
    """⚠️ TROIS LIVRABLES, UN INVARIANT : la date du jour ne se fait jamais
    passer pour l'arrete quand aucun arrete n'est communique."""

    @classmethod
    def setUpClass(cls):
        # ⚠️ UN RUN REEL, RE-EXPORTE AVEC L'ARRETE CHOISI. Bricoler un `n4`
        # partiel ferait lever l'Excel (`n4['reserve_p75']`, acces direct) sur
        # une cause etrangere au lot. On part des vrais objets et on ne fait
        # varier QUE l'arrete.
        with kaleido_declare(True), rendeur_substitue():
            r = AgentA7Provisionnement(verbose=False).run(
                source=np.array(GENINS, dtype=float), mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=200, seed=42,
                generer_graphiques=False)
        cls.n1, cls.n2 = r['n1'], r['n2']
        cls.n3, cls.n4 = r['n3'], r['n4']
        cls.C = np.array(GENINS, dtype=float)

    def _formes_fabriquees(self):
        t = _today_fr()
        return (f'Arrêté au {t}', f'Arrêté {t}', f'Arrêté : {t}')

    def test_html_nomme_l_absence_au_lieu_de_la_combler(self):
        html = export_html(self.n1, self.n2, self.n3, self.n4,
                           arrete='', ref_client='ACME')
        self.assertIn(ARRETE_ABSENT, html)
        for forme in self._formes_fabriquees():
            self.assertNotIn(forme, html, f'HTML fabrique : {forme!r}')
        print('    OK LIV-B1 le HTML dit « non communiqué », pas la date du jour')

    def test_html_garde_la_vraie_date_quand_elle_existe(self):
        html = export_html(self.n1, self.n2, self.n3, self.n4,
                           arrete='30/06/2026', ref_client='ACME')
        # ⚠️ NE PAS SUR-CORRIGER : un arrete fourni doit s'afficher, avec « au ».
        self.assertIn('Arrêté au 30/06/2026', html)
        print('    OK LIV-B2 le HTML garde l arrete fourni, avec « au »')

    def test_word_nomme_l_absence(self):
        import io

        from docx import Document
        w = export_word(self.n1, self.n2, self.n3, self.n4,
                        arrete='', ref_client='ACME')
        cells = [c.text for t in Document(io.BytesIO(w)).tables
                 for r in t.rows for c in r.cells]
        self.assertIn('Arrêté', cells)                    # l'etiquette existe
        self.assertIn(ARRETE_ABSENT, cells)               # la valeur nomme l'absence
        self.assertNotIn(_today_fr(), cells, 'le Word porte la date du jour')
        print('    OK LIV-B3 le Word dit « non communiqué » dans sa table')

    def test_excel_n_etiquette_jamais_la_date_du_jour_en_arrete(self):
        # ⚠️ L'EXCEL N'EST PAS LE MEME FAUX (cf. l'en-tete). On verrouille qu'il
        # ne le DEVIENNE pas : la date du jour peut figurer, mais jamais sous
        # une etiquette d'arrete.
        import io

        from openpyxl import load_workbook
        xl = export_excel(self.C, self.n1, self.n2, self.n3, self.n4,
                          ref_client='ACME', arrete='')
        wb = load_workbook(io.BytesIO(xl))
        textes = [str(c.value) for ws in wb.worksheets
                  for row in ws.iter_rows() for c in row if c.value is not None]
        self.assertTrue(any('Date rapport' in t for t in textes),
                        "l'etiquette « Date rapport » a disparu")
        for forme in self._formes_fabriquees():
            self.assertFalse(any(forme in t for t in textes),
                             f'Excel etiquette la date du jour en arrete : {forme!r}')
        print('    OK LIV-B4 l Excel etiquette la date du jour « Date rapport »')


# =============================================================================
#  F-cal (releve B2) — UN TEST CALENDAIRE NON FAIT N'EST PAS UN << AUCUN EFFET >>
# =============================================================================
#
#  ⚠️ LE FAUX DE F3, TRANSPOSE AU CALENDAIRE. Le Word (document SIGNE) et le
#  prompt LLM deduisaient << Aucun effet calendaire significatif >> / << non
#  significatif >> du SEUL compteur d'effets, sans regarder si le GLM Poisson
#  age-cohorte avait pu s'ajuster. Sur un triangle ou il ne s'ajuste pas, ils
#  affirmaient une absence d'EFFET la ou il n'y avait qu'absence de TEST. Le
#  HTML distinguait deja les deux (FIX 2) -- encore une correction qui n'avait
#  pas atteint tous ses sites.
#
#  ⚠️ LE FILET PORTE SUR LES TROIS CONSOMMATEURS, PAS SUR LE MODULE, et le
#  remede est une SOURCE UNIQUE (`etat_calendaire`) plutot qu'une troisieme
#  copie de la logique -- c'est la duplication qui a cause la divergence.

#: Les quatre etats du test calendaire et un `bz` (glm_apc) qui les produit.
#: Au niveau du module, comme N4_GARDE : un dict mutable en attribut de classe
#: demanderait un ClassVar, et il n'appartient pas plus a une classe qu'a une
#: autre.
ETATS_CALENDAIRE = {
    'indisponible': {'glm_disponible': False},
    'aucun':        {'glm_disponible': True, 'n_effets_significatifs': 0,
                     'cal_significatif': False},
    'diffus':       {'glm_disponible': True, 'n_effets_significatifs': 0,
                     'cal_significatif': True},
    'present':      {'glm_disponible': True, 'n_effets_significatifs': 2},
}


class T_Calendaire_Non_Teste_N_Est_Pas_Aucun_Effet(unittest.TestCase):
    """⚠️ TROIS RENDUS (HTML, Word, prompt LLM), UNE SOURCE : un test absent
    ne se rend jamais comme une absence d'effet."""

    @classmethod
    def setUpClass(cls):
        with kaleido_declare(True), rendeur_substitue():
            cls.r = AgentA7Provisionnement(verbose=False).run(
                source=np.array(GENINS, dtype=float), mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=200, seed=42,
                generer_graphiques=False)

    def _word_paras(self, bz):
        import io

        from docx import Document
        n3 = dict(self.r['n3']); n3['glm_apc'] = bz
        w = export_word(self.r['n1'], self.r['n2'], n3, self.r['n4'],
                        arrete='30/06/2026', ref_client='ACME')
        return ' '.join(p.text for p in Document(io.BytesIO(w)).paragraphs)

    def _html(self, bz):
        n3 = dict(self.r['n3']); n3['glm_apc'] = bz
        return export_html(self.r['n1'], self.r['n2'], n3, self.r['n4'],
                           arrete='30/06/2026', ref_client='ACME')

    def test_l_etat_a_quatre_valeurs_exactes(self):
        for attendu, bz in ETATS_CALENDAIRE.items():
            self.assertEqual(etat_calendaire(bz), attendu, f'etat {attendu}')
        print('    OK CAL-1 les quatre etats de etat_calendaire sont exacts')

    def test_word_ne_dit_pas_aucun_effet_quand_le_test_est_absent(self):
        joined = self._word_paras(ETATS_CALENDAIRE['indisponible'])
        self.assertNotIn('Aucun effet calendaire significatif', joined)
        self.assertIn('indisponible', joined.lower())
        print('    OK CAL-2 le Word dit indisponible, jamais aucun effet')

    def test_word_dit_bien_aucun_effet_quand_le_test_a_conclu(self):
        # ⚠️ NE PAS SUR-CORRIGER : un vrai << aucun effet >> reste vert.
        joined = self._word_paras(ETATS_CALENDAIRE['aucun'])
        self.assertIn('Aucun effet calendaire significatif', joined)
        print('    OK CAL-3 le Word dit aucun effet quand le test a conclu')

    def test_le_prompt_llm_ne_conclut_pas_sur_un_test_absent(self):
        ctx = _construire_contexte({}, {'glm_apc': ETATS_CALENDAIRE['indisponible']},
                                   {}, 'RC', '30/06')
        ligne = next(x for x in ctx.splitlines() if 'test calendaire' in x)
        self.assertIn('INDISPONIBLE', ligne)
        self.assertNotIn('non significatif', ligne)
        print('    OK CAL-4 le prompt LLM ne conclut pas sur un test absent')

    def test_html_distingue_toujours_indisponible_de_aucun(self):
        # ⚠️ NON-REGRESSION : le HTML distinguait deja, il doit rendre pareil.
        self.assertIn('Test calendaire indisponible',
                      self._html(ETATS_CALENDAIRE['indisponible']))
        self.assertIn('Aucun effet significatif',
                      self._html(ETATS_CALENDAIRE['aucun']))
        print('    OK CAL-5 le HTML distingue toujours indisponible et aucun')


# =============================================================================
#  AVIS-COULEUR (releve B2) — LA COULEUR SUIT LE STATUT, PAS UN MOT
# =============================================================================
#
#  ⚠️⚠️ LA COULEUR ETAIT TOUJOURS VERTE, PAS << verte sauf si >>. Les deux
#  renderers coloraient l'avis en rouge SI le mot << DEFAVORABLE >> figurait
#  dans `avis_actuariel` -- un mot que ce champ NE PRODUIT JAMAIS. Ses trois
#  valeurs (n4_best_estimate l. 1315-1322) commencent TOUTES par
#  << FAVORABLE >>. Le test etait donc toujours faux : l'avis d'un dossier
#  ROUGE (<< FAVORABLE SOUS RESERVE -- revisions requises avant bilan S2 >>)
#  s'affichait EN VERT dans le HTML ET dans le Word signe.
#
#  Le controle lisait le vocabulaire d'un AUTRE champ : c'est `jugement`
#  (l. 1919) qui ecrit << AVIS DEFAVORABLE >>.
#
#  ⚠️ LE REMEDE TRANSPOSE `_RAG_CELLULE` : match exact sur un vocabulaire FINI
#  (VERT / AMBRE / ROUGE), jamais une recherche de mot dans du texte libre.
#  Un statut inconnu tombe sur l'orange, jamais sur le vert.
#
#  ⚠️ ET CE LOT A REVELE UN BUG DU LOT C : la boucle des cartes d'hypotheses
#  ECRASAIT le parametre `statut` de `_build_blocks`. Invisible tant que rien
#  ne lisait `statut` en aval ; avis-couleur l'a fait pour la premiere fois.
#  Le verrou ci-dessous est STRUCTUREL -- aucun parametre reaffecte -- parce
#  qu'un test sur la seule couleur n'aurait pas empeche le prochain ecrasement.

#: Les couleurs attendues, par statut RAG. Cote Word ce sont les RGB de
#: `n5_rapport` (VR / AR / RgR), cote HTML les variables CSS.
_AVIS_ATTENDU = {
    'VERT':  ('var(--vert)',   '1E8449'),
    'AMBRE': ('var(--orange)', 'E67E22'),
    'ROUGE': ('var(--rouge)',  'C0392B'),
}


class T_Avis_Couleur_Suit_Le_Statut(unittest.TestCase):
    """⚠️ DEUX LIVRABLES, UN INVARIANT : la couleur de l'avis vient du statut
    RAG, jamais d'un mot cherche dans le texte."""

    @classmethod
    def setUpClass(cls):
        with kaleido_declare(True), rendeur_substitue():
            cls.r = AgentA7Provisionnement(verbose=False).run(
                source=np.array(GENINS, dtype=float), mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=200, seed=42,
                generer_graphiques=False)

    def _n4(self, statut):
        n4 = dict(self.r['n4']); n4['statut'] = statut
        return n4

    def test_le_champ_avis_est_produit(self):
        # ⚠️ CE TEST CONSTATAIT << l'avis ne dit JAMAIS DEFAVORABLE >> -- la
        # mesure qui fondait le lot avis-couleur. Le lot du double vocabulaire
        # a rendu ce constat FAUX A DESSEIN : un dossier ROUGE dit desormais
        # << DEFAVORABLE >>. Le test a donc fait son office et cede la place ;
        # ce qui reste verifiable ici, c'est qu'un avis EXISTE. La coherence
        # avec `jugement` est verrouillee par la classe dediee plus bas.
        avis = str(self.r['n4'].get('avis_actuariel', ''))
        self.assertTrue(avis, "n4 ne produit aucun avis actuariel")
        print(f'    OK AVIS-1 un avis est produit : {avis[:44]}…')

    def test_html_colore_l_avis_selon_le_statut(self):
        import re
        # ⚠️ INDEX PLUTOT QUE DEBALLAGE : `for statut, (css, _rgb)` laissait une
        # locale assignee et jamais relue -- l'angle mort de ruff, que le
        # controle de proprete du depot attrape.
        for statut, attendu in _AVIS_ATTENDU.items():
            css = attendu[0]
            b = _build_blocks(self.r['n2'], self.r['n3'], self._n4(statut),
                              '', 'aucune', 'RC', 'X', '30/06', '18/08',
                              'A7-1', 'chain_ladder', statut, {})
            trouve = re.search(r'background:(var\(--[a-z]+\))', b['avis'])
            self.assertIsNotNone(trouve, f'{statut} : encadre d avis absent')
            self.assertEqual(trouve.group(1), css, f'statut {statut}')
        print('    OK AVIS-2 le HTML colore l avis selon VERT/AMBRE/ROUGE')

    def test_word_colore_l_avis_selon_le_statut(self):
        import io

        from docx import Document
        avis = str(self.r['n4'].get('avis_actuariel', '')).strip()[:20]
        for statut, attendu in _AVIS_ATTENDU.items():
            rgb = attendu[1]
            w = export_word(self.r['n1'], self.r['n2'], self.r['n3'],
                            self._n4(statut), arrete='30/06/2026',
                            ref_client='ACME')
            cols = [str(run.font.color.rgb)
                    for p in Document(io.BytesIO(w)).paragraphs
                    for run in p.runs if run.text.strip()[:20] == avis]
            self.assertTrue(cols, f'{statut} : avis introuvable dans le Word')
            self.assertEqual(cols[0], rgb, f'statut {statut}')
        print('    OK AVIS-3 le Word colore l avis selon VERT/AMBRE/ROUGE')

    def test_un_dossier_rouge_n_est_jamais_vert(self):
        # ⚠️ LE FAUX EXACT, ENONCE EN TOUTES LETTRES pour qu'une recherche
        # plein texte le retrouve s'il revenait.
        b = _build_blocks(self.r['n2'], self.r['n3'], self._n4('ROUGE'),
                          '', 'aucune', 'RC', 'X', '30/06', '18/08',
                          'A7-1', 'chain_ladder', 'ROUGE', {})
        self.assertNotIn('background:var(--vert)', b['avis'])
        print('    OK AVIS-4 un dossier ROUGE n affiche jamais un avis vert')

    def test_aucun_parametre_de_build_blocks_n_est_ecrase(self):
        # ⚠️ VERROU STRUCTUREL, PAS COSMETIQUE. Le lot C avait ecrase `statut`
        # dans la boucle des cartes d'hypotheses : tout code lisant `statut`
        # en aval recevait le verdict de la DERNIERE hypothese. Un test sur la
        # seule couleur n'aurait pas empeche le prochain ecrasement.
        import ast
        import inspect

        from direction_non_vie.provisionnement.a7_provisionnement import (
            n5_rapport as _n5r_mod,
        )
        fn = ast.parse(inspect.getsource(_n5r_mod._build_blocks)).body[0]
        params = {a.arg for a in fn.args.args}
        ecrases = sorted({
            t.id for n in ast.walk(fn) if isinstance(n, ast.Assign)
            for t in n.targets if isinstance(t, ast.Name) and t.id in params
        })
        self.assertEqual(ecrases, [],
                         f'parametres de _build_blocks reaffectes : {ecrases}')
        print('    OK AVIS-5 aucun parametre de _build_blocks n est ecrase')


# =============================================================================
#  N2/N3/N4 (releve narration) — LE NETTOYAGE N'EFFACE PLUS DU VRAI
# =============================================================================
#
#  ⚠️ CE DEFAUT N'AFFIRMAIT RIEN DE FAUX -- IL SUPPRIMAIT DU VRAI, en silence,
#  dans le commentaire actuariel remis au CAC. Trois pertes, toutes mesurees
#  AVANT correction :
#    N2 : une ligne de tableau markdown etait effacee AVEC ses chiffres.
#         Mesure : narration 81 -> 25 caracteres, << 7 746 000 EUR >> disparu.
#    N3 : six motifs de faux-titre s'appliquaient a TOUT le texte. Mesure :
#         << Ce rapport sera transmis a l'ACPR avec une reserve de 12 M EUR >>,
#         ecrite en §5, etait supprimee -- le chiffre avec.
#    N4 : le preambule etait TRONQUE en bloc, ouverture legitime comprise.
#
#  ⚠️ LE FILET PORTE SUR `_md_to_html`, PAS SEULEMENT SUR `_nettoyer_narration`.
#  La trace par consommateur l'impose : les deux services Sante-Prevoyance
#  (`sp_rapport_prevoyance`, `sp_rapport_sante`) importent ces fonctions mais
#  n'appellent QUE `_md_to_html` -- ils consomment donc la correction
#  INDIRECTEMENT. Un filet borne a la fonction interne ne prouverait rien pour
#  eux. TROIS livrables, d'ou la gate Vie+SP.

_NL = chr(10)


class T_Nettoyage_Narration_N_Efface_Plus_Du_Vrai(unittest.TestCase):
    """⚠️ CE QUI DOIT PARTIR PART, CE QUI DOIT RESTER RESTE."""

    TABLEAU = ('§1 — CONTEXTE' + _NL * 2 + '| Methode | Reserve |' + _NL
               + '|---|---|' + _NL + '| CL | 7 746 000 EUR |' + _NL * 2 + '§2')
    CORPS = ('§1 — CONTEXTE' + _NL * 2
             + "Ce rapport sera transmis a l ACPR avec une reserve de 12 M EUR."
             + _NL * 2 + '§2')
    OUVERTURE = ('Analyse du portefeuille auto, primes acquises 45 M EUR.'
                 + _NL * 2 + '§1 — CONTEXTE' + _NL + 'suite')
    FAUX_TITRE = ('# RAPPORT ACTUARIEL' + _NL + 'Document destine au Conseil'
                  + _NL + 'Arrete au 30/06/2026' + _NL * 2 + '§1 — CONTEXTE'
                  + _NL + 'corps')

    # ── N2 ───────────────────────────────────────────────────────────────────
    def test_n2_le_chiffre_d_un_tableau_survit(self):
        r = _nettoyer_narration(self.TABLEAU)
        self.assertIn('7 746 000', r, 'le Best Estimate est efface')
        self.assertNotIn('|', r, 'les pipes markdown restent')
        self.assertIn('CL · 7 746 000 EUR', r)
        print('    OK NAR-1 le contenu du tableau est converti, pas efface')

    def test_n2_la_separatrice_seule_disparait(self):
        # ⚠️ ELLE NE PORTE AUCUN CONTENU : la garder afficherait des tirets nus.
        self.assertNotIn('---', _nettoyer_narration(self.TABLEAU))
        print('    OK NAR-2 la separatrice |---|---| part, elle seule')

    # ── N3 ───────────────────────────────────────────────────────────────────
    def test_n3_une_phrase_du_corps_n_est_pas_filtree(self):
        self.assertIn('12 M EUR', _nettoyer_narration(self.CORPS),
                      'un motif de faux-titre mange une phrase du corps')
        print('    OK NAR-3 le corps du rapport n est plus filtre')

    # ── N4 ───────────────────────────────────────────────────────────────────
    def test_n4_une_ouverture_legitime_survit(self):
        self.assertIn('45 M EUR', _nettoyer_narration(self.OUVERTURE),
                      'le preambule est encore tronque en bloc')
        print('    OK NAR-4 une ouverture legitime avant le §1 survit')

    # ── NON-REGRESSION : ce que la fonction doit TOUJOURS retirer ────────────
    def test_le_faux_titre_part_toujours(self):
        r = _nettoyer_narration(self.FAUX_TITRE)
        self.assertNotIn('RAPPORT ACTUARIEL', r)
        self.assertNotIn('Document destine', r)
        self.assertIn('§1', r); self.assertIn('corps', r)
        print('    OK NAR-5 le faux-titre part, le corps reste')

    def test_un_texte_vide_reste_vide(self):
        self.assertEqual(_nettoyer_narration(''), '')
        print('    OK NAR-6 un texte vide reste vide')

    # ── LE CHEMIN DES TROIS LIVRABLES ───────────────────────────────────────
    def test_md_to_html_le_point_d_entree_reel_conserve_tout(self):
        # ⚠️ C'EST CE CHEMIN QUE LES DEUX SERVICES SP CONSOMMENT.
        html = _md_to_html(self.TABLEAU + _NL * 2
                           + "Ce rapport transmis a l ACPR avec 12 M EUR.")
        self.assertIn('7 746 000', html)
        self.assertIn('12 M EUR', html)
        self.assertNotIn('|', html)
        print('    OK NAR-7 _md_to_html conserve tout (chemin A7 + 2 SP)')


# =============================================================================
#  N1 — L'AVIS ET LE JUGEMENT NE SE CONTREDISENT JAMAIS
# =============================================================================
#
#  ⚠️⚠️ LA RAISON PREMIERE DE CE LOT N'EST PAS L'HARMONISATION : c'est que le
#  module IMPOSAIT AU MODELE UNE REGLE QU'IL VIOLAIT LUI-MEME. Le
#  `SYSTEM_PROMPT` d'A7 porte en regle 9 :
#
#      << INTERDIT : [...] FAVORABLE si H1 rejetee ET BT ROUGE. >>
#
#  et la premiere branche de l'avis (`n4_best_estimate`) est le MOT POUR MOT de
#  cet interdit -- elle ecrivait << FAVORABLE SOUS RESERVE >>.
#
#  ⚠️ ET LE DOCUMENT SE CONTREDISAIT A UNE PAGE D'ECART. Sur un ROUGE :
#      section 7 (narration en repli) : << AVIS DEFAVORABLE -- Ne pas inscrire
#                                         au bilan S2 sans validation formelle >>
#      section 8 (conclusion)         : << FAVORABLE SOUS RESERVE >>
#  Le mot rassurant etait dans la section de CONCLUSION.
#
#  ⚠️⚠️ CE VERROU PORTE SUR LA PROPRIETE, PAS SUR LES LIBELLES. Un test qui
#  comparerait trois chaines passerait le jour ou l'un des deux producteurs
#  changerait de vocabulaire sans l'autre -- exactement la divergence qu'on
#  ferme. Ce qui est verrouille : les deux verdicts ne se contredisent jamais.

#: Les trois etats RAG et le mot que l'avis doit porter. Vocabulaire du depot
#: (Sante-Prevoyance le produit deja dans ce meme champ).
_AVIS_PAR_STATUT = {
    'VERT':  'FAVORABLE',
    'AMBRE': 'AVEC RÉSERVES',
    'ROUGE': 'DÉFAVORABLE',
}


class T_Avis_Et_Jugement_Ne_Se_Contredisent_Pas(unittest.TestCase):
    """⚠️ UN DOSSIER QU'ON NE PEUT INSCRIRE A AUCUN BILAN N'EST PAS FAVORABLE,
    meme sous reserve."""

    @classmethod
    def setUpClass(cls):
        with kaleido_declare(True), rendeur_substitue():
            cls.r = AgentA7Provisionnement(verbose=False).run(
                source=np.array(GENINS, dtype=float), mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=200, seed=42,
                generer_graphiques=False)

    def test_un_rouge_ne_dit_plus_favorable(self):
        # ⚠️ LE CAS QUE LA REGLE 9 DU PROMPT NOMME. On force le statut ROUGE et
        # on relit l'avis produit pour CE cas, sans toucher au calcul.
        from direction_non_vie.provisionnement.a7_provisionnement import (
            n4_best_estimate as _n4m,
        )
        src = inspect.getsource(_n4m)
        i = src.find("if statut == 'ROUGE' or (not h1_ok")
        self.assertGreater(i, 0, 'la branche ROUGE de l avis est introuvable')
        branche = src[i:i + 260]
        self.assertIn('DÉFAVORABLE', branche)
        self.assertNotIn('FAVORABLE SOUS RÉSERVE', branche)
        print('    OK N1-1 la branche ROUGE dit DEFAVORABLE')

    def test_les_trois_etats_portent_le_vocabulaire_du_depot(self):
        from direction_non_vie.provisionnement.a7_provisionnement import (
            n4_best_estimate as _n4m,
        )
        src = inspect.getsource(_n4m)
        for statut, mot in _AVIS_PAR_STATUT.items():
            self.assertIn(mot, src, f'{statut} : « {mot} » absent du module')
        print('    OK N1-2 FAVORABLE / AVEC RESERVES / DEFAVORABLE presents')

    def test_le_prompt_interdit_ce_que_le_code_n_ecrit_plus(self):
        # ⚠️ LE VERROU QUI COMPTE : la regle du prompt et le code doivent rester
        # d'accord. Si l'un des deux bouge, ce test le dit.
        from direction_non_vie.provisionnement.a7_provisionnement import (
            n5_rapport as _n5m,
        )
        self.assertIn('FAVORABLE si H1 rejetée ET BT ROUGE',
                      _n5m.SYSTEM_PROMPT,
                      "la regle 9 du prompt a change : revoir l'avis de n4")
        from direction_non_vie.provisionnement.a7_provisionnement import (
            n4_best_estimate as _n4m,
        )
        src = inspect.getsource(_n4m)
        i = src.find("if statut == 'ROUGE' or (not h1_ok")
        self.assertNotIn('FAVORABLE', src[i:i + 260].replace('DÉFAVORABLE', ''),
                         'le code ecrit encore FAVORABLE dans le cas interdit')
        print('    OK N1-3 le code respecte la regle 9 du SYSTEM_PROMPT')

    def test_avis_et_jugement_ne_se_contredisent_jamais(self):
        # ⚠️ LA PROPRIETE, PAS LES MOTS. Les deux champs sont produits par le
        # MEME statut : quand l'un dit << ne pas inscrire >>, l'autre ne peut
        # pas dire << favorable >>. On lit les deux sur le run reel.
        avis = str(self.r['n4'].get('avis_actuariel', ''))
        jug = str(self.r['n4'].get('jugement', ''))
        self.assertTrue(avis and jug, 'avis ou jugement absent du run')
        contredit = ('DÉFAVORABLE' in jug.upper()
                     and avis.upper().startswith('FAVORABLE'))
        self.assertFalse(contredit,
                         f'jugement dit DEFAVORABLE, avis dit : {avis[:60]}')
        print('    OK N1-4 avis et jugement ne se contredisent pas')


# =============================================================================
#  LE JUGEMENT DIT DE QUELLE GRANDEUR IL PARLE, ET NE DIRIGE PLUS VERS LE
#  MAUVAIS NOMBRE
# =============================================================================
#
#  ⚠️ CES TROIS DEFAUTS ONT ETE TROUVES PAR LA MESURE DU VERROU C2, AVANT MEME
#  QUE LE VERROU EXISTE : en comparant les nombres publies a la charge utile,
#  trois provisions sont ressorties orphelines. Aucun releve ne les avait vues.
#
#  P1/P2 -- `_documenter_jugement` publiait les percentiles de MACK SEUL sous
#  l'etiquette generique << Provision P75 / P90 / P99.5 >>, la MEME que le
#  rapport emploie pour les percentiles COMPOSES de N4. Deux grandeurs justes
#  chacune, sous un nom identique, dans le meme document. Ecart mesure :
#      P75   Mack 20 226 075   compose 17 660 196   -12,7 %
#      P90   Mack 21 892 882   compose 21 832 044    -0,3 %
#      P99.5 Mack 25 918 951   compose 34 310 605   +32,4 %  (8,4 M EUR)
#  ⚠️ AUCUNE VALEUR N'A CHANGE : le defaut etait l'ETIQUETTE. La vue Mack
#  native est diagnostique et le HTML la publie sciemment.
#
#  P3 -- LE PLUS GRAVE, ET SUR LE CHEMIN LE MOINS RELU. La section
#  << DECISION ET RECOMMANDATIONS >>, branche VERT SEULE, disait :
#      << Utiliser 21 892 882 EUR pour le calcul du SCR provisions. >>
#  Or `scr_prov = 3.0 * sigma_eiopa * be` -- AUCUN percentile n'y entre. Le
#  SCR reel vaut 4 894 197 EUR : un actuaire qui aurait suivi l'instruction
#  aurait obtenu un SCR 4,5x trop eleve. Et l'instruction n'avait pas d'objet,
#  le SCR etant deja calcule par ce meme module.

class T_Jugement_Nomme_Ses_Grandeurs(unittest.TestCase):
    """⚠️ DEUX GRANDEURS JUSTES SOUS UN MEME NOM RESTENT UN DEFAUT."""

    @classmethod
    def setUpClass(cls):
        with kaleido_declare(True), rendeur_substitue():
            cls.r = AgentA7Provisionnement(verbose=False).run(
                source=np.array(GENINS, dtype=float), mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=200, seed=42,
                generer_graphiques=False)

    def _jugement(self, statut):
        """Le jugement pour un statut donne -- la branche VERT n'est atteignable
        que par cette voie, et c'est justement celle qu'on relit le moins."""
        from direction_non_vie.provisionnement.a7_provisionnement.n4_best_estimate import (
            BestEstimateS2,
        )
        n4 = self.r['n4']
        return BestEstimateS2()._documenter_jugement(
            {}, {}, n4.get('poids', {}), n4['best_estimate'], 2.0,
            self.r['n2'], self.r['n3'], statut, n4['scr'], {'label': 'RC'})

    # ── P1 / P2 ──────────────────────────────────────────────────────────────
    #
    # ⚠️⚠️ CES TROIS TESTS ONT ETE REECRITS PAR LE LOT PERCENTILES, ET LA RAISON
    # DOIT ETRE LUE AVANT DE LES CROIRE. Ils verrouillaient la decision du lot
    # `fcfb3d3` : le jugement publie MACK NATIF, et on le NOMME sans rien
    # deplacer. Cette decision a ete explicitement renversee -- Mack natif est
    # centre sur la reserve de Mack, pas sur le BE publie, donc ses percentiles
    # ne decrivent aucune distribution du Best Estimate que le meme bloc
    # annonce. Le jugement republie desormais `reserve_p*`, la meme grandeur
    # que les quatre livrables.
    #
    # ⚠️ ILS VERROUILLENT DESORMAIS UNE PROPRIETE, PLUS UN LIBELLE : non pas
    # << le jugement dit "Mack natif" >>, mais << le jugement dit ce que
    # `cle_percentiles` designe, et publie les nombres que le rapport publie >>.
    # Un libelle peut etre rearbitre ; la coherence entre deux documents, non.
    def test_les_percentiles_nomment_leur_source(self):
        n4 = self.r['n4']
        attendu = libelle_percentiles(n4)
        j = n4['jugement']
        for cle in ('P75', 'P90', 'P99.5'):
            ligne = next((x for x in j.splitlines()
                          if f'Provision {cle}' in x), None)
            self.assertIsNotNone(ligne, f'{cle} absent du jugement')
            self.assertIn(attendu, ligne,
                          f'{cle} ne dit pas de quelle grandeur il parle')
        # Et ce nom n'est pas une constante de test : il vient de l'arbitrage.
        self.assertIn(n4.get('cle_percentiles'),
                      (CLE_MACK, CLE_BOOT, CLE_COMPOSE))
        print(f'    OK JUG-1 les trois percentiles disent « {attendu} »')

    def test_le_jugement_nomme_sa_source_en_toutes_lettres(self):
        # ⚠️ NOMMER L'APPROCHE NE SUFFIT PAS : le lecteur doit savoir POURQUOI
        # c'est celle-la. La phrase de `source_percentiles` porte la raison --
        # et, quand Mack est retenu, la reserve sur l'hypothese d'independance
        # du compose. Elle doit arriver INTACTE dans le jugement.
        n4 = self.r['n4']
        source = n4.get('source_percentiles') or ''
        self.assertTrue(source.strip(), 'aucune source publiee')
        self.assertIn(source, n4['jugement'],
                      'le jugement ne porte pas la raison de son arbitrage')
        print('    OK JUG-2 le jugement porte la raison de son arbitrage')

    def test_le_jugement_publie_les_memes_nombres_que_le_rapport(self):
        # ⚠️ LE VERROU DU LOT, ET IL REMPLACE L'ANCIEN. On ne verifie plus que
        # << rien n'a bouge >> -- quelque chose a bouge, deliberement. On
        # verifie que le jugement et le rapport ne peuvent plus DIVERGER.
        n4 = self.r['n4']
        for cle, attendu in (('P75',   n4['reserve_p75']),
                             ('P90',   n4['reserve_p90']),
                             ('P99.5', n4['reserve_p99_5'])):
            ligne = next(x for x in n4['jugement'].splitlines()
                         if f'Provision {cle}' in x)
            self.assertIn(f'{attendu:,.0f}', ligne,
                          f'{cle} : le jugement ne publie pas le chiffre du rapport')
        # Et la contre-epreuve : Mack NATIF, l'ancienne source, a bien disparu
        # du bloc -- sinon les deux grandeurs cohabiteraient encore.
        mk = self.r['n3']['mack']
        bloc = [x for x in n4['jugement'].splitlines() if 'Provision P' in x]
        self.assertNotIn(f"{mk.get('reserve_p99_5') or 0:,.0f}", ' '.join(bloc),
                         'le P99.5 de Mack natif est encore publie')
        print('    OK JUG-3 jugement et rapport publient les memes nombres')

    # ── P3 ───────────────────────────────────────────────────────────────────
    def test_le_vert_ne_dirige_plus_vers_un_percentile_pour_le_scr(self):
        j = self._jugement('VERT')
        self.assertNotIn('pour le calcul du SCR provisions', j,
                         "l'instruction fausse est encore la")
        print('    OK JUG-4 l instruction « utiliser X pour le SCR » a disparu')

    def test_le_vert_publie_le_scr_reellement_calcule(self):
        j = self._jugement('VERT')
        scr = self.r['n4']['scr'].get('scr_provisions') or 0
        self.assertIn(f'{scr:,.0f}', j, 'le SCR publie n est pas celui de n4')
        self.assertIn('Art. 115', j)
        # ⚠️ ET SURTOUT : le P90 ne doit PLUS figurer comme cible du SCR.
        p90 = self.r['n3']['mack'].get('reserve_p90') or 0
        ligne = next(x for x in j.splitlines() if 'SCR provisions' in x)
        self.assertNotIn(f'{p90:,.0f}', ligne,
                         'la ligne SCR porte encore le P90')
        print('    OK JUG-5 le VERT publie le SCR reel, pas un percentile')


# =============================================================================
#  LOT PERCENTILES — LA REFERENCE SUIT L'ARBITRAGE, ET LES 4 FORMATS SUIVENT
# =============================================================================
#
#  ⚠️ LE COMPOSE NE MESURAIT PAS L'INCERTITUDE, IL MESURAIT LE DESACCORD ENTRE
#  METHODES. sigma_modele est l'ecart-type des reserves des methodes retenues :
#  79 % de la variance sur GenIns (methodes divergentes), 18 % sur RAA (elles
#  convergent). Et son hypothese d'independance est FAUSSE -- les methodes
#  partagent `pct_dev`, BF recoit `ultimates_cl`. Sous covariance positive la
#  somme quadratique SOUS-ESTIME : le compose n'est pas prudent, il est
#  INDETERMINE.
#
#  ⚠️⚠️ CE FILET VERROUILLE DES PROPRIETES, JAMAIS DES LIBELLES. Un libelle
#  peut etre rearbitre demain -- ce lot en rearbitre trois. Ce qui ne doit
#  jamais changer : quelle grandeur `reserve_p*` porte, que le sigma publie la
#  REGENERE, que la bascule reste conditionnelle a CLM-H3, et que les quatre
#  livrables nomment tous la meme approche.


class T_Percentiles_La_Reference_Suit_L_Arbitrage(unittest.TestCase):
    """⚠️ LA BASCULE EST CONDITIONNELLE, ET C'EST LE POINT LE PLUS IMPORTANT."""

    @classmethod
    def setUpClass(cls):
        cls.C = np.array(GENINS, dtype=float)
        with kaleido_declare(True), rendeur_substitue():
            cls.r = AgentA7Provisionnement(verbose=False).run(
                source=cls.C, mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=200, seed=42,
                generer_graphiques=False)
        cls.n4 = cls.r['n4']

    def _calculer(self, mack_ok, boot_ok):
        """N4 recalcule avec les DEUX portes d'hypothese forcees.

        C'est la seule facon d'atteindre les branches de repli : sur les
        triangles de reference, CLM-H3 passe toujours."""
        from direction_non_vie.provisionnement.a7_provisionnement.n4_best_estimate import (
            BestEstimateS2,
        )
        n2 = self.r['n2']
        n2b = {
            **n2,
            'clm': {**n2.get('clm', {}),
                    'percentiles_mack_publiables': mack_ok},
            'bootstrap_hyp': {**n2.get('bootstrap_hyp', {}),
                              'percentiles_publiables': boot_ok},
        }
        return BestEstimateS2().calculer(n2b, self.r['n3'], self.C)

    @staticmethod
    def _log_normale(be, sig, z):
        """QIS5 TP.5.26, REECRITE ICI. Le filet ne doit pas emprunter la
        fonction qu'il verifie -- sinon une erreur commune passerait deux
        fois inapercue."""
        cv = sig / be
        s2 = float(np.log(1.0 + cv ** 2))
        return float(np.exp(np.log(be) - s2 / 2.0 + z * np.sqrt(s2)))

    # ── PCT-1 ────────────────────────────────────────────────────────────────
    def test_la_reference_est_mack_recentre_quand_clm_h3_valide(self):
        n4 = self.n4
        self.assertTrue(self.r['n2'].get('clm', {})
                        .get('percentiles_mack_publiables', True),
                        'ce triangle ne valide pas CLM-H3 : test sans objet')
        self.assertEqual(n4['cle_percentiles'], CLE_MACK)
        for p in ('p75', 'p90', 'p99_5'):
            self.assertEqual(n4[f'reserve_{p}'], n4[f'reserve_{p}_mack'],
                             f'{p} publie n est pas le Mack recentre')
        print('    OK PCT-1 la reference publiee est Mack recentre')

    # ── PCT-2 ────────────────────────────────────────────────────────────────
    def test_le_sigma_publie_regenere_les_percentiles_publies(self):
        # ⚠️ LA PROPRIETE QUI REND LE CHIFFRE OPPOSABLE : un tiers doit
        # pouvoir refaire le calcul avec les seuls nombres publies.
        n4 = self.n4
        be, sig = n4['best_estimate'], n4['sigma_percentiles']
        for p, z in (('p75', 0.6745), ('p90', 1.2816), ('p99_5', 2.5758)):
            self.assertAlmostEqual(
                self._log_normale(be, sig, z), n4[f'reserve_{p}'], delta=1.0,
                msg=f'{p} n est pas regenerable depuis (BE, sigma_percentiles)')
        print('    OK PCT-2 (BE, sigma_percentiles) regenerent les 3 publies')

    # ── PCT-3 ────────────────────────────────────────────────────────────────
    def test_le_compose_reste_publie_et_differe_de_la_reference(self):
        # Une colonne comparative qui vaut la reference ne compare rien.
        n4 = self.n4
        for p in ('p75', 'p90', 'p99_5'):
            self.assertIsNotNone(n4.get(f'reserve_{p}_compose'),
                                 f'{p} compose retire du livrable')
        self.assertNotEqual(n4['reserve_p99_5'], n4['reserve_p99_5_compose'],
                            'le compose vaut la reference : rien n a bascule')
        # Et le sens mesure : retirer sigma_modele RETRECIT la queue.
        self.assertLess(n4['reserve_p99_5'], n4['reserve_p99_5_compose'])
        print('    OK PCT-3 le compose reste publie, et il differe')

    # ── PCT-4 : LA BASCULE EST CONDITIONNELLE ────────────────────────────────
    def test_clm_h3_rejetee_retire_mack_de_la_reference(self):
        # ⚠️⚠️ LE TEST LE PLUS IMPORTANT DU LOT. Si CLM-H3 est NON VALIDEE,
        # sigma_Mack ne mesure plus l'erreur de prediction : en faire la
        # reference publierait sigma_Mack exactement la ou ce module a etabli
        # qu'il ne vaut rien. La bascule DOIT rendre la main.
        r = self._calculer(mack_ok=False, boot_ok=True)
        self.assertNotEqual(r['cle_percentiles'], CLE_MACK,
                            'Mack reste la reference alors que CLM-H3 le rejette')
        self.assertEqual(r['cle_percentiles'], CLE_BOOT)
        for p in ('p75', 'p90', 'p99_5'):
            self.assertEqual(r[f'reserve_{p}'], r[f'reserve_{p}_boot'],
                             f'{p} ne suit pas le relais Bootstrap')
        print('    OK PCT-4 CLM-H3 rejetee : la reference passe au Bootstrap')

    def test_sans_relais_bootstrap_le_repli_est_le_compose_signale(self):
        r = self._calculer(mack_ok=False, boot_ok=False)
        self.assertEqual(r['cle_percentiles'], CLE_COMPOSE)
        for p in ('p75', 'p90', 'p99_5'):
            self.assertEqual(r[f'reserve_{p}'], r[f'reserve_{p}_compose'])
        # Et le repli se DIT : une case juste mais contestee doit le declarer.
        self.assertIn('CONTESTÉE', r['source_percentiles'])
        print('    OK PCT-5 sans relais : compose, et le livrable le signale')

    # ── PCT-6 : l'hypothese d'independance est declaree NON VERIFIEE ─────────
    def test_l_hypothese_d_independance_du_compose_est_declaree(self):
        # ⚠️ LA RESERVE DE FOND VIT DANS LE LIVRABLE, PAS SEULEMENT EN
        # COMMENTAIRE : le compose suppose sigma_Mack et sigma_modele
        # independants, et rien ne le verifie.
        src = self.n4['source_percentiles']
        self.assertIn("indépendance", src)
        self.assertIn("PAS vérifiée", src)
        print('    OK PCT-6 l hypothese d independance est declaree non verifiee')

    # ── PCT-7 : une seule approche marquee « retenue », la bonne ─────────────
    def test_une_seule_approche_est_marquee_retenue(self):
        n4 = self.n4
        marques = [c for c in (CLE_MACK, CLE_BOOT, CLE_COMPOSE)
                   if marque_retenue(n4, c, 'X') != 'X']
        self.assertEqual(marques, [n4['cle_percentiles']],
                         'la marque « retenue » ne suit pas l arbitrage')
        # Mack natif ne peut JAMAIS l'etre : il est centre ailleurs.
        self.assertEqual(marque_retenue(n4, '', 'Mack natif'), 'Mack natif')
        print('    OK PCT-7 une seule approche marquee, celle qui est publiee')

    # ── PCT-8 : les quatre livrables nomment la MEME approche ────────────────
    def test_les_quatre_livrables_nomment_la_meme_approche(self):
        # ⚠️ LE DEFAUT QUE LE LOT FERME : chaque format ecrivait « composé » en
        # dur. Le nom vient desormais d'une source unique -- on verifie qu'il
        # ARRIVE dans les trois formats textuels, et qu'aucun ne contredit.
        n4, n3, n2, n1 = (self.r['n4'], self.r['n3'],
                          self.r['n2'], self.r['n1'])
        appr = libelle_percentiles(n4)
        self.assertEqual(appr, LIBELLE_APPROCHE[CLE_MACK])
        html = export_html(n1, n2, n3, n4, {})
        self.assertIn(appr, html, 'le HTML ne nomme pas l approche publiee')
        # ⚠️ CONTRE-EPREUVE : plus aucune etiquette « (composé) » collee a un
        # percentile publie. On cherche la FORME EXACTE qui etait en dur.
        for forme in ('P99,5 (composé)', 'P90 (composé)', 'P75 (composé)'):
            self.assertNotIn(forme, html,
                             f'etiquette figee encore presente : {forme}')
        comm = n4.get('jugement') or ''
        self.assertIn(appr, comm)
        print(f'    OK PCT-8 les livrables nomment tous « {appr} »')

    # ── PCT-9 : le chemin LLT ne ressuscite pas le compose ───────────────────
    def test_le_recentrage_llt_repart_du_sigma_de_la_bascule(self):
        # ⚠️ LE SECOND PRODUCTEUR. `agent.py` recalcule `reserve_p*` apres le
        # LLT ; il lisait `sigma_total_compose` EN DUR. Le dossier divergeait
        # donc selon qu'un grand sinistre etait detecte ou non. On verifie sur
        # le CODE que la cle de la bascule est bien celle qui est lue -- un
        # test de valeur ne le verrait pas tant qu'aucun LLT ne se declenche.
        import direction_non_vie.provisionnement.a7_provisionnement.agent as ag
        src = inspect.getsource(ag)
        # On isole L'EXPRESSION elle-meme, pas une fenetre de caracteres : un
        # commentaire un peu plus long deplacerait une fenetre, jamais ceci.
        i = src.index('_sig_pct = float(')
        expr = src[i:src.index('0.0)', i) + 4]
        self.assertIn("n4.get('sigma_percentiles')", expr,
                      'le recentrage LLT ne lit pas le sigma de la bascule')
        self.assertLess(expr.index("sigma_percentiles"),
                        expr.index("sigma_total_compose"),
                        'le compose est lu AVANT le sigma de la bascule')
        print('    OK PCT-9 le recentrage LLT repart du sigma de la bascule')

    # ── PCT-10 : rien d'autre n'a bouge ──────────────────────────────────────
    def test_ni_le_be_ni_le_scr_ni_la_rm_ne_dependent_des_percentiles(self):
        # ⚠️ MESURE, PAS RAISONNEMENT : on rejoue les trois arbitrages et on
        # verifie que les agregats du bilan sont RIGOUREUSEMENT identiques.
        ref = self.n4
        for mack_ok, boot_ok in ((False, True), (False, False)):
            r = self._calculer(mack_ok, boot_ok)
            for cle in ('best_estimate', 'risk_margin',
                        'provisions_techniques_s2', 'cv_inter_methodes'):
                self.assertEqual(r[cle], ref[cle],
                                 f'{cle} depend de l arbitrage des percentiles')
            self.assertEqual(r['scr']['scr_provisions'],
                             ref['scr']['scr_provisions'])
        print('    OK PCT-10 BE / SCR / RM / PT invariants aux 3 arbitrages')


# =============================================================================
#  LOT « CORRECTIONS PARTIELLES » — UN SITE RESTANT, TROIS FOIS
# =============================================================================
#
#  ⚠️⚠️ LES TROIS DEFAUTS DE CE LOT AVAIENT DEJA ETE CORRIGES AILLEURS.
#  Le facteur 3 : juste dans la docstring de N4, juste dans le commentaire,
#  juste dans une fixture -- FAUX dans l'Excel. Le faux zero : ferme aux
#  onglets 1 et 4, ouvert au 5. Le motif d'exclusion : juste sur quatre
#  sites, impossible sur le cinquieme.
#
#  ⚠️ UNE CORRECTION PARTIELLE LAISSE UNE TRACE QUI ATTESTE QU'ELLE EST
#  COMPLETE. C'est ce qui la rend plus dangereuse qu'une faute jamais vue :
#  le releve suivant lit la trace et passe. Le remede n'est donc PAS la
#  correction du site restant -- c'est la SOURCE UNIQUE, qui rend le
#  sixieme site impossible.


class T_Corrections_Partielles_Une_Seule_Source(unittest.TestCase):
    """⚠️ LES SITES NE REDIGENT PLUS, DONC ILS NE PEUVENT PLUS DIVERGER."""

    @classmethod
    def setUpClass(cls):
        cls.C = np.array(GENINS, dtype=float)
        with kaleido_declare(True), rendeur_substitue():
            # SANS exposition : BF et Cape Cod ne sont PAS calculables. C'est
            # la seule configuration ou les trois defauts se manifestent.
            cls.r = AgentA7Provisionnement(verbose=False).run(
                source=cls.C, mode_declare='cumule', n_sim_bootstrap=60,
                seed=42, generer_graphiques=False)
        cls.n1, cls.n2 = cls.r['n1'], cls.r['n2']
        cls.n3, cls.n4 = cls.r['n3'], cls.r['n4']
        cls.xl = export_excel(cls.C, cls.n1, cls.n2, cls.n3, cls.n4)

    def _onglet(self, titre):
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(self.xl))
        return [[c.value for c in row] for row in wb[titre].iter_rows()]

    # ── CP-1 : le facteur 3 ─────────────────────────────────────────────────
    def test_le_facteur_3_n_est_plus_dit_quantile_d_une_normale(self):
        from direction_non_vie.provisionnement.a7_provisionnement.n4_best_estimate import (  # noqa: E501
            MSG_FACTEUR_3,
        )
        # ⚠️ LA PROPRIETE, PAS LE LIBELLE : le quantile 99,5 % d'une loi
        # normale vaut 2,5758 -- 16,5 % d'ecart avec le facteur 3. Aucun
        # livrable ne doit plus attribuer ce facteur a une loi normale.
        plat = ' '.join(str(c) for lg in self._onglet('9. SCR formule standard')
                        for c in lg if c)
        self.assertNotIn('Quantile 99.5% loi normale', plat)
        self.assertIn('2,576', plat, "l'Excel ne nomme pas le vrai quantile")
        # Et la source est UNE : le commentaire ne redige plus la sienne.
        self.assertIn(MSG_FACTEUR_3,
                      generer_commentaire(self.n1, self.n2, self.n3, self.n4))
        src = inspect.getsource(COMM_MOD)
        self.assertNotIn('Ce n\'est pas le quantile d\'une loi normale', src,
                         'le commentaire redige de nouveau sa propre version')
        print('    OK CP-1 le facteur 3 vient d une source unique')

    # ── CP-2 : le faux zero de l'onglet 5 ───────────────────────────────────
    def test_l_onglet_ibnr_ne_publie_pas_zero_pour_une_methode_absente(self):
        self.assertFalse(disponible(self.n3, 'bornhuetter_ferguson'),
                         'ce run calcule BF : le test serait sans objet')
        lignes = self._onglet('5. IBNR par année')
        total = next(lg for lg in lignes if lg and lg[0] == 'TOTAL')
        self.assertEqual([total[2], total[3]], ['—', '—'],
                         'le TOTAL republie un zero pour une methode absente')
        corps = [lg for lg in lignes
                 if lg and isinstance(lg[0], str) and lg[0].startswith('An. ')]
        self.assertTrue(corps, 'aucune ligne de detail')
        for lg in corps:
            self.assertEqual([lg[2], lg[3]], ['—', '—'], str(lg[:4]))
        # ⚠️ CONTRE-EPREUVE : quand la methode EST calculee, le montant sort.
        # Un garde qui masque tout serait aussi faux que le zero.
        with kaleido_declare(True), rendeur_substitue():
            r2 = AgentA7Provisionnement(verbose=False).run(
                source=self.C, mode_declare='cumule', seed=42,
                primes=_exposition(GENINS), n_sim_bootstrap=60,
                generer_graphiques=False)
        xl2 = export_excel(self.C, r2['n1'], r2['n2'], r2['n3'], r2['n4'])
        import openpyxl
        wb2 = openpyxl.load_workbook(io.BytesIO(xl2))
        t2 = next(([c.value for c in row]
                   for row in wb2['5. IBNR par année'].iter_rows()
                   if row[0].value == 'TOTAL'))
        self.assertIsInstance(t2[2], float)
        self.assertGreater(t2[2], 0)
        print('    OK CP-2 IBNR : tiret si non calculee, montant sinon')

    # ── CP-3 : le motif d'exclusion atteint TOUS les formats ────────────────
    def test_le_motif_d_exclusion_atteint_le_html_comme_l_excel(self):
        # ⚠️ LE HTML CHERCHAIT PAR LIBELLE DANS UN DICT INDEXE PAR CLE :
        # intersection vide, mesuré — le repli était pris à TOUS les coups.
        #
        # ⚠️⚠️ ET LA MESURE A REFUTE MA PROPRE DESCRIPTION DU DEFAUT. J'avais
        # ecrit que << l'Excel publie le motif detaille, le HTML jamais >>.
        # FAUX : `methodes_exclues_motifs` porte litteralement
        # << non calculée >> -- LE MEME TEXTE QUE LE REPLI. Le bug du lookup
        # etait donc DORMANT : les deux formats disaient deja la meme chose,
        # par hasard. Il cessera de l'etre le jour ou le dict portera un vrai
        # motif -- c'est-a-dire au premier lot qui remonte
        # `n3[cle]['message']`, aujourd'hui publie nulle part (SIGNALE, non
        # traite ici).
        #
        # LA PROPRIETE VERROUILLEE EST DONC CELLE QUI TIENDRA ENCORE APRES :
        # ce que publie un format ne depend plus de la FORME de la cle, et
        # les deux formats publient la MEME chose.
        motif = motif_exclusion(self.n4, 'bornhuetter_ferguson')
        self.assertIn('bornhuetter_ferguson', self.n4['methodes_exclues'])
        html = export_html(self.n1, self.n2, self.n3, self.n4, {})
        plat = ' '.join(str(c) for lg in self._onglet('1. Synthèse')
                        for c in lg if c)
        for ou, txt in (('HTML', html), ('Excel', plat)):
            self.assertIn(motif, txt, f'{ou} : motif absent')
        # La contre-epreuve du bug ferme : chercher par LIBELLE ne rend plus
        # rien, chercher par CLE rend le motif. Le code ne fait plus le premier.
        self.assertEqual(
            motif_exclusion(self.n4, libelle('bornhuetter_ferguson')),
            'non calculée',
            'le libelle serait devenu une cle valide : le test perd son sens')
        print('    OK CP-3 le motif ne depend plus de la forme de la cle')

    # ── CP-4 : la source est unique, et personne ne la contourne ────────────
    def test_plus_aucun_site_ne_lit_le_dict_des_motifs_en_direct(self):
        # ⚠️ LA PROPRIETE STRUCTURELLE QUI EMPECHE LE SIXIEME SITE. Un `.get`
        # ecrit a la main peut se tromper de cle ; `motif_exclusion` non.
        for mod in (RAPP_MOD, COMM_MOD, XL_MOD):
            src = inspect.getsource(mod)
            self.assertNotIn("methodes_exclues_motifs", src,
                             f'{mod.__name__} lit le dict en direct')
        print('    OK CP-4 un seul acces au dict des motifs')

    # ── CP-5 : aucun euro deplace ───────────────────────────────────────────
    def test_aucun_agregat_n_a_bouge(self):
        # ⚠️ MESURE, PAS RAISONNEMENT : ce lot ne touche que des ETIQUETTES et
        # des GARDES d'affichage. Les agregats du bilan doivent etre inchanges.
        self.assertEqual(self.n4['best_estimate'], 18680856.0)
        self.assertEqual(self.n4['scr']['scr_provisions'], 6164682.0)
        self.assertEqual(self.n4['reserve_p90'], 21892882.0)
        print('    OK CP-5 BE / SCR / P90 inchanges')


# =============================================================================
#  COMPARATIF N-1/N — UNE COULEUR NE SIGNALE PLUS UN ARTEFACT DE METHODE
# =============================================================================
#
#  ⚠️⚠️ REGRESSION INTRODUITE PAR LE LOT PERCENTILES (`6b630d2`), ET C'EST
#  MOI QUI L'AI FAITE. `reserve_p90` a change de NATURE -- compose avant,
#  sigma_Mack recentre depuis. La valeur N-1 est SAISIE A LA MAIN par
#  l'actuaire : rien ne dit de quelle approche elle vient.
#
#  Un ecart de -17,3 % s'affichait donc en ROUGE (seuil 15 %) : le rapport
#  signalait une DERIVE DE PROVISIONNEMENT la ou rien n'avait bouge. C'est
#  plus grave que l'imprecision -- une couleur se lit avant un chiffre.


class T_Comparatif_N1_La_Couleur_Suit_La_Comparabilite(unittest.TestCase):
    """⚠️ LA COULEUR SUIT LA COMPARABILITE, PAS LA LIGNE."""

    #: Un N-1 dont le P90 est le COMPOSE -- exactement le cas qui produit
    #: l'artefact. Les autres grandeurs sont proches, pour qu'un ecart sur
    #: elles reste dans les seuils VERT/AMBRE et prouve que leur couleur vit.
    _PREC = {'best_estimate': 14000000.0, 'reserve_p90': 21832044.0,
             'cv_inter_methodes': 38.0, 'sigma_mack': 2400000.0}

    @classmethod
    def setUpClass(cls):
        cls.C = np.array(GENINS, dtype=float)
        with kaleido_declare(True), rendeur_substitue():
            cls.r = AgentA7Provisionnement(verbose=False).run(
                source=cls.C, mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
                generer_graphiques=False)
        cls.xl = export_excel(cls.C, cls.r['n1'], cls.r['n2'], cls.r['n3'],
                              cls.r['n4'], resultats_precedents=cls._PREC)

    def _lignes(self):
        import openpyxl
        ws = openpyxl.load_workbook(io.BytesIO(self.xl))['8. Comparatif N-1 N']
        return [row for row in ws.iter_rows(min_row=3, max_row=6)
                if row[0].value]

    def test_la_ligne_p90_n_est_plus_coloree(self):
        # ⚠️ LA PROPRIETE : l'ecart EXISTE et reste publie -- on ne cache pas
        # l'information, on retire la LECTURE fausse qu'une couleur impose.
        p90 = next(r for r in self._lignes() if 'P90' in str(r[0].value))
        self.assertLess(p90[4].value, -0.15,
                        "l ecart n atteint plus le seuil ROUGE : test sans objet")
        self.assertNotIn(p90[4].font.color.rgb[-6:], ('C0392B', 'B87A00'),
                         'un artefact de methode est encore colore')
        print('    OK CN-1 la ligne P90 publie son ecart, sans couleur')

    def test_les_autres_lignes_gardent_leur_couleur(self):
        # ⚠️ CONTRE-EPREUVE INDISPENSABLE : un correctif qui eteint TOUT le
        # tableau serait aussi faux que la couleur de trop. BE, sigma et CV
        # n'ont pas change de nature, leur variation reste interpretable.
        vivantes = [r for r in self._lignes() if 'P90' not in str(r[0].value)]
        self.assertEqual(len(vivantes), 3)
        for r in vivantes:
            self.assertIn(r[4].font.color.rgb[-6:], ('1D7A3A', 'B87A00', 'C0392B'),
                          f'{r[0].value} a perdu sa couleur')
        print('    OK CN-2 BE, sigma et CV gardent leur couleur')

    def test_la_raison_est_publiee_dans_les_deux_formats(self):
        # ⚠️ UNE COULEUR RETIREE SANS EXPLICATION SE LIT COMME UN OUBLI.
        from direction_non_vie.provisionnement.a7_provisionnement.n4_best_estimate import (  # noqa: E501
            MSG_P90_NON_COMPARABLE,
        )
        import openpyxl
        ws = openpyxl.load_workbook(io.BytesIO(self.xl))['8. Comparatif N-1 N']
        plat = ' '.join(str(c.value) for row in ws.iter_rows()
                        for c in row if c.value)
        self.assertIn(MSG_P90_NON_COMPARABLE, plat, 'Excel : reserve absente')
        com = generer_commentaire(self.r['n1'], self.r['n2'], self.r['n3'],
                                  self.r['n4'],
                                  resultats_precedents=self._PREC)
        self.assertIn(MSG_P90_NON_COMPARABLE, com, 'commentaire : absente')
        print('    OK CN-3 la raison est publiee dans les deux formats')

    def test_la_raison_vit_a_l_endroit_ou_la_couleur_sautait(self):
        # ⚠️ SANS CELA, QUELQU'UN LA RETABLIRA DANS SIX MOIS en croyant
        # combler un oubli. Le lot avis-couleur a deja paye exactement ca.
        src = inspect.getsource(XL_MOD)
        i = src.index('# Couleur variation')
        bloc = src[i:i + 1200]
        self.assertIn('CHANGÉ DE NATURE', bloc)
        self.assertIn('comparable', bloc)
        print('    OK CN-4 la raison est ecrite la ou la couleur sautait')

    def test_aucun_agregat_n_a_bouge(self):
        n4 = self.r['n4']
        self.assertEqual(n4['best_estimate'], 14830899.0)
        self.assertEqual(n4['scr']['scr_provisions'], 4894197.0)
        self.assertEqual(n4['reserve_p90'], 18053284.0)
        print('    OK CN-5 BE / SCR / P90 inchanges')


# =============================================================================
#  UNE SEULE NARRATION POUR LES DEUX DOCUMENTS SIGNES
# =============================================================================
#
#  ⚠️⚠️ CHAQUE EXPORT APPELAIT `_generer_narration` DE SON COTE. Mesure sur un
#  run reel : DEUX appels. Tant que le texte deterministe passait devant, les
#  deux etaient identiques par construction -- 12 699 caracteres, mesures a
#  l'identique -- et le doublon ne coutait rien.
#
#  ⚠️ C'EST LA DECISION DE FAIRE PASSER LA NARRATION LLM DEVANT QUI REND LE
#  DEFAUT ACTIF : deux appels LLM independants produisent deux textes
#  differents. Le HTML et le Word du MEME dossier, tous deux transmis au
#  commissaire aux comptes, ne diraient pas la meme chose -- et l'appel serait
#  paye deux fois. Un doublon inerte devient une divergence entre documents
#  signes.


class T_Une_Seule_Narration_Pour_Les_Deux_Formats(unittest.TestCase):
    """⚠️ DEUX DOCUMENTS SIGNES NE PEUVENT PAS PORTER DEUX TEXTES."""

    def _run_espionne(self, **kw):
        """Un run complet, en comptant les appels a la narration."""
        import direction_non_vie.provisionnement.a7_provisionnement.agent as AG
        appels = []
        vrai = RAPP_MOD._generer_narration

        # ⚠️ TROIS VALEURS, PAS DEUX. `_generer_narration` rend desormais la
        # CHARGE UTILE en plus du texte et de sa source, pour que le verrou C2
        # n'ait pas a la reconstruire. Cet espion en deballait deux : il levait
        # AVANT de compter, le compteur restait a zero, et le test annoncait
        # << 0 appels : les deux formats divergeraient >> -- un diagnostic
        # faux sur un code juste. La traversee de couches, payee une fois de
        # plus ; c'est la gate qui l'a vue, pas la relecture.
        def espion(n2, n3, n4, com, lob, arr):
            txt, src, charge = vrai(n2, n3, n4, com, lob, arr)
            appels.append((src, txt))
            return txt, src, charge

        RAPP_MOD._generer_narration = espion
        AG._generer_narration = espion
        try:
            with kaleido_declare(True), rendeur_substitue():
                r = AG.AgentA7Provisionnement(verbose=False).run(
                    source=np.array(GENINS, dtype=float),
                    mode_declare='cumule', primes=_exposition(GENINS),
                    n_sim_bootstrap=60, seed=42, **kw)
        finally:
            RAPP_MOD._generer_narration = vrai
            AG._generer_narration = vrai
        return r, appels

    def test_un_run_ne_genere_la_narration_qu_une_fois(self):
        # ⚠️ LA PROPRIETE, ET ELLE TIENT MEME SANS CLE API : c'est le NOMBRE
        # d'appels qui garantit l'identite des deux documents, pas le fait que
        # le repli deterministe rende deux fois le meme texte.
        r, appels = self._run_espionne(generer_graphiques=True,
                                       generer_word=True)
        self.assertEqual(len(appels), 1,
                         f'{len(appels)} appels : les deux formats divergeraient')
        self.assertTrue(r.get('html'), 'le HTML n a pas ete produit')
        self.assertTrue(r.get('word_bytes'), 'le Word n a pas ete produit')
        print('    OK NAR-1 un seul appel, deux documents produits')

    def test_la_narration_recue_est_celle_qui_est_publiee(self):
        # ⚠️ UN PARAMETRE RECU ET IGNORE SERAIT PIRE QUE LE DOUBLON : les deux
        # formats regenereraient en silence. On PLANTE un texte reconnaissable
        # et on verifie qu'il ressort des DEUX cotes.
        r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False)
        temoin = 'TEMOIN-NARRATION-PARTAGEE-4718'
        html = export_html(r['n1'], r['n2'], r['n3'], r['n4'], {},
                           narration=temoin, source_narration='templates')
        mot = export_word(r['n1'], r['n2'], r['n3'], r['n4'], {},
                          narration=temoin, source_narration='templates')
        self.assertIn(temoin, html, 'le HTML a ignore la narration recue')
        import docx
        txt = '\n'.join(p.text for p in docx.Document(io.BytesIO(mot)).paragraphs)
        self.assertIn(temoin, txt, 'le Word a ignore la narration recue')
        print('    OK NAR-2 la narration recue ressort des deux cotes')

    def test_sans_narration_fournie_le_comportement_est_inchange(self):
        # ⚠️ CONTRE-EPREUVE INDISPENSABLE : une trentaine d'appels de tests et
        # l'application n'ont PAS ete modifies. Le repli doit rester intact.
        r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False)
        html = export_html(r['n1'], r['n2'], r['n3'], r['n4'], {})
        self.assertGreater(len(html), _TAILLE_MIN_LIVRABLE)
        self.assertNotIn(MARQUEUR_ECHEC_RAPPORT, html)
        print('    OK NAR-3 sans narration fournie, rien ne change')


# =============================================================================
#  LE VERROU C2 EST BRANCHE — ET IL MORD, PROUVE PAR UNE REPONSE SIMULEE
# =============================================================================
#
#  ⚠️⚠️ SANS CLE API, LE VERROU NE S'EXERCE JAMAIS EN CONDITIONS REELLES.
#  C'est la contrepartie assumee du branchement : il REND POSSIBLE la mesure
#  du taux sur du LLM, il ne la PRODUIT pas. Un mecanisme jamais exerce est
#  un mecanisme qui atteste surveiller sans avoir jamais surveille -- le
#  defaut exact que ce chantier vient de nommer sur `vocabulaire_scr_fautif`.
#
#  D'OU LA REPONSE SIMULEE : on substitue la frontiere LLM pour lui faire
#  rendre un texte porteur d'un nombre ABSENT de la charge utile. C'est la
#  SEULE preuve que le verrou mord, et elle est indispensable, pas optionnelle.


class T_Verrou_C2_Branche_Sur_L_Audit(unittest.TestCase):
    """⚠️ UN JOURNAL QUE PERSONNE NE CONSULTE N'EST PAS UN CONTROLE."""

    @classmethod
    def setUpClass(cls):
        cls.C = np.array(GENINS, dtype=float)

    def _run(self, **kw):
        import tempfile
        d = tempfile.mkdtemp()
        r = AgentA7Provisionnement(models_path=d, audit_path=d,
                                   verbose=False).run(
            source=self.C, mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False, **kw)
        return r, d

    def _audit_sur_disque(self, dossier):
        import glob
        import json
        fichiers = glob.glob(os.path.join(dossier, '*.json'))
        self.assertTrue(fichiers, 'aucun audit ecrit sur disque')
        with open(fichiers[0], encoding='utf-8') as f:
            return json.load(f)

    # ── C2B-1 : le controle atteint le fichier archive ──────────────────────
    def test_le_controle_est_ecrit_dans_l_audit_sur_disque(self):
        # ⚠️ SUR DISQUE, pas dans le dict retourne : c'est le FICHIER qui est
        # archive pour l'ACPR, et c'est lui que quelqu'un relira.
        _, d = self._run()
        audit = self._audit_sur_disque(d)
        self.assertIn('controle_narration', audit)
        self.assertIn('porte', audit['controle_narration'])
        print('    OK C2B-1 le controle est dans l audit archive')

    # ── C2B-2 : hors LLM, il rend sa RAISON, jamais un zero ─────────────────
    def test_hors_chemin_llm_il_dit_pourquoi_il_ne_s_applique_pas(self):
        # ⚠️ UN ZERO SE LIRAIT << controle, rien trouve >>. Sans transmission
        # il n'y a pas de charge utile : le controle n'a pas d'objet, et le
        # dire est la seule sortie honnete.
        _, d = self._run()
        c = self._audit_sur_disque(d)['controle_narration']
        self.assertFalse(c['applicable'])
        self.assertNotIn('taux_orphelins', c)
        self.assertIn('charge utile', c['raison'])
        print('    OK C2B-2 hors LLM : la raison, jamais un zero')

    # ── C2B-3 : LA PREUVE QU'IL MORD ────────────────────────────────────────
    def test_un_nombre_invente_par_le_modele_est_signale(self):
        # ⚠️ LA SEULE PREUVE POSSIBLE AUJOURD'HUI. On substitue la frontiere
        # LLM : elle rend un texte qui porte 999 111 777 EUR -- un montant
        # qu'AUCUNE charge utile ne peut contenir.
        invente = '999 111 777'
        faux = (f"Le Best Estimate retenu s'eleve a {invente} EUR pour "
                f"l'exercice, en hausse de 4,2 % sur l'arrete precedent.")

        # ⚠️ `type = 'text'` EST INDISPENSABLE : `texte_des_blocs` ne
        # concatene que les blocs de ce type et LEVE sur une chaine vide. Un
        # bloc sans type ferait retomber le run sur `templates` EN SILENCE,
        # et le test aurait valide un chemin jamais pris -- mesure : c'est
        # exactement ce qui s'est produit a la premiere ecriture.
        class _Bloc:
            type = 'text'
            text = faux

        class _Resp:
            content = [_Bloc()]

        vrai_appeler = RAPP_MOD.frontiere_llm.appeler
        vraie_cle = RAPP_MOD.frontiere_llm.cle_api_ou_secrets
        RAPP_MOD.frontiere_llm.appeler = lambda **kw: _Resp()
        RAPP_MOD.frontiere_llm.cle_api_ou_secrets = lambda *a, **k: 'cle-test'
        try:
            _, d = self._run()
            c = self._audit_sur_disque(d)['controle_narration']
        finally:
            RAPP_MOD.frontiere_llm.appeler = vrai_appeler
            RAPP_MOD.frontiere_llm.cle_api_ou_secrets = vraie_cle

        self.assertTrue(c['applicable'], 'le chemin LLM n a pas ete pris')
        self.assertEqual(c['source'], 'claude_api')
        self.assertGreaterEqual(c['n_orphelins'], 1,
                                'le nombre invente n a PAS ete signale')
        plat = ' '.join(str(o) for o in c['orphelins'])
        self.assertIn('999', plat, f'orphelins releves : {c["orphelins"]}')
        print('    OK C2B-3 un nombre invente est signale — le verrou mord')

    # ── C2B-4 : sa reserve vit dans le fichier, pas seulement en commentaire ─
    def test_la_reserve_de_fond_est_ecrite_dans_l_audit(self):
        # ⚠️ IL PROUVE LA PROVENANCE, JAMAIS LA JUSTESSE. Un relecteur qui
        # ouvre l'audit doit le lire LA, pas dans le code source.
        _, d = self._run()
        c = self._audit_sur_disque(d)['controle_narration']
        self.assertIn('provenance', c['porte'])
        self.assertIn('justesse', c['porte'])
        print('    OK C2B-4 la reserve est ecrite dans le fichier archive')

    # ── C2B-5 : il ne leve jamais ───────────────────────────────────────────
    def test_le_verrou_ne_fait_jamais_tomber_le_run(self):
        # ⚠️ DECISION ARBITREE : il JOURNALISE, il ne LEVE pas. Meme avec une
        # narration integralement inventee, le run rend ses livrables.
        class _Bloc:
            type = 'text'
            text = 'Provision de 123 456 789 EUR, ratio 42,42 %, sur 77 ans.'

        class _Resp:
            content = [_Bloc()]

        vrai_appeler = RAPP_MOD.frontiere_llm.appeler
        vraie_cle = RAPP_MOD.frontiere_llm.cle_api_ou_secrets
        RAPP_MOD.frontiere_llm.appeler = lambda **kw: _Resp()
        RAPP_MOD.frontiere_llm.cle_api_ou_secrets = lambda *a, **k: 'cle-test'
        try:
            r, d = self._run()
        finally:
            RAPP_MOD.frontiere_llm.appeler = vrai_appeler
            RAPP_MOD.frontiere_llm.cle_api_ou_secrets = vraie_cle
        self.assertTrue(r.get('success'), 'le run est tombe sur le verrou')
        c = self._audit_sur_disque(d)['controle_narration']
        self.assertGreaterEqual(c['n_orphelins'], 1)
        print('    OK C2B-5 il journalise, il ne leve pas')


# =============================================================================
#  C3 — LE DOSSIER CONSERVE ET VERIFIABLE
# =============================================================================
#
#  ⚠️⚠️ MESURE QUI A OUVERT CE BLOC : un run produit 344 118 octets -- 1,07 Mo
#  avec les figures reelles -- et en ecrivait 7 169, l'audit trail SEUL.
#  97,9 % de ce qui est produit disparaissait a la fin du run. Les trois
#  documents ne survivaient que si l'actuaire cliquait sur un bouton de
#  telechargement : la conservation du dossier signe reposait sur un clic.
#
#  ⚠️ ET L'EMPREINTE EST LE CŒUR, PAS L'ECRITURE. Un fichier ecrit sans
#  empreinte ne prouve rien : rien n'interdit de le remplacer. C'est pourquoi
#  l'ALTERATION PLANTEE est la condition d'acceptation du lot -- modifier UN
#  octet doit faire echouer la verification.
#
#  ⚠️ CE QUE CE LOT NE FAIT PAS, et le module le dit aux trois endroits ou on
#  pourrait croire le contraire : un dossier CONSERVE ET VERIFIABLE n'est pas
#  un dossier OPPOSABLE au sens juridique. La signature reste hors du systeme.


class T_C3_Le_Dossier_Est_Conserve_Et_Verifiable(unittest.TestCase):
    """⚠️ SANS EMPREINTE, UN FICHIER ECRIT NE PROUVE RIEN."""

    @classmethod
    def setUpClass(cls):
        import tempfile
        cls.dir = tempfile.mkdtemp()
        with kaleido_declare(True), rendeur_substitue():
            cls.r = AgentA7Provisionnement(
                models_path=cls.dir, audit_path=cls.dir, verbose=False).run(
                source=np.array(GENINS, dtype=float), mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
                generer_graphiques=True, generer_word=True, archiver=True)
        cls.audit = cls._audit(cls.dir)
        cls.archive = cls.audit['archive']

    @staticmethod
    def _audit(dossier):
        import glob
        import json
        f = [x for x in glob.glob(os.path.join(dossier, '*.json'))]
        with open(f[0], encoding='utf-8') as fh:
            return json.load(fh)

    # ── P1 : le document EXISTE apres le run ────────────────────────────────
    def test_p1_les_livrables_sont_ecrits_sur_disque(self):
        fichiers = self.archive['fichiers']
        self.assertEqual(sorted(fichiers), ['commentaire.txt', 'donnees.xlsx',
                                            'rapport.docx', 'rapport.html'])
        for nom, meta in fichiers.items():
            chemin = os.path.join(self.archive['dossier'], nom)
            self.assertTrue(os.path.exists(chemin), nom)
            self.assertEqual(os.path.getsize(chemin), meta['octets'], nom)
        total = sum(m['octets'] for m in fichiers.values())
        self.assertGreater(total, 300_000, 'le dossier est anormalement leger')
        print(f'    OK C3-1 {len(fichiers)} livrables ecrits, {total:,} octets')

    # ── P2 : il est RETROUVABLE par son audit_id ────────────────────────────
    def test_p2_le_dossier_porte_l_audit_id(self):
        self.assertIn(self.audit['audit_id'],
                      os.path.basename(self.archive['dossier']))
        print('    OK C3-2 le dossier porte l audit_id')

    # ── P3 : LE CŒUR — il est INALTERE, et l'alteration se voit ─────────────
    def test_p3_l_alteration_d_un_seul_octet_est_detectee(self):
        # ⚠️ LA CONDITION D'ACCEPTATION DU LOT. Un archivage qui ne detecte
        # pas l'alteration ne vaut pas mieux qu'un fichier sans empreinte --
        # il ATTESTE la conservation sans la prouver.
        v = verifier_archive(self.archive)
        self.assertTrue(v['verifiable'])
        self.assertTrue(v['intact'], v['ecarts'])

        cible = os.path.join(self.archive['dossier'], 'commentaire.txt')
        octets = open(cible, 'rb').read()
        try:
            with open(cible, 'wb') as fh:      # UN SEUL octet change
                fh.write(octets[:-1] + b'X')
            v2 = verifier_archive(self.archive)
            self.assertFalse(v2['intact'], 'une alteration passe inapercue')
            self.assertEqual(len(v2['ecarts']), 1)
            self.assertIn('altere', v2['ecarts'][0])
        finally:
            with open(cible, 'wb') as fh:
                fh.write(octets)
        self.assertTrue(verifier_archive(self.archive)['intact'],
                        'le fichier n a pas ete restaure')
        print('    OK C3-3 un octet modifie fait echouer la verification')

    def test_p3bis_un_fichier_retire_est_detecte(self):
        import shutil
        cible = os.path.join(self.archive['dossier'], 'donnees.xlsx')
        garde = cible + '.garde'
        shutil.move(cible, garde)
        try:
            v = verifier_archive(self.archive)
            self.assertFalse(v['intact'])
            self.assertIn('absent', v['ecarts'][0])
        finally:
            shutil.move(garde, cible)
        print('    OK C3-3b un fichier retire est detecte')

    # ── P4 : il est RELIE a son audit trail ─────────────────────────────────
    def test_p4_les_empreintes_vivent_dans_l_audit_archive(self):
        # ⚠️ DANS L'AUDIT, PAS A COTE : c'est le seul fichier ecrit a chaque
        # run, donc le seul qui puisse porter de quoi verifier les autres.
        for nom, meta in self.archive['fichiers'].items():
            self.assertEqual(len(meta['sha256']), 64, nom)
            self.assertRegex(meta['sha256'], r'^[0-9a-f]{64}$', nom)
        print('    OK C3-4 les empreintes SHA-256 sont dans l audit')

    # ── LA PORTEE, ECRITE LA OU ON POURRAIT CROIRE LE CONTRAIRE ─────────────
    def test_la_portee_est_publiee_et_ne_promet_pas_l_opposabilite(self):
        # ⚠️ TROIS ENDROITS : la constante, l'audit archive, et le verdict de
        # `verifier_archive`. Un mecanisme qui laisserait croire qu'il rend un
        # dossier OPPOSABLE serait exactement ce que ce chantier combat.
        for ou, texte in (('archive', self.archive.get('porte', '')),
                          ('verdict', verifier_archive(self.archive)['porte']),
                          ('constante', PORTEE_ARCHIVE)):
            self.assertIn('verifiable', texte, ou)
            self.assertIn('signature', texte, ou)
            self.assertIn('hors du systeme', texte, ou)
        print('    OK C3-5 la portee est publiee aux trois endroits')

    # ── SANS ARCHIVAGE : rien n'est ecrit, et rien n'est promis ─────────────
    def test_sans_archiver_rien_n_est_ecrit_et_rien_n_est_affirme(self):
        # ⚠️ CONTRE-EPREUVE : `archiver=False` par defaut, sinon la gate
        # ecrirait 1,07 Mo a chacun de ses 1 232 tests. Et `verifier_archive`
        # doit rendre `verifiable: False` -- JAMAIS `intact: True`, qui se
        # lirait << verifie, tout va bien >>.
        import tempfile
        d = tempfile.mkdtemp()
        r = AgentA7Provisionnement(models_path=d, audit_path=d,
                                   verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False)
        self.assertTrue(r.get('success'))
        ecrits = [x for x in os.listdir(d) if not x.endswith('.json')]
        self.assertEqual(ecrits, [], f'des fichiers ont ete ecrits : {ecrits}')
        v = verifier_archive(self._audit(d)['archive'])
        self.assertFalse(v['verifiable'])
        self.assertIsNone(v['intact'], 'un dossier absent se dit INTACT')
        print('    OK C3-6 sans archivage : rien ecrit, rien affirme')

    # ── AUCUN EURO DEPLACE ──────────────────────────────────────────────────
    def test_aucun_agregat_n_a_bouge(self):
        n4 = self.r['n4']
        self.assertEqual(n4['best_estimate'], 14830899.0)
        self.assertEqual(n4['scr']['scr_provisions'], 4894197.0)
        self.assertEqual(n4['reserve_p90'], 18053284.0)
        print('    OK C3-7 BE / SCR / P90 inchanges')


# =============================================================================
#  D3 — UN SILENCE DOCUMENTE SE PUBLIE, IL NE SE CACHE PAS
# =============================================================================
#
#  ⚠️⚠️ << BRANCHE NON IDENTIFIEE >> ETAIT FAUX POUR 7 LoB SUR 15. Mesure :
#  quinze branches sont CONFIGUREES -- seuils H2, sigma_EIOPA, segment S2,
#  reference de marche -- et seules HUIT ont un bloc de prose. Les sept autres
#  tombaient sur un repli qui les declarait INCONNUES alors que le module lit
#  leurs parametres et les applique.
#
#  ⚠️ DEUX FAITS DIFFERENTS, ET LE SECOND SE PUBLIE : << branche inconnue >>
#  n'est pas << branche configuree sans prose documentee >>. Le taire derriere
#  le premier ferait croire a une lacune de PARAMETRAGE qui n'existe pas.
#
#  ⚠️ ET L'INTITULE DU PLAN ETAIT FAUX : D3 disait << sourcer les 14 branches >>.
#  Il n'y a rien a sourcer -- le referentiel est complet. Il y a une VALEUR a
#  CONSOMMER et un SILENCE a DECLARER.


class T_D3_Branche_Configuree_Sans_Prose(unittest.TestCase):
    """⚠️ CONFIGUREE SANS PROSE N'EST PAS INCONNUE."""

    @classmethod
    def setUpClass(cls):
        cls.r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False)

    def _contexte(self, lob):
        return _narration_lob(lob, self.r['n1'], self.r['n2'],
                              self.r['n3'], self.r['n4'])['contexte']

    # ── D3-1 : l'ecart qui a ouvert le lot, mesure ──────────────────────────
    def test_sept_lob_sont_configurees_sans_prose(self):
        # ⚠️ SI CET ECART DISPARAIT, LE LOT PERD SON OBJET : le test le dit
        # plutot que de le supposer. Selasse peut rediger ces proses quand il
        # veut -- ce compte baissera, et c'est le but.
        sans = set(LOB_CONFIG) - set(_LOB_HANDLERS)
        self.assertTrue(sans, 'toutes les LoB ont une prose : lot sans objet')
        self.assertEqual(len(LOB_CONFIG), 15)
        print(f'    OK D3-1 {len(sans)} LoB configurees sans prose redigee')

    # ── D3-2 : LE CŒUR — les deux etats ne se confondent plus ───────────────
    def test_une_branche_configuree_n_est_plus_dite_inconnue(self):
        for lob in sorted(set(LOB_CONFIG) - set(_LOB_HANDLERS)):
            c = self._contexte(lob)
            self.assertIn('CONFIGURÉE', c, lob)
            self.assertNotIn('NON IDENTIFIÉE', c,
                             f'{lob} est configuree et se dit inconnue')
            # Et le silence est NOMME, pas comble par du texte plausible.
            self.assertIn('SANS ANALYSE RÉDIGÉE', c, lob)
            self.assertIn(LOB_CONFIG[lob]['label'], c, lob)
        print('    OK D3-2 une branche configuree ne se dit plus inconnue')

    def test_une_branche_inconnue_se_dit_inconnue(self):
        # ⚠️ CONTRE-EPREUVE : un correctif qui declarerait TOUT << configure >>
        # serait aussi faux que l'inverse.
        c = self._contexte('branche_qui_n_existe_pas')
        self.assertIn('NON IDENTIFIÉE', c)
        self.assertNotIn('SANS ANALYSE RÉDIGÉE', c)
        print('    OK D3-3 une branche inconnue se dit inconnue')

    # ── D3-4 : les seuils sont LUS, pas reecrits ────────────────────────────
    def test_les_seuils_publies_sont_ceux_du_referentiel(self):
        # ⚠️ L'ANCIEN TEXTE RECOPIAIT << H2 CV=15%, derive=20% >> A LA MAIN :
        # juste pour `generique`, faux pour toute autre branche passant ici.
        # Cat-Nat applique 25 % / 35 % -- l'ecart le plus large des quinze.
        for lob in ('catastrophes_naturelles', 'mrh_inexistant', 'generique'):
            cfg = get_lob_config(lob)
            c = self._contexte(lob)
            self.assertIn(f"CV<{cfg['h2_seuil_cv']:.0%}", c, lob)
            self.assertIn(f"dérive<{cfg['h2_seuil_derive']:.0%}", c, lob)
        cat = self._contexte('catastrophes_naturelles')
        self.assertIn('CV<25%', cat)
        self.assertIn('dérive<35%', cat)
        print('    OK D3-4 les seuils viennent du referentiel, pas de la prose')

    # ── D3-5 : aucune prose n'est inventee ──────────────────────────────────
    def test_aucune_prose_de_branche_n_est_fabriquee(self):
        # ⚠️ LE REFUS QUI FONDE CE LOT : ecrire << sur Credit/Caution, le rejet
        # de H1 est frequent parce que... >> serait une affirmation sur un
        # portefeuille reel, du meme genre que celles qu'un modele a inventees
        # le 2026-06-24 et que ce chantier a passe une semaine a retirer.
        for lob in sorted(set(LOB_CONFIG) - set(_LOB_HANDLERS)):
            blocs = _narration_lob(lob, self.r['n1'], self.r['n2'],
                                   self.r['n3'], self.r['n4'])
            for cle in ('hypotheses', 'methodes', 'recommandations'):
                self.assertEqual(blocs[cle], '',
                                 f'{lob} : une prose a ete fabriquee en {cle}')
        print('    OK D3-5 aucune prose de branche n est fabriquee')

    # ── D3-6 : les 8 LoB redigees gardent leur prose ────────────────────────
    def test_les_lob_redigees_ne_perdent_rien(self):
        for lob in sorted(k for k in _LOB_HANDLERS if k != 'generique'):
            c = self._contexte(lob)
            self.assertNotIn('SANS ANALYSE RÉDIGÉE', c, lob)
            self.assertGreater(len(c), 200, lob)
        print('    OK D3-6 les LoB redigees gardent leur analyse')

    def test_aucun_agregat_n_a_bouge(self):
        # ⚠️ LES VALEURS DU RUN AVEC EXPOSITION, celui de `setUpClass`. La
        # premiere version portait celles du run SANS primes (18 680 856) --
        # une erreur de recopie, attrapee par le filet lui-meme.
        n4 = self.r['n4']
        self.assertEqual(n4['best_estimate'], 14830899.0)
        self.assertEqual(n4['scr']['scr_provisions'], 4894197.0)
        print('    OK D3-7 BE / SCR inchanges')


# =============================================================================
#  VERROU C2 — AUCUN NOMBRE PUBLIE QUI NE SOIT DANS LA CHARGE UTILE
# =============================================================================
#
#  ⚠️⚠️ CE VERROU PROUVE LA PROVENANCE, JAMAIS LA JUSTESSE -- et c'est ecrit
#  dans le module, pas seulement ici. Un modele qui INVERSERAIT deux valeurs
#  justes passerait : les deux sont dans la charge utile.
#
#  ⚠️ IL A TROUVE TROIS DEFAUTS AVANT D'EXISTER. Sa mesure preparatoire a fait
#  ressortir les trois provisions du jugement (`fcfb3d3`) -- un ecart de
#  8,4 M EUR sur le P99.5 -- que SIX releves successifs n'avaient pas vus.
#
#  ⚠️ ET LE DETECTEUR EST TESTE POUR LUI-MEME, AVANT D'ETRE CRU. Une premiere
#  version ne reconnaissait pas la virgule comme separateur de milliers : elle
#  scindait << 8,057,830 >> en trois nombres, tous orphelins, et annoncait 42 %
#  de faux positifs. La narration n'inventait rien -- le detecteur fabriquait
#  des orphelins. Un verrou ne vaut pas mieux que son instrument.

#: Formes mesurees dans la narration et la charge utile reelles. Au niveau du
#: module : un attribut de classe mutable demanderait un `ClassVar`, et ces
#: formes n'appartiennent pas plus a une classe qu'a une autre.
_FORMES_MESUREES = {
    '8,057,830':   '8057830',      # virgule -- celle qui m'avait trompe
    '1.234.567':   '1234567',      # point separateur de milliers
    '4 894 197':   '4894197',      # espace fine insecable
    '11.0':        '11',           # decimale nulle == entier
    '17,3':        '17.3',         # decimale a la francaise
    '0.5':         '0.5',
    '7':           '7',
    # ⚠️⚠️ LE TROU DE L'ASSIETTE, MESURE SUR DU LLM REEL. Aucune forme testee
    # n'avait EXACTEMENT TROIS decimales avec une partie entiere nulle :
    # `0,309` devenait `309`. Le detecteur etait teste, pas sur ce point.
    '0,309':       '0.309',
    '0,131':       '0.131',
    '0,500':       '0.5',        # la decimale nulle tombe, la valeur reste
    # Et la contre-epreuve : un zero NON isole reste un separateur de milliers.
    '10,500':      '10500',
}


class T_Detecteur_De_Nombres(unittest.TestCase):
    """⚠️ LE DETECTEUR AVANT LE VERROU. S'il se trompe, tout le reste ment."""

    def test_les_formes_mesurees_sont_reconnues(self):
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            _cle_nombre,
            nombres_publies,
        )
        for brut, attendu in _FORMES_MESUREES.items():
            trouves = nombres_publies(brut)
            self.assertEqual(trouves, [brut], f'{brut!r} mal decoupe')
            self.assertEqual(_cle_nombre(brut), attendu, f'{brut!r} mal normalise')
        print(f'    OK C2-1 les {len(_FORMES_MESUREES)} formes mesurees '
              f'sont reconnues')

    def test_une_decimale_non_nulle_ne_s_apparie_pas_a_l_entier(self):
        # ⚠️ SANS CETTE EPREUVE, LA NORMALISATION POURRAIT TOUT APLATIR et le
        # verrou laisserait passer n'importe quel arrondi.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            _cle_nombre,
        )
        self.assertEqual(_cle_nombre('11.0'), _cle_nombre('11'))
        self.assertNotEqual(_cle_nombre('11.3'), _cle_nombre('11'))
        self.assertNotEqual(_cle_nombre('17'), _cle_nombre('17,3'))
        print('    OK C2-2 un arrondi ne s apparie pas a sa valeur pleine')

    def test_une_ponctuation_ne_soude_pas_deux_nombres(self):
        # ⚠️⚠️ MESURE SUR DU LLM REEL : << Score 83,2/100, 2 alertes >> rendait
        # << 100, 2 >> comme UN SEUL nombre. Un separateur de milliers ne
        # porte JAMAIS d'espace apres lui -- << 100, 2 >> est une virgule de
        # ponctuation suivie d'un second nombre.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            nombres_publies,
        )
        self.assertEqual(nombres_publies('100, 2'), ['100', '2'])
        self.assertEqual(nombres_publies('Score 83,2/100, 2 alertes'),
                         ['83,2', '100', '2'])
        print('    OK C2-2c une ponctuation ne soude pas deux nombres')

    def test_les_separateurs_colles_restent_des_separateurs(self):
        # ⚠️ CONTRE-EPREUVE : le correctif ne doit rien casser de ce qui
        # marchait. Les quatre espaces de milliers sont testes un par un.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            nombres_publies,
        )
        for brut in ('8,057,830', '1.234.567', '0,309', '11.0',
                     '4 894 197', '14 830 899',
                     '2 447 095', '18 680 856'):
            self.assertEqual(nombres_publies(brut), [brut],
                             f'{brut!r} scinde a tort')
        print('    OK C2-2d les 4 espaces et les separateurs colles tiennent')

    def test_les_espaces_du_motif_sont_lisibles_dans_le_source(self):
        # ⚠️ ILS ETAIENT ECRITS EN CLAIR : invisibles, ils ont fait echouer une
        # edition de ce fichier. Personne ne peut relire ce qu'il ne voit pas.
        src = inspect.getsource(RAPP_MOD)
        ligne = [x for x in src.splitlines()
                 if x.startswith('_ESPACES_MILLIERS =')]
        self.assertEqual(len(ligne), 1)
        self.assertFalse(any(ord(c) > 127 for c in ligne[0]),
                         'un espace invisible est revenu dans le source')
        print('    OK C2-2e les espaces du motif sont nommes, pas invisibles')

    def test_une_decimale_a_partie_entiere_nulle_n_est_pas_des_milliers(self):
        # ⚠️⚠️ LE DEFAUT MESURE SUR LE PREMIER RUN LLM REEL. La charge portait
        # `p = 0.3090`, le modele publiait `0,309` : MEME VALEUR, deux cles.
        # Trois des 34 signalements etaient FABRIQUES par le detecteur.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            _cle_nombre,
        )
        self.assertEqual(_cle_nombre('0.3090'), _cle_nombre('0,309'),
                         'la charge et la narration ne se correspondent plus')
        self.assertNotEqual(_cle_nombre('0,309'), _cle_nombre('309'),
                            'FAUX NEGATIF : un entier passe pour une decimale')
        # ⚠️ ET LE CAS OPPOSE NE DOIT PAS BOUGER : un zero non isole ouvre
        # bien des milliers.
        self.assertEqual(_cle_nombre('10,500'), '10500')
        self.assertEqual(_cle_nombre('8,057,830'), '8057830')
        print('    OK C2-2b 0,309 est une decimale, 10,500 des milliers')

    def test_les_formes_de_reference_du_llm_reel_sont_muettes(self):
        # ⚠️ TOUTES MESUREES SUR LA PREMIERE NARRATION LLM DU DEPOT, jamais
        # devinees. Chacune ressortait orpheline.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            nombres_publies,
        )
        for s in ('articles 77 et 115', 'article 77 S2', '(Clark, 2003)',
                  'Guide IA 2023', "Guide de l'Institut des Actuaires 2023",
                  'Mack (1993)', 'Art. 115', '§7', 'H1', 'BFCC-H5'):
            self.assertEqual(nombres_publies(s), [],
                             f'{s!r} publie encore un nombre')
        print('    OK C2-3b les 10 formes de reference mesurees sont muettes')

    def test_l_elargissement_ne_rend_aucune_grandeur_muette(self):
        # ⚠️⚠️ LA REGLE D'ASYMETRIE : une liste qui EXEMPTE ouvre un trou.
        # Le sens 2 de la calibration est donc OBLIGATOIRE -- et il porte sur
        # les DEUX vrais defauts de la narration reelle, qui doivent rester
        # signales : un montant fabrique et une soustraction fausse.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            nombres_publies,
        )
        for s, attendu in (('soit 910 000 € minimum', '910 000'),
                           ('est de 10 622 026 €', '10 622 026'),
                           ('BE de 14 830 899 €', '14 830 899'),
                           ('2 alertes rouges', '2'),
                           ('CV de 13,1 %', '13,1'),
                           ('au taux de 6 %', '6')):
            self.assertIn(attendu, nombres_publies(s),
                          f'{s!r} : la grandeur est devenue muette')
        print('    OK C2-3c aucune grandeur ne tombe dans la zone franche')

    def test_la_zone_franche_ne_couvre_que_des_references(self):
        # ⚠️ LA REGLE DE SELASSE : aucun seuil numerique n'y entre.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            nombres_publies,
        )
        for ref in ('§5 — CONTEXTE', 'Mack 1993', 'Mack (1993)', 'BFCC-H4',
                    'H2', 'Art. 115', 'QIS5 TP.5.26', '2015/35'):
            self.assertEqual(nombres_publies(ref), [],
                             f'{ref!r} devrait etre en zone franche')
        # ⚠️ ET CE QUI N'EST PAS UNE REFERENCE DOIT RESTER CONTROLE.
        for grandeur in ('5 annees', '3 observations', 'CV de 11.0 %'):
            self.assertTrue(nombres_publies(grandeur),
                            f'{grandeur!r} ne doit PAS etre exempte')
        print('    OK C2-3 zone franche = references seules, aucun seuil')


class T_Verrou_Charge_Utile(unittest.TestCase):
    """⚠️ UN NOMBRE PUBLIE QUI N'EST PAS DANS LA CHARGE EST SOIT INVENTE, SOIT
    PERIME."""

    @classmethod
    def setUpClass(cls):
        with kaleido_declare(True), rendeur_substitue():
            cls.r = AgentA7Provisionnement(verbose=False).run(
                source=np.array(GENINS, dtype=float), mode_declare='cumule',
                primes=_exposition(GENINS), n_sim_bootstrap=200, seed=42,
                generer_graphiques=False)

    def test_un_nombre_invente_est_detecte(self):
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            orphelins_narration,
        )
        charge = 'BEST ESTIMATE : 14 830 899 € | CV=2.0 %'
        propre = 'Le Best Estimate ressort a 14 830 899 € (CV 2.0 %).'
        self.assertEqual(orphelins_narration(propre, charge), [])
        invente = 'Le Best Estimate ressort a 99 999 999 € (CV 2.0 %).'
        self.assertEqual(orphelins_narration(invente, charge), ['99 999 999'])
        print('    OK C2-4 un nombre absent de la charge est signale')

    def test_le_taux_reste_praticable_sur_la_narration_reelle(self):
        # ⚠️ MESURE SUR LE DETERMINISTE, PAS SUR LE LLM -- aucune cle API
        # disponible, aucune narration Claude archivee. Ordre de grandeur.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            _construire_contexte,
            nombres_publies,
            orphelins_narration,
        )
        n2, n3, n4 = self.r['n2'], self.r['n3'], self.r['n4']
        texte = n4.get('jugement', '')
        charge = _construire_contexte(n2, n3, n4, 'RC', '30/06/2026')
        pub = nombres_publies(texte)
        orph = orphelins_narration(texte, charge)
        taux = 100 * len(orph) / max(len(pub), 1)
        self.assertGreater(len(pub), 20, 'narration trop pauvre pour mesurer')
        # ⚠️ SEUIL LARGE ET ASSUME : il verrouille que le detecteur ne se
        # DEREGLE pas (retour a 42 %), pas que la narration soit parfaite.
        self.assertLess(taux, 30,
                        f'taux d orphelins {taux:.0f} % -- detecteur deregle ?')
        print(f'    OK C2-5 taux mesure {taux:.0f} % '
              f'({len(orph)}/{len(pub)}) sur le deterministe')

    def test_le_verrou_ne_juge_pas_la_justesse(self):
        # ⚠️ LA RESERVE DE FOND, EPROUVEE PLUTOT QUE SEULEMENT ECRITE. Deux
        # valeurs justes ECHANGEES passent le verrou : c'est sa limite, et
        # elle doit etre demontree, pas promise.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            orphelins_narration,
        )
        charge = 'BE : 14 830 899 € | SCR : 4 894 197 €'
        inverse = 'Le BE vaut 4 894 197 € et le SCR 14 830 899 €.'
        self.assertEqual(orphelins_narration(inverse, charge), [],
                         'le verrou pretendrait juger la justesse')
        print('    OK C2-6 le verrou prouve la provenance, PAS la justesse')


class T_Le_Repli_Du_Repli_Dit_Qu_Il_Est_Une_Panne(unittest.TestCase):
    """⚠️ UN RAPPORT DEGRADE QUI NE LE DIT PAS FAIT PASSER UNE PANNE POUR UN
    FONCTIONNEMENT NORMAL.

    Le repli `commentaire or jugement` rendait les deux etats sous une meme
    source `templates` : le badge d'un rapport ampute etait celui d'un
    rapport complet.
    """

    @classmethod
    def setUpClass(cls):
        cls.r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False)
        cls.n4 = cls.r['n4']
        cls.comm = cls.r['commentaire']

    # ── les trois etats sont distincts ────────────────────────────────────
    def test_commentaire_present_aucun_signal_de_panne(self):
        txt, src = RAPP_MOD._narration_templates(self.n4, self.comm)
        self.assertEqual(src, 'templates')
        self.assertNotIn(RAPP_MOD.SIGNAL_NARRATION_DEGRADEE, txt,
                         'un rapport complet se declarerait degrade')
        print('    OK REPLI-1 commentaire present -> templates, sans signal')

    def test_commentaire_absent_le_signal_est_dans_le_texte(self):
        txt, src = RAPP_MOD._narration_templates(self.n4, '')
        self.assertEqual(src, 'jugement_degrade',
                         'la panne passe pour un mode normal')
        self.assertIn(RAPP_MOD.SIGNAL_NARRATION_DEGRADEE, txt,
                      'le texte degrade ne dit pas ce qu il est')
        # ⚠️ DANS LE TEXTE, PAS SEULEMENT DANS LE BADGE : le signal doit
        # survivre au copier-coller et a la lecture du fichier archive.
        self.assertTrue(txt.startswith(RAPP_MOD.SIGNAL_NARRATION_DEGRADEE),
                        'le signal doit precede le contenu, pas le suivre')
        print('    OK REPLI-2 commentaire absent -> signal EN TETE du texte')

    def test_le_jugement_est_conserve_pas_remplace(self):
        # ⚠️ ON AVERTIT, ON NE PRIVE PAS. Quand le commentaire manque, le
        # jugement est tout ce qui reste : le jeter couterait au lecteur.
        txt, _ = RAPP_MOD._narration_templates(self.n4, '')
        for temoin in ('BEST ESTIMATE', 'SCR PROVISIONS'):
            self.assertIn(temoin, txt,
                          f'{temoin} perdu : le repli prive au lieu d avertir')
        print('    OK REPLI-3 le jugement est conserve sous le signal')

    def test_ni_commentaire_ni_jugement_rend_aucune(self):
        txt, src = RAPP_MOD._narration_templates({}, '')
        self.assertEqual((txt, src), ('', 'aucune'))
        print('    OK REPLI-4 rien a publier -> aucune, sans signal orphelin')

    # ── le badge suit, et il nomme la panne ───────────────────────────────
    def test_le_badge_distingue_la_panne_du_mode_normal(self):
        degrade = RAPP_MOD.badge_narration('jugement_degrade')
        normal = RAPP_MOD.badge_narration('templates')
        self.assertTrue(degrade, 'aucun badge sur un rapport degrade')
        self.assertNotEqual(degrade, normal,
                            'le badge ne distingue pas la panne')
        self.assertIn('dégradé', degrade.lower())
        print(f'    OK REPLI-5 badge degrade distinct : {degrade[:52]}')

    def test_l_entree_templates_n_a_pas_change(self):
        # ⚠️ TROIS MODULES HORS A7 CITENT CE LIBELLE EN PROSE pour dire qu il
        # est « deja honnete ». Le reecrire perimerait leurs mentions.
        self.assertEqual(RAPP_MOD._LIBELLE_SOURCE['templates'],
                         '📝 Mode standard')
        print('    OK REPLI-6 `templates` intact — 3 mentions hors A7 saines')

    # ── le signal atteint les DEUX documents signes ───────────────────────
    def test_le_signal_atteint_le_html_et_le_word(self):
        from docx import Document
        txt, src = RAPP_MOD._narration_templates(self.n4, '')
        html = RAPP_MOD.export_html(
            self.r['n1'], self.r['n2'], self.r['n3'], self.n4,
            narration=txt, source_narration=src)
        self.assertIn('RAPPORT DÉGRADÉ', html,
                      'le HTML signe ne dit pas qu il est degrade')
        w = RAPP_MOD.export_word(
            self.r['n1'], self.r['n2'], self.r['n3'], self.n4,
            narration=txt, source_narration=src)
        mot = '\n'.join(p.text for p in Document(io.BytesIO(w)).paragraphs)
        self.assertIn('RAPPORT DÉGRADÉ', mot,
                      'le Word transmis au CAC ne dit pas qu il est degrade')
        print('    OK REPLI-7 le signal est dans le HTML ET dans le Word')

    # ── l audit trail porte l etat ────────────────────────────────────────
    def test_l_audit_trail_porte_la_source_degradee(self):
        ctrl = RAPP_MOD.controle_narration('peu importe', 'jugement_degrade', '')
        self.assertEqual(ctrl['source'], 'jugement_degrade',
                         'l audit ne conserve pas la trace de la panne')
        print('    OK REPLI-8 la source degradee entre dans l audit trail')

    # ── et le cas est RARE : la production ne l atteint pas ───────────────
    def test_un_run_reel_ne_produit_jamais_l_etat_degrade(self):
        # ⚠️ CE LOT FERME UN CAS RARE, ET LE VERROUILLER EVITE DE LE
        # SURESTIMER. `generer_commentaire` ne peut pas rendre vide : huit
        # sections en dur, un seul `return`. Sur donnees degradees il LEVE,
        # et l agent ne l entoure d aucun `except` — le run tombe avant.
        self.assertTrue(self.comm.strip(),
                        'un run reel a rendu un commentaire vide')
        _, src = RAPP_MOD._narration_templates(self.n4, self.comm)
        self.assertEqual(src, 'templates',
                         'un run reel est passe par le chemin degrade')
        with self.assertRaises(KeyError):
            COMM_MOD.generer_commentaire(n1={}, n2={}, n3={}, n4={})
        print('    OK REPLI-9 cas rare : un run reel ne degrade jamais')


class T_D2_Le_Socle_N_Est_Pas_Substituable(unittest.TestCase):
    """⚠️ LA NARRATION LLM NE REMPLACE PLUS LE COMMENTAIRE DETERMINISTE.

    Mesure d'ouverture, narration LLM SIMULEE : la ligne du Best Estimate
    avec sa base reglementaire (`Art. 77`) DISPARAISSAIT des deux documents
    signes des qu'une cle API existait.
    """

    FAUSSE = ("§1 — CONTEXTE\n\nLe triangle porte dix annees.\n\n"
              "§7 — CONCLUSION\n\nLe provisionnement est adequat.")
    TEMOIN = 'BEST ESTIMATE'

    @classmethod
    def setUpClass(cls):
        cls.r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False)
        cls.comm = cls.r['commentaire']

    def _html(self, narration, source):
        return RAPP_MOD.export_html(
            self.r['n1'], self.r['n2'], self.r['n3'], self.r['n4'],
            commentaire=self.comm, narration=narration,
            source_narration=source)

    # ── la regle, la ou elle vit ──────────────────────────────────────────
    def test_le_socle_ne_se_republie_pas_s_il_est_deja_la(self):
        # ⚠️ SANS CLE API, LA NARRATION *EST* LE COMMENTAIRE. Publier les deux
        # livrerait 14 150 caracteres EN DOUBLE a tous les utilisateurs
        # actuels — une regression introduite par le lot cense proteger.
        deja = RAPP_MOD._clean(self.comm)
        self.assertEqual(RAPP_MOD.socle_a_publier(self.comm, deja), '',
                         'le socle serait publie DEUX fois')
        print('    OK D2-1 narration == socle -> aucune republication')

    def test_le_socle_est_rendu_quand_la_narration_ne_le_porte_pas(self):
        rendu = RAPP_MOD.socle_a_publier(self.comm, self.FAUSSE)
        self.assertIn(self.TEMOIN, rendu,
                      'le socle manque alors que la narration ne le porte pas')
        print('    OK D2-2 narration LLM -> le socle est rendu')

    def test_sans_commentaire_il_n_y_a_pas_de_socle(self):
        self.assertEqual(RAPP_MOD.socle_a_publier('', self.FAUSSE), '')
        self.assertEqual(RAPP_MOD.socle_a_publier('   ', self.FAUSSE), '')
        print('    OK D2-3 pas de commentaire -> pas de bloc socle vide')

    # ── l invariant, sur le document ──────────────────────────────────────
    def test_le_socle_atteint_le_html_exactement_une_fois(self):
        # ⚠️ L'INVARIANT DU LOT, ET IL PORTE SUR LES TROIS CHEMINS : le
        # commentaire deterministe atteint le document signe UNE fois.
        for narration, source, attendu in (
                (self.FAUSSE, 'claude_api', 1),
                (RAPP_MOD._clean(self.comm), 'templates', 0),
                ('', 'aucune', 1)):
            html = self._html(narration, source)
            self.assertEqual(html.count(RAPP_MOD.PORTEE_SOCLE), attendu,
                             f'bloc socle mal compte sur {source!r}')
            self.assertIn(self.TEMOIN, html,
                          f'le socle n atteint pas le HTML sur {source!r}')
        print('    OK D2-4 socle une seule fois, sur les TROIS chemins')

    def test_la_ligne_du_be_revient_dans_le_html_signe(self):
        # ⚠️ LE TEMOIN DE LA MESURE D'OUVERTURE : cette ligne disparaissait.
        html = self._html(self.FAUSSE, 'claude_api')
        self.assertIn('Art. 77', html,
                      'la base reglementaire du BE manque au document signe')
        print('    OK D2-5 la ligne du BE avec Art. 77 est de retour')

    def test_le_socle_atteint_le_word_avec_sa_portee(self):
        from docx import Document
        w = RAPP_MOD.export_word(
            self.r['n1'], self.r['n2'], self.r['n3'], self.r['n4'],
            commentaire=self.comm, narration=self.FAUSSE,
            source_narration='claude_api')
        txt = '\n'.join(p.text for p in Document(io.BytesIO(w)).paragraphs)
        self.assertIn(self.TEMOIN, txt,
                      'le Word transmis au CAC a perdu le socle')
        self.assertIn(RAPP_MOD.PORTEE_SOCLE, txt,
                      'le Word publie le socle SANS sa portee')
        print('    OK D2-6 le socle et sa portee atteignent le Word')

    # ── la portee, et ce qu elle empeche de croire ────────────────────────
    def test_la_portee_dit_qu_elle_ne_valide_rien(self):
        # ⚠️ DEUX TEXTES COTE A COTE LAISSENT CROIRE QUE L'UN VALIDE L'AUTRE.
        # Rien ne compare les deux : la phrase doit le dire, pas le suggerer.
        p = RAPP_MOD.PORTEE_SOCLE
        self.assertIn('ne vérifie PAS', p)
        self.assertIn("n'atteste de rien", p)
        print('    OK D2-7 la portee refuse explicitement le controle')

    def test_la_portee_n_affirme_que_ce_qui_est_mesure(self):
        # ⚠️⚠️ MA PREMIERE VERSION DISAIT « il se reproduit A L'IDENTIQUE d'un
        # run a l'autre ». JAMAIS MESURE, ET FAUX : le socle porte deux
        # horodatages. Ce test tient les DEUX moities du fait.
        p = RAPP_MOD.PORTEE_SOCLE
        self.assertNotIn("à l'identique", p,
                         'la portee revendique une reproductibilite exacte')
        self.assertIn('date de génération', p,
                      'la portee tait la seule variation qui existe')
        # La moitie mesuree : a donnees egales, les enonces ne bougent pas.
        autre = COMM_MOD.generer_commentaire(
            n1=self.r['n1'], n2=self.r['n2'], n3=self.r['n3'],
            n4=self.r['n4'])
        sans_date = re.compile(r'\d{2}/\d{2}/\d{4}')
        self.assertEqual(sans_date.sub('', autre),
                         sans_date.sub('', self.comm),
                         'a donnees egales, les enonces du socle ont bouge')
        print('    OK D2-10 la portee dit le fait mesure, et lui seul')

    def test_la_portee_vient_d_une_seule_constante(self):
        # ⚠️ RECOPIEE DANS LES DEUX FORMATS, ELLE DIVERGERAIT — c'est ce qui
        # etait arrive aux deux tables de libelles avant qu'on les fonde.
        src = inspect.getsource(RAPP_MOD)
        litteral = "Socle dérivé du calcul par du code déterministe"
        self.assertEqual(src.count(litteral), 1,
                         'la phrase de portee est recopiee : elle divergera')
        print('    OK D2-8 une seule occurrence litterale de la portee')

    # ── deux blocs, deux badges ───────────────────────────────────────────
    def test_le_socle_a_son_propre_badge(self):
        # ⚠️ TOUT L'ARGUMENT DE Q1 : le conteneur porte UN badge d'origine ;
        # deux textes de provenances differentes le rendraient faux.
        socle = RAPP_MOD.badge_narration('socle')
        self.assertTrue(socle, 'le bloc socle ne dit pas son origine')
        self.assertNotEqual(socle, RAPP_MOD.badge_narration('claude_api'))
        html = self._html(self.FAUSSE, 'claude_api')
        self.assertIn(RAPP_MOD._LIBELLE_SOURCE['socle'], html)
        print(f'    OK D2-9 badge propre au socle : {socle[:44]}')


class T_C_Une_Prescription_Chiffree_Non_Sourcee_Est_Marquee(unittest.TestCase):
    """⚠️ LE MODELE DECRIT, IL NE PRESCRIT PAS DE CHIFFRE SANS SOURCE.

    Mesure d'ouverture : sur cinq montants orphelins d'un run reel, QUATRE
    etaient des soustractions justes au centime. Le seul fabrique etait dans
    la seule phrase qui PRESCRIT — celle adressee au Conseil.
    """

    CHARGE = 'BE : 14 830 899 € | SCR : 4 894 197 € | P75 : 16 343 466 €'

    def _mq(self, texte):
        return RAPP_MOD.marquer_prescriptions_chiffrees(texte, self.CHARGE)

    # ── ce qui est marque, et ce qui ne l est pas ─────────────────────────
    def test_prescription_avec_chiffre_orphelin_est_marquee(self):
        _, n = self._mq('Il est recommandé de provisionner 999 999 999 €.')
        self.assertEqual(n, 1, 'la prescription non sourcee passe')
        print('    OK C-1 prescription + chiffre orphelin -> marquee')

    def test_une_phrase_descriptive_n_est_jamais_marquee(self):
        # ⚠️ LA DISTINCTION EST TOUT LE LOT : un chiffre descriptif faux est
        # refutable par son contexte, un chiffre prescriptif faux engage.
        _, n = self._mq("L'écart atteint 999 999 999 € sur la période.")
        self.assertEqual(n, 0, 'une description est marquee a tort')
        print('    OK C-2 une phrase descriptive n est pas marquee')

    def test_une_prescription_sourcee_n_est_pas_marquee(self):
        _, n = self._mq('Il est recommandé de retenir 14 830 899 €.')
        self.assertEqual(n, 0, 'une prescription sourcee est marquee a tort')
        print('    OK C-3 prescription dont le chiffre est transmis -> muette')

    def test_sans_charge_utile_rien_n_est_marque(self):
        # ⚠️ Hors du chemin LLM il n'y a pas de charge : le controle n'a pas
        # d'objet, et une marque y affirmerait un defaut qu'on ne mesure pas.
        _, n = RAPP_MOD.marquer_prescriptions_chiffrees(
            'Il est recommandé de provisionner 999 999 999 €.', '')
        self.assertEqual(n, 0)
        print('    OK C-4 pas de charge utile -> aucune marque')

    # ── L INVARIANT : ON N ENLEVE RIEN ────────────────────────────────────
    def test_le_marquage_n_enleve_ni_ne_modifie_rien(self):
        # ⚠️⚠️ C'EST L'ARBITRAGE DU LOT, EPROUVE PLUTOT QU'ECRIT. On a REFUSE
        # de retirer la phrase : elle portait aussi un avertissement legitime
        # et un calcul juste. Retirer la marque doit rendre le texte D'ORIGINE.
        src = ("Le BE atteint 14 830 899 €. Il est recommandé de "
               "provisionner 999 999 999 € de plus.\nAutre ligne.")
        marque, n = self._mq(src)
        self.assertEqual(n, 1)
        restitue = marque.replace(
            ' ' + RAPP_MOD.PORTEE_MARQUE_PRESCRIPTION, '')
        self.assertEqual(restitue, src,
                         'le marquage a altere le texte du modele')
        print('    OK C-5 marque retiree -> texte identique a l original')

    # ── LA PHRASE DE PORTEE SE MESURE COMME UN CHIFFRE ────────────────────
    def test_la_marque_n_affirme_pas_que_le_chiffre_est_faux(self):
        # ⚠️⚠️ LE VERROU NE SAIT PAS SI LE NOMBRE EST JUSTE. Quatre des cinq
        # montants orphelins du run reel etaient EXACTS : une marque
        # accusatrice aurait denonce quatre calculs justes.
        p = RAPP_MOD.PORTEE_MARQUE_PRESCRIPTION
        for interdit in ('est faux', 'inventé', 'erroné', 'incorrect'):
            self.assertNotIn(interdit, p.lower(),
                             f'la marque accuse : {interdit!r}')
        self.assertIn('jamais sur la justesse', p)
        self.assertIn("n'a pas été vérifié", p)
        print('    OK C-6 la marque dit ce qu elle ignore, sans accuser')

    def test_les_deux_affirmations_de_la_marque_sont_vraies(self):
        # ⚠️⚠️ CE QUI LIMITE EST SUR, CE QUI AFFIRME EST UNE DETTE. La marque
        # affirme DEUX choses : que la phrase RECOMMANDE, et qu'un nombre NE
        # FIGURE PAS dans le dossier. Les deux se mesurent.
        src = 'Il est recommandé de provisionner 999 999 999 €.'
        marque, _ = self._mq(src)
        phrase = marque.split(RAPP_MOD.PORTEE_MARQUE_PRESCRIPTION)[0]
        self.assertTrue(RAPP_MOD._PRESCRIPTION_ACTUARIELLE.search(phrase),
                        "la marque dit RECOMMANDATION sur une phrase qui ne "
                        "prescrit pas")
        self.assertTrue(
            RAPP_MOD.orphelins_narration(phrase, self.CHARGE),
            "la marque dit NON SOURCE sur une phrase entierement sourcee")
        print('    OK C-7 les deux affirmations de la marque sont mesurees')

    # ── elle atteint les documents signes et l audit ──────────────────────
    def test_la_marque_atteint_le_html_et_l_audit(self):
        r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False)
        marque, n = self._mq(
            'Il est recommandé de provisionner 999 999 999 €.')
        self.assertEqual(n, 1)
        html = RAPP_MOD.export_html(
            r['n1'], r['n2'], r['n3'], r['n4'], commentaire=r['commentaire'],
            narration=marque, source_narration='claude_api')
        self.assertIn('RECOMMANDATION CHIFFRÉE NON SOURCÉE', html,
                      'le document signe ne porte pas la marque')
        ctrl = RAPP_MOD.controle_narration(marque, 'claude_api', self.CHARGE)
        self.assertEqual(ctrl['n_prescriptions_marquees'], 1,
                         'l audit ne compte pas les prescriptions marquees')
        print('    OK C-8 la marque atteint le HTML et l audit trail')

    def test_les_identifiants_de_recommandation_ne_sont_pas_des_nombres(self):
        # ⚠️ SANS CETTE EXEMPTION, << R6 -- Documenter ... doivent etre
        # justifies >> etait marquee a tort : le `6` de `R6` ressortait
        # orphelin. Mesure : 8 occurrences dans la narration reelle, TOUTES
        # des identifiants.
        self.assertEqual(RAPP_MOD.nombres_publies('R6 — Documenter'), [])
        self.assertEqual(RAPP_MOD.nombres_publies('des points R1 à R4'), [])
        # ⚠️ Et la valeur qui SUIT l etiquette reste visible.
        self.assertIn('0,85', RAPP_MOD.nombres_publies('R2 = 0,85'))
        print('    OK C-9 R1..R9 exemptes, la valeur qui suit reste visible')


class T_Le_Document_N_Affirme_Une_Origine_Qu_Une_Fois(unittest.TestCase):
    """⚠️ UN DOCUMENT NE PEUT PAS DIRE DEUX CHOSES DE SA PROPRE ORIGINE.

    Le pied de la section 7 ecrivait << Narration generee par ActuarIA
    Intelligence >> SANS CONDITION, pendant que le badge du meme bloc disait
    << Mode standard >>. Sur `templates` -- le chemin par defaut, sans cle
    API -- la seconde etait juste et la premiere fausse.
    """

    @classmethod
    def setUpClass(cls):
        cls.r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False)

    def _html(self, source):
        return RAPP_MOD.export_html(
            self.r['n1'], self.r['n2'], self.r['n3'], self.r['n4'],
            commentaire=self.r['commentaire'],
            narration='§1 — CONTEXTE\n\nUn texte.', source_narration=source)

    def test_aucune_generation_par_modele_affirmee_hors_du_chemin_llm(self):
        # ⚠️ LE CAS COURANT, ET C'EST CE QUI REND LE DEFAUT GRAVE : sans cle
        # API, TOUS les rapports produits portaient cette phrase.
        for source in ('templates', 'jugement_degrade', 'aucune'):
            html = self._html(source)
            self.assertNotIn('Narration générée par ActuarIA Intelligence',
                             html,
                             f'{source!r} affirme une generation par modele')
        print('    OK PIED-1 aucune origine LLM affirmee hors du chemin LLM')

    def test_l_origine_est_dite_par_le_badge_et_par_lui_seul(self):
        # ⚠️ SOURCE UNIQUE : l'origine vient de `_LIBELLE_SOURCE`, jamais d'une
        # phrase ecrite en dur ailleurs dans le meme bloc.
        html = self._html('templates')
        self.assertIn(RAPP_MOD._LIBELLE_SOURCE['templates'], html,
                      'le badge d origine a disparu du document')
        print('    OK PIED-2 l origine vient du badge, source unique')

    def test_le_pied_identifie_toujours_le_producteur_et_la_date(self):
        # ⚠️ ON RETIRE UNE AFFIRMATION FAUSSE, PAS L'IDENTIFICATION. L'agent
        # et la date ne dependent d'AUCUNE source : ils restent.
        html = self._html('templates')
        self.assertIn('Agent A7 Ibrahim v5.0', html,
                      'le rapport ne dit plus quel agent l a produit')
        print('    OK PIED-3 le producteur et la date restent publies')

    def test_le_chemin_llm_dit_bien_son_origine(self):
        # ⚠️ CONTRE-EPREUVE : le correctif ne doit pas rendre le document MUET
        # sur le chemin ou l'affirmation etait JUSTE.
        html = self._html('claude_api')
        self.assertIn(RAPP_MOD._LIBELLE_SOURCE['claude_api'], html,
                      'le chemin LLM ne declare plus son origine')
        print('    OK PIED-4 sur claude_api, l origine reste declaree')


class T_Une_Reference_Hors_Liste_Est_Signalee(unittest.TestCase):
    """⚠️⚠️ AUTORISEE PAR L'ENTITE, PAS VERIFIEE AU TEXTE.

    Le depot ne detient pas le Reglement delegue. Ce controle verifie une
    APPARTENANCE A UNE LISTE DECLAREE, jamais une exactitude — et c'est la
    meme limite que le verrou sur les nombres.
    """

    def test_un_article_hors_liste_est_releve(self):
        # ⚠️ MESURE : le modele a cite << article 260 >> DEUX FOIS dans la
        # premiere narration reelle, sans qu'on le lui demande.
        self.assertEqual(
            RAPP_MOD.references_hors_liste("conformément à l'article 260 du "
                                           "Règlement Délégué"), ['260'])
        print('    OK REF-1 un article hors liste est releve')

    def test_les_articles_autorises_ne_sont_jamais_releves(self):
        for txt in ('Art. 77 S2', "l'article 115", 'articles 77 et 115',
                    'Art. 116', 'article 117(2)'):
            self.assertEqual(RAPP_MOD.references_hors_liste(txt), [],
                             f'{txt!r} signale a tort')
        print('    OK REF-2 les articles autorises restent muets')

    def test_l_enumeration_est_lue_en_entier(self):
        # ⚠️ << articles 77 et 260 >> : le SECOND doit ressortir. Un motif qui
        # ne lirait que le premier laisserait passer la moitie des citations.
        self.assertEqual(
            RAPP_MOD.references_hors_liste('articles 77 et 260'), ['260'])
        print('    OK REF-3 le second membre d une enumeration est lu')

    # ── la marque, et ce qu elle refuse de dire ───────────────────────────
    def test_la_marque_dit_hors_liste_et_jamais_faux(self):
        # ⚠️⚠️ LE DEPOT NE PEUT PAS SAVOIR SI L'ARTICLE EST EXACT. Dire
        # << faux >> serait affirmer plus que ce qu'on mesure.
        p = RAPP_MOD.PORTEE_MARQUE_REFERENCE
        for interdit in ('est faux', 'inexact', 'inventé', 'erroné'):
            self.assertNotIn(interdit, p.lower(), f'la marque accuse : {interdit}')
        self.assertIn('AUTORISATION, pas une vérification', p)
        self.assertIn("n'a pas été vérifiée", p)
        print('    OK REF-4 la marque dit hors liste, jamais faux')

    def test_le_marquage_n_enleve_rien(self):
        src = ("Le BE atteint 14 830 899 €. Conformément à l'article 260, "
               "un inventaire est requis.\nAutre ligne.")
        marque, n = RAPP_MOD.marquer_references_hors_liste(src)
        self.assertEqual(n, 1)
        self.assertEqual(
            marque.replace(' ' + RAPP_MOD.PORTEE_MARQUE_REFERENCE, ''), src,
            'le marquage a altere le texte du modele')
        print('    OK REF-5 marque retiree -> texte identique a l original')

    def test_une_narration_conforme_n_est_pas_marquee(self):
        src = "Le BE relève de l'article 77, le SCR de l'article 115."
        marque, n = RAPP_MOD.marquer_references_hors_liste(src)
        self.assertEqual((marque, n), (src, 0))
        print('    OK REF-6 une narration conforme reste intacte')

    # ── la liste part AVEC le dossier ─────────────────────────────────────
    def test_la_liste_autorisee_est_transmise_au_modele(self):
        # ⚠️⚠️ LA MOITIE QUI MANQUAIT : le prompt PRESCRIVAIT cinq references
        # sans en transmettre AUCUNE. Mesure : 9 citees, 0 dans la charge.
        r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False)
        ctx = RAPP_MOD._construire_contexte(
            r['n2'], r['n3'], r['n4'], 'Automobile', '31/12/2025')
        self.assertIn('RÉFÉRENCES AUTORISÉES', ctx)
        for a in ('77', '115', '116', '117'):
            self.assertIn('Art. ' + a, ctx, f'Art. {a} non transmis')
        # ⚠️ ET LE CONTEXTE DIT SA PROPRE LIMITE au modele.
        self.assertIn('NON vérifiée au texte', ctx,
                      'la charge presente la liste comme une garantie')
        self.assertNotIn('Art. 260', ctx)
        print('    OK REF-7 la liste part avec le dossier, limite comprise')

    def test_l_audit_porte_les_references_hors_liste(self):
        ctrl = RAPP_MOD.controle_narration(
            "selon l'article 260", 'claude_api', 'charge')
        self.assertEqual(ctrl['references_hors_liste'], ['260'])
        print('    OK REF-8 l audit trail porte les numeros hors liste')

    def test_la_liste_est_declaree_autorisee_et_non_verifiee(self):
        # ⚠️ LA MENTION VIT DANS LE MODULE, a l'endroit exact ou quelqu'un
        # pourrait prendre une AUTORISATION pour une PREUVE.
        src = inspect.getsource(RAPP_MOD)
        self.assertIn('AUTORISEES PAR L\'ENTITE -- PAS VERIFIEES AU TEXTE',
                      src)
        self.assertNotIn('260', RAPP_MOD.ARTICLES_AUTORISES)
        print('    OK REF-9 la liste se declare autorisee, pas verifiee')


class T_F3b_Aucun_Seuil_Generique_Sur_Une_Hypothese_Non_Testee(
        unittest.TestCase):
    """⚠️ `seuil_UTILISE` EST AU PASSE : sur un chemin non testable, aucun
    seuil n'a servi, et l'Excel en affichait un.

    Mesure : le repli `0,50` s'appliquait a 6 LoB sur 15 qui en ont un
    autre — Cat-Nat 0,40, MRH et Credit/Caution 0,45, RC medicale,
    Construction et Accidents corporels 0,55.
    """

    PETIT = np.array([[100., 180., 220.],
                      [120., 200., np.nan],
                      [130., np.nan, np.nan]])

    def _h1(self, lob):
        from direction_non_vie.provisionnement.a7_provisionnement.n2_hypotheses import (
            HypothesesValidator,
        )
        return HypothesesValidator()._tester_h1_independance(
            self.PETIT, [], [], get_lob_config(lob))

    def test_le_chemin_non_testable_ne_publie_aucun_seuil(self):
        # ⚠️ LE FAIT DE DEPART : la cle est ABSENTE, pas nulle. C'est
        # delibere cote `_h1`, et c'est au lecteur de la cle de le respecter.
        for lob in ('rc_medicale', 'catastrophes_naturelles', 'mrh'):
            r = self._h1(lob)
            self.assertEqual(str(r.get('statut')), 'NON TESTABLE')
            self.assertNotIn('seuil_utilise', r,
                             f'{lob} publie un seuil sans avoir teste')
        print('    OK F3b-1 aucun seuil publie quand rien n a ete teste')

    def test_l_excel_n_affiche_plus_le_seuil_generique(self):
        # ⚠️⚠️ ON LIT LE VRAI CLASSEUR, PAS LA LOGIQUE REPRODUITE. Une
        # premiere version de ce test recalculait `'—' if None else ...`
        # dans le test lui-meme : il aurait passe au vert avec `n5_excel`
        # inchange. UN CONTROLE QUI ATTESTE SANS SURVEILLER — le motif que
        # ce chantier ferme, commis dans son propre filet.
        from openpyxl import load_workbook
        r = AgentA7Provisionnement(verbose=False).run(
            source=self.PETIT, mode_declare='cumule', n_sim_bootstrap=30,
            seed=42, generer_graphiques=False, lob='rc_medicale')
        self.assertTrue(r.get('success'), r.get('erreur'))
        wb = load_workbook(io.BytesIO(r['excel_bytes']))
        lignes = [[str(c) for c in row if c is not None]
                  for ws in wb for row in ws.iter_rows(values_only=True)]
        h1 = [x for x in lignes if x and x[0].startswith('H1 — Indépendance')]
        self.assertTrue(h1, 'la ligne H1 du tableau est introuvable')
        for ligne in h1:
            self.assertNotIn('0.50', ligne,
                             'le seuil generique est publie sur un test '
                             'qui n a pas eu lieu')
            self.assertIn('—', ligne, 'le seuil n est pas declare absent')
        print(f'    OK F3b-2 classeur reel, ligne H1 : {h1[0]}')

    def test_un_seuil_reellement_utilise_reste_publie(self):
        # ⚠️ CONTRE-EPREUVE : sur un triangle testable, le seuil de la BRANCHE
        # doit sortir — et il doit differer d'une branche a l'autre, sinon le
        # correctif aurait simplement tout rendu muet.
        r = AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            primes=_exposition(GENINS), n_sim_bootstrap=60, seed=42,
            generer_graphiques=False, lob='rc_medicale')
        h = r['n2'].get('h1_independance', {})
        if str(h.get('statut')) != 'NON TESTABLE':
            self.assertAlmostEqual(
                h.get('seuil_utilise'),
                get_lob_config('rc_medicale')['h1_seuil_corr'],
                msg='le seuil publie n est pas celui de la branche')
            print(f"    OK F3b-3 seuil publie = celui de la branche "
                  f"({h.get('seuil_utilise')})")
        else:
            self.skipTest('GenIns non testable pour H1 sur cette branche')

    def test_les_seuils_h1_different_reellement_entre_branches(self):
        # ⚠️ SANS CETTE MESURE, LE LOT N'AURAIT PAS D'OBJET : si toutes les
        # branches valaient 0,50, le repli aurait ete juste.
        vus = {get_lob_config(lob)['h1_seuil_corr'] for lob in LOB_CONFIG}
        self.assertGreater(len(vus), 1, 'toutes les branches au meme seuil')
        self.assertIn(0.55, vus)
        self.assertIn(0.40, vus)
        print(f'    OK F3b-4 {len(vus)} seuils H1 distincts : {sorted(vus)}')


class T_La_Portee_De_L_Archive_Est_EPROUVEE_Pas_Seulement_Ecrite(
        unittest.TestCase):
    """⚠️⚠️ UNE PHRASE DE PORTEE SE MESURE COMME UN CHIFFRE.

    `PORTEE_ARCHIVE` affirme deux choses. La premiere — l'empreinte prouve
    l'integrite — etait deja eprouvee par une alteration plantee. LA
    SECONDE — << la signature de l'actuaire reste hors du systeme >> — ne
    l'etait PAS : le test C3-5 verifie que le MOT << signature >> figure
    dans le texte, jamais que le fait est vrai. UN CONTROLE QUI CHERCHE LE
    MOT NE PROUVE RIEN.
    """

    def test_le_nom_publie_n_est_verifie_par_personne(self):
        # ⚠️ L'EPREUVE DU FAIT, ET ELLE EST BRUTALE : n'importe quelle chaine
        # est acceptee telle quelle. C'est precisement ce qui fait que ce
        # n'est PAS une signature -- rien n'authentifie qui a saisi quoi.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            trace_relecture,
        )
        for faux in ('Personne Inexistante', 'AAAA', '???', 'Ibrahim'):
            t = trace_relecture(faux, 'X-999')
            self.assertIn(faux, t.texte,
                          'un nom arbitraire devrait passer tel quel')
            self.assertFalse(t.alerte)
        print('    OK PORT-1 tout nom passe : ce n est pas une signature')

    def test_sans_nom_l_etat_negatif_est_actif(self):
        # ⚠️ CONTRE-EPREUVE : l'absence ne doit pas etre SILENCIEUSE, sinon
        # << non relu >> et << champ non transmis >> seraient indistincts.
        from direction_non_vie.provisionnement.a7_provisionnement.n5_rapport import (
            trace_relecture,
        )
        t = trace_relecture('', '')
        self.assertTrue(t.alerte)
        self.assertIn('non enregistrée', t.texte)
        print('    OK PORT-2 sans nom, l etat negatif est ACTIF')

    def test_aucune_cle_du_rapport_ne_se_dit_signature(self):
        # ⚠️⚠️ LE DEFAUT TROUVE : la cle s'appelait `signature_actuaire` et
        # portait un ETAT DE RELECTURE. Le nom affirmait plus que le code, a
        # cote de la phrase qui disait l'inverse.
        src = inspect.getsource(RAPP_MOD)
        self.assertNotIn("b['signature_actuaire']", src,
                         'une cle se dit signature alors que la portee dit '
                         "qu'il n'y en a pas")
        self.assertIn("b['relecture_actuaire']", src)
        print('    OK PORT-3 aucune cle ne se dit signature')

    def test_la_portee_declare_que_la_mention_n_est_pas_authentifiee(self):
        # ⚠️ CE QUI LIMITE NE PEUT QU'AFFAIBLIR LA REVENDICATION : cette
        # clause s'ajoute sans arbitrage. Elle dit ce que la portee taisait.
        self.assertIn('DECLARATIVE', PORTEE_ARCHIVE)
        self.assertIn("n'est verifie par personne", PORTEE_ARCHIVE)
        print('    OK PORT-4 la portee declare la mention non authentifiee')

    def test_la_clause_neuve_atteint_les_trois_sites_de_la_portee(self):
        # ⚠️ LA PORTEE VIT A TROIS ENDROITS depuis C3 : la constante, l'audit
        # archive, et le verdict de `verifier_archive`. Une clause ajoutee a
        # la constante SEULE laisserait deux des trois sites incomplets.
        vide = verifier_archive({})
        self.assertIn('DECLARATIVE', vide['porte'],
                      'le verdict ne porte pas la clause neuve')
        self.assertIn('DECLARATIVE', PORTEE_ARCHIVE)
        print('    OK PORT-5 la clause atteint la constante et le verdict')


class T_B1_La_Regle_De_H1_Est_Celle_Du_Code(unittest.TestCase):
    """⚠️⚠️ UNE DOCSTRING NE SE TESTE PAS — UN COMPORTEMENT, SI.

    `_tester_h1_independance` annoncait << rejetee si corr_moy > seuil ET au
    moins 2 colonnes significatives >>. Le code valide sur
    `corr_moy < seuil AND n_sig <= 2` : il rejette donc sur l'une OU
    l'autre, et a partir de TROIS colonnes. Mesure : 8 combinaisons sur 20
    divergent.

    Ce test verrouille la regle REELLE. La prose peut deriver a nouveau ;
    lui ne le peut pas.
    """

    SEUIL = 0.50

    @staticmethod
    def _valide(corr_moy, n_sig, seuil):
        """La regle du code, recopiee depuis `n2_hypotheses` ligne 326."""
        return corr_moy < seuil and n_sig <= 2

    def test_le_rejet_se_declenche_sur_l_une_OU_l_autre_condition(self):
        # ⚠️ LE COEUR DE L'ECART : la docstring disait ET, le code fait OU.
        self.assertFalse(self._valide(0.60, 0, self.SEUIL),
                         'corr seule au-dessus du seuil : le code rejette')
        self.assertFalse(self._valide(0.30, 3, self.SEUIL),
                         'n_sig seul a 3 : le code rejette')
        self.assertTrue(self._valide(0.30, 2, self.SEUIL))
        print('    OK B1-1 le rejet est un OU, pas un ET')

    def test_le_seuil_de_colonnes_significatives_est_trois_pas_deux(self):
        # ⚠️ SECOND ECART, PLUS DISCRET : << au moins 2 >> dans la prose,
        # `n_sig <= 2` dans le code -- donc le rejet commence a TROIS.
        self.assertTrue(self._valide(0.30, 2, self.SEUIL),
                        'deux colonnes significatives ne rejettent PAS')
        self.assertFalse(self._valide(0.30, 3, self.SEUIL))
        print('    OK B1-2 le rejet commence a 3 colonnes, pas a 2')

    def test_la_docstring_decrit_desormais_la_regle_du_code(self):
        # ⚠️ ON VERIFIE QUE LA PROSE A CEDE, pas qu'elle contient un mot.
        from direction_non_vie.provisionnement.a7_provisionnement.n2_hypotheses import (
            HypothesesValidator,
        )
        doc = HypothesesValidator._tester_h1_independance.__doc__ or ''
        self.assertIn('corr_moy >= seuil  OU  n_sig >= 3', doc,
                      'la docstring ne publie pas la regle du code')
        self.assertIn('SIGNALÉ, NON TRANCHÉ', doc,
                      'la question actuarielle ouverte n est pas declaree')
        print('    OK B1-3 la docstring publie la regle reellement appliquee')

    def test_le_code_n_a_pas_bouge(self):
        # ⚠️⚠️ AUCUN EURO DEPLACE : c'est la PROSE qui a cede. Si le code
        # avait change, des verdicts publies auraient bouge.
        from direction_non_vie.provisionnement.a7_provisionnement.n2_hypotheses import (
            HypothesesValidator,
        )
        src = inspect.getsource(HypothesesValidator._tester_h1_independance)
        self.assertIn('ok    = corr_moy < seuil and n_sig <= 2', src,
                      'la regle du code a ete modifiee : des verdicts bougent')
        print('    OK B1-4 la regle du code est intacte, seule la prose cede')


class T_La_Date_D_Arrete_Est_Normalisee_A_La_Frontiere(unittest.TestCase):
    """⚠️⚠️ LE MEME PARAMETRE, DEUX CONSOMMATEURS, DEUX EXIGENCES OPPOSEES.

    `n5_commentaire` lisait `date_arrete` par `core.arrete` — tolerant.
    `courbe_rfr.date_reference` n'accepte que 'AAAA-MM-JJ'. Avec
    << 30/06/2026 >>, le commentaire sortait JUSTE et le run TOMBAIT.
    """

    def _run(self, date_arrete):
        return AgentA7Provisionnement(verbose=False).run(
            source=np.array(GENINS, dtype=float), mode_declare='cumule',
            n_sim_bootstrap=20, seed=42, generer_graphiques=False,
            date_arrete=date_arrete)

    def test_le_format_francais_ne_fait_plus_tomber_le_run(self):
        # ⚠️ C'EST CE QU'UN ACTUAIRE TAPE SPONTANEMENT. Le rapport tombait
        # sur la forme la plus naturelle.
        r = self._run('30/06/2026')
        self.assertTrue(r.get('success'), r.get('erreur'))
        print('    OK DATE-1 30/06/2026 ne fait plus tomber le run')

    def test_tous_les_formats_de_la_table_sont_acceptes(self):
        # ⚠️ AUCUN COMPTE ECRIT ICI : on itere la TABLE. Un format ajoute
        # demain entre dans ce test tout seul.
        from core.arrete import FORMATS
        d = date(2026, 6, 30)
        # ⚠️ PAS `libelle` COMME VARIABLE : ce nom est deja importe de
        # `methodes_be` dans ce fichier, et l'ombrer aurait rendu muet tout
        # test ecrit APRES celui-ci qui l'utiliserait.
        for motif, forme in FORMATS:
            texte = d.strftime(motif)
            with self.subTest(format=forme):
                r = self._run(texte)
                self.assertTrue(r.get('success'),
                                f'{forme} ({texte}) fait tomber le run')
        print(f'    OK DATE-2 les {len(FORMATS)} formats de la table passent')

    def test_le_verdict_ne_depend_pas_de_la_forme_ecrite(self):
        # ⚠️⚠️ L'INVARIANT DU LOT : normaliser ne doit RIEN deplacer. La meme
        # date ecrite de deux facons doit rendre le MEME verdict.
        iso = self._run('2026-06-30')['n4'].get('peremption_courbe', {})
        fra = self._run('30/06/2026')['n4'].get('peremption_courbe', {})
        self.assertEqual(iso.get('statut'), fra.get('statut'),
                         'la forme ecrite change le verdict')
        self.assertEqual(iso.get('statut'), 'ROUGE',
                         'le controle d anachronisme ne s exerce plus')
        print(f"    OK DATE-3 meme verdict quelle que soit la forme "
              f"({iso.get('statut')})")

    def test_une_date_indechiffrable_est_refusee_en_nommant_les_formats(self):
        # ⚠️ ELLE TOMBAIT DEJA — mais SIX appels plus bas, sur un ValueError
        # de `strptime` qui ne disait ni quel champ ni quelles formes.
        # ⚠️ ET ELLE NE LEVE PAS JUSQU'A L'APPELANT : `run` a son except
        # global et rend `success=False`. Ma premiere version de ce test
        # attendait une exception — le comportement reel est MEILLEUR, et
        # c'est lui qu'il faut verrouiller, pas celui que j'imaginais.
        r = self._run('Q2 2026')
        self.assertFalse(r.get('success'))
        msg = str(r.get('erreur', ''))
        self.assertIn('JJ/MM/AAAA', msg, 'le refus ne nomme pas les formats')
        self.assertIn('Q2 2026', msg, 'le refus ne cite pas la valeur recue')
        print('    OK DATE-4 date illisible : refus tot, formats nommes')

    def test_sans_date_le_controle_juge_a_la_date_du_jour(self):
        # ⚠️ CONTRE-EPREUVE : le lot ne doit pas rendre le controle
        # obligatoire. Sans arrete fourni, le comportement d'avant tient.
        r = self._run(None)
        self.assertTrue(r.get('success'))
        self.assertEqual(
            r['n4'].get('peremption_courbe', {}).get('statut'), 'VERT')
        print('    OK DATE-5 sans date, le comportement d avant est intact')

    def test_le_compte_de_formats_n_est_ecrit_nulle_part(self):
        # ⚠️ LA DOCSTRING DE `arrete.lire` ANNONCAIT << les quatre formats >>
        # quand la table en portait SEPT. Un nombre recopie derive.
        from core import arrete as _a
        self.assertNotIn('quatre formats', inspect.getsource(_a),
                         'un compte de formats est recopie dans la prose')
        print('    OK DATE-6 aucun compte de formats recopie dans la prose')


if __name__ == '__main__':
    unittest.main(verbosity=1)
