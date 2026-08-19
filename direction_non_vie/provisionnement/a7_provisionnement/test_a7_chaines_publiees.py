# -*- coding: utf-8 -*-
"""
=============================================================================
 A7 Ibrahim — verrou GÉNÉRIQUE sur les chaînes publiées (lot C1)
=============================================================================

 ⚠️ CE VERROU EST ÉCRIT EN PREMIER DU LOT C, ET C'EST DÉLIBÉRÉ.

 Le même défaut s'est produit SIX fois dans ce dépôt : un séparateur de
 milliers `.replace(',', ' ')` appliqué à la PHRASE entière au lieu du seul
 nombre. Il transforme « p < 0,01 » en « p < 0 01 » et « seule, soit » en
 « seule  soit ». Cinq fois dans la série B10, une sixième au lot A2 — alors
 que je venais de relire la note qui le décrivait.

 Un test par occurrence ne l'arrêtera jamais : il faut un verrou qui inspecte
 TOUT ce que l'agent publie. Le lot C va produire des centaines de lignes de
 texte ; écrit maintenant, ce verrou les protège à mesure qu'elles s'écrivent.
 Écrit à la fin, il n'aurait fait que constater.

 DEUX RÈGLES, ET LEUR PORTÉE EST DIFFÉRENTE — c'est la calibration qui l'a
 imposée, pas une intuition :

  · LE NOMBRE, partout, sans exception. Dans un nombre à séparateur d'espace,
    tous les groupes APRÈS le premier font exactement trois chiffres.
    « 1 564 926 » est légitime ; « 0 01 » (deux chiffres) et « 0 0294 »
    (quatre) sont le défaut. Zéro faux positif mesuré.

  · LA PONCTUATION, sur les lignes de PROSE seulement. Un double espace entre
    deux mots trahit une virgule mangée. Mais `n4['jugement']` est un TABLEAU
    à colonnes alignées où l'espacement est voulu : une ligne qui contient un
    alignement de trois espaces ou plus n'est pas une phrase. Sans cette
    distinction, le verrou criait 12 fois sur des alignements légitimes.

 CALIBRATION MESURÉE AVANT D'ÊTRE FIGÉE : 4 défauts sur 4 détectés, 3 cas
 légitimes sur 3 silencieux, et 0 signalement sur 1 619 chaînes réellement
 publiées par trois scénarios.
=============================================================================
"""

import inspect
import io
import re
import unittest

import numpy as np

from direction_non_vie.provisionnement.a7_provisionnement import (
    n4_best_estimate as N4)
from direction_non_vie.provisionnement.a7_provisionnement import (
    n5_commentaire as COMM)
from direction_non_vie.provisionnement.a7_provisionnement.n4_best_estimate import (
    MSG_ASSIETTE_SCR)
from direction_non_vie.provisionnement.a7_provisionnement.n5_commentaire import (
    _s6_incertitude, _s8_recommandations, generer_commentaire)
from direction_non_vie.provisionnement.a7_provisionnement.agent import (
    AgentA7Provisionnement)
from direction_non_vie.provisionnement.a7_provisionnement.test_a7_ibrahim import (
    GENINS, RAA, _TRI_RECOURS)
from direction_non_vie.provisionnement.a7_provisionnement.test_a7_graphiques import (
    kaleido_declare, rendeur_substitue)

#: Un nombre à séparateur d'espace : « 1 564 926 ».
#:
#: ⚠️ LE CHIFFRE NE DOIT PAS ÊTRE PRÉCÉDÉ D'UNE LETTRE, ET C'EST UN FAUX
#: POSITIF MESURÉ QUI L'A IMPOSÉ. Sans cette garde, « Q1 2025 » déclenchait le
#: verrou : la regex attrapait le « 1 » de « Q1 » et voyait un groupe de quatre
#: chiffres là où il n'y a qu'un trimestre suivi d'une année. L'angle mort
#: était GÉNÉRAL — « S2 2025 » tombait pareil, et « S2 » désigne Solvabilité 2
#: dans tout ce dépôt. Il a fallu qu'une chaîne de cette forme atteigne le
#: `n4` publié pour qu'il se révèle.
#: Un chiffre précédé d'une lettre appartient à un jeton alphanumérique — Q1,
#: S2, P99 — et n'ouvre jamais un montant formaté. Les quatre défauts réels
#: restent détectés : dans « p < 0 01 », « 0 0294 » et « 12 34 € », le premier
#: groupe est précédé d'une espace, jamais d'une lettre.
_NOMBRE_ESPACE = re.compile(r'(?<![A-Za-zÀ-ÿ])\d+(?: \d+)+')
#: Double espace entre deux caractères de mot — une virgule a sauté.
_DOUBLE_ESPACE = re.compile(r'(?<=[^\W\d_])  +(?=[^\W\d_])')
#: Trois espaces ou plus : la ligne aligne des colonnes, ce n'est pas de la prose.
_ALIGNEMENT = re.compile(r'   ')


def defauts_de_separateur(texte):
    """Rend les (genre, extrait) suspects d'une chaîne. Vide = propre."""
    trouves = []
    for m in _NOMBRE_ESPACE.finditer(texte):
        groupes = m.group(0).split(' ')
        mauvais = [g for g in groupes[1:] if len(g) != 3]
        if mauvais:
            debut, fin = m.span()
            trouves.append(
                ('groupe de %s chiffres au lieu de 3'
                 % ','.join(str(len(g)) for g in mauvais),
                 texte[max(0, debut - 30):fin + 18]))
    for ligne in texte.split('\n'):
        if _ALIGNEMENT.search(ligne):
            continue
        for m in _DOUBLE_ESPACE.finditer(ligne):
            d = m.start()
            trouves.append(('virgule remplacée par un espace',
                            ligne[max(0, d - 30):d + 30]))
    return trouves


def _chaines(obj, chemin='', acc=None):
    """Toutes les chaînes d'une structure, avec le chemin qui y mène."""
    acc = acc if acc is not None else []
    if isinstance(obj, str):
        acc.append((chemin, obj))
    elif isinstance(obj, dict):
        for cle, val in obj.items():
            _chaines(val, f'{chemin}.{cle}', acc)
    elif isinstance(obj, (list, tuple)):
        for i, val in enumerate(obj):
            _chaines(val, f'{chemin}[{i}]', acc)
    return acc


def _run_complet(triangle, **kw):
    """Un run avec TOUS les livrables — le verrou de vocabulaire les inspecte.

    ⚠️ CE VERROU PORTE SUR LES MOTS PUBLIÉS, pas sur les pixels. Il faisait
    rasteriser 26 figures — 160 s mesurées — pour relire du vocabulaire. Le
    rendeur substitué met une image dans le document sans la dessiner.
    """
    src = np.asarray(triangle, dtype=float)
    with kaleido_declare(True), rendeur_substitue():
        return AgentA7Provisionnement(verbose=False).run(
            source=src, mode_declare='cumule', generer_graphiques=True,
            generer_word=True, n_sim_bootstrap=60, seed=42, **kw)


def _run(triangle, **kw):
    src = np.asarray(triangle, dtype=float)
    kw.setdefault('primes', np.full(src.shape[0],
                                    float(np.nanmean(src[:, 0])) * 8.0))
    return AgentA7Provisionnement(verbose=False).run(
        source=src, mode_declare='cumule', generer_graphiques=False,
        generer_word=False, n_sim_bootstrap=60, seed=42, **kw)


# =============================================================================
#  T1 — LE DÉTECTEUR DISCRIMINE, DANS LES DEUX SENS
# =============================================================================

class T1_Le_Detecteur_Discrimine(unittest.TestCase):
    """Un verrou qui ne peut pas se déclencher ne protège rien."""

    #: Les défauts RÉELS, repris tels qu'ils sont sortis du code.
    _DEFAUTS = (
        "L'hypothèse est non validée sur la colonne 0 (p < 0 01) que ces "
        "années doivent traverser",
        "Elles sont portées par Chain Ladder seule  soit 4 625 811 € du total",
        "réserve portée de 17 469 539 € à 59 400 660 €, écart de 0 0294",
        "provision de 12 34 € sur l'exercice",
    )

    #: Ce qui est LÉGITIME et ne doit jamais crier.
    _SAINS = (
        "réserve 1 564 926 €, p < 0,01, soit 26,3 % du Best Estimate",
        # ⚠️ LES DEUX FORMES QUI ONT RÉVÉLÉ L'ANGLE MORT (lot « courbe »).
        "EIOPA RFR Term Structures — Q1 2025",
        "SCR PROVISIONS (Art. 115 S2 2025)",
        "BE = 18 680 856 € · σ = 2 447 095 € · CV = 13,1 %",
        "  Chain Ladder      réserve=18 680 856   poids=53%",
        "H2 Stabilité      : VALIDÉE       CV=7.9%  dérive=6.8",
        "Années [5, 6, 7, 8] : couverture à justifier",
    )

    def test_il_attrape_les_defauts_reels(self):
        for s in self._DEFAUTS:
            self.assertTrue(defauts_de_separateur(s),
                            f"défaut non détecté : {s[:60]}")
        print(f"    OK C1-1 les {len(self._DEFAUTS)} défauts réels sont "
              f"détectés")

    def test_il_se_tait_sur_ce_qui_est_legitime(self):
        for s in self._SAINS:
            self.assertEqual(defauts_de_separateur(s), [],
                             f"faux positif sur : {s[:60]}")
        print(f"    OK C1-2 les {len(self._SAINS)} cas légitimes ne "
              f"déclenchent rien — dont un tableau à colonnes alignées")


# =============================================================================
#  T2 — TOUT CE QUE L'AGENT PUBLIE EST INSPECTÉ
# =============================================================================

class T2_Aucune_Chaine_Publiee_N_Est_Fautive(unittest.TestCase):

    def test_les_resultats_de_trois_scenarios(self):
        """n1, n2, n3, n4 et le commentaire, récursivement."""
        total, fautes = 0, []
        for nom, tri in (('GenIns', GENINS), ('RAA', RAA),
                         ('Recours', _TRI_RECOURS)):
            r = _run(tri)
            bloc = {'n1': r.get('n1'), 'n2': r['n2'], 'n3': r['n3'],
                    'n4': r['n4'], 'commentaire': r.get('commentaire')}
            for chemin, s in _chaines(bloc):
                total += 1
                for genre, extrait in defauts_de_separateur(s):
                    fautes.append(f'{nom}{chemin} — {genre} : {extrait!r}')
        self.assertEqual(fautes, [], '\n'.join(fautes[:10]))
        print(f"    OK C1-3 {total} chaînes publiées sur 3 scénarios, "
              f"aucune fautive")

    def test_les_cellules_de_l_excel(self):
        """Le format que le défaut atteindrait sans qu'on le voie."""
        import openpyxl
        r = _run(GENINS)
        octets = r.get('excel_bytes') or b''
        if not octets:
            self.skipTest('openpyxl absent')
        wb = openpyxl.load_workbook(io.BytesIO(octets))
        total, fautes = 0, []
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for c in row:
                    if not isinstance(c.value, str):
                        continue
                    total += 1
                    for genre, extrait in defauts_de_separateur(c.value):
                        fautes.append(f'{ws.title}!{c.coordinate} — {genre} : '
                                      f'{extrait!r}')
        self.assertEqual(fautes, [], '\n'.join(fautes[:10]))
        print(f"    OK C1-4 {total} cellules de texte dans l'Excel, "
              f"aucune fautive")



# =============================================================================
#  T3 — LE VERROU DE VOCABULAIRE « SCR »  (lot C3b)
# =============================================================================
#
#  ⚠️ MÊME DISPOSITIF QUE LE VERROU DE SÉPARATEUR CI-DESSUS, ET POUR LA MÊME
#  RAISON : le mot « SCR » désignait QUATRE grandeurs différentes dans les
#  livrables. Le relevé exhaustif du lot C3b a compté 110 occurrences, 70
#  formulations distinctes — et sept d'entre elles nommaient « SCR » un NIVEAU
#  de réserve. Corriger les sept sans poser de verrou, c'est attendre la
#  huitième : C3c et C3d vont ajouter des graphiques et du routage.
#
#  LA CONVENTION, EN TROIS RÈGLES :
#    1. « SCR » désigne UNE grandeur, la charge de capital de l'article 115
#       (3·σ·V). C'est une MARGE, jamais un niveau de réserve.
#    2. Un niveau de percentile se nomme par son percentile — « Réserve au
#       P99,5 » — jamais « SCR ».
#    3. Une marge issue d'un percentile se nomme comme telle, « Marge
#       P99,5 − BE » : c'est elle, et non le niveau, qui se compare au SCR.
#
#  L'UNITÉ D'ANALYSE EST ATOMIQUE, ET C'EST LA CALIBRATION QUI L'A IMPOSÉ :
#  une CELLULE d'un tableau, une étiquette de figure, une LIGNE de prose —
#  jamais une ligne de tableau entière. Le Word aligne « P90 (composé) » et
#  « SCR Provisions » dans deux cellules voisines : les aplatir ferait crier le
#  verrou sur une mise en page parfaitement correcte.
#
#  CALIBRÉ AVANT D'ÊTRE FIGÉ, DANS LES DEUX SENS : 7 défauts réels sur 7
#  détectés, 8 formulations légitimes sur 8 silencieuses, et 0 signalement sur
#  9 045 unités réellement publiées par deux configurations d'exposition.
# =============================================================================

#: Le mot, isolé.
_SCR = re.compile(r'\bSCR\b')
#: Un percentile, sous toutes ses écritures — « P99.5 », « 99,5 », « 99.5th ».
_PERCENTILE = re.compile(
    r'\bP\s?\d{2,3}([.,]\d)?\b|\bpercentile\b|\bquantile\b|\b\d{2}[.,]\d',
    re.I)
#: ⚠️⚠️ L'EXEMPTION SE SCINDE EN DEUX ÉTAGES, ET C'EST LE CŒUR D'UN LOT.
#:
#: Elle était UNE seule alternative, où la FORMULE de l'article 115 et sa
#: CITATION valaient exemption à égalité. Conséquence mesurée : la phrase
#:
#:     « Retenir 18 053 284 € pour le calcul du SCR provisions
#:       (stress test P90, formule standard Art. 115). »
#:
#: passait le contrôle — avec DEUX mots d'échappement, « formule standard »
#: et « Art. 115 ». **Citer l'article suffisait à échapper au contrôle de
#: l'article.** C'est le motif du lot avis-couleur : un contrôle qui cherche
#: un mot est battu par le mot.
#:
#: A — LA FORMULE. Elle DÉMONTRE qu'on parle de l'exigence de capital :
#: personne n'écrit « 3 × σ » pour désigner un percentile. Exemption forte.
_FORMULE_SCR = re.compile(
    r'3\s*[×x]\s*σ|3\s*[×x]\s*\d|SCR\s*/\s*BE|ratio SCR', re.IGNORECASE)

#: B — LA CITATION. Elle n'ATTESTE rien : n'importe quelle phrase peut la
#: porter, y compris celle qui contredit l'article cité. Elle n'exempte plus
#: qu'à condition que la phrase ne PRESCRIVE rien.
#: ⚠️ `σ de réserve` A REJOINT CE TIER, ET PAS L'AUTRE — la mesure a tranche.
#: Le bloc LoB Marine publie « Le SCR de la ligne d'activite 6 (σ de réserve =
#: 11,0 %) est eleve » : phrase JUSTE, signalee a tort des que l'assiette l'a
#: atteinte, parce que `11,0` ressemble a un percentile. Le reflexe etait de
#: l'exempter par la FORMULE, qui exempte SANS CONDITION. Mesure :
#:
#:     « σ de réserve : retenir le P99.5 pour le calcul du SCR provisions. »
#:
#: passait alors le controle. Nommer sigma aurait suffi a prescrire n'importe
#: quoi. La regle posee au lot precedent le disait deja : `3 × σ` est une
#: OPERATION, qui demontre ; `σ de réserve` est un NOM, qui coexiste avec
#: n'importe quelle phrase. Un nom atteste aussi peu qu'une citation d'article.
_CITATION_SCR = re.compile(
    r'[Aa]rt(icle)?\.?\s*11[57]|formule standard|σ\s+de\s+réserve',
    re.IGNORECASE)

#: C — LA PRESCRIPTION : un verbe qui DIRIGE une grandeur vers le SCR.
#:
#: ⚠️ C'EST UNE LISTE DE MOTS, ET C'EST ASSUMÉ — PARCE QU'ELLE NE PEUT QUE
#: RENDRE LE CONTRÔLE PLUS SENSIBLE. Une liste qui EXEMPTE ouvre un trou (on
#: vient d'en payer un) ; une liste qui ACCUSE ne peut qu'ajouter des
#: signalements, jamais en retirer. Un mot qui manque ici laisse le contrôle
#: exactement où il était sans cette étape — il ne crée aucune cécité neuve.
_PRESCRIPTION_SCR = re.compile(
    r'\bretenir\b|\butiliser\b|\bpour le calcul d|\bchiffre critique\b'
    r'|\bmaximum des deux\b|\bse calcule sur\b', re.IGNORECASE)


def vocabulaire_scr_fautif(unite):
    """True si cette unité publiée nomme « SCR » un niveau de percentile.

    Calibré dans les DEUX SENS après la fermeture de l'exemption : les 7
    défauts du relevé restent détectés, les 8 formulations légitimes restent
    silencieuses — dont « c'est cette marge, et non le niveau, qui se compare
    au SCR de l'article 115 », qui aurait crié si la citation avait été
    retirée sans l'étage de prescription."""
    if not _SCR.search(unite) or not _PERCENTILE.search(unite):
        return False
    if _FORMULE_SCR.search(unite):
        return False
    # La citation n'exempte QUE la phrase qui ne prescrit rien.
    exempt = bool(_CITATION_SCR.search(unite)
                  and not _PRESCRIPTION_SCR.search(unite))
    return not exempt


def _unites_publiees(r):
    """Les unités ATOMIQUES d'un run — cellules, étiquettes, lignes de prose."""
    out = []
    for cle in ('n2', 'n3', 'n4'):
        for chemin, s in _chaines(r.get(cle), cle):
            out += [(chemin, ligne) for ligne in s.split('\n')]
    for ligne in (r.get('commentaire') or '').split('\n'):
        out.append(('commentaire', ligne))
    for nom, fig in (r.get('graphiques') or {}).items():
        for chemin, s in _chaines(fig.to_plotly_json(), 'figure:' + nom):
            out.append((chemin, s))
    octets = r.get('excel_bytes') or b''
    if octets:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(octets))
        for ws in wb.worksheets:
            for row in ws.iter_rows():
                for c in row:
                    if isinstance(c.value, str):
                        out.append(('excel:%s!%s' % (ws.title, c.coordinate),
                                    c.value))
    mot = r.get('word_bytes') or b''
    if mot:
        import docx
        doc = docx.Document(io.BytesIO(mot))
        for p in doc.paragraphs:
            out.append(('word', p.text))
        for tbl in doc.tables:
            for row in tbl.rows:
                for c in row.cells:       # CELLULE, jamais la ligne entière
                    out.append(('word:cellule', c.text))
    return out


class T3_Le_Mot_SCR_Ne_Nomme_Qu_Une_Grandeur(unittest.TestCase):

    #: Les défauts RÉELS, tels qu'ils sortaient du code avant le lot C3b.
    _DEFAUTS = (
        'SCR (P99.5)',                          # g7, segment n°4
        'Bootstrap P99.5 (SCR)',                # Excel, onglet Sensibilités
        'P99.5 — SCR provisions',               # Excel, onglet Bootstrap
        'P99.5 — SCR provisions (composé)',     # Excel, onglet Résultats
        'SCR provisions — extrême (99.5th)',    # Excel, colonne Lecture
        '<b>P99.5 (SCR)</b><br>25,040,191€',    # g6, annotation
        "Le P99.5 Bootstrap de 25 040 191 € constitue l'estimation "
        "stochastique du SCR provisions — comparable du P99.5 Mack.",
    )

    #: Ce qui est LÉGITIME : le SCR de l'article 115, sous ses formes réelles.
    _SAINS = (
        'SCR_prov = 3 × 11.0% × 17,571,609€ = 5,798,631€ (ratio SCR/BE = 33.0%)',
        'Ratio SCR/BE',
        'SCR PROVISIONS (Art. 115 S2)',
        'CALCUL SCR PROVISIONS (LoB unique)',
        # ⚠️ CETTE FIXTURE PORTAIT UNE PHRASE FAUSSE, DANS UNE LISTE NOMMÉE
        # « SAINS ». Elle disait « le facteur 3 correspond au quantile 99,5 %
        # d'une loi normale » — ce quantile vaut 2,5758, soit 16,5 % d'écart.
        # `_SAINS` veut dire « sain POUR LA RÈGLE SCR » ; ça se lisait « sain,
        # point ». La propriété testée est inchangée : le détecteur reste muet
        # sur l'ancienne comme sur la nouvelle. Seul l'exemple cesse d'endosser
        # une fausseté. Ne pas restaurer l'ancienne formulation.
        # ⚠️ PARENTHÉSÉE (ISC004) : dans un tuple de fixtures, une
        # concaténation implicite non parenthésée est à une virgule oubliée de
        # fusionner deux entrées en silence. Dans un fichier dont le métier est
        # de détecter des chaînes fausses, ce serait un vrai défaut.
        ("Le facteur 3 provient de la calibration EIOPA du risque de réserve, "
         "qui suppose une distribution log-normale : pour les σ retenus, le "
         "rapport entre le quantile 99,5 % et la moyenne vaut environ 3σ. "
         "Ce n'est pas le quantile d'une loi normale, qui vaut 2,576."),
        'P99.5 — Provision extrême (composé)',
        'SCR Provisions — Formule standard Art. 115 Règlement Délégué (UE) 2015/35',
        "c'est cette marge, et non le niveau, qui se compare au SCR de "
        "l'article 115. Elle est proche de celle du P99.5 Mack.",
        # ⚠️ LE CAS QUE L'ELARGISSEMENT DE L'ASSIETTE A REVELE (bloc LoB
        # Marine). σ de réserve EST le facteur de l'article 115 ; « 11,0 »
        # n'est pas un percentile. La phrase est JUSTE, et elle n'avait
        # jamais ete balayee.
        ("4. Le SCR de la ligne d'activité 6 (σ de réserve = 11,0 %) est "
         "élevé — anticiper une exigence de capital importante."),
    )

    #: ⚠️ LES DEUX CONTRE-EPREUVES DU MEME MOT, ET ELLES SONT INDISPENSABLES.
    #: Exempter `σ de réserve` par la FORMULE aurait fait passer ces deux
    #: prescriptions : nommer sigma aurait suffi a diriger n'importe quelle
    #: grandeur vers le SCR. Elles verrouillent le tier ou le motif a ete mis.
    _DEFAUTS_SIGMA_NOMME = (
        "σ de réserve : retenir le P99.5 pour le calcul du SCR provisions.",
        "σ de réserve = 11,0 % — utiliser le P90 comme assiette du SCR.",
    )

    def test_il_attrape_les_sept_defauts_du_releve(self):
        for s in self._DEFAUTS:
            self.assertTrue(vocabulaire_scr_fautif(s),
                            'défaut non détecté : %s' % s[:70])
        print('    OK C3b-1 les %d défauts du relevé exhaustif sont détectés'
              % len(self._DEFAUTS))

    def test_il_se_tait_sur_l_article_115(self):
        for s in self._SAINS:
            self.assertFalse(vocabulaire_scr_fautif(s),
                             'faux positif sur : %s' % s[:70])
        print('    OK C3b-2 les %d formulations légitimes restent silencieuses'
              % len(self._SAINS))

    def test_nommer_sigma_ne_permet_pas_de_prescrire(self):
        """⚠️ LE TROU QU'UNE EXEMPTION PAR LA FORMULE AURAIT OUVERT.

        `σ de réserve` a rejoint le tier CITATION, qui n'exempte QUE ce qui
        ne prescrit rien. Dans le tier FORMULE — inconditionnel — ces deux
        phrases seraient passées : nommer sigma aurait suffi à diriger
        n'importe quelle grandeur vers le SCR. Mesuré avant d'être écrit."""
        for s in self._DEFAUTS_SIGMA_NOMME:
            self.assertTrue(vocabulaire_scr_fautif(s),
                            f'prescription non detectee : {s[:70]}')
        print('    OK C3b-2b nommer sigma n exempte aucune prescription')

    def test_aucune_unite_publiee_ne_nomme_scr_un_percentile(self):
        """Tout ce que l'agent publie, avec et sans exposition."""
        n = np.asarray(GENINS).shape[0]
        total, fautes = 0, []
        for kw in ({'primes': np.full(n, 4e6)}, {}):
            r = _run_complet(GENINS, **kw)
            for ou, u in _unites_publiees(r):
                total += 1
                if vocabulaire_scr_fautif(u):
                    fautes.append('%s — %r' % (ou, re.sub(r'\s+', ' ', u)[:90]))
        self.assertEqual(fautes, [], '\n'.join(fautes[:10]))
        print('    OK C3b-3 %d unités publiées, aucune ne nomme « SCR » un '
              'niveau de percentile' % total)

# =============================================================================
#  T4 — AUCUNE PRESCRIPTION DE PERCENTILE POUR LE SCR
# =============================================================================
#
#  ⚠️⚠️ CE CONTRÔLE SURVEILLAIT, ET SON ASSIETTE NE COUVRAIT PAS LES FAUTES.
#  `test_aucune_unite_publiee_ne_nomme_scr_un_percentile` balaie DÉJÀ tout ce
#  que l'agent publie — n2/n3/n4, commentaire, 26 figures, Excel, Word. Son
#  balayage est complet. Son ASSIETTE ne l'est pas : UN triangle (GenIns), UNE
#  LoB (générique), deux expositions. Les QUATRE prescriptions fausses
#  vivaient toutes sur des chemins que cette assiette n'atteint pas :
#
#      · branche VERT          — GenIns sort ROUGE
#      · écart P99.5 > 15 %    — sort sur RAA, pas sur GenIns
#      · LoB rc_auto_corporels — la LoB par défaut est « générique »
#      · BOOT-H3 NON VALIDÉE   — validée sur les triangles de référence
#
#  ⚠️ C'EST « UN VERDICT NE SE PUBLIE PAS SANS SON ASSIETTE », APPLIQUÉ AU
#  CONTRÔLE LUI-MÊME. Le test imprime « N unités publiées, aucune ne nomme
#  SCR un percentile » — vrai de CES unités-là, lu comme une propriété du
#  logiciel. Ce bloc atteint les chemins, au lieu d'élargir un balayage qui
#  était déjà complet.


class T4_Aucune_Prescription_De_Percentile_Pour_Le_SCR(unittest.TestCase):
    """⚠️ LE SCR NE SE CALCULE SUR AUCUN PERCENTILE, NI SUR σ_MACK."""

    @classmethod
    def setUpClass(cls):
        cls.r_gen = _run(GENINS)
        cls.r_raa = _run(RAA)

    def _aucune_faute(self, texte, ou):
        fautes = [ligne for ligne in texte.split('\n')
                  if vocabulaire_scr_fautif(ligne)]
        self.assertEqual(fautes, [], f'{ou} : {fautes[:3]}')

    # ── SCR-1 : la phrase qui ÉCHAPPAIT au contrôle est désormais vue ────────
    def test_la_citation_de_l_article_n_exempte_plus_une_prescription(self):
        # ⚠️ LA VIOLATION PLANTÉE DU LOT, SOUS SA FORME EXACTE. Elle portait
        # DEUX mots d'échappement — « formule standard » ET « Art. 115 » — et
        # passait le contrôle. Citer l'article suffisait à échapper au
        # contrôle de l'article.
        echappait = ('2. Retenir 18 053 284 € pour le calcul du SCR '
                     'provisions (stress test P90, formule standard Art. 115).')
        self.assertTrue(vocabulaire_scr_fautif(echappait),
                        "la citation exempte encore une prescription")
        # ET LA CONTRE-ÉPREUVE, dans les deux sens : la citation exempte
        # toujours ce qui ne prescrit rien, et la formule exempte toujours.
        self.assertFalse(vocabulaire_scr_fautif(
            "c'est cette marge, et non le niveau, qui se compare au SCR de "
            "l'article 115. Elle est proche de celle du P99.5 Mack."))
        self.assertFalse(vocabulaire_scr_fautif(
            'SCR_prov = 3 × 11.0% × 17,571,609€ (ratio SCR/BE = 33.0%)'))
        print('    OK SCR-1 la citation n exempte plus une prescription')

    # ── SCR-2 : les quatre chemins, atteints un par un ──────────────────────
    def test_la_branche_verte_ne_dirige_plus_un_percentile_vers_le_scr(self):
        n4 = dict(self.r_gen['n4'], statut='VERT')
        txt = _s8_recommandations(self.r_gen['n1'], self.r_gen['n2'],
                                  self.r_gen['n3'], n4, 'generique')
        self.assertNotIn('pour le calcul du SCR provisions', txt)
        self._aucune_faute(txt, 'recommandations VERT')
        self.assertIn(MSG_ASSIETTE_SCR, txt)
        print('    OK SCR-2 la branche VERT publie le SCR, pas un percentile')

    def test_l_ecart_p995_n_appelle_plus_une_regle_de_maximum(self):
        # RAA le déclenche réellement ; GenIns non — d'où les deux.
        for nom, r in (('RAA', self.r_raa), ('GenIns', self.r_gen)):
            txt = _s6_incertitude(r['n3'], r['n4'])
            self.assertNotIn('maximum des deux', txt, nom)
            self._aucune_faute(txt, 'incertitude ' + nom)
        print('    OK SCR-3 l ecart P99.5 appelle un examen, pas une provision')

    def test_les_blocs_lob_ne_nomment_plus_un_percentile_critique(self):
        for lob in ('rc_auto_corporels', 'rc_medicale', 'generique'):
            txt = generer_commentaire(
                self.r_gen['n1'], self.r_gen['n2'], self.r_gen['n3'],
                self.r_gen['n4'], lob=lob)
            self.assertNotIn('chiffre critique pour le SCR', txt, lob)
            self._aucune_faute(txt, 'commentaire ' + lob)
        print('    OK SCR-4 les blocs LoB ne prescrivent plus de percentile')

    def test_la_recommandation_boot_h3_ne_dirige_plus_vers_sigma_mack(self):
        # ⚠️ CHEMIN NON ATTEIGNABLE SUR LES TRIANGLES DE RÉFÉRENCE (BOOT-H3 y
        # est validée) : on verrouille donc la SOURCE, pas une sortie. Un test
        # de valeur ne verrait rien tant qu'aucun triangle ne rejette BOOT-H3.
        src = inspect.getsource(N4)
        self.assertNotIn("Retenir\"\n                    f\"\n", src)
        self.assertNotIn("l'incertitude de Mack (σ) pour le SCR", src,
                         'la prescription vers sigma_Mack est revenue')
        i = src.index("ne pèse pas dans sa")
        self.assertIn('MSG_ASSIETTE_SCR', src[i:i + 400])
        print('    OK SCR-5 BOOT-H3 ne dirige plus vers sigma_Mack')

    # ── SCR-3 : la phrase est UNE, et elle n'est rédigée qu'une fois ─────────
    def test_la_phrase_du_scr_a_une_source_unique(self):
        # ⚠️ LA PROPRIÉTÉ QUI EMPÊCHE LA CORRECTION PARTIELLE DE REVENIR : les
        # sites ne rédigent plus, donc ils ne peuvent plus diverger.
        for mod in (N4, COMM):
            src = inspect.getsource(mod)
            self.assertEqual(
                src.count('"Le SCR provisions ne se calcule sur aucun'),
                1 if mod is N4 else 0,
                'la phrase est réécrite ailleurs que dans sa source')
        self.assertIn('σ(LoB)', MSG_ASSIETTE_SCR)
        self.assertFalse(vocabulaire_scr_fautif(MSG_ASSIETTE_SCR),
                         'la phrase de correction déclenche le contrôle')
        print('    OK SCR-6 une seule redaction, et elle ne se signale pas')


# =============================================================================
#  T5 — L'ASSIETTE DES CONTROLES : PAR LES CHEMINS, PAS PAR LES TRIANGLES
# =============================================================================
#
#  ⚠️⚠️ LE BALAYAGE ETAIT COMPLET, L'ASSIETTE NE L'ETAIT PAS. `_unites_publiees`
#  parcourt DEJA n2/n3/n4, le commentaire, 26 figures, l'Excel et le Word.
#  Mais on ne lui donnait a balayer qu'UN triangle, UNE LoB, deux expositions.
#  Les QUATRE prescriptions fausses du lot SCR vivaient toutes sur des chemins
#  que cette assiette n'atteignait pas :
#
#      branche VERT           GenIns sort ROUGE
#      ecart P99.5 > 15 %     sort sur RAA, pas sur GenIns
#      LoB specifique         la LoB par defaut est << generique >>
#      BOOT-H3 rejetee        validee sur les triangles de reference
#
#  ⚠️ ON ELARGIT PAR LES CHEMINS, PAS PAR LES TRIANGLES, ET C'EST MESURE.
#  `_run_complet` produit Excel + Word + 26 figures : le multiplier par N
#  quadruplerait le controle le plus lourd du depot, sur une gate deja a
#  30 minutes. Or les quatre fautes vivaient dans le COMMENTAIRE et dans N4 --
#  `_run` leger les atteint toutes. Un chemin coute un appel ; un balayage
#  coute une gate.
#
#  ⚠️ ET LA MEME ASSIETTE SERT LES DEUX DETECTEURS. Elargir pour
#  `vocabulaire_scr_fautif` fait passer gratuitement `defauts_de_separateur`
#  sur les memes chemins -- T2 les avait deja sur trois triangles, jamais sur
#  ces etats-la.


class T5_Les_Chemins_Que_L_Assiette_N_Atteignait_Pas(unittest.TestCase):
    """⚠️ UN CONTROLE NE VAUT QUE SUR CE QU'ON LUI DONNE A REGARDER."""

    @classmethod
    def setUpClass(cls):
        cls.gen = _run(GENINS)
        cls.raa = _run(RAA)

    @staticmethod
    def _unites_textuelles(n1, n2, n3, n4, commentaire):
        """Les unites ou vivaient les quatre fautes : N4 et la prose."""
        bloc = {'n2': n2, 'n3': n3, 'n4': n4, 'commentaire': commentaire}
        out = []
        for chemin, s in _chaines(bloc):
            out += [(chemin, ligne) for ligne in s.split('\n')]
        return out

    def _balayer(self, nom, unites):
        """Les DEUX detecteurs sur les memes unites, et le compte rendu."""
        fautes = []
        for chemin, u in unites:
            if vocabulaire_scr_fautif(u):
                fautes.append(f'{nom}{chemin} — SCR/percentile : {u[:90]!r}')
            for genre, extrait in defauts_de_separateur(u):
                fautes.append(f'{nom}{chemin} — {genre} : {extrait!r}')
        return fautes

    # ── CHEMIN 1 : la branche VERT ──────────────────────────────────────────
    def test_chemin_branche_verte(self):
        # ⚠️ GenIns et RAA sortent ROUGE : ce chemin n'etait JAMAIS exerce, et
        # c'est celui ou vivait << Retenir {P90} pour le calcul du SCR >>.
        r = self.gen
        n4 = dict(r['n4'], statut='VERT')
        com = generer_commentaire(r['n1'], r['n2'], r['n3'], n4)
        unites = self._unites_textuelles(r['n1'], r['n2'], r['n3'], n4, com)
        unites += [('.s8', l) for l in
                   _s8_recommandations(r['n1'], r['n2'], r['n3'], n4,
                                       'generique').split('\n')]
        self.assertTrue(unites, 'aucune unite produite')
        self.assertEqual(self._balayer('VERT', unites), [])
        print(f'    OK ASS-1 branche VERT : {len(unites)} unites, aucune '
              f'fautive')

    # ── CHEMIN 2 : l'ecart P99.5 significatif ───────────────────────────────
    def test_chemin_ecart_p995_significatif(self):
        # RAA le declenche reellement ; GenIns non. C'est ce chemin qui
        # publiait << Pour le SCR, retenir le maximum des deux >>.
        r = self.raa
        s6 = _s6_incertitude(r['n3'], r['n4'])
        unites = [('.s6', l) for l in s6.split('\n')]
        self.assertEqual(self._balayer('RAA', unites), [])
        print(f'    OK ASS-2 ecart P99.5 : {len(unites)} unites, aucune '
              f'fautive')

    # ── CHEMIN 3 : les LoB specifiques ──────────────────────────────────────
    def test_chemin_lob_specifiques(self):
        # ⚠️ LA LoB PAR DEFAUT EST << generique >>, qui rend des blocs VIDES.
        # Les blocs LoB sont le chemin le moins relu du module -- et l'un
        # d'eux disait du P99.5 qu'il etait << le chiffre critique pour le
        # SCR >>. Les sept handlers sont exerces, plus le repli.
        r = self.gen
        total, fautes = 0, []
        for lob in ('mrh', 'rc_auto_materiel', 'rc_auto_corporels',
                    'rc_generale', 'rc_medicale', 'construction',
                    'marine_aviation_transport', 'generique'):
            com = generer_commentaire(r['n1'], r['n2'], r['n3'], r['n4'],
                                      lob=lob)
            unites = [(f'.{lob}', l) for l in com.split('\n')]
            total += len(unites)
            fautes += self._balayer(lob, unites)
        self.assertEqual(fautes, [], '\n'.join(fautes[:10]))
        print(f'    OK ASS-3 8 LoB : {total} unites, aucune fautive')

    # ── CHEMIN 4 : BOOT-H3 rejetee ──────────────────────────────────────────
    def test_chemin_boot_h3_rejetee(self):
        # ⚠️ VALIDEE SUR TOUS LES TRIANGLES DE REFERENCE : ce chemin n'est
        # atteignable qu'en forcant la porte d'hypothese. C'est lui qui
        # publiait << Retenir l'incertitude de Mack pour le SCR >>.
        from direction_non_vie.provisionnement.a7_provisionnement.n4_best_estimate import (
            BestEstimateS2,
        )
        r = self.gen
        n2 = dict(r['n2'],
                  bootstrap_hyp={**r['n2'].get('bootstrap_hyp', {}),
                                 'percentiles_publiables': False})
        n4 = BestEstimateS2().calculer(n2, r['n3'],
                                       np.asarray(GENINS, dtype=float))
        unites = self._unites_textuelles(r['n1'], n2, r['n3'], n4, '')
        self.assertTrue(any('recommandations' in c for c, _ in unites),
                        'les recommandations ne sont pas produites')
        self.assertEqual(self._balayer('BOOTH3', unites), [])
        print(f'    OK ASS-4 BOOT-H3 rejetee : {len(unites)} unites, aucune '
              f'fautive')

    # ── ET LE CONTROLE DE L'ASSIETTE ELLE-MEME ──────────────────────────────
    def test_les_quatre_chemins_sont_bien_DISTINCTS_du_balayage_existant(self):
        # ⚠️ SANS CE TEST, RIEN NE PROUVE QUE L'ELARGISSEMENT ELARGIT. Un
        # chemin qui reproduirait l'etat deja balaye ne couvrirait rien de
        # neuf, et le lot n'aurait fait qu'allonger la gate.
        self.assertEqual(self.gen['n4']['statut'], 'ROUGE',
                         'GenIns sort VERT : le chemin 1 doublonne')
        self.assertTrue(self.raa['n2']['bootstrap_hyp']
                        .get('percentiles_publiables', True),
                        'BOOT-H3 deja rejetee : le chemin 4 doublonne')
        print('    OK ASS-5 les quatre chemins sont bien hors du balayage')


if __name__ == '__main__':
    unittest.main(verbosity=2)
