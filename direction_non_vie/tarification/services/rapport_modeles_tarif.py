"""
ActuarIA — Rapport Tarification Non-Vie
Direction Non-Vie · Équipe Tarification · BLOC 5

4 formats : HTML · Word (.docx) · PDF (weasyprint) · Excel (via tarif_excel.py)
Narration : Claude API → commentaire agent → données seules (3 niveaux)

Structure du rapport : voir `CHAPITRES`, plus bas. ⚠️ CET EN-TÊTE PORTAIT UNE
TROISIÈME NUMÉROTATION — un plan en huit « §N » que le code n'a jamais produit,
différent de celui de l'HTML comme de celui du Word. Les titres vivent
désormais à un seul endroit, et les deux formats les y lisent.

Auteur   : ActuarIA v1.0
Version  : 1.0.0
"""

from __future__ import annotations
import base64, io, logging, re
from core.conformite_reglementaire import (
    avertissement_walk_forward, synthese_exclusions, synthese_alertes_experience,
    synthese_modele_dl,
)
from core.qualite_donnees import synthese_qualite_donnees
from core.plan_tarifaire import synthese_colonnes_plan_manquantes
from core.mapping_client import synthese_mapping
from datetime import datetime
from typing import Dict, List, NamedTuple, Optional, Tuple

from core import format_fr as F
from core import narration as _md
from core import frontiere_llm


logger = logging.getLogger('actuaria.tarif.rapport')

# ── Palette ───────────────────────────────────────────────────────────────────
NAVY      = '#0F2E52'
NAVY_MID  = '#1B3A5C'
NAVY_L    = '#1E3A5F'
GOLD      = '#C9A84C'
GOLD_L    = '#E2C97E'
ROUGE     = '#C0392B'
VERT      = '#1E8449'
ORANGE    = '#E67E22'
SLATE     = '#8A9BB0'
BG        = '#F5F7FA'
WHITE     = '#FFFFFF'
TEXT      = '#1C2B3A'

# ── Graphiques V3 (core/charts_tarif) — rendu HTML natif interactif ─────────────
# plotly.js chargé UNE FOIS en CDN dans le <head> ; chaque figure rendue via
# to_html(include_plotlyjs=False). Import gardé : si plotly/charts absents, les
# figures sont simplement omises (rétro-compatible, comme synthese_mapping).
try:
    import plotly.offline as _po
    from core.charts_tarif import CONFIG_PLOTLY as _CFG_PLOTLY
    _PLOTLYJS_CDN = (
        f'<script src="https://cdn.plot.ly/plotly-{_po.get_plotlyjs_version()}'
        f'.min.js" charset="utf-8"></script>')
    _CHARTS_HTML_OK = True
except Exception:
    _CFG_PLOTLY = {'displayModeBar': False}
    _PLOTLYJS_CDN = ''
    _CHARTS_HTML_OK = False

# ── Helpers formatage ─────────────────────────────────────────────────────────
# ⚠️ TROIS FORMATEURS ONT ÉTÉ RETIRÉS ICI — `_f`, `_pct` et `_s`. Ils
# reproduisaient la convention d'A7 et n'étaient appelés NULLE PART : vulture
# les signalait, et les 53 nombres du rapport étaient formatés en ligne, chacun
# à sa façon. Le rapport RESSEMBLAIT à A7 sans en appliquer un seul formatage —
# d'où « 2,435 » à l'anglaise à côté de « 10308 » sans séparateur. La
# convention vit désormais dans `core.format_fr`, en un seul endroit.

def _clean(txt) -> str:
    if not txt: return ''
    txt = re.sub(r'[■□▪▸►═╔╗╚╝║─]+', '', str(txt))
    txt = re.sub(r'={4,}', '', txt)
    txt = re.sub(r'\n{3,}', '\n\n', txt)
    return txt.strip()

def _statut_col(s: str) -> str:
    return {'VERT': VERT, 'AMBRE': ORANGE, 'ROUGE': ROUGE}.get(s.upper(), SLATE)

def _statut_emoji(s: str) -> str:
    return {'VERT': '🟢', 'AMBRE': '🟡', 'ROUGE': '🔴'}.get(s.upper(), '⚪')


# =============================================================================
#  SYSTEM PROMPT CLAUDE API
# =============================================================================
_PROMPT_TARIF_GABARIT = """\
Tu es un actuaire Non-Vie senior certifié par l'Institut des Actuaires (IA), \
expert en tarification Solvabilité 2 et IFRS 17, avec 20 ans d'expérience \
auprès de mutuelles et grands groupes d'assurance.

Tu rédiges le commentaire actuariel d'un rapport de tarification Non-Vie \
destiné à l'actuaire désigné, au Comité des risques et à l'ACPR.

RÈGLES ABSOLUES :
0. FORMAT : les sections s'écrivent « §N — TITRE », séparées par une ligne vide.
{CONSIGNE_MARKDOWN}
1. LANGUE : Français professionnel actuariel.
2. RIGUEUR : Chaque affirmation est justifiée par les données fournies.
3. CHIFFRES : Gini à 4 décimales, AIC arrondi, p-values à 4 décimales.
4. RÉFÉRENCES : Mildenhall (1999) GLM, Goldburd (2016) GLM P&C, \
Denuit (2019) calibration, Siddiqi (2006) PSI, Wüthrich & Merz (2019) CANN.
5. ALERTES : Ne jamais minimiser. Présenter avec l'implication tarifaire réelle.

STRUCTURE OBLIGATOIRE EN 7 SECTIONS :
§1 — CONTEXTE ET QUALITÉ DES DONNÉES
§2 — VALIDATION DES HYPOTHÈSES ACTUARIELLES (H1–H4)
§3 — RÉSULTATS GLM ET RELATIVITÉS TARIFAIRES
§4 — COMPARAISON DES MODÈLES ML ET SÉLECTION
§5 — BACKTESTING WALK-FORWARD ET TEST A/E
§6 — RISQUES IDENTIFIÉS ET POINTS D'ATTENTION
§7 — CONCLUSION ET RECOMMANDATIONS POUR L'ACTUAIRE DÉSIGNÉ\
"""

# ⚠️ LA CONSIGNE ANTI-MARKDOWN VIENT DE `core.narration` — un seul endroit du
# dépôt la porte, comme la conversion qui la double. CEINTURE ET BRETELLES :
# la consigne évite le marqueur, la conversion traite celui qui passe quand
# même. Une consigne peut se relâcher — le modèle glisse un marqueur, ou son
# comportement change d'une version à l'autre ; une conversion, non.
# ⚠️ ET LA RÈGLE 0 A ÉTÉ RÉÉCRITE : elle demandait « ### pour les
# sous-titres, **gras** pour les termes importants », c'est-à-dire
# exactement les marqueurs qui ressortaient en clair dans le livrable.
SYSTEM_PROMPT_TARIF = _PROMPT_TARIF_GABARIT.replace(
    '{CONSIGNE_MARKDOWN}', _md.CONSIGNE_SANS_MARKDOWN)


# =============================================================================
#  LE VOCABULAIRE DU RAPPORT — DEFINI UNE FOIS, CONSOMME PAR LES DEUX FORMATS
# =============================================================================
#
#  ⚠️ UN EN-TÊTE DÉFINI DEUX FOIS DIVERGE. Mesuré avant ce lot sur les quinze
#  tableaux du rapport : 21 libellés communs, 12 propres à l'HTML, 10 propres
#  au Word. Les mêmes colonnes portaient des noms différents — « IC 95% bas »
#  et « IC bas », « Relativité exp(β) » et « Relativité », « Significatif » et
#  « Sig. », « Score global » et « Score ». Ce ne sont pas deux rendus d'un
#  document, ce sont deux documents. C'est le motif que T1, T3 et T6 ont
#  chacun fermé sur leur objet ; celui-ci le ferme sur les mots.
#
#  ⚠️ ET SIX COLONNES MANQUAIENT AU WORD, PLUS UN TABLEAU ENTIER — les
#  contrôles de sélection C1/C2/C3. OUBLI, PAS CHOIX, et c'est mesuré : le
#  tableau de backtesting n'occupait que 10,0 cm sur les 16,5 cm utiles de la
#  page, et aucun commentaire du dépôt ne déclarait l'abrègement. Un choix se
#  déclare ; ce qui ne se déclare pas est un oubli. Les quatre exceptions qui
#  restent sont désormais écrites, juste en dessous.

class Colonne(NamedTuple):
    """Un en-tête de tableau, et la nature de ce qu'il porte.

    ⚠️ L'ALIGNEMENT SUIT LE LIBELLÉ, il ne se redéclare pas à chaque appel :
    une colonne dite numérique ici l'est dans les deux formats.
    """
    titre: str
    numerique: bool = False


def _c(titre: str, numerique: bool = False) -> Colonne:
    return Colonne(titre, numerique)


COLONNES: Dict[str, Tuple[Colonne, ...]] = {
    'glm': (
        _c('Modèle'), _c('Gini test', True), _c('AIC', True),
        _c('Déviance', True), _c('Pseudo-R²', True),
        _c('Vars retenues', True)),
    'relativites': (
        _c('Variable'), _c('β', True), _c('Relativité exp(β)', True),
        _c('IC 95% bas', True), _c('IC 95% haut', True), _c('p-value', True),
        _c('Significatif'), _c('Sens')),
    'classement': (
        _c('Rang'), _c('Modèle'), _c('Famille'), _c('Gini test', True),
        _c('RMSE test', True), _c('Overfit', True),
        _c('Score global', True)),
    'hypotheses': (
        _c('Hypothèse'), _c('Statut'), _c('Valeur', True), _c('Message'),
        _c('Conseil')),
    'backtest': (
        _c('Année test'), _c('N train', True), _c('N test', True),
        _c('Moy train', True), _c('Moy test', True), _c('A/E ratio', True),
        _c('Statut')),
    'controles': (
        _c('Contrôle sélection'), _c('Statut'), _c('Message')),
    'audit': (
        _c('Élément'), _c('Valeur')),
    # ── Les deux tableaux propres au Word, et pourquoi ────────────────────
    # ⚠️ EXCEPTIONS DÉCLARÉES, PAS OUBLIS. L'HTML porte les mêmes informations
    # sous une autre forme : la page de garde y est la bande d'en-tête, et le
    # modèle retenu y est une grille de six vignettes. Un `.docx` n'a ni l'une
    # ni l'autre. Ce sont les deux seuls endroits où la structure diffère —
    # partout ailleurs les deux formats portent les mêmes colonnes.
    'garde': (
        _c('Client'), _c('Branche'), _c('Arrêté'), _c('Identifiant d\'audit')),
    'production': (
        _c('Attribut'), _c('Valeur'), _c('Attribut'), _c('Valeur')),
}

# ⚠️ TROISIÈME EXCEPTION, DÉCORATIVE : l'étoile du premier du classement
# n'existe qu'en HTML. Elle ne dit rien que la colonne « Rang » ne dise déjà —
# c'est une mise en valeur, pas une donnée. Le Word met la même ligne en gras.


def titres(cle: str) -> List[str]:
    """Les en-têtes d'un tableau. La MÊME liste pour l'HTML et pour le Word."""
    return [c.titre for c in COLONNES[cle]]


def colonnes_numeriques(cle: str) -> Tuple[int, ...]:
    """Les index à aligner à droite — la règle posée par T3, déclarée ici."""
    return tuple(i for i, c in enumerate(COLONNES[cle]) if c.numerique)


# ── Les chapitres du rapport ────────────────────────────────────────────────
# ⚠️ « §4 » DÉSIGNAIT DEUX CHOSES DANS LE MÊME DOCUMENT : le chapitre
# « Validation des hypothèses » et, dans le commentaire actuariel, la section
# « COMPARAISON DES MODÈLES ML ET SÉLECTION » que le prompt impose au modèle.
# Un renvoi « voir §4 » n'y voulait rien dire — et le commentaire réel en
# contient. ⚠️ ET IL Y AVAIT UNE TROISIÈME NUMÉROTATION : celle de l'en-tête
# de ce fichier, qui décrivait un plan que le code n'a jamais produit.
# Le rapport parle donc de CHAPITRES et la narration garde ses « §N » — forme
# que son prompt lui impose et que T6 verrouille. Les deux ne se croisent plus.
CHAPITRES: Tuple[str, ...] = (
    'Résultats GLM — Poisson / Gamma / Tweedie',
    'Relativités tarifaires GLM Poisson — exp(β)',
    'Classement ML — grille multicritères '
    '(Gini 40 % · Stabilité 30 % · Interprétabilité 20 % · RMSE 10 %)',
    'Validation des hypothèses actuarielles',
    'Backtesting walk-forward et test A/E',
    'Modèle de production retenu',
    'Commentaire actuariel',
    'Piste d\'audit et traçabilité ACPR',
)


def chapitre(numero: int) -> str:
    """« Chapitre N — Titre », rendu à l'identique dans les deux formats.

    ⚠️ LA PONDÉRATION DE LA GRILLE VIT DANS LE TITRE, et non plus dans le seul
    HTML : le Word la perdait, alors qu'elle explique le classement.
    """
    return f'Chapitre {numero} — {CHAPITRES[numero - 1]}'


# ── Les huit hypothèses, nommées sans ambiguïté ─────────────────────────────
# ⚠️ QUATRE LIGNES S'APPELAIENT « H1 » À « H4 » ET QUATRE AUTRES « H1 ML » À
# « H4 ML », SOUS UN TITRE QUI ANNONÇAIT « H1–H4 ». Le lecteur avait huit
# lignes pour quatre numéros. Le Word préfixait déjà « H1 GLM » — l'HTML non :
# c'est cette forme-là qui est retenue, elle nomme la famille des deux côtés.
# Le troisième membre est la clé de la valeur mesurée, propre à chaque test.
HYPOTHESES: Tuple[Tuple[str, str, str], ...] = (
    ('h1_poisson',     'H1 GLM — Sur-dispersion Poisson',           'ratio_disp'),
    ('h2_homosc',      'H2 GLM — Homoscédasticité résidus Pearson',
     'ratio_variance'),
    ('h3_ajustement',  'H3 GLM — Qualité ajustement (Gini)',        'gini_max'),
    ('h4_stabilite',   'H4 GLM — Stabilité relativités bootstrap',  'cv_max'),
    ('h1_overfitting', 'H1 ML — Overfitting',                       'ratio'),
    ('h2_psi',         'H2 ML — PSI réel',                          'psi'),
    ('h3_gini',        'H3 ML — Performance Gini',                  'gini'),
    ('h4_calibration', 'H4 ML — Calibration',              'ecart_moy_pct'),
)

# ── Les contrôles de sélection ──────────────────────────────────────────────
# ⚠️ CE TABLEAU N'EXISTAIT QU'EN HTML. Il porte les trois contrôles qui
# justifient le choix du modèle de production : le fichier qui part chez un
# commissaire ne les avait pas.
CONTROLES_SELECTION: Tuple[Tuple[str, str], ...] = (
    ('c1_nb_modeles', 'C1 — Nombre de modèles'),
    ('c2_ecart_gini', 'C2 — Écart Gini'),
    ('c3_coherence',  'C3 — Cohérence'),
)


# ── Un modèle porte UN nom ──────────────────────────────────────────────────
# ⚠️ LE MODÈLE RETENU S'APPELAIT « GLM Poisson (référence A3) » au chapitre 3
# et « GLM_POISSON » au chapitre 6 — mesuré sur le rapport réel, 5 fois sous
# la seconde forme. Rien ne dit au lecteur qu'il s'agit du même. Six autres
# lignes du classement portaient leur identifiant technique tel quel :
# « lightgbm », « lineaire_regularise », « xgboost_tweedie ».
# La provenance entre parenthèses est CONSERVÉE : elle dit que ce modèle vient
# d'A3 et n'a pas été réentraîné par A6 — c'est une information, pas un nom.
_NOMS_MODELES = {
    'glm_poisson': 'GLM Poisson',
    'glm_gamma': 'GLM Gamma',
    'glm_tweedie': 'GLM Tweedie',
    'lineaire_regularise': 'Linéaire régularisé',
    'elasticnet': 'ElasticNet',
    'xgboost': 'XGBoost',
    'xgboost_tweedie': 'XGBoost Tweedie',
    'lightgbm': 'LightGBM',
    'catboost': 'CatBoost',
    'gbm': 'GBM',
    'random_forest': 'Random Forest',
    'cann': 'CANN',
    'tabnet': 'TabNet',
    'poisson': 'Poisson',
    'gamma': 'Gamma',
    'tweedie': 'Tweedie',
}


def nom_modele(brut) -> str:
    """Le nom lisible d'un modèle, quelle que soit son écriture technique.

    ⚠️ UN NOM INCONNU EST RENDU TEL QUEL, jamais remplacé ni masqué : la
    liste ci-dessus ne prétend pas connaître tous les modèles que la chaîne
    pourra produire, et un nom inventé serait pire qu'un nom technique.
    """
    if brut in (None, ''):
        return F.ABSENT
    texte = str(brut).strip()
    identifiant, _, precision = texte.partition('(')
    cle = identifiant.strip().lower().replace(' ', '_').replace('-', '_')
    lisible = _NOMS_MODELES.get(cle)
    # ⚠️ LA MÊME CHAÎNE PRODUIT DEUX ÉCRITURES SELON LE PASSAGE : le rapport
    # réel du 08/08 porte « lightgbm », un rejeu de la même chaîne porte
    # « ML_LIGHTGBM ». Le préfixe de famille est déjà dans la colonne
    # « Famille » — il ne fait pas partie du nom.
    if lisible is None and cle.startswith('ml_'):
        lisible = _NOMS_MODELES.get(cle[3:])
    if lisible is None:
        return texte
    return f'{lisible} ({precision.strip()}' if precision else lisible


# ── La piste d'audit se lit en français ─────────────────────────────────────
# ⚠️ LE CHAPITRE 8 AFFICHAIT « Ae Ratio », « Stabilite Wf », « Nb Modeles »,
# « Gouvernance Ok : True » et un horodatage à la microseconde. Ce sont des
# noms de variables mis en capitales par une substitution mécanique, pas la
# langue d'un rapport signé. Quatorze clés, mesurées sur le rapport réel.
_LIBELLES_AUDIT = {
    'agent': 'Agent producteur',
    'version': 'Version du moteur',
    'audit_id': 'Identifiant d\'audit',
    'timestamp': 'Date et heure de production',
    'branche': 'Branche',
    'statut_rag': 'Statut RAG',
    'nb_modeles': 'Nombre de modèles comparés',
    'modele_production': 'Modèle retenu pour la production',
    'ae_ratio': 'Ratio A/E',
    'stabilite_wf': 'Stabilité walk-forward',
    'profil_ponderation': 'Profil de pondération',
    'profil_valide_par': 'Profil validé par',
    'environnement': 'Environnement',
    'gouvernance_ok': 'Gouvernance validée',
}

_HORODATAGE = re.compile(r'^(\d{4})-(\d{2})-(\d{2})[T ](\d{2}):(\d{2})')


def libelle_audit(cle) -> str:
    """Le libellé lisible d'une clé de piste d'audit.

    ⚠️ UNE CLÉ INCONNUE RESTE VISIBLE — mise en forme, jamais effacée : une
    piste d'audit amputée ne serait plus une piste d'audit.
    """
    brut = str(cle)
    return _LIBELLES_AUDIT.get(brut, brut.replace('_', ' ').capitalize())


def valeur_audit(valeur, cle=None) -> str:
    """La valeur d'une entrée de piste d'audit, en français.

    ⚠️ « True » N'EST PAS DU FRANÇAIS, et la microseconde d'un horodatage
    n'apporte rien à un lecteur — elle encombre une colonne.

    ⚠️ ET LA PISTE D'AUDIT NOMMAIT LE MODÈLE UNE TROISIÈME FOIS, sous sa
    forme technique : c'est le même modèle que celui des chapitres 3 et 6, il
    porte donc le même nom.
    """
    if cle == 'modele_production':
        return nom_modele(valeur)
    if isinstance(valeur, bool):
        return 'oui' if valeur else 'non'
    texte = str(valeur)
    h = _HORODATAGE.match(texte)
    if h:
        return (f'{h.group(3)}/{h.group(2)}/{h.group(1)} '
                f'à {h.group(4)} h {h.group(5)}')
    if isinstance(valeur, float):
        return F.nombre(valeur, F.DEC_GINI)
    return F.tronque(texte, 120)


# =============================================================================
#  LE CATALOGUE DES FIGURES — UN PLAN UNIQUE, DEUX FORMATS
# =============================================================================
#
#  ⚠️ MESURÉ SUR UN REJEU RÉEL DE LA CHAÎNE : 30 figures existent en mémoire,
#  le rapport en rendait SIX. Et elles vivent dans DEUX dictionnaires — les
#  agents remplissent `graphiques` (19) et `graphiques_validation` (11). Le
#  second n'était lu par PERSONNE : relevé exhaustif, 25 sites de production
#  dans le dépôt, ZÉRO lecture.
#
#  ⚠️ CE N'EST PAS UN DÉFAUT D'APPARENCE. Le chapitre 4 s'appelle « Validation
#  des hypothèses actuarielles », publie huit verdicts et ne montrait aucune
#  preuve — pendant que la chaîne dessinait `sur_dispersion_poisson` (H1 GLM)
#  et `overfitting_train_test` (H1 ML) puis les jetait. Le chapitre 3 annonce
#  dans son titre une grille multicritères et n'en montrait pas la figure.
#  Un rapport qui affirme un verdict et cache le raisonnement est celui qu'on
#  ne peut pas défendre devant l'ACPR.
#
#  ⚠️ ET TOUT RENDRE SERAIT PIRE. Le dépôt produit NEUF variantes de
#  Lorenz/Gini et SEPT de scores. Le catalogue est donc DÉCIDÉ : quatorze
#  figures, douze concepts, aucun doublon — et ce qui est écarté l'est avec sa
#  raison, dans le code, sous un test qui tombe si une figure nouvelle
#  n'était pas classée.

TITRES_FIGURES: Dict[str, str] = {
    # chapitre 1 — les GLM
    'aic_comparaison': 'Comparaison des GLM par AIC',
    'chart_distribution_predictions': 'Distribution des fréquences prédites',
    # chapitre 2 — les relativités
    'chart_relativites_glm': 'Relativités tarifaires GLM Poisson — exp(β)',
    # chapitre 3 — le classement
    'scores_multicriteres': 'Grille multicritères — score par modèle',
    'scatter_gini_stabilite': 'Pouvoir discriminant et stabilité — '
                              'l\'arbitrage du classement',
    'chart_lorenz_gini': 'Courbe de Lorenz et coefficient de Gini',
    # chapitre 4 — les hypothèses
    'sur_dispersion_poisson': 'H1 GLM — sur-dispersion de Poisson',
    'chart_residus_qq': 'H2 GLM — QQ-plot des résidus de Pearson',
    'overfitting_train_test': 'H1 ML — surapprentissage, apprentissage '
                              'contre test',
    'monitoring_gini': 'H2 ML — dérive du pouvoir discriminant',
    # chapitre 5 — le backtesting
    'chart_walkforward_ae': 'Backtesting walk-forward — ratio A/E par fenêtre',
    'chart_lift_decile': 'Lift par décile de risque prédit',
    # chapitre 6 — le modèle retenu
    'radar_modele_retenu': 'Profil du modèle retenu',
    'chart_shap_summary': 'Importance SHAP — contribution des variables',
}

#: Le plan : quel chapitre porte quelles figures, dans quel ordre.
#:
#: ⚠️ C'EST LA SOURCE UNIQUE, ET LES DEUX FORMATS LA CONSOMMENT. A7 tient
#: deux compteurs positionnels indépendants — un dans l'HTML, un dans le Word
#: — que seul un test rattrape après coup. Ici la numérotation est calculée
#: UNE fois et les deux formats lisent le même résultat : « Figure 7 » ne
#: PEUT plus désigner deux choses. C'est le motif que ce chantier a fermé
#: quatre fois (T1 sur la narration, T5 sur les en-têtes).
PLAN_FIGURES: Tuple[Tuple[int, Tuple[str, ...]], ...] = (
    (1, ('aic_comparaison', 'chart_distribution_predictions')),
    (2, ('chart_relativites_glm',)),
    (3, ('scores_multicriteres', 'scatter_gini_stabilite',
         'chart_lorenz_gini')),
    (4, ('sur_dispersion_poisson', 'chart_residus_qq',
         'overfitting_train_test', 'monitoring_gini')),
    (5, ('chart_walkforward_ae', 'chart_lift_decile')),
    (6, ('radar_modele_retenu', 'chart_shap_summary')),
)

#: Où chaque figure du catalogue se trouve : (résultat d'agent, dictionnaire).
#: ⚠️ CINQ DES QUATORZE VIENNENT DE `graphiques_validation`, celui que
#: personne n'ouvrait.
SOURCES_FIGURES: Dict[str, Tuple[str, str]] = {
    'aic_comparaison': ('a3', 'graphiques'),
    'chart_distribution_predictions': ('a3', 'graphiques'),
    'chart_relativites_glm': ('a3', 'graphiques'),
    'chart_residus_qq': ('a3', 'graphiques'),
    'sur_dispersion_poisson': ('a3', 'graphiques_validation'),
    'chart_shap_summary': ('a4', 'graphiques'),
    'overfitting_train_test': ('a4', 'graphiques_validation'),
    'monitoring_gini': ('a4', 'graphiques_validation'),
    'chart_lorenz_gini': ('a6', 'graphiques'),
    'chart_lift_decile': ('a6', 'graphiques'),
    'chart_walkforward_ae': ('a6', 'graphiques'),
    'scatter_gini_stabilite': ('a6', 'graphiques'),
    'scores_multicriteres': ('a6', 'graphiques_validation'),
    'radar_modele_retenu': ('a6', 'graphiques_validation'),
}

#: Ce que la chaîne produit et que le rapport NE PUBLIE PAS, avec la raison.
#:
#: ⚠️ LA RAISON VIT DANS LE CODE, PAS DANS UN COMMENTAIRE. Un test compare
#: cette table au relevé des figures réellement produites par les agents :
#: une figure nouvelle qui ne serait ni au plan ni ici fait tomber la gate.
#: C'est ce qui empêche la liste de se périmer en silence.
FIGURES_ECARTEES: Dict[str, str] = {
    # les neuf variantes de Lorenz / Gini — une seule est publiée
    'lorenz': 'doublon — chart_lorenz_gini porte la courbe de Lorenz',
    'lorenz_glm': 'doublon — chart_lorenz_gini porte la courbe de Lorenz',
    'gini_comparaison': 'doublon — le Gini par modèle est au tableau du '
                        'chapitre 3',
    'gini_comparaison_glm': 'doublon — le Gini des GLM est au tableau du '
                            'chapitre 1',
    'comparaison_gini': 'doublon — et produite par A5, hors de la chaîne du '
                        'rapport',
    'scatter_gini_rmse': 'doublon — scatter_gini_stabilite porte l\'arbitrage '
                         'qui décide du classement',
    # les variantes de relativités
    'coefficients_glm': 'doublon — chart_relativites_glm porte les exp(β) '
                        'avec leurs intervalles',
    'relativites_poisson': 'doublon — chart_relativites_glm, même contenu',
    # les résidus
    'residus_deviance': 'doublon — chart_residus_qq porte le même diagnostic '
                        'sous une forme lisible',
    'durbin_watson': 'hors catalogue — l\'autocorrélation n\'est pas une des '
                     'huit hypothèses que le chapitre 4 publie',
    # le lift
    'lift_chart': 'doublon — chart_lift_decile, même mesure',
    # les scorecards : des tableaux déguisés en figures
    'scorecard_validation_glm': 'tableau déguisé — le chapitre 4 porte déjà '
                                'les huit hypothèses en tableau',
    'scorecard_validation_ml': 'tableau déguisé — même tableau, chapitre 4',
    'scorecard_validation_dl': 'tableau déguisé — et produite par A5',
    'scorecard_selection': 'tableau déguisé — les trois contrôles de '
                           'sélection sont au tableau du chapitre 6',
    # les trois vues supplémentaires du même classement
    'radar_multicriteres': 'doublon — scores_multicriteres porte la grille',
    'scores_profils': 'doublon — la grille du profil retenu suffit ; les '
                      'autres profils ne sont pas ceux qui décident',
    'scores_finaux': 'doublon — scores_multicriteres, même grandeur',
    # hors périmètre du rapport de modèles
    'optimisation_tarifaire': 'hors périmètre — c\'est du pilotage tarifaire, '
                              'pas la validation d\'un modèle',
    # A5 n'entre pas dans la chaîne du rapport
    'apprentissage_cann': 'A5 n\'entre pas dans la chaîne du rapport',
    'apprentissage_tabnet': 'A5 n\'entre pas dans la chaîne du rapport',
    'importance_tabnet': 'A5 n\'entre pas dans la chaîne du rapport',
    'convergence_loss': 'A5 n\'entre pas dans la chaîne du rapport',
    'comparaison_dl_glm': 'A5 n\'entre pas dans la chaîne du rapport',
    'jauge_surapprentissage': 'A5 n\'entre pas dans la chaîne du rapport',
}


# ── CE QUE LE WORD NE MONTRE PAS, ET QU'IL DOIT DIRE ────────────────────────
#
#  ⚠️ MESURÉ, ET INVISIBLE JUSQU'ICI : le Word coupe les relativités à quinze
#  lignes et le classement à dix, quand l'HTML les rend TOUTES. Mêmes
#  colonnes, mêmes libellés, même ordre — et pas les mêmes lignes. T5 avait
#  comparé les EN-TÊTES, pas le NOMBRE DE LIGNES : le défaut lui a échappé.
#  Un GLM à vingt variables retenues perdait ses cinq dernières relativités
#  dans le fichier qui part chez un commissaire aux comptes, SANS RIEN DIRE.
#
#  ⚠️ ARBITRAGE : DÉCLARER, PAS ALIGNER. Le tri est déjà par |β| décroissant
#  et par rang de classement : les premières lignes sont celles qui pèsent le
#  plus. C'est un choix éditorial défendable — À CONDITION D'ÊTRE ÉCRIT.
#  Aligner alourdirait un livrable papier sans rien apprendre au lecteur.
#
#  ⚠️ ET LE NOMBRE AFFICHÉ EST CELUI QUI COUPE : la limite et la phrase
#  lisent la même constante. Deux copies auraient fini par dire deux choses.
MAX_RELATIVITES_WORD = 15
MAX_CLASSEMENT_WORD = 10


def note_troncature(montrees: int, total: int, quoi: str, critere: str) -> str:
    """La phrase qui déclare une troncature, ou rien s'il n'y en a pas."""
    if total <= montrees:
        return ''
    return (f'{montrees} {quoi} sur {total}, {critere}. '
            f'Le rapport HTML les porte toutes.')


class FigureNumerotee(NamedTuple):
    """Une figure du plan, avec le numéro qu'elle porte dans le document."""
    numero: int
    cle: str
    titre: str
    objet: object


def figures_disponibles(result_a3, result_a4, result_a6) -> Dict[str, object]:
    """Rassemble les figures du catalogue, où qu'elles se trouvent.

    ⚠️ ELLE OUVRE LES DEUX DICTIONNAIRES. `graphiques_validation` n'avait
    aucun lecteur dans tout le dépôt ; cinq des quatorze figures publiées en
    viennent, dont les preuves de H1 GLM et de H1 ML.
    """
    resultats = {'a3': result_a3 or {}, 'a4': result_a4 or {},
                 'a6': result_a6 or {}}
    trouvees: Dict[str, object] = {}
    for cle, (agent, dictionnaire) in SOURCES_FIGURES.items():
        objet = ((resultats.get(agent) or {}).get(dictionnaire) or {}).get(cle)
        if objet is not None:
            trouvees[cle] = objet
    # ⚠️ UNE FIGURE NI PUBLIÉE NI ÉCARTÉE SE SIGNALE, ICI AUSSI. Un test
    # statique relit le code des agents ; ce contrôle-ci voit ce qui arrive
    # RÉELLEMENT dans les résultats — une figure ajoutée par un autre chemin
    # ne passerait pas devant lui.
    for resultat in resultats.values():
        for dictionnaire in ('graphiques', 'graphiques_validation'):
            for cle in (resultat.get(dictionnaire) or {}):
                if cle not in SOURCES_FIGURES and cle not in FIGURES_ECARTEES:
                    logger.warning(
                        'Figure « %s » produite par la chaîne mais ni publiée '
                        'ni écartée : le catalogue est à instruire.', cle)
    return trouvees


def numeroter(disponibles: Dict[str, object]) -> Dict[int, Tuple[
        FigureNumerotee, ...]]:
    """Numérote le plan POSITIONNELLEMENT — un seul passage, deux formats.

    ⚠️ LE COMPTEUR N'AVANCE QUE SUR UNE FIGURE RÉELLEMENT PRÉSENTE. Une
    figure absente — SHAP désactivé, une hypothèse non calculée — ne laisse
    donc aucun trou : la suivante prend son numéro. La continuité est
    garantie PAR CONSTRUCTION, pas par une liste qu'il faudrait tenir à jour.
    """
    numero = 0
    par_chapitre: Dict[int, Tuple[FigureNumerotee, ...]] = {}
    for chapitre, cles in PLAN_FIGURES:
        du_chapitre = []
        for cle in cles:
            objet = disponibles.get(cle)
            if objet is None:
                continue
            numero += 1
            du_chapitre.append(FigureNumerotee(
                numero, cle, TITRES_FIGURES[cle], objet))
        if du_chapitre:
            par_chapitre[chapitre] = tuple(du_chapitre)
    return par_chapitre


def legende(figure: FigureNumerotee) -> str:
    """« Figure N — titre ». La MÊME chaîne dans les deux formats."""
    return f'Figure {figure.numero} — {figure.titre}'


# ── LE RENDU : le format le plus léger, décidé par la mesure ────────────────
#
#  ⚠️ MESURÉ SUR LES FIGURES RÉELLES, ET LE RÉSULTAT N'EST PAS UNIFORME :
#  douze figures pèsent moins en SVG (6,7 à 21,8 ko), et le QQ-plot des
#  résidus — quatre cents points, donc quatre cents éléments de dessin —
#  passe de 49,8 ko en PNG à 940,0 ko en SVG. Une règle « tout en SVG »
#  aurait multiplié le poids du rapport par cinq sur cette seule figure.
#
#  ⚠️ LA RÈGLE COMPARE, ELLE NE LISTE PAS. Une figure future qui changerait
#  de nature — un nuage de points là où il y avait des barres — bascule
#  d'elle-même. Une liste figée aurait menti au premier changement.
#
#  ⚠️ ET LE POIDS COMPARÉ EST CELUI DU DOCUMENT, pas celui du fichier : un
#  PNG inséré dans un HTML autonome passe par base64, qui l'alourdit d'un
#  tiers. C'est ce chiffre-là qu'il faut comparer au SVG.

TAILLE_FIGURE = (1100, 520)


def rasteriser(figure, format: str) -> bytes:
    """Convertit une figure Plotly. LE SEUL point du module qui le fasse.

    ⚠️ SUBSTITUABLE PAR LES TESTS, ET C'EST UNE EXIGENCE, pas un confort :
    une conversion réelle coûte ~4 s par figure et par format. Le lot
    précédent a montré ce que devient une gate qui les paie.
    """
    largeur, hauteur = TAILLE_FIGURE
    if format == 'png':
        return figure.to_image(format='png', width=largeur, height=hauteur,
                               scale=1)
    return figure.to_image(format='svg', width=largeur, height=hauteur)


def kaleido_disponible() -> bool:
    """Le module qui convertit une figure Plotly en image est-il là ?"""
    import importlib.util
    return importlib.util.find_spec('kaleido') is not None


class RenduFigure(NamedTuple):
    """Ce qu'une figure devient dans chaque format, et pourquoi.

    `genre` vaut 'svg', 'png' ou 'interactif' ; `note` est la phrase que le
    document porte quand le rendu est dégradé — jamais un silence.
    """
    html: str
    png: Optional[bytes]
    genre: str
    note: str = ''


def _svg_propre(octets: bytes) -> str:
    """Le SVG rendu par kaleido porte un en-tête XML : on garde la balise."""
    texte = octets.decode('utf-8', 'replace')
    debut = texte.find('<svg')
    return texte[debut:] if debut >= 0 else texte


def rendre_figure(objet) -> RenduFigure:
    """Le rendu d'une figure dans les deux formats, avec le repli en cascade.

    ⚠️ SANS kaleido, ON DÉGRADE VERS LE COMPORTEMENT D'AUJOURD'HUI, NOMMÉ.
    Le fragment interactif reste — il s'affiche, mais il exige une connexion
    pour charger sa bibliothèque. Le document le DIT. Dégrader vers rien
    aurait été plus grave que l'état antérieur ; dégrader en silence, comme
    aujourd'hui où l'absence de réseau laisse un cadre vide, l'est déjà.
    """
    if not kaleido_disponible():
        return RenduFigure(
            html=_fragment_interactif(objet), png=None, genre='interactif',
            note='Figure interactive : le module « kaleido » n\'est pas '
                 'installé sur la machine de production. Son affichage exige '
                 'une connexion Internet.')
    try:
        svg = rasteriser(objet, 'svg')
        png = rasteriser(objet, 'png')
    except Exception as erreur:
        return RenduFigure(
            html=_fragment_interactif(objet), png=None, genre='interactif',
            note='Figure non convertie en image (%s) : elle reste '
                 'interactive et exige une connexion Internet.'
                 % type(erreur).__name__)
    # ⚠️ base64 alourdit un PNG d'un tiers : c'est le poids DANS LE DOCUMENT
    # qu'on compare, pas celui du fichier.
    if len(svg) <= len(png) * 4 // 3:
        return RenduFigure(html=_svg_propre(svg), png=png, genre='svg')
    # ⚠️ PAS DE `%` DE FORMATAGE ICI : le style porte « width:100% », que
    # l'interpolation prendrait pour un spécificateur. Le test de la règle a
    # attrapé exactement cela — cette branche n'avait jamais été exercée.
    b64 = base64.b64encode(png).decode('ascii')
    return RenduFigure(
        html=('<img alt="" style="width:100%;height:auto;" '
              'src="data:image/png;base64,' + b64 + '">'),
        png=png, genre='png')


def _fragment_interactif(objet) -> str:
    """Le repli : la figure Plotly telle que le rapport la rendait avant."""
    try:
        return objet.to_html(full_html=False, include_plotlyjs=False,
                             config=_CFG_PLOTLY)
    except Exception as erreur:
        logger.debug('Fragment interactif non rendu : %s', erreur)
        return ''


def _construire_contexte_tarif(
    result_a3: Dict, result_a4: Dict, result_a6: Dict,
    branche: str, arrete: str
) -> str:
    met3  = result_a3.get('metriques', {}) if result_a3 else {}
    hyp3  = result_a3.get('hypotheses', {}) if result_a3 else {}
    rels  = result_a3.get('relativites_poisson', {}) if result_a3 else {}
    # ⚠️ LE SCORE EXISTAIT, LE RAPPORT LISAIT LA MAUVAISE SOURCE. A4 RANGE
    # les modèles ; A6 les SCORE (grille multicritères, 40 % Gini / 30 %
    # stabilité / 20 % interprétabilité / 10 % RMSE) et publie SON propre
    # `classement`, enrichi de `score_global`. Le rapport lisait celui d'A4 —
    # donc sans score — pendant que sa section « Modèle retenu » affichait
    # 1.0000, venu d'A6 : deux sections du même document, deux vérités sur le
    # même modèle. Mesuré : la colonne Score était vide sur les SEPT lignes.
    cl4   = ((result_a6 or {}).get('classement')
             or (result_a4.get('classement', []) if result_a4 else []))
    hyp4  = result_a4.get('hypotheses', {}) if result_a4 else {}
    bt6   = result_a6.get('backtest', {}) if result_a6 else {}
    val6  = result_a6.get('validation_selection', {}) if result_a6 else {}
    prod  = result_a6.get('modele_production', {}) if result_a6 else {}

    # ⚠️ CE QUI N'EST PAS CALCULÉ NE VAUT PAS ZÉRO — surtout ici : ce texte
    # est lu par le modèle qui rédige le commentaire actuariel. « Gini=0.0000 »
    # lui AFFIRME un pouvoir discriminant nul ; « Gini=— » lui dit la vérité.
    m_poi = met3.get('poisson', {})
    vars_r = m_poi.get('vars_retenues', [])

    # Modules avancés P2 (A3) et gouvernance (A6)
    cred6 = result_a3.get('credibilite', {}) if result_a3 else {}
    geo6  = result_a3.get('lissage_geo', {}) if result_a3 else {}
    at6   = result_a6.get('audit_trail', {}) if result_a6 else {}

    lines = [
        f"DOSSIER TARIFICATION — {branche.upper()} — Arrêté {arrete}",
        "",
        "=== DONNÉES ===",
        f"Nb obs train : {m_poi.get('nb_obs_train', '—')} | Nb obs test : {m_poi.get('nb_obs_test', '—')}",
        f"Variables retenues : {', '.join(vars_r) or '—'}",
        "",
        "=== GLM ===",
        f"Poisson : Gini={F.nombre(m_poi.get('gini'), F.DEC_GINI)} | "
        f"AIC={F.nombre(m_poi.get('aic'), 0)} | "
        f"Pseudo-R²={F.nombre(m_poi.get('pseudo_r2'), F.DEC_GINI)}",
        f"Gamma   : Gini={F.nombre(met3.get('gamma',{}).get('gini'), F.DEC_GINI)} | "
        f"AIC={F.nombre(met3.get('gamma',{}).get('aic'), 0)}",
        f"Tweedie : Gini={F.nombre(met3.get('tweedie',{}).get('gini'), F.DEC_GINI)} | "
        f"AIC={F.nombre(met3.get('tweedie',{}).get('aic'), 0)}",
        "",
        "=== HYPOTHÈSES GLM ===",
        f"H1 Sur-dispersion : {hyp3.get('h1_poisson',{}).get('statut','?')} | Var/E={hyp3.get('h1_poisson',{}).get('ratio_disp','—')}",
        f"H2 Homoscédasticité : {hyp3.get('h2_homosc',{}).get('statut','?')} | Ratio var={hyp3.get('h2_homosc',{}).get('ratio_variance','—')}",
        f"H3 Gini : {hyp3.get('h3_ajustement',{}).get('statut','?')} | Gini max={hyp3.get('h3_ajustement',{}).get('gini_max','—')}",
        f"H4 Stabilité relativités : {hyp3.get('h4_stabilite',{}).get('statut','?')} | CV max={hyp3.get('h4_stabilite',{}).get('cv_max','—')} | Instables={hyp3.get('h4_stabilite',{}).get('vars_instables',[])}",
        "",
        "=== RELATIVITÉS POISSON (top 5) ===",
    ]
    for var, d in list(rels.items())[:5]:
        lines.append(
            f"  {var} : exp(β)={F.nombre(d.get('relativite'), F.DEC_GINI)} | "
            f"p={F.nombre(d.get('pvalue'), F.DEC_GINI)} | {d.get('sens','')} | "
            f"sig={'oui' if d.get('significatif') else 'non'}")

    # === CRÉDIBILITÉ BÜHLMANN-STRAUB (P2) ===
    lines += ["", "=== CRÉDIBILITÉ BÜHLMANN-STRAUB ==="]
    if cred6.get('appliquee'):
        lines.append(
            f"Appliquée : k={cred6.get('k',0):.4f} | Z_moyen={cred6.get('z_moyen',0):.4f} "
            f"| n_groupes={cred6.get('n_groupes',0)} | colonne={cred6.get('col_groupe','—')}"
        )
    else:
        lines.append(f"Non applicable : {cred6.get('raison', '—')}")

    # === LISSAGE GÉOGRAPHIQUE (P2) ===
    lines += ["", "=== LISSAGE GÉOGRAPHIQUE ==="]
    if geo6.get('applique'):
        lines.append(
            f"Appliqué : méthode={geo6.get('methode','—')} | n_zones={geo6.get('n_zones',0)} "
            f"| source={geo6.get('source_prime','—')}"
        )
    else:
        lines.append(f"Non applicable : {geo6.get('raison', '—')}")

    # Mapping client (couche 2) — section CONDITIONNELLE (rien si aucun mapping,
    # contrairement aux sections qualité/colonnes qui affichent un repli). Même
    # source-unique partagée avec Excel A6 / rapport équipe.
    _synth_map6 = synthese_mapping(result_a6.get('rapport_mapping') if result_a6 else None)
    _lignes_map6 = (["=== MAPPING CLIENT (renommage du fichier avant tarification) ===",
                     _synth_map6, ""] if _synth_map6 else [])

    # ⚠️ LE CLASSEMENT ÉTAIT EXTRAIT PUIS JETÉ — `ruff F841` le signalait depuis
    # toujours. Le prompt ordonne pourtant « §4 — COMPARAISON DES MODÈLES ML ET
    # SÉLECTION », et sa règle 9 interdit « les phrases génériques sans
    # données » : le modèle recevait UN modèle et devait en comparer sept.
    # Mesuré sur le rapport réel produit avant ce lot : §4 citait 1 modèle sur
    # 7, aucun contrôle de sélection, et pas une occurrence de « comparaison »
    # hors de son propre titre. Il n'inventait pas — il remplissait la section
    # avec autre chose, ce qui est le bon comportement face à un contexte vide.
    lines += ["", f"=== CLASSEMENT ML ({len(cl4)} modèle(s) comparé(s)) ==="]
    for rang, m in enumerate(cl4, 1):
        # ⚠️ « NON CALCULÉ » NE VAUT PAS ZÉRO : un Gini à 0,0000 affirmerait un
        # pouvoir discriminant nul. Même garde que le tableau du rapport.
        def _n(cle, dec=4):
            v = m.get(cle)
            return f'{float(v):.{dec}f}' if isinstance(v, (int, float)) else '—'
        lines.append(
            f"{rang}. {m.get('modele','—')} | famille={m.get('famille','—')} | "
            f"Gini={_n('gini_test')} | RMSE={_n('rmse_test')} | "
            f"overfit={_n('overfit_ratio', 3)} | score={_n('score_global')}")
    if not cl4:
        lines.append("aucun classement transmis — la comparaison des modèles "
                     "ne peut pas être commentée")

    lines += [
        "",
        "=== ML — MODÈLE RETENU ===",
        f"Modèle production : {prod.get('modele','—')} | Score={prod.get('score_global',0):.4f} | Gini={prod.get('gini_test',0):.4f}",
        f"Overfit ratio : {prod.get('overfit_ratio',0):.3f}",
    ]

    # ⚠️ LES TROIS CONTRÔLES DE SÉLECTION, EUX AUSSI EXTRAITS PUIS JETÉS. Ils
    # disent si la comparaison était robuste, si les modèles se distinguaient,
    # et si le choix est cohérent avec le classement. Sans eux, la SÉLECTION du
    # titre de la section n'a rien derrière elle.
    lines += ["", "=== CONTRÔLES DE SÉLECTION ==="]
    for _cle, _libelle in (('c1_nb_modeles', 'C1 — nombre de modèles comparés'),
                           ('c2_ecart_gini', 'C2 — écart de Gini entre modèles'),
                           ('c3_coherence', 'C3 — cohérence du modèle retenu')):
        _c = val6.get(_cle) or {}
        lines.append(f"{_libelle} : {_c.get('statut', 'non calculé')} | "
                     f"{_c.get('message', '—')}")

    lines += [
        "",
        "=== HYPOTHÈSES ML ===",
        f"H1 Overfitting : {hyp4.get('h1_overfitting',{}).get('statut','?')} | ratio={hyp4.get('h1_overfitting',{}).get('ratio','—')}",
        f"H2 PSI réel : {hyp4.get('h2_psi',{}).get('statut','?')} | PSI={hyp4.get('h2_psi',{}).get('psi','—')}",
        f"H3 Gini : {hyp4.get('h3_gini',{}).get('statut','?')} | Gini={hyp4.get('h3_gini',{}).get('gini','—')}",
        f"H4 Calibration : {hyp4.get('h4_calibration',{}).get('statut','?')} | écart moy={hyp4.get('h4_calibration',{}).get('ecart_moy_pct','—')}%",
        "",
        "=== BACKTESTING A/E (walk-forward recalibré) ===",
        f"A/E ratio : {bt6.get('ae_ratio','—')} | {bt6.get('interpretation','—')}",
        f"Walk-forward : {bt6.get('n_fenetres','—')} fenêtres | stabilité={bt6.get('stabilite_wf','—')} | CV={bt6.get('ae_cv_wf','—')}",
        f"Modèle recalibré par fenêtre : {bt6.get('modele_recalibre','—')} | Gini WF moyen={bt6.get('gini_wf_moyen','—')}",
        # Audit V11 : sans cette ligne, le modèle de narration ignorait que la
        # validation temporelle avait pu porter sur un AUTRE modèle que celui
        # retenu — et rédigeait un commentaire faussement rassurant.
        (avertissement_walk_forward(bt6) or
         "Validation temporelle : porte bien sur le modèle de production."),
        "",
        "=== COLONNES ÉCARTÉES DE LA MATRICE X (conformité) ===",
        (synthese_exclusions(result_a6.get('exclusions_conformite') if result_a6 else None)
         or "Aucune colonne écartée : toutes les variables candidates sont conformes."),
        "",
        "=== SINISTRALITÉ PASSÉE CONSERVÉE (à vérifier par l'actuaire) ===",
        (synthese_alertes_experience(result_a6.get('alertes_conformite') if result_a6 else None)
         or "Aucune variable d'expérience passée à signal atypique."),
        "",
        "=== MODÈLE DEEP LEARNING EN PRODUCTION (validation humaine) ===",
        (synthese_modele_dl(result_a6.get('modele_production') if result_a6 else None,
                            result_a6.get('valide_par_actuaire_dl') if result_a6 else None,
                            at6.get('timestamp'))
         or "Aucun modèle Deep Learning retenu : pas de validation humaine spécifique requise."),
        "",
        "=== QUALITÉ DES DONNÉES (traitements de la couche générique, chemin déclaratif) ===",
        (synthese_qualite_donnees(result_a6.get('rapport_qualite') if result_a6 else None)
         or "Aucun traitement de qualité de données à signaler (ou couche non exécutée sur ce chemin)."),
        "",
        "=== COLONNES DU PLAN NON PRODUITES (modèle amputé) ===",
        (synthese_colonnes_plan_manquantes(
            result_a6.get('colonnes_plan_manquantes') if result_a6 else None)
         or "Aucune : toutes les colonnes déclarées au plan ont été produites."),
        "",
        *_lignes_map6,
        "=== GOUVERNANCE DU PROFIL DE PONDÉRATION ===",
        f"Profil retenu : {at6.get('profil_ponderation','—')} | Environnement={at6.get('environnement','—')}",
        f"Validé par : {at6.get('profil_valide_par') or 'NON VALIDÉ'} | Conforme={at6.get('gouvernance_ok','—')}",
        "",
        "Rédige le commentaire actuariel complet en 7 sections, en intégrant "
        "la crédibilité, le lissage géographique et la gouvernance du profil "
        "si ces éléments sont pertinents pour la branche analysée.",
    ]
    return '\n'.join(lines)


# ⚠⚠ CE BLOC AFFIRMAIT UNE GARANTIE QUE LE MODÈLE NE PERMET PLUS. Il
# prescrivait `temperature=0` comme socle de reproductibilité réglementaire.
# MESURÉ CONTRE L'API le 2026-08-07 : ce modèle REFUSE le paramètre (400,
# « deprecated for this model »), quelle que soit sa valeur. La prescription
# ne produisait donc pas du déterminisme : elle produisait un appel rejeté,
# et un repli muet sur le commentaire d'agent.
#
# CE QUI RESTE VRAI, ET QUI ÉTAIT DÉJÀ ÉCRIT ICI : la reproductibilité d'une
# narration n'a JAMAIS été garantie bit-à-bit. Des variations d'infrastructure
# côté fournisseur (répartition de charge, arrondis flottants liés au batching,
# mises à jour internes non annoncées par un changement de nom de modèle)
# peuvent produire un texte différent sur les mêmes entrées. Le paramètre
# retiré apportait donc moins qu'annoncé, et son retrait ne coûte pas la
# garantie — celle-ci n'existait pas.
#
# Réf. : GL EIOPA ORSA GL 56 (déterminisme des processus de calcul S2) ;
# ACPR-2022-P-01 (auditabilité des sorties IA). Ce qui satisfait ces exigences
# n'est pas un réglage d'échantillonnage, c'est la RELECTURE par l'actuaire
# responsable, ci-dessous.
#
# Conséquence pour l'usage réglementaire : le commentaire narratif généré
# doit être traité comme une AIDE À LA RÉDACTION, relue et validée par
# l'actuaire responsable avant diffusion — jamais comme une sortie de
# calcul déterministe au même titre que les modèles GLM/ML eux-mêmes
# (qui, eux, sont reproductibles bit-à-bit sur la même version de
# statsmodels/scikit-learn et les mêmes données).
# L'identifiant vient de la frontière (C1) ; ce nom reste parce que le pied
# de page du rapport le cite au lecteur.
CLAUDE_MODEL_TARIF = frontiere_llm.MODELE_RECENT

# ⚠️ TEMPERATURE RETIRÉE — MESURÉ CONTRE L'API le 2026-08-07 : ce modèle REFUSE
# le paramètre (400, « deprecated for this model »), quelle que soit sa valeur.
# CE SITE ÉTAIT EN PANNE SILENCIEUSE : chaque rapport de tarification tentait
# l'appel, était rejeté, et retombait sur le commentaire d'agent — l'actuaire
# lisait « Source : commentaire_agent » sans savoir que la plateforme avait
# essayé et échoué. `None` = paramètre non transmis.
CLAUDE_TEMPERATURE_TARIF = None

# Sources de narration. TROIS REPLIS, TROIS NOMS : un repli d'ENVIRONNEMENT est
# attendu ; un repli sur REQUÊTE REFUSÉE est un défaut de configuration ; un
# repli sur RÉPONSE INEXPLOITABLE veut dire que le service a répondu et qu'on
# n'a rien su en faire. Les confondre, c'est ce qui a coûté six réponses reçues.
SRC_API = 'claude_api'
SRC_REPLI = 'commentaire_agent'
SRC_DEFAUT = 'commentaire_agent (appel refuse — defaut de configuration)'
SRC_REPONSE = 'commentaire_agent (reponse recue mais inexploitable)'

def _narration_claude(result_a3, result_a4, result_a6, branche, arrete) -> Tuple[str, str]:
    repli = SRC_REPLI
    try:
        ctx = _construire_contexte_tarif(result_a3, result_a4, result_a6, branche, arrete)
        resp = frontiere_llm.appeler(
            modele=CLAUDE_MODEL_TARIF, max_tokens=6000,
            temperature=CLAUDE_TEMPERATURE_TARIF,
            systeme=SYSTEM_PROMPT_TARIF,
            messages=[{'role': 'user', 'content': ctx}],
            cle=frontiere_llm.cle_api_ou_secrets(),
        )
        return frontiere_llm.texte_des_blocs(resp), SRC_API
    except frontiere_llm.RequeteRefusee as e:
        # ⚠️ CE N'EST PAS UNE DÉGRADATION, C'EST UN DÉFAUT. Journalisé en
        # ERREUR, et le livrable le dit — c'est ce qui manquait.
        logger.error(f'Narration NON generee — requete refusee par la '
                     f'frontiere : {e}')
        repli = SRC_DEFAUT
    except frontiere_llm.ReponseInexploitable as e:
        # ⚠️ L'API A REPONDU. Le dire est le contraire de ce que ce site
        # faisait : il journalisait « Claude API indisponible » sur un 200.
        logger.error(f'Narration NON generee — le service a repondu mais sa '
                     f'reponse est inexploitable : {e}')
        repli = SRC_REPONSE
    except Exception as e:
        logger.warning(f'Narration non generee : {e}')
    # Fallback : commentaire agent
    for r in [result_a6, result_a4, result_a3]:
        if r and r.get('commentaire'):
            return _clean(r['commentaire']), repli
    return '', 'aucune'


# =============================================================================
#  HTML
# =============================================================================

# ⚠️ L'ANCIENNE CONVERSION LOCALE A ÉTÉ RETIRÉE ICI (T6). Elle ne traitait
# que 3 formes de markdown sur 7 et produisait un HTML MALFORMÉ : elle
# enveloppait les <li> dans un <ul> par une substitution de regex, puis
# passait chaque LIGNE au découpeur de paragraphes — la balise </ul> se
# retrouvait donc À L'INTÉRIEUR d'un <p>. La conversion vit désormais dans
# `core.narration`, partagée avec le rendu Word, et un test valide que les
# balises se ferment.


def export_html(
    result_a3: Dict = None, result_a4: Dict = None, result_a6: Dict = None,
    ref_client: str = '', arrete: str = '', audit_id: str = '',
    narration_calculee: Optional[Tuple[str, str]] = None,
) -> str:
    """Génère le rapport HTML tarification. Retourne str HTML ou ''.

    `narration_calculee` : le couple (texte, source) déjà obtenu, à réutiliser.
    ⚠️ SANS LUI, CHAQUE FORMAT PAIE SA PROPRE NARRATION — et rien ne garantit
    qu'ils obtiennent la même. Mesuré sur un rapport réel : le HTML portait les
    18 089 caractères de commentaire actuariel, le Word portait le dépôt
    technique de l'agent. Le fichier qui part chez un commissaire n'avait pas
    le commentaire. `generer_rapport_tarification` calcule donc UNE fois et
    transmet ; appelé seul, cet export calcule pour lui-même comme avant.
    """
    now    = datetime.now().strftime('%d/%m/%Y %H:%M')
    arr    = arrete or datetime.now().strftime('%d/%m/%Y')
    branche = (result_a6 or result_a3 or {}).get('branche', 'non_vie')
    statut  = (result_a6 or result_a3 or {}).get('statut_rag', 'AMBRE')
    s_col   = _statut_col(statut)
    s_emoji = _statut_emoji(statut)

    narration, n_src = (
        narration_calculee if narration_calculee is not None
        else _narration_claude(result_a3, result_a4, result_a6, branche, arr))
    # ⚠️ MÊME ANALYSE QUE LE WORD (T6). L'ancienne conversion locale ne
    # traitait que 3 formes sur 7 et produisait un HTML MALFORMÉ : elle
    # enveloppait les <li> par une regex puis passait chaque LIGNE au
    # découpeur de paragraphes, si bien que la balise </ul> se retrouvait
    # DANS un <p>.
    narr_html = (_md.en_html(narration) if narration
                 else '<p><em>Narration non disponible.</em></p>')

    met3  = (result_a3 or {}).get('metriques', {})
    rels  = (result_a3 or {}).get('relativites_poisson', {})
    # ⚠️ LE SCORE EXISTAIT, LE RAPPORT LISAIT LA MAUVAISE SOURCE. A4 RANGE
    # les modèles ; A6 les SCORE (grille multicritères, 40 % Gini / 30 %
    # stabilité / 20 % interprétabilité / 10 % RMSE) et publie SON propre
    # `classement`, enrichi de `score_global`. Le rapport lisait celui d'A4 —
    # donc sans score — pendant que sa section « Modèle retenu » affichait
    # 1.0000, venu d'A6 : deux sections du même document, deux vérités sur le
    # même modèle. Mesuré : la colonne Score était vide sur les SEPT lignes.
    cl4   = ((result_a6 or {}).get('classement')
             or (result_a4 or {}).get('classement', []))
    hyp3  = (result_a3 or {}).get('hypotheses', {})
    hyp4  = (result_a4 or {}).get('hypotheses', {})
    bt6   = (result_a6 or {}).get('backtest', {})
    prod  = (result_a6 or {}).get('modele_production', {})
    val6  = (result_a6 or {}).get('validation_selection', {})
    # Priorité A6 (le plus complet — gouvernance + WF recalibré) > A4 > A3
    at    = (
        (result_a6 or {}).get('audit_trail')
        or (result_a4 or {}).get('audit_trail')
        or (result_a3 or {}).get('audit_trail')
        or {}
    )

    # ── LES FIGURES — le plan unique, numéroté une fois ──────────────────
    # ⚠️ CE BLOC LISAIT SEPT CLÉS EN DUR, TOUTES DANS `graphiques`, et
    # ignorait `graphiques_validation` en entier. Il consomme désormais
    # `PLAN_FIGURES` — la même source que le Word, si bien que « Figure 7 »
    # ne peut plus désigner deux choses.
    _figures = {}
    _rendus = {}
    _genres = []
    if _CHARTS_HTML_OK:
        _figures = numeroter(figures_disponibles(result_a3, result_a4,
                                                 result_a6))
        for _fs in _figures.values():
            for _f in _fs:
                _r = rendre_figure(_f.objet)
                _rendus[_f.cle] = _r
                _genres.append(_r.genre)

    def _bloc_figure(f):
        """Une figure numérotée, et la raison de sa dégradation s'il y en a."""
        r = _rendus.get(f.cle)
        if r is None or not r.html:
            return ''
        note = (f'<div style="font-size:10px;color:{SLATE};font-style:italic;'
                f'margin-top:4px;">✦ {r.note}</div>') if r.note else ''
        return (f'<div class="figure" style="margin:16px 0;">'
                f'<div style="color:{NAVY};font-weight:600;font-size:12px;'
                f'margin-bottom:6px;">{legende(f)}</div>'
                f'<div style="max-width:100%;overflow:hidden;">{r.html}</div>'
                f'{note}</div>')

    # ⚠️ plotly.js n'est chargé QUE si une figure reste interactive : une
    # image n'a besoin d'aucune bibliothèque, et un rapport qui n'en porte
    # plus ne doit pas réclamer le réseau pour rien.
    _plotly_cdn = _PLOTLYJS_CDN if 'interactif' in _genres else ''

    def _row(cells, header=False, num=()):
        """Une ligne de tableau. `num` = les index des colonnes NUMÉRIQUES.

        ⚠️ L'ALIGNEMENT SE DÉCLARE, IL NE SE DEVINE PAS. Un tiret « — » d'une
        colonne de chiffres doit s'aligner comme un chiffre ; une heuristique
        sur le contenu de la cellule ne peut pas le savoir, la colonne si.
        Mesuré avant ce lot : 41 cellules numériques sur 95 alignées à droite,
        et les deux tableaux GLM — ceux qu'un actuaire lit en premier — les
        avaient TOUTES à gauche, sous des en-têtes qui annonçaient « right ».
        """
        tag = 'th' if header else 'td'
        return '<tr>' + ''.join(
            f'<{tag} class="right">{c}</{tag}>' if i in num
            else f'<{tag}>{c}</{tag}>'
            for i, c in enumerate(cells)) + '</tr>'

    def _figures_du_chapitre(numero):
        return ''.join(_bloc_figure(f) for f in _figures.get(numero, ()))

    def _fermer_chapitre(numero, avec_table=True):
        """La clôture d'un chapitre : SES figures, puis la fermeture.

        ⚠️ ELLES ÉTAIENT GROUPÉES APRÈS LE CHAPITRE 6, dans une section sans
        numéro, hors de la numérotation des chapitres. Le lecteur devait
        quitter la section pour aller chercher la preuve de ce qu'elle
        affirmait — et le chapitre 4, qui publie huit verdicts, renvoyait sa
        preuve trois chapitres plus loin.

        ⚠️ LE NUMÉRO DE FIGURE NE BOUGE PAS D'UN CRAN : `numeroter` parcourt
        le plan une fois, avant tout rendu. Déplacer le GESTE de rendu ne
        touche pas au compteur — c'est ce qu'un compteur positionnel calculé
        en amont garantit, et ce qu'une numérotation faite au fil du document
        n'aurait pas garanti.
        """
        return (('    </table>\n' if avec_table else '\n')
                + _figures_du_chapitre(numero) + '  </div>\n</div>\n')

    def _ouvrir_chapitre(numero, classe=''):
        """L'ouverture d'un chapitre — son titre vient de `CHAPITRES`.

        ⚠️ CE GABARIT ÉTAIT RECOPIÉ HUIT FOIS, chaque copie portant son titre
        en dur, et le Word en portait huit autres. C'est ainsi que les deux
        formats ont fini par annoncer des chapitres différents.
        """
        return (f'\n<!-- CHAPITRE {numero} -->\n'
                f'<div class="section">\n'
                f'  <div class="section-head">{chapitre(numero)}</div>\n'
                f'  <div class="section-body{classe}">\n')

    def _hyp_row(key, label, val_key='', val=''):
        # Cherche dans hyp3 (GLM), puis hyp4 (ML) — clés toujours distinctes.
        # ⚠️ ELLE DISPARAISSAIT EN SILENCE. Le titre de la section promet
        # « H1–H4 » (douze endroits du dépôt le promettent) et la ligne
        # s'effaçait : le lecteur ne pouvait pas distinguer « vérifiée » de
        # « jamais calculée ». Publier ce qui n'a pas pu l'être vaut mieux que
        # de le cacher — c'est la règle du projet, sans exception.
        h = hyp3.get(key) or hyp4.get(key) or {}
        if not h:
            return _row([label,
                         '<span class="badge-non-calcule">○ NON CALCULÉE</span>',
                         F.ABSENT, F.ABSENT,
                         'Hypothèse non produite par la chaîne — '
                         'à instruire avant validation'],
                        num=colonnes_numeriques('hypotheses'))
        st = h.get('statut', '?')
        em = _statut_emoji(st)
        brut = h.get(val_key, val)
        # ⚠️ LA MÊME COLONNE AFFICHAIT 1, 3 ET 4 DÉCIMALES SELON LA LIGNE —
        # « 1.0 » à côté de « 1.2176 ». Une précision se justifie, elle ne se
        # choisit pas ligne par ligne. Une valeur non numérique (un libellé)
        # passe telle quelle ; une absence rend le tiret de la convention.
        valeur = (F.nombre(brut, F.DEC_GINI)
                  if isinstance(brut, (int, float)) and not isinstance(brut, bool)
                  else (str(brut)[:60] if brut else F.ABSENT))
        # ⚠️ LA COLONNE « MESSAGE » N'EXISTAIT QU'EN WORD. Elle porte le
        # diagnostic du test ; l'HTML n'affichait que le conseil qui en
        # découle, c'est-à-dire la conclusion sans son motif.
        return _row([label,
                     f'<span class="badge-{st.lower()}">{em} {st}</span>',
                     valeur,
                     F.tronque(h.get('message'), 70),
                     F.tronque(h.get('conseil'), 80)],
                    num=colonnes_numeriques('hypotheses'))

    html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Rapport Tarification — ActuarIA</title>
{_plotly_cdn}
<style>
:root {{
  --navy:{NAVY}; --navy-mid:{NAVY_MID}; --gold:{GOLD}; --gold-l:{GOLD_L};
  --rouge:{ROUGE}; --vert:{VERT}; --orange:{ORANGE}; --slate:{SLATE};
  --bg:{BG}; --white:{WHITE}; --text:{TEXT};
}}
*{{box-sizing:border-box; margin:0; padding:0;}}
body{{font-family:'Segoe UI',Arial,sans-serif; background:var(--bg); color:var(--text); font-size:13px; line-height:1.6;}}
.page{{max-width:1060px; margin:0 auto; padding:32px 24px;}}
/* Header */
.header{{background:var(--navy); color:#fff; padding:28px 32px; border-radius:8px; margin-bottom:28px;}}
.header h1{{font-size:22px; color:var(--gold); letter-spacing:.5px;}}
.header .sub{{font-size:13px; color:var(--gold-l); margin-top:4px;}}
.header .meta{{font-size:11px; color:#8a9bb0; margin-top:10px; display:flex; gap:24px; flex-wrap:wrap;}}
.badge{{display:inline-block; padding:3px 10px; border-radius:4px; font-size:11px; font-weight:700; color:#fff;}}
.badge-vert{{background:var(--vert);}} .badge-ambre{{background:var(--orange);}} .badge-rouge{{background:var(--rouge);}}
/* Une hypothese NON CALCULEE se distingue d'une hypothese verte : gris,
   sans couleur de statut, parce qu'elle n'en a pas. */
.badge-non-calcule{{display:inline-block; padding:3px 10px; border-radius:4px;
  font-size:11px; font-weight:700; color:var(--slate);
  border:1px dashed var(--slate); background:transparent;}}
/* Sections */
.section{{background:var(--white); border-radius:8px; margin-bottom:20px; box-shadow:0 1px 4px rgba(0,0,0,.07); overflow:hidden;}}
.section-head{{background:var(--navy-mid); color:var(--gold); padding:10px 18px; font-size:13px; font-weight:700;}}
.section-body{{padding:16px 18px;}}
/* Tables */
table{{width:100%; border-collapse:collapse; font-size:12px; margin-top:8px;}}
th{{background:var(--navy); color:var(--gold); padding:7px 10px; text-align:left;}}
td{{padding:6px 10px; border-bottom:1px solid #e8edf2;}}
tr:nth-child(even) td{{background:#f7f9fc;}}
.right{{text-align:right;}}
.center{{text-align:center;}}
/* Narration */
.narration{{line-height:1.7; font-size:12.5px;}}
.narration h3.s-head{{color:var(--navy); font-size:14px; margin:18px 0 6px; border-left:3px solid var(--gold); padding-left:8px;}}
.narration h4{{color:var(--navy-mid); font-size:13px; margin:10px 0 4px;}}
.narration p{{margin:4px 0;}}
.narration ul{{margin:4px 0 4px 16px;}}
/* Footer */
.footer{{text-align:center; font-size:10px; color:var(--slate); margin-top:28px; padding:10px;}}

/* ══════════════════════════════════════════════════════════════════════
   L'IMPRESSION — le rapport se lit A L'ECRAN et se signe SUR PAPIER
   ══════════════════════════════════════════════════════════════════════
   ⚠️ MESURE AVANT CE LOT : ZERO regle d'impression dans ce rapport, contre
   vingt et une dans celui du provisionnement. Un livrable destine a etre
   imprime, converti en PDF et archive n'en portait aucune.
*/
@media print {{
  /* ⚠️ LA REGLE QUI EMPECHE UN ROUGE DE S'IMPRIMER COMME UN VERT. Un
     navigateur SUPPRIME par defaut les fonds colores a l'impression — la
     case << Graphiques d'arriere-plan >> du dialogue Ctrl+P est DECOCHEE.
     Or le signal RAG de ce rapport ne vit QUE dans le fond des badges :
     << VERT >>, << AMBRE >> et << ROUGE >> s'impriment sinon en blanc sur
     blanc, avec le meme texte noir. C'est le seul defaut de ce chantier
     qui puisse TROMPER un lecteur plutot que le gener.
     Le prefixe -webkit- reste necessaire : Chrome et Edge, qui font la
     conversion PDF, ne lisent encore que lui dans certaines versions. */
  *, *::before, *::after {{
    -webkit-print-color-adjust: exact !important;
    print-color-adjust: exact !important;
  }}

  /* ⚠️ DES MARGES, PAS ZERO. Sans marge le contenu court jusqu'au bord
     physique de la feuille, zone que la plupart des imprimantes et des
     convertisseurs rognent — un tableau y perd sa derniere colonne. */
  @page {{ size: A4; margin: 14mm 12mm 16mm 12mm; }}

  body {{ background: white; }}
  .page {{ max-width: none; margin: 0; padding: 0; }}
  .section {{ box-shadow: none; }}

  /* CE QUI NE DOIT PAS ETRE COUPE EN DEUX PAR UNE FIN DE PAGE.
     Un tableau scinde perd son en-tete ; une FIGURE scindee ne veut plus
     rien dire ; une vignette coupee separe le chiffre de son libelle.
     `break-inside` est la propriete moderne, `page-break-inside` son
     ancetre : les deux sont posees parce que les moteurs d'impression
     n'ont pas tous migre. */
  table, .kpi, .kpi-grid, .figure, .plotly-graph-div {{
    break-inside: avoid; page-break-inside: avoid;
  }}
  tr {{ break-inside: avoid; page-break-inside: avoid; }}

  /* Un tableau qui franchit une page reimprime son en-tete. Sans cela, la
     suite du tableau des huit hypotheses arrive sans ses colonnes. */
  thead {{ display: table-header-group; }}

  /* UN TITRE NE RESTE JAMAIS SEUL EN BAS DE PAGE, separe de ce qu'il
     annonce — c'est la faute de mise en page la plus visible d'un
     rapport. */
  .section-head, .narration h3.s-head, .narration h4 {{
    break-after: avoid; page-break-after: avoid;
  }}

  /* Pas de ligne isolee en haut ou en bas de page dans la prose. */
  p, li {{ orphans: 3; widows: 3; }}

  /* Une figure a une largeur fixee en pixels par le navigateur ; sur une
     feuille A4 elle deborderait de la zone imprimable. */
  .figure img, .figure svg, .plotly-graph-div {{
    max-width: 100% !important; height: auto !important;
  }}
}}
.kpi-grid{{display:grid; grid-template-columns:repeat(auto-fill,minmax(160px,1fr)); gap:10px; margin-top:8px;}}
.kpi{{background:var(--bg); border:1px solid #dde4ee; border-radius:6px; padding:10px 12px;}}
.kpi-label{{font-size:10px; color:var(--slate); text-transform:uppercase; letter-spacing:.4px;}}
.kpi-value{{font-size:18px; font-weight:700; color:var(--navy); margin-top:2px;}}
</style>
</head>
<body>
<div class="page">

<!-- HEADER -->
<div class="header">
  <div style="display:flex; justify-content:space-between; align-items:flex-start; flex-wrap:wrap; gap:12px;">
    <div>
      <h1>📊 Rapport de Tarification Non-Vie</h1>
      <div class="sub">{branche.replace('_',' ').title()} · Arrêté {arr}</div>
    </div>
    <div style="text-align:right;">
      <span class="badge badge-{statut.lower()}">{s_emoji} {statut}</span>
      <div class="sub" style="margin-top:6px;">{ref_client or 'Client à renseigner'}</div>
    </div>
  </div>
  <div class="meta">
    <span>🏢 ActuarIA v1.0</span>
    <span>📅 {now}</span>
    <span>🔑 {audit_id or 'N/A'}</span>
    <span>🧠 Narration : {n_src}</span>
  </div>
</div>

{_ouvrir_chapitre(1)}    <table>
      {_row(titres('glm'), header=True, num=colonnes_numeriques('glm'))}
"""
    for modele in ['poisson', 'gamma', 'tweedie']:
        m = met3.get(modele, {})
        if m:
            html += _row([
                nom_modele(modele),
                F.nombre(m.get('gini'), F.DEC_GINI),
                F.nombre(m.get('aic'), 0),
                F.nombre(m.get('deviance'), 2),
                F.nombre(m.get('pseudo_r2'), F.DEC_GINI),
                F.nombre(m.get('nb_vars_retenues'), 0),
            ], num=colonnes_numeriques('glm'))
    html += _fermer_chapitre(1)
    html += _ouvrir_chapitre(2) + '    <table>\n      ' + _row(
        titres('relativites'), header=True,
        num=colonnes_numeriques('relativites')) + '\n'
    for var, d in sorted(rels.items(), key=lambda x: -abs(x[1].get('beta',0))):
        sens_col = f'color:{VERT}' if d.get('sens')=='allegant' else f'color:{ROUGE}'
        html += _row([
            var,
            F.nombre(d.get('beta'), F.DEC_GINI),
            f"<strong>{F.nombre(d.get('relativite'), F.DEC_GINI)}</strong>",
            F.nombre(d.get('ic95_low'), F.DEC_GINI),
            F.nombre(d.get('ic95_high'), F.DEC_GINI),
            # Audit V7 IMPORTANT : garde NA — était toujours affiché "0.0000"
            # (indiscernable d'une vraie p-value nulle) si la clé était absente.
            F.nombre(d.get('pvalue'), F.DEC_GINI),
            '✓' if d.get('significatif') else '·',
            f'<span style="{sens_col};font-weight:600">{d.get("sens","")}</span>',
        ], num=colonnes_numeriques('relativites'))
    html += _fermer_chapitre(2)
    html += _ouvrir_chapitre(3) + '    <table>\n      ' + _row(
        titres('classement') + ['⭐'], header=True,
        num=colonnes_numeriques('classement')) + '\n'
    for rank, m in enumerate(cl4, 1):
        # ⚠️ L'ÉTOILE EST LA TROISIÈME EXCEPTION DÉCLARÉE : décorative, propre
        # à l'HTML, elle ne dit rien que « Rang 1 » ne dise. Le Word met la
        # même ligne en gras.
        star = '⭐' if rank == 1 else ''
        style = 'font-weight:700; background:#f0f7ff;' if rank == 1 else ''
        html += f'<tr style="{style}">'
        html += f'<td>{rank}</td><td>{nom_modele(m.get("modele"))}</td><td>{m.get("famille","")}</td>'
        html += f'<td class="right">{F.nombre(m.get("gini_test"), F.DEC_GINI)}</td>'
        html += f'<td class="right">{F.nombre(m.get("rmse_test"), F.DEC_GINI)}</td>'
        html += f'<td class="right">{F.nombre(m.get("overfit_ratio"), F.DEC_RATIO)}</td>'
        # Audit V7 IMPORTANT : garde NA — était toujours affiché "0.0000".
        html += f'<td class="right">{F.nombre(m.get("score_global"), F.DEC_GINI)}</td>'
        html += f'<td class="center">{star}</td></tr>\n'
    html += _fermer_chapitre(3)
    html += _ouvrir_chapitre(4) + '    <table>\n      ' + _row(
        titres('hypotheses'), header=True,
        num=colonnes_numeriques('hypotheses')) + '\n'
    # ⚠️ HUIT LIGNES POUR QUATRE NUMÉROS, sous un titre qui annonçait
    # « H1–H4 » : les quatre tests GLM s'appelaient « H1 » à « H4 » et les
    # quatre tests ML « H1 ML » à « H4 ML ». Elles sont désormais nommées à
    # un seul endroit, et le titre ne promet plus quatre hypothèses.
    for _cle, _libelle, _val in HYPOTHESES:
        html += _hyp_row(_cle, _libelle, _val)
    html += _fermer_chapitre(4)
    html += _ouvrir_chapitre(5)
    if bt6.get('disponible'):
        ae = bt6.get('ae_ratio', 0)
        ae_col = VERT if 0.95<=ae<=1.05 else ORANGE if 0.90<=ae<=1.10 else ROUGE
        html += f"""
    <div class="kpi-grid">
      <div class="kpi"><div class="kpi-label">A/E ratio (N-1→N)</div><div class="kpi-value" style="color:{ae_col}">{ae:.4f}</div></div>
      <div class="kpi"><div class="kpi-label">Interprétation</div><div class="kpi-value" style="font-size:13px;">{bt6.get('interpretation','—')}</div></div>
      <div class="kpi"><div class="kpi-label">Stabilité Walk-Forward</div><div class="kpi-value" style="font-size:13px;">{bt6.get('stabilite_wf','—')}</div></div>
      <div class="kpi"><div class="kpi-label">Fenêtres testées</div><div class="kpi-value">{bt6.get('n_fenetres','—')}</div></div>
      <div class="kpi"><div class="kpi-label">CV A/E WF</div><div class="kpi-value">{bt6.get('ae_cv_wf','—')}</div></div>
      <div class="kpi"><div class="kpi-label">Fenêtres ROUGE</div><div class="kpi-value" style="color:{ROUGE}">{bt6.get('n_fenetres_rouge','—')}</div></div>
    </div>
"""
        wf = bt6.get('walk_forward', [])
        if wf:
            html += """
    <table style="margin-top:12px;">
      """ + _row(titres('backtest'), header=True,
                 num=colonnes_numeriques('backtest')) + """
"""
            for w in wf:
                st_w = w.get('statut','')
                bg   = {'VERT':'#eaf3de','AMBRE':'#faeeda','ROUGE':'#fcebeb'}.get(st_w,'')
                html += f'<tr style="background:{bg};">'
                html += f'<td>{w.get("annee_test","")}</td>'
                # ⚠️ « {v:,} » RENDAIT « 2,435 » — le séparateur ANGLAIS,
                # dans un rapport français, à côté de « 10308 » sans aucun
                # séparateur. Trois conventions coexistaient dans ce document.
                html += f'<td class="right">{F.nombre(w.get("n_train"), 0)}</td>'
                html += f'<td class="right">{F.nombre(w.get("n_test"), 0)}</td>'
                html += f'<td class="right">{F.nombre(w.get("moy_train"), F.DEC_GINI)}</td>'
                html += f'<td class="right">{F.nombre(w.get("moy_test"), F.DEC_GINI)}</td>'
                html += f'<td class="right"><strong>{F.nombre(w.get("ae_ratio"), F.DEC_GINI)}</strong></td>'
                html += f'<td class="center"><span class="badge badge-{st_w.lower()}">{st_w}</span></td></tr>\n'
            html += '    </table>'
    else:
        html += f'<p><em>Backtesting non disponible. {bt6.get("note","")}</em></p>'
    html += _fermer_chapitre(5, avec_table=False)
    html += _ouvrir_chapitre(6)
    if prod:
        prod_col = _statut_col((result_a6 or {}).get('statut_rag','AMBRE'))
        _score_txt = f"{prod.get('score_global'):.4f}" if 'score_global' in prod else '—'
        html += f"""
    <div class="kpi-grid">
      <div class="kpi"><div class="kpi-label">Modèle retenu</div><div class="kpi-value" style="font-size:15px;color:{NAVY}">{nom_modele(prod.get('modele'))}</div></div>
      <div class="kpi"><div class="kpi-label">Famille</div><div class="kpi-value" style="font-size:15px;">{prod.get('famille','—')}</div></div>
      <div class="kpi"><div class="kpi-label">Score global</div><div class="kpi-value">{_score_txt}</div></div>
      <div class="kpi"><div class="kpi-label">Gini test</div><div class="kpi-value">{F.nombre(prod.get('gini_test'), F.DEC_GINI)}</div></div>
      <div class="kpi"><div class="kpi-label">Overfit ratio</div><div class="kpi-value">{F.nombre(prod.get('overfit_ratio'), F.DEC_RATIO)}</div></div>
      <div class="kpi"><div class="kpi-label">Interprétabilité</div><div class="kpi-value">{F.nombre(prod.get('interpretabilite'), 2)}/1.0</div></div>
    </div>
    <p style="margin-top:8px; font-size:10px; color:{SLATE}; font-style:italic;">
      ✦ Le « Score global » est une normalisation RELATIVE au meilleur
      modèle du profil de pondération retenu (le meilleur modèle vaut
      toujours ≈ 1,0000) — ce n'est PAS une mesure de performance absolue
      en pourcentage. Réf. : audit V7, recommandation IMPORTANT #1.
    </p>
"""
    if val6:
        html += """
    <table style="margin-top:14px;">
      """ + _row(titres('controles'), header=True) + """
"""
        for ck, cl in CONTROLES_SELECTION:
            cv = val6.get(ck, {})
            st = cv.get('statut','?')
            html += _row([cl,
                          f'<span class="badge badge-{st.lower()}">{_statut_emoji(st)} {st}</span>',
                          F.tronque(cv.get('message'), 80)])
        html += '    </table>'
    html += _fermer_chapitre(6, avec_table=False)
    html += _ouvrir_chapitre(7, ' narration')
    html += narr_html
    html += f"""
    <p style="margin-top:12px; font-size:10px; color:{SLATE}; font-style:italic;">
      ✦ {f'Narration générée par ActuarIA Intelligence ({CLAUDE_MODEL_TARIF})' if n_src=='claude_api' else f'Source : {n_src}'}
    </p>
  </div>
</div>

{_ouvrir_chapitre(8)}    <table>
      {_row(titres('audit'), header=True)}
"""
    # ⚠️ « Ae Ratio », « Stabilite Wf », « Gouvernance Ok : True » — la
    # substitution mécanique `.title()` produisait des noms de variables
    # capitalisés, pas du français.
    for k, v in at.items():
        if isinstance(v, (str, int, float, bool)):
            html += _row([libelle_audit(k), valeur_audit(v, k)])
    html += f"""    </table>
  </div>
</div>

<div class="footer">
  ActuarIA · {ref_client or 'Client'} · {branche.replace('_',' ').title()} · Arrêté {arr} · {now} · CONFIDENTIEL
</div>
</div>
</body>
</html>"""
    return html


# =============================================================================
#  WORD
# =============================================================================

def export_word(
    result_a3: Dict = None, result_a4: Dict = None, result_a6: Dict = None,
    ref_client: str = '', arrete: str = '', audit_id: str = '',
    narration_calculee: Optional[Tuple[str, str]] = None,
) -> bytes:
    """Génère le rapport Word tarification (.docx). Retourne bytes ou b''.

    `narration_calculee` : voir `export_html`. C'est ce format-ci qui payait le
    plus cher l'appel séparé — il partait sans le commentaire actuariel.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        logger.error(f'python-docx absent : {e}'); return b''

    try:
        now    = datetime.now().strftime('%d/%m/%Y')
        arr    = arrete or now
        branche= (result_a6 or result_a3 or {}).get('branche', 'non_vie')
        statut = (result_a6 or result_a3 or {}).get('statut_rag', 'AMBRE')
        met3   = (result_a3 or {}).get('metriques', {})
        rels   = (result_a3 or {}).get('relativites_poisson', {})
        # ⚠️ Même source que l'HTML : le classement SCORÉ d'A6 en premier.
        cl4    = ((result_a6 or {}).get('classement')
                  or (result_a4 or {}).get('classement', []))
        hyp3   = (result_a3 or {}).get('hypotheses', {})
        hyp4   = (result_a4 or {}).get('hypotheses', {})
        bt6    = (result_a6 or {}).get('backtest', {})
        prod   = (result_a6 or {}).get('modele_production', {})
        # ⚠️ CETTE SOURCE N'ÉTAIT PAS LUE PAR LE WORD — c'est pour cela que
        # les trois contrôles de sélection n'y figuraient pas. Même clé, même
        # repli que l'HTML.
        val6   = (result_a6 or {}).get('validation_selection', {})
        # ⚠️ LE MÊME PLAN QUE L'HTML, NUMÉROTÉ PAR LA MÊME FONCTION. Deux
        # compteurs indépendants — le choix d'A7 — se rattrapent par un test ;
        # une numérotation unique ne peut pas diverger.
        figures_numerotees = numeroter(
            figures_disponibles(result_a3, result_a4, result_a6))
        rendus_figures = {f.cle: rendre_figure(f.objet)
                          for fs in figures_numerotees.values() for f in fs}
        # Priorité A6 (le plus complet — gouvernance + WF recalibré) > A4 > A3
        at     = (
            (result_a6 or {}).get('audit_trail')
            or (result_a4 or {}).get('audit_trail')
            or (result_a3 or {}).get('audit_trail')
            or {}
        )

        narration, n_src = (
            narration_calculee if narration_calculee is not None
            else _narration_claude(result_a3, result_a4, result_a6, branche, arr))

        def rgb(h):
            h = h.lstrip('#')
            return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))

        NR = rgb(NAVY); GR = rgb(GOLD); BR = rgb('#FFFFFF')
        GrR= rgb(SLATE); RgR= rgb(ROUGE); VR = rgb(VERT); AR = rgb(ORANGE)

        doc = Document()
        for s in doc.sections:
            s.top_margin=Cm(2); s.bottom_margin=Cm(2)
            s.left_margin=Cm(2.5); s.right_margin=Cm(2.5)

        def _bg(cell, hex6):
            tc=cell._tc; tcp=tc.get_or_add_tcPr()
            sd=OxmlElement('w:shd')
            sd.set(qn('w:fill'), hex6.lstrip('#'))
            sd.set(qn('w:color'), 'auto'); sd.set(qn('w:val'), 'clear')
            tcp.append(sd)

        def _run(p, txt, bold=False, italic=False, sz=10, col=None):
            r=p.add_run(str(txt)); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
            if col: r.font.color.rgb=col
            return r

        def _h(txt, lv=1, col=None):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(3)
            sz={1:16,2:13,3:10}.get(lv,10)
            c=col or (NR if lv==1 else GR)
            _run(p, txt, bold=True, sz=sz, col=c)

        def _sep():
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
            bo=OxmlElement('w:bottom')
            bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'6')
            bo.set(qn('w:space'),'1'); bo.set(qn('w:color'),'C9A84C')
            pBdr.append(bo); pPr.append(pBdr)

        def _tbl(heads, rows, ws=None, num=(), gras_premiere=False):
            # ⚠️ MESURÉ AVANT CE LOT : 0 cellule numérique sur 69 alignée à
            # droite dans le Word. `num` = les index des colonnes numériques.
            t=doc.add_table(rows=1+len(rows), cols=len(heads)); t.style='Table Grid'
            for i, hd in enumerate(heads):
                c=t.rows[0].cells[i]; _bg(c, '0F2E52')
                pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=pp.add_run(str(hd)); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=BR
            for ri, row in enumerate(rows):
                for ci, v in enumerate(row):
                    c=t.rows[ri+1].cells[ci]
                    if ri%2==1: _bg(c, 'EEF2F7')
                    pp=c.paragraphs[0]
                    if ci in num:
                        pp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
                    r=pp.add_run(str(v) if v is not None else '—'); r.font.size=Pt(9)
                    # ⚠️ L'ÉTOILE DU PREMIER DU CLASSEMENT EST PROPRE À L'HTML
                    # — décorative, elle ne dit rien que « Rang 1 » ne dise.
                    # Le Word marque la même ligne en gras : même information,
                    # moyen propre au format. C'est une exception DÉCLARÉE.
                    if gras_premiere and ri == 0:
                        r.bold = True
            if ws:
                for i, w in enumerate(ws):
                    for row in t.rows: row.cells[i].width=Cm(w)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)

        def _figures_du_chapitre(numero):
            """Les figures du chapitre, à sa fin — MÊME plan que l'HTML.

            ⚠️ ELLES ÉTAIENT GROUPÉES DANS UNE SECTION « Figures » placée
            après le chapitre 6. Elles rejoignent le chapitre qu'elles
            illustrent, et le numéro qu'elles portent est celui que `numeroter`
            leur a donné avant tout rendu : il ne bouge pas d'un cran.
            """
            for f in figures_numerotees.get(numero, ()):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                _run(p, legende(f), bold=True, sz=9, col=GR)
                rendu = rendus_figures.get(f.cle)
                # ⚠️ UN .DOCX NE PORTE QUE DU RASTER : il prend le PNG dans
                # tous les cas, là où l'HTML choisit le plus léger des deux.
                if rendu is not None and rendu.png:
                    doc.add_picture(io.BytesIO(rendu.png), width=Cm(16.5))
                else:
                    q = doc.add_paragraph()
                    q.paragraph_format.space_after = Pt(6)
                    _run(q, (rendu.note if rendu is not None and rendu.note
                             else 'Figure non rendue dans ce format.'),
                         sz=8, italic=True, col=GrR)

        # ── PAGE DE GARDE ─────────────────────────────────────────────────────
        p=doc.add_paragraph()
        _run(p, 'Actuar', bold=True, sz=28, col=NR)
        _run(p, 'IA',     bold=True, sz=28, col=GR)
        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p, 'RAPPORT DE TARIFICATION NON-VIE\n', bold=True, sz=20, col=NR)
        _run(p, branche.replace('_',' ').title(), bold=True, sz=14, col=GR)
        doc.add_paragraph()
        s_col_r = VR if statut=='VERT' else AR if statut=='AMBRE' else RgR
        p=doc.add_paragraph()
        _run(p, 'Statut : ', sz=11, col=NR)
        _run(p, statut, bold=True, sz=11, col=s_col_r)
        doc.add_paragraph()
        _tbl(titres('garde'),
             [[ref_client or 'À renseigner', branche.replace('_',' ').title(), arr, audit_id or 'N/A']],
             ws=[3.5, 4.0, 3.0, 3.5])
        doc.add_page_break()

        # ── CHAPITRE 1 : SYNTHÈSE GLM ────────────────────────────────────────
        _h(chapitre(1)); _sep()
        rows_glm = []
        for modele in ['poisson','gamma','tweedie']:
            m = met3.get(modele, {})
            if m:
                rows_glm.append([
                    nom_modele(modele),
                    F.nombre(m.get('gini'), F.DEC_GINI),
                    F.nombre(m.get('aic'), 0),
                    # ⚠️ LA DÉVIANCE MANQUAIT AU WORD. C'est la mesure
                    # d'écart du GLM ; l'AIC seul ne la remplace pas.
                    F.nombre(m.get('deviance'), 2),
                    F.nombre(m.get('pseudo_r2'), F.DEC_GINI),
                    F.nombre(m.get('nb_vars_retenues'), 0),
                ])
        if rows_glm:
            _tbl(titres('glm'), rows_glm, ws=[2.8,2.3,2.3,2.6,2.3,3.0],
                 num=colonnes_numeriques('glm'))
        _figures_du_chapitre(1)
        doc.add_page_break()

        # ── CHAPITRE 2 : RELATIVITÉS ─────────────────────────────────────────
        _h(chapitre(2)); _sep()
        rows_rel = []
        for var, d in sorted(rels.items(),
                             key=lambda x: -abs(x[1].get('beta', 0))
                             )[:MAX_RELATIVITES_WORD]:
            rows_rel.append([
                var,
                F.nombre(d.get('beta'), F.DEC_GINI),
                F.nombre(d.get('relativite'), F.DEC_GINI),
                F.nombre(d.get('ic95_low'), F.DEC_GINI),
                F.nombre(d.get('ic95_high'), F.DEC_GINI),
                # Audit V7 IMPORTANT : garde NA cohérent avec le HTML.
                f"{d.get('pvalue'):.4f}" if 'pvalue' in d else '—',
                'Oui' if d.get('significatif') else 'Non',
                d.get('sens',''),
            ])
        if rows_rel:
            _tbl(titres('relativites'), rows_rel,
                 ws=[3.0,1.6,2.4,1.9,1.9,1.7,1.4,1.6],
                 num=colonnes_numeriques('relativites'))
        _note = note_troncature(MAX_RELATIVITES_WORD, len(rels),
                                'relativités', 'triées par |β| décroissant')
        if _note:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _run(p, '✦ ' + _note, sz=8, italic=True, col=GrR)
        _figures_du_chapitre(2)
        doc.add_page_break()

        # ── CHAPITRE 3 : CLASSEMENT ML ───────────────────────────────────────
        _h(chapitre(3)); _sep()
        # Audit V7 IMPORTANT : garde NA cohérent avec le HTML.
        # ⚠️ LE RMSE TEST MANQUAIT AU WORD — il pèse 10 % de la grille que le
        # titre du chapitre annonce désormais dans les deux formats.
        rows_ml = [[str(i), nom_modele(m.get('modele')), m.get('famille',''),
                    F.nombre(m.get('gini_test'), F.DEC_GINI),
                    F.nombre(m.get('rmse_test'), F.DEC_GINI),
                    F.nombre(m.get('overfit_ratio'), F.DEC_RATIO),
                    F.nombre(m.get('score_global'), F.DEC_GINI)]
                   for i, m in enumerate(cl4[:MAX_CLASSEMENT_WORD], 1)]
        if rows_ml:
            _tbl(titres('classement'), rows_ml,
                 ws=[1.3,3.6,2.8,2.3,2.3,1.9,2.3],
                 num=colonnes_numeriques('classement'), gras_premiere=True)
        _note = note_troncature(MAX_CLASSEMENT_WORD, len(cl4), 'modèles',
                                'dans l\'ordre du classement')
        if _note:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _run(p, '✦ ' + _note, sz=8, italic=True, col=GrR)
        _figures_du_chapitre(3)
        doc.add_page_break()

        # ── CHAPITRE 4 : LES HUIT HYPOTHÈSES ─────────────────────────────────
        _h(chapitre(4)); _sep()
        rows_hyp = []
        for hkey, hlabel, hval in HYPOTHESES:
            h = hyp3.get(hkey) or hyp4.get(hkey) or {}
            if not h:
                # ⚠️ LE WORD LA FAISAIT DISPARAÎTRE EN SILENCE, là où l'HTML
                # la publie « NON CALCULÉE » depuis T4. Une hypothèse absente
                # doit se voir dans les deux formats — sinon le livrable qui
                # part chez un commissaire est le plus indulgent des deux.
                rows_hyp.append([hlabel, '○ NON CALCULÉE', F.ABSENT, F.ABSENT,
                                 'Hypothèse non produite par la chaîne — '
                                 'à instruire avant validation'])
                continue
            brut = h.get(hval)
            # ⚠️ LA VALEUR MESURÉE MANQUAIT AU WORD : il donnait le verdict
            # sans le chiffre qui le fonde. Même précision que l'HTML.
            valeur = (F.nombre(brut, F.DEC_GINI)
                      if isinstance(brut, (int, float))
                      and not isinstance(brut, bool)
                      else (str(brut)[:60] if brut else F.ABSENT))
            rows_hyp.append([hlabel, h.get('statut','?'), valeur,
                             F.tronque(h.get('message'), 70),
                             F.tronque(h.get('conseil'), 60)])
        if rows_hyp:
            _tbl(titres('hypotheses'), rows_hyp,
                 ws=[4.2,1.7,1.6,4.5,4.0],
                 num=colonnes_numeriques('hypotheses'))
        _figures_du_chapitre(4)
        doc.add_page_break()

        # ── CHAPITRE 5 : BACKTESTING ─────────────────────────────────────────
        _h(chapitre(5)); _sep()
        if bt6.get('disponible'):
            ae = bt6.get('ae_ratio', 0)
            bt_col = VR if 0.95<=ae<=1.05 else AR if 0.90<=ae<=1.10 else RgR
            p=doc.add_paragraph()
            _run(p, f"A/E ratio (N-1→N) : ", sz=10, col=NR)
            _run(p, f"{ae:.4f}", bold=True, sz=10, col=bt_col)
            _run(p, f" | {bt6.get('interpretation','')} | Stabilité : {bt6.get('stabilite_wf','')}",
                 sz=9, col=NR)
            wf = bt6.get('walk_forward', [])
            if wf:
                # ⚠️ TROIS COLONNES MANQUAIENT — l'effectif de test et les
                # deux moyennes. Le ratio A/E est leur QUOTIENT : sans elles
                # le lecteur ne peut ni le recalculer ni juger s'il repose sur
                # 2 371 contrats ou sur 40. Le tableau n'occupait que 10,0 cm
                # sur les 16,5 utiles : la place n'était pas en cause.
                rows_wf = [[str(w.get('annee_test','')),
                             F.nombre(w.get('n_train'), 0),
                             F.nombre(w.get('n_test'), 0),
                             F.nombre(w.get('moy_train'), F.DEC_GINI),
                             F.nombre(w.get('moy_test'), F.DEC_GINI),
                             F.nombre(w.get('ae_ratio'), F.DEC_GINI),
                             w.get('statut','')]
                           for w in wf]
                _tbl(titres('backtest'), rows_wf,
                     ws=[2.2,2.4,2.4,2.4,2.4,2.4,2.3],
                     num=colonnes_numeriques('backtest'))
        else:
            p=doc.add_paragraph()
            _run(p, f"Backtesting non disponible. {bt6.get('note','')}", sz=9, italic=True)
        _figures_du_chapitre(5)
        doc.add_page_break()

        # ── CHAPITRE 6 : MODÈLE DE PRODUCTION ────────────────────────────────
        _h(chapitre(6)); _sep()
        if prod:
            p_col = VR if statut=='VERT' else AR if statut=='AMBRE' else RgR
            _score_txt = f"{prod.get('score_global'):.4f}" if 'score_global' in prod else '—'
            _tbl(titres('production'),
                 [['Modèle retenu', nom_modele(prod.get('modele')),
                   'Famille', prod.get('famille','—')],
                  ['Score global', _score_txt,
                   'Gini test', f"{prod.get('gini_test',0):.4f}"],
                  ['Overfit ratio', f"{prod.get('overfit_ratio',0):.3f}",
                   'Interprétabilité', f"{prod.get('interpretabilite',0):.2f}/1.0"]],
                 ws=[4.0,4.0,4.0,4.0])
            # Audit V7 IMPORTANT #1 : qualification du score composite —
            # normalisation RELATIVE au meilleur modèle du profil retenu
            # (le meilleur vaut toujours ≈ 1,0000), PAS une performance
            # absolue. Absente du livrable avant ce correctif (ne figurait
            # que dans un commentaire de code, jamais vue par le client).
            p = doc.add_paragraph()
            _run(p,
                 "✦ Le « Score global » est une normalisation relative au "
                 "meilleur modèle du profil de pondération retenu (le "
                 "meilleur modèle vaut toujours ≈ 1,0000) — ce n'est pas "
                 "une mesure de performance absolue en pourcentage.",
                 sz=9, italic=True)
        # ⚠️ CE TABLEAU N'EXISTAIT PAS DANS LE WORD. Les trois contrôles
        # justifient le choix du modèle de production : le fichier qui part
        # chez un commissaire aux comptes ne les portait pas.
        if val6:
            rows_ct = []
            for ck, cl in CONTROLES_SELECTION:
                cv = val6.get(ck, {})
                rows_ct.append([cl, cv.get('statut','?'),
                                F.tronque(cv.get('message'), 80)])
            _tbl(titres('controles'), rows_ct, ws=[5.0,2.5,9.0])
        _figures_du_chapitre(6)
        doc.add_page_break()

        # ── CHAPITRE 7 : COMMENTAIRE ACTUARIEL ───────────────────────────────
        _h(chapitre(7)); _sep()
        if narration:
            # ⚠️ CE BLOC N'INTERPRÉTAIT AUCUN MARKDOWN. Il découpait sur « §N »
            # et écrivait chaque ligne TELLE QUELLE : mesuré, 7 formes sur 7
            # ressortaient brutes — « # », « ## », « **gras** », « - », « | »…
            # Il consomme désormais la MÊME analyse que l'HTML, si bien que les
            # deux formats ne peuvent plus rendre des structures différentes.
            for bloc in _md.analyser(_clean(narration)):
                if bloc.genre == _md.SECTION or bloc.genre == _md.TITRE1:
                    _h(bloc.texte, lv=2)
                    continue
                if bloc.genre == _md.REGLE:
                    _sep()
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(
                    0.6 if bloc.genre in (_md.PUCE, _md.CITATION) else 0.3)
                if bloc.genre == _md.TITRE2:
                    _run(p, bloc.texte, bold=True, sz=10, col=NR)
                    continue
                if bloc.genre == _md.TITRE3:
                    _run(p, bloc.texte, bold=True, sz=9, col=NR)
                    continue
                if bloc.genre == _md.PUCE:
                    _run(p, '• ', sz=9, col=NR)
                elif bloc.genre == _md.CITATION:
                    _run(p, '« ', sz=9, italic=True, col=GrR)
                # Le gras et l'italique deviennent des passages, pas des
                # étoiles : un `.docx` n'a pas de balises, il a des runs.
                for seg in _md.segments(bloc.texte):
                    _run(p, seg.texte, bold=seg.gras, italic=seg.italique,
                         sz=9, col=NR)
                if bloc.genre == _md.CITATION:
                    _run(p, ' »', sz=9, italic=True, col=GrR)
            # ⚠️ LE WORD NE NOMMAIT SA SOURCE QU'EN CAS DE SUCCÈS. Un repli y
            # était donc INDISCERNABLE d'une narration réussie, là où l'HTML
            # l'affichait. Une narration calculée une fois qui échoue doit le
            # dire dans les DEUX formats — sinon on remplace deux commentaires
            # divergents par un silence dans l'un des deux.
            p = doc.add_paragraph()
            _run(p, (f'✦ Narration générée par ActuarIA Intelligence '
                     f'({CLAUDE_MODEL_TARIF})' if n_src == SRC_API
                     else f'✦ Source de la narration : {n_src}'),
                 sz=7, italic=True, col=GrR)
        else:
            p=doc.add_paragraph()
            _run(p, f'Narration non disponible — source : {n_src}',
                 sz=9, italic=True)
        doc.add_page_break()

        # ── CHAPITRE 8 : PISTE D'AUDIT ───────────────────────────────────────
        _h(chapitre(8)); _sep()
        rows_at = [[libelle_audit(k), valeur_audit(v, k)]
                   for k, v in at.items() if isinstance(v, (str,int,float,bool))]
        if rows_at:
            _tbl(titres('audit'), rows_at, ws=[5.5,10.5])

        _sep()
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f'ActuarIA · {ref_client or "Client"} · {branche.replace("_"," ").title()} · '
                f'Arrêté {arr} · {now} · CONFIDENTIEL',
             sz=7, italic=True, col=GrR)

        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        wb = buf.read()
        logger.info(f'Word tarification : {len(wb):,} bytes')
        return wb

    except Exception as e:
        logger.error(f'export_word tarification : {e}', exc_info=True)
        return b''


# =============================================================================
#  PDF
# =============================================================================

def export_pdf(
    result_a3: Dict = None, result_a4: Dict = None, result_a6: Dict = None,
    ref_client: str = '', arrete: str = '', audit_id: str = '',
    narration_calculee: Optional[Tuple[str, str]] = None,
) -> bytes:
    """Génère le rapport PDF via weasyprint (HTML→PDF). Retourne bytes ou b''.

    `narration_calculee` : transmis à `export_html` — le PDF est un rendu de
    l'HTML, il doit porter exactement le même commentaire.
    """
    try:
        from weasyprint import HTML as WP_HTML
        html = export_html(result_a3, result_a4, result_a6, ref_client, arrete,
                           audit_id, narration_calculee)
        if not html:
            return b''
        buf = io.BytesIO()
        WP_HTML(string=html).write_pdf(buf)
        buf.seek(0)
        pdf = buf.read()
        logger.info(f'PDF tarification : {len(pdf):,} bytes')
        return pdf
    except ImportError:
        logger.warning('weasyprint absent — PDF non généré')
        return b''
    except Exception as e:
        logger.error(f'export_pdf tarification : {e}', exc_info=True)
        return b''


# =============================================================================
#  POINT D'ENTRÉE UNIFIÉ
# =============================================================================

def generer_rapport_tarification(
    result_a3: Dict = None,
    result_a4: Dict = None,
    result_a6: Dict = None,
    ref_client: str = '',
    arrete: str = '',
    audit_id: str = '',
    formats: List[str] = None,
) -> Dict[str, bytes]:
    """
    Génère tous les formats demandés en un seul appel.
    formats = ['html','word','pdf','excel'] (défaut = html + word)
    Retourne {'html_bytes': ..., 'word_bytes': ..., 'pdf_bytes': ..., 'excel_bytes': ...}
    """
    if formats is None:
        formats = ['html', 'word']

    out: Dict[str, bytes] = {
        'html_bytes':  b'',
        'word_bytes':  b'',
        'pdf_bytes':   b'',
        'excel_bytes': b'',
    }

    # ⚠️ LA NARRATION EST CALCULÉE UNE FOIS, ICI, POUR TOUS LES FORMATS.
    # Avant ce lot, chaque export appelait `_narration_claude` pour son propre
    # compte. Mesuré sur un rapport réel : le HTML portait les 18 089 caractères
    # du commentaire actuariel, le Word portait le dépôt technique de l'agent —
    # le format qui part chez un commissaire n'avait pas le commentaire. Deux
    # narrations divergentes dans un même livrable deviennent désormais
    # IMPOSSIBLES PAR CONSTRUCTION, et le nombre d'appels est divisé par deux.
    arr_narr = arrete or datetime.now().strftime('%d/%m/%Y')
    branche_narr = (result_a6 or result_a3 or {}).get('branche', 'non_vie')
    narration_calculee = _narration_claude(
        result_a3, result_a4, result_a6, branche_narr, arr_narr)

    # HTML d'abord (réutilisé pour PDF)
    html_str = ''
    if 'html' in formats or 'pdf' in formats:
        html_str = export_html(result_a3, result_a4, result_a6, ref_client, arrete,
                               audit_id, narration_calculee)
        out['html_bytes'] = html_str.encode('utf-8') if html_str else b''

    if 'word' in formats:
        out['word_bytes'] = export_word(result_a3, result_a4, result_a6, ref_client,
                                        arrete, audit_id, narration_calculee)

    if 'pdf' in formats:
        if html_str:
            try:
                from weasyprint import HTML as WP_HTML
                buf = io.BytesIO()
                WP_HTML(string=html_str).write_pdf(buf)
                buf.seek(0)
                out['pdf_bytes'] = buf.read()
            except ImportError:
                logger.warning('weasyprint absent — PDF non généré')
        else:
            out['pdf_bytes'] = export_pdf(result_a3, result_a4, result_a6, ref_client,
                                          arrete, audit_id, narration_calculee)

    if 'excel' in formats:
        try:
            from .tarif_excel import export_excel_a6
            if result_a6:
                out['excel_bytes'] = export_excel_a6(result_a6, audit_id)
        except ImportError:
            try:
                from direction_non_vie.tarification.services.tarif_excel import export_excel_a6
                if result_a6:
                    out['excel_bytes'] = export_excel_a6(result_a6, audit_id)
            except ImportError:
                pass

    logger.info(
        f"Rapport tarification : "
        f"HTML={len(out['html_bytes']):,}b "
        f"Word={len(out['word_bytes']):,}b "
        f"PDF={len(out['pdf_bytes']):,}b "
        f"Excel={len(out['excel_bytes']):,}b"
    )
    return out
