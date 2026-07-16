"""
ActuarIA — Rapport Consolidé Équipe
Direction Non-Vie · Équipe Tarification — Vue d'ensemble A1 → A6

4 formats : HTML · Word (.docx) · PDF (weasyprint) · Excel (openpyxl)
Consolide les résultats des 6 agents en un rapport de synthèse unique,
destiné à l'actuaire responsable et à la Direction pour validation finale.

Structure du rapport (7 sections) :
  §1 — Page de garde : statuts RAG des 6 agents en un coup d'œil
  §2 — Qualité des données (A1)
  §3 — Preprocessing & Feature Engineering (A2)
  §4 — Tarification : GLM vs ML vs DL (A3/A4/A5) + Crédibilité + Lissage géo
  §5 — Décision finale (A6) : modèle retenu, backtesting, gouvernance
  §6 — Audit trail consolidé
  §7 — Commentaire actuariel de référence (réutilise le commentaire A6,
       aucun nouvel appel API Claude)

Auteur   : ActuarIA v1.0
Version  : 1.0.0
"""

from __future__ import annotations
import io, logging, re
from core.conformite_reglementaire import (
    avertissement_walk_forward, synthese_exclusions, synthese_alertes_experience,
    synthese_modele_dl,
)
from core.qualite_donnees import synthese_qualite_donnees
from core.plan_tarifaire import synthese_colonnes_plan_manquantes
from datetime import datetime
from typing import Dict, List, Optional, Any

logger = logging.getLogger('actuaria.tarif.rapport_equipe')

try:
    from openpyxl import Workbook
except ImportError:
    pass

# ── Helpers Excel partagés — audit V4 point #12 ──────────────────────────────
# Palette (_HEX) et fonctions factorisées dans excel_helpers.py (source
# unique, éliminait une duplication avec tarif_excel.py).
try:
    from .excel_helpers import (
        OPENPYXL_OK, NAVY_HEX, GOLD_HEX, VERT_HEX, AMBRE_HEX, ROUGE_HEX,
        GRIS_HEX, GRIS_L_HEX, NOIR_HEX, BLANC_HEX,
        FMT_EUR, FMT_PCT, FMT_DEC4, FMT_NB,
        _font, _fill, _border, _align, _statut_fill, _col_w,
        _cell, _header, _section, _kpi, _bandeau,
    )
except ImportError:
    from direction_non_vie.tarification.services.excel_helpers import (
        OPENPYXL_OK, NAVY_HEX, GOLD_HEX, VERT_HEX, AMBRE_HEX, ROUGE_HEX,
        GRIS_HEX, GRIS_L_HEX, NOIR_HEX, BLANC_HEX,
        FMT_EUR, FMT_PCT, FMT_DEC4, FMT_NB,
        _font, _fill, _border, _align, _statut_fill, _col_w,
        _cell, _header, _section, _kpi, _bandeau,
    )

# ── Palette ActuarIA — variantes HTML/Word (avec '#', non factorisées :
# ── absentes de tarif_excel.py, propres à ce module) ─────────────────────────
NAVY      = '#0F2E52'
NAVY_MID  = '#1B3A5C'
GOLD      = '#C9A84C'
GOLD_L    = '#E2C97E'
ROUGE     = '#C0392B'
VERT      = '#1E8449'
ORANGE    = '#E67E22'
SLATE     = '#8A9BB0'
BG        = '#F5F7FA'
WHITE     = '#FFFFFF'
TEXT      = '#1C2B3A'
FMT_NB   = '# ##0'


# =============================================================================
#  HELPERS COMMUNS
# =============================================================================

def _statut_col(s: str) -> str:
    return {'VERT': VERT, 'AMBRE': ORANGE, 'ROUGE': ROUGE}.get((s or '').upper(), SLATE)

def _statut_emoji(s: str) -> str:
    return {'VERT': '🟢', 'AMBRE': '🟡', 'ROUGE': '🔴'}.get((s or '').upper(), '⚪')

def _s(v) -> str:
    if v is None:
        return '—'
    return re.sub(r'\s+', ' ', str(v)).strip() or '—'

def _statut_global(statuts: List[str]) -> str:
    """Statut consolidé : ROUGE si au moins 1 ROUGE, AMBRE si au moins 1 AMBRE, sinon VERT."""
    statuts_up = [s.upper() for s in statuts if s]
    if 'ROUGE' in statuts_up:
        return 'ROUGE'
    if 'AMBRE' in statuts_up:
        return 'AMBRE'
    if statuts_up:
        return 'VERT'
    return 'AMBRE'


def _collecter_statuts(results: Dict[str, Dict]) -> Dict[str, str]:
    """Extrait le statut_rag de chaque agent présent dans results."""
    return {
        agent: r.get('statut_rag', 'N/A')
        for agent, r in results.items()
        if r and r.get('success')
    }


def _md_to_html_light(txt: str) -> str:
    """
    Conversion Markdown minimal → HTML — réutilisée pour afficher le
    commentaire actuariel déjà généré par A6 (rapport_modeles_tarif.py),
    SANS effectuer de nouvel appel à l'API Claude.
    """
    import re as _re
    if not txt:
        return ''
    txt = _re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', txt)
    txt = _re.sub(r'\*(.+?)\*', r'<em>\1</em>', txt)
    txt = _re.sub(r'^§(\d+)\s*[—\-–]\s*(.+)$', r'<h4 class="s-head">§\1 — \2</h4>', txt, flags=_re.MULTILINE)
    txt = _re.sub(r'^###\s+(.+)$', r'<h5>\1</h5>', txt, flags=_re.MULTILINE)
    txt = _re.sub(r'^-\s+(.+)$', r'<li>\1</li>', txt, flags=_re.MULTILINE)
    txt = _re.sub(r'(<li>.*</li>\n?)+', r'<ul>\g<0></ul>', txt)
    paragraphs = []
    for line in txt.split('\n'):
        line = line.strip()
        if not line or line.startswith('<h') or line.startswith('<ul') or line.startswith('<li'):
            paragraphs.append(line)
        else:
            paragraphs.append(f'<p>{line}</p>')
    return '\n'.join(paragraphs)


# =============================================================================
#  EXCEL — Rapport consolidé équipe
# =============================================================================

def export_excel_equipe(results: Dict[str, Dict], branche: str = '',
                          arrete: str = '', audit_id: str = '') -> bytes:
    """
    Génère le rapport Excel consolidé équipe (6 onglets).
    results : dict avec clés 'a1','a2','a3','a4','a5','a6' → dict de retour de chaque agent.
    Retourne bytes ou b''.
    """
    if not OPENPYXL_OK:
        return b''
    try:
        r1 = results.get('a1') or {}
        r2 = results.get('a2') or {}
        r3 = results.get('a3') or {}
        r4 = results.get('a4') or {}
        r5 = results.get('a5') or {}
        r6 = results.get('a6') or {}

        now = datetime.now().strftime("%d/%m/%Y %H:%M")
        arr = arrete or now
        branche_f = branche or (r6 or r3 or r1).get('branche', 'non_vie')
        statuts = _collecter_statuts(results)
        statut_global = _statut_global(list(statuts.values()))

        # Fonctions _font/_fill/_border/_align/_statut_fill/_col_w/_cell/
        # _header/_section/_kpi/_bandeau importées de excel_helpers.py
        # (audit V4 point #12 — anciennement redéfinies ici à l'identique).

        wb = Workbook()

        # ── Onglet 1 : Vue d'ensemble équipe ──────────────────────────────────
        ws1 = wb.active
        ws1.title = "1-Vue d'ensemble"
        _bandeau(ws1, "Rapport Consolidé Équipe Tarification", "Synthèse A1 → A6",
                 "Équipe Tarification Non-Vie", audit_id, now)
        r = 7
        _section(ws1, r, "▶ STATUT GLOBAL ÉQUIPE"); r += 1
        _kpi(ws1, r, "Statut consolidé", statut_global, statut=statut_global); r += 1
        _kpi(ws1, r, "Branche", branche_f); r += 1
        _kpi(ws1, r, "Arrêté", arr); r += 1
        r += 1
        _section(ws1, r, "▶ STATUT PAR AGENT"); r += 1
        for col, txt, w in [(1, "Agent", 30), (2, "Statut", 16), (3, "Rôle", 40)]:
            _header(ws1, r, col, txt, width=w)
        r += 1
        agents_roles = {
            'a1': 'A1 — Ingestion & Validation qualité',
            'a2': 'A2 — Preprocessing & Feature Engineering',
            'a3': 'A3 — GLM Poisson/Gamma/Tweedie + Crédibilité + Lissage géo',
            'a4': 'A4 — Machine Learning ×8 modèles',
            'a5': 'A5 — Deep Learning (CANN/TabNet)',
            'a6': 'A6 — Comparaison, sélection & gouvernance',
        }
        for agent_key, role in agents_roles.items():
            st = statuts.get(agent_key, 'N/A')
            bg = _statut_fill(st) if st != 'N/A' else GRIS_HEX
            _cell(ws1, r, 1, agent_key.upper(), cf=NOIR_HEX, fill=GRIS_L_HEX, bold=True)
            _cell(ws1, r, 2, st, cf=BLANC_HEX, fill=bg, ah="center")
            _cell(ws1, r, 3, role, cf=NOIR_HEX)
            r += 1

        # ── Onglet 2 : Qualité des données (A1) ───────────────────────────────
        ws2 = wb.create_sheet("2-Qualité Données (A1)")
        _bandeau(ws2, "Qualité des Données", "Agent A1 — Ingestion & Validation",
                 "A1", audit_id, now)
        r = 7
        qualite = r1.get('qualite', {})
        _section(ws2, r, "▶ SCORE QUALITÉ"); r += 1
        _kpi(ws2, r, "Score global", round(qualite.get('score_global', 0), 1),
             statut=r1.get('statut_rag'), fmt=FMT_DEC4); r += 1
        _kpi(ws2, r, "Nb lignes", qualite.get('nb_lignes', 0), fmt=FMT_NB); r += 1
        _kpi(ws2, r, "Taux complétude", qualite.get('taux_completude', 0) / 100, fmt=FMT_PCT); r += 1
        _kpi(ws2, r, "Nb types d'anomalies", qualite.get('nb_types_aberrants', 0),
             statut="VERT" if qualite.get('nb_types_aberrants', 0) == 0 else "AMBRE",
             fmt=FMT_NB); r += 1
        r += 1
        _section(ws2, r, "▶ ALERTES ABERRANTS"); r += 1
        for alerte in qualite.get('alertes_aberrants', []):
            _cell(ws2, r, 1, f"• {alerte}", cf=NOIR_HEX, wrap=True)
            ws2.merge_cells(start_row=r, start_column=1, end_row=r, end_column=6)
            ws2.row_dimensions[r].height = 30
            r += 1
        if not qualite.get('alertes_aberrants'):
            _kpi(ws2, r, "Alertes", "Aucune", statut="VERT"); r += 1

        # ── Onglet 3 : Preprocessing (A2) ─────────────────────────────────────
        ws3 = wb.create_sheet("3-Preprocessing (A2)")
        _bandeau(ws3, "Preprocessing & Feature Engineering", "Agent A2",
                 "A2", audit_id, now)
        r = 7
        data_dict = r2.get('data_dictionnaire', {})
        _section(ws3, r, "▶ VARIABLES DÉRIVÉES DOCUMENTÉES"); r += 1
        _kpi(ws3, r, "Statut", r2.get('statut_rag', 'N/A'), statut=r2.get('statut_rag')); r += 1
        _kpi(ws3, r, "Nb variables tracées", len(data_dict), fmt=FMT_NB); r += 1
        r += 1
        for col, txt, w in [(1, "Variable", 26), (2, "Usage", 30)]:
            _header(ws3, r, col, txt, width=w)
        r += 1
        for var_nom, meta in data_dict.items():
            _cell(ws3, r, 1, var_nom, cf=NOIR_HEX, fill=GRIS_L_HEX, bold=True)
            _cell(ws3, r, 2, meta.get('usage', ''), cf=NOIR_HEX, wrap=True)
            r += 1

        # ── Onglet 4 : Tarification GLM/ML/DL (A3/A4/A5) ──────────────────────
        ws4 = wb.create_sheet("4-Tarification (A3-A5)")
        _bandeau(ws4, "Tarification — GLM vs ML vs DL", "Agents A3, A4, A5",
                 "A3/A4/A5", audit_id, now)
        r = 7
        met3 = r3.get('metriques', {})
        _section(ws4, r, "▶ GLM (A3)"); r += 1
        _kpi(ws4, r, "Poisson — Gini", round(met3.get('poisson', {}).get('gini', 0), 4), fmt=FMT_DEC4); r += 1
        _kpi(ws4, r, "Gamma — Gini", round(met3.get('gamma', {}).get('gini', 0), 4), fmt=FMT_DEC4); r += 1
        _kpi(ws4, r, "Tweedie — Gini", round(met3.get('tweedie', {}).get('gini', 0), 4), fmt=FMT_DEC4); r += 1
        r += 1
        cred = r3.get('credibilite', {})
        geo  = r3.get('lissage_geo', {})
        _section(ws4, r, "▶ CRÉDIBILITÉ BÜHLMANN-STRAUB"); r += 1
        _kpi(ws4, r, "Appliquée",
             "✓ Oui" if cred.get('appliquee') else "○ Non applicable",
             statut="VERT" if cred.get('appliquee') else "AMBRE"); r += 1
        if cred.get('appliquee'):
            _kpi(ws4, r, "Z moyen / k", f"{cred.get('z_moyen',0):.4f} / {cred.get('k',0):.4f}"); r += 1
        r += 1
        _section(ws4, r, "▶ LISSAGE GÉOGRAPHIQUE"); r += 1
        _kpi(ws4, r, "Appliqué",
             "✓ Oui" if geo.get('applique') else "○ Non applicable",
             statut="VERT" if geo.get('applique') else "AMBRE"); r += 1
        if geo.get('applique'):
            _kpi(ws4, r, "Méthode", geo.get('methode', '')); r += 1
        r += 1
        cl4 = r4.get('classement', [])
        _section(ws4, r, "▶ MACHINE LEARNING — CLASSEMENT (A4)"); r += 1
        if cl4:
            for col, txt, w in [(1,"Modèle",22),(2,"Gini",12),(3,"Score",12),(4,"Overfit",12)]:
                _header(ws4, r, col, txt, width=w)
            r += 1
            for m in cl4[:8]:
                _cell(ws4, r, 1, m.get('modele',''), cf=NOIR_HEX, fill=GRIS_L_HEX)
                _cell(ws4, r, 2, round(m.get('gini_test',0),4), cf=NOIR_HEX, ah="right", fmt=FMT_DEC4)
                # Audit V7 IMPORTANT : garde NA — 'N/A' si absent, jamais 0.0000.
                _cell(ws4, r, 3, round(m.get('score_global'),4) if 'score_global' in m else 'N/A',
                      cf=NOIR_HEX, ah="right", fmt=FMT_DEC4 if 'score_global' in m else None)
                _cell(ws4, r, 4, round(m.get('overfit_ratio',0),3), cf=NOIR_HEX, ah="right", fmt=FMT_DEC4)
                r += 1
        r += 1
        _section(ws4, r, "▶ DEEP LEARNING (A5)"); r += 1
        _kpi(ws4, r, "Statut", r5.get('statut_rag', 'N/A') if r5 else 'Non exécuté',
             statut=r5.get('statut_rag') if r5 else None); r += 1

        # ── Onglet 5 : Décision finale (A6) ───────────────────────────────────
        ws5 = wb.create_sheet("5-Décision Finale (A6)")
        _bandeau(ws5, "Décision Finale", "Agent A6 — Comparaison & Gouvernance",
                 "A6", audit_id, now)
        r = 7
        prod = r6.get('modele_production', {})
        bt6  = r6.get('backtest', {})
        at6  = r6.get('audit_trail', {})
        _section(ws5, r, "▶ MODÈLE DE PRODUCTION RETENU"); r += 1
        _kpi(ws5, r, "Modèle", prod.get('modele', 'N/A')); r += 1
        _kpi(ws5, r, "Score global",
             round(prod.get('score_global'), 4) if 'score_global' in prod else 'N/A',
             fmt=FMT_DEC4 if 'score_global' in prod else None); r += 1
        _kpi(ws5, r, "  ↳ Nature du score",
             "Relatif au meilleur modèle du profil (≈1,0000 = meilleur) — "
             "pas une performance absolue.", wrap=True); r += 1
        _kpi(ws5, r, "Gini test", round(prod.get('gini_test', 0), 4), fmt=FMT_DEC4); r += 1
        r += 1
        _section(ws5, r, "▶ BACKTESTING WALK-FORWARD (recalibré)"); r += 1
        _kpi(ws5, r, "Modèle recalibré", bt6.get('modele_recalibre', 'N/A'),
             statut=("VERT" if bt6.get('modele_recalibre_fidele') else "AMBRE")); r += 1
        # Source UNIQUE partagée avec Excel A6 / Word / HTML (audit V11 — le
        # correctif V10 B3 n'avait été appliqué qu'à l'export Excel d'A6).
        _avert_wf6 = avertissement_walk_forward(bt6)
        if _avert_wf6:
            _kpi(ws5, r, "⚠ Portée de la validation", _avert_wf6,
                 statut="AMBRE", wrap=True); r += 1
        if bt6.get('gini_wf_moyen') is not None:
            _kpi(ws5, r, "Gini WF moyen", round(bt6.get('gini_wf_moyen', 0), 4), fmt=FMT_DEC4); r += 1
        _kpi(ws5, r, "A/E ratio", round(bt6.get('ae_ratio', 0), 4), fmt=FMT_DEC4); r += 1
        r += 1
        _synth_exc6 = synthese_exclusions(r6.get('exclusions_conformite')
                                          if isinstance(r6, dict) else None)
        if _synth_exc6:
            _kpi(ws5, r, "⚠ Colonnes écartées de la matrice X", _synth_exc6,
                 statut=("AMBRE" if "ACTION REQUISE" in _synth_exc6 else "VERT"),
                 wrap=True); r += 1
        _synth_alert6 = synthese_alertes_experience(
            r6.get('alertes_conformite') if isinstance(r6, dict) else None)
        if _synth_alert6:
            _kpi(ws5, r, "ℹ Sinistralité passée conservée — à vérifier",
                 _synth_alert6, statut="AMBRE", wrap=True); r += 1
        r += 1
        _section(ws5, r, "▶ GOUVERNANCE"); r += 1
        _kpi(ws5, r, "Profil retenu", at6.get('profil_ponderation', 'N/A')); r += 1
        _kpi(ws5, r, "Validé par", at6.get('profil_valide_par') or "⚠ NON VALIDÉ",
             statut="VERT" if at6.get('profil_valide_par') else "AMBRE"); r += 1
        _kpi(ws5, r, "Environnement", at6.get('environnement', 'N/A')); r += 1
        _kpi(ws5, r, "Conforme", "✓ Oui" if at6.get('gouvernance_ok') else "✗ Non",
             statut="VERT" if at6.get('gouvernance_ok') else "AMBRE"); r += 1
        # Validation humaine d'un modèle Deep Learning en production (2026-07) —
        # source UNIQUE partagée avec Excel A6 / Word / HTML. Ne s'affiche que si
        # le modèle retenu est un DL ; sinon rien (pas de bruit).
        _synth_dl6 = synthese_modele_dl(
            r6.get('modele_production') if isinstance(r6, dict) else None,
            r6.get('valide_par_actuaire_dl') if isinstance(r6, dict) else None,
            at6.get('timestamp'))
        if _synth_dl6:
            _kpi(ws5, r, "Modèle Deep Learning — validation actuarielle", _synth_dl6,
                 statut=("AMBRE" if "ACTION REQUISE" in _synth_dl6 else "VERT"),
                 wrap=True); r += 1
        # Qualité des données (couche générique, chemin déclaratif) — source UNIQUE
        # partagée avec Excel A6 / Word / HTML. Rien affiché si la couche n'a pas tourné.
        _synth_q6 = synthese_qualite_donnees(
            r6.get('rapport_qualite') if isinstance(r6, dict) else None)
        if _synth_q6:
            _kpi(ws5, r, "Qualité des données — traitements appliqués", _synth_q6,
                 statut=("AMBRE" if ("EXCLUE" in _synth_q6 or "SIGNALEE" in _synth_q6
                                     or "BLOQUE" in _synth_q6) else "VERT"),
                 wrap=True); r += 1
        # Colonnes du plan non produites (fichier client incomplet → modèle
        # amputé) — source UNIQUE partagée avec Excel A6 / Word / HTML.
        _synth_cp6 = synthese_colonnes_plan_manquantes(
            r6.get('colonnes_plan_manquantes') if isinstance(r6, dict) else None)
        if _synth_cp6:
            _kpi(ws5, r, "Colonnes du plan non produites — modèle amputé",
                 _synth_cp6, statut="AMBRE", wrap=True); r += 1

        # ── Onglet 6 : Audit trail consolidé ──────────────────────────────────
        ws6 = wb.create_sheet("6-Audit Trail Consolidé")
        _bandeau(ws6, "Audit Trail Consolidé", "Traçabilité complète A1 → A6",
                 "Équipe", audit_id, now)
        r = 7
        for col, txt, w in [(1,"Agent",14),(2,"Audit ID",26),(3,"Statut",14),(4,"Timestamp",22)]:
            _header(ws6, r, col, txt, width=w)
        r += 1
        for agent_key in ['a1','a2','a3','a4','a5','a6']:
            res = results.get(agent_key) or {}
            if not res:
                continue
            at = res.get('audit_trail', {})
            st = res.get('statut_rag', 'N/A')
            _cell(ws6, r, 1, agent_key.upper(), cf=NOIR_HEX, fill=GRIS_L_HEX, bold=True)
            _cell(ws6, r, 2, res.get('audit_id', at.get('audit_id', 'N/A')), cf=NOIR_HEX)
            _cell(ws6, r, 3, st, cf=BLANC_HEX, fill=_statut_fill(st), ah="center")
            _cell(ws6, r, 4, at.get('timestamp', 'N/A'), cf=NOIR_HEX)
            r += 1

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    except Exception as e:
        logger.error(f"export_excel_equipe échoué : {e}", exc_info=True)
        return b''


# =============================================================================
#  HTML — Rapport consolidé équipe
# =============================================================================

def export_html_equipe(results: Dict[str, Dict], branche: str = '',
                         arrete: str = '', audit_id: str = '') -> str:
    """Génère le rapport HTML consolidé équipe. Retourne str HTML ou ''."""
    try:
        r1 = results.get('a1') or {}
        r2 = results.get('a2') or {}
        r3 = results.get('a3') or {}
        r4 = results.get('a4') or {}
        r5 = results.get('a5') or {}
        r6 = results.get('a6') or {}

        now = datetime.now().strftime('%d/%m/%Y %H:%M')
        arr = arrete or datetime.now().strftime('%d/%m/%Y')
        branche_f = branche or (r6 or r3 or r1).get('branche', 'non_vie')
        statuts = _collecter_statuts(results)
        statut_global = _statut_global(list(statuts.values()))
        s_col   = _statut_col(statut_global)
        s_emoji = _statut_emoji(statut_global)

        def _row(cells, header=False):
            tag = 'th' if header else 'td'
            return '<tr>' + ''.join(f'<{tag}>{c}</{tag}>' for c in cells) + '</tr>'

        agents_roles = {
            'a1': 'A1 — Ingestion & Validation qualité',
            'a2': 'A2 — Preprocessing & Feature Engineering',
            'a3': 'A3 — GLM Poisson/Gamma/Tweedie + Crédibilité + Lissage géo',
            'a4': 'A4 — Machine Learning ×8 modèles',
            'a5': 'A5 — Deep Learning (CANN/TabNet)',
            'a6': 'A6 — Comparaison, sélection & gouvernance',
        }

        rows_statuts = ''
        for agent_key, role in agents_roles.items():
            st = statuts.get(agent_key, 'N/A')
            em = _statut_emoji(st) if st != 'N/A' else '⚪'
            rows_statuts += _row([
                agent_key.upper(), role,
                f'<span class="badge badge-{st.lower()}">{em} {st}</span>' if st != 'N/A'
                    else '<span class="badge" style="background:#8A9BB0">— Non exécuté</span>'
            ])

        # Section 2 — Qualité données
        qualite = r1.get('qualite', {})
        section_a1 = f"""
        <div class="kpi-grid">
          <div class="kpi"><b>Score qualité</b><br>{qualite.get('score_global',0):.1f}/100</div>
          <div class="kpi"><b>Nb lignes</b><br>{qualite.get('nb_lignes',0):,}</div>
          <div class="kpi"><b>Complétude</b><br>{qualite.get('taux_completude',0):.1f}%</div>
          <div class="kpi"><b>Anomalies</b><br>{qualite.get('nb_types_aberrants',0)} type(s)</div>
        </div>
        """
        if qualite.get('alertes_aberrants'):
            section_a1 += '<table>' + _row(['Alerte'], header=True)
            for al in qualite.get('alertes_aberrants', []):
                section_a1 += _row([al])
            section_a1 += '</table>'

        # Section 3 — Preprocessing
        data_dict = r2.get('data_dictionnaire', {})
        section_a2 = f"""
        <div class="kpi-grid">
          <div class="kpi"><b>Variables tracées</b><br>{len(data_dict)}</div>
          <div class="kpi"><b>Statut</b><br>{r2.get('statut_rag','N/A')}</div>
        </div>
        <table>{_row(['Variable','Usage'], header=True)}
        """
        for var_nom, meta in list(data_dict.items())[:10]:
            section_a2 += _row([var_nom, meta.get('usage','')])
        section_a2 += '</table>'

        # Section 4 — Tarification GLM/ML/DL
        met3 = r3.get('metriques', {})
        cred = r3.get('credibilite', {})
        geo  = r3.get('lissage_geo', {})
        cl4  = r4.get('classement', [])
        section_a3_a5 = f"""
        <div class="kpi-grid">
          <div class="kpi"><b>Poisson Gini</b><br>{met3.get('poisson',{}).get('gini',0):.4f}</div>
          <div class="kpi"><b>Gamma Gini</b><br>{met3.get('gamma',{}).get('gini',0):.4f}</div>
          <div class="kpi"><b>Tweedie Gini</b><br>{met3.get('tweedie',{}).get('gini',0):.4f}</div>
          <div class="kpi"><b>Crédibilité B-S</b><br>{'✓ Appliquée' if cred.get('appliquee') else '○ N/A'}</div>
          <div class="kpi"><b>Lissage géo</b><br>{'✓ Appliqué' if geo.get('applique') else '○ N/A'}</div>
          <div class="kpi"><b>Deep Learning</b><br>{r5.get('statut_rag','Non exécuté') if r5 else 'Non exécuté'}</div>
        </div>
        <table>{_row(['Modèle ML','Gini','Score','Overfit'], header=True)}
        """
        for m in cl4[:8]:
            # Audit V7 IMPORTANT : garde NA cohérent avec les autres formats.
            section_a3_a5 += _row([
                m.get('modele',''), f"{m.get('gini_test',0):.4f}",
                f"{m.get('score_global'):.4f}" if 'score_global' in m else '—',
                f"{m.get('overfit_ratio',0):.3f}"
            ])
        section_a3_a5 += '</table>'

        # Section 5 — Décision finale A6
        prod = r6.get('modele_production', {})
        bt6  = r6.get('backtest', {})
        at6  = r6.get('audit_trail', {})
        _score_txt6 = f"{prod.get('score_global'):.4f}" if 'score_global' in prod else '—'
        section_a6 = f"""
        <div class="kpi-grid">
          <div class="kpi"><b>Modèle retenu</b><br>{prod.get('modele','N/A')}</div>
          <div class="kpi"><b>Score global</b><br>{_score_txt6}</div>
          <div class="kpi"><b>Gini WF moyen</b><br>{bt6.get('gini_wf_moyen') if bt6.get('gini_wf_moyen') is not None else '—'}</div>
          <div class="kpi"><b>A/E ratio</b><br>{bt6.get('ae_ratio','—')}</div>
        </div>
        <p style="font-size:10px;color:#8A9BB0;font-style:italic;margin-top:4px;">
          ✦ Score global : normalisation relative au meilleur modèle du
          profil retenu — pas une performance absolue. Réf. audit V7.
        </p>
        <table>{_row(['Gouvernance','Valeur'], header=True)}
          {_row(['Profil retenu', at6.get('profil_ponderation','N/A')])}
          {_row(['Validé par', at6.get('profil_valide_par') or '⚠ NON VALIDÉ'])}
          {_row(['Environnement', at6.get('environnement','N/A')])}
          {_row(['Conforme', '✓ Oui' if at6.get('gouvernance_ok') else '✗ Non'])}
        </table>
        """

        # Section 6 — Audit trail consolidé
        section_audit = f"<table>{_row(['Agent','Audit ID','Statut'], header=True)}"
        for agent_key in ['a1','a2','a3','a4','a5','a6']:
            res = results.get(agent_key) or {}
            if not res:
                continue
            st = res.get('statut_rag', 'N/A')
            em = _statut_emoji(st)
            section_audit += _row([
                agent_key.upper(), res.get('audit_id','N/A'),
                f'<span class="badge badge-{st.lower()}">{em} {st}</span>'
            ])
        section_audit += '</table>'

        # Section 7 — Commentaire actuariel de référence (A6)
        # Réutilise le commentaire déjà généré par A6 (rapport_modeles_tarif.py)
        # AUCUN nouvel appel à l'API Claude n'est effectué ici.
        commentaire_a6 = r6.get('commentaire', '')
        if commentaire_a6:
            section_commentaire = (
                _md_to_html_light(commentaire_a6)
                + '<p style="margin-top:12px; font-size:10px; color:#8A9BB0; '
                  'font-style:italic;">✦ Commentaire généré par A6 — voir le '
                  'rapport détaillé A6 (Word/HTML) pour l\'analyse complète en 7 sections.</p>'
            )
        else:
            section_commentaire = '<p><em>Commentaire non disponible — exécuter A6 avec ' \
                                   'generer_graphiques=True pour le générer.</em></p>'

        html = f"""<!DOCTYPE html>
<html lang="fr">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Rapport Consolidé Équipe — ActuarIA</title>
<style>
:root {{
  --navy:{NAVY}; --navy-mid:{NAVY_MID}; --gold:{GOLD}; --gold-l:{GOLD_L};
  --rouge:{ROUGE}; --vert:{VERT}; --orange:{ORANGE}; --slate:{SLATE};
  --bg:{BG}; --white:{WHITE}; --text:{TEXT};
}}
*{{box-sizing:border-box; margin:0; padding:0;}}
body{{font-family:'Segoe UI',Arial,sans-serif; background:var(--bg); color:var(--text); font-size:13px; line-height:1.6;}}
.page{{max-width:1060px; margin:0 auto; padding:32px 24px;}}
.header{{background:var(--navy); color:#fff; padding:28px 32px; border-radius:8px; margin-bottom:28px;}}
.header h1{{font-size:22px; color:var(--gold); letter-spacing:.5px;}}
.header .sub{{font-size:13px; color:var(--gold-l); margin-top:4px;}}
.header .meta{{font-size:11px; color:#8a9bb0; margin-top:10px; display:flex; gap:24px; flex-wrap:wrap;}}
.badge{{display:inline-block; padding:3px 10px; border-radius:4px; font-size:11px; font-weight:700; color:#fff;}}
.badge-vert{{background:var(--vert);}} .badge-ambre{{background:var(--orange);}} .badge-rouge{{background:var(--rouge);}}
.section{{background:var(--white); border-radius:8px; margin-bottom:20px; box-shadow:0 1px 4px rgba(0,0,0,.07); overflow:hidden;}}
.section-head{{background:var(--navy-mid); color:var(--gold); padding:10px 18px; font-size:13px; font-weight:700;}}
.section-body{{padding:16px 18px;}}
table{{width:100%; border-collapse:collapse; font-size:12px; margin-top:8px;}}
th{{background:var(--navy); color:var(--gold); padding:7px 10px; text-align:left;}}
td{{padding:6px 10px; border-bottom:1px solid #e8edf2;}}
tr:nth-child(even) td{{background:#f7f9fc;}}
.footer{{text-align:center; font-size:10px; color:var(--slate); margin-top:28px; padding:10px;}}
.kpi-grid{{display:grid; grid-template-columns:repeat(auto-fill,minmax(160px,1fr)); gap:10px; margin-top:8px;}}
.kpi{{background:var(--bg); border:1px solid #dde4ee; border-radius:6px; padding:10px 12px;}}
</style>
</head>
<body>
<div class="page">
  <div class="header">
    <h1>ActuarIA — Rapport Consolidé Équipe Tarification</h1>
    <div class="sub">Vue d'ensemble A1 → A6 — {branche_f.replace('_',' ').title()}</div>
    <div class="meta">
      <span>Arrêté : {arr}</span>
      <span>Généré le : {now}</span>
      <span>Audit ID : {audit_id or 'N/A'}</span>
      <span class="badge badge-{statut_global.lower()}">{s_emoji} {statut_global}</span>
    </div>
  </div>

  <div class="section">
    <div class="section-head">§1 — Statut Global Équipe</div>
    <div class="section-body">
      <table>{_row(['Agent','Rôle','Statut'], header=True)}{rows_statuts}</table>
    </div>
  </div>

  <div class="section">
    <div class="section-head">§2 — Qualité des Données (A1)</div>
    <div class="section-body">{section_a1}</div>
  </div>

  <div class="section">
    <div class="section-head">§3 — Preprocessing & Feature Engineering (A2)</div>
    <div class="section-body">{section_a2}</div>
  </div>

  <div class="section">
    <div class="section-head">§4 — Tarification GLM · ML · DL (A3, A4, A5)</div>
    <div class="section-body">{section_a3_a5}</div>
  </div>

  <div class="section">
    <div class="section-head">§5 — Décision Finale & Gouvernance (A6)</div>
    <div class="section-body">{section_a6}</div>
  </div>

  <div class="section">
    <div class="section-head">§6 — Audit Trail Consolidé</div>
    <div class="section-body">{section_audit}</div>
  </div>

  <div class="section">
    <div class="section-head">§7 — Commentaire Actuariel de Référence (A6)</div>
    <div class="section-body narration">{section_commentaire}</div>
  </div>

  <div class="footer">
    ActuarIA · Équipe Tarification Non-Vie · {branche_f.replace('_',' ').title()} · Arrêté {arr} · {now} · CONFIDENTIEL
  </div>
</div>
</body>
</html>"""
        return html

    except Exception as e:
        logger.error(f"export_html_equipe échoué : {e}", exc_info=True)
        return ''


# =============================================================================
#  WORD — Rapport consolidé équipe
# =============================================================================

def export_word_equipe(results: Dict[str, Dict], branche: str = '',
                         arrete: str = '', audit_id: str = '') -> bytes:
    """Génère le rapport Word consolidé équipe (.docx). Retourne bytes ou b''."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        logger.error(f'python-docx absent : {e}'); return b''

    try:
        r1 = results.get('a1') or {}
        r2 = results.get('a2') or {}
        r3 = results.get('a3') or {}
        r4 = results.get('a4') or {}
        r5 = results.get('a5') or {}
        r6 = results.get('a6') or {}

        now = datetime.now().strftime('%d/%m/%Y')
        arr = arrete or now
        branche_f = branche or (r6 or r3 or r1).get('branche', 'non_vie')
        statuts = _collecter_statuts(results)
        statut_global = _statut_global(list(statuts.values()))

        def rgb(h):
            h = h.lstrip('#')
            return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))
        NR = rgb(NAVY); GR = rgb(GOLD); BR = rgb('#FFFFFF')
        GrR = rgb(SLATE); VR = rgb(VERT); AR = rgb(ORANGE); RgR = rgb(ROUGE)

        def statut_rgb(s):
            return {'VERT': VR, 'AMBRE': AR, 'ROUGE': RgR}.get((s or '').upper(), GrR)

        doc = Document()
        for s in doc.sections:
            s.left_margin = Cm(2); s.right_margin = Cm(2)

        # Titre
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("ActuarIA — Rapport Consolidé Équipe Tarification")
        run.font.size = Pt(18); run.font.bold = True; run.font.color.rgb = NR

        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(f"Vue d'ensemble A1 → A6 — {branche_f.replace('_',' ').title()} — Arrêté {arr}")
        run2.font.size = Pt(12); run2.font.color.rgb = GrR

        doc.add_paragraph()

        # §1 — Statut global
        h1 = doc.add_heading("§1 — Statut Global Équipe", level=1)
        h1.runs[0].font.color.rgb = NR
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = 'Light Grid Accent 1'
        hdr = tbl.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = 'Agent', 'Rôle', 'Statut'
        agents_roles = {
            'a1': 'A1 — Ingestion & Validation qualité',
            'a2': 'A2 — Preprocessing & Feature Engineering',
            'a3': 'A3 — GLM + Crédibilité + Lissage géo',
            'a4': 'A4 — Machine Learning ×8 modèles',
            'a5': 'A5 — Deep Learning',
            'a6': 'A6 — Comparaison & Gouvernance',
        }
        for agent_key, role in agents_roles.items():
            st = statuts.get(agent_key, 'N/A')
            row = tbl.add_row().cells
            row[0].text = agent_key.upper()
            row[1].text = role
            row[2].text = st

        # §2 — Qualité données (A1)
        doc.add_heading("§2 — Qualité des Données (A1)", level=1).runs[0].font.color.rgb = NR
        qualite = r1.get('qualite', {})
        doc.add_paragraph(f"Score qualité : {qualite.get('score_global',0):.1f}/100")
        doc.add_paragraph(f"Nb lignes : {qualite.get('nb_lignes',0):,}")
        doc.add_paragraph(f"Taux complétude : {qualite.get('taux_completude',0):.1f}%")
        doc.add_paragraph(f"Types d'anomalies détectées : {qualite.get('nb_types_aberrants',0)}")
        for alerte in qualite.get('alertes_aberrants', []):
            doc.add_paragraph(f"• {alerte}", style='List Bullet')

        # §3 — Preprocessing (A2)
        doc.add_heading("§3 — Preprocessing & Feature Engineering (A2)", level=1).runs[0].font.color.rgb = NR
        data_dict = r2.get('data_dictionnaire', {})
        doc.add_paragraph(f"Nombre de variables dérivées tracées : {len(data_dict)}")
        for var_nom, meta in list(data_dict.items())[:10]:
            doc.add_paragraph(f"• {var_nom} — {meta.get('usage','')}", style='List Bullet')

        # §4 — Tarification GLM/ML/DL
        doc.add_heading("§4 — Tarification GLM · ML · DL (A3, A4, A5)", level=1).runs[0].font.color.rgb = NR
        met3 = r3.get('metriques', {})
        doc.add_paragraph(f"GLM Poisson — Gini : {met3.get('poisson',{}).get('gini',0):.4f}")
        doc.add_paragraph(f"GLM Gamma — Gini : {met3.get('gamma',{}).get('gini',0):.4f}")
        doc.add_paragraph(f"GLM Tweedie — Gini : {met3.get('tweedie',{}).get('gini',0):.4f}")
        cred = r3.get('credibilite', {})
        if cred.get('appliquee'):
            doc.add_paragraph(
                f"Crédibilité Bühlmann-Straub : appliquée "
                f"(Z moyen={cred.get('z_moyen', 0):.4f}, k={cred.get('k', 0):.4f})"
            )
        else:
            doc.add_paragraph(
                f"Crédibilité Bühlmann-Straub : non applicable ({cred.get('raison', '—')})"
            )
        geo = r3.get('lissage_geo', {})
        doc.add_paragraph(
            f"Lissage géographique : {'appliqué (' + geo.get('methode','') + ')' if geo.get('applique') else 'non applicable'}"
        )
        cl4 = r4.get('classement', [])
        if cl4:
            tbl4 = doc.add_table(rows=1, cols=4)
            tbl4.style = 'Light Grid Accent 1'
            h = tbl4.rows[0].cells
            h[0].text, h[1].text, h[2].text, h[3].text = 'Modèle', 'Gini', 'Score', 'Overfit'
            for m in cl4[:8]:
                row = tbl4.add_row().cells
                row[0].text = m.get('modele','')
                row[1].text = f"{m.get('gini_test',0):.4f}"
                # Audit V7 IMPORTANT : garde NA cohérent avec les autres formats.
                row[2].text = f"{m.get('score_global'):.4f}" if 'score_global' in m else '—'
                row[3].text = f"{m.get('overfit_ratio',0):.3f}"

        # §5 — Décision finale
        doc.add_heading("§5 — Décision Finale & Gouvernance (A6)", level=1).runs[0].font.color.rgb = NR
        prod = r6.get('modele_production', {})
        bt6  = r6.get('backtest', {})
        at6  = r6.get('audit_trail', {})
        doc.add_paragraph(f"Modèle de production retenu : {prod.get('modele','N/A')}")
        _score_txt_eq = f"{prod.get('score_global'):.4f}" if 'score_global' in prod else '—'
        doc.add_paragraph(f"Score global : {_score_txt_eq}")
        # Audit V7 IMPORTANT #1 : qualification du score composite, absente
        # du livrable avant ce correctif (ne figurait que dans un
        # commentaire de code, jamais vue par le client). Pattern
        # add_run()/.italic identique à celui déjà utilisé plus bas dans
        # ce même fichier (note "Commentaire généré par A6").
        p_score_note = doc.add_paragraph()
        r_score_note = p_score_note.add_run(
            "  ↳ Score relatif au meilleur modèle du profil de pondération "
            "retenu (≈1,0000 = meilleur) — pas une performance absolue."
        )
        r_score_note.italic = True
        r_score_note.font.size = Pt(8)
        doc.add_paragraph(f"Gini walk-forward moyen (recalibré) : {bt6.get('gini_wf_moyen','—')}")
        doc.add_paragraph(f"Profil de pondération retenu : {at6.get('profil_ponderation','N/A')}")
        doc.add_paragraph(f"Validé par : {at6.get('profil_valide_par') or '⚠ NON VALIDÉ'}")
        doc.add_paragraph(f"Gouvernance conforme : {'Oui' if at6.get('gouvernance_ok') else 'Non'}")

        # §6 — Audit trail
        doc.add_heading("§6 — Audit Trail Consolidé", level=1).runs[0].font.color.rgb = NR
        tbl6 = doc.add_table(rows=1, cols=3)
        tbl6.style = 'Light Grid Accent 1'
        h6 = tbl6.rows[0].cells
        h6[0].text, h6[1].text, h6[2].text = 'Agent', 'Audit ID', 'Statut'
        for agent_key in ['a1','a2','a3','a4','a5','a6']:
            res = results.get(agent_key) or {}
            if not res:
                continue
            row = tbl6.add_row().cells
            row[0].text = agent_key.upper()
            row[1].text = res.get('audit_id','N/A')
            row[2].text = res.get('statut_rag','N/A')

        # §7 — Commentaire actuariel de référence (A6)
        # Réutilise le commentaire déjà généré par A6 — AUCUN nouvel appel API.
        # Cohérence avec export_html_equipe (même contenu, formats différents).
        doc.add_heading("§7 — Commentaire Actuariel de Référence (A6)", level=1).runs[0].font.color.rgb = NR
        commentaire_a6 = r6.get('commentaire', '')
        if commentaire_a6:
            sections_n = re.split(r'(?=§\d+\s*[—\-–])', commentaire_a6.strip())
            for sec in sections_n:
                sec = sec.strip()
                if not sec:
                    continue
                lignes = sec.split('\n', 1)
                if lignes[0]:
                    ph = doc.add_paragraph()
                    rh = ph.add_run(lignes[0].strip())
                    rh.bold = True; rh.font.size = Pt(10); rh.font.color.rgb = GR
                if len(lignes) > 1:
                    for ln in lignes[1].split('\n'):
                        ln = ln.strip()
                        if ln:
                            pl = doc.add_paragraph()
                            pl.paragraph_format.left_indent = Cm(0.3)
                            rl = pl.add_run(ln)
                            rl.font.size = Pt(9); rl.font.color.rgb = NR
            p_note = doc.add_paragraph()
            r_note = p_note.add_run(
                "✦ Commentaire généré par A6 — voir le rapport détaillé A6 "
                "(Word/HTML) pour l'analyse complète en 7 sections."
            )
            r_note.font.size = Pt(7); r_note.italic = True; r_note.font.color.rgb = GrR
        else:
            p_nc = doc.add_paragraph()
            r_nc = p_nc.add_run(
                "Commentaire non disponible — exécuter A6 avec "
                "generer_graphiques=True pour le générer."
            )
            r_nc.italic = True; r_nc.font.size = Pt(9)

        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = footer.add_run(f"ActuarIA · Équipe Tarification Non-Vie · {branche_f} · Arrêté {arr} · {now} · CONFIDENTIEL")
        fr.font.size = Pt(8); fr.font.color.rgb = GrR

        buf = io.BytesIO()
        doc.save(buf)
        word_bytes = buf.getvalue()
        logger.info(f'Word équipe : {len(word_bytes):,} bytes')
        return word_bytes

    except Exception as e:
        logger.error(f'export_word_equipe échoué : {e}', exc_info=True)
        return b''


# =============================================================================
#  PDF — Rapport consolidé équipe (via HTML → weasyprint)
# =============================================================================

def export_pdf_equipe(results: Dict[str, Dict], branche: str = '',
                        arrete: str = '', audit_id: str = '') -> bytes:
    """Génère le rapport PDF consolidé équipe via weasyprint (HTML→PDF)."""
    try:
        from weasyprint import HTML as WP_HTML
        html = export_html_equipe(results, branche, arrete, audit_id)
        if not html:
            return b''
        buf = io.BytesIO()
        WP_HTML(string=html).write_pdf(buf)
        buf.seek(0)
        pdf = buf.read()
        logger.info(f'PDF équipe : {len(pdf):,} bytes')
        return pdf
    except ImportError:
        logger.warning('weasyprint absent — PDF équipe non généré')
        return b''
    except Exception as e:
        logger.error(f'export_pdf_equipe échoué : {e}', exc_info=True)
        return b''


# =============================================================================
#  POINT D'ENTRÉE UNIFIÉ
# =============================================================================

def generer_rapport_equipe_tarification(
    results: Dict[str, Dict],
    branche: str = '',
    arrete: str = '',
    audit_id: str = '',
    formats: List[str] = None,
) -> Dict[str, bytes]:
    """
    Génère le rapport consolidé équipe dans tous les formats demandés.

    results : dict {'a1': result_a1, 'a2': result_a2, ..., 'a6': result_a6}
    formats : ['html','word','pdf','excel'] (défaut = tous)
    Retourne {'html_bytes':..., 'word_bytes':..., 'pdf_bytes':..., 'excel_bytes':...}
    """
    if formats is None:
        formats = ['html', 'word', 'pdf', 'excel']

    out: Dict[str, bytes] = {
        'html_bytes': b'', 'word_bytes': b'', 'pdf_bytes': b'', 'excel_bytes': b'',
    }

    html_str = ''
    if 'html' in formats or 'pdf' in formats:
        html_str = export_html_equipe(results, branche, arrete, audit_id)
        out['html_bytes'] = html_str.encode('utf-8') if html_str else b''

    if 'word' in formats:
        out['word_bytes'] = export_word_equipe(results, branche, arrete, audit_id)

    if 'pdf' in formats:
        if html_str:
            try:
                from weasyprint import HTML as WP_HTML
                buf = io.BytesIO()
                WP_HTML(string=html_str).write_pdf(buf)
                buf.seek(0)
                out['pdf_bytes'] = buf.read()
            except ImportError:
                logger.warning('weasyprint absent — PDF équipe non généré')
        else:
            out['pdf_bytes'] = export_pdf_equipe(results, branche, arrete, audit_id)

    if 'excel' in formats:
        out['excel_bytes'] = export_excel_equipe(results, branche, arrete, audit_id)

    logger.info(
        f"Rapport équipe : HTML={len(out['html_bytes']):,}b "
        f"Word={len(out['word_bytes']):,}b PDF={len(out['pdf_bytes']):,}b "
        f"Excel={len(out['excel_bytes']):,}b"
    )
    return out
