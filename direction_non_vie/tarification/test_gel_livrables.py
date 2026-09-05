# -*- coding: utf-8 -*-
"""LE GEL DES LIVRABLES : COMPARER DU CONTENU, PAS DES OCTETS.

Cet outil rend mesurable la phrase << aucun euro n'a bouge >>. Il produit
l'empreinte normalisee des surfaces signees avant et apres un correctif, et
nomme chaque ecart par sa feuille et sa coordonnee.

Trois mesures du 05/09/2026 fondent sa forme, et chacune a sa sentinelle :

  1. DEUX RUNS IDENTIQUES NE RENDENT PAS LES MEMES OCTETS. Un `.docx` de
     contenu identique produit a deux instants differe : le ZIP horodate ses
     entrees. Et sa TAILLE bouge aussi (41 588 puis 41 589 octets) parce que
     `deflate` comprime differemment un horodatage qui a change de minute.
     -> comparer des octets, ou des tailles, c'est un rouge par minute.

  2. IL Y A DEUX CONVENTIONS D'HEURE dans les surfaces signees :
     `entete_livrable.genere_le()` rend << 05/09/2026 01:46 >>, et
     `rapport_modeles_tarif.valeur_audit()` rend << 05/09/2026 a 01 h 46 >>.
     Un normaliseur qui n'en connait qu'une rougit une minute sur deux.

  3. ET ON NE PEUT PAS EFFACER TOUTE DATE. `libelle_arrete()` rend
     << 30/06/2026 >> -- la date d'ARRETE, du contenu signe, de forme
     identique a une impression. La distinction retenue est STRUCTURELLE :
     une date SUIVIE D'UNE HEURE est une impression, une date SEULE est du
     metier.

Ce que cette sentinelle exige :
  GEL-1   l'horodatage du VRAI producteur est neutralise, aucun chiffre ne
          survit ;
  GEL-2   le rendu francais de la piste d'audit l'est aussi ;
  GEL-3   la date d'ARRETE, elle, SURVIT -- sinon le gel est aveugle au champ
          le plus lourd du document ;
  GEL-4   deux Word de contenu identique, d'octets differents, ont la MEME
          empreinte ;
  GEL-5   l'extraction Word ne lit QUE `<w:t>` : un changement de style ne
          produit aucun ecart, un changement de texte en produit un ;
  GEL-6   une cellule Excel changee est vue, NOMMEE par feuille et coordonnee ;
  GEL-7   une surface presente d'un seul cote est un ECART, pas un silence ;
  GEL-8   une surface devenue ILLISIBLE est un ecart, et le verdict publie son
          assiette ;
  GEL-9   l'inventaire des livrables ENUMERE : une cle `_bytes` ajoutee demain
          entre dans la mesure sans toucher au code ;
  GEL-10  aucune taille en octets n'entre dans l'empreinte ;
  GEL-11  la chaine est reellement DETERMINISTE : A3 deux fois, meme empreinte.

Tout en `unittest.TestCase` : la gate lance `unittest discover`.
"""
import io
import os
import sys
import unittest
import warnings
import zipfile

_ICI = os.path.dirname(os.path.abspath(__file__))
_RACINE = os.path.dirname(os.path.dirname(_ICI))
for _c in (_RACINE, _ICI):
    if _c not in sys.path:
        sys.path.insert(0, _c)

from direction_non_vie.tarification.services import entete_livrable as EL
from direction_non_vie.tarification.services import gel_livrables as G
from direction_non_vie.tarification.services.rapport_modeles_tarif import (
    valeur_audit,
)

# =============================================================================
#  FABRIQUES — de vrais fichiers, jamais des octets simules
# =============================================================================

def _docx(paragraphes, largeur=2400, fond='EEF2F7'):
    """Un vrai .docx : un tableau d'une colonne, style parametrable."""
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    table = doc.add_table(rows=len(paragraphes), cols=1)
    for indice, texte in enumerate(paragraphes):
        cellule = table.rows[indice].cells[0]
        cellule.paragraphs[0].add_run(str(texte))
        proprietes = cellule._tc.get_or_add_tcPr()
        largeur_el = OxmlElement('w:tcW')
        largeur_el.set(qn('w:w'), str(largeur))
        largeur_el.set(qn('w:type'), 'dxa')
        proprietes.append(largeur_el)
        ombre = OxmlElement('w:shd')
        ombre.set(qn('w:val'), 'clear')
        ombre.set(qn('w:fill'), fond)
        proprietes.append(ombre)
    flux = io.BytesIO()
    doc.save(flux)
    return flux.getvalue()


def _xlsx(feuilles):
    """Un vrai .xlsx : {nom de feuille: {coordonnee: valeur}}."""
    from openpyxl import Workbook

    classeur = Workbook()
    classeur.remove(classeur.active)
    for nom, cellules in feuilles.items():
        feuille = classeur.create_sheet(nom)
        for coordonnee, valeur in cellules.items():
            feuille[coordonnee] = valeur
    flux = io.BytesIO()
    classeur.save(flux)
    return flux.getvalue()


def _rehorodater(octets, decalage_s=120):
    """Le MEME contenu, un ZIP horodate autrement.

    ⚠️⚠️ POURQUOI PAS UN `time.sleep()` : **mesure du 05/09/2026**, l'horodatage
    ZIP avance par pas de DEUX secondes (58 -> 0 -> 2). Une pause d'une seconde
    retombe une fois sur deux dans le meme palier, les octets sortent
    identiques, et le test rougit au hasard. C'est le sceau qui l'a trouve, en
    relancant la suite six fois. Ici le decalage est impose, donc reproductible.
    """
    lu = zipfile.ZipFile(io.BytesIO(octets))
    flux = io.BytesIO()
    with zipfile.ZipFile(flux, 'w', zipfile.ZIP_DEFLATED) as ecrit:
        for info in lu.infolist():
            annee, mois, jour, heure, minute, seconde = info.date_time
            minute = (minute + decalage_s // 60) % 60
            neuf = zipfile.ZipInfo(info.filename,
                                   (annee, mois, jour, heure, minute, seconde))
            neuf.compress_type = info.compress_type
            ecrit.writestr(neuf, lu.read(info.filename))
    return flux.getvalue()


def _empreinte(livrables):
    return G.empreinte(livrables)


def _ecarts(avant, apres):
    return G.comparer(_empreinte(avant), _empreinte(apres))


# =============================================================================
#  GEL-1 a GEL-3 — LA NORMALISATION, LIEE AUX VRAIS PRODUCTEURS
# =============================================================================

class TestNormalisationLieeAuxProducteurs(unittest.TestCase):
    """La regle n'est pas recopiee ici : elle est verifiee sur la source."""

    def test_GEL1_l_horodatage_de_genere_le_ne_survit_pas(self):
        """⚠️ On appelle le VRAI producteur. Une liste de formats recopiee se
        met a jour a moitie ; un controle qui appelle la source suit."""
        brut = EL.genere_le()
        neutre = G.neutraliser(brut)
        self.assertNotEqual(brut, neutre,
                            f"l'horodatage de production {brut!r} n'a pas ete "
                            "neutralise")
        self.assertFalse(any(c.isdigit() for c in neutre),
                         f"des chiffres survivent dans {neutre!r} : deux runs "
                         "a une minute d'ecart rougiraient")

    def test_GEL1b_la_phrase_complete_du_livrable_est_neutralisee(self):
        """C'est la phrase telle qu'elle figure dans les classeurs."""
        phrase = f'Généré le : {EL.genere_le()}'
        neutre = G.neutraliser(phrase)
        self.assertNotIn(':', neutre.replace('Généré le :', ''),
                         f'une heure survit dans {neutre!r}')
        self.assertFalse(any(c.isdigit() for c in neutre), neutre)

    def test_GEL2_le_rendu_francais_de_la_piste_d_audit_ne_survit_pas(self):
        """`valeur_audit` rend << 05/09/2026 a 01 h 46 >>. Mesure : c'est le
        SEUL ecart qui restait entre deux runs de la chaine complete."""
        rendu = valeur_audit('2026-09-05T01:46:12.123456')
        self.assertIn(' h ', rendu,
                      f'le producteur a change de forme : {rendu!r} -- la '
                      'sentinelle doit etre relue avec lui')
        neutre = G.neutraliser(rendu)
        self.assertFalse(any(c.isdigit() for c in neutre),
                         f'des chiffres survivent dans {neutre!r}')

    def test_GEL3_la_date_d_arrete_SURVIT_a_la_neutralisation(self):
        """⚠️⚠️ LE COEUR. Une date d'arrete a la meme forme qu'une impression.
        L'effacer rendrait le gel aveugle au champ le plus lourd du document.

        ⚠️⚠️ ET CE TEST A DEJA ATTESTE SANS SURVEILLER. Ecrit d'abord sur
        `30/06/2026`, il passait -- mais `core/arrete.py` DERIVE une fin de
        trimestre en << T2 2026 >>, qui n'a AUCUNE collision de forme avec un
        horodatage. Le cas dangereux est la date HORS fin de trimestre, seule
        a ressortir en `jj/mm/aaaa`. Mesure du 05/09/2026 :
            30/06/2026 -> << T2 2026 >>     (libelle derive, aucun risque)
            15/05/2026 -> << 15/05/2026 >>  (COLLISION -- le vrai cas)
        """
        derive = EL.libelle_arrete('30/06/2026')
        self.assertNotRegex(derive, r'\d{2}/\d{2}/\d{4}',
                            'core/arrete a change de rendu : ce test doit '
                            'etre relu avec lui')
        collision = EL.libelle_arrete('15/05/2026')
        self.assertRegex(collision, r'^15/05/2026$',
                         "le cas de COLLISION n'existe plus sous cette forme : "
                         'sans lui ce test ne surveille rien')
        self.assertEqual(G.neutraliser(collision), collision,
                         "la date d'arrete a ete neutralisee : un changement "
                         "d'arrete passerait desormais inapercu")

    def test_GEL3b_un_arrete_qui_change_produit_un_ecart(self):
        """La contre-epreuve, sur le cas de COLLISION exclusivement."""
        avant = {'Word': _docx([f'Arrêté : {EL.libelle_arrete("15/05/2026")}',
                                f'Généré le : {EL.genere_le()}'])}
        apres = {'Word': _docx([f'Arrêté : {EL.libelle_arrete("16/05/2026")}',
                                f'Généré le : {EL.genere_le()}'])}
        ecarts = _ecarts(avant, apres)
        self.assertEqual(len(ecarts), 1, G.rapport_ecarts(
            ecarts, _empreinte(avant), _empreinte(apres)))
        self.assertIn('15/05/2026', str(ecarts[0].avant))
        self.assertIn('16/05/2026', str(ecarts[0].apres))

    def test_GEL3c_un_identifiant_d_audit_PREFIXE_est_neutralise(self):
        """⚠️⚠️ `\\b\\d{8}_\\d{6}\\b` NE MORD PAS `A3_20260905_015505` : entre
        `_` et `2` il n'y a pas de frontiere de mot. Le motif ne voyait que la
        forme nue. Trouve par GEL-11, pas par relecture."""
        for identifiant in ('A3_20260905_015505', 'Audit ID : A6_20260905_015505',
                            '20260905_015505', 'A5_20260905_015505_v2'):
            neutre = G.neutraliser(identifiant)
            self.assertIn('<horodatage>', neutre,
                          f'{identifiant!r} reste horodate : deux runs a une '
                          'seconde d ecart rougiraient')


# =============================================================================
#  GEL-4, GEL-5, GEL-10 — LE WORD : DU CONTENU, NI OCTETS NI STYLE
# =============================================================================

class TestWordContenuSeul(unittest.TestCase):

    def test_GEL4_meme_contenu_octets_differents_empreinte_identique(self):
        """⚠️⚠️ LA MESURE FONDATRICE. Deux .docx de contenu identique produits
        a deux instants n'ont PAS les memes octets (le ZIP horodate ses
        entrees). Un gel par empreinte SHA-256 serait rouge en permanence."""
        textes = ['Prime pure moyenne', '412,55 EUR', 'Arrêté : 15/05/2026']
        premier = _docx(textes)
        second = _rehorodater(premier)
        self.assertNotEqual(premier, second,
                            'les octets sont identiques : la mesure qui fonde '
                            'cet outil ne se reproduit plus, relire le lot')
        self.assertEqual(_ecarts({'Word': premier}, {'Word': second}), [],
                         'du contenu identique produit un ecart')

    def test_GEL5_un_changement_de_STYLE_ne_produit_aucun_ecart(self):
        """⚠️⚠️ `<w:t[^>]*>` mord aussi `<w:tcPr>` et `<w:tcW>`. Mesure du
        05/09/2026 : l'extraction remontait des largeurs de colonne comme du
        contenu signe. Une largeur qui change n'est pas un tarif qui change."""
        textes = ['Modèle retenu', 'GLM Poisson']
        etroit = _docx(textes, largeur=2400, fond='EEF2F7')
        large = _docx(textes, largeur=5953, fond='FFFFFF')
        self.assertNotEqual(etroit, large, 'les deux docx sont identiques : '
                                           'le plant ne teste rien')
        ecarts = _ecarts({'Word': etroit}, {'Word': large})
        self.assertEqual(ecarts, [], 'un changement de style a ete annonce '
                                     'comme un changement de contenu :\n'
                         + G.rapport_ecarts(ecarts, _empreinte({'Word': etroit}),
                                            _empreinte({'Word': large})))

    def test_GEL5b_un_changement_de_TEXTE_produit_un_ecart_nomme(self):
        """La contre-epreuve du meme plant : le filet voit-il encore ?"""
        avant = _docx(['Modèle retenu', 'GLM Poisson'])
        apres = _docx(['Modèle retenu', 'LightGBM'])
        ecarts = _ecarts({'Word': avant}, {'Word': apres})
        self.assertEqual(len(ecarts), 1, G.rapport_ecarts(
            ecarts, _empreinte({'Word': avant}), _empreinte({'Word': apres})))
        self.assertEqual(ecarts[0].avant, 'GLM Poisson')
        self.assertEqual(ecarts[0].apres, 'LightGBM')
        self.assertIn('word/document.xml', ecarts[0].emplacement)

    def test_GEL10_aucune_taille_n_entre_dans_l_empreinte(self):
        """⚠️ Le resultat d'A6 publie `livrables_tailles`. Mesure : cette
        table bouge entre deux runs identiques. Elle ne doit atteindre ni
        l'inventaire des livrables, ni l'empreinte."""
        resultat = {'word_bytes': _docx(['a']),
                    'livrables_tailles': {'Word tarification': 41588}}
        inventaire = G.livrables_d_un_resultat(resultat)
        self.assertNotIn('livrables_tailles', inventaire)
        self.assertEqual(sorted(inventaire), ['Word'])
        textuel = repr(_empreinte(inventaire).contenus)
        self.assertNotIn('41588', textuel)


# =============================================================================
#  GEL-6 — L'EXCEL : LA COORDONNEE VOYAGE AVEC LA VALEUR
# =============================================================================

class TestExcelCoordonnee(unittest.TestCase):

    def test_GEL6_une_cellule_changee_est_vue_et_NOMMEE(self):
        avant = _xlsx({'1-Synthèse': {'A1': 'Prime pure', 'B1': 412.55,
                                      'A2': 'Gini', 'B2': 0.2651}})
        apres = _xlsx({'1-Synthèse': {'A1': 'Prime pure', 'B1': 418.90,
                                      'A2': 'Gini', 'B2': 0.2651}})
        ecarts = _ecarts({'Excel': avant}, {'Excel': apres})
        self.assertEqual(len(ecarts), 1, G.rapport_ecarts(
            ecarts, _empreinte({'Excel': avant}), _empreinte({'Excel': apres})))
        self.assertIn('1-Synthèse', ecarts[0].emplacement)
        self.assertIn('B1', ecarts[0].emplacement)
        self.assertEqual((ecarts[0].avant, ecarts[0].apres), (412.55, 418.90))

    def test_GEL6b_une_feuille_ajoutee_est_vue(self):
        avant = _xlsx({'1-Synthèse': {'A1': 'x'}})
        apres = _xlsx({'1-Synthèse': {'A1': 'x'}, '2-Classement': {'A1': 'y'}})
        ecarts = _ecarts({'Excel': avant}, {'Excel': apres})
        self.assertTrue(ecarts, 'une feuille entiere ajoutee est passee')
        self.assertTrue(any('2-Classement' in e.emplacement for e in ecarts),
                        [str(e) for e in ecarts])

    def test_GEL6c_deux_classeurs_identiques_ne_produisent_rien(self):
        cellules = {'1-Synthèse': {'A1': 'Prime', 'B1': 412.55,
                                   'A4': f'Généré le : {EL.genere_le()}',
                                   'A5': 'Audit ID : A3_20260905_015505'}}
        premier = _xlsx(cellules)
        second = _rehorodater(_xlsx(cellules))
        self.assertNotEqual(premier, second)
        self.assertEqual(_ecarts({'Excel': premier}, {'Excel': second}), [])


# =============================================================================
#  GEL-7, GEL-8 — L'ABSENCE ET L'ILLISIBLE SONT DES VALEURS
# =============================================================================

class TestAbsenceEtIllisible(unittest.TestCase):

    def test_GEL7_une_surface_presente_d_un_seul_cote_est_un_ecart(self):
        """⚠️⚠️ LE PIEGE DU RELAIS FIDELE A UNE ABSENCE. Un controle qui
        compare deux `None` reste vert sans avoir rien surveille."""
        avant = {'Excel': _xlsx({'S': {'A1': 'x'}}), 'Word': _docx(['y'])}
        apres = {'Excel': _xlsx({'S': {'A1': 'x'}}), 'Word': b''}
        ecarts = _ecarts(avant, apres)
        self.assertEqual(len(ecarts), 1, [str(e) for e in ecarts])
        self.assertEqual(ecarts[0].surface, 'Word')
        self.assertEqual(ecarts[0].apres, G.ABSENT)

    def test_GEL7b_deux_absences_reelles_ne_sont_pas_un_ecart(self):
        """La contre-epreuve : le gel ne crie pas sur un PDF jamais demande."""
        self.assertEqual(_ecarts({'PDF': b''}, {'PDF': None}), [])

    def test_GEL7c_une_surface_qui_n_existe_meme_plus_est_un_ecart(self):
        """⚠️ CAS DISTINCT DE GEL-7, et il etait NON COUVERT : la surface n'est
        pas vide, elle a disparu de l'inventaire. Le sceau l'a montre -- le
        plant qui neutralisait cette branche ne faisait rougir personne."""
        avant = {'Excel': _xlsx({'S': {'A1': 'x'}})}
        apres = {'Excel': _xlsx({'S': {'A1': 'x'}}),
                 'Rapport equipe PDF': _docx(['nouveau livrable'])}
        ecarts = _ecarts(avant, apres)
        self.assertEqual(len(ecarts), 1, [str(e) for e in ecarts])
        self.assertEqual(ecarts[0].surface, 'Rapport equipe PDF')
        self.assertEqual(ecarts[0].avant, G.ABSENT)

    def test_GEL8_une_surface_devenue_illisible_est_un_ecart(self):
        """Sinon un export casse passerait pour << rien n'a change >>."""
        bon = _xlsx({'S': {'A1': 'x'}})
        casse = b'PK\x03\x04' + b'\x00' * 60
        ecarts = _ecarts({'Excel': bon}, {'Excel': casse})
        self.assertTrue(ecarts, "un classeur devenu illisible n'a produit "
                                'aucun ecart')

    def test_GEL8b_le_verdict_publie_son_assiette_et_ses_non_lues(self):
        """⚠️ << 0 ecart >> ne peut pas taire une surface illisible.
        Un chiffre se publie avec la methode qui l'a produit."""
        illisible = b'%PDF-1.4 corrompu'
        emp = _empreinte({'Excel': _xlsx({'S': {'A1': 'x'}}), 'PDF': illisible})
        texte = G.rapport_ecarts(G.comparer(emp, emp), emp, emp)
        self.assertIn('0 ecart', texte)
        self.assertIn('assiette', texte)
        self.assertIn('NON LUES', texte, texte)
        self.assertIn('PDF', texte)


# =============================================================================
#  GEL-9 — L'INVENTAIRE ENUMERE, IL NE DECLARE PAS
# =============================================================================

class TestInventaireEnumere(unittest.TestCase):

    def test_GEL9_une_cle_bytes_INCONNUE_entre_dans_la_mesure(self):
        """⚠️⚠️ Une table tenue a la main aurait diverge le jour ou un export
        nouveau apparait -- et l'assiette se serait retrecie en silence."""
        resultat = {'excel_bytes': b'x', 'csv_bytes': b'y', 'statut_rag': 'VERT'}
        inventaire = G.livrables_d_un_resultat(resultat)
        self.assertIn('csv', inventaire,
                      "une cle `_bytes` inconnue n'entre pas dans la mesure : "
                      "l'inventaire DECLARE au lieu d'ENUMERER")
        self.assertNotIn('statut_rag', inventaire)

    def test_GEL9b_le_rapport_equipe_est_ouvert_d_un_cran(self):
        resultat = {'rapport_equipe': {'excel_bytes': b'a', 'pdf_bytes': b'b'}}
        inventaire = G.livrables_d_un_resultat(resultat)
        self.assertEqual(sorted(inventaire),
                         ['Rapport equipe Excel', 'Rapport equipe PDF'])

    def test_GEL9c_le_format_se_derive_des_octets_pas_de_la_cle(self):
        """Une cle qui ment sur son contenu ne trompe pas la lecture."""
        self.assertEqual(G.format_livrable(_xlsx({'S': {'A1': 1}})), 'xlsx')
        self.assertEqual(G.format_livrable(_docx(['a'])), 'docx')
        self.assertEqual(G.format_livrable(b'<!DOCTYPE html><p>x</p>'), 'html')
        self.assertEqual(G.format_livrable(b''), 'vide')


# =============================================================================
#  GEL-11 — LA CHAINE EST-ELLE REELLEMENT DETERMINISTE ?
# =============================================================================

class TestDeterminismeReel(unittest.TestCase):
    """Sans cette propriete l'outil ment : il attribuerait a un correctif un
    ecart que la chaine produit toute seule."""

    def test_GEL11_A3_deux_fois_rend_la_meme_empreinte(self):
        warnings.filterwarnings('ignore')
        import numpy as np

        from core.qualite_donnees import preambule_qualite
        from direction_non_vie.tarification import test_pipeline_agents as T
        from direction_non_vie.tarification.a1_ingestion.agent import (
            AgentA1Ingestion,
        )
        from direction_non_vie.tarification.a2_preprocessing.agent import (
            AgentA2Preprocessing,
        )
        from direction_non_vie.tarification.a3_glm.agent import AgentA3GLM

        def un_run():
            np.random.seed(7)
            donnees = T._portefeuille_auto(1200)
            plan = T._PLAN_AUTO
            base = {'audit_path': '/tmp', 'verbose': False}
            r1 = AgentA1Ingestion(**base).run(branche='non_vie',
                                              sous_branche='auto',
                                              dataframe=donnees)
            qualite = preambule_qualite(r1.get('dataframe'), plan,
                                        qualite_validee_par='Actuaire Test',
                                        horodatage=None)
            r1 = {**r1, 'dataframe': qualite.dataframe_propre}
            r2 = AgentA2Preprocessing(**base).run(result_a1=r1, plan=plan)
            r3 = AgentA3GLM(models_path='/tmp', audit_path='/tmp').run(
                result_a2=r2, plan=plan, col_frequence=plan.cible_frequence,
                col_cout=plan.cible_cout, generer_graphiques=False)
            return G.empreinte(G.livrables_d_un_resultat(r3))

        premier, second = un_run(), un_run()
        # ⚠️⚠️ L'ASSIETTE SE DECLARE. Sans ces trois lignes, un run ou A3 ne
        # produirait PLUS AUCUN classeur laisserait ce test vert : trois
        # surfaces `<livrable absent>` des deux cotes sont fidelement egales.
        # Un relais fidele a une absence est encore une absence.
        self.assertEqual(premier.non_lues, {}, premier.non_lues)
        reelles = [nom for nom, contenu in premier.contenus.items()
                   if contenu != G.ABSENT]
        self.assertIn('Excel', reelles,
                      "A3 n'a produit aucun classeur lisible : ce controle ne "
                      f'mesurerait rien. Surfaces vues : {premier.surfaces()}')
        self.assertGreaterEqual(
            sum(len(cellules) for cellules in premier.contenus['Excel'].values()),
            50, 'le classeur mesure est presque vide')
        ecarts = G.comparer(premier, second)
        self.assertEqual(ecarts, [], G.rapport_ecarts(ecarts, premier, second))


if __name__ == '__main__':
    unittest.main(verbosity=2)
