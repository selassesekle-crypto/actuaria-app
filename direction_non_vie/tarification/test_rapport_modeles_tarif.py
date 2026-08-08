# -*- coding: utf-8 -*-
"""Tests T1 — une seule narration pour tous les formats du rapport tarif.

⚠️ POURQUOI CE FICHIER EXISTE. `rapport_modeles_tarif.py` produit un livrable
SIGNÉ — celui qui part chez un commissaire aux comptes — et n'avait aucun
test. Le défaut que T1 corrige a donc vécu jusqu'à ce qu'un rapport réel soit
regardé : mesuré, le HTML portait 18 089 caractères de commentaire actuariel
et le Word portait le dépôt technique de l'agent.

⚠️ AUCUN TEST D'ICI N'ATTEINT LE RÉSEAU : le client est simulé.
"""
import io
import os
import re
import struct
import sys
import unittest
import zipfile
import zlib
from unittest.mock import patch

RACINE = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
if RACINE not in sys.path:
    sys.path.insert(0, RACINE)

from direction_non_vie.tarification.services import (  # noqa: E402
    rapport_modeles_tarif as R)

NARRATION = ('§1 — CONTEXTE ET QUALITE DES DONNEES\n'
             'Le portefeuille compte 12 000 contrats. Reference : Goldburd.\n'
             '§7 — CONCLUSION\nAvis favorable sous reserve.')

# Le repli : ce que le Word portait SEUL avant ce lot.
A3 = {'metriques': {'poisson': {'gini': 0.1775, 'aic': 10308}},
      'commentaire': 'REPLI AGENT'}
A4 = {'classement': [], 'commentaire': 'REPLI AGENT'}
A6 = {'branche': 'auto', 'statut_rag': 'VERT'}


class _Bloc:
    def __init__(self, texte):
        self.text, self.type = texte, 'text'


def _client_simule(appels, texte=NARRATION, echouer=False):
    """Un client `anthropic` de substitution qui compte les appels."""
    class _Messages:
        def create(self, **kw):
            appels.append(kw.get('model'))
            if echouer:
                raise RuntimeError('panne simulée du service')

            class _Reponse:
                content = [_Bloc(texte)]
            return _Reponse()

    class _Client:
        def __init__(self, **_):
            self.messages = _Messages()
    return _Client


def _texte_docx(octets):
    if not octets:
        return ''
    xml = zipfile.ZipFile(io.BytesIO(octets)).read(
        'word/document.xml').decode('utf-8')
    return re.sub(r'<[^>]+>', '', xml)


def _produire(appels, formats=('html', 'word'), echouer=False):
    with patch.dict(os.environ, {'ANTHROPIC_API_KEY': 'cle-de-test'}):
        with patch('anthropic.Anthropic',
                   _client_simule(appels, echouer=echouer)):
            rap = R.generer_rapport_tarification(
                A3, A4, A6, 'DEMO', '31/12/2025', 'T1', list(formats))
    return ((rap.get('html_bytes') or b'').decode('utf-8', 'replace'),
            _texte_docx(rap.get('word_bytes')))


class T1_UneSeuleNarration(unittest.TestCase):
    """T1 — le format qui part chez un commissaire doit porter le même
    commentaire que celui qu'on lit à l'écran."""

    def test_un_seul_appel_pour_les_deux_formats(self):
        """⚠️ DEUX AVANT CE LOT : `export_html` et `export_word` appelaient
        chacun la narration pour leur propre compte."""
        appels = []
        _produire(appels)
        self.assertEqual(len(appels), 1, f'appels : {appels}')
        print('    OK T1 : 1 appel pour html+word (2 avant ce lot)')

    def test_les_deux_formats_portent_le_MEME_commentaire(self):
        """⚠️ LE POINT DU LOT. Mesuré sur un rapport réel : le HTML portait la
        narration Claude, le Word le dépôt technique de l'agent."""
        html, word = _produire([])
        for marqueur in ('CONTEXTE ET QUALITE DES DONNEES', 'Goldburd',
                         'Avis favorable sous reserve'):
            self.assertIn(marqueur, html, f'HTML : {marqueur}')
            self.assertIn(marqueur, word, f'WORD : {marqueur}')
        # et AUCUN des deux ne retombe sur le dépôt de l'agent
        self.assertNotIn('REPLI AGENT', html)
        self.assertNotIn('REPLI AGENT', word)
        print('    OK T1b : le titre, la référence et la conclusion sont dans '
              'les DEUX formats')

    def test_la_source_est_nommee_dans_les_deux_formats(self):
        """Le succès se dit des deux côtés, pas seulement à l'écran."""
        html, word = _produire([])
        self.assertIn('ActuarIA Intelligence', html)
        self.assertIn('ActuarIA Intelligence', word)
        print('    OK T1c : le marqueur de source est dans les deux formats')

    def test_un_echec_se_voit_dans_les_deux_formats(self):
        """⚠️ SINON ON REMPLACE DEUX COMMENTAIRES DIVERGENTS PAR UN SILENCE.
        Le Word ne nommait sa source qu'en cas de SUCCÈS : un repli y était
        indiscernable d'une narration réussie."""
        appels = []
        html, word = _produire(appels, echouer=True)
        self.assertEqual(len(appels), 1, 'le repli est calculé une seule fois')
        for texte in (html, word):
            self.assertIn('commentaire_agent', texte)   # la source est nommée
            self.assertIn('REPLI AGENT', texte)         # le même texte
            self.assertNotIn('ActuarIA Intelligence', texte)
        print('    OK T1d : un échec est NOMMÉ dans les deux formats, et le '
              'repli y est identique')

    def test_appeles_seuls_les_exports_restent_autonomes(self):
        """L'API publique ne change pas : sans narration fournie, chaque
        export calcule la sienne, comme avant."""
        appels = []
        with patch.dict(os.environ, {'ANTHROPIC_API_KEY': 'cle-de-test'}):
            with patch('anthropic.Anthropic', _client_simule(appels)):
                html = R.export_html(A3, A4, A6, 'DEMO', '31/12/2025', 'T1')
        self.assertEqual(len(appels), 1)
        self.assertIn('Goldburd', html)
        print('    OK T1e : export_html appelé seul calcule sa narration')


class T2_LeContexteDeNarration(unittest.TestCase):
    """T2 — le prompt ordonne « §4 — COMPARAISON DES MODÈLES ML ET SÉLECTION ».

    ⚠️ MESURÉ SUR LE RAPPORT RÉEL PRODUIT AVANT CE LOT : le §4 citait
    1 modèle sur 7, aucun contrôle de sélection, et pas une occurrence de
    « comparaison » hors de son propre titre. `cl4` et `val6` étaient extraits
    puis jetés — `ruff F841` les signalait depuis toujours.
    """

    CLASSEMENT = [
        {'modele': 'GLM Poisson (référence A3)', 'famille': 'GLM',
         'gini_test': 0.1775, 'rmse_test': 0.4504, 'overfit_ratio': 1.000},
        {'modele': 'lightgbm', 'famille': 'Arbres / Boosting',
         'gini_test': 0.1729, 'rmse_test': 0.4579, 'overfit_ratio': 2.757},
        {'modele': 'xgboost_tweedie', 'famille': 'Arbres / Boosting',
         'gini_test': 0.1120, 'rmse_test': 0.4629, 'overfit_ratio': 5.248},
    ]
    CONTROLES = {
        'c1_nb_modeles': {'statut': 'VERT',
                          'message': '7 modèles comparés ≥ 3 → robuste'},
        'c2_ecart_gini': {'statut': 'VERT',
                          'message': 'Écart Gini = 0.0655 ≥ 2%'},
        'c3_coherence': {'statut': 'AMBRE',
                         'message': 'GLM_POISSON rang #1'},
    }

    def _contexte(self, classement=None, controles=None):
        a4 = {'classement': self.CLASSEMENT if classement is None else classement}
        a6 = {'branche': 'auto',
              'validation_selection': (self.CONTROLES if controles is None
                                       else controles)}
        return R._construire_contexte_tarif(A3, a4, a6, 'auto', '31/12/2025')

    def test_les_modeles_compares_entrent_TOUS_dans_le_contexte(self):
        ctx = self._contexte()
        for modele in ('GLM Poisson (référence A3)', 'lightgbm',
                       'xgboost_tweedie'):
            self.assertIn(modele, ctx, modele)
        self.assertIn('CLASSEMENT ML (3 modèle(s) comparé(s))', ctx)
        print('    OK T2 : les 3 modèles comparés sont dans le contexte '
              '(1 sur 7 avant ce lot)')

    def test_les_indicateurs_qui_permettent_de_comparer_y_sont(self):
        """⚠️ C'EST L'ÉCART D'OVERFIT QUI JUSTIFIE LE CHOIX : lightgbm est à
        0,0046 de Gini du GLM, mais à 2,757 de surapprentissage contre 1,000.
        Sans ces chiffres, la « comparaison » n'a rien à dire."""
        ctx = self._contexte()
        for valeur in ('Gini=0.1729', 'overfit=2.757', 'overfit=5.248',
                       'RMSE=0.4504'):
            self.assertIn(valeur, ctx, valeur)
        print('    OK T2b : Gini, RMSE et surapprentissage de chaque modèle')

    def test_les_trois_controles_de_selection_y_sont(self):
        ctx = self._contexte()
        self.assertIn('=== CONTRÔLES DE SÉLECTION ===', ctx)
        for attendu in ('C1 — nombre de modèles comparés : VERT',
                        'C2 — écart de Gini entre modèles : VERT',
                        'C3 — cohérence du modèle retenu : AMBRE'):
            self.assertIn(attendu, ctx, attendu)
        print('    OK T2c : les 3 contrôles de sélection, avec leur statut')

    def test_une_valeur_non_calculee_ne_vaut_pas_zero(self):
        """⚠️ UN GINI À 0,0000 AFFIRMERAIT UN POUVOIR DISCRIMINANT NUL. Sur le
        rapport réel, `score_global` est absent des sept entrées : il doit
        sortir « — », jamais « 0.0000 »."""
        ctx = self._contexte(classement=[{'modele': 'sans_metrique'}])
        self.assertIn('Gini=— | RMSE=— | overfit=— | score=—', ctx)
        self.assertNotIn('0.0000', ctx.split('CLASSEMENT ML')[1].split('===')[0])
        print('    OK T2d : une métrique absente sort « — », pas « 0.0000 »')

    def test_un_classement_vide_SE_DIT(self):
        """L'absence se nomme — elle ne laisse pas une section muette."""
        ctx = self._contexte(classement=[])
        self.assertIn('aucun classement transmis', ctx)
        ctx2 = self._contexte(controles={})
        self.assertIn('non calculé', ctx2)
        print('    OK T2e : un classement vide et un contrôle absent se disent')

    def test_aucun_identifiant_n_est_reintroduit(self):
        """⚠️ CE LOT AJOUTE DES DONNÉES AU CONTEXTE — le chantier C1-C3 vient
        d'établir qu'aucun identifiant n'en sort. Noms d'algorithmes et
        indicateurs, oui ; référence client ou entité, non."""
        motif = re.compile(
            r'entite|entité|ref_client|client_nom|societe|société|'
            r'raison_sociale|siren|siret|numero_police|matricule', re.I)
        fautives = [l for l in self._contexte().split('\n') if motif.search(l)]
        self.assertEqual(fautives, [], '; '.join(fautives))
        print('    OK T2f : aucun motif identifiant dans le contexte enrichi')


class T4_CeQuiNEstPasCalcule(unittest.TestCase):
    """T4 — ne jamais fabriquer un chiffre pour combler un trou.

    ⚠️ RÈGLE DU PROJET, SANS EXCEPTION : une case vide honnête vaut mieux
    qu'un nombre faux ; publier ce qui n'a pas pu être calculé plutôt que de
    le cacher.
    """

    def _html(self, a3=None, a4=None, a6=None):
        with patch.dict(os.environ, {}, clear=False):
            return R.export_html(a3 or {}, a4 or {}, a6 or {},
                                 'DEMO', '31/12/2025', 'T4')

    def test_le_score_vient_du_classement_SCORE_d_A6(self):
        """⚠️ LE SCORE EXISTAIT — LE RAPPORT LISAIT LA MAUVAISE SOURCE. A4
        range les modèles, A6 les SCORE et publie son propre classement. La
        colonne Score était vide sur les SEPT lignes pendant que la section
        « Modèle retenu » affichait 1.0000 : deux vérités sur le même modèle.
        """
        a4 = {'classement': [{'modele': 'GLM', 'gini_test': 0.1775}]}
        a6 = {'classement': [{'modele': 'GLM', 'gini_test': 0.1775,
                              'score_global': 1.0}],
              'modele_production': {'modele': 'GLM', 'score_global': 1.0}}
        html = self._html(a4=a4, a6=a6)
        self.assertIn('1.0000', html)
        # sans A6, on lit A4 — et le score absent sort en tiret, pas en zéro
        html_sans_a6 = self._html(a4=a4)
        self.assertIn('—', html_sans_a6)
        self.assertNotIn('0.0000', html_sans_a6.split('Classement ML')[1][:900])
        print('    OK T4 : le score vient du classement SCORÉ d\'A6')

    def test_une_metrique_absente_ne_vaut_pas_zero(self):
        """⚠️ UN GINI À 0,0000 AFFIRME UN POUVOIR DISCRIMINANT NUL. Le code
        portait déjà un commentaire disant qu'UNE colonne sur quatre avait été
        corrigée ; les trois autres portaient encore le défaut."""
        html = self._html(a3={'metriques': {'poisson': {}}},
                          a6={'modele_production': {'modele': 'GLM'}})
        bloc = html.split('Résultats GLM')[1][:600]
        self.assertNotIn('0.0000', bloc)
        self.assertIn('—', bloc)
        print('    OK T4b : une métrique absente sort « — », jamais « 0.0000 »')

    def test_une_hypothese_non_calculee_NE_DISPARAIT_PLUS(self):
        """⚠️ DOUZE ENDROITS DU DÉPÔT PROMETTENT « H1–H4 » et trois sites
        effaçaient la ligne : le lecteur ne pouvait pas distinguer
        « vérifiée » de « jamais calculée »."""
        a3 = {'hypotheses': {'h1_poisson': {'statut': 'VERT',
                                            'ratio_disp': 1.022}}}
        html = self._html(a3=a3, a6={'modele_production': {}})
        bloc = html.split(R.chapitre(4))[1].split('</table>')[0]
        self.assertIn('NON CALCULÉE', bloc)
        # les quatre hypothèses GLM figurent, calculées ou non
        # ⚠️ T5 les a NOMMÉES : elles s'appelaient « H1 » à « H4 » à côté de
        # « H1 ML » à « H4 ML », soit huit lignes pour quatre numéros.
        for h in ('H1 GLM —', 'H2 GLM —', 'H3 GLM —', 'H4 GLM —'):
            self.assertIn(h, bloc, h)
        print('    OK T4c : les hypothèses non calculées sont NOMMÉES, pas '
              'effacées')

    def test_un_texte_trop_long_est_coupe_SUR_UN_MOT_et_le_dit(self):
        """⚠️ LE RAPPORT COUPAIT EN PLEIN MOT, SANS RIEN DIRE : « le GLM est
        bien spécifié sur toute ». La phrase semblait mal écrite, pas
        tronquée."""
        from core.format_fr import tronque
        long = ('Dispersion homogène entre bandes de risque — le GLM est bien '
                'spécifié sur toute la plage des primes prédites')
        coupe = tronque(long, 80)
        self.assertTrue(coupe.endswith('…'))
        self.assertLessEqual(len(coupe), 81)
        self.assertFalse(coupe[:-1].endswith(' '))
        # la coupure tombe sur une frontière de mot
        self.assertTrue(long.startswith(coupe[:-1]))
        self.assertIn(' ', coupe)
        mot_coupe = coupe[:-1].rsplit(' ', 1)[-1]
        self.assertIn(mot_coupe, long.split())
        print(f'    OK T4d : « …{coupe[-28:]} » — coupé sur un mot, et dit')

    def test_un_texte_court_n_est_pas_touche(self):
        from core.format_fr import tronque
        self.assertEqual(tronque('GLM Poisson adapté', 80), 'GLM Poisson adapté')
        self.assertEqual(tronque(None, 80), '')
        print('    OK T4e : un texte court passe intact')


# ── T5 : un payload complet, pour voir les DEUX formats en entier ───────────
# ⚠️ LES TESTS PRÉCÉDENTS SE CONTENTAIENT D'UN A3 MINIMAL : ils ne rendaient
# ni classement, ni backtesting, ni contrôles de sélection — c'est-à-dire
# précisément les tableaux où les deux formats divergeaient.
A3_COMPLET = {
    'metriques': {
        'poisson': {'gini': 0.1775, 'aic': 10308.0, 'deviance': 8421.5,
                    'pseudo_r2': 0.0412, 'nb_vars_retenues': 9},
        'gamma': {'gini': 0.0912, 'aic': 20114.0, 'deviance': 15012.3,
                  'pseudo_r2': 0.0188, 'nb_vars_retenues': 7},
    },
    'relativites_poisson': {
        'age_conducteur': {'beta': -0.2841, 'relativite': 0.7527,
                           'ic95_low': 0.6902, 'ic95_high': 0.8208,
                           'pvalue': 0.0001, 'significatif': True,
                           'sens': 'allegant'},
    },
    'hypotheses': {
        'h1_poisson': {'statut': 'VERT', 'ratio_disp': 1.022,
                       'message': 'Var/E = 1.02 < 2', 'conseil': 'GLM adapte'},
    },
    'commentaire': 'REPLI AGENT',
}
A4_COMPLET = {'classement': [], 'commentaire': 'REPLI AGENT'}
A6_COMPLET = {
    'branche': 'auto', 'statut_rag': 'VERT',
    'classement': [
        {'modele': 'GLM_POISSON', 'famille': 'GLM', 'gini_test': 0.1775,
         'rmse_test': 0.4504, 'overfit_ratio': 1.0, 'score_global': 1.0},
        {'modele': 'ML_LIGHTGBM', 'famille': 'ML', 'gini_test': 0.1729,
         'rmse_test': 0.4579, 'overfit_ratio': 2.757, 'score_global': 0.6696},
    ],
    'modele_production': {'modele': 'GLM_POISSON', 'famille': 'GLM',
                          'score_global': 1.0, 'gini_test': 0.1775,
                          'overfit_ratio': 1.0, 'interpretabilite': 1.0},
    'validation_selection': {
        'c1_nb_modeles': {'statut': 'VERT', 'message': '7 modeles compares'},
        'c2_ecart_gini': {'statut': 'VERT', 'message': 'Ecart = 0.0655'},
        'c3_coherence': {'statut': 'VERT', 'message': 'Rang 1 coherent'},
    },
    'backtest': {
        'disponible': True, 'ae_ratio': 0.9947, 'interpretation': 'Bon',
        'stabilite_wf': 'Stable', 'n_fenetres': 4,
        'walk_forward': [
            {'annee_test': 2023, 'n_train': 9528, 'n_test': 2472,
             'moy_train': 0.21, 'moy_test': 0.21, 'ae_ratio': 0.9947,
             'statut': 'VERT'},
        ],
    },
    'audit_trail': {'agent': 'A6_COMPARAISON', 'ae_ratio': 0.9947,
                    'stabilite_wf': 'Stable', 'nb_modeles': 7,
                    'gouvernance_ok': True, 'modele_production': 'GLM_POISSON',
                    'timestamp': '2026-08-08T06:34:34.813112'},
}


def _les_deux_formats():
    """Le MÊME payload dans les deux formats, sans réseau."""
    appels = []
    with patch.dict(os.environ, {'ANTHROPIC_API_KEY': 'cle-de-test'}):
        with patch('anthropic.Anthropic', _client_simule(appels)):
            rap = R.generer_rapport_tarification(
                A3_COMPLET, A4_COMPLET, A6_COMPLET, 'DEMO', '31/12/2025',
                'T5', ['html', 'word'])
    html = (rap.get('html_bytes') or b'').decode('utf-8', 'replace')
    return html, rap.get('word_bytes') or b''


def _entetes_html(html):
    return {re.sub(r'<[^>]+>', '', c).strip()
            for c in re.findall(r'<th[^>]*>(.*?)</th>', html, re.S)}


def _entetes_word(octets):
    """Les en-têtes du Word : la 1re ligne de chaque tableau du document."""
    from docx import Document
    doc = Document(io.BytesIO(octets))
    return {c.text.strip() for t in doc.tables for c in t.rows[0].cells}


class T5_LeVocabulaire(unittest.TestCase):
    """T5 — un en-tête défini deux fois diverge.

    ⚠️ MESURÉ AVANT CE LOT sur un payload réel : 21 libellés communs aux deux
    formats, 12 propres à l'HTML, 10 propres au Word. Les mêmes colonnes
    portaient « IC 95% bas » d'un côté et « IC bas » de l'autre.
    """

    def test_les_titres_de_colonnes_vivent_a_UN_SEUL_endroit(self):
        """Le motif du chantier : la source est unique, les formats la lisent.

        Un test qui recopierait les libellés attendus recréerait la deuxième
        définition qu'on vient de supprimer — il lit donc la source.
        """
        html, word = _les_deux_formats()
        h, w = _entetes_html(html), _entetes_word(word)
        for cle in ('glm', 'relativites', 'classement', 'hypotheses',
                    'backtest', 'controles', 'audit'):
            for titre in R.titres(cle):
                self.assertIn(titre, h, f'HTML · {cle} · {titre}')
                self.assertIn(titre, w, f'WORD · {cle} · {titre}')
        print(f'    OK T5 : {sum(len(R.titres(c)) for c in R.COLONNES)} '
              f'libellés, une seule définition, lue par les deux formats')

    def test_ce_qui_reste_propre_a_UN_format_est_DECLARE(self):
        """⚠️ UN CHOIX SE DÉCLARE ; CE QUI NE SE DÉCLARE PAS EST UN OUBLI.
        Avant ce lot, six colonnes et un tableau entier manquaient au Word
        sans qu'aucun commentaire du dépôt ne le mentionne."""
        html, word = _les_deux_formats()
        h, w = _entetes_html(html), _entetes_word(word)
        # les deux tableaux propres au Word, page de garde et fiche modèle
        declares_word = set(R.titres('garde')) | set(R.titres('production'))
        self.assertEqual(w - h - declares_word, set())
        # l'étoile décorative, propre à l'HTML
        self.assertEqual(h - w, {'⭐'})
        print(f'    OK T5b : {len(h - w)} propre HTML + '
              f'{len(w - h)} propres Word, tous déclarés')

    def test_les_huit_hypotheses_portent_HUIT_noms(self):
        """⚠️ QUATRE S'APPELAIENT « H1 » À « H4 » ET QUATRE « H1 ML » À
        « H4 ML », sous un titre qui annonçait « H1–H4 »."""
        libelles = [lib for _, lib, _ in R.HYPOTHESES]
        self.assertEqual(len(libelles), 8)
        self.assertEqual(len(set(libelles)), 8)
        for lib in libelles:
            self.assertRegex(lib, r'^H[1-4] (GLM|ML) — ')
        html, _ = _les_deux_formats()
        self.assertNotRegex(html, r'>H[1-4] — ')
        print('    OK T5c : 8 hypothèses, 8 noms, la famille dans chacun')

    def test_les_deux_formats_annoncent_les_MEMES_chapitres(self):
        """⚠️ « §1 — Résultats GLM Poisson / Gamma / Tweedie » en HTML et
        « 1. Résultats GLM — Poisson / Gamma / Tweedie » en Word : 8 titres,
        0 identique."""
        html, word = _les_deux_formats()
        texte_word = _texte_docx(word)
        for n in range(1, len(R.CHAPITRES) + 1):
            titre = R.chapitre(n)
            self.assertIn(titre, html, f'HTML · {titre}')
            self.assertIn(titre, texte_word, f'WORD · {titre}')
        print(f'    OK T5d : {len(R.CHAPITRES)} chapitres, mot pour mot dans '
              f'les deux formats')

    def test_le_rapport_ne_dit_plus_paragraphe_N_comme_la_narration(self):
        """⚠️ « §4 » DÉSIGNAIT DEUX CHOSES : le chapitre « Hypothèses » du
        rapport et, dans le commentaire, la section « COMPARAISON DES MODÈLES »
        que le prompt impose. Un renvoi « voir §4 » n'y voulait rien dire."""
        for titre in R.CHAPITRES:
            self.assertNotIn('§', titre)
        html, _ = _les_deux_formats()
        # les « §N » restants sont ceux de la narration, et d'elle seule
        narration = html.split('class="section-body narration"')[1]
        for m in re.finditer(r'§\s*\d', html):
            self.assertGreaterEqual(
                m.start(), html.index('class="section-body narration"'),
                'un « §N » hors du commentaire actuariel')
        self.assertIn('§1', narration)
        print('    OK T5e : « §N » n\'appartient plus qu\'au commentaire')

    def test_un_modele_porte_UN_nom(self):
        """⚠️ « GLM Poisson (référence A3) » au classement et « GLM_POISSON »
        deux chapitres plus loin. Et la même chaîne produit deux écritures
        selon le passage : « lightgbm » ici, « ML_LIGHTGBM » là."""
        self.assertEqual(R.nom_modele('GLM_POISSON'), 'GLM Poisson')
        self.assertEqual(R.nom_modele('ML_LIGHTGBM'), 'LightGBM')
        self.assertEqual(R.nom_modele('lightgbm'), 'LightGBM')
        self.assertEqual(R.nom_modele('ML_LINEAIRE_REGULARISE'),
                         'Linéaire régularisé')
        # la provenance est une information, pas un nom : elle est conservée
        self.assertEqual(R.nom_modele('GLM Poisson (référence A3)'),
                         'GLM Poisson (référence A3)')
        # ⚠️ UN NOM INCONNU N'EST JAMAIS REMPLACÉ — un nom inventé serait pire
        self.assertEqual(R.nom_modele('MODELE_FUTUR'), 'MODELE_FUTUR')
        self.assertEqual(R.nom_modele(None), '—')
        html, word = _les_deux_formats()
        texte_word = _texte_docx(word)
        for forme in ('GLM_POISSON', 'ML_LIGHTGBM'):
            self.assertNotIn(forme, html, f'HTML · {forme}')
            self.assertNotIn(forme, texte_word, f'WORD · {forme}')
        print('    OK T5f : les identifiants techniques ne sortent plus')

    def test_la_piste_d_audit_se_lit_en_francais(self):
        """⚠️ « Ae Ratio », « Stabilite Wf », « Gouvernance Ok : True » et un
        horodatage à la microseconde — des noms de variables capitalisés."""
        self.assertEqual(R.libelle_audit('ae_ratio'), 'Ratio A/E')
        self.assertEqual(R.valeur_audit(True), 'oui')
        self.assertEqual(R.valeur_audit(False), 'non')
        self.assertEqual(R.valeur_audit('2026-08-08T06:34:34.813112'),
                         '08/08/2026 à 06 h 34')
        # ⚠️ UNE CLÉ INCONNUE RESTE VISIBLE : une piste d'audit amputée n'en
        # serait plus une.
        self.assertEqual(R.libelle_audit('cle_inconnue'), 'Cle inconnue')
        html, word = _les_deux_formats()
        texte_word = _texte_docx(word)
        for brut in ('Ae Ratio', 'Stabilite Wf', 'Nb Modeles',
                     'Gouvernance Ok'):
            self.assertNotIn(brut, html, f'HTML · {brut}')
            self.assertNotIn(brut, texte_word, f'WORD · {brut}')
        self.assertIn('Ratio A/E', html)
        self.assertIn('Ratio A/E', texte_word)
        print('    OK T5g : 14 clés techniques devenues des libellés')

    def test_le_Word_porte_les_colonnes_et_le_tableau_qui_lui_manquaient(self):
        """⚠️ SIX COLONNES ET UN TABLEAU ENTIER. Le fichier qui part chez un
        commissaire n'avait ni la déviance, ni le RMSE, ni l'effectif de test,
        ni les moyennes, ni les trois contrôles qui justifient la sélection."""
        _, word = _les_deux_formats()
        texte = _texte_docx(word)
        for manquant in ('Déviance', 'RMSE test', 'N test', 'Moy train',
                         'Moy test', 'Contrôle sélection'):
            self.assertIn(manquant, texte, manquant)
        for controle in ('C1 — Nombre de modèles', 'C2 — Écart Gini',
                         'C3 — Cohérence'):
            self.assertIn(controle, texte, controle)
        print('    OK T5h : 6 colonnes + 3 contrôles rendus au Word')

    def test_aucun_tableau_du_Word_ne_deborde_de_la_page(self):
        """⚠️ AJOUTER UNE COLONNE PEUT COÛTER LA LISIBILITÉ : la largeur utile
        est de 16,5 cm, marges déduites. Elle se vérifie, elle ne s'espère
        pas."""
        from docx import Document
        from docx.shared import Cm
        _, word = _les_deux_formats()
        doc = Document(io.BytesIO(word))
        for n, t in enumerate(doc.tables, 1):
            largeur = sum((c.width or 0) / Cm(1) for c in t.rows[0].cells)
            self.assertLessEqual(round(largeur, 1), 16.5,
                                 f'tableau {n} : {largeur:.1f} cm')
        print(f'    OK T5i : {len(doc.tables)} tableaux, tous dans 16,5 cm')

    def test_une_hypothese_absente_se_voit_AUSSI_dans_le_Word(self):
        """⚠️ LE WORD LA FAISAIT DISPARAÎTRE EN SILENCE là où l'HTML la publie
        « NON CALCULÉE » depuis T4 : le livrable signé était le plus indulgent
        des deux."""
        _, word = _les_deux_formats()
        texte = _texte_docx(word)
        # A3_COMPLET ne porte que h1_poisson : les sept autres sont absentes
        self.assertIn('NON CALCULÉE', texte)
        for _, libelle, _ in R.HYPOTHESES:
            self.assertIn(libelle, texte, libelle)
        print('    OK T5j : 8 hypothèses nommées dans le Word, calculées ou '
              'non')


# ── T7a : le catalogue des figures ─────────────────────────────────────────
#
# ⚠️ AUCUN TEST D'ICI N'APPELLE kaleido POUR DE VRAI. Le lot précédent a
# mesuré ce que coûte une gate qui rasterise : 94 conversions, 562 s, et une
# suite qui n'exerçait pas le même code selon la machine. `rasteriser` est
# une couture ; les tests la substituent.
#
# ⚠️ ET LE PNG DE SUBSTITUTION EST ÉCRIT ICI PLUTÔT QU'IMPORTÉ D'A7 : un
# test de tarification qui importerait un test de provisionnement coupleroit
# les deux sous-directions, alors que l'indépendance de la tarification est
# précisément ce qui permet de ne pas faire tourner `provisionnement` sur un
# lot qui ne le touche pas. Dix lignes valent mieux qu'un couplage.

def _png_minimal(largeur, hauteur):
    """Un PNG valide, sans aucune dépendance — zlib et struct suffisent."""
    def _bloc(typ, donnees):
        return (struct.pack('>I', len(donnees)) + typ + donnees
                + struct.pack('>I', zlib.crc32(typ + donnees) & 0xFFFFFFFF))

    entete = struct.pack('>IIBBBBB', largeur, hauteur, 8, 2, 0, 0, 0)
    ligne = b'\x00' + bytes((11, 30, 61)) * largeur
    return (b'\x89PNG\r\n\x1a\n' + _bloc(b'IHDR', entete)
            + _bloc(b'IDAT', zlib.compress(ligne * hauteur, 9))
            + _bloc(b'IEND', b''))


class rendu_simule:
    """Substitue le rasteriseur et déclare kaleido présent.

    `svg_lourd` reproduit le cas mesuré du QQ-plot : un SVG plus lourd que
    le PNG, où la règle doit basculer.

    ⚠️ LES PNG SONT DISTINCTS : OOXML déduplique les médias identiques —
    quatorze images identiques ne laissent qu'UN fichier dans `word/media/`.
    """

    def __init__(self, present=True, svg_lourd=False):
        self._present, self._lourd = present, svg_lourd
        self.appels = []

    def __enter__(self):
        self._vk, self._vr = R.kaleido_disponible, R.rasteriser
        R.kaleido_disponible = lambda: self._present

        def _faux(figure, format):
            self.appels.append((getattr(figure, 'nom', '?'), format))
            n = len(self.appels)
            if format == 'png':
                return _png_minimal(200 + n, 100 + n)
            corps = b'x' * (40000 if self._lourd else 200)
            return (b'<?xml version="1.0"?><svg xmlns="http://www.w3.org/'
                    b'2000/svg"><!--' + corps + b'--></svg>')
        R.rasteriser = _faux
        return self

    def __exit__(self, *_):
        R.kaleido_disponible, R.rasteriser = self._vk, self._vr
        return False


class _FigureFactice:
    """Une figure Plotly de substitution — elle sait rendre son fragment."""

    def __init__(self, nom):
        self.nom = nom

    def to_html(self, **_):
        return '<div class="plotly-graph-div">%s</div>' % self.nom


def _payload_figures(cles=None):
    """Les trois résultats d'agent, portant les figures demandées."""
    cles = list(R.SOURCES_FIGURES) if cles is None else list(cles)
    res = {'a3': {}, 'a4': {}, 'a6': {}}
    for cle in cles:
        agent, dico = R.SOURCES_FIGURES[cle]
        res[agent].setdefault(dico, {})[cle] = _FigureFactice(cle)
    return res['a3'], res['a4'], res['a6']


class T7a_LeCatalogueDesFigures(unittest.TestCase):
    """T7a — 30 figures produites, 6 publiées, et un dictionnaire entier que
    personne ne lisait."""

    def test_toute_figure_produite_par_la_chaine_est_CLASSEE(self):
        """⚠️ LA RAISON D'UN ÉCART VIT DANS LE CODE, PAS DANS UN COMMENTAIRE.

        Ce test relève les figures que les agents produisent réellement et
        vérifie que chacune est soit au plan, soit écartée AVEC SA RAISON.
        Une figure nouvelle qui ne serait ni l'un ni l'autre fait tomber la
        gate — c'est ce qui empêche le catalogue de se périmer en silence.
        """
        racine = os.path.abspath(os.path.join(os.path.dirname(__file__), '..',
                                              '..'))
        motif = re.compile(r"""graphiques(?:_validation)?\s*\[\s*['"]"""
                           r"""([A-Za-z0-9_]+)['"]\s*\]\s*=""")
        produites = set()
        for sous in ('a3_glm', 'a4_ml', 'a5_deep_learning', 'a6_comparaison'):
            chemin = os.path.join(racine, 'direction_non_vie', 'tarification',
                                  sous, 'agent.py')
            if not os.path.exists(chemin):
                continue
            with io.open(chemin, encoding='utf-8', errors='replace') as f:
                produites |= set(motif.findall(f.read()))
        au_plan = {c for _, cles in R.PLAN_FIGURES for c in cles}
        self.assertTrue(produites, 'le relevé ne trouve aucune figure')
        non_classees = produites - au_plan - set(R.FIGURES_ECARTEES)
        self.assertEqual(non_classees, set(),
                         'figures ni publiées ni écartées : %s' % non_classees)
        # et rien d'écarté ne doit être au plan en même temps
        self.assertEqual(au_plan & set(R.FIGURES_ECARTEES), set())
        print('    OK T7a : %d figures produites, %d publiées, %d écartées '
              'avec leur raison'
              % (len(produites), len(au_plan), len(produites - au_plan)))

    def test_chaque_figure_du_plan_a_un_titre_et_une_source(self):
        au_plan = [c for _, cles in R.PLAN_FIGURES for c in cles]
        self.assertEqual(len(au_plan), len(set(au_plan)), 'doublon au plan')
        for cle in au_plan:
            self.assertIn(cle, R.TITRES_FIGURES, cle)
            self.assertIn(cle, R.SOURCES_FIGURES, cle)
        self.assertEqual(set(R.TITRES_FIGURES), set(au_plan))
        self.assertEqual(set(R.SOURCES_FIGURES), set(au_plan))
        print('    OK T7a-b : %d figures, un titre et une source chacune, '
              'aucune orpheline' % len(au_plan))

    def test_cinq_figures_viennent_du_dictionnaire_que_personne_ne_lisait(self):
        """⚠️ `graphiques_validation` : 25 sites de production dans le dépôt,
        ZÉRO lecture. Le chapitre 4 publiait huit verdicts sans preuve."""
        validation = [c for c, (_, d) in R.SOURCES_FIGURES.items()
                      if d == 'graphiques_validation']
        self.assertEqual(len(validation), 5)
        for attendue in ('sur_dispersion_poisson', 'overfitting_train_test',
                         'scores_multicriteres'):
            self.assertIn(attendue, validation)
        a3, a4, a6 = _payload_figures()
        trouvees = R.figures_disponibles(a3, a4, a6)
        self.assertEqual(len(trouvees), len(R.SOURCES_FIGURES))
        print('    OK T7a-c : %d figures tirées de `graphiques_validation`'
              % len(validation))

    def test_la_numerotation_reste_CONTINUE_quand_une_figure_s_efface(self):
        """⚠️ LE COMPTEUR N'AVANCE QUE SUR UNE FIGURE PRÉSENTE. Une
        numérotation précalculée laisserait un trou : « Figure 3 » puis
        « Figure 5 »."""
        toutes = [c for _, cles in R.PLAN_FIGURES for c in cles]
        for retiree in ('chart_shap_summary', 'sur_dispersion_poisson',
                        'aic_comparaison'):
            restantes = [c for c in toutes if c != retiree]
            plan = R.numeroter(R.figures_disponibles(
                *_payload_figures(restantes)))
            numeros = [f.numero for fs in plan.values() for f in fs]
            self.assertEqual(sorted(numeros),
                             list(range(1, len(restantes) + 1)),
                             'trou dans la numérotation sans %s' % retiree)
        vide = R.numeroter({})
        self.assertEqual(vide, {})
        print('    OK T7a-d : numérotation 1..N sans trou, quelle que soit la '
              'figure qui manque')

    def test_les_numeros_suivent_l_ordre_du_PLAN(self):
        plan = R.numeroter(R.figures_disponibles(*_payload_figures()))
        attendu = [c for _, cles in R.PLAN_FIGURES for c in cles]
        obtenu = [f.cle for _, fs in sorted(plan.items()) for f in fs]
        self.assertEqual(obtenu, attendu)
        print('    OK T7a-e : les %d numéros suivent le plan, chapitre par '
              'chapitre' % len(obtenu))

    def test_le_format_est_DECIDE_PAR_LA_MESURE_pas_par_une_liste(self):
        """⚠️ MESURÉ SUR LES FIGURES RÉELLES : douze pèsent moins en SVG, et
        le QQ-plot — 400 points — passe de 49,8 ko en PNG à 940,0 ko en SVG.
        Une règle « tout en SVG » aurait multiplié son poids par dix-neuf."""
        fig = _FigureFactice('t')
        with rendu_simule(svg_lourd=False):
            leger = R.rendre_figure(fig)
        with rendu_simule(svg_lourd=True):
            lourd = R.rendre_figure(fig)
        self.assertEqual(leger.genre, 'svg')
        self.assertTrue(leger.html.startswith('<svg'))
        self.assertEqual(lourd.genre, 'png')
        self.assertIn('data:image/png;base64,', lourd.html)
        # le Word reçoit un PNG dans les deux cas
        self.assertTrue(leger.png and lourd.png)
        print('    OK T7a-f : SVG quand il est plus léger, PNG quand il ne '
              'l\'est pas — la règle compare, elle ne liste pas')

    def test_sans_kaleido_le_repli_est_INTERACTIF_et_DECLARE(self):
        """⚠️ ON DÉGRADE VERS LE COMPORTEMENT D'AUJOURD'HUI, NOMMÉ — pas vers
        rien. Et aujourd'hui, hors ligne, le cadre reste vide EN SILENCE."""
        with rendu_simule(present=False):
            r = R.rendre_figure(_FigureFactice('g'))
        self.assertEqual(r.genre, 'interactif')
        self.assertIsNone(r.png)
        self.assertIn('plotly-graph-div', r.html)
        self.assertIn('kaleido', r.note)
        self.assertIn('connexion', r.note)
        print('    OK T7a-g : repli interactif, et la raison est DANS le '
              'document')

    def test_les_deux_formats_portent_les_MEMES_figures_numerotees(self):
        """⚠️ LE POINT DU LOT. A7 tient deux compteurs positionnels
        indépendants, rattrapés par un test. Ici la numérotation est calculée
        UNE fois : « Figure 7 » ne PEUT pas désigner deux choses."""
        a3, a4, a6 = _payload_figures()
        a3.update(A3_COMPLET)
        a6.update({k: v for k, v in A6_COMPLET.items()
                   if k not in ('graphiques', 'graphiques_validation')})
        with rendu_simule():
            html = R.export_html(a3, a4, a6, 'DEMO', '31/12/2025', 'T7a',
                                 narration_calculee=('Texte.', 'temoin'))
            word = R.export_word(a3, a4, a6, 'DEMO', '31/12/2025', 'T7a',
                                 narration_calculee=('Texte.', 'temoin'))
        legendes_html = re.findall(r'Figure (\d+) — ([^<]{4,120})', html)
        texte_word = _texte_docx(word)
        self.assertTrue(legendes_html, 'aucune figure dans l\'HTML')
        self.assertEqual([n for n, _ in legendes_html],
                         [str(i) for i in range(1, len(legendes_html) + 1)])
        for numero, titre in legendes_html:
            self.assertIn('Figure %s — %s' % (numero, titre), texte_word,
                          'Figure %s absente ou différente dans le Word'
                          % numero)
        print('    OK T7a-h : %d figures, mêmes numéros et mêmes titres dans '
              'les deux formats' % len(legendes_html))

    def test_le_word_porte_bien_les_images(self):
        """⚠️ LE WORD NE PORTAIT AUCUNE FIGURE — zéro image mesurée sur
        39 414 octets, quand l'HTML en portait six."""
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            word = R.export_word(a3, a4, a6, 'DEMO', '31/12/2025', 'T7a',
                                 narration_calculee=('Texte.', 'temoin'))
        medias = [n for n in zipfile.ZipFile(io.BytesIO(word)).namelist()
                  if n.startswith('word/media/')]
        self.assertEqual(len(medias), len(R.SOURCES_FIGURES),
                         'les images ne sont pas toutes entrées — et OOXML '
                         'déduplique les médias identiques')
        print('    OK T7a-i : %d images dans le .docx (0 avant ce lot)'
              % len(medias))

    def test_plotly_n_est_charge_QUE_si_une_figure_reste_interactive(self):
        """Une image n'a besoin d'aucune bibliothèque : un rapport qui n'en
        porte plus ne doit pas réclamer le réseau pour rien."""
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            html_images = R.export_html(a3, a4, a6, 'D', '31/12/2025', 'T',
                                        narration_calculee=('T.', 'temoin'))
        with rendu_simule(present=False):
            html_replis = R.export_html(a3, a4, a6, 'D', '31/12/2025', 'T',
                                        narration_calculee=('T.', 'temoin'))
        self.assertNotIn('cdn.plot.ly', html_images)
        self.assertIn('cdn.plot.ly', html_replis)
        print('    OK T7a-j : aucun appel réseau quand les figures sont des '
              'images')


class T7b_LePlacementDesFigures(unittest.TestCase):
    """T7b — une figure vit dans le chapitre qu'elle illustre.

    ⚠️ ELLES ÉTAIENT GROUPÉES dans une section « Figures » placée après le
    chapitre 6, hors de la numérotation des chapitres. Le chapitre 4 publiait
    huit verdicts et renvoyait sa preuve trois chapitres plus loin.
    """

    def _positions(self, texte, chapitres):
        """L'index de chaque titre de chapitre dans le document rendu."""
        return {n: texte.index(R.chapitre(n)) for n in chapitres
                if R.chapitre(n) in texte}

    def test_chaque_figure_est_DANS_son_chapitre_en_HTML(self):
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            html = R.export_html(a3, a4, a6, 'DEMO', '31/12/2025', 'T7b',
                                 narration_calculee=('Texte.', 'temoin'))
        bornes = self._positions(html, range(1, 9))
        plan = R.numeroter(R.figures_disponibles(a3, a4, a6))
        for numero, figures in plan.items():
            debut = bornes[numero]
            # la borne haute : le titre du chapitre suivant qui existe
            suivants = [p for _, p in sorted(bornes.items()) if p > debut]
            fin = suivants[0] if suivants else len(html)
            for f in figures:
                place = html.find(R.legende(f))
                self.assertNotEqual(place, -1, R.legende(f))
                self.assertTrue(debut < place < fin,
                                '%s hors du chapitre %d' % (R.legende(f),
                                                            numero))
        print('    OK T7b : %d figures, chacune entre son chapitre et le '
              'suivant' % sum(len(f) for f in plan.values()))

    def test_chaque_figure_est_DANS_son_chapitre_en_WORD(self):
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            word = R.export_word(a3, a4, a6, 'DEMO', '31/12/2025', 'T7b',
                                 narration_calculee=('Texte.', 'temoin'))
        texte = _texte_docx(word)
        bornes = self._positions(texte, range(1, 9))
        plan = R.numeroter(R.figures_disponibles(a3, a4, a6))
        for numero, figures in plan.items():
            debut = bornes[numero]
            suivants = [p for _, p in sorted(bornes.items()) if p > debut]
            fin = suivants[0] if suivants else len(texte)
            for f in figures:
                place = texte.find(R.legende(f))
                self.assertNotEqual(place, -1, R.legende(f))
                self.assertTrue(debut < place < fin,
                                '%s hors du chapitre %d' % (R.legende(f),
                                                            numero))
        print('    OK T7b-b : les mêmes figures dans les mêmes chapitres, '
              'côté Word')

    def test_la_section_groupee_a_DISPARU(self):
        """⚠️ SANS CE VERROU, LES DEUX PLACEMENTS POURRAIENT COEXISTER —
        chaque figure publiée deux fois, et le lecteur incapable de savoir
        laquelle fait foi."""
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            html = R.export_html(a3, a4, a6, 'DEMO', '31/12/2025', 'T7b',
                                 narration_calculee=('Texte.', 'temoin'))
            word = R.export_word(a3, a4, a6, 'DEMO', '31/12/2025', 'T7b',
                                 narration_calculee=('Texte.', 'temoin'))
        self.assertNotIn('<div class="section-head">Figures</div>', html)
        for f in R.numeroter(R.figures_disponibles(a3, a4, a6))[1]:
            self.assertEqual(html.count(R.legende(f)), 1,
                             '%s publiée deux fois' % R.legende(f))
            self.assertEqual(_texte_docx(word).count(R.legende(f)), 1)
        print('    OK T7b-c : aucune section « Figures », aucune figure en '
              'double')

    def test_le_deplacement_N_A_PAS_TOUCHE_A_LA_NUMEROTATION(self):
        """⚠️ LE CAS OÙ UN COMPTEUR POSITIONNEL SE PROUVE. `numeroter`
        parcourt le plan AVANT tout rendu : déplacer le geste de rendu ne peut
        pas décaler un numéro. Une numérotation faite au fil du document
        n'aurait pas offert cette garantie."""
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            html = R.export_html(a3, a4, a6, 'DEMO', '31/12/2025', 'T7b',
                                 narration_calculee=('Texte.', 'temoin'))
        numeros = [int(n) for n in re.findall(r'Figure (\d+) — ', html)]
        self.assertEqual(numeros, list(range(1, len(numeros) + 1)),
                         'numérotation à trous : %s' % numeros)
        # et elle reste continue quand une figure s'efface d'un chapitre
        for retiree in ('aic_comparaison', 'chart_lorenz_gini',
                        'radar_modele_retenu'):
            restantes = [c for c in R.SOURCES_FIGURES if c != retiree]
            b3, b4, b6 = _payload_figures(restantes)
            with rendu_simule():
                h = R.export_html(b3, b4, b6, 'D', '31/12/2025', 'T',
                                  narration_calculee=('T.', 'temoin'))
            n = [int(x) for x in re.findall(r'Figure (\d+) — ', h)]
            self.assertEqual(n, list(range(1, len(n) + 1)),
                             'trou sans %s : %s' % (retiree, n))
        print('    OK T7b-d : numérotation 1..%d continue après le '
              'déplacement, et à chaque figure retirée' % len(numeros))


class T7c_LImpressionEtLesTroncatures(unittest.TestCase):
    """T7c — ce que le document TAIT : une couleur qui disparaît au papier,
    des lignes coupées sans le dire."""

    def _html(self, **kw):
        a3, a4, a6 = _payload_figures()
        a3.update(kw.pop('a3', {}))
        a6.update(kw.pop('a6', {}))
        with rendu_simule():
            return R.export_html(a3, a4, a6, 'DEMO', '31/12/2025', 'T7c',
                                 narration_calculee=('Texte.', 'temoin'))

    def test_les_fonds_colores_survivent_a_l_impression(self):
        """⚠️ LE SEUL DÉFAUT DE CE CHANTIER QUI PUISSE TROMPER UN LECTEUR.

        Un navigateur supprime par défaut les fonds à l'impression. Le signal
        RAG de ce rapport ne vit QUE dans le fond des badges : sans cette
        règle, « ROUGE » s'imprime en blanc sur blanc, avec le même texte
        noir que « VERT ».
        """
        html = self._html()
        self.assertIn('@media print', html)
        for regle in ('-webkit-print-color-adjust: exact',
                      'print-color-adjust: exact'):
            self.assertIn(regle, html, regle)
        # et le signal RAG dépend bien d'un fond, ce qui justifie la règle
        self.assertIn('.badge-rouge{background:var(--rouge);}', html)
        print('    OK T7c : les fonds du signal RAG survivent au papier')

    def test_les_regles_d_impression_essentielles_sont_TOUTES_la(self):
        """Mesuré avant ce lot : ZÉRO règle d'impression, contre 21 dans le
        rapport de provisionnement."""
        html = self._html()
        attendues = {
            '@page': 'size: A4',
            'marges': 'margin: 14mm 12mm 16mm 12mm',
            'rupture interne': 'break-inside: avoid',
            'ancienne rupture': 'page-break-inside: avoid',
            'en-tête répété': 'display: table-header-group',
            'titre non orphelin': 'break-after: avoid',
            'lignes veuves': 'orphans: 3',
            'lignes isolées': 'widows: 3',
        }
        for nom, regle in attendues.items():
            self.assertIn(regle, html, nom)
        print(f'    OK T7c-b : les {len(attendues)} règles d\'impression '
              f'sont émises (0 avant ce lot)')

    def test_le_bloc_media_print_est_BIEN_FORME(self):
        """⚠️ UNE ACCOLADE MANQUANTE DANS UN `@media` AVALE TOUT CE QUI SUIT.
        Le style est construit par une f-string à accolades doublées : une
        seule oubliée casse le bloc sans que rien ne le dise."""
        html = self._html()
        style = html.split('<style>')[1].split('</style>')[0]
        self.assertEqual(style.count('{'), style.count('}'),
                         'accolades déséquilibrées dans le style')
        bloc = style[style.index('@media print'):]
        # le bloc se referme, et il contient bien ses règles
        profondeur, fin = 0, None
        for i, c in enumerate(bloc):
            if c == '{':
                profondeur += 1
            elif c == '}':
                profondeur -= 1
                if profondeur == 0:
                    fin = i
                    break
        self.assertIsNotNone(fin, 'le bloc @media print ne se referme pas')
        self.assertIn('@page', bloc[:fin])
        print('    OK T7c-c : le bloc @media print est équilibré et complet')

    def test_le_word_DECLARE_qu_il_coupe_les_relativites(self):
        """⚠️ MÊMES COLONNES, MÊMES LIBELLÉS, ET PAS LES MÊMES LIGNES. Le
        Word coupait à quinze sans rien dire ; l'HTML les rend toutes."""
        rels = {f'variable_{i:02d}': {
            'beta': 0.5 - i * 0.01, 'relativite': 1.0 + i * 0.01,
            'ic95_low': 0.9, 'ic95_high': 1.1, 'pvalue': 0.01,
            'significatif': True, 'sens': 'aggravant'} for i in range(22)}
        a3, a4, a6 = _payload_figures()
        a3['relativites_poisson'] = rels
        with rendu_simule():
            word = _texte_docx(R.export_word(
                a3, a4, a6, 'DEMO', '31/12/2025', 'T7c',
                narration_calculee=('Texte.', 'temoin')))
            html = R.export_html(a3, a4, a6, 'DEMO', '31/12/2025', 'T7c',
                                 narration_calculee=('Texte.', 'temoin'))
        self.assertIn('15 relativités sur 22', word)
        self.assertIn('|β| décroissant', word)
        # l'HTML les porte toutes, et n'a donc rien à déclarer
        self.assertEqual(sum(1 for v in rels if v in html), 22)
        self.assertNotIn('relativités sur 22', html)
        print('    OK T7c-d : « 15 relativités sur 22 » — la coupe est DITE')

    def test_le_word_DECLARE_qu_il_coupe_le_classement(self):
        cl = [{'modele': f'MODELE_{i:02d}', 'famille': 'ML',
               'gini_test': 0.2 - i * 0.01, 'rmse_test': 0.4,
               'overfit_ratio': 1.0, 'score_global': 1.0 - i * 0.05}
              for i in range(14)]
        a3, a4, a6 = _payload_figures()
        a6['classement'] = cl
        with rendu_simule():
            word = _texte_docx(R.export_word(
                a3, a4, a6, 'DEMO', '31/12/2025', 'T7c',
                narration_calculee=('Texte.', 'temoin')))
        self.assertIn('10 modèles sur 14', word)
        self.assertIn('ordre du classement', word)
        print('    OK T7c-e : « 10 modèles sur 14 » — la coupe est DITE')

    def test_rien_n_est_declare_quand_rien_n_est_coupe(self):
        """⚠️ UNE DÉCLARATION INUTILE EST UN BRUIT. Sous le seuil, le
        document ne dit rien parce qu'il n'y a rien à dire."""
        self.assertEqual(R.note_troncature(15, 8, 'relativités', 'triées'), '')
        self.assertEqual(R.note_troncature(15, 15, 'relativités', 'triées'), '')
        self.assertIn('sur 16', R.note_troncature(15, 16, 'x', 'y'))
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            word = _texte_docx(R.export_word(
                a3, a4, a6, 'DEMO', '31/12/2025', 'T7c',
                narration_calculee=('Texte.', 'temoin')))
        self.assertNotIn('sur 0', word)
        self.assertNotIn('relativités sur', word)
        print('    OK T7c-f : aucune déclaration là où rien n\'est coupé')

    def test_la_limite_et_la_phrase_lisent_la_MEME_constante(self):
        """⚠️ DEUX COPIES AURAIENT FINI PAR DIRE DEUX CHOSES : le tableau
        coupé à quinze et la phrase annonçant douze. C'est le motif que ce
        chantier a fermé cinq fois."""
        rels = {f'v{i:02d}': {'beta': 0.5 - i * 0.01, 'relativite': 1.0,
                              'ic95_low': 0.9, 'ic95_high': 1.1,
                              'pvalue': 0.01, 'significatif': True,
                              'sens': 'aggravant'} for i in range(30)}
        a3, a4, a6 = _payload_figures()
        a3['relativites_poisson'] = rels
        with rendu_simule():
            octets = R.export_word(a3, a4, a6, 'DEMO', '31/12/2025', 'T7c',
                                   narration_calculee=('Texte.', 'temoin'))
        from docx import Document
        doc = Document(io.BytesIO(octets))
        # le tableau des relativités : autant de lignes que la constante
        rel = [t for t in doc.tables
               if [c.text.strip() for c in t.rows[0].cells]
               == R.titres('relativites')]
        self.assertEqual(len(rel), 1)
        self.assertEqual(len(rel[0].rows) - 1, R.MAX_RELATIVITES_WORD)
        self.assertIn('%d relativités sur 30' % R.MAX_RELATIVITES_WORD,
                      _texte_docx(octets))
        print('    OK T7c-g : %d lignes rendues, %d annoncées — une seule '
              'constante' % (len(rel[0].rows) - 1, R.MAX_RELATIVITES_WORD))


class T7d_LaPageDeGardeEtLesBordures(unittest.TestCase):
    """T7d — le Word ouvre sur une page de garde, l'HTML enchaînait."""

    def _html(self):
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            return R.export_html(a3, a4, a6, 'CLIENT DE DEMONSTRATION',
                                 '31/12/2025', 'A6_20260809_101112',
                                 narration_calculee=('Texte.', 'temoin'))

    def test_les_quatre_champs_d_identite_sont_NOMMES(self):
        """⚠️ « 🔑 A6_20260808_063434 » NE DIT PAS À UN LECTEUR CE QU'IL
        REGARDE. Le Word les présente en tableau, avec des en-têtes, depuis
        toujours ; l'HTML les portait sans libellé."""
        html = self._html()
        for libelle in R.titres('garde'):
            self.assertIn(f'<span class="lbl">{libelle}</span>', html,
                          libelle)
        for valeur in ('CLIENT DE DEMONSTRATION', '31/12/2025',
                       'A6_20260809_101112'):
            self.assertIn(valeur, html, valeur)
        print(f'    OK T7d : les {len(R.titres("garde"))} champs de garde '
              f'sont nommés, pas devinés')

    def test_les_libelles_viennent_de_LA_MEME_source_que_le_Word(self):
        """⚠️ SANS CELA, LES DEUX PAGES DE GARDE DÉRIVERAIENT — c'est le
        motif que T5 a fermé sur les en-têtes de tableaux."""
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            word = _texte_docx(R.export_word(
                a3, a4, a6, 'CLIENT', '31/12/2025', 'AUDIT-1',
                narration_calculee=('Texte.', 'temoin')))
        html = self._html()
        for libelle in R.titres('garde'):
            self.assertIn(libelle, html, f'HTML · {libelle}')
            self.assertIn(libelle, word, f'WORD · {libelle}')
        print('    OK T7d-b : mêmes libellés de garde dans les deux formats, '
              'lus dans `titres(\'garde\')`')

    def test_la_page_de_garde_occupe_SA_page_a_l_impression(self):
        """⚠️ L'HTML ENCHAÎNAIT SUR LE CHAPITRE 1 au milieu de la première
        feuille : le document imprimé n'avait pas de page de titre, quand le
        Word en a une depuis toujours."""
        html = self._html()
        bloc = html[html.index('@media print'):]
        self.assertIn('.header { break-after: page; '
                      'page-break-after: always; }',
                      bloc.replace('\n', ' ').replace('  ', ' '))
        print('    OK T7d-c : la page de garde se termine par un saut de page')

    def test_les_tableaux_sont_bordes_comme_ceux_du_Word(self):
        """⚠️ LE .DOCX GRILLE TOUTES SES CELLULES (`Table Grid`) quand l'HTML
        ne traçait que des filets horizontaux. C'est l'HTML qui s'aligne : le
        document SIGNÉ garde son apparence."""
        html = self._html()
        style = html.split('<style>')[1].split('</style>')[0]
        self.assertIn('td{padding:6px 10px; border:1px solid #e2e8f0;}', style)
        self.assertIn('border:1px solid var(--navy-mid);', style)
        # l'ancienne regle, qui ne bordait qu'en bas, a disparu
        self.assertNotIn('border-bottom:1px solid #e8edf2', style)
        print('    OK T7d-d : cellules bordées des quatre côtés, comme au Word')


class T7e_LaNavigationDansLeWord(unittest.TestCase):
    """T7e — 23 paragraphes de titre, zéro niveau de plan : Word ne pouvait
    construire ni volet de navigation ni table des matières."""

    def _word(self):
        a3, a4, a6 = _payload_figures()
        with rendu_simule():
            return R.export_word(a3, a4, a6, 'DEMO', '31/12/2025', 'T7e',
                                 narration_calculee=(
                                     '§1 — CONTEXTE\nTexte.\n'
                                     '§2 — SUITE\nEncore.', 'temoin'))

    def _niveaux(self, octets):
        xml = zipfile.ZipFile(io.BytesIO(octets)).read(
            'word/document.xml').decode('utf-8')
        return re.findall(r'<w:outlineLvl w:val="(\d+)"/>', xml)

    def test_les_huit_chapitres_portent_un_niveau_de_plan(self):
        """⚠️ SANS NIVEAU DE PLAN, un rapport de neuf pages se parcourt à la
        molette : Word n'a rien pour bâtir sa navigation."""
        niveaux = self._niveaux(self._word())
        self.assertEqual(niveaux.count('0'), len(R.CHAPITRES),
                         'un chapitre sans niveau de plan')
        self.assertGreaterEqual(niveaux.count('1'), 2,
                                'les sections du commentaire n\'ont pas de '
                                'niveau')
        print(f'    OK T7e : {niveaux.count("0")} chapitres au niveau 0, '
              f'{niveaux.count("1")} sections au niveau 1 (0 avant ce lot)')

    def test_le_niveau_est_pose_SANS_TOUCHER_A_LA_TYPOGRAPHIE(self):
        """⚠️ LE POINT DU LOT, ET LA RAISON DE NE PAS UTILISER « Heading 1 ».

        Basculer les paragraphes sur un style Heading aurait apporté la
        typographie de Word avec — Calibri Light, bleu #2F5496 — et changé
        l'apparence de TOUS les titres d'un document signé. Le niveau de plan
        seul suffit à Word ; la typographie ne bouge pas d'un point.
        """
        from docx import Document
        doc = Document(io.BytesIO(self._word()))
        titres = [p for p in doc.paragraphs
                  if p.text.strip().startswith('Chapitre ')]
        self.assertEqual(len(titres), len(R.CHAPITRES))
        for p in titres:
            self.assertEqual(p.style.name, 'Normal',
                             'le style a changé — la typographie avec')
            r0 = [r for r in p.runs if r.text.strip()][0]
            self.assertEqual(r0.font.size.pt, 16.0)
            self.assertTrue(r0.bold)
            self.assertEqual(str(r0.font.color.rgb), '0F2E52')
        print('    OK T7e-b : style « Normal », 16 pt, navy, gras — '
              'inchangés sur les 8 chapitres')

    def test_le_niveau_de_plan_reste_dans_les_bornes_de_Word(self):
        """Word n'accepte que 0..8 ; au-delà le document est refusé."""
        for niveau in self._niveaux(self._word()):
            self.assertTrue(0 <= int(niveau) <= 8, niveau)
        print('    OK T7e-c : tous les niveaux dans 0..8')


class T8_LaRelectureActuarielle(unittest.TestCase):
    """T8 — le rapport se déclare destiné à l'actuaire désigné et à l'ACPR, et
    n'avait nulle part où enregistrer qu'un actuaire l'avait relu."""

    HORODATAGE = '2026-08-09T10:11:12.131415'

    def _rendu(self, nom='', ia='', environnement='production'):
        a3, a4, a6 = _payload_figures()
        a6['audit_trail'] = {'timestamp': self.HORODATAGE,
                             'environnement': environnement}
        with rendu_simule():
            html = R.export_html(a3, a4, a6, 'CLIENT', '31/12/2025', 'AUD',
                                 narration_calculee=('Texte.', 'temoin'),
                                 actuaire_nom=nom, actuaire_numero_ia=ia)
            word = _texte_docx(R.export_word(
                a3, a4, a6, 'CLIENT', '31/12/2025', 'AUD',
                narration_calculee=('Texte.', 'temoin'),
                actuaire_nom=nom, actuaire_numero_ia=ia))
        return html, word

    def test_relu_et_valide_porte_QUI_et_QUAND(self):
        """⚠️ FACTUEL : qui, et quand. Rien de plus."""
        html, word = self._rendu('Marie Dupont', 'IA-2024-1234')
        attendu = ('Relu et validé par Marie Dupont — N° IA IA-2024-1234, '
                   'le 09/08/2026 à 10 h 11')
        for format_, texte in (('HTML', html), ('WORD', word)):
            self.assertIn(attendu, texte, format_)
        print('    OK T8 : « %s » — dans les deux formats' % attendu[:52])

    def test_le_libelle_est_RELU_ET_VALIDE_jamais_SIGNE(self):
        """⚠️ UN CHAMP DE TEXTE REMPLI PAR CELUI QUI CLIQUE N'EST PAS UNE
        SIGNATURE au sens juridique. Le dire autrement serait une fausse
        déclaration."""
        html, word = self._rendu('Marie Dupont')
        for texte in (html, word):
            self.assertIn('Relu et validé par', texte)
            self.assertNotIn('Signé par', texte)
            self.assertNotIn('Signature', texte)
        print('    OK T8-b : « Relu et validé par », jamais « Signé par »')

    def test_l_absence_de_relecture_SE_DIT(self):
        """⚠️ « Ce qui n'est que dans un champ technique n'existe pas pour
        l'actuaire qui relit le dossier plus tard. » Le danger n'a jamais été
        l'erreur, c'était le silence."""
        html, word = self._rendu('')
        for format_, texte in (('HTML', html), ('WORD', word)):
            self.assertIn('Relecture actuarielle non enregistrée', texte,
                          format_)
            self.assertIn('n\'a pas été relu par un actuaire identifié',
                          texte, format_)
            self.assertNotIn('Relu et validé par', texte, format_)
        print('    OK T8-c : l\'absence est ÉCRITE dans les deux formats')

    def test_hors_production_ne_reclame_pas_de_relecture(self):
        """Une exécution de développement n'a pas à porter une alerte de
        gouvernance : elle dit ce qu'elle est."""
        html, word = self._rendu('', environnement='developpement')
        for texte in (html, word):
            self.assertIn('Environnement de développement', texte)
            self.assertNotIn('non enregistrée', texte)
        print('    OK T8-d : hors production, la trace dit ce qu\'elle est')

    def test_les_trois_etats_et_leurs_codes(self):
        t = R.trace_relecture('Marie Dupont', 'IA-1', self.HORODATAGE)
        self.assertEqual(t.etat, R.RELECTURE_VALIDEE)
        self.assertFalse(t.alerte)
        t = R.trace_relecture('', None, self.HORODATAGE)
        self.assertEqual(t.etat, R.RELECTURE_NON_ENREGISTREE)
        self.assertTrue(t.alerte)
        t = R.trace_relecture('   ', None, None, 'developpement')
        self.assertEqual(t.etat, R.RELECTURE_HORS_PRODUCTION)
        self.assertFalse(t.alerte)
        # un nom fait d'espaces n'est pas un nom
        self.assertEqual(R.trace_relecture('  ').etat,
                         R.RELECTURE_NON_ENREGISTREE)
        print('    OK T8-e : trois états, et un nom vide reste vide')

    def test_AUCUNE_DATE_N_EST_GENEREE(self):
        """⚠️ LA DATE VIENT DU DOSSIER, pas du moment du rendu. Une date
        fabriquée dirait quand le document a été produit, pas quand il a été
        relu."""
        sans = R.trace_relecture('Marie Dupont', None, None)
        self.assertEqual(sans.texte, 'Relu et validé par Marie Dupont')
        self.assertNotIn(', le ', sans.texte)
        avec = R.trace_relecture('Marie Dupont', None, self.HORODATAGE)
        self.assertIn('le 09/08/2026 à 10 h 11', avec.texte)
        # et c'est le MÊME formateur que le chapitre 8
        self.assertIn(R.valeur_audit(self.HORODATAGE), avec.texte)
        print('    OK T8-f : la date vient de `audit_trail[timestamp]`, '
              'jamais du rendu')

    def test_la_relecture_NE_TOUCHE_PAS_au_statut_RAG(self):
        """⚠️ LE STATUT DIT « CE MODÈLE EST-IL DÉPLOYABLE », LA RELECTURE DIT
        « CE DOCUMENT EST-IL DIFFUSABLE ». Les huit plafonds d'A6 portent sur
        ce qui change la valeur du résultat ; la relecture d'un document n'en
        est pas."""
        a3, a4, a6 = _payload_figures()
        a6['statut_rag'] = 'VERT'
        a6['audit_trail'] = {'timestamp': self.HORODATAGE,
                             'environnement': 'production'}
        rendus = {}
        for cle, nom in (('avec', 'Marie Dupont'), ('sans', '')):
            with rendu_simule():
                rendus[cle] = R.export_html(
                    a3, a4, a6, 'CLIENT', '31/12/2025', 'AUD',
                    narration_calculee=('Texte.', 'temoin'),
                    actuaire_nom=nom)
        for html in rendus.values():
            self.assertIn('badge badge-vert', html,
                          'le statut publié a changé avec la relecture')
        print('    OK T8-g : statut VERT dans les deux cas — la relecture ne '
              'plafonne rien')

    def test_A6_fait_transiter_la_relecture(self):
        """A6 ne l'utilise pas : il la fait passer, comme `rapport_qualite`."""
        import inspect
        from direction_non_vie.tarification.a6_comparaison.agent import (
            AgentA6Comparaison)
        params = inspect.signature(AgentA6Comparaison.run).parameters
        for nom in ('actuaire_nom', 'actuaire_numero_ia'):
            self.assertIn(nom, params, nom)
            self.assertIsNone(params[nom].default)
        # et l'orchestrateur les accepte aussi
        for fonction in (R.generer_rapport_tarification, R.export_html,
                         R.export_word, R.export_pdf):
            p = inspect.signature(fonction).parameters
            self.assertIn('actuaire_nom', p, fonction.__name__)
            self.assertIn('actuaire_numero_ia', p, fonction.__name__)
        print('    OK T8-h : le champ traverse A6 et les quatre entrées du '
              'générateur')


if __name__ == '__main__':
    unittest.main(verbosity=2)
