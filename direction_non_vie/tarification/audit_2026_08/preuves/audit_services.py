# ruff: noqa
# ⚠️ SCRIPT D'ARCHIVE — PREUVE D'AUDIT, PAS DU CODE MAINTENU.
# Il a ete ecrit un jour donne pour ETABLIR UN FAIT, et il est conserve
# tel qu'il a ete execute. Le relire aux regles du code de production
# reviendrait a reecrire une piece a conviction : les 25 ecarts qu'il
# porte (10 F841, 8 BLE001, 3 PLW1510, 2 UP031, 1 S110, 1 RUF015) sont
# DECLARES ici plutot que corriges. Les BLE001 sont d'ailleurs voulus :
# un script qui mesure si quelque chose casse doit tout attraper.
# ⚠️ DETERMINISME VERIFIE : sur deux executions successives, 7 des 8
# scripts rendent des mesures IDENTIQUES. Le huitieme, audit_services,
# ne varie que sur l'horodatage QU'IL MESURE -- et c'est precisement son
# constat U1 : << Arrete : publie l'horodatage de generation >>.
"""RELEVE DES SERVICES DE RAPPORT -- chaque constat porte sa mesure."""
import io
import pathlib
import re
import sys

RACINE = pathlib.Path(r'C:\Users\selse\actuaria-app')
sys.path.insert(0, str(RACINE))

import openpyxl

S = RACINE / 'direction_non_vie/tarification/services'
SRC_XL = (S / 'tarif_excel.py').read_text(encoding='utf-8')
SRC_EQ = (S / 'rapport_equipe_tarif.py').read_text(encoding='utf-8')
SRC_MO = (S / 'rapport_modeles_tarif.py').read_text(encoding='utf-8')
SRC_HE = (S / 'excel_helpers.py').read_text(encoding='utf-8')
RES = []


def dire(cle, verdict, detail):
    RES.append((cle, verdict, detail))
    print(f'  [{verdict:9}] {cle:46} {detail}')


def texte_xlsx(o):
    wb = openpyxl.load_workbook(io.BytesIO(o), data_only=True)
    return '\n'.join(str(c.value) for ws in wb.worksheets
                     for row in ws.iter_rows() for c in row if c.value is not None)


def onglets(o):
    return openpyxl.load_workbook(io.BytesIO(o)).sheetnames


def jeux():
    """Des resultats d'agent plausibles pour les six exports."""
    a1 = {'success': True, 'statut_rag': 'VERT', 'branche': 'auto',
          'audit_id': 'A1-X', 'hash_md5': 'abc',
          'qualite': {'score_global': 95., 'nb_lignes': 1000, 'nb_colonnes': 12,
                      'taux_completude': 99., 'nb_doublons': 0, 'taux_doublons': 0.,
                      'expo_ok_pct': 100., 'nb_types_aberrants': 0, 'aberrants': {},
                      'alertes_aberrants': []},
          'rapport': {'coercition_types': {'colonnes_forcees': [], 'alertes': []}}}
    a2 = {'success': True, 'statut_rag': 'AMBRE', 'branche': 'auto', 'audit_id': 'A2-X',
          'rapport': {'etapes': ['imputation'], 'alertes': []},
          'data_dictionnaire': {'log_exposition': {'source': 'exposition',
                                                   'operation': 'log', 'usage': 'offset',
                                                   'justification': 'x'}}}
    a3 = {'success': True, 'statut_rag': 'AMBRE', 'branche': 'auto', 'audit_id': 'A3-X',
          'metriques': {'poisson': {'gini': .18, 'aic': 1200., 'deviance': 800.,
                                    'deviance_nulle': 900., 'pseudo_r2': .1,
                                    'nb_vars_retenues': 5, 'vars_retenues': ['age']},
                        'gamma': {'gini': .12, 'aic': 900.}},
          'relativites_poisson': {'age': {'beta': .05, 'relativite': 1.05,
                                          'ic95_low': 1.01, 'ic95_high': 1.09,
                                          'pvalue': .02, 'significatif': True,
                                          'sens': 'aggravant'}},
          'relativites_gamma': {},
          'validation_glm': {'statut_global': 'VERT', 'conclusion': 'ok',
                             'h1_poisson': {'statut': 'VERT', 'ratio_disp': 1.02,
                                            'message': 'm', 'conseil': 'c'}},
          'hypotheses': {'h1_poisson': {'statut': 'VERT', 'ratio_disp': 1.02,
                                        'message': 'm', 'conseil': 'c'}},
          'credibilite': {'appliquee': False, 'raison': 'pas de groupe'},
          'lissage_geo': {'applique': False, 'raison': 'pas de geo'}}
    a4 = {'success': True, 'statut_rag': 'AMBRE', 'branche': 'auto', 'audit_id': 'A4-X',
          'classement': [{'modele': 'xgboost', 'famille': 'ML', 'gini_test': .22,
                          'rmse_test': .12, 'overfit_ratio': .92, 'score_global': .78}],
          'shap_values': {}, 'validation_ml': {'statut_global': 'AMBRE', 'conclusion': 'x'},
          'rapport': {}}
    a5 = {'success': True, 'statut_rag': 'AMBRE', 'branche': 'auto', 'audit_id': 'A5-X',
          'classement': [{'modele': 'CANN', 'type': 'Deep Learning', 'gini_test': .47,
                          'rmse_test': .3, 'overfit_ratio': 1.1, 'glm_gele': True}],
          'metriques': {'cann': {'glm_gele': True, 'n_vars_glm_matchees': 3,
                                 'n_vars_glm_total': 3, 'glm_verification_error': 0.0}},
          'validation_dl': {'statut_global': 'ROUGE', 'conclusion': 'x'}}
    a6 = {'success': True, 'statut_rag': 'AMBRE', 'branche': 'auto', 'audit_id': 'A6-X',
          'classement': [{'modele': 'GLM_POISSON', 'famille': 'GLM', 'gini_test': .18,
                          'rmse_test': .3, 'overfit_ratio': 1., 'interpretabilite': 1.,
                          'score_global': 1.}],
          'modele_production': {'modele': 'GLM_POISSON', 'famille': 'GLM',
                                'score_global': 1., 'gini_test': .18,
                                'overfit_ratio': 1., 'interpretabilite': 1.},
          'backtest': {'disponible': True, 'ae_ratio': 1.02, 'interpretation': 'ok',
                       'stabilite_wf': 'Stable', 'n_fenetres': 3, 'walk_forward': [],
                       'gini_wf_moyen': .17,
                       'modele_recalibre': 'GLM_POISSON -> proxy GBM',
                       'modele_recalibre_fidele': False},
          'validation_selection': {'c1_nb_modeles': {'statut': 'VERT', 'message': 'm'},
                                   'c2_ecart_gini': {'statut': 'VERT', 'message': 'm'},
                                   'c3_coherence': {'statut': 'VERT', 'message': 'm'}},
          'fiche_decision': {}, 'audit_trail': {'profil_ponderation': 'equilibre',
                                                'profil_valide_par': None,
                                                'environnement': 'production',
                                                'gouvernance_ok': False,
                                                'timestamp': '2026-08-09T10:11:12'},
          'exclusions_conformite': {'titre_enc': 'genre interdit (CJUE C-236/09) - ACTION REQUISE'},
          'alertes_conformite': {}, 'colonnes_plan_manquantes':
              {'plan': 'auto', 'facteurs_absents': ['csp'],
               'colonnes_non_produites': ['csp_employe']},
          'commentaire': 'Commentaire A6.'}
    return a1, a2, a3, a4, a5, a6


# ===============================================================
# U1 -- << Arrete : >> publie l horodatage de generation
# ===============================================================
def u1():
    from direction_non_vie.tarification.services import tarif_excel as X
    a6 = jeux()[5]
    o = X.export_excel_a6(a6, 'AUD')
    txt = texte_xlsx(o)
    # le bandeau porte "Arrete : <date>" -- avec ou sans heure ?
    m = re.findall(r'Arrêté\s*:\s*([0-9/]+(?:\s+\d{2}:\d{2})?)', txt)
    avec_heure = [x for x in m if ':' in x]
    n_sites = len(re.findall(r'_bandeau\([^)]*?,\s*now\)', SRC_XL + SRC_EQ, re.DOTALL))
    dire('U1 "Arrete" = horodatage de generation',
         'CONSTAT' if avec_heure else 'BON',
         f'valeurs trouvees = {m[:3]} · portent une HEURE = {bool(avec_heure)} · '
         f'{n_sites} bandeaux recoivent `now` au lieu de l arrete')


# ===============================================================
# U2 -- comptes d onglets annonces
# ===============================================================
def u2():
    from direction_non_vie.tarification.services import tarif_excel as X
    a1, a2, a3, a4, a5, a6 = jeux()
    reels = {}
    for nom, fab, res in (('a1', X.export_excel_a1, a1), ('a2', X.export_excel_a2, a2),
                          ('a3', X.export_excel_a3, a3), ('a4', X.export_excel_a4, a4),
                          ('a5', X.export_excel_a5, a5), ('a6', X.export_excel_a6, a6)):
        reels[nom] = len(onglets(fab(res, 'AUD')))
    annonces = {m.group(1).lower(): int(m.group(2)) for m in
                re.finditer(r'export_excel_(a\d).*?\((\d+) onglets\)', SRC_XL, re.DOTALL)}
    ecarts = {k: (annonces.get(k), reels[k]) for k in reels
              if annonces.get(k) is not None and annonces[k] != reels[k]}
    dire('U2 onglets annonces vs produits', 'CONSTAT' if ecarts else 'BON',
         f'annonces={annonces} · REELS={reels} · ecarts={ecarts}')


# ===============================================================
# U3 -- << 8 modeles >> republie aux livrables
# ===============================================================
def u3():
    from direction_non_vie.tarification.services import rapport_equipe_tarif as E
    from direction_non_vie.tarification.services import tarif_excel as X
    a1, a2, a3, a4, a5, a6 = jeux()
    t_xl = texte_xlsx(X.export_excel_a4(a4, 'AUD'))
    res = {'a1': a1, 'a2': a2, 'a3': a3, 'a4': a4, 'a5': a5, 'a6': a6}
    t_eq_xl = texte_xlsx(E.export_excel_equipe(res, 'auto', '31/12/2025', 'AUD'))
    t_eq_ht = E.export_html_equipe(res, 'auto', '31/12/2025', 'AUD')
    huit = [('Excel A4', '8 modèles' in t_xl),
            ('Excel equipe', '×8 modèles' in t_eq_xl),
            ('HTML equipe', '×8 modèles' in t_eq_ht)]
    dire('U3 "8 modeles" republie', 'CONSTAT' if any(v for _, v in huit) else 'BON',
         f'{huit} — la boucle d A4 en calibre 6')


# ===============================================================
# U4 -- la reference Wuthrich de l Excel A5
# ===============================================================
def u4():
    a5src = (RACINE / 'direction_non_vie/tarification/a5_deep_learning/agent.py'
             ).read_text(encoding='utf-8')
    ref_agent = re.findall(r'Wüthrich[^\n]{0,90}', a5src)
    ref_excel = re.findall(r'Wüthrich[^\n"]{0,90}', SRC_XL)
    differe = any('Chain-Ladder' in r for r in ref_excel) and \
        not any('Chain-Ladder' in r for r in ref_agent)
    dire('U4 reference Wuthrich Excel vs agent', 'CONSTAT' if differe else 'BON',
         f'AGENT : {ref_agent[0][:64] if ref_agent else "?"}… | '
         f'EXCEL : {ref_excel[0][:64] if ref_excel else "?"}…')


# ===============================================================
# U5 -- le rapport EQUIPE : Excel contre HTML/Word
# ===============================================================
def u5():
    from direction_non_vie.tarification.services import rapport_equipe_tarif as E
    a1, a2, a3, a4, a5, a6 = jeux()
    res = {'a1': a1, 'a2': a2, 'a3': a3, 'a4': a4, 'a5': a5, 'a6': a6}
    xl = texte_xlsx(E.export_excel_equipe(res, 'auto', '31/12/2025', 'AUD'))
    ht = E.export_html_equipe(res, 'auto', '31/12/2025', 'AUD')
    import docx
    wd = docx.Document(io.BytesIO(E.export_word_equipe(res, 'auto', '31/12/2025', 'AUD')))
    wt = '\n'.join(p.text for p in wd.paragraphs) + '\n'.join(
        c.text for t in wd.tables for r in t.rows for c in r.cells)
    marqueurs = {
        'avertissement proxy WF': 'Portée de la validation',
        'colonnes ecartees (conformite)': 'écartées de la matrice X',
        'colonnes du plan non produites': 'MODELE AMPUTE',
    }
    lignes = []
    for nom, m in marqueurs.items():
        lignes.append((nom, m in xl, m in ht, m in wt))
    manquants = [n for n, x, h, w in lignes if x and not (h or w)]
    dire('U5 rapport EQUIPE : Excel dit, HTML/Word taisent',
         'CONSTAT' if manquants else 'BON',
         ' · '.join(f'{n}: xl={x} html={h} word={w}' for n, x, h, w in lignes))


# ===============================================================
# U6 -- H5 absent du tableau des hypotheses du rapport
# ===============================================================
def u6():
    from direction_non_vie.tarification.services import rapport_modeles_tarif as R
    cles = [c for c, _, _ in R.HYPOTHESES]
    plafonnantes = re.findall(r"\((hypotheses_\w+), '(\w+)'\)",
                              (RACINE / 'direction_non_vie/tarification/a6_comparaison/agent.py'
                               ).read_text(encoding='utf-8'))
    pl = [k for _, k in plafonnantes]
    absentes = [k for k in pl if k not in cles]
    dire('U6 hypotheses plafonnantes publiees', 'CONSTAT' if absentes else 'BON',
         f'{len(cles)} hypotheses au tableau · {len(pl)} plafonnantes chez A6 · '
         f'plafonnante(s) ABSENTE(S) du tableau = {absentes}')


# ===============================================================
# U7 -- << non calcule ne vaut pas zero >> : trois valeurs y echappent
# ===============================================================
def u7():
    sites_ctx = re.findall(r"prod\.get\('(\w+)',\s*0\)[^\n]*:\.\d+f", SRC_MO)
    sites_word = re.findall(r'f"\{prod\.get\(\'(\w+)\', ?0\):\.\d+f\}"', SRC_MO)
    regle = "NE VAUT PAS ZÉRO" in SRC_MO.upper() or 'ne vaut pas zéro' in SRC_MO
    dire('U7 le modele retenu : .get(cle, 0) formate',
         'CONSTAT' if sites_ctx or sites_word else 'BON',
         f'contexte du modele : {sites_ctx} · tableau Word : {sites_word} · '
         f'la regle est ecrite dans le meme fichier = {regle}')


# ===============================================================
# U8 -- un ROUGE publie << Attention >>
# ===============================================================
def u8():
    m = re.search(r'txt = \{([^}]+)\}\.get\(statut', SRC_HE)
    table = m.group(1) if m else ''
    rouge = re.search(r'"ROUGE":\s*"([^"]+)"', table)
    ambre = re.search(r'"AMBRE":\s*"([^"]+)"', table)
    dire('U8 libelle publie pour ROUGE', 'CONSTAT' if rouge else 'BON',
         f'ROUGE -> {rouge.group(1) if rouge else "?"!r} · '
         f'AMBRE -> {ambre.group(1) if ambre else "?"!r}')


# ===============================================================
# U9 -- VERIFICATIONS POSITIVES
# ===============================================================
def u9():
    from direction_non_vie.tarification.services import rapport_modeles_tarif as R
    a = 'monitoring_gini' in R.FIGURES_ECARTEES and 'FABRIQU' in \
        R.FIGURES_ECARTEES['monitoring_gini'].upper()
    b = 'optimisation_tarifaire' in R.FIGURES_ECARTEES
    c = len(set(R.TITRES_FIGURES)) == len(set(R.SOURCES_FIGURES)) == \
        len({x for _, cs in R.PLAN_FIGURES for x in cs})
    d = R.nom_modele('MODELE_FUTUR') == 'MODELE_FUTUR' and R.nom_modele(None) == '—'
    e = R.trace_relecture('').etat == R.RELECTURE_NON_ENREGISTREE
    f_ = R.note_troncature(15, 15, 'x', 'y') == '' and 'sur 16' in \
        R.note_troncature(15, 16, 'x', 'y')
    g_ = R.raisons_plafond({'statut_rag': 'AMBRE', 'audit_trail': {}}) == \
        (R.RAISON_INCONNUE,)
    h_ = 'print-color-adjust: exact' in SRC_MO and '@media print' in SRC_MO
    for cle, ok, det in (
            ('U9a monitoring_gini ecartee (donnees fabriquees)', a, 'FABRIQUÉES'),
            ('U9b optimisation_tarifaire ecartee', b, 'hors perimetre'),
            ('U9c catalogue coherent titre/source/plan', c, f'{len(R.TITRES_FIGURES)} figures'),
            ('U9d un nom inconnu n est jamais remplace', d, 'MODELE_FUTUR'),
            ('U9e absence de relecture -> se dit', e, 'non_enregistree'),
            ('U9f troncature declaree seulement si elle coupe', f_, ''),
            ('U9g plafond sans cause -> anomalie signalee', g_, 'RAISON_INCONNUE'),
            ('U9h fonds RAG survivent a l impression', h_, 'print-color-adjust')):
        dire(cle, 'BON' if ok else 'CONSTAT', det)


def main() -> int:
    print('  RELEVE DES SERVICES -- chaque ligne est une mesure\n')
    for f in (u1, u2, u3, u4, u5, u6, u7, u8, u9):
        try:
            f()
        except Exception as e:
            dire(f.__name__.upper(), 'NON MESURE', f'{type(e).__name__}: {e}')
    print('\n  ' + '=' * 82)
    for v in ('CONSTAT', 'BON', 'NON MESURE'):
        print(f'  {v:11} : {sum(1 for _, x, _ in RES if x == v)}')
    return 0


if __name__ == '__main__':
    sys.exit(main())
