# -*- coding: utf-8 -*-
"""DEUX BADGES VERTS SUR UNE ABSENCE DE DONNEE (A6.7).

`A6.run` accepte `result_a1` depuis toujours et le relaie au rapport d'equipe.
**Aucun appelant ne le passait** -- releve par AST le 05/09/2026 : ni
`pipeline_agents`, ni les demos, ni les scripts. L'onglet << Qualite des
donnees (A1) >> lisait donc `{}` et publiait, sur une ABSENCE :

    Score global 0 | Nb lignes 0 | Taux completude 0 %
    Nb types d'anomalies 0   -> badge VERT
    Alertes : Aucune         -> badge VERT

*Un lecteur y voyait un fichier parfait ; il n'y avait pas de fichier.* C'est
la forme exacte du defaut que ce chantier ferme partout : **une valeur absente
remplacee par un litteral**, ici avec un badge qui la confirme.

⚠️⚠️ ZERO EURO, ET C'EST MESURE. `result_a1` n'est lu, dans tout A6, QUE pour
construire le dictionnaire du rapport d'equipe (releve AST : 3 occurrences --
la signature, la docstring, la ligne 930). Ni le modele de production, ni le
classement, ni le statut RAG n'en dependent. Ce cablage change le CONTENU d'un
rapport, jamais un prix.

Ce que cette sentinelle exige :
  QT-1  une qualite absente n'est PAS publiee en zeros ;
  QT-2  elle n'est PAS publiee en VERT ;
  QT-3  les TROIS surfaces se comportent pareil -- Excel, HTML, Word ;
  QT-4  second sens : une qualite PRESENTE s'affiche normalement ;
  QT-5  ⚠️ ET LE CABLAGE SE VERIFIE PAR EXECUTION : `pipeline_agents` fait
        arriver la VRAIE qualite jusqu'au rapport. Un controle qui lirait le
        code au lieu du comportement passerait sur un `importlib` ou un
        renommage.

Tout en `unittest.TestCase` : la gate lance `unittest discover`.
"""
import io
import logging
import os
import sys
import unittest
import warnings

_ICI = os.path.dirname(os.path.abspath(__file__))
_RACINE = os.path.dirname(os.path.dirname(_ICI))
for _c in (_RACINE, _ICI):
    if _c not in sys.path:
        sys.path.insert(0, _c)

import numpy as np

from core.conformite_reglementaire import (
    NON_TRANSMIS,
    phrase_qualite_non_transmise,
)

#: Une qualite REELLE, telle qu'A1 la produit.
_QUALITE = {'score_global': 92.5, 'nb_lignes': 1200, 'taux_completude': 99.1,
            'nb_types_aberrants': 2, 'alertes_aberrants': ['ages hors bornes']}


def _textes(results):
    """Le texte des TROIS livrables d'equipe. ⚠️ Contenu, pas appel : c'est ce
    que l'actuaire lit qui compte."""
    from docx import Document
    from openpyxl import load_workbook

    from direction_non_vie.tarification.services.rapport_equipe_tarif import (
        export_excel_equipe,
        export_html_equipe,
        export_word_equipe,
    )
    classeur = load_workbook(
        io.BytesIO(export_excel_equipe(results, branche='auto',
                                       audit_id='T')), data_only=True)
    excel = '\n'.join(str(c.value) for f in classeur.worksheets
                      for ligne in f.iter_rows() for c in ligne
                      if c.value is not None)
    html = export_html_equipe(results, branche='auto', audit_id='T')
    if isinstance(html, bytes):
        html = html.decode('utf-8', 'replace')
    mot = Document(io.BytesIO(
        export_word_equipe(results, branche='auto', audit_id='T')))
    word = '\n'.join(p.text for p in mot.paragraphs)
    return {'Excel': excel, 'HTML': html, 'Word': word}


class TestUneAbsenceNeSePubliePasEnVert(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        warnings.filterwarnings('ignore')
        logging.disable(logging.CRITICAL)
        cls.absente = _textes({'a6': {'statut_rag': 'VERT'}})
        cls.presente = _textes({'a1': {'statut_rag': 'VERT',
                                       'qualite': dict(_QUALITE)},
                                'a6': {'statut_rag': 'VERT'}})

    @classmethod
    def tearDownClass(cls):
        logging.disable(logging.NOTSET)

    def test_QT1_QT3_une_qualite_absente_est_DITE_sur_les_TROIS_surfaces(self):
        """⚠️⚠️ LES TROIS, PAS UNE. Elles lisent le MEME `r1` : n'en reparer
        que deux laisserait la troisieme affirmer << 0 anomalie >>."""
        for surface, texte in self.absente.items():
            with self.subTest(surface=surface):
                self.assertIn(NON_TRANSMIS, texte,
                              f'{surface} ne dit pas que la qualite manque')

    def test_QT1b_aucun_ZERO_fabrique_ne_subsiste(self):
        """*Un `0` sur une absence se lit comme une mesure.*"""
        for surface, texte in self.absente.items():
            with self.subTest(surface=surface):
                for invente in ('0.0/100', '0 type(s)', ': 0.0%'):
                    self.assertNotIn(
                        invente, texte,
                        f'{surface} publie {invente!r} sur une donnee que '
                        f'personne ne lui a transmise')

    def test_QT2_elle_n_est_PAS_publiee_en_VERT(self):
        """⚠️⚠️ LE COEUR DU CONSTAT. << Alertes : Aucune >> avec un badge vert
        affirme qu'on a CHERCHE et RIEN TROUVE. On n'avait pas cherche."""
        excel = self.absente['Excel']
        self.assertNotIn('Aucune', excel,
                         "le classeur publie encore << Aucune >> sur une "
                         "absence : un lecteur y lit un fichier sain")

    def test_QT2b_la_phrase_dit_POURQUOI_les_compteurs_ne_valent_rien(self):
        phrase = phrase_qualite_non_transmise()
        self.assertIn('NON TRANSMISE', phrase)
        self.assertIn('result_a1', phrase, "la phrase ne dit pas QUOI "
                                           'rebrancher')
        for surface, texte in self.absente.items():
            with self.subTest(surface=surface):
                self.assertIn('NON TRANSMISE', texte)

    def test_QT4_SECOND_SENS_une_qualite_presente_s_affiche_normalement(self):
        """⚠️ Sans ce sens, on aurait pu remplacer TOUS les chiffres par le
        mot et le test QT-1 serait passe."""
        for surface, texte in self.presente.items():
            with self.subTest(surface=surface):
                self.assertNotIn(NON_TRANSMIS, texte,
                                 f'{surface} dit << non transmis >> alors '
                                 f'que la qualite est fournie')
                self.assertIn('92.5', texte, 'le score reel manque')
                self.assertIn('ages hors bornes', texte,
                              "l'alerte reelle n'est pas publiee")


class TestLeRelaisSeVerifieParExecution(unittest.TestCase):
    """⚠️⚠️ PAR EXECUTION, PAS PAR LECTURE DU CODE. Un controle qui verifierait
    par AST qu'un appelant ecrit `result_a1=...` passerait sur un renommage,
    un `**kwargs`, ou un appel construit a l'execution. Ce test appelle A6 et
    LIT LE LIVRABLE.

    ⚠️⚠️ ET IL EXERCE LE MAILLON COMMUN, PAS UN APPELANT. Mesure du
    05/09/2026 : `pipeline_agents` force `generer_rapport_equipe=False`, et
    `scripts/rapport_tarif_local` aussi -- **le seul producteur reel du
    rapport d'equipe est `actuaria_app.py`**, qui ne passe pas le drapeau (donc
    True par defaut) et qu'on ne modifie pas. Tester le relais d'A6, c'est
    donc tester ce que TOUS traversent, celui-la compris.
    """

    def test_QT5_A6_relaie_la_VRAIE_qualite_jusqu_au_rapport(self):
        from docx import Document

        from core.qualite_donnees import preambule_qualite
        from direction_non_vie.tarification import test_pipeline_agents as T
        from direction_non_vie.tarification.a1_ingestion.agent import (
            AgentA1Ingestion,
        )
        from direction_non_vie.tarification.a2_preprocessing.agent import (
            AgentA2Preprocessing,
        )
        from direction_non_vie.tarification.a3_glm.agent import AgentA3GLM
        from direction_non_vie.tarification.a6_comparaison.agent import (
            AgentA6Comparaison,
        )
        warnings.filterwarnings('ignore')
        logging.disable(logging.CRITICAL)
        try:
            np.random.seed(7)
            plan = T._PLAN_AUTO
            base = {'audit_path': '/tmp', 'verbose': False}
            r1 = AgentA1Ingestion(**base).run(
                branche='non_vie', sous_branche='auto',
                dataframe=T._portefeuille_auto(900))
            q = preambule_qualite(r1.get('dataframe'), plan,
                                  qualite_validee_par='Test', horodatage=None)
            r2 = AgentA2Preprocessing(**base).run(
                result_a1={**r1, 'dataframe': q.dataframe_propre}, plan=plan)
            r3 = AgentA3GLM(models_path='/tmp', audit_path='/tmp').run(
                result_a2=r2, plan=plan,
                col_frequence=plan.cible_frequence,
                col_cout=plan.cible_cout, generer_graphiques=False)
            r6 = AgentA6Comparaison(models_path='/tmp', audit_path='/tmp').run(
                result_a2=r2, result_a1=r1, result_a3=r3,
                col_cible='nb_sinistres', plan=plan,
                environnement='production', profil_valide_par='Test',
                generer_graphiques=False, generer_rapport_equipe=True)
        finally:
            logging.disable(logging.NOTSET)

        qualite = (r1 or {}).get('qualite') or {}
        self.assertTrue(
            qualite, "A1 ne publie plus de qualite : ce test ne mesure plus "
                     'rien')
        octets = (r6.get('rapport_equipe') or {}).get('word_bytes')
        self.assertTrue(
            octets,
            "aucun rapport d'equipe Word : le chemin que ce test surveille "
            "n'est plus exerce, et un test qui ne s'execute pas ne prouve "
            'rien')
        texte = '\n'.join(p.text
                          for p in Document(io.BytesIO(octets)).paragraphs)
        self.assertNotIn(
            NON_TRANSMIS, texte,
            "le rapport dit << non transmis >> alors qu'A6 a bien recu "
            '`result_a1` : le relais est rompu entre A6 et le rapport')
        self.assertIn(f"{qualite.get('nb_lignes', 0):,}", texte,
                      'le nombre de lignes REELLEMENT ingere n atteint pas '
                      'le rapport signe')

    def test_QT5b_SANS_result_a1_le_MEME_appel_dit_non_transmis(self):
        """⚠️ Le second sens, sur le MEME chemin : c'est lui qui prouve que
        QT-5 mesure le relais et non l'existence d'un rapport."""
        from docx import Document

        from direction_non_vie.tarification.services.rapport_equipe_tarif import (
            export_word_equipe,
        )
        octets = export_word_equipe({'a6': {'statut_rag': 'VERT'}},
                                    branche='auto', audit_id='T')
        texte = '\n'.join(p.text
                          for p in Document(io.BytesIO(octets)).paragraphs)
        self.assertIn(NON_TRANSMIS, texte)


if __name__ == '__main__':
    unittest.main(verbosity=2)
