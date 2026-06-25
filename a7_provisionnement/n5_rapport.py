# =============================================================================
#  ActuarIA — Agent A7 Ibrahim v5.0
#  n5_rapport.py  —  Export Word (.docx) + PDF professionnel
#
#  Générateur : python-docx (Word) + reportlab (PDF)
#  Graphiques : kaleido (PNG embarqués) avec fallback gracieux si absent
#
#  Palette : Navy #0F2E52 / Gold #C9A84C / Blanc #FFFFFF / Gris #8A9BB0
# =============================================================================

from __future__ import annotations

import io
import logging
import re
from datetime import datetime
from typing import Dict, Optional

import numpy as np

logger = logging.getLogger('actuaria.a7.rapport')

# ── Couleurs hex ──────────────────────────────────────────────────────────────
NAVY    = '0F2E52'
NAVY_L  = '1A3F6B'
GOLD    = 'C9A84C'
BLANC   = 'FFFFFF'
GRIS    = '8A9BB0'
ROUGE   = 'C0392B'
VERT    = '27AE60'
AMBRE   = 'F39C12'
BLEU    = '2980B9'


# =============================================================================
#  HELPERS
# =============================================================================

def _f(v, decimals=0) -> str:
    if v is None: return '—'
    try:
        fv = float(v)
        if np.isnan(fv) or np.isinf(fv): return '—'
        if decimals == 0:
            return f"{fv:,.0f} €".replace(',', ' ')
        return f"{fv:,.{decimals}f}".replace(',', ' ')
    except Exception: return '—'

def _pct(v, decimals=1) -> str:
    if v is None: return '—'
    try:
        fv = float(v)
        if np.isnan(fv) or np.isinf(fv): return '—'
        return f"{fv:.{decimals}f} %"
    except Exception: return '—'

def _statut_txt(b) -> str:
    return 'VALIDÉE' if b else 'REJETÉE'

def _nettoyer(texte: str) -> str:
    """Supprimer les caractères spéciaux illisibles."""
    if not texte: return ''
    texte = re.sub(r'[■□▪▸►]+', '•', texte)
    texte = re.sub(r'─+', '', texte)
    texte = re.sub(r'\n{3,}', '\n\n', texte)
    return texte.strip()


# =============================================================================
#  EXPORT PNG DES GRAPHIQUES via kaleido
# =============================================================================

def _exporter_pngs(graphiques: Dict, width=900, height=380) -> Dict[str, bytes]:
    pngs = {}
    if not graphiques:
        return pngs
    try:
        import plotly.io as pio
        for nom, fig in graphiques.items():
            try:
                png = pio.to_image(fig, format='png', width=width, height=height, scale=2)
                pngs[nom] = png
            except Exception as e:
                logger.debug(f"PNG {nom} ignoré : {e}")
    except ImportError:
        logger.info("kaleido absent — graphiques omis dans le rapport")
    return pngs


# =============================================================================
#  EXPORT WORD  (python-docx)
# =============================================================================

def export_word(
    n1: Dict, n2: Dict, n3: Dict, n4: Dict,
    commentaire : str  = '',
    ref_client  : str  = '',
    arrete      : str  = '',
    audit_id    : str  = '',
    lob_label   : str  = '',
    graphiques  : Dict = None,
) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        logger.error(f"python-docx non disponible : {e}")
        return b''

    try:
        def _rgb(hex6: str) -> RGBColor:
            h = hex6.lstrip('#')
            return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

        NAVY_RGB  = _rgb(NAVY);  GOLD_RGB  = _rgb(GOLD)
        BLANC_RGB = _rgb(BLANC); GRIS_RGB  = _rgb(GRIS)
        ROUGE_RGB = _rgb(ROUGE); VERT_RGB  = _rgb(VERT)
        AMBRE_RGB = _rgb(AMBRE)

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

        # ── Données ───────────────────────────────────────────────────────────
        cl   = n3.get('chain_ladder', {});  mack = n3.get('mack', {})
        bf   = n3.get('bf', {});            cc   = n3.get('cape_cod', {})
        boot = n3.get('bootstrap', {});     scr  = n4.get('scr', {})
        h1   = n2.get('h1_independance', {}); h2 = n2.get('h2_stabilite', {})
        h3   = n2.get('h3_apriori_bf', {});   h4 = n2.get('h4_homosc_bootstrap', {})
        poids = n4.get('poids', {})

        be_val   = float(n4.get('best_estimate', 0) or 0)
        sigma    = float(mack.get('sigma_total', n4.get('sigma_mack', 0)) or 0)
        cv_val   = float(n4.get('cv_inter_methodes', 0) or 0)
        p75      = float(n4.get('reserve_p75', 0) or 0)
        p90      = float(n4.get('reserve_p90', mack.get('reserve_p90', 0)) or 0)
        p99      = float(n4.get('reserve_p99_5', 0) or 0)
        scr_prov = float(scr.get('scr_prov', be_val*0.30) if scr else be_val*0.30)
        scr_ratio = scr_prov / be_val * 100 if be_val else 0

        date_str  = datetime.now().strftime('%d/%m/%Y')
        arrete_s  = arrete or date_str
        client_s  = ref_client or 'ActuarIA'
        lob_s     = lob_label or n2.get('lob_label', '—')
        methode_s = n4.get('methode_facteurs', n2.get('methode_recommandee', '—'))
        statut_s  = n4.get('statut', 'AMBRE')
        statut_rgb = VERT_RGB if statut_s=='VERT' else AMBRE_RGB if statut_s=='AMBRE' else ROUGE_RGB

        pngs = _exporter_pngs(graphiques or {})

        # ── Helpers ───────────────────────────────────────────────────────────
        def _cell_bg(cell, hex_color):
            tc = cell._tc; tcp = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), hex_color); shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:val'), 'clear'); tcp.append(shd)

        def _run(para, text, bold=False, italic=False, size=10, color=None):
            r = para.add_run(str(text))
            r.bold = bold; r.italic = italic; r.font.size = Pt(size)
            if color: r.font.color.rgb = color
            return r

        def _h(text, level=1, color=None, sb=10, sa=4):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(sb)
            p.paragraph_format.space_after  = Pt(sa)
            sz = {1:16, 2:13, 3:11}.get(level, 10)
            cl = color or (NAVY_RGB if level==1 else GOLD_RGB)
            _run(p, text, bold=True, size=sz, color=cl)
            return p

        def _p(text, size=9, color=None, bold=False, italic=False, sa=3, align=None):
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(sa)
            p.paragraph_format.space_before = Pt(0)
            if align: p.alignment = align
            _run(p, _nettoyer(str(text)), bold=bold, italic=italic,
                 size=size, color=color or NAVY_RGB)
            return p

        def _sep(color=NAVY):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            b = OxmlElement('w:bottom')
            b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'6')
            b.set(qn('w:space'),'1'); b.set(qn('w:color'), color)
            pBdr.append(b); pPr.append(pBdr)

        def _tbl(headers, rows, widths=None, hdr=NAVY):
            from docx.enum.table import WD_ALIGN_VERTICAL
            t = doc.add_table(rows=1+len(rows), cols=len(headers))
            t.style = 'Table Grid'
            # Header
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]; _cell_bg(c, hdr)
                p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(h)); r.bold=True; r.font.size=Pt(9)
                r.font.color.rgb=BLANC_RGB; c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            # Rows
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    c = t.rows[ri+1].cells[ci]
                    if ri%2==1: _cell_bg(c, 'EEF2F7')
                    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run(str(val) if val is not None else '—')
                    r.font.size=Pt(9); c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
            if widths:
                for i, w in enumerate(widths):
                    for row in t.rows: row.cells[i].width = Cm(w)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            return t

        def _img(png_bytes, width_cm=14.5):
            if not png_bytes: return
            try:
                doc.add_picture(io.BytesIO(png_bytes), width=Cm(width_cm))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                _p(f'[Graphique non disponible : {e}]', size=8, italic=True)

        # =========================================================================
        #  PAGE DE GARDE
        # =========================================================================
        doc.add_paragraph(); doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, 'RAPPORT ACTUARIEL\n', bold=True, size=26, color=NAVY_RGB)
        _run(p, 'Provisionnement Non-Vie', bold=True, size=18, color=GOLD_RGB)
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f'{client_s}\n', bold=True, size=14, color=NAVY_RGB)
        _run(p, f'Arrêté : {arrete_s}  ·  Branche : {lob_s}', size=12, color=GRIS_RGB)
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, 'Statut global : ', size=11, color=NAVY_RGB)
        emoji = '✅' if statut_s=='VERT' else '⚠️' if statut_s=='AMBRE' else '❌'
        _run(p, f'{emoji} {statut_s}', bold=True, size=11, color=statut_rgb)
        doc.add_paragraph()
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, 'CONFIDENTIEL — USAGE STRICTEMENT ACTUARIEL', bold=True, size=10, color=ROUGE_RGB)
        doc.add_paragraph()
        _tbl(['Référence client','Branche (LoB)','Méthode retenue','Date','Audit ID'],
             [[client_s, lob_s, methode_s.replace('_',' ').title(), date_str, audit_id or '—']],
             widths=[3.0,3.0,3.5,2.5,4.0])
        doc.add_page_break()

        # =========================================================================
        #  1. SYNTHÈSE
        # =========================================================================
        _h('1. Synthèse exécutive'); _sep(GOLD)
        _tbl(['Indicateur','Valeur','Indicateur','Valeur'],
             [['Best Estimate S2 (Art.77)', _f(be_val), 'σ Mack total', _f(sigma)],
              ['Provision P75', _f(p75), 'CV inter-méthodes', _pct(cv_val)],
              ['Provision P90', _f(p90), 'SCR Provisions', _f(scr_prov)],
              ['Provision P99.5', _f(p99), 'Ratio SCR/BE', _pct(scr_ratio)]],
             widths=[4.5,3.5,4.5,3.5])
        if pngs.get('g5_convergence'):
            _h('Convergence des méthodes', level=2)
            _img(pngs['g5_convergence'])
        if commentaire:
            _h('Points saillants', level=2)
            extrait = _nettoyer(commentaire)
            for ligne in [l.strip() for l in extrait.split('\n') if l.strip()][:10]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Cm(0.3)
                _run(p, ligne, size=9,
                     color=GOLD_RGB if ligne[0].isdigit() or ligne.startswith('§') else NAVY_RGB,
                     bold=ligne[0].isdigit() or ligne.startswith('§'))
        doc.add_page_break()

        # =========================================================================
        #  2. RÉSULTATS PAR MÉTHODE
        # =========================================================================
        _h('2. Résultats par méthode actuarielle'); _sep(GOLD)
        _tbl(['Méthode','Réserve IBNR (€)','Poids BE','Score /100','Statut'],
             [['Chain Ladder', _f(cl.get('reserve_totale')),
               _pct(poids.get('chain_ladder',0)*100),
               str(n2.get('scores_confiance',{}).get('chain_ladder','—')), '✓'],
              ['Mack 1993', _f(mack.get('reserve_best_estimate')),
               _pct(poids.get('mack',0)*100),
               str(n2.get('scores_confiance',{}).get('mack','—')), '✓'],
              ['Bornhuetter-Ferguson', _f(bf.get('reserve_totale')),
               _pct(poids.get('bf',0)*100),
               str(n2.get('scores_confiance',{}).get('bf','—')), '✓'],
              ['Cape Cod', _f(cc.get('reserve_totale')),
               _pct(poids.get('cape_cod',0)*100),
               str(n2.get('scores_confiance',{}).get('cape_cod','—')), '✓'],
              ['⭐ BEST ESTIMATE S2', _f(be_val), '100 %', '—', '→ Bilan S2']],
             widths=[4.5,3.5,2.5,2.5,3.0], hdr=NAVY_L)
        for nom_g, titre_g in [('g1_heatmap','Triangle de développement cumulé'),
                                ('g13_paiements','Paiements cumulés par année de survenance'),
                                ('g4_ibnr','IBNR par année de survenance'),
                                ('g2_cadences','Cadences cumulées — Chain Ladder')]:
            if pngs.get(nom_g):
                _h(titre_g, level=2); _img(pngs[nom_g])
        doc.add_page_break()

        # =========================================================================
        #  3. VALIDATION DES HYPOTHÈSES
        # =========================================================================
        _h('3. Validation des hypothèses actuarielles'); _sep(GOLD)
        rows_hyp = []
        for label, h in [('H1 — Indépendance (Mack 1993)', h1),
                          ('H2 — Stabilité des facteurs', h2),
                          ('H3 — A priori BF/Cape Cod', h3),
                          ('H4 — Homoscédasticité ODP', h4)]:
            if not h: continue
            ok = bool(h.get('ok', True))
            indic = (f"corr={h.get('corr_moy','—')}" if 'corr_moy' in h else
                     f"CV={h.get('cv_moy','—')}"     if 'cv_moy'   in h else
                     f"LR={h.get('lr_apriori','—')}" if 'lr_apriori' in h else
                     f"φ={h.get('phi','—')}")
            rows_hyp.append([label, _statut_txt(ok), f"{h.get('score','—')}/100", indic])
        _tbl(['Hypothèse','Résultat','Score','Indicateur'], rows_hyp,
             widths=[6.0,2.5,2.0,5.5])
        for label, h in [('H1 — Indépendance (Mack 1993)', h1),
                          ('H2 — Stabilité des facteurs', h2),
                          ('H3 — A priori BF/Cape Cod', h3),
                          ('H4 — Homoscédasticité ODP', h4)]:
            if not h: continue
            ok  = bool(h.get('ok', True))
            msg = _nettoyer(h.get('message', ''))
            if not msg: continue
            _h(label, level=2, color=VERT_RGB if ok else AMBRE_RGB)
            _p(msg, size=9, sa=4)
        for nom_g, titre_g in [('g8_h1','H1 — Indépendance Spearman'),
                                ('g9_h2','H2 — Stabilité facteurs'),
                                ('g6_bootstrap','Distribution Bootstrap ODP')]:
            if pngs.get(nom_g):
                _h(titre_g, level=2); _img(pngs[nom_g])
        doc.add_page_break()

        # =========================================================================
        #  4. SCR
        # =========================================================================
        _h('4. SCR Provisions — Art. 105 Solvabilité 2'); _sep(GOLD)
        _tbl(['Composante','Valeur','Référence réglementaire'],
             [['Best Estimate S2', _f(be_val), 'Art. 77 Directive S2'],
              ['Facteur σ EIOPA', '10 %', f'LoB : {lob_s} (Annexe II, Rgt 2015/35)'],
              ['SCR Provisions', _f(scr_prov), 'SCR = 3 × σ(LoB) × BE'],
              ['Ratio SCR/BE', _pct(scr_ratio), 'Cible < 35 % (pratique marché)']],
             widths=[4.5,3.5,8.0])
        scr_msg = scr.get('message','') if scr else ''
        if scr_msg: _p(_nettoyer(scr_msg), size=9)
        doc.add_page_break()

        # =========================================================================
        #  5. COMMENTAIRE
        # =========================================================================
        _h('5. Commentaire actuariel complet'); _sep(GOLD)
        if commentaire:
            comm = _nettoyer(commentaire)
            sections = re.split(r'(?=§\d+\s*—|\d+\.\s+[A-ZÀÂÉÈÊ])', comm)
            for sec in sections:
                sec = sec.strip()
                if not sec: continue
                lignes = sec.split('\n', 1)
                titre_s = lignes[0].strip()
                corps_s = lignes[1].strip() if len(lignes) > 1 else ''
                if titre_s: _h(titre_s, level=2)
                if corps_s:
                    for para_txt in corps_s.split('\n'):
                        para_txt = para_txt.strip()
                        if para_txt:
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(3)
                            p.paragraph_format.left_indent = Cm(0.3)
                            _run(p, para_txt, size=9, color=NAVY_RGB)
        else:
            _p('Commentaire non disponible.', size=9, italic=True)
        doc.add_page_break()

        # =========================================================================
        #  6. JUGEMENT ACTUARIEL
        # =========================================================================
        _h('6. Jugement actuariel documenté'); _sep(GOLD)
        jugement = n4.get('jugement', '')
        if jugement:
            jug = _nettoyer(jugement)
            sections_j = re.split(r'(?=\d+\.\s+[A-ZÀÂÉÈÊ])', jug)
            for sec in sections_j:
                sec = sec.strip()
                if not sec: continue
                lignes = sec.split('\n', 1)
                titre_j = lignes[0].strip()
                corps_j = lignes[1].strip() if len(lignes) > 1 else ''
                if titre_j: _h(titre_j, level=2)
                if corps_j:
                    for para_txt in corps_j.split('\n'):
                        para_txt = para_txt.strip()
                        if para_txt:
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(3)
                            p.paragraph_format.left_indent = Cm(0.3)
                            _run(p, para_txt, size=9, color=NAVY_RGB)
        else:
            _p('Jugement actuariel non disponible.', size=9, italic=True)
        doc.add_page_break()

        # =========================================================================
        #  7. RECOMMANDATIONS
        # =========================================================================
        _h('7. Recommandations'); _sep(GOLD)
        alertes = n4.get('alertes', n2.get('alertes', []))
        if alertes:
            _h('Points de vigilance', level=2)
            for alerte in alertes:
                a = _nettoyer(str(alerte))
                if not a: continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.5)
                _run(p, '⚠️  ', size=10, color=AMBRE_RGB)
                _run(p, a, size=9, color=NAVY_RGB)
        recs = n4.get('recommandations', [])
        if recs:
            _h('Actions recommandées', level=2)
            for i, rec in enumerate(recs, 1):
                r = _nettoyer(str(rec))
                if not r: continue
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.left_indent = Cm(0.5)
                _run(p, f'{i}.  ', bold=True, size=10, color=GOLD_RGB)
                _run(p, r, size=9, color=NAVY_RGB)
        avis = n4.get('avis_actuariel', '')
        if avis:
            _h('Avis actuariel final', level=2)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Cm(0.3)
            _run(p, _nettoyer(avis), size=9,
                 color=ROUGE_RGB if 'DÉFAVORABLE' in avis.upper() else VERT_RGB)

        # ── Footer ────────────────────────────────────────────────────────────
        _sep(NAVY)
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f'ActuarIA · {client_s} · Arrêté {arrete_s} · '
                f'Audit ID : {audit_id or "—"} · Généré le {date_str} · CONFIDENTIEL',
             size=7, color=GRIS_RGB, italic=True)

        buf = io.BytesIO()
        doc.save(buf); buf.seek(0)
        word_bytes = buf.read()
        logger.info(f"Word généré : {len(word_bytes):,} bytes — {len(pngs)} graphiques")
        return word_bytes

    except Exception as e:
        logger.error(f"export_word échoué : {e}", exc_info=True)
        return b''


# =============================================================================
#  EXPORT PDF  (reportlab)
# =============================================================================

def export_pdf(
    n1: Dict = None, n2: Dict = None, n3: Dict = None, n4: Dict = None,
    commentaire : str  = '',
    ref_client  : str  = '',
    arrete      : str  = '',
    audit_id    : str  = '',
    lob_label   : str  = '',
    graphiques  : Dict = None,
    **kwargs,
) -> bytes:
    n1=n1 or {}; n2=n2 or {}; n3=n3 or {}; n4=n4 or {}
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.lib import colors
        from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
            Table, TableStyle, PageBreak, HRFlowable)
        from reportlab.platypus import Image as RLImage
    except ImportError as e:
        logger.error(f"reportlab non disponible : {e}"); return b''

    try:
        cl   = n3.get('chain_ladder',{}); mack = n3.get('mack',{})
        bf   = n3.get('bf',{});           cc   = n3.get('cape_cod',{})
        boot = n3.get('bootstrap',{});    scr  = n4.get('scr',{})
        poids = n4.get('poids',{})
        h1=n2.get('h1_independance',{}); h2=n2.get('h2_stabilite',{})
        h3=n2.get('h3_apriori_bf',{});   h4=n2.get('h4_homosc_bootstrap',{})

        be_val   = float(n4.get('best_estimate',0) or 0)
        sigma    = float(mack.get('sigma_total',n4.get('sigma_mack',0)) or 0)
        cv_val   = float(n4.get('cv_inter_methodes',0) or 0)
        p75      = float(n4.get('reserve_p75',0) or 0)
        p90      = float(n4.get('reserve_p90',0) or 0)
        p99      = float(n4.get('reserve_p99_5',0) or 0)
        scr_prov = float(scr.get('scr_prov',be_val*0.30) if scr else be_val*0.30)
        scr_ratio = scr_prov/be_val*100 if be_val else 0

        date_str = datetime.now().strftime('%d/%m/%Y')
        arrete_s = arrete or date_str
        client_s = ref_client or 'ActuarIA'
        lob_s    = lob_label or n2.get('lob_label','—')
        methode_s = n4.get('methode_facteurs', n2.get('methode_recommandee','—'))

        NAVY_RL  = colors.HexColor(f'#{NAVY}')
        GOLD_RL  = colors.HexColor(f'#{GOLD}')
        GRIS_RL  = colors.HexColor(f'#{GRIS}')
        ROUGE_RL = colors.HexColor(f'#{ROUGE}')
        VERT_RL  = colors.HexColor(f'#{VERT}')
        AMBRE_RL = colors.HexColor(f'#{AMBRE}')
        BLANC_RL = colors.white
        LGRIS_RL = colors.HexColor('#EEF2F7')

        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
            topMargin=2*cm, bottomMargin=2*cm,
            leftMargin=2.5*cm, rightMargin=2.5*cm,
            title=f"Rapport Actuariel {arrete_s}", author="ActuarIA")

        def S(name, **kw): return ParagraphStyle(name, **kw)
        styles = {
            'titre': S('titre', fontSize=22, fontName='Helvetica-Bold',
                       textColor=NAVY_RL, alignment=1, spaceAfter=8),
            'sous':  S('sous',  fontSize=13, fontName='Helvetica',
                       textColor=GOLD_RL, alignment=1, spaceAfter=6),
            'conf':  S('conf',  fontSize=9,  fontName='Helvetica-Bold',
                       textColor=ROUGE_RL, alignment=1, spaceAfter=12),
            'h1':    S('h1',    fontSize=14, fontName='Helvetica-Bold',
                       textColor=NAVY_RL, spaceBefore=12, spaceAfter=4),
            'h2':    S('h2',    fontSize=11, fontName='Helvetica-Bold',
                       textColor=GOLD_RL, spaceBefore=8, spaceAfter=3),
            'h3':    S('h3',    fontSize=10, fontName='Helvetica-Bold',
                       textColor=NAVY_RL, spaceBefore=4, spaceAfter=2),
            'body':  S('body',  fontSize=9,  fontName='Helvetica',
                       textColor=colors.black, spaceAfter=4, leading=13),
            'footer':S('footer',fontSize=7,  fontName='Helvetica',
                       textColor=GRIS_RL, alignment=1),
        }

        def _tbl(headers, rows, col_w=None, hdr_bg=None):
            hdr_bg = hdr_bg or NAVY_RL
            data = [[Paragraph(f'<b>{h}</b>',
                               ParagraphStyle('th',fontSize=9,fontName='Helvetica-Bold',
                                              textColor=BLANC_RL,alignment=1))
                     for h in headers]]
            for row in rows:
                data.append([Paragraph(str(c or '—'),
                                       ParagraphStyle('td',fontSize=9,fontName='Helvetica',
                                                      alignment=1))
                             for c in row])
            t = Table(data, colWidths=col_w, repeatRows=1)
            t.setStyle(TableStyle([
                ('BACKGROUND',(0,0),(-1,0),hdr_bg),
                ('TEXTCOLOR',(0,0),(-1,0),BLANC_RL),
                ('FONTNAME',(0,0),(-1,0),'Helvetica-Bold'),
                ('FONTSIZE',(0,0),(-1,-1),9),
                ('ALIGN',(0,0),(-1,-1),'CENTER'),
                ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
                ('GRID',(0,0),(-1,-1),0.5,colors.HexColor('#D0D8E4')),
                ('ROWBACKGROUNDS',(0,1),(-1,-1),[BLANC_RL,LGRIS_RL]),
                ('TOPPADDING',(0,0),(-1,-1),5),
                ('BOTTOMPADDING',(0,0),(-1,-1),5),
            ]))
            return t

        pngs = _exporter_pngs(graphiques or {})
        story = []

        # Page de garde
        story += [Spacer(1,2*cm),
                  Paragraph('RAPPORT ACTUARIEL<br/>Provisionnement Non-Vie', styles['titre']),
                  Spacer(1,0.5*cm),
                  Paragraph(f'{client_s}<br/>'
                            f'<font size="11">Arrêté : {arrete_s}  ·  Branche : {lob_s}</font>',
                            styles['sous']),
                  Spacer(1,0.3*cm),
                  Paragraph('CONFIDENTIEL — USAGE STRICTEMENT ACTUARIEL', styles['conf']),
                  Spacer(1,0.8*cm),
                  _tbl(['Référence','Méthode','Date','Audit ID'],
                       [[client_s, methode_s.replace('_',' ').title(), date_str, audit_id or '—']],
                       col_w=[4*cm,5*cm,3*cm,4*cm]),
                  PageBreak()]

        # 1. Synthèse
        story += [Paragraph('1. Synthèse exécutive', styles['h1']),
                  HRFlowable(width='100%',thickness=1.5,color=GOLD_RL,spaceAfter=8),
                  _tbl(['Indicateur','Valeur','Indicateur','Valeur'],
                       [['Best Estimate S2',_f(be_val),'σ Mack total',_f(sigma)],
                        ['Provision P75',_f(p75),'CV inter-méthodes',_pct(cv_val)],
                        ['Provision P90',_f(p90),'SCR Provisions',_f(scr_prov)],
                        ['Provision P99.5',_f(p99),'Ratio SCR/BE',_pct(scr_ratio)]],
                       col_w=[4*cm,4*cm,4*cm,4*cm])]
        if pngs.get('g5_convergence'):
            story += [Spacer(1,0.3*cm), Paragraph('Convergence des méthodes',styles['h2']),
                      RLImage(io.BytesIO(pngs['g5_convergence']),width=15*cm,height=6*cm)]
        story.append(PageBreak())

        # 2. Résultats
        story += [Paragraph('2. Résultats par méthode',styles['h1']),
                  HRFlowable(width='100%',thickness=1.5,color=GOLD_RL,spaceAfter=8),
                  _tbl(['Méthode','Réserve (€)','Poids','Score'],
                       [['Chain Ladder',_f(cl.get('reserve_totale')),
                         _pct(poids.get('chain_ladder',0)*100),
                         str(n2.get('scores_confiance',{}).get('chain_ladder','—'))],
                        ['Mack 1993',_f(mack.get('reserve_best_estimate')),
                         _pct(poids.get('mack',0)*100),
                         str(n2.get('scores_confiance',{}).get('mack','—'))],
                        ['Bornhuetter-Ferguson',_f(bf.get('reserve_totale')),
                         _pct(poids.get('bf',0)*100),
                         str(n2.get('scores_confiance',{}).get('bf','—'))],
                        ['Cape Cod',_f(cc.get('reserve_totale')),
                         _pct(poids.get('cape_cod',0)*100),
                         str(n2.get('scores_confiance',{}).get('cape_cod','—'))],
                        ['⭐ BEST ESTIMATE S2',_f(be_val),'100 %','—']],
                       col_w=[5*cm,4*cm,3*cm,4*cm])]
        for nom_g, titre_g in [('g1_heatmap','Triangle'),('g13_paiements','Paiements cumulés'),
                                ('g4_ibnr','IBNR par année')]:
            if pngs.get(nom_g):
                story += [Spacer(1,0.3*cm), Paragraph(titre_g,styles['h2']),
                          RLImage(io.BytesIO(pngs[nom_g]),width=15*cm,height=5.5*cm)]
        story.append(PageBreak())

        # 3. Hypothèses
        story += [Paragraph('3. Validation des hypothèses',styles['h1']),
                  HRFlowable(width='100%',thickness=1.5,color=GOLD_RL,spaceAfter=8),
                  _tbl(['Hypothèse','Résultat','Score','Message'],
                       [['H1 — Indépendance',_statut_txt(h1.get('ok',True)),
                         f"{h1.get('score','—')}/100",_nettoyer(h1.get('message',''))[:80]],
                        ['H2 — Stabilité',_statut_txt(h2.get('ok',True)),
                         f"{h2.get('score','—')}/100",_nettoyer(h2.get('message',''))[:80]],
                        ['H3 — A priori BF',_statut_txt(h3.get('ok',True)),
                         f"{h3.get('score','—')}/100",_nettoyer(h3.get('message',''))[:80]],
                        ['H4 — Homoscédasticité',_statut_txt(h4.get('ok',True)),
                         f"{h4.get('score','—')}/100",_nettoyer(h4.get('message',''))[:80]]],
                       col_w=[4*cm,2.5*cm,2*cm,7.5*cm])]
        story.append(PageBreak())

        # 4. SCR
        story += [Paragraph('4. SCR Provisions — Art. 105 S2',styles['h1']),
                  HRFlowable(width='100%',thickness=1.5,color=GOLD_RL,spaceAfter=8),
                  _tbl(['Composante','Valeur','Référence'],
                       [['Best Estimate S2',_f(be_val),'Art. 77'],
                        ['Facteur σ EIOPA','10 %',f'LoB : {lob_s}'],
                        ['SCR Provisions',_f(scr_prov),'SCR = 3 × σ × BE'],
                        ['Ratio SCR/BE',_pct(scr_ratio),'Cible < 35 %']],
                       col_w=[5*cm,4*cm,7*cm]),
                  PageBreak()]

        # 5. Commentaire + jugement
        story.append(Paragraph('5. Commentaire actuariel',styles['h1']))
        story.append(HRFlowable(width='100%',thickness=1.5,color=GOLD_RL,spaceAfter=8))
        if commentaire:
            for sec in re.split(r'(?=§\d+\s*—|\d+\.\s+[A-ZÀÂÉÈÊ])', _nettoyer(commentaire)):
                sec = sec.strip()
                if not sec: continue
                lignes = sec.split('\n',1)
                if lignes[0]: story.append(Paragraph(lignes[0],styles['h2']))
                if len(lignes)>1:
                    for pt in lignes[1].split('\n'):
                        pt = pt.strip()
                        if pt: story.append(Paragraph(pt,styles['body']))
        jugement = n4.get('jugement','')
        if jugement:
            story.append(Spacer(1,0.3*cm))
            story.append(Paragraph('Jugement actuariel documenté',styles['h2']))
            for sec in re.split(r'(?=\d+\.\s+[A-ZÀÂÉÈÊ])', _nettoyer(jugement)):
                sec = sec.strip()
                if not sec: continue
                lignes = sec.split('\n',1)
                if lignes[0]: story.append(Paragraph(lignes[0],styles['h3']))
                if len(lignes)>1:
                    for pt in lignes[1].split('\n'):
                        pt = pt.strip()
                        if pt: story.append(Paragraph(pt,styles['body']))

        # Footer
        story += [Spacer(1,0.5*cm),
                  HRFlowable(width='100%',thickness=0.5,color=GRIS_RL),
                  Paragraph(f'ActuarIA · {client_s} · Arrêté {arrete_s} · '
                            f'Audit ID : {audit_id or "—"} · {date_str} · CONFIDENTIEL',
                            styles['footer'])]

        doc.build(story)
        buf.seek(0)
        pdf_bytes = buf.read()
        logger.info(f"PDF généré : {len(pdf_bytes):,} bytes")
        return pdf_bytes

    except Exception as e:
        logger.error(f"export_pdf échoué : {e}", exc_info=True)
        return b''
