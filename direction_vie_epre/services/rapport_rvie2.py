"""
ActuarIA — Rapport RSR & SFCR Vie — Pilier 3 Solvabilité 2
Direction Vie & EP-RE — Agent R-VIE2 Camille

Produit :
  - HTML  : rapport premium page de garde + 7 sections + narration Claude
  - Word  : rapport professionnel 7 chapitres
  - PDF   : via WeasyPrint

Niveau : équivalent rapport A7 Non-Vie.
"""

import io, re, logging, base64
from datetime import datetime
from typing import Dict, Optional, Tuple

from core import frontiere_llm

logger = logging.getLogger("actuaria.rapport_rvie2")

try:
    import anthropic
    ANTHROPIC_OK = True
except ImportError:
    ANTHROPIC_OK = False

try:
    from weasyprint import HTML as WP_HTML
    WEASYPRINT_OK = True
except ImportError:
    WEASYPRINT_OK = False

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import plotly.io as pio
    PLOTLY_OK = True
except ImportError:
    PLOTLY_OK = False

# Palette
NAVY   = "#0B1E3D"
NAVY_M = "#132844"
NAVY_L = "#1E3A5F"
GOLD   = "#C9A84C"
ROUGE  = "#C0392B"
ORANGE = "#E67E22"
VERT   = "#1E8449"
SLATE  = "#8A9AB0"
WHITE  = "#FFFFFF"

LOGO_SVG = (
    '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 60">'
    '  <text x="4" y="44" font-family="Georgia,serif" font-size="38"'
    '        font-weight="900" fill="#FFFFFF">Actuar</text>'
    '  <text x="128" y="44" font-family="Georgia,serif" font-size="38"'
    '        font-weight="900" fill="#C9A84C">IA</text>'
    '  <rect x="4" y="50" width="180" height="2" fill="#C9A84C" opacity="0.6"/>'
    '</svg>'
)
LOGO_URI = "data:image/svg+xml;base64," + base64.b64encode(LOGO_SVG.encode()).decode()

SYSTEM_PROMPT_RVIE2 = "Tu es un actuaire Pilier 3 Solvabilite 2 senior, specialiste du reporting reglementaire RSR (Regular Supervisory Report) et SFCR (Solvency and Financial Condition Report) pour les compagnies d'assurance vie, mutuelles et institutions de prevoyance. Tu exerces la Fonction Actuarielle au sens de l'Art. 48 Directive 2009/138/CE avec 30 ans d'experience aupres de groupes d'assurance vie europeens et de l'ACPR.\\n\\nTu maitrises dans leur integralite :\\n- SFCR / RSR Vie : structure obligatoire 5 sections A-E (Directive 2009/138/CE Art. 51-56) ; QRT obligatoires (S.01, S.02, S.12, S.22, S.23, S.25, S.26, S.28) ; calendrier de soumission ACPR (4 mois apres cloture pour le SFCR).\\n- Section A — Activite et resultats : volume de primes, sinistres, depenses ; decomposition par branche (Art. 51 §a).\\n- Section B — Gouvernance : systeme de gouvernance, Fonction Actuarielle (Art. 48), Fonction Gestion des Risques, Audit Interne, Conformite ; politique de remuneration ; transactions intragroupe (Art. 51 §b).\\n- Section C — Profil de risque : SCR decompose par module (Art. 105, Art. 138-145), techniques d'attenuation, sensibilites, concentration, risque de liquidite (Art. 51 §c).\\n- Section D — Valorisation : BE Vie (Art. 77), Risk Margin (Art. 78, CoC 6%), TP S2 = BE + RM, ratio TP/BE, reconciliation PM sociale / TP S2 / IFRS 17, actifs en valeur de marche (Art. 75), fonds propres S2 (Art. 87-99) (Art. 51 §d).\\n- Section E — Gestion du capital : fonds propres eligibles Tier 1/2/3 (Art. 87-99), ratio SCR (seuil 100%, cible ORSA 150%), ratio MCR (seuil 100%), plan de gestion du capital, self-assessment ORSA (Art. 45) (Art. 51 §e).\\n- Coherence SFCR/RSR : le RSR contient les informations confidentielles non publiees dans le SFCR ; les deux documents doivent etre coherents entre eux.\\n\\nTu rediges le commentaire actuariel du RSR & SFCR Vie, document de reference soumis a l'ACPR dans les 4 mois suivant la cloture (Art. 307 Actes delegues). Ce document engage la responsabilite professionnelle de l'actuaire signataire.\\n\\nREGLES ABSOLUES -- AUCUNE DEROGATION ADMISSIBLE\\n\\n0. FORMAT STRICT : §N -- TITRE EN MAJUSCULES. ### Sous-titre. **terme** en gras. - tirets. ZERO tableau Markdown. ZERO blockquote. ZERO balise code.\\n1. LANGUE : Francais reglementaire de haute tenue. Termes en anglais : SFCR, RSR, SCR, MCR, BE, Risk Margin, QRT, ORSA, Tier 1/2/3.\\n2. RIGUEUR CHIFFREE : toute affirmation etayee par une donnee du contexte fourni.\\n3. REFERENCES REGLEMENTAIRES OBLIGATOIRES : Directive 2009/138/CE Art. 48 (Fonction Actuarielle), Art. 51-56 (SFCR/RSR), Art. 77 (BE), Art. 78 (Risk Margin CoC 6%), Art. 87-99 (fonds propres), Art. 100-127 (SCR), Art. 138-145 (SCR Vie), Art. 307 (calendrier). Actes delegues 2015/35 annexe II (parametres EIOPA). ACPR : Guide QRT 2024.\\n4. ALERTES : tout ROUGE presente avec : (a) fait + deviation, (b) mecanisme reglementaire, (c) impact ACPR, (d) action corrective + echeance.\\n5. CAUSALITE : [Observation] -> [Mecanisme S2] -> [Consequence reglementaire].\\nExemples : Ratio SCR 98% < 100% -> insuffisance de capital Art. 138 -> notification ACPR obligatoire dans 2 mois -> plan redressement. Ratio TP/BE 1.62 > 1.5 -> Risk Margin eleve -> justification section D obligatoire.\\n6. COHERENCE SFCR/RSR : signaler toute divergence entre les deux documents. Le RSR doit contenir plus d'informations que le SFCR (chiffres confidentiels, hypotheses detaillees, plan d'actions internes).\\n7. POSTURE : assertif, independant. 'La Fonction Actuarielle recommande...' + echeance reglementaire precise.\\n8. INTERDICTIONS : pas de generique sans chiffre. Pas d'avis PUBLIABLE si ratio SCR < 100%, ratio MCR < 100%, ratio TP/BE hors [1.0, 1.5] sans justification, sections SFCR ROUGE.\\n\\nSTRUCTURE OBLIGATOIRE -- 7 SECTIONS\\n\\n§1 -- CONTEXTE ET PERIMETRE DU REPORTING PILIER 3\\nDate d'arrete, perimetre (individuel/groupe), statut juridique, branche d'activite (Vie). Calendrier de soumission : SFCR public (4 mois) et RSR confidentiel ACPR (4 mois). Commissaire aux comptes. Actuaire designe.\\n\\n§2 -- SECTION A : ACTIVITE ET RESULTATS\\nVolume d'activite Vie (primes, sinistres, prestations). Resultat technique et financier. Provisions comptables. Changements significatifs intervenus dans l'exercice.\\n\\n§3 -- SECTION B : SYSTEME DE GOUVERNANCE\\nStructure de gouvernance : CA, Direction Generale, Comite des Risques. Quatre fonctions cles : Actuarielle (Art. 48), Gestion des Risques, Audit Interne, Conformite. Politique de remuneration. Adequation de la gouvernance au profil de risque.\\n\\n§4 -- SECTION C : PROFIL DE RISQUE\\nSCR decompose par module (7 sous-modules Vie, Art. 138-145). Module dominant et effet de diversification (matrice EIOPA 7x7). Techniques d'attenuation du risque (reassurance, couvertures). Risque de liquidite : projection flux tresorerie. Sensibilites : +/-200bp taux, +/-10% mortalite.\\n\\n§5 -- SECTION D : VALORISATION A DES FINS DE SOLVABILITE\\nBE Vie (Art. 77) : methode, hypotheses, courbe RFR EIOPA. Risk Margin (Art. 78) : methode CoC 6%, durée de run-off. Ratio TP/BE [1.0, 1.5]. Reconciliation PM sociale / TP S2 / IFRS 17. Actifs en valeur de marche (Art. 75). Fonds propres S2 (Art. 87-99) : Tier 1/2/3.\\n\\n§6 -- SECTION E : GESTION DU CAPITAL\\nFonds propres eligibles par Tier. Ratio SCR (seuil 100%, cible ORSA 150%) et evolution N vs N-1. Ratio MCR (seuil 100%). Plan de gestion du capital. Conclusions ORSA (Art. 45) : adequation du capital au profil de risque. Alerte si ratio SCR < 150% avec recommandation.\\n\\n§7 -- CONCLUSION ET AVIS DE LA FONCTION ACTUARIELLE\\nAvis global : FAVORABLE A PUBLICATION / FAVORABLE AVEC RESERVE / DEFAVORABLE. Conditions FAVORABLE A PUBLICATION : toutes sections SFCR VERTES, ratio SCR >= 100%, ratio MCR >= 100%, ratio TP/BE in [1.0, 1.5], gouvernance conforme Art. 41-49. Points de vigilance pour la soumission ACPR. Recommandations par ordre de priorite + echeance reglementaire. Confirmation de la coherence SFCR/RSR."


def _f(v, dec=0) -> str:
    try:
        n = float(v or 0)
        s = f"{n:,.{dec}f}".replace(",", "\u202f")
        return s + "\u202f€"
    except Exception:
        return str(v) if v is not None else "—"

def _pct(v, dec=1) -> str:
    try:
        return f"{float(v or 0):.{dec}f}\u202f%"
    except Exception:
        return "—"

def _clean(txt) -> str:
    if not txt:
        return ""
    txt = re.sub(r"<[^>]+>", " ", str(txt))
    txt = re.sub(r"\s+", " ", txt)
    return txt.strip()

def _statut_cls(s: str) -> str:
    s = (s or "").upper()
    if "ROUGE" in s:
        return "rouge"
    if "AMBRE" in s or "RESERVE" in s:
        return "ambre"
    return "vert"

def _g(d, *keys, default=None):
    v = d or {}
    for k in keys:
        if isinstance(v, dict):
            v = v.get(k, default)
        else:
            return default
    return v if v is not None else default

def _css() -> str:
    return CSS_RVIE2

CSS_RVIE2 = "<style>\n@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,700;0,900;1,400&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');\n:root{--navy:#0B1E3D;--navy-mid:#132844;--navy-light:#1E3A5F;--gold:#C9A84C;--gold-pale:rgba(201,168,76,0.12);--rouge:#C0392B;--orange:#E67E22;--vert:#1E8449;--slate:#8A9AB0;--slate-light:#B8C5D3;--bg:#F5F7FA;--white:#FFFFFF;--text:#1C2B3A;--text-mid:#3D5166;--border:#DDE4EE;}\n*{box-sizing:border-box;margin:0;padding:0;}\nbody{font-family:'Inter',sans-serif;font-size:10pt;color:var(--text);background:var(--bg);line-height:1.7;}\n.rapport-container{max-width:960px;margin:0 auto;background:var(--white);box-shadow:0 4px 40px rgba(11,30,61,0.14);}\n.page-garde{position:relative;min-height:720px;display:flex;flex-direction:column;overflow:hidden;background:var(--navy);page-break-after:always;}\n.garde-bg{position:absolute;inset:0;overflow:hidden;pointer-events:none;}\n.garde-bg::before{content:'';position:absolute;top:-120px;right:-80px;width:560px;height:560px;border:1px solid rgba(201,168,76,0.12);border-radius:50%;}\n.garde-bg::after{content:'';position:absolute;top:-60px;right:-20px;width:380px;height:380px;border:1px solid rgba(201,168,76,0.20);border-radius:50%;}\n.garde-dots{position:absolute;top:0;right:0;width:340px;height:100%;background-image:radial-gradient(circle,rgba(201,168,76,0.25) 1px,transparent 1px);background-size:22px 22px;mask-image:linear-gradient(to left,rgba(0,0,0,0.7),transparent);}\n.garde-diagonal{position:absolute;inset:0;background:linear-gradient(135deg,transparent 62%,rgba(201,168,76,0.06) 62%);}\n.garde-accent-line{position:absolute;left:52px;top:0;bottom:0;width:2px;background:linear-gradient(to bottom,transparent,var(--gold) 20%,var(--gold) 80%,transparent);opacity:0.5;}\n.garde-inner{position:relative;z-index:2;display:flex;flex-direction:column;height:100%;}\n.garde-header{display:flex;justify-content:space-between;align-items:flex-start;padding:36px 56px 0;}\n.garde-logo-wrap img{height:52px;}\n.garde-badges{display:flex;flex-direction:column;align-items:flex-end;gap:8px;}\n.badge-confidentiel{font-size:7.5pt;font-weight:700;color:var(--rouge);letter-spacing:2.5px;border:1.5px solid var(--rouge);padding:5px 12px;text-transform:uppercase;}\n.badge-direction{font-size:7pt;font-weight:600;color:var(--gold);letter-spacing:1.5px;border:1px solid rgba(201,168,76,0.4);padding:4px 10px;text-transform:uppercase;}\n.garde-hero{flex:1;display:flex;flex-direction:column;justify-content:center;padding:48px 56px 0 80px;}\n.garde-eyebrow{font-size:8pt;font-weight:600;color:var(--gold);letter-spacing:4px;text-transform:uppercase;margin-bottom:20px;}\n.garde-titre{font-family:'Playfair Display',serif;font-size:42pt;font-weight:900;color:var(--white);line-height:1.05;margin-bottom:8px;}\n.garde-titre em{font-style:italic;color:var(--gold);}\n.garde-subtitle{font-family:'Playfair Display',serif;font-size:15pt;font-weight:400;font-style:italic;color:var(--slate-light);margin-bottom:8px;}\n.garde-sousdirection{font-size:9pt;font-weight:500;color:var(--gold);letter-spacing:1.5px;margin-bottom:28px;}\n.garde-sep{display:flex;align-items:center;gap:16px;margin-bottom:28px;}\n.garde-sep-line{height:1px;width:80px;background:var(--gold);opacity:0.6;}\n.garde-sep-diamond{width:8px;height:8px;background:var(--gold);transform:rotate(45deg);}\n.garde-statut{display:inline-flex;align-items:center;gap:12px;padding:10px 24px;}\n.garde-statut-rouge{background:rgba(192,57,43,0.15);border:1px solid rgba(192,57,43,0.5);}\n.garde-statut-ambre{background:rgba(230,126,34,0.15);border:1px solid rgba(230,126,34,0.5);}\n.garde-statut-vert{background:rgba(30,132,73,0.15);border:1px solid rgba(30,132,73,0.5);}\n.statut-dot{width:10px;height:10px;border-radius:50%;animation:pulse 2s ease-in-out infinite;}\n.statut-dot-rouge{background:var(--rouge);box-shadow:0 0 8px rgba(192,57,43,0.7);}\n.statut-dot-ambre{background:var(--orange);box-shadow:0 0 8px rgba(230,126,34,0.7);}\n.statut-dot-vert{background:var(--vert);box-shadow:0 0 8px rgba(30,132,73,0.7);}\n@keyframes pulse{0%,100%{opacity:1;transform:scale(1)}50%{opacity:0.7;transform:scale(1.2)}}\n.statut-label-rouge{font-size:8pt;font-weight:700;color:var(--rouge);letter-spacing:3px;text-transform:uppercase;}\n.statut-label-ambre{font-size:8pt;font-weight:700;color:var(--orange);letter-spacing:3px;text-transform:uppercase;}\n.statut-label-vert{font-size:8pt;font-weight:700;color:var(--vert);letter-spacing:3px;text-transform:uppercase;}\n.garde-footer{border-top:1px solid rgba(201,168,76,0.25);margin:36px 56px 0;padding:24px 0 40px;}\n.garde-kpis{display:grid;grid-template-columns:repeat(6,1fr);gap:0;}\n.garde-kpi{padding:0 20px 0 0;border-right:1px solid rgba(255,255,255,0.08);}\n.garde-kpi:last-child{border-right:none;}\n.kpi-label{font-size:6.5pt;font-weight:600;color:var(--slate);text-transform:uppercase;letter-spacing:1.5px;margin-bottom:6px;}\n.kpi-value{font-family:'Playfair Display',serif;font-size:11pt;font-weight:700;color:var(--white);line-height:1.2;}\n.kpi-value.highlight{color:var(--gold);font-size:13pt;}\n.kpi-sub{font-size:7pt;color:var(--slate);margin-top:3px;}\n.rapport-body{padding:0;}\n.section-header{background:linear-gradient(to right,var(--navy),var(--navy-light));padding:22px 56px;display:flex;align-items:center;gap:20px;}\n.section-num{font-family:'JetBrains Mono',monospace;font-size:9pt;font-weight:500;color:var(--gold);opacity:0.8;min-width:28px;}\n.section-titre{font-family:'Playfair Display',serif;font-size:15pt;font-weight:700;color:var(--white);}\n.section-body{padding:36px 56px 48px;}\n.section-divider{height:1px;background:var(--border);}\n.kpi-grid{display:grid;grid-template-columns:repeat(4,1fr);gap:12px;margin-bottom:32px;}\n.kpi-card{background:var(--navy);border-radius:6px;padding:20px 18px;position:relative;overflow:hidden;}\n.kpi-card::before{content:'';position:absolute;top:0;left:0;right:0;height:3px;background:var(--gold);}\n.kpi-card-rouge::before{background:var(--rouge);}\n.kpi-card-ambre::before{background:var(--orange);}\n.kpi-card-vert::before{background:var(--vert);}\n.kpi-card-label{font-size:7pt;font-weight:600;color:var(--slate);text-transform:uppercase;letter-spacing:1.5px;margin-bottom:8px;}\n.kpi-card-value{font-family:'Playfair Display',serif;font-size:16pt;font-weight:700;color:var(--white);line-height:1.1;margin-bottom:6px;}\n.kpi-card-value-rouge{color:var(--rouge)!important;}\n.kpi-card-value-ambre{color:var(--orange)!important;}\n.kpi-card-value-vert{color:#2ECC71!important;}\n.kpi-card-sub{font-size:7.5pt;color:var(--gold);}\n.section-card{border:1px solid var(--border);border-radius:8px;overflow:hidden;margin:12px 0;}\n.section-card-header{padding:14px 20px;background:var(--navy);display:flex;align-items:center;justify-content:space-between;}\n.section-card-lettre{font-family:'JetBrains Mono',monospace;font-size:20pt;font-weight:700;color:var(--gold);opacity:0.9;}\n.section-card-titre{font-size:11pt;font-weight:700;color:var(--white);}\n.section-card-statut{font-size:8pt;font-weight:700;letter-spacing:2px;text-transform:uppercase;padding:4px 12px;border-radius:3px;}\n.statut-vert{background:rgba(30,132,73,0.2);color:#2ECC71;}\n.statut-ambre{background:rgba(230,126,34,0.2);color:var(--orange);}\n.statut-rouge{background:rgba(192,57,43,0.2);color:var(--rouge);}\n.section-card-body{padding:16px 20px;background:#FAFBFC;font-size:9pt;color:var(--text-mid);line-height:1.7;}\n.avis-box{padding:24px 28px;border-radius:8px;margin:24px 0;display:flex;align-items:flex-start;gap:20px;}\n.avis-favorable{background:rgba(30,132,73,0.08);border:2px solid rgba(30,132,73,0.4);}\n.avis-reserve{background:rgba(230,126,34,0.08);border:2px solid rgba(230,126,34,0.4);}\n.avis-defavorable{background:rgba(192,57,43,0.08);border:2px solid rgba(192,57,43,0.5);}\n.avis-icon{font-size:24pt;flex-shrink:0;}\n.avis-label{font-size:7pt;font-weight:700;letter-spacing:2px;text-transform:uppercase;margin-bottom:6px;}\n.avis-value{font-family:'Playfair Display',serif;font-size:15pt;font-weight:700;margin-bottom:8px;}\n.avis-value-favorable{color:var(--vert);}\n.avis-value-reserve{color:var(--orange);}\n.avis-value-defavorable{color:var(--rouge);}\n.avis-text{font-size:9pt;color:var(--text-mid);line-height:1.6;}\n.commentaire-wrap{border:1px solid var(--border);border-radius:6px;overflow:hidden;margin-top:8px;}\n.commentaire-header{background:var(--navy);padding:14px 20px;display:flex;align-items:center;justify-content:space-between;}\n.commentaire-header-title{font-size:8pt;font-weight:600;color:var(--gold);letter-spacing:2px;text-transform:uppercase;}\n.commentaire-ai-badge{font-size:7pt;color:var(--slate);font-family:'JetBrains Mono',monospace;}\n.commentaire-body{padding:28px 32px;background:#FAFBFC;}\n.comm-section-title{font-family:'Playfair Display',serif;font-size:11pt;font-weight:700;color:var(--navy);border-bottom:1.5px solid var(--gold);padding-bottom:6px;margin:24px 0 12px;}\n.comm-section-title:first-child{margin-top:0;}\n.comm-h4{font-size:9.5pt;font-weight:700;color:var(--navy);margin:16px 0 8px;}\n.comm-p{font-size:9.5pt;line-height:1.85;color:var(--text-mid);margin-bottom:12px;}\n.comm-p strong{color:var(--navy);font-weight:700;}\n.comm-ul{margin:10px 0 12px 20px;list-style:none;}\n.comm-ul li{position:relative;padding-left:16px;font-size:9.5pt;line-height:1.7;color:var(--text-mid);margin-bottom:6px;}\n.comm-ul li::before{content:'—';position:absolute;left:0;color:var(--gold);font-weight:700;}\n.comm-ul li strong{color:var(--navy);}\n.comm-divider{border:none;border-top:1px solid var(--border);margin:20px 0;}\n.comm-footer{font-size:7.5pt;color:var(--slate);font-style:italic;text-align:right;border-top:1px solid var(--border);padding-top:10px;margin-top:8px;}\n.graph-wrap{border:1px solid var(--border);border-radius:6px;overflow:hidden;margin:20px 0;}\n.graph-titre{background:var(--navy-light);padding:10px 18px;font-size:8pt;font-weight:600;color:var(--gold);letter-spacing:1px;text-transform:uppercase;}\n.graph-body{padding:16px;background:#FAFBFC;}\n.pied-de-page{background:var(--navy);padding:20px 56px;display:flex;justify-content:space-between;align-items:center;border-top:2px solid var(--gold);}\n.pied-logo img{height:26px;}\n.pied-meta{text-align:right;font-size:7.5pt;color:var(--slate);line-height:1.7;}\n.pied-meta .confidentiel-footer{color:var(--rouge);font-weight:700;letter-spacing:1px;}\n@media print{body{background:white}.rapport-container{box-shadow:none;max-width:none;margin:0}.page-garde{page-break-after:always;min-height:100vh}@page{margin:0}}\n</style>"


def _md_to_html(texte: str) -> str:
    if not texte:
        return ""
    html = []
    for l in texte.split("\n"):
        l = l.rstrip()
        m = re.match(r"^§\s*(\d+)\s*[-—–]\s*(.+)$", l)
        if m:
            html.append(f'<div class="comm-section-title">§{m.group(1)} — {m.group(2).strip()}</div>')
            continue
        if l.startswith("###"):
            html.append(f'<div class="comm-h4">{l.lstrip("#").strip()}</div>')
            continue
        if not l.strip():
            html.append('<hr class="comm-divider">')
            continue
        if l.strip().startswith("- "):
            item = l.strip()[2:]
            item = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", item)
            html.append(f'<ul class="comm-ul"><li>{item}</li></ul>')
            continue
        line = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", l)
        html.append(f'<p class="comm-p">{line}</p>')
    return "\n".join(html)


def _construire_contexte_rvie2(
    result_rvie2, result_v5, result_rvie1,
    arrete,
) -> str:
    r2  = result_rvie2 or {}
    r5  = result_v5   or {}
    rr1 = result_rvie1 or {}

    sfcr = r2.get("sfcr", {}) or {}
    rsr  = r2.get("rsr",  {}) or {}
    ratio_scr_n  = float(r2.get("ratio_scr_pct", 0) or 0)
    ratio_scr_n1 = float(r2.get("ratio_scr_n1",  200) or 200)
    ratio_mcr    = float(r2.get("ratio_mcr_pct", 0)  or 0)
    ratio_tp_be  = float(r2.get("ratio_tp_be",   1.05) or 1.05)
    variation    = float(r2.get("variation_scr_pp", 0) or 0)
    be_vie       = float(rr1.get("be_vie", r5.get("tp_vie", 0)) or 0)
    scr_vie      = float(rr1.get("scr_vie_total", 0) or 0)
    fonds_propres= float(r2.get("fonds_propres", 0) or 0)
    avis_fa      = r2.get("avis_fa", "FAVORABLE AVEC RESERVE")

    # Statuts sections SFCR
    sections_str = ""
    for sec, data in sfcr.items():
        if isinstance(data, dict):
            sections_str += f"Section {sec} : statut={data.get('statut','?')} | {data.get('titre','')[:40]}\n"

    lignes = [
        f"RAPPORT RSR & SFCR VIE -- Arrete {arrete}",
        "(Identite de l'organisme non transmise : ne jamais la nommer ni l'inventer.)",
        f"Date generation : {datetime.now().strftime('%d/%m/%Y')}",
        f"Type rapport : SFCR (public) + RSR (confidentiel ACPR)",
        "",
        "=== SECTION E -- GESTION DU CAPITAL ===",
        f"Fonds propres eligibles          : {_f(fonds_propres)}",
        f"SCR Vie total                    : {_f(scr_vie)}",
        f"Ratio SCR N                      : {_pct(ratio_scr_n)} (seuil 100%, cible ORSA 150%)",
        f"Ratio SCR N-1                    : {_pct(ratio_scr_n1)}",
        f"Variation SCR N vs N-1           : {variation:+.1f} pp",
        f"Ratio MCR                        : {_pct(ratio_mcr)} (seuil 100%)",
        f"Plan capital                     : {rsr.get('plan_capital', 'non disponible')}",
        "",
        "=== SECTION D -- VALORISATION ===",
        f"Best Estimate Vie (Art. 77)      : {_f(be_vie)}",
        f"Ratio TP/BE                      : {ratio_tp_be:.4f} (seuil conforme : 1.0-1.5)",
        f"Risk Margin (Art. 78 CoC 6%)     : integre dans TP",
        f"Mode BE IFRS 17                  : {r5.get('be_ifrs17_mode', 'approximation S2') if r5 else 'non disponible'}",
        "",
        "=== SCR VIE PAR MODULE (ART. 138-145) ===",
        f"SCR mortalite (Art. 137 +15%)    : {_f(float((rr1.get('scr_par_module') or {}).get('mortalite', 0) or 0))}",
        f"SCR longevite (Art. 138 -20%)    : {_f(float((rr1.get('scr_par_module') or {}).get('longevite', 0) or 0))}",
        f"SCR rachat   (Art. 142)          : {_f(float((rr1.get('scr_par_module') or {}).get('rachat', 0) or 0))}",
        f"SCR frais    (Art. 143)          : {_f(float((rr1.get('scr_par_module') or {}).get('frais', 0) or 0))}",
        f"Effet diversification            : {_f(float(rr1.get('effet_diversification', 0) or 0))} economise",
        "",
        "=== SECTIONS SFCR A-E ===",
        sections_str or "Sections non disponibles -- alimenter result_rvie2",
        "",
        "=== AVIS ACTUAIRE DESIGNE ===",
        f"Avis Fonction Actuarielle        : {avis_fa}",
        f"Statut global SFCR               : {r2.get('statut_rag', 'non disponible')}",
        "",
        "=== STATUTS VALIDATION ===",
        f"V5 Nia (QRT S.12/S.23)           : {r5.get('statut_rag', 'non disponible')}",
        f"R-VIE1 Eric (SCR S.26)           : {rr1.get('statut_rag', 'non disponible')}",
        f"R-VIE2 Camille (RSR/SFCR)        : {r2.get('statut_rag', 'non disponible')}",
        "",
        "Redige le commentaire actuariel RSR & SFCR en 7 sections (§1 a §7) selon les regles du prompt.",
    ]
    return "\n".join(lignes)


def _narration_claude_api_rvie2(contexte: str) -> str:
    resp = frontiere_llm.appeler(
        modele=frontiere_llm.MODELE_ETABLI,
        max_tokens=12000,
        systeme=SYSTEM_PROMPT_RVIE2,
        messages=[{"role": "user", "content": contexte}],
        cle=frontiere_llm.cle_api_ou_secrets(),
    )
    return frontiere_llm.texte_des_blocs(resp)


def _generer_narration_rvie2(contexte: str, commentaire: str) -> Tuple[str, str]:
    try:
        txt = _narration_claude_api_rvie2(contexte)
        if txt:
            return txt, "claude_api"
    except Exception as e:
        logger.warning(f"Narration non générée : {e}")
    if commentaire and commentaire.strip():
        return _clean(commentaire), "manuel"
    return "", "aucune"


def _kpi_card(label, value, sub="", cls="") -> str:
    cls_val  = f" kpi-card-value-{cls}" if cls else ""
    cls_card = f" kpi-card-{cls}" if cls else ""
    return (
        f'<div class="kpi-card{cls_card}">'''
        f'<div class="kpi-card-label">{label}</div>'''
        f'<div class="kpi-card-value{cls_val}">{value}</div>'''
        f'<div class="kpi-card-sub">{sub}</div></div>'
    )

def _section_sfcr_card(lettre, titre, statut, details) -> str:
    sc = statut.lower() if statut else "ambre"
    return (
        f'<div class="section-card">'''
        f'<div class="section-card-header">'''
        f'<div><span class="section-card-lettre">{lettre}</span></div>'''
        f'<div class="section-card-titre">{titre}</div>'''
        f'<div class="section-card-statut statut-{sc}">{statut}</div>'''
        f'</div>'''
        f'<div class="section-card-body">{details}</div>'''
        f'</div>'
    )

def _wrap_graph(html_g: str, titre: str) -> str:
    if not html_g:
        return ""
    return (
        f'<div class="graph-wrap">'''
        f'<div class="graph-titre">{titre}</div>'''
        f'<div class="graph-body">{html_g}</div></div>'
    )

def _avis_box(avis: str, texte: str) -> str:
    a = avis.upper()
    if "DEFAV" in a:
        cls, icon, vcls = "avis-defavorable", "❌", "avis-value-defavorable"
    elif "RESERVE" in a:
        cls, icon, vcls = "avis-reserve",    "⚠️", "avis-value-reserve"
    else:
        cls, icon, vcls = "avis-favorable",  "✅", "avis-value-favorable"
    return (
        f'<div class="avis-box {cls}">'''
        f'<div class="avis-icon">{icon}</div>'''
        f'<div>'''
        f'<div class="avis-label">Avis de la Fonction Actuarielle — Art. 48 §g S2</div>'''
        f'<div class="avis-value {vcls}">{avis}</div>'''
        f'<div class="avis-text">{texte}</div>'''
        f'</div></div>'
    )


def export_html(
    result_rvie2=None, result_v5=None, result_rvie1=None,
    commentaire="", ref_client="", arrete="",
    audit_id="", graphiques=None,
    actuaire_nom="", actuaire_numero_ia="",
) -> str:
    """Genere le rapport HTML RSR & SFCR Vie — niveau premium A7."""
    try:
        r2  = result_rvie2 or {}
        r5  = result_v5   or {}
        rr1 = result_rvie1 or {}

        dt  = datetime.now().strftime("%d/%m/%Y %H:%M")
        arr = arrete or datetime.now().strftime("%d/%m/%Y")
        cli = ref_client or "A renseigner"

        ratio_scr_n  = float(r2.get("ratio_scr_pct", 0) or 0)
        ratio_scr_n1 = float(r2.get("ratio_scr_n1",  200) or 200)
        ratio_mcr    = float(r2.get("ratio_mcr_pct", 0)  or 0)
        ratio_tp_be  = float(r2.get("ratio_tp_be",   1.05) or 1.05)
        variation    = float(r2.get("variation_scr_pp", 0) or 0)
        be_vie       = float(rr1.get("be_vie", 0) or 0)
        fonds_propres= float(r2.get("fonds_propres", 0) or 0)
        avis_fa      = r2.get("avis_fa", "FAVORABLE AVEC RESERVE")
        sfcr         = r2.get("sfcr", {}) or {}

        sg = r2.get("statut_rag", "AMBRE")
        sc = _statut_cls(sg)
        dot_cls    = f"statut-dot-{sc}"
        statut_cls_str = f"garde-statut-{sc}"
        label_cls  = f"statut-label-{sc}"
        label_txt  = {"rouge":"Non publiable","ambre":"Reserves signalees","vert":"Publiable"}[sc]

        ghtml = {}
        if graphiques and PLOTLY_OK:
            for nom, fig_or_html in graphiques.items():
                try:
                    ghtml[nom] = fig_or_html if isinstance(fig_or_html, str) else pio.to_html(
                        fig_or_html, full_html=False, include_plotlyjs=False,
                        config={"displayModeBar": False})
                except Exception:
                    pass

        contexte = _construire_contexte_rvie2(r2, r5, rr1, arr)
        narration, source = _generer_narration_rvie2(contexte, commentaire)

        # KPI page de garde
        garde_kpis = (
            f'<div class="garde-kpi"><div class="kpi-label">Ratio SCR</div>'''
            f'<div class="kpi-value{"  highlight" if ratio_scr_n >= 150 else ""}">{_pct(ratio_scr_n)}</div>'''
            f'<div class="kpi-sub">Cible ORSA ≥ 150%</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Variation SCR</div>'''
            f'<div class="kpi-value">{variation:+.1f} pp</div>'''
            f'<div class="kpi-sub">vs N-1 : {_pct(ratio_scr_n1)}</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Ratio MCR</div>'''
            f'<div class="kpi-value">{_pct(ratio_mcr)}</div>'''
            f'<div class="kpi-sub">Seuil ≥ 100%</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Ratio TP/BE</div>'''
            f'<div class="kpi-value">{ratio_tp_be:.4f}</div>'''
            f'<div class="kpi-sub">Seuil : 1.0–1.5</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">BE Vie</div>'''
            f'<div class="kpi-value">{_f(be_vie)}</div>'''
            f'<div class="kpi-sub">Art. 77 S2</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Fonds propres</div>'''
            f'<div class="kpi-value">{_f(fonds_propres)}</div>'''
            f'<div class="kpi-sub">Tier 1/2/3</div></div>'
        )

        kpi_grid = '<div class="kpi-grid">'
        kpi_grid += _kpi_card("Ratio SCR N", _pct(ratio_scr_n), "Seuil 100% · Cible ORSA 150%",
                               cls="rouge" if ratio_scr_n < 100 else "ambre" if ratio_scr_n < 150 else "vert")
        kpi_grid += _kpi_card("Ratio SCR N-1", _pct(ratio_scr_n1), f"Variation : {variation:+.1f} pp")
        kpi_grid += _kpi_card("Ratio MCR", _pct(ratio_mcr), "Seuil ≥ 100%",
                               cls="rouge" if ratio_mcr < 100 else "vert")
        kpi_grid += _kpi_card("Ratio TP/BE", f"{ratio_tp_be:.4f}", "Seuil : 1.0–1.5",
                               cls="rouge" if not (1.0 <= ratio_tp_be <= 1.5) else "vert")
        kpi_grid += _kpi_card("BE Vie", _f(be_vie), "Art. 77 S2")
        kpi_grid += _kpi_card("Fonds propres", _f(fonds_propres), "Art. 87-99 S2")
        kpi_grid += '</div>'

        avis_html = _avis_box(avis_fa,
            r2.get("conseil_global", "Rapport RSR & SFCR produit par ActuarIA Direction Vie & EP-RE."))

        # Cards sections SFCR A-E
        sfcr_sections_html = ""
        SFCR_LABELS = {
            "A": "Activité et Résultats",
            "B": "Système de Gouvernance",
            "C": "Profil de Risque",
            "D": "Valorisation",
            "E": "Gestion du Capital",
        }
        for lettre, label in SFCR_LABELS.items():
            data = sfcr.get(lettre, {}) or {}
            statut = data.get("statut", "AMBRE")
            details_list = data.get("details", data.get("points_cles", []))
            details = " · ".join(str(d) for d in details_list[:3]) if details_list else data.get("titre","")
            sfcr_sections_html += _section_sfcr_card(lettre, label, statut, details or "En cours de consolidation")

        narration_html = _md_to_html(narration) if narration else (
            '<p class="comm-p" style="color:var(--slate);font-style:italic;">'''
            "Narration non disponible — alimenter result_rvie2 et configurer ANTHROPIC_API_KEY.</p>"
        )
        ai_badge = ("❆ Narration générée par ActuarIA Intelligence"
                    if source == "claude_api" else "❆ Commentaire actuariel manuel")

        H = []
        H.append('<!DOCTYPE html>')
        H.append('<html lang="fr"><head><meta charset="UTF-8">')
        H.append(f'<title>RSR &amp; SFCR Vie — {cli} — {arr}</title>')
        H.append('<script src="https://cdn.jsdelivr.net/npm/plotly.js-dist@2.26.0/plotly.min.js"></script>')
        H.append(_css())
        H.append('</head><body><div class="rapport-container">')

        H.append('<div class="page-garde">')
        H.append('<div class="garde-bg"><div class="garde-dots"></div>')
        H.append('<div class="garde-diagonal"></div><div class="garde-accent-line"></div></div>')
        H.append('<div class="garde-inner">')
        H.append(f'<div class="garde-header"><div class="garde-logo-wrap"><img src="{LOGO_URI}" alt="ActuarIA"/></div>')
        H.append('<div class="garde-badges"><span class="badge-confidentiel">⬛ Confidentiel ACPR</span>')
        H.append('<span class="badge-direction">Pilier 3 Solvabilité 2</span></div></div>')
        H.append('<div class="garde-hero">')
        H.append('<div class="garde-eyebrow">Rapport Réglementaire Pilier 3</div>')
        H.append('<div class="garde-titre">RSR &amp; <em>SFCR</em></div>')
        H.append(f'<div class="garde-subtitle">Vie — Arrêté au {arr}</div>')
        H.append(f'<div class="garde-sousdirection">Regular Supervisory Report · Solvency &amp; Financial Condition Report · {cli}</div>')
        H.append('<div class="garde-sep"><div class="garde-sep-line"></div>')
        H.append('<div class="garde-sep-diamond"></div><div class="garde-sep-line"></div></div>')
        H.append(f'<div class="garde-statut {statut_cls_str}"><div class="statut-dot {dot_cls}"></div>')
        H.append(f'<div class="{label_cls}">{label_txt}</div></div>')
        H.append('</div>')
        H.append(f'<div class="garde-footer"><div class="garde-kpis">{garde_kpis}</div></div>')
        H.append('</div></div>')

        H.append('<div class="rapport-body">')

        # §1 Synthese
        H.append('<div class="section-header"><span class="section-num">01</span>')
        H.append('<span class="section-titre">Synthèse exécutive RSR &amp; SFCR</span></div>')
        H.append(f'<div class="section-body">{kpi_grid}')
        H.append(_wrap_graph(ghtml.get("scorecard_rsr_sfcr",""), "Scorecard RSR/SFCR — Sections A-E"))
        H.append(avis_html)
        H.append('</div><div class="section-divider"></div>')

        # §2 Section A
        H.append('<div class="section-header"><span class="section-num">02</span>')
        H.append('<span class="section-titre">Section A — Activité et Résultats</span></div>')
        H.append(f'<div class="section-body">{_section_sfcr_card("A","Activité et Résultats",(sfcr.get("A") or {}).get("statut","AMBRE"),str((sfcr.get("A") or {}).get("titre",""))[:200])}')
        H.append(_wrap_graph(ghtml.get("sections_sfcr",""), "Sections SFCR A-E — Statuts de conformité"))
        H.append('</div><div class="section-divider"></div>')

        # §3 Section B
        H.append('<div class="section-header"><span class="section-num">03</span>')
        H.append('<span class="section-titre">Section B — Système de Gouvernance</span></div>')
        H.append(f'<div class="section-body">{_section_sfcr_card("B","Système de Gouvernance",(sfcr.get("B") or {}).get("statut","AMBRE"),str((sfcr.get("B") or {}).get("titre",""))[:200])}')
        H.append('</div><div class="section-divider"></div>')

        # §4 Section C
        H.append('<div class="section-header"><span class="section-num">04</span>')
        H.append('<span class="section-titre">Section C — Profil de Risque</span></div>')
        H.append(f'<div class="section-body">{_section_sfcr_card("C","Profil de Risque",(sfcr.get("C") or {}).get("statut","AMBRE"),str((sfcr.get("C") or {}).get("titre",""))[:200])}')
        H.append(_wrap_graph(ghtml.get("evolution_ratio_scr",""), "Evolution Ratio SCR N vs N-1"))
        H.append('</div><div class="section-divider"></div>')

        # §5 Section D
        H.append('<div class="section-header"><span class="section-num">05</span>')
        H.append('<span class="section-titre">Section D — Valorisation</span></div>')
        H.append(f'<div class="section-body"><div class="kpi-grid">')
        H.append(_kpi_card("BE Vie", _f(be_vie), "Art. 77 S2"))
        H.append(_kpi_card("Ratio TP/BE", f"{ratio_tp_be:.4f}", "Seuil 1.0–1.5",
                            cls="rouge" if not (1.0 <= ratio_tp_be <= 1.5) else "vert"))
        H.append(_kpi_card("Fonds propres", _f(fonds_propres), "Art. 87-99"))
        H.append(_kpi_card("Mode BE", "Art. 77 S2", "Methode prospective"))
        H.append('</div>')
        H.append(_section_sfcr_card("D","Valorisation",(sfcr.get("D") or {}).get("statut","AMBRE"),
                                     str((sfcr.get("D") or {}).get("titre",""))[:200]))
        H.append(_wrap_graph(ghtml.get("jauge_tp_be_sfcr",""), "Ratio TP/BE — Jauge de conformité"))
        H.append('</div><div class="section-divider"></div>')

        # §6 Section E
        H.append('<div class="section-header"><span class="section-num">06</span>')
        H.append('<span class="section-titre">Section E — Gestion du Capital</span></div>')
        H.append(f'<div class="section-body"><div class="kpi-grid">')
        H.append(_kpi_card("Ratio SCR", _pct(ratio_scr_n), "Seuil 100% · Cible 150%",
                            cls="rouge" if ratio_scr_n < 100 else "ambre" if ratio_scr_n < 150 else "vert"))
        H.append(_kpi_card("Ratio MCR", _pct(ratio_mcr), "Seuil ≥ 100%",
                            cls="rouge" if ratio_mcr < 100 else "vert"))
        H.append(_kpi_card("Variation SCR", f"{variation:+.1f} pp", f"N-1 : {_pct(ratio_scr_n1)}"))
        H.append(_kpi_card("Plan capital", "OK" if ratio_scr_n >= 150 else "Surveillance",
                            "ORSA Art. 45"))
        H.append('</div>')
        H.append(_section_sfcr_card("E","Gestion du Capital",(sfcr.get("E") or {}).get("statut","AMBRE"),
                                     str((sfcr.get("E") or {}).get("titre",""))[:200]))
        H.append(_wrap_graph(ghtml.get("sections_sfcr",""), "Sections SFCR — Vue d'ensemble"))
        H.append('</div><div class="section-divider"></div>')

        # §7 Commentaire
        H.append('<div class="section-header"><span class="section-num">07</span>')
        H.append('<span class="section-titre">Commentaire Actuariel — Avis de la Fonction Actuarielle</span></div>')
        H.append('<div class="section-body"><div class="commentaire-wrap">')
        H.append('<div class="commentaire-header">')
        H.append('<div class="commentaire-header-title">Analyse RSR &amp; SFCR Vie — Pilier 3 Solvabilité 2</div>')
        H.append(f'<div class="commentaire-ai-badge">{ai_badge}</div></div>')
        H.append('<div class="commentaire-body">')
        H.append(narration_html)
        H.append(f'<div class="comm-footer">Rédigé le {dt} · Actuaire : {actuaire_nom or "A renseigner"} · N° IA : {actuaire_numero_ia or "—"} · CONFIDENTIEL ACPR</div>')
        H.append('</div></div></div>')

        H.append('</div>')  # fin rapport-body

        H.append('<div class="pied-de-page">')
        H.append(f'<div class="pied-logo"><img src="{LOGO_URI}" alt="ActuarIA"/></div>')
        H.append(f'<div class="pied-meta"><div>{cli} · RSR &amp; SFCR Vie · Arrêté {arr}</div>')
        H.append(f'<div>Audit ID : {audit_id or "—"} · Généré le {dt}</div>')
        H.append('<div class="confidentiel-footer">⬛ CONFIDENTIEL ACPR — USAGE REGLEMENTAIRE EXCLUSIF</div>')
        H.append('</div></div>')
        H.append('</div></body></html>')

        html = "\n".join(H)
        logger.info(f"HTML RSR/SFCR genere : {len(html):,} chars, source={source}")
        return html

    except Exception as e:
        logger.error(f"export_html RSR/SFCR : {e}", exc_info=True)
        return f"<html><body><h1>Erreur</h1><p>{e}</p></body></html>"


def export_pdf(**kwargs) -> bytes:
    try:
        if not WEASYPRINT_OK:
            raise ImportError("weasyprint requis")
        html_str = export_html(**kwargs)
        pdf = WP_HTML(string=html_str).write_pdf()
        logger.info(f"PDF RSR/SFCR : {len(pdf):,} bytes")
        return pdf
    except Exception as e:
        logger.error(f"export_pdf RSR/SFCR : {e}", exc_info=True)
        return b""


def export_word(
    result_rvie2=None, result_v5=None, result_rvie1=None,
    commentaire="", ref_client="", arrete="",
    audit_id="", actuaire_nom="", actuaire_numero_ia="",
) -> bytes:
    """Genere le rapport Word RSR & SFCR — 7 chapitres."""
    if not DOCX_OK:
        logger.error("python-docx absent"); return b""
    try:
        def rgb(h):
            h = h.lstrip("#")
            return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))
        NR=rgb(NAVY); GR=rgb(GOLD); BR=rgb(WHITE)
        SR=rgb(SLATE); RgR=rgb(ROUGE); VR=rgb(VERT); AR=rgb(ORANGE)

        dt  = datetime.now().strftime("%d/%m/%Y")
        arr = arrete or dt
        cli = ref_client or "A renseigner"

        r2  = result_rvie2 or {}
        r5  = result_v5   or {}
        rr1 = result_rvie1 or {}

        ratio_scr_n  = float(r2.get("ratio_scr_pct", 0) or 0)
        ratio_scr_n1 = float(r2.get("ratio_scr_n1",  200) or 200)
        ratio_mcr    = float(r2.get("ratio_mcr_pct", 0)  or 0)
        ratio_tp_be  = float(r2.get("ratio_tp_be",   1.05) or 1.05)
        variation    = float(r2.get("variation_scr_pp", 0) or 0)
        be_vie       = float(rr1.get("be_vie", 0) or 0)
        fonds_propres= float(r2.get("fonds_propres", 0) or 0)
        avis_fa      = r2.get("avis_fa", "FAVORABLE AVEC RESERVE")
        sfcr         = r2.get("sfcr", {}) or {}

        contexte = _construire_contexte_rvie2(r2, r5, rr1, arr)
        narration, source = _generer_narration_rvie2(contexte, commentaire)

        doc = Document()
        for s in doc.sections:
            s.top_margin=Cm(2.2); s.bottom_margin=Cm(2.2)
            s.left_margin=Cm(2.8); s.right_margin=Cm(2.8)

        def _bg(cell, hex6):
            tc=cell._tc; tcp=tc.get_or_add_tcPr()
            sd=OxmlElement("w:shd")
            sd.set(qn("w:fill"), hex6.lstrip("#"))
            sd.set(qn("w:color"),"auto"); sd.set(qn("w:val"),"clear")
            tcp.append(sd)

        def _run(p, txt, bold=False, italic=False, sz=10, col=None):
            r=p.add_run(str(txt)); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
            if col: r.font.color.rgb=col
            return r

        def _h(txt, lv=1, col=None):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(10 if lv==1 else 6)
            p.paragraph_format.space_after=Pt(4)
            sz={1:16,2:12,3:10}.get(lv,10)
            c=col or (NR if lv==1 else GR)
            _run(p, txt, bold=True, sz=sz, col=c)

        def _sep():
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
            bo=OxmlElement("w:bottom")
            bo.set(qn("w:val"),"single"); bo.set(qn("w:sz"),"6")
            bo.set(qn("w:space"),"1"); bo.set(qn("w:color"),"C9A84C")
            pBdr.append(bo); pPr.append(pBdr)

        def _tbl(heads, rows, ws=None):
            from docx.enum.table import WD_ALIGN_VERTICAL
            t=doc.add_table(rows=1+len(rows), cols=len(heads)); t.style="Table Grid"
            for i,hd in enumerate(heads):
                c=t.rows[0].cells[i]; _bg(c,"0B1E3D")
                pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=pp.add_run(str(hd)); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=BR
            for ri,row in enumerate(rows):
                for ci,v in enumerate(row):
                    c=t.rows[ri+1].cells[ci]
                    if ri%2==1: _bg(c,"EEF2F7")
                    pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r=pp.add_run(str(v) if v is not None else "—"); r.font.size=Pt(9)
            if ws:
                for i,w in enumerate(ws):
                    for row in t.rows: row.cells[i].width=Cm(w)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)

        # Page de garde
        p=doc.add_paragraph()
        _run(p,"Actuar",bold=True,sz=28,col=NR); _run(p,"IA",bold=True,sz=28,col=GR)
        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p,"RSR & SFCR VIE\n",bold=True,sz=22,col=NR)
        _run(p,"Regular Supervisory Report & Solvency and Financial Condition Report",bold=True,sz=12,col=GR)
        doc.add_paragraph()
        avis_col = VR if "FAVORABLE" in avis_fa.upper() and "RESERVE" not in avis_fa.upper() else (
                   AR if "RESERVE" in avis_fa.upper() else RgR)
        p=doc.add_paragraph()
        _run(p,"Avis Fonction Actuarielle : ",sz=11,col=NR)
        _run(p,avis_fa,bold=True,sz=11,col=avis_col)
        doc.add_paragraph()
        _tbl(["Societe","Type rapport","Arrete","Actuaire designe"],
             [[cli,"SFCR (public) + RSR (confidentiel)",arr,actuaire_nom or "A renseigner"]],
             ws=[3.5,5.0,3.0,4.5])
        doc.add_page_break()

        # §1 Contexte
        _h("§1 — Contexte et Perimetre du Reporting Pilier 3"); _sep()
        _tbl(["Indicateur","Valeur","Reference"],
             [["Type rapport","SFCR (public) + RSR (confidentiel ACPR)","Art. 51-56 S2"],
              ["Calendrier SFCR","4 mois apres cloture","Art. 307 Actes delegues"],
              ["Calendrier RSR","4 mois apres cloture","Art. 307"],
              ["Perimetre","Assurance vie individuelle et collective","Directive 2009/138/CE"],
              ["Date arrete",arr,"—"]],
             ws=[5.0,6.0,5.0])
        doc.add_page_break()

        # §2-§6 Sections A-E
        SFCR_SECTIONS = [
            ("A","Activite et Resultats","Art. 51 §a"),
            ("B","Systeme de Gouvernance","Art. 41-49 · Art. 51 §b"),
            ("C","Profil de Risque","Art. 100-127 · Art. 51 §c"),
            ("D","Valorisation","Art. 75-86 · Art. 51 §d"),
            ("E","Gestion du Capital","Art. 87-131 · Art. 51 §e"),
        ]
        for i, (lettre, titre_sec, ref_sec) in enumerate(SFCR_SECTIONS, 2):
            _h(f"§{i} — Section {lettre} — {titre_sec}"); _sep()
            data = sfcr.get(lettre, {}) or {}
            statut_sec = data.get("statut", "AMBRE")
            statut_col = VR if statut_sec == "VERT" else AR if statut_sec == "AMBRE" else RgR
            p=doc.add_paragraph()
            _run(p,f"Statut : ",sz=9,col=NR)
            _run(p,statut_sec,bold=True,sz=9,col=statut_col)
            _run(p,f" | Reference : {ref_sec}",sz=9,col=NR)
            titre_data = str(data.get("titre",""))
            if titre_data:
                p=doc.add_paragraph()
                p.paragraph_format.left_indent=Cm(0.4)
                _run(p,titre_data[:300],sz=9,col=NR)
            doc.add_page_break()

        # Indicateurs cles D et E
        _h("§7 — Indicateurs Cles Sections D &amp; E"); _sep()
        _tbl(
            ["Indicateur","Valeur","Reference","Seuil"],
            [["BE Vie (Art. 77)",_f(be_vie),"Art. 77 S2","—"],
             ["Ratio TP/BE",f"{ratio_tp_be:.4f}","Art. 75-86","[1.0, 1.5]"],
             ["Fonds propres eligibles",_f(fonds_propres),"Art. 87-99","—"],
             ["Ratio SCR N",_pct(ratio_scr_n),"Art. 100-127","≥ 100% (cible 150%)"],
             ["Ratio SCR N-1",_pct(ratio_scr_n1),"Exercice precedent","—"],
             ["Variation SCR",f"{variation:+.1f} pp","Evolution","—"],
             ["Ratio MCR",_pct(ratio_mcr),"Art. 129","≥ 100%"]],
            ws=[5.0,3.5,4.0,4.0])
        doc.add_page_break()

        # §8 Commentaire actuariel
        _h("§8 — Commentaire Actuariel — Avis RSR/SFCR"); _sep()
        if narration:
            secs = re.split(r"(?=§\d+\s*[-—–])", _clean(narration))
            for sec in secs:
                sec = sec.strip()
                if not sec: continue
                ls = sec.split("\n",1)
                if ls[0]: _h(ls[0],lv=2)
                if len(ls) > 1:
                    for ln in ls[1].split("\n"):
                        ln = ln.strip()
                        if ln:
                            p=doc.add_paragraph()
                            p.paragraph_format.space_after=Pt(3)
                            p.paragraph_format.left_indent=Cm(0.4)
                            _run(p,ln,sz=9,col=NR)
            if source=="claude_api":
                p=doc.add_paragraph()
                _run(p,"❆ Narration generee par ActuarIA Intelligence",sz=7,italic=True,col=SR)
        else:
            p=doc.add_paragraph()
            _run(p,"Narration non disponible — configurer ANTHROPIC_API_KEY.",sz=9,italic=True,col=SR)

        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p,"Avis Fonction Actuarielle (Art. 48 §g Directive 2009/138/CE) : ",sz=10,col=NR)
        _run(p,avis_fa,bold=True,sz=12,col=avis_col)

        _sep()
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p,f"ActuarIA · {cli} · RSR & SFCR Vie · Arrêté {arr} · {dt} · CONFIDENTIEL ACPR",
             sz=7,italic=True,col=SR)

        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        wb=buf.read()
        logger.info(f"Word RSR/SFCR : {len(wb):,} bytes")
        return wb

    except Exception as e:
        logger.error(f"export_word RSR/SFCR : {e}", exc_info=True)
        return b""
