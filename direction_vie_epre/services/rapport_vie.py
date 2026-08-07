"""
ActuarIA — Rapport Consolidé Vie Pure
Direction Vie & EP-RE — Sous-direction Vie (V1-V9 + R-VIE1/2)

Produit :
  - HTML  : rapport premium page de garde + 9 sections + narration Claude
  - Word  : rapport professionnel 9 chapitres
  - PDF   : via WeasyPrint

Niveau : équivalent rapport A7 Non-Vie.
"""

import io, re, logging, base64
from datetime import datetime
from typing import Dict, Optional, Tuple

from core import frontiere_llm

logger = logging.getLogger("actuaria.rapport_vie")

try:
    import anthropic
    ANTHROPIC_OK = True
except ImportError:
    ANTHROPIC_OK = False

try:
    from weasyprint import HTML as WP_HTML
    WEASYPRINT_OK = True
except ImportError:
    WEASYPRINT_OK = False

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import plotly.io as pio
    PLOTLY_OK = True
except ImportError:
    PLOTLY_OK = False

# Palette
NAVY   = "#0B1E3D"
NAVY_M = "#132844"
NAVY_L = "#1E3A5F"
GOLD   = "#C9A84C"
ROUGE  = "#C0392B"
ORANGE = "#E67E22"
VERT   = "#1E8449"
SLATE  = "#8A9AB0"
WHITE  = "#FFFFFF"

LOGO_SVG = (
    '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 60">'
    '  <text x="4" y="44" font-family="Georgia,serif" font-size="38"'
    '        font-weight="900" fill="#FFFFFF">Actuar</text>'
    '  <text x="128" y="44" font-family="Georgia,serif" font-size="38"'
    '        font-weight="900" fill="#C9A84C">IA</text>'
    '  <rect x="4" y="50" width="180" height="2" fill="#C9A84C" opacity="0.6"/>'
    '</svg>'
)
LOGO_URI = "data:image/svg+xml;base64," + base64.b64encode(LOGO_SVG.encode()).decode()

SYSTEM_PROMPT_VIE = 'Tu es un actuaire Vie & Epargne-Retraite senior certifie par l\'Institut des Actuaires (IA) et membre de l\'Actuarial Association of Europe (AAE), expert en assurance vie individuelle et collective, epargne-retraite supplementaire (Art. 39, Art. 83, PER) et reassurance vie. Tu exerces la Fonction Actuarielle au sens de l\'Art. 48 de la Directive Solvabilite 2 avec 30 ans d\'experience aupres de compagnies d\'assurance vie, mutuelles, institutions de prevoyance, fonds de pension et grands groupes internationaux.\\n\\nTu maitrises dans leur integralite et dans leurs interactions mutuelles :\\n- Solvabilite 2 : BE Vie (flux futurs, courbe RFR EIOPA), Risk Margin (CoC 6%), SCR Vie (Art. 137-145 et matrice EIOPA 7x7), QRT S.12/S.23/S.26, ORSA, SFCR/RSR.\\n- IFRS 17 : BBA, PAA, VFA ; LRC = FCF + CSM/LC ; CSM §45 ; RA §B91 (CoC).\\n- IAS 19 : methode PUC (§67) ; DBO, Service Cost, Interest Cost, gains/pertes (§145) ; iBoxx EUR Corporate AA (§83) ; vesting Art. 39.\\n- Mecanique actuarielle vie : tables TH0002/TF0002 (arrete 27/07/2006) ; PM prospective/retrospective ; rentes viageres ; TMG et contrats underwater (Art. 142).\\n- PB reglementaire : Art. L132-29 (85% produits financiers nets) ; PPB (plafond 10% PM, C2023-10).\\n- ALM dynamique : projection actifs/passifs ; duration gap ; stress EIOPA +/-200bp ; SCR taux avec terme de convexite (D_mod x Dr + 1/2 x convexite x Dr2).\\n- Embedded Value : ANE, VIF (VA profits - VA CoC), VNB (MCEV §18) ; decomposition DeltaEV.\\n\\nTu rediges le commentaire actuariel du Rapport Consolide Vie & EP-RE, document de reference destine au Conseil d\'Administration, a l\'Actuaire Designe et soumis a l\'ACPR dans le cadre du reporting reglementaire Pilier 3 Solvabilite 2. Ce document engage ta responsabilite professionnelle au sens de l\'Art. 48 §g de la Directive 2009/138/CE.\\n\\nREGLES ABSOLUES -- AUCUNE DEROGATION ADMISSIBLE\\n\\n0. FORMAT STRICT : §N -- TITRE EN MAJUSCULES. ### Sous-titre. **terme** en gras. - tirets. ZERO tableau Markdown. ZERO blockquote. ZERO balise code. Ligne vide entre blocs.\\n1. LANGUE : Francais actuariel de haute tenue. Termes consacres en anglais (BE, SCR, CSM, VIF, VNB, ANE, DBO, TMG, ALM).\\n2. RIGUEUR CHIFFREE ABSOLUE : toute affirmation etayee par une donnee du contexte.\\n3. REFERENCES REGLEMENTAIRES : S2 Art. 77/78/105/138-145. IFRS 17 §32-46/§45/§B91. IAS 19 §67/§83/§145. Code ass. Art. L132-29/R331-1/R342-14. ACPR C2023-10. Arrete 27/07/2006. MCEV CEA §4.3.\\n4. ALERTES : tout ROUGE presente avec (a) fait + deviation vs seuil, (b) mecanisme actuariel, (c) impact bilantaire, (d) action corrective + echeance. Jamais minimiser.\\n5. CAUSALITE : [Observation] -> [Mecanisme actuariel] -> [Consequence quantifiee].\\n6. INCERTITUDE : Vie +/-50bp -> impact PM. SCR +/-10% mortalite. ALM duration gap -> stress +/-200bp. EV +/-25bp rendement -> deltaVIF. IAS 19 deltaDBo/deltataux = D_mod x DBO (§83).\\n7. POSTURE : assertif, precis, independant. "La Fonction Actuarielle recommande..." + echeance.\\n8. INTERDICTIONS : pas de generique sans chiffre ; pas d\'avis FAVORABLE si SCR<100%, PM<0 non justifiee, ecart PM/TP>30%, backtesting ROUGE >=2 params, underwater>30%.\\n9. COHERENCE INTER-SECTIONS : chiffres coherents entre §2 et §3, DBO §6 coherente avec PM §2.\\n\\nSTRUCTURE OBLIGATOIRE -- 9 SECTIONS\\n\\n§1 -- CONTEXTE, PERIMETRE ET GOUVERNANCE ACTUARIELLE\\nPerimetre exact, date arrete, encours PM, contrats, actuaire designe, calendrier ACPR.\\n\\n§2 -- PROVISIONS TECHNIQUES -- RECONCILIATION TRICOLONNE PM / TP S2 / IFRS 17\\nPM sociale (Art. R331-1). TP S2 (Art. 77). IFRS 17 LRC = FCF + CSM ou LC. Decomposer ecart PM->TP S2 : effet taux, mortalite, frais. Alerte si ecart PM/TP > 20%.\\n\\n§3 -- SCR VIE ET COUVERTURE EN CAPITAL (ART. 138-145)\\n7 sous-modules avec chocs reglementaires. Module dominant. Diversification matrice EIOPA 7x7. Ratio SCR (seuil 100%, cible ORSA 150%).\\n\\n§4 -- RISQUE TMG ET GESTION DU PORTEFEUILLE UNDERWATER (ART. 142)\\nPar tranche TMG : spread vs rendement, statut, cout annuel, PPB allouee, horizon epuisement. SCR_taux = PM x (D_mod x Dr + 1/2 x convexite x Dr2). Stress +/-200bp. Leviers d\'action.\\n\\n§5 -- PROJECTION ALM ET ADEQUATION ACTIF-PASSIF\\nProjection pluriannuelle. Ratio couverture. Duration gap. DeltaSurplus = -(D_A x A - D_P x PM) x Dtaux. Rachats. Recommandations ORSA.\\n\\n§6 -- ENGAGEMENTS EP-RE -- IAS 19, SOLVABILITE 2 ET OPTIMISATION PB\\nDBO + SC+IC (§67). iBoxx AA (§83) +/-25bp. Gains/pertes OCI. Sensibilites individuelles Somme poids_i x d_i. Couverture Art. R342-14. Optimisation PB : Art. L132-29 (85%/90%), PPB vs C2023-10.\\n\\n§7 -- EMBEDDED VALUE ET MESURE DE CREATION DE VALEUR\\nANE = Actifs - TP S2. VIF = VA profits - VA CoC (6%). VNB (MCEV §18). DeltaEV = VNB + unwinding + gains experience. Ratio VNB/primes vs benchmark (8-15%).\\n\\n§8 -- BACK-TESTING ET ROBUSTESSE DES HYPOTHESES\\nEcarts experience par parametre. DeltaDBO_p = DBO x (sensibilite_p/10000) x ecart_bps (IAS 19.145). Gain/perte actuarielle et impact fonds propres IFRS.\\n\\n§9 -- CONCLUSION, AVIS FORMEL ET RECOMMANDATIONS\\nAvis : FAVORABLE / FAVORABLE AVEC RESERVE / DEFAVORABLE. Conditions FAVORABLE : SCR>=150%, ecart PM/TP<20%, underwater<20%, backtesting <=1 ROUGE, EV positive. Recommandations par priorite + echeance. Soutenabilite 5 ans.'


def _f(v, dec=0) -> str:
    try:
        n = float(v or 0)
        s = f"{n:,.{dec}f}".replace(",", "\u202f")
        return s + "\u202f€"
    except Exception:
        return str(v) if v is not None else "—"

def _pct(v, dec=1) -> str:
    try:
        return f"{float(v or 0):.{dec}f}\u202f%"
    except Exception:
        return "—"

def _s(v) -> str:
    return str(v) if v is not None else "—"

def _clean(txt) -> str:
    if not txt:
        return ""
    txt = re.sub(r"<[^>]+>", " ", str(txt))
    txt = re.sub(r"\s+", " ", txt)
    return txt.strip()

def _statut_cls(s: str) -> str:
    s = (s or "").upper()
    if "ROUGE" in s or "FAVO" not in s and "DEF" in s:
        return "rouge"
    if "AMBRE" in s or "RESERVE" in s:
        return "ambre"
    return "vert"

def _g(d, *keys, default=None):
    v = d or {}
    for k in keys:
        if isinstance(v, dict):
            v = v.get(k, default)
        else:
            return default
    return v if v is not None else default

def _css() -> str:
    # CSS premium identique palette A7 — adapte direction Vie
    return CSS_VIE

CSS_VIE = '''
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,700;0,900;1,400&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');
:root {
  --navy:#0B1E3D; --navy-mid:#132844; --navy-light:#1E3A5F;
  --gold:#C9A84C; --gold-light:#E2C97E; --gold-pale:rgba(201,168,76,0.12);
  --rouge:#C0392B; --rouge-pale:rgba(192,57,43,0.08);
  --orange:#E67E22; --vert:#1E8449;
  --slate:#8A9AB0; --slate-light:#B8C5D3;
  --bg:#F5F7FA; --white:#FFFFFF; --text:#1C2B3A; --text-mid:#3D5166; --border:#DDE4EE;
}
* { box-sizing:border-box; margin:0; padding:0; }
body { font-family:'Inter',sans-serif; font-size:10pt; color:var(--text); background:var(--bg); line-height:1.7; }
.rapport-container { max-width:960px; margin:0 auto; background:var(--white); box-shadow:0 4px 40px rgba(11,30,61,0.14); }
.page-garde { position:relative; min-height:720px; display:flex; flex-direction:column; overflow:hidden; background:var(--navy); page-break-after:always; }
.garde-bg { position:absolute; inset:0; overflow:hidden; pointer-events:none; }
.garde-bg::before { content:''; position:absolute; top:-120px; right:-80px; width:560px; height:560px; border:1px solid rgba(201,168,76,0.12); border-radius:50%; }
.garde-bg::after  { content:''; position:absolute; top:-60px; right:-20px; width:380px; height:380px; border:1px solid rgba(201,168,76,0.20); border-radius:50%; }
.garde-dots { position:absolute; top:0; right:0; width:340px; height:100%; background-image:radial-gradient(circle,rgba(201,168,76,0.25) 1px,transparent 1px); background-size:22px 22px; mask-image:linear-gradient(to left,rgba(0,0,0,0.7),transparent); }
.garde-diagonal { position:absolute; inset:0; background:linear-gradient(135deg,transparent 62%,rgba(201,168,76,0.06) 62%); }
.garde-accent-line { position:absolute; left:52px; top:0; bottom:0; width:2px; background:linear-gradient(to bottom,transparent,var(--gold) 20%,var(--gold) 80%,transparent); opacity:0.5; }
.garde-inner { position:relative; z-index:2; display:flex; flex-direction:column; height:100%; }
.garde-header { display:flex; justify-content:space-between; align-items:flex-start; padding:36px 56px 0; }
.garde-logo-wrap img { height:52px; }
.garde-badges { display:flex; flex-direction:column; align-items:flex-end; gap:8px; }
.badge-confidentiel { font-size:7.5pt; font-weight:700; color:var(--rouge); letter-spacing:2.5px; border:1.5px solid var(--rouge); padding:5px 12px; text-transform:uppercase; }
.badge-direction { font-size:7pt; font-weight:600; color:var(--gold); letter-spacing:1.5px; border:1px solid rgba(201,168,76,0.4); padding:4px 10px; text-transform:uppercase; }
.garde-hero { flex:1; display:flex; flex-direction:column; justify-content:center; padding:48px 56px 0 80px; }
.garde-eyebrow { font-size:8pt; font-weight:600; color:var(--gold); letter-spacing:4px; text-transform:uppercase; margin-bottom:20px; }
.garde-titre { font-family:'Playfair Display',serif; font-size:48pt; font-weight:900; color:var(--white); line-height:1.05; margin-bottom:8px; }
.garde-titre em { font-style:italic; color:var(--gold); }
.garde-subtitle { font-family:'Playfair Display',serif; font-size:16pt; font-weight:400; font-style:italic; color:var(--slate-light); margin-bottom:8px; }
.garde-sousdirection { font-size:9pt; font-weight:500; color:var(--gold); letter-spacing:1.5px; margin-bottom:28px; }
.garde-sep { display:flex; align-items:center; gap:16px; margin-bottom:28px; }
.garde-sep-line { height:1px; width:80px; background:var(--gold); opacity:0.6; }
.garde-sep-diamond { width:8px; height:8px; background:var(--gold); transform:rotate(45deg); }
.garde-statut { display:inline-flex; align-items:center; gap:12px; padding:10px 24px; }
.garde-statut-rouge { background:rgba(192,57,43,0.15); border:1px solid rgba(192,57,43,0.5); }
.garde-statut-ambre { background:rgba(230,126,34,0.15); border:1px solid rgba(230,126,34,0.5); }
.garde-statut-vert  { background:rgba(30,132,73,0.15);  border:1px solid rgba(30,132,73,0.5); }
.statut-dot { width:10px; height:10px; border-radius:50%; animation:pulse 2s ease-in-out infinite; }
.statut-dot-rouge { background:var(--rouge); box-shadow:0 0 8px rgba(192,57,43,0.7); }
.statut-dot-ambre { background:var(--orange); box-shadow:0 0 8px rgba(230,126,34,0.7); }
.statut-dot-vert  { background:var(--vert);  box-shadow:0 0 8px rgba(30,132,73,0.7); }
@keyframes pulse { 0%,100%{opacity:1;transform:scale(1)} 50%{opacity:0.7;transform:scale(1.2)} }
.statut-label-rouge { font-size:8pt; font-weight:700; color:var(--rouge); letter-spacing:3px; text-transform:uppercase; }
.statut-label-ambre { font-size:8pt; font-weight:700; color:var(--orange); letter-spacing:3px; text-transform:uppercase; }
.statut-label-vert  { font-size:8pt; font-weight:700; color:var(--vert);  letter-spacing:3px; text-transform:uppercase; }
.garde-footer { border-top:1px solid rgba(201,168,76,0.25); margin:36px 56px 0; padding:24px 0 40px; }
.garde-kpis { display:grid; grid-template-columns:repeat(6,1fr); gap:0; }
.garde-kpi { padding:0 20px 0 0; border-right:1px solid rgba(255,255,255,0.08); }
.garde-kpi:last-child { border-right:none; }
.kpi-label { font-size:6.5pt; font-weight:600; color:var(--slate); text-transform:uppercase; letter-spacing:1.5px; margin-bottom:6px; }
.kpi-value { font-family:'Playfair Display',serif; font-size:11pt; font-weight:700; color:var(--white); line-height:1.2; }
.kpi-value.highlight { color:var(--gold); font-size:13pt; }
.kpi-sub { font-size:7pt; color:var(--slate); margin-top:3px; }
.rapport-body { padding:0; }
.section-header { background:linear-gradient(to right,var(--navy),var(--navy-light)); padding:22px 56px; display:flex; align-items:center; gap:20px; }
.section-num { font-family:'JetBrains Mono',monospace; font-size:9pt; font-weight:500; color:var(--gold); opacity:0.8; min-width:28px; }
.section-titre { font-family:'Playfair Display',serif; font-size:15pt; font-weight:700; color:var(--white); }
.section-body { padding:36px 56px 48px; }
.section-divider { height:1px; background:var(--border); }
.kpi-grid { display:grid; grid-template-columns:repeat(4,1fr); gap:12px; margin-bottom:32px; }
.kpi-card { background:var(--navy); border-radius:6px; padding:20px 18px; position:relative; overflow:hidden; }
.kpi-card::before { content:''; position:absolute; top:0; left:0; right:0; height:3px; background:var(--gold); }
.kpi-card-rouge::before { background:var(--rouge); }
.kpi-card-ambre::before { background:var(--orange); }
.kpi-card-vert::before  { background:var(--vert); }
.kpi-card-label { font-size:7pt; font-weight:600; color:var(--slate); text-transform:uppercase; letter-spacing:1.5px; margin-bottom:8px; }
.kpi-card-value { font-family:'Playfair Display',serif; font-size:16pt; font-weight:700; color:var(--white); line-height:1.1; margin-bottom:6px; }
.kpi-card-value-rouge { color:var(--rouge) !important; }
.kpi-card-value-ambre { color:var(--orange) !important; }
.kpi-card-value-vert  { color:#2ECC71 !important; }
.kpi-card-sub { font-size:7.5pt; color:var(--gold); }
.table-section-title { font-family:'Playfair Display',serif; font-size:10pt; font-weight:600; color:var(--gold); margin:28px 0 12px; padding-bottom:6px; border-bottom:1px solid var(--gold-pale); }
table.premium { width:100%; border-collapse:collapse; font-size:9pt; }
table.premium thead tr { background:var(--navy); }
table.premium thead th { padding:11px 14px; color:var(--white); font-weight:600; font-size:8pt; text-align:left; }
table.premium thead th.right { text-align:right; }
table.premium tbody tr:nth-child(even) { background:#F8FAFB; }
table.premium tbody tr.highlight-gold { background:var(--gold-pale); font-weight:700; }
table.premium tbody td { padding:9px 14px; border-bottom:1px solid var(--border); color:var(--text); }
table.premium tbody td.right { text-align:right; font-family:'JetBrains Mono',monospace; font-size:8.5pt; }
table.premium tbody td.label { font-weight:600; color:var(--navy); }
.badge { display:inline-flex; align-items:center; gap:4px; font-size:7.5pt; font-weight:700; padding:2px 8px; border-radius:3px; }
.badge-vert  { background:rgba(30,132,73,0.12);  color:var(--vert); }
.badge-ambre { background:rgba(230,126,34,0.12); color:var(--orange); }
.badge-rouge { background:rgba(192,57,43,0.12);  color:var(--rouge); }
.tricolonne { display:grid; grid-template-columns:repeat(3,1fr); gap:16px; margin:24px 0; }
.tricolonne-card { border-radius:8px; overflow:hidden; border:1px solid var(--border); }
.tricolonne-header { padding:14px 20px; background:var(--navy); }
.tricolonne-ref { font-size:7pt; font-weight:600; color:var(--gold); letter-spacing:1.5px; text-transform:uppercase; margin-bottom:4px; }
.tricolonne-label { font-size:9pt; font-weight:700; color:var(--white); }
.tricolonne-body { padding:18px 20px; background:#FAFBFC; }
.tricolonne-value { font-family:'Playfair Display',serif; font-size:18pt; font-weight:700; color:var(--navy); margin-bottom:6px; }
.tricolonne-delta { font-size:8pt; color:var(--slate); }
.tricolonne-delta strong { color:var(--orange); }
.scr-grid { display:grid; grid-template-columns:repeat(4,1fr); gap:10px; margin:20px 0; }
.scr-card { padding:14px 16px; border-radius:6px; background:var(--navy); position:relative; }
.scr-card::before { content:''; position:absolute; top:0; left:0; right:0; height:3px; background:var(--gold); border-radius:3px 3px 0 0; }
.scr-card-dominant::before { background:var(--rouge); }
.scr-card-label { font-size:7pt; color:var(--slate); text-transform:uppercase; letter-spacing:1px; margin-bottom:6px; }
.scr-card-value { font-family:'Playfair Display',serif; font-size:13pt; font-weight:700; color:var(--white); }
.scr-card-pct { font-size:8pt; color:var(--gold); margin-top:4px; }
.scr-card-art { font-size:7pt; color:var(--slate); margin-top:2px; font-style:italic; }
.commentaire-wrap { border:1px solid var(--border); border-radius:6px; overflow:hidden; margin-top:8px; }
.commentaire-header { background:var(--navy); padding:14px 20px; display:flex; align-items:center; justify-content:space-between; }
.commentaire-header-title { font-size:8pt; font-weight:600; color:var(--gold); letter-spacing:2px; text-transform:uppercase; }
.commentaire-ai-badge { font-size:7pt; color:var(--slate); font-family:'JetBrains Mono',monospace; }
.commentaire-body { padding:28px 32px; background:#FAFBFC; }
.comm-section-title { font-family:'Playfair Display',serif; font-size:11pt; font-weight:700; color:var(--navy); border-bottom:1.5px solid var(--gold); padding-bottom:6px; margin:24px 0 12px; }
.comm-section-title:first-child { margin-top:0; }
.comm-h4 { font-size:9.5pt; font-weight:700; color:var(--navy); margin:16px 0 8px; }
.comm-p { font-size:9.5pt; line-height:1.85; color:var(--text-mid); margin-bottom:12px; }
.comm-p strong { color:var(--navy); font-weight:700; }
.comm-ul { margin:10px 0 12px 20px; list-style:none; }
.comm-ul li { position:relative; padding-left:16px; font-size:9.5pt; line-height:1.7; color:var(--text-mid); margin-bottom:6px; }
.comm-ul li::before { content:'—'; position:absolute; left:0; color:var(--gold); font-weight:700; }
.comm-ul li strong { color:var(--navy); }
.comm-divider { border:none; border-top:1px solid var(--border); margin:20px 0; }
.comm-footer { font-size:7.5pt; color:var(--slate); font-style:italic; text-align:right; border-top:1px solid var(--border); padding-top:10px; margin-top:8px; }
.avis-box { padding:24px 28px; border-radius:8px; margin:24px 0; display:flex; align-items:flex-start; gap:20px; }
.avis-favorable   { background:rgba(30,132,73,0.08);  border:2px solid rgba(30,132,73,0.4); }
.avis-reserve     { background:rgba(230,126,34,0.08); border:2px solid rgba(230,126,34,0.4); }
.avis-defavorable { background:rgba(192,57,43,0.08);  border:2px solid rgba(192,57,43,0.5); }
.avis-icon { font-size:24pt; flex-shrink:0; }
.avis-label { font-size:7pt; font-weight:700; letter-spacing:2px; text-transform:uppercase; margin-bottom:6px; }
.avis-value { font-family:'Playfair Display',serif; font-size:16pt; font-weight:700; margin-bottom:8px; }
.avis-value-favorable   { color:var(--vert); }
.avis-value-reserve     { color:var(--orange); }
.avis-value-defavorable { color:var(--rouge); }
.avis-text { font-size:9pt; color:var(--text-mid); line-height:1.6; }
.graph-wrap { border:1px solid var(--border); border-radius:6px; overflow:hidden; margin:20px 0; }
.graph-titre { background:var(--navy-light); padding:10px 18px; font-size:8pt; font-weight:600; color:var(--gold); letter-spacing:1px; text-transform:uppercase; }
.graph-body { padding:16px; background:#FAFBFC; }
.pied-de-page { background:var(--navy); padding:20px 56px; display:flex; justify-content:space-between; align-items:center; border-top:2px solid var(--gold); }
.pied-logo img { height:26px; }
.pied-meta { text-align:right; font-size:7.5pt; color:var(--slate); line-height:1.7; }
.pied-meta .confidentiel-footer { color:var(--rouge); font-weight:700; letter-spacing:1px; }
@media print { body{background:white} .rapport-container{box-shadow:none;max-width:none;margin:0} .page-garde{page-break-after:always;min-height:100vh} @page{margin:0} }
</style>
'''


def _md_to_html(texte: str) -> str:
    """Convertit la narration markdown actuarielle en HTML premium."""
    if not texte:
        return ""
    html = []
    for l in texte.split("\n"):
        l = l.rstrip()
        m = re.match(r"^§\s*(\d+)\s*[-—–]\s*(.+)$", l)
        if m:
            html.append(f'<div class="comm-section-title">§{m.group(1)} — {m.group(2).strip()}</div>')
            continue
        if l.startswith("###"):
            html.append(f'<div class="comm-h4">{l.lstrip("#").strip()}</div>')
            continue
        if not l.strip():
            html.append('<hr class="comm-divider">')
            continue
        if l.strip().startswith("- "):
            item = l.strip()[2:]
            item = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", item)
            html.append(f'<ul class="comm-ul"><li>{item}</li></ul>')
            continue
        line = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", l)
        html.append(f'<p class="comm-p">{line}</p>')
    return "\n".join(html)


def _construire_contexte_vie(
    result_v3, result_v5, result_v7, result_v8, result_v9,
    result_v6, result_rvie1, result_rvie2,
    result_ep1, result_ep3, result_ep4, result_ep6, result_ep7,
    ref_client, arrete,
) -> str:
    r3=result_v3 or {}; r5=result_v5 or {}; r6=result_v6 or {}
    r7=result_v7 or {}; r8=result_v8 or {}; r9=result_v9 or {}
    rr1=result_rvie1 or {}; rr2=result_rvie2 or {}
    ep1=result_ep1 or {}; ep3=result_ep3 or {}; ep4=result_ep4 or {}
    ep6=result_ep6 or {}; ep7=result_ep7 or {}

    pm      = float(r3.get("pm_prospective", 0) or 0)
    be      = float(rr1.get("be_vie", pm) or pm)
    tp      = float(r5.get("tp_vie", be * 1.05) or be * 1.05)
    tp_be   = float(r5.get("ratio_tp_be", tp/be if be else 0) or 0)
    scr     = float(rr1.get("scr_vie_total", 0) or 0)
    ratio_s = float(r5.get("ratio_scr_pct", 0) or 0)
    ev      = float(r9.get("embedded_value", 0) or 0)
    ane     = float(r9.get("ane", 0) or 0)
    vif     = float(r9.get("vif", 0) or 0)
    vnb     = float(r9.get("vnb", 0) or 0)
    dbo     = float(_g(ep1, "ias19", "dbo_total", default=0) or 0)
    pm_ep   = float(_g(ep3, "provisions", "pm_encours", default=0) or 0)
    ppb_ep  = float(_g(ep3, "provisions", "ppb", default=0) or 0)
    scr_mod = rr1.get("scr_par_module", {}) or {}
    dom_mod = max(scr_mod, key=lambda k: float(scr_mod[k] or 0)) if scr_mod else "non disponible"

    lignes = [
        f"RAPPORT CONSOLIDE VIE & EP-RE -- {ref_client.upper()} -- Arrete {arrete}",
        f"Date generation : {datetime.now().strftime('%d/%m/%Y')}",
        "",
        "=== VIE PURE -- PROVISIONS ===",
        f"PM prospective (Art. R331-1)    : {_f(pm)}",
        f"Best Estimate S2 (Art. 77)      : {_f(be)}",
        f"TP Solvabilite 2 (BE + RM)      : {_f(tp)}",
        f"Ratio TP/BE                     : {tp_be:.4f}  (seuil conforme : 1.0-1.5)",
        f"CSM estime (IFRS 17)            : {_f(float(r8.get('csm', 0) or 0))}",
        f"be_ifrs17_mode                  : {r8.get('be_ifrs17_mode', 'approximation S2')}",
        "",
        "=== SCR VIE (ART. 138-145) ===",
        f"SCR Vie total                   : {_f(scr)}",
        f"Ratio SCR / fonds propres       : {_pct(ratio_s)}  (seuil 100%, cible ORSA 150%)",
        f"Module dominant                 : {dom_mod}",
        f"SCR mortalite  (Art. 137 +15%)  : {_f(float(scr_mod.get('mortalite', 0) or 0))}",
        f"SCR longevite  (Art. 138 -20%)  : {_f(float(scr_mod.get('longevite', 0) or 0))}",
        f"SCR rachat     (Art. 142 max3)  : {_f(float(scr_mod.get('rachat', 0) or 0))}",
        f"SCR frais      (Art. 143 +10%)  : {_f(float(scr_mod.get('frais', 0) or 0))}",
        f"SCR invalidite (Art. 139 +35%)  : {_f(float(scr_mod.get('invalidite', 0) or 0))}",
        f"SCR revision   (Art. 144 +3%)   : {_f(float(scr_mod.get('revision', 0) or 0))}",
        f"SCR catastrophe (Art. 145 1.5%) : {_f(float(scr_mod.get('catastrophe', 0) or 0))}",
        f"Effet diversification           : {_f(float(rr1.get('effet_diversification', 0) or 0))} economise",
        "",
        "=== TMG ET PORTEFEUILLE UNDERWATER (ART. 142) ===",
        f"Tranches TMG detectees          : {r7.get('nb_tranches', 0)}",
        f"Part portefeuille underwater    : {_pct(float(r7.get('pct_underwater', 0) or 0))}",
        f"Cout annuel TMG garanti         : {_f(float(r7.get('cout_annuel_garanti', 0) or 0))}",
        f"PPB disponible                  : {_f(float(r7.get('ppb_disponible', 0) or 0))}",
        f"Horizon epuisement PPB          : {r7.get('horizon_epuisement_ppb', 'non disponible')} ans",
        f"SCR taux (avec convexite)       : {_f(float(r7.get('scr_taux', 0) or 0))}",
        f"Statut TMG                      : {r7.get('statut_rag', 'non disponible')}",
        "",
        "=== ALM DYNAMIQUE ===",
        f"Duration gap (D_passif-D_actif) : {float(r6.get('duration_gap', 0) or 0):.2f} ans",
        f"Ratio couverture initial        : {_pct(float(r6.get('ratio_couverture_initial', 0) or 0))}",
        f"Impact stress +200bp sur surplus: {_f(float(r6.get('impact_stress_hausse', 0) or 0))}",
        f"Impact stress -200bp sur surplus: {_f(float(r6.get('impact_stress_baisse', 0) or 0))}",
        f"Statut ALM                      : {r6.get('statut_rag', 'non disponible')}",
        "",
        "=== EP-RE -- IAS 19 ET PROVISIONS ===",
        f"DBO totale (IAS 19 PUC §67)     : {_f(dbo)}",
        f"Service Cost exercice           : {_f(float(_g(ep1, 'ias19', 'service_cost', default=0) or 0))}",
        f"Interest Cost exercice          : {_f(float(_g(ep1, 'ias19', 'interest_cost', default=0) or 0))}",
        f"Taux actualisation iBoxx AA §83 : {_pct(float(_g(ep1, 'ias19', 'taux_actu', default=0) or 0))}",
        f"Sensibilite DBO a -100bp        : {_pct(float(_g(ep1, 'ias19', 'sensibilite_dbo_taux_pct', default=0) or 0))}",
        f"PM epargne-retraite             : {_f(pm_ep)}",
        f"PPB EP-RE                       : {_f(ppb_ep)}",
        f"Ratio PPB/PM                    : {_pct(ppb_ep/pm_ep*100 if pm_ep else 0)} (plafond C2023-10 : 10%)",
        f"Ratio couverture actifs/PM      : {_pct(float(ep4.get('ratio_base', 0) or 0))}  (seuil 100%)",
        "",
        "=== BACK-TESTING EP-RE (IAS 19.145) ===",
        f"Statut backtesting              : {ep6.get('statut_rag', 'non disponible')}",
        f"Impact DBO total                : {_f(float(ep6.get('impact_dbo_total', 0) or 0))}",
        "",
        "=== OPTIMISATION PB (ART. L132-29, C2023-10) ===",
        f"Taux PB optimal retenu          : {_pct(float(ep7.get('taux_pb_optimal', 0) or 0))}",
        f"PPB finale apres optimisation   : {_f(float(ep7.get('ppb_finale', 0) or 0))}",
        f"Ratio SCR post-distribution     : {_pct(float(ep7.get('ratio_scr_post_pb', 0) or 0))}",
        "",
        "=== EMBEDDED VALUE (MCEV §18) ===",
        f"ANE (Actifs marche - TP S2)     : {_f(ane)}",
        f"VIF (VA profits - VA CoC 6%)    : {_f(vif)}",
        f"EV totale                       : {_f(ev)}",
        f"VNB (nouvelles affaires)        : {_f(vnb)}",
        f"Ratio VNB/primes                : {_pct(float(r9.get('ratio_vnb_primes', 0) or 0))} (benchmark 8-15%)",
        f"Gain experience                 : {_f(float(r9.get('gain_experience', 0) or 0))}",
        "",
        "=== STATUTS GLOBAUX ===",
        f"Vie pure (V5)                   : {r5.get('statut_rag', 'non disponible')}",
        f"SCR (R-VIE1)                    : {rr1.get('statut_rag', 'non disponible')}",
        f"RSR/SFCR (R-VIE2)               : {rr2.get('statut_rag', 'non disponible')}",
        f"ALM (V6)                        : {r6.get('statut_rag', 'non disponible')}",
        f"TMG (V7)                        : {r7.get('statut_rag', 'non disponible')}",
        f"EV (V9)                         : {r9.get('statut_rag', 'non disponible')}",
        f"IAS 19 (EP1)                    : {ep1.get('statut_rag', 'non disponible')}",
        f"Stress EP-RE (EP4)              : {ep4.get('statut_rag', 'non disponible')}",
        f"Backtesting (EP6)               : {ep6.get('statut_rag', 'non disponible')}",
        f"Optimisation PB (EP7)           : {ep7.get('statut_rag', 'non disponible')}",
        "",
        "Redige le commentaire actuariel complet en 9 sections (§1 a §9) selon les regles du prompt.",
    ]
    return "\n".join(lignes)


def _narration_claude_api_vie(contexte: str) -> str:
    resp = frontiere_llm.appeler(
        modele=frontiere_llm.MODELE_ETABLI,
        max_tokens=12000,
        systeme=SYSTEM_PROMPT_VIE,
        messages=[{"role": "user", "content": contexte}],
        cle=frontiere_llm.cle_api_ou_secrets(),
    )
    return frontiere_llm.texte_du_premier_bloc(resp)


def _generer_narration_vie(contexte: str, commentaire: str) -> Tuple[str, str]:
    try:
        txt = _narration_claude_api_vie(contexte)
        if txt:
            return txt, "claude_api"
    except Exception as e:
        logger.warning(f"Claude API : {e}")
    if commentaire and commentaire.strip():
        return _clean(commentaire), "manuel"
    return "", "aucune"


def _kpi_card(label, value, sub="", cls="") -> str:
    cls_val  = f" kpi-card-value-{cls}" if cls else ""
    cls_card = f" kpi-card-{cls}" if cls else ""
    return (
        f'<div class="kpi-card{cls_card}">'''
        f'<div class="kpi-card-label">{label}</div>'''
        f'<div class="kpi-card-value{cls_val}">{value}</div>'''
        f'<div class="kpi-card-sub">{sub}</div></div>'
    )

def _scr_card(label, value, pct, art, dominant=False) -> str:
    cls = " scr-card-dominant" if dominant else ""
    return (
        f'<div class="scr-card{cls}">'''
        f'<div class="scr-card-label">{label}</div>'''
        f'<div class="scr-card-value">{value}</div>'''
        f'<div class="scr-card-pct">{pct}</div>'''
        f'<div class="scr-card-art">{art}</div></div>'
    )

def _tricolonne_card(ref, label, value, delta="") -> str:
    return (
        f'<div class="tricolonne-card">'''
        f'<div class="tricolonne-header">'''
        f'<div class="tricolonne-ref">{ref}</div>'''
        f'<div class="tricolonne-label">{label}</div></div>'''
        f'<div class="tricolonne-body">'''
        f'<div class="tricolonne-value">{value}</div>'''
        f'<div class="tricolonne-delta">{delta}</div>'''
        f'</div></div>'
    )

def _wrap_graph(html_g: str, titre: str) -> str:
    if not html_g:
        return ""
    return (
        f'<div class="graph-wrap">'''
        f'<div class="graph-titre">{titre}</div>'''
        f'<div class="graph-body">{html_g}</div></div>'
    )

def _avis_box(avis: str, texte: str) -> str:
    a = avis.upper()
    if "DEFAV" in a or ("FAVOR" not in a and "DEF" in a):
        cls, icon, vcls = "avis-defavorable", "❌", "avis-value-defavorable"
    elif "RESERVE" in a or "RESERV" in a:
        cls, icon, vcls = "avis-reserve",    "⚠️", "avis-value-reserve"
    else:
        cls, icon, vcls = "avis-favorable",  "✅", "avis-value-favorable"
    return (
        f'<div class="avis-box {cls}">'''
        f'<div class="avis-icon">{icon}</div>'''
        f'<div>'''
        f'<div class="avis-label">Avis de la Fonction Actuarielle — Art. 48 §g Directive S2</div>'''
        f'<div class="avis-value {vcls}">{avis}</div>'''
        f'<div class="avis-text">{texte}</div>'''
        f'</div></div>'
    )


def export_html(
    result_v3=None, result_v5=None, result_v6=None, result_v7=None,
    result_v8=None, result_v9=None, result_rvie1=None, result_rvie2=None,
    result_ep1=None, result_ep3=None, result_ep4=None,
    result_ep6=None, result_ep7=None,
    commentaire="", ref_client="", arrete="",
    audit_id="", graphiques=None,
    actuaire_nom="", actuaire_numero_ia="",
) -> str:
    """Genere le rapport HTML consolide Vie & EP-RE — niveau premium A7."""
    try:
        r3=result_v3 or {}; r5=result_v5 or {}; r6=result_v6 or {}
        r7=result_v7 or {}; r8=result_v8 or {}; r9=result_v9 or {}
        rr1=result_rvie1 or {}; rr2=result_rvie2 or {}
        ep1=result_ep1 or {}; ep3=result_ep3 or {}; ep4=result_ep4 or {}
        ep6=result_ep6 or {}; ep7=result_ep7 or {}

        dt  = datetime.now().strftime("%d/%m/%Y %H:%M")
        arr = arrete or datetime.now().strftime("%d/%m/%Y")
        cli = ref_client or "A renseigner"

        pm      = float(r3.get("pm_prospective", 0) or 0)
        be      = float(rr1.get("be_vie", pm) or pm)
        tp      = float(r5.get("tp_vie", be * 1.05) or be * 1.05)
        tp_be   = float(r5.get("ratio_tp_be", tp/be if be else 1) or 1)
        scr     = float(rr1.get("scr_vie_total", 0) or 0)
        ratio_s = float(r5.get("ratio_scr_pct", 0) or 0)
        ev      = float(r9.get("embedded_value", 0) or 0)
        ane     = float(r9.get("ane", 0) or 0)
        vif     = float(r9.get("vif", 0) or 0)
        vnb     = float(r9.get("vnb", 0) or 0)
        dbo     = float(_g(ep1, "ias19", "dbo_total", default=0) or 0)
        pm_ep   = float(_g(ep3, "provisions", "pm_encours", default=0) or 0)
        ppb_ep  = float(_g(ep3, "provisions", "ppb", default=0) or 0)
        ratio_b = float(ep4.get("ratio_base", 0) or 0)
        avis_pa = r5.get("avis_pa", "FAVORABLE AVEC RESERVE")
        csm     = float(r8.get("csm", 0) or 0)
        lrc     = tp + csm

        statuts = [r5.get("statut_rag","AMBRE"), rr1.get("statut_rag","AMBRE"),
                   r7.get("statut_rag","AMBRE"), ep4.get("statut_rag","AMBRE")]
        if "ROUGE" in statuts:
            sg = "ROUGE"
        elif "AMBRE" in statuts:
            sg = "AMBRE"
        else:
            sg = "VERT"
        sc = _statut_cls(sg)
        dot_cls    = f"statut-dot-{sc}"
        statut_cls_str = f"garde-statut-{sc}"
        label_cls  = f"statut-label-{sc}"
        label_txt  = {"rouge":"Vigilance requise","ambre":"Reserves signalees","vert":"Favorable"}[sc]

        # Graphiques
        ghtml = {}
        if graphiques and PLOTLY_OK:
            for nom, fig_or_html in graphiques.items():
                try:
                    if isinstance(fig_or_html, str):
                        ghtml[nom] = fig_or_html
                    else:
                        ghtml[nom] = pio.to_html(fig_or_html, full_html=False,
                                                  include_plotlyjs=False,
                                                  config={"displayModeBar": False})
                except Exception:
                    pass

        # Narration
        contexte = _construire_contexte_vie(
            r3,r5,r7,r8,r9,r6,rr1,rr2,ep1,ep3,ep4,ep6,ep7,cli,arr)
        narration, source = _generer_narration_vie(contexte, commentaire)

        # SCR modules
        scr_mod = rr1.get("scr_par_module", {}) or {}
        dom_mod = max(scr_mod, key=lambda k: float(scr_mod[k] or 0)) if scr_mod else "rachat"

        # KPI page de garde
        garde_kpis = (
            f'<div class="garde-kpi"><div class="kpi-label">PM Vie pure</div>'''
            f'<div class="kpi-value highlight">{_f(pm)}</div>'''
            f'<div class="kpi-sub">Art. R331-1</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Best Estimate S2</div>'''
            f'<div class="kpi-value">{_f(be)}</div>'''
            f'<div class="kpi-sub">Art. 77 S2</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Ratio SCR</div>'''
            f'<div class="kpi-value">{_pct(ratio_s)}</div>'''
            f'<div class="kpi-sub">Cible ORSA ≥ 150%</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">DBO IAS 19</div>'''
            f'<div class="kpi-value">{_f(dbo)}</div>'''
            f'<div class="kpi-sub">§67 IAS 19 PUC</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Embedded Value</div>'''
            f'<div class="kpi-value">{_f(ev)}</div>'''
            f'<div class="kpi-sub">ANE + VIF (MCEV §18)</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">VNB</div>'''
            f'<div class="kpi-value">{_f(vnb)}</div>'''
            f'<div class="kpi-sub">Nouvelles affaires</div></div>'
        )

        # KPI grid synthese
        kpi_grid = '<div class="kpi-grid">'
        kpi_grid += _kpi_card("PM Vie Pure", _f(pm), "Art. R331-1")
        kpi_grid += _kpi_card("Best Estimate S2", _f(be), "Art. 77 S2")
        kpi_grid += _kpi_card("Ratio TP/BE", f"{tp_be:.4f}", "Seuil : 1.0-1.5",
                               cls="rouge" if not (1.0 <= tp_be <= 1.5) else "vert")
        kpi_grid += _kpi_card("Ratio SCR", _pct(ratio_s), "Cible ORSA ≥ 150%",
                               cls="rouge" if ratio_s < 100 else "ambre" if ratio_s < 150 else "vert")
        kpi_grid += _kpi_card("DBO IAS 19", _f(dbo), "§67 PUC IAS 19")
        kpi_grid += _kpi_card("EV totale", _f(ev), "ANE + VIF")
        kpi_grid += _kpi_card("VIF", _f(vif), "VA profits – VA CoC 6%")
        kpi_grid += _kpi_card("VNB", _f(vnb), "Nouvelles affaires MCEV §18")
        kpi_grid += '</div>'

        # Tricolonne
        tricolonne = (
            '<div class="tricolonne">'
            + _tricolonne_card("Art. R331-1 Code des assurances", "PM Sociale", _f(pm),
                               f"Δ vs TP S2 : {_pct((tp-pm)/pm*100 if pm else 0)}")
            + _tricolonne_card("Art. 77 S2 — BE + Risk Margin", "TP Solvabilité 2", _f(tp),
                               f"Ratio TP/BE : {tp_be:.4f} — {'Conforme' if 1.0<=tp_be<=1.5 else 'Hors seuil'}")
            + _tricolonne_card("IFRS 17 §32-46 — FCF + CSM/LC", "LRC IFRS 17", _f(lrc),
                               f"CSM : {_f(csm)} — {r8.get('be_ifrs17_mode', 'approx. S2')}")
            + '</div>'
        )

        # SCR modules
        mods = {
            "mortalite":  ("Mortalité",  "Art. 137 +15%"),
            "longevite":  ("Longévité", "Art. 138 -20%"),
            "rachat":     ("Rachat",    "Art. 142 max3"),
            "frais":      ("Frais",     "Art. 143 +10%"),
            "invalidite": ("Invalidité","Art. 139 +35%"),
            "revision":   ("Révision",  "Art. 144 +3%"),
            "catastrophe":("Catastro.", "Art. 145 1.5‰"),
        }
        scr_cards = '<div class="scr-grid">'
        for mkey, (mlabel, mart) in mods.items():
            mv = float(scr_mod.get(mkey, 0) or 0)
            mpct = mv/scr*100 if scr else 0
            scr_cards += _scr_card(mlabel, _f(mv), _pct(mpct), mart, dominant=(mkey==dom_mod))
        scr_cards += '</div>'

        # Avis formel
        avis_html = _avis_box(avis_pa, r5.get("conseil_global","Rapport produit par ActuarIA Direction Vie & EP-RE."))

        # Narration HTML
        narration_html = _md_to_html(narration) if narration else (
            '<p class="comm-p" style="color:var(--slate);font-style:italic;">'''
            "Narration non disponible — alimenter les agents et configurer ANTHROPIC_API_KEY.</p>"
        )
        ai_badge = ("❆ Narration générée par ActuarIA Intelligence"
                    if source == "claude_api" else "❆ Commentaire actuariel manuel")

        # Assemblage HTML
        H = []
        H.append('<!DOCTYPE html>')
        H.append('<html lang="fr">')
        H.append('<head><meta charset="UTF-8">')
        H.append(f'<title>Rapport Consolide Vie & EP-RE — {cli} — {arr}</title>')
        H.append('<script src="https://cdn.jsdelivr.net/npm/plotly.js-dist@2.26.0/plotly.min.js"></script>')
        H.append(_css())
        H.append('</head><body><div class="rapport-container">')

        # Page de garde
        H.append('<div class="page-garde">')
        H.append('<div class="garde-bg"><div class="garde-dots"></div>')
        H.append('<div class="garde-diagonal"></div><div class="garde-accent-line"></div></div>')
        H.append('<div class="garde-inner">')
        H.append(f'<div class="garde-header"><div class="garde-logo-wrap"><img src="{LOGO_URI}" alt="ActuarIA"/></div>')
        H.append('<div class="garde-badges"><span class="badge-confidentiel">⬛ Confidentiel</span>')
        H.append('<span class="badge-direction">Direction Vie &amp; EP-RE</span></div></div>')
        H.append('<div class="garde-hero">')
        H.append('<div class="garde-eyebrow">Rapport Consolide Actuariel</div>')
        H.append('<div class="garde-titre">Vie &amp; <em>EP-RE</em></div>')
        H.append(f'<div class="garde-subtitle">Arrêté au {arr}</div>')
        H.append(f'<div class="garde-sousdirection">Vie Pure (V1–V9) · Épargne-Retraite (EP1–EP7) · {cli}</div>')
        H.append('<div class="garde-sep"><div class="garde-sep-line"></div>')
        H.append('<div class="garde-sep-diamond"></div><div class="garde-sep-line"></div></div>')
        H.append(f'<div class="garde-statut {statut_cls_str}"><div class="statut-dot {dot_cls}"></div>')
        H.append(f'<div class="{label_cls}">{label_txt}</div></div>')
        H.append('</div>')
        H.append(f'<div class="garde-footer"><div class="garde-kpis">{garde_kpis}</div></div>')
        H.append('</div></div>')  # fin garde-inner, page-garde

        # Corps
        H.append('<div class="rapport-body">')

        # §1 Synthese
        H.append('<div class="section-header"><span class="section-num">01</span>')
        H.append('<span class="section-titre">Synthèse exécutive</span></div>')
        H.append('<div class="section-body">')
        H.append(kpi_grid)
        H.append(_wrap_graph(ghtml.get("graphiques_reconc",""), "Réconciliation PM / TP S2 / IFRS 17"))
        H.append(avis_html)
        H.append('</div><div class="section-divider"></div>')

        # §2 Provisions
        H.append('<div class="section-header"><span class="section-num">02</span>')
        H.append('<span class="section-titre">Provisions Techniques — Réconciliation PM / TP S2 / IFRS 17</span></div>')
        H.append(f'<div class="section-body">{tricolonne}')
        H.append(_wrap_graph(ghtml.get("graphiques_tricolonne",""), "Décomposition tricolonne"))
        H.append('</div><div class="section-divider"></div>')

        # §3 SCR
        H.append('<div class="section-header"><span class="section-num">03</span>')
        H.append('<span class="section-titre">SCR Vie — Couverture en Capital (Art. 138-145)</span></div>')
        H.append(f'<div class="section-body"><div class="table-section-title">')
        H.append(f'SCR total : {_f(scr)} · Ratio : {_pct(ratio_s)} · Module dominant : {dom_mod.title()}</div>')
        H.append(scr_cards)
        H.append(_wrap_graph(ghtml.get("graphiques_scr",""), "Décomposition SCR Vie — Matrice EIOPA 7×7"))
        H.append('</div><div class="section-divider"></div>')

        # §4 TMG
        H.append('<div class="section-header"><span class="section-num">04</span>')
        H.append('<span class="section-titre">Risque TMG — Portefeuille Underwater (Art. 142)</span></div>')
        H.append(f'<div class="section-body"><div class="table-section-title">')
        H.append(f'SCR taux (convexité 2ᵉ ordre) : {_f(float(r7.get("scr_taux",0) or 0))}</div>')
        H.append(_wrap_graph(ghtml.get("graphiques_tmg",""), "Tranches TMG — Spread vs Rendement actifs"))
        H.append('</div><div class="section-divider"></div>')

        # §5 ALM
        H.append('<div class="section-header"><span class="section-num">05</span>')
        H.append('<span class="section-titre">Projection ALM — Adéquation Actif-Passif</span></div>')
        H.append(f'<div class="section-body"><div class="table-section-title">')
        H.append(f'Duration gap : {float(r6.get("duration_gap",0) or 0):.2f} ans')
        H.append(f' · Impact ±200bp : {_f(abs(float(r6.get("impact_stress_hausse",0) or 0)))}</div>')
        H.append(_wrap_graph(ghtml.get("graphiques_alm",""), "Projection ALM — Ratio couverture annuel"))
        H.append('</div><div class="section-divider"></div>')

        # §6 EP-RE
        H.append('<div class="section-header"><span class="section-num">06</span>')
        H.append('<span class="section-titre">Engagements EP-RE — IAS 19 &amp; Optimisation PB</span></div>')
        H.append('<div class="section-body"><div class="kpi-grid">')
        H.append(_kpi_card("DBO IAS 19", _f(dbo), "§67 PUC IAS 19"))
        H.append(_kpi_card("PM EP-RE", _f(pm_ep), "Art. R342-14"))
        H.append(_kpi_card("PPB EP-RE", _f(ppb_ep),
                            f"Ratio {_pct(ppb_ep/pm_ep*100 if pm_ep else 0)} · Plafond 10%"))
        H.append(_kpi_card("Couverture actifs", _pct(ratio_b), "Seuil ≥ 100%",
                            cls="rouge" if ratio_b < 100 else "vert"))
        H.append('</div>')
        H.append(_wrap_graph(ghtml.get("graphiques_ias19",""), "IAS 19 — DBO, Service Cost, Sensibilités"))
        H.append(_wrap_graph(ghtml.get("graphiques_ep7",""), "Optimisation PB — Projection PPB pluriannuelle"))
        H.append('</div><div class="section-divider"></div>')

        # §7 EV
        H.append('<div class="section-header"><span class="section-num">07</span>')
        H.append('<span class="section-titre">Embedded Value — Mesure de Création de Valeur</span></div>')
        H.append('<div class="section-body"><div class="kpi-grid">')
        H.append(_kpi_card("ANE", _f(ane), "Actifs marché − TP S2"))
        H.append(_kpi_card("VIF", _f(vif), "VA profits − VA CoC 6%"))
        H.append(_kpi_card("EV totale", _f(ev), "ANE + VIF (MCEV §18)"))
        H.append(_kpi_card("VNB", _f(vnb),
                            f"Ratio {_pct(float(r9.get('ratio_vnb_primes',0) or 0))} · Bench. 8-15%"))
        H.append('</div>')
        H.append(_wrap_graph(ghtml.get("graphiques_ev",""), "Décomposition ΔEV — VNB / Unwinding / Gains expérience"))
        H.append('</div><div class="section-divider"></div>')

        # §8 Back-testing
        H.append('<div class="section-header"><span class="section-num">08</span>')
        H.append('<span class="section-titre">Back-Testing — Robustesse des Hypothèses (IAS 19.145)</span></div>')
        H.append(f'<div class="section-body">')
        H.append(_wrap_graph(ghtml.get("graphiques_bt",""), "Écarts d'expérience — Impact DBO par paramètre"))
        H.append('</div><div class="section-divider"></div>')

        # §9 Commentaire actuariel
        H.append('<div class="section-header"><span class="section-num">09</span>')
        H.append('<span class="section-titre">Commentaire Actuariel — Avis de la Fonction Actuarielle</span></div>')
        H.append('<div class="section-body">')
        H.append('<div class="commentaire-wrap">')
        H.append('<div class="commentaire-header">')
        H.append('<div class="commentaire-header-title">Analyse actuarielle consolidée Vie &amp; EP-RE</div>')
        H.append(f'<div class="commentaire-ai-badge">{ai_badge}</div></div>')
        H.append('<div class="commentaire-body">')
        H.append(narration_html)
        H.append(f'<div class="comm-footer">Rédigé le {dt} · Actuaire : {actuaire_nom or "A renseigner"} · N° IA : {actuaire_numero_ia or "—"} · CONFIDENTIEL</div>')
        H.append('</div></div></div>')

        H.append('</div>')  # fin rapport-body

        # Pied de page
        H.append('<div class="pied-de-page">')
        H.append(f'<div class="pied-logo"><img src="{LOGO_URI}" alt="ActuarIA"/></div>')
        H.append(f'<div class="pied-meta"><div>{cli} · Arrêté {arr}</div>')
        H.append(f'<div>Audit ID : {audit_id or "—"} · Généré le {dt}</div>')
        H.append('<div class="confidentiel-footer">⬛ CONFIDENTIEL — USAGE INTERNE EXCLUSIF</div>')
        H.append('</div></div>')

        H.append('</div></body></html>')

        html = "\n".join(H)
        logger.info(f"HTML Vie genere : {len(html):,} chars, source={source}")
        return html

    except Exception as e:
        logger.error(f"export_html Vie : {e}", exc_info=True)
        return f"<html><body><h1>Erreur</h1><p>{e}</p></body></html>"


def export_pdf(**kwargs) -> bytes:
    """Genere le PDF du rapport Vie via WeasyPrint."""
    try:
        if not WEASYPRINT_OK:
            raise ImportError("weasyprint requis")
        html_str = export_html(**kwargs)
        pdf = WP_HTML(string=html_str).write_pdf()
        logger.info(f"PDF Vie : {len(pdf):,} bytes")
        return pdf
    except Exception as e:
        logger.error(f"export_pdf Vie : {e}", exc_info=True)
        return b""


def export_word(
    result_v3=None, result_v5=None, result_v6=None, result_v7=None,
    result_v8=None, result_v9=None, result_rvie1=None, result_rvie2=None,
    result_ep1=None, result_ep3=None, result_ep4=None,
    result_ep6=None, result_ep7=None,
    commentaire="", ref_client="", arrete="",
    audit_id="", actuaire_nom="", actuaire_numero_ia="",
) -> bytes:
    """Genere le rapport Word professionnel Vie & EP-RE — 9 chapitres."""
    if not DOCX_OK:
        logger.error("python-docx absent"); return b""
    try:
        def rgb(h):
            h = h.lstrip("#")
            return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))
        NR=rgb(NAVY); GR=rgb(GOLD); BR=rgb(WHITE)
        SR=rgb(SLATE); RgR=rgb(ROUGE); VR=rgb(VERT); AR=rgb(ORANGE)

        dt  = datetime.now().strftime("%d/%m/%Y")
        arr = arrete or dt
        cli = ref_client or "A renseigner"

        r3=result_v3 or {}; r5=result_v5 or {}; r6=result_v6 or {}
        r7=result_v7 or {}; r8=result_v8 or {}; r9=result_v9 or {}
        rr1=result_rvie1 or {}; ep1=result_ep1 or {}; ep3=result_ep3 or {}
        ep4=result_ep4 or {}; ep6=result_ep6 or {}; ep7=result_ep7 or {}

        pm  = float(r3.get("pm_prospective",0) or 0)
        be  = float(rr1.get("be_vie",pm) or pm)
        tp  = float(r5.get("tp_vie",be*1.05) or be*1.05)
        tp_be = float(r5.get("ratio_tp_be",tp/be if be else 1) or 1)
        scr = float(rr1.get("scr_vie_total",0) or 0)
        ratio_s = float(r5.get("ratio_scr_pct",0) or 0)
        ev  = float(r9.get("embedded_value",0) or 0)
        ane = float(r9.get("ane",0) or 0)
        vif = float(r9.get("vif",0) or 0)
        vnb = float(r9.get("vnb",0) or 0)
        dbo = float(_g(ep1,"ias19","dbo_total",default=0) or 0)
        pm_ep = float(_g(ep3,"provisions","pm_encours",default=0) or 0)
        ppb_ep= float(_g(ep3,"provisions","ppb",default=0) or 0)
        avis  = r5.get("avis_pa","FAVORABLE AVEC RESERVE")
        csm   = float(r8.get("csm",0) or 0)

        contexte = _construire_contexte_vie(
            r3,r5,r7,r8,r9,r6,rr1,{},ep1,ep3,ep4,ep6,ep7,cli,arr)
        narration, source = _generer_narration_vie(contexte, commentaire)

        doc = Document()
        for s in doc.sections:
            s.top_margin=Cm(2.2); s.bottom_margin=Cm(2.2)
            s.left_margin=Cm(2.8); s.right_margin=Cm(2.8)

        def _bg(cell, hex6):
            tc=cell._tc; tcp=tc.get_or_add_tcPr()
            sd=OxmlElement("w:shd")
            sd.set(qn("w:fill"), hex6.lstrip("#"))
            sd.set(qn("w:color"),"auto"); sd.set(qn("w:val"),"clear")
            tcp.append(sd)

        def _run(p, txt, bold=False, italic=False, sz=10, col=None):
            r=p.add_run(str(txt)); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
            if col: r.font.color.rgb=col
            return r

        def _h(txt, lv=1, col=None):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(10 if lv==1 else 6)
            p.paragraph_format.space_after=Pt(4)
            sz={1:16,2:12,3:10}.get(lv,10)
            c=col or (NR if lv==1 else GR)
            _run(p, txt, bold=True, sz=sz, col=c)

        def _sep():
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
            bo=OxmlElement("w:bottom")
            bo.set(qn("w:val"),"single"); bo.set(qn("w:sz"),"6")
            bo.set(qn("w:space"),"1"); bo.set(qn("w:color"),"C9A84C")
            pBdr.append(bo); pPr.append(pBdr)

        def _tbl(heads, rows, ws=None):
            from docx.enum.table import WD_ALIGN_VERTICAL
            t=doc.add_table(rows=1+len(rows), cols=len(heads)); t.style="Table Grid"
            for i,hd in enumerate(heads):
                c=t.rows[0].cells[i]; _bg(c,"0B1E3D")
                pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=pp.add_run(str(hd)); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=BR
            for ri,row in enumerate(rows):
                for ci,v in enumerate(row):
                    c=t.rows[ri+1].cells[ci]
                    if ri%2==1: _bg(c,"EEF2F7")
                    pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r=pp.add_run(str(v) if v is not None else "—"); r.font.size=Pt(9)
            if ws:
                for i,w in enumerate(ws):
                    for row in t.rows: row.cells[i].width=Cm(w)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)

        # Page de garde
        p=doc.add_paragraph()
        _run(p,"Actuar",bold=True,sz=28,col=NR); _run(p,"IA",bold=True,sz=28,col=GR)
        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p,"RAPPORT CONSOLIDE ACTUARIEL\n",bold=True,sz=22,col=NR)
        _run(p,"Direction Vie & Epargne-Retraite",bold=True,sz=14,col=GR)
        doc.add_paragraph()
        avis_col = VR if "FAVORABLE" in avis.upper() and "RESERVE" not in avis.upper() else (
                   AR if "RESERVE" in avis.upper() else RgR)
        p=doc.add_paragraph()
        _run(p,"Avis Fonction Actuarielle : ",sz=11,col=NR)
        _run(p,avis,bold=True,sz=11,col=avis_col)
        doc.add_paragraph()
        _tbl(["Client","Sous-direction","Arrete","Actuaire"],
             [[cli,"Vie Pure + EP-RE",arr,actuaire_nom or "A renseigner"]],
             ws=[3.5,4.0,3.0,5.5])
        doc.add_page_break()

        # §1
        _h("§1 — Synthese executive"); _sep()
        _tbl(
            ["Indicateur","Valeur","Reference","Seuil"],
            [["PM Vie pure",_f(pm),"Art. R331-1","—"],
             ["Best Estimate S2",_f(be),"Art. 77 S2","—"],
             ["TP Solvabilite 2",_f(tp),"BE + Risk Margin","—"],
             ["Ratio TP/BE",f"{tp_be:.4f}","QRT S.12","1.0-1.5"],
             ["SCR Vie total",_f(scr),"Art. 105 S2","—"],
             ["Ratio SCR",_pct(ratio_s),"QRT S.23",">= 100% (cible 150%)"],
             ["DBO IAS 19",_f(dbo),"§67 PUC","—"],
             ["PM epargne-retraite",_f(pm_ep),"Art. R342-14","—"],
             ["Embedded Value",_f(ev),"MCEV §18","—"],
             ["ANE",_f(ane),"Actifs - TP S2","—"],
             ["VIF",_f(vif),"VA profits - CoC 6%","—"],
             ["VNB",_f(vnb),"CEA EV §4.3","Benchmark 8-15%"]],
            ws=[4.5,3.5,4.0,4.0])
        doc.add_page_break()

        # §2
        _h("§2 — Provisions Techniques — Reconciliation Tricolonne"); _sep()
        _tbl(
            ["Referentiel","Base legale","Montant","Delta vs PM Sociale"],
            [["PM Sociale","Art. R331-1",_f(pm),"—"],
             ["TP Solvabilite 2","Art. 77 S2 (BE+RM)",_f(tp),_pct((tp-pm)/pm*100 if pm else 0)],
             ["LRC IFRS 17","§32-46 (FCF + CSM)",_f(tp+csm),_pct((tp+csm-pm)/pm*100 if pm else 0)],
             ["CSM IFRS 17","§45",_f(csm),"Integre dans LRC"],
             ["Mode BE IFRS 17",r8.get("be_ifrs17_mode","approx. S2"),"—","—"]],
            ws=[4.0,4.5,3.5,4.0])
        doc.add_page_break()

        # §3
        _h("§3 — SCR Vie — Couverture en Capital (Art. 138-145)"); _sep()
        scr_mod_w = rr1.get("scr_par_module",{}) or {}
        dom_w = max(scr_mod_w, key=lambda k: float(scr_mod_w[k] or 0)) if scr_mod_w else "—"
        rows_scr = [
            [k.title(), _f(float(v or 0)), _pct(float(v or 0)/scr*100 if scr else 0),
             "DOMINANT" if k==dom_w else ""]
            for k,v in scr_mod_w.items()
        ]
        rows_scr.append(["SCR Total",_f(scr),"100%","Ratio : "+_pct(ratio_s)])
        _tbl(["Sous-module","SCR (euros)","Part (%)","Remarque"],rows_scr,ws=[4.0,3.5,2.5,6.0])
        doc.add_page_break()

        # §4
        _h("§4 — Risque TMG — Portefeuille Underwater (Art. 142)"); _sep()
        p=doc.add_paragraph()
        _run(p,f"SCR taux (convexite 2e ordre) : {_f(float(r7.get('scr_taux',0) or 0))} — Art. 166",sz=9,col=NR)
        doc.add_page_break()

        # §5
        _h("§5 — Projection ALM — Adequation Actif-Passif"); _sep()
        dur = float(r6.get("duration_gap",0) or 0)
        ih  = float(r6.get("impact_stress_hausse",0) or 0)
        ib  = float(r6.get("impact_stress_baisse",0) or 0)
        _tbl(
            ["Indicateur ALM","Valeur","Reference","Statut"],
            [["Duration gap",f"{dur:.2f} ans","EIOPA ALM",r6.get("statut_rag","—")],
             ["Impact +200bp",_f(ih),"DeltaSurplus = -(D_A*A - D_P*PM)*Dr","EIOPA"],
             ["Impact -200bp",_f(ib),"Art. 166 Actes delegues S2","EIOPA"],
             ["Ratio couverture",_pct(float(r6.get("ratio_couverture_initial",0) or 0)),"Actifs/PM",r6.get("statut_rag","—")]],
            ws=[5.5,3.0,5.0,2.5])
        doc.add_page_break()

        # §6
        _h("§6 — Engagements EP-RE — IAS 19 & Optimisation PB"); _sep()
        ias = ep1.get("ias19",{}) or {}
        _tbl(
            ["Indicateur EP-RE","Valeur","Reference"],
            [["DBO totale",_f(dbo),"§67 PUC IAS 19"],
             ["Service Cost",_f(float(ias.get("service_cost",0) or 0)),"§67"],
             ["Interest Cost",_f(float(ias.get("interest_cost",0) or 0)),"§67"],
             ["Taux iBoxx AA",_pct(float(ias.get("taux_actu",0) or 0)),"§83 IAS 19"],
             ["Sensibilite DBO -100bp",_pct(float(ias.get("sensibilite_dbo_taux_pct",0) or 0)),"§83 duree modifiee"],
             ["PM epargne-retraite",_f(pm_ep),"Art. R342-14"],
             ["PPB EP-RE",_f(ppb_ep),f"Plafond 10% PM C2023-10 : {_pct(ppb_ep/pm_ep*100 if pm_ep else 0)}"],
             ["Couverture actifs/PM",_pct(float(ep4.get("ratio_base",0) or 0)),"Seuil 100%"],
             ["Taux PB optimal",_pct(float(ep7.get("taux_pb_optimal",0) or 0)),"Art. L132-29 (85%/90%)"],
             ["SCR post-PB",_pct(float(ep7.get("ratio_scr_post_pb",0) or 0)),"Seuil >= 100%"]],
            ws=[5.5,3.5,7.0])
        doc.add_page_break()

        # §7
        _h("§7 — Embedded Value — Creation de Valeur (MCEV §18)"); _sep()
        _tbl(
            ["Composante EV","Valeur","Formule"],
            [["ANE",_f(ane),"Actifs marche - TP S2"],
             ["VIF",_f(vif),"VA profits - VA CoC (6% x SCR)"],
             ["EV totale",_f(ev),"ANE + VIF — MCEV §18"],
             ["VNB",_f(vnb),"CEA EV Principles §4.3"],
             ["Ratio VNB/primes",_pct(float(r9.get("ratio_vnb_primes",0) or 0)),"Benchmark 8-15%"],
             ["Unwinding",_f(float(r9.get("unwinding",0) or 0)),"EV_{N-1} x taux_rf"],
             ["Gain experience",_f(float(r9.get("gain_experience",0) or 0)),"DeltaEV - VNB - Unwinding"]],
            ws=[4.5,3.5,8.0])
        doc.add_page_break()

        # §8
        _h("§8 — Back-Testing — Robustesse des Hypotheses (IAS 19.145)"); _sep()
        ecarts = ep6.get("ecarts",{}) or {}
        rows_bt = [
            [param,str(ecarts.get(f"{param}_bps","—"))+" bp",
             _f(float(ep6.get(f"impact_{param}",0) or 0)),
             ep6.get(f"statut_{param}","—")]
            for param in ["mortalite","rotation","revalorisation","taux"]
        ]
        rows_bt.append(["Impact total DBO","—",_f(float(ep6.get("impact_dbo_total",0) or 0)),ep6.get("statut_rag","—")])
        _tbl(["Parametre","Ecart observe","Impact DBO","Statut"],rows_bt,ws=[4.0,3.0,3.5,5.5])
        doc.add_page_break()

        # §9
        _h("§9 — Commentaire Actuariel — Avis de la Fonction Actuarielle"); _sep()
        if narration:
            secs = re.split(r"(?=§\d+\s*[-—–])", _clean(narration))
            for sec in secs:
                sec = sec.strip()
                if not sec: continue
                ls = sec.split("\n",1)
                if ls[0]: _h(ls[0],lv=2)
                if len(ls) > 1:
                    for ln in ls[1].split("\n"):
                        ln = ln.strip()
                        if ln:
                            p=doc.add_paragraph()
                            p.paragraph_format.space_after=Pt(3)
                            p.paragraph_format.left_indent=Cm(0.4)
                            _run(p,ln,sz=9,col=NR)
            if source=="claude_api":
                p=doc.add_paragraph()
                _run(p,"❆ Narration generee par ActuarIA Intelligence",sz=7,italic=True,col=SR)
        else:
            p=doc.add_paragraph()
            _run(p,"Narration non disponible — configurer ANTHROPIC_API_KEY.",sz=9,italic=True,col=SR)

        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p,"Avis Fonction Actuarielle (Art. 48 §g Directive 2009/138/CE) : ",sz=10,col=NR)
        _run(p,avis,bold=True,sz=12,col=avis_col)

        for reco in (r5.get("recommandations") or []):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
            _run(p,"→  ",bold=True,sz=10,col=GR)
            _run(p,_clean(str(reco)),sz=9,col=NR)

        _sep()
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p,f"ActuarIA · {cli} · Direction Vie & EP-RE · Arrete {arr} · {dt} · CONFIDENTIEL",
             sz=7,italic=True,col=SR)

        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        wb=buf.read()
        logger.info(f"Word Vie : {len(wb):,} bytes")
        return wb

    except Exception as e:
        logger.error(f"export_word Vie : {e}", exc_info=True)
        return b""
