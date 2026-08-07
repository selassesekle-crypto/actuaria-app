"""
ActuarIA — Rapport Consolidé Épargne-Retraite
Direction Vie & EP-RE — Sous-direction EP-RE (EP1-EP7)

Produit :
  - HTML  : rapport premium page de garde + 8 sections + narration Claude
  - Word  : rapport professionnel 8 chapitres
  - PDF   : via WeasyPrint

Niveau : équivalent rapport A7 Non-Vie.
Auteur : ActuarIA — Direction Paul (Vie/EP-RE)

ActuarIA — Rapport Consolidé Vie Pure
Direction Vie & EP-RE — Sous-direction Vie (V1-V9 + R-VIE1/2)

Produit :
  - HTML  : rapport premium page de garde + 9 sections + narration Claude
  - Word  : rapport professionnel 9 chapitres
  - PDF   : via WeasyPrint

Niveau : équivalent rapport A7 Non-Vie.
"""

import io, re, logging, base64
from datetime import datetime
from typing import Dict, Optional, Tuple

from core import frontiere_llm

logger = logging.getLogger("actuaria.rapport_epre")

try:
    import anthropic
    ANTHROPIC_OK = True
except ImportError:
    ANTHROPIC_OK = False

try:
    from weasyprint import HTML as WP_HTML
    WEASYPRINT_OK = True
except ImportError:
    WEASYPRINT_OK = False

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    import plotly.io as pio
    PLOTLY_OK = True
except ImportError:
    PLOTLY_OK = False

# Palette
NAVY   = "#0B1E3D"
NAVY_M = "#132844"
NAVY_L = "#1E3A5F"
GOLD   = "#C9A84C"
ROUGE  = "#C0392B"
ORANGE = "#E67E22"
VERT   = "#1E8449"
SLATE  = "#8A9AB0"
WHITE  = "#FFFFFF"

LOGO_SVG = (
    '<svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 60">'
    '  <text x="4" y="44" font-family="Georgia,serif" font-size="38"'
    '        font-weight="900" fill="#FFFFFF">Actuar</text>'
    '  <text x="128" y="44" font-family="Georgia,serif" font-size="38"'
    '        font-weight="900" fill="#C9A84C">IA</text>'
    '  <rect x="4" y="50" width="180" height="2" fill="#C9A84C" opacity="0.6"/>'
    '</svg>'
)
LOGO_URI = "data:image/svg+xml;base64," + base64.b64encode(LOGO_SVG.encode()).decode()

SYSTEM_PROMPT_EPRE = "Tu es un actuaire Epargne-Retraite senior certifie par l'Institut des Actuaires (IA), expert en regimes de retraite supplementaire d'entreprise (Article 39, Article 83, PER Collectif), actuaire conseil aupres de grandes entreprises, groupes industriels et institutions de prevoyance. Tu exerces la Fonction Actuarielle au sens de l'Art. 48 de la Directive Solvabilite 2 avec 30 ans d'experience dans l'evaluation des engagements sociaux, le provisionnement des regimes de retraite et la conformite IAS 19 / Solvabilite 2 / Loi PACTE.\\n\\nTu maitrises dans leur integralite :\\n- IAS 19 : methode Projected Unit Credit (PUC, §67) ; DBO = Σ UC_i × ä_i × facteur_actu ; Service Cost (droit acquis exercice) ; Interest Cost (desactualisation DBO) ; gains/pertes actuariels (§145, retraitement OCI) ; taux d'actualisation iBoxx EUR Corporate AA (§83) ; vesting progressif Art. 39 ; droits acquis a la date de depart.\\n- Solvabilite 2 : SCR Vie sous-module longevite (-20% qx, Art. 138) et retraite ; BE provisions epargne-retraite ; Risk Margin (CoC 6%) ; QRT S.12 retraite, S.26.\\n- Loi PACTE 2019 / PER : transferabilite, portabilite, cas de deblocage anticipe, Fiche d'Information Annuelle (FIA), information beneficiaire Art. L224-32.\\n- Provisionnement : PM viagere = ä_x × rente_garantie ; PPB EP-RE (Art. R342-14, plafond 10% PM) ; reserve de capitalisation ; dotation annuelle = rendement - PB distribuee.\\n- Optimisation PB : Art. L132-29 (85% produits financiers nets, 90% benefice technique) ; C2023-10 ACPR (PPB reprise 8 ans) ; contrainte SCR post-distribution.\\n- Stress et sensibilites : stress longevite +5 ans ; stress inflation revalorisation +100bp ; stress taux -100bp (impact DBO = D_mod × DBO × Δr, §83 IAS 19) ; stress rotation -20% (droits non acquis) ; scenario combine (pire cas).\\n- Back-testing IAS 19.145 : ecarts d'experience par parametre (mortalite, rotation, revalorisation salariale, taux) ; delta_DBO_p = DBO × (sensibilite_p/10000) × ecart_bps ; duree residuelle ponderee = Σ poids_i × (age_retraite_i - age_i) ; gain/perte actuarielle globale et impact OCI / fonds propres IFRS.\\n- Tables : TH0002/TF0002 (arrete 27/07/2006) ; CTIP 2023 pour taux arret de travail (LR_ITT = 0.68).\\n\\nTu rediges le commentaire actuariel du Rapport Consolide Epargne-Retraite, document de reference destine au Conseil d'Administration, au Comite Social et Economique (CSE), a l'Actuaire Designe et soumis a l'ACPR dans le cadre du reporting reglementaire. Ce rapport engage ta responsabilite professionnelle au sens de l'Art. 48 §g Directive 2009/138/CE et de l'Art. R142-4 Code des assurances.\\n\\nREGLES ABSOLUES -- AUCUNE DEROGATION ADMISSIBLE\\n\\n0. FORMAT STRICT : §N -- TITRE EN MAJUSCULES. ### Sous-titre. **terme** en gras. - tirets. ZERO tableau Markdown. ZERO blockquote. ZERO balise code.\\n1. LANGUE : Francais actuariel de haute tenue. Termes en anglais : DBO, IAS 19, PUC, CoC, SCR, BE, PER.\\n2. RIGUEUR CHIFFREE : toute affirmation etayee par une donnee du contexte. Si donnee absente : noter explicitement comme lacune a combler.\\n3. REFERENCES REGLEMENTAIRES OBLIGATOIRES : IAS 19 §67 (PUC), §83 (iBoxx AA), §145 (gains/pertes). S2 Art. 138 (longevite -20%), Art. 77 (BE), Art. R342-14 (PPB). Code ass. Art. L132-29 (PB 85%/90%), Art. L224-32 (PER/Loi PACTE). ACPR C2023-10 (PPB plafond 10%, reprise 8 ans). Tables : arrete 27/07/2006 (TH0002/TF0002). CTIP 2023 (LR_ITT = 0.68).\\n4. ALERTES : tout indicateur ROUGE presente avec : (a) fait + deviation vs seuil, (b) mecanisme actuariel, (c) impact bilantaire/OCI, (d) action corrective + echeance.\\n5. CAUSALITE : [Observation] -> [Mecanisme IAS 19 / S2] -> [Consequence chiffree].\\nExemples : Taux iBoxx -100bp -> DBO +D_mod% -> OCI -X€ -> fonds propres IFRS reduits. Ratio couverture 92% -> sous-provisionnement X€ -> dotation complementaire requise avant cloture. Ecart rotation +10bp -> gain actuariel +Y€ -> impact OCI positif N+1.\\n6. INCERTITUDE : sensibilite DBO a -100bp = D_mod x DBO (§83). Stress longevite +5 ans -> PM + Z%. Stress inflation +100bp -> DBO + W%. Scenario combine = pire cas des trois stress.\\n7. POSTURE : assertif, independant. 'La Fonction Actuarielle recommande...' + echeance. Les avis sont ceux de l'actuaire, pas de la direction financiere.\\n8. INTERDICTIONS : pas de generique sans chiffre. Pas d'avis FAVORABLE si : ratio couverture < 100%, DBO non certifiee conforme iBoxx, backtesting ROUGE sur >= 2 params, PPB > plafond C2023-10 sans plan de reprise.\\n9. COHERENCE : DBO §2 coherente avec IAS 19 §6. PM §3 coherente avec provisions §5. SCR §4 coherent avec ratio couverture §3.\\n\\nSTRUCTURE OBLIGATOIRE -- 8 SECTIONS\\n\\n§1 -- CONTEXTE, PERIMETRE ET GOUVERNANCE DU REGIME\\nRegime concerne (Art. 39 / Art. 83 / PER), entreprise, date arrete, effectifs couverts, encours PM totaux, taux de cotisation, actuaire designe, calendrier reglementaire.\\n\\n§2 -- ENGAGEMENTS IAS 19 -- DBO ET DECOMPOSITION\\nDBO totale et decomposition : Service Cost (droit acquis exercice, §67), Interest Cost (desactualisation, §67), gains/pertes actuariels (§145, OCI). Taux iBoxx EUR AA utilise (§83) et conformite ±25bp. Sensibilite : delta_DBO/delta_taux = D_mod x DBO. Alerte si non-conformite iBoxx ou sensibilite > 15%.\\n\\n§3 -- PROVISIONS ET COUVERTURE DU REGIME\\nPM viagere = ä_x x rente_garantie. PPB et ratio PPB/PM vs plafond C2023-10 (10%). Reserve de capitalisation. Ratio couverture actifs/PM (seuil 100%). Alerte si sous-couverture et quantification de la dotation complementaire requise.\\n\\n§4 -- ANALYSE SCR ET EXIGENCES SOLVABILITE 2\\nSCR longevite (Art. 138 : -20% qx). SCR retraite global. Ratio SCR / fonds propres (seuil 100%, cible ORSA 150%). Impact du stress longevite +5 ans sur PM et SCR.\\n\\n§5 -- OPTIMISATION DE LA PARTICIPATION AUX BENEFICES\\nTaux PB optimal retenu. Contraintes Art. L132-29 (85% produits financiers / 90% benefice technique). PPB residuelle vs plafond C2023-10. Ratio SCR post-distribution. Projection PPB pluriannuelle et horizon de reprise obligatoire (8 ans).\\n\\n§6 -- ANALYSE DE STRESS ET SENSIBILITES\\nStress longevite +5 ans : impact PM. Stress inflation +100bp : impact DBO. Stress taux -100bp : delta_DBO = D_mod x DBO (§83). Stress rotation -20%. Scenario combine : pire cas avec ratio de couverture stresse.\\n\\n§7 -- BACK-TESTING DES HYPOTHESES ACTUARIELLES (IAS 19.145)\\nEcarts d'experience par parametre. Quantification : delta_DBO_p = DBO x (sensibilite_p/10000) x ecart_bps. Gain/perte actuarielle globale et impact OCI. Fiabilite des hypotheses pour N+1. Ajustements recommandes.\\n\\n§8 -- CONCLUSION, AVIS FORMEL ET RECOMMANDATIONS\\nAvis : FAVORABLE / FAVORABLE AVEC RESERVE / DEFAVORABLE. Conditions FAVORABLE : ratio couverture >= 100%, DBO conforme iBoxx, backtesting <= 1 parametre ROUGE, PPB dans le plafond, SCR >= 100%. Recommandations par priorite + echeance (immmediat / prochain arrete / plan triennal). Point de vigilance CSE et Comite des Risques. Soutenabilite du regime a horizon 10 ans."


def _f(v, dec=0) -> str:
    try:
        n = float(v or 0)
        s = f"{n:,.{dec}f}".replace(",", "\u202f")
        return s + "\u202f€"
    except Exception:
        return str(v) if v is not None else "—"

def _pct(v, dec=1) -> str:
    try:
        return f"{float(v or 0):.{dec}f}\u202f%"
    except Exception:
        return "—"

def _s(v) -> str:
    return str(v) if v is not None else "—"

def _clean(txt) -> str:
    if not txt:
        return ""
    txt = re.sub(r"<[^>]+>", " ", str(txt))
    txt = re.sub(r"\s+", " ", txt)
    return txt.strip()

def _statut_cls(s: str) -> str:
    s = (s or "").upper()
    if "ROUGE" in s or "FAVO" not in s and "DEF" in s:
        return "rouge"
    if "AMBRE" in s or "RESERVE" in s:
        return "ambre"
    return "vert"

def _g(d, *keys, default=None):
    v = d or {}
    for k in keys:
        if isinstance(v, dict):
            v = v.get(k, default)
        else:
            return default
    return v if v is not None else default

def _css() -> str:
    # CSS premium identique palette A7 — adapte direction Vie
    return CSS_VIE

CSS_VIE = '''
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@0,700;0,900;1,400&family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');
:root {
  --navy:#0B1E3D; --navy-mid:#132844; --navy-light:#1E3A5F;
  --gold:#C9A84C; --gold-light:#E2C97E; --gold-pale:rgba(201,168,76,0.12);
  --rouge:#C0392B; --rouge-pale:rgba(192,57,43,0.08);
  --orange:#E67E22; --vert:#1E8449;
  --slate:#8A9AB0; --slate-light:#B8C5D3;
  --bg:#F5F7FA; --white:#FFFFFF; --text:#1C2B3A; --text-mid:#3D5166; --border:#DDE4EE;
}
* { box-sizing:border-box; margin:0; padding:0; }
body { font-family:'Inter',sans-serif; font-size:10pt; color:var(--text); background:var(--bg); line-height:1.7; }
.rapport-container { max-width:960px; margin:0 auto; background:var(--white); box-shadow:0 4px 40px rgba(11,30,61,0.14); }
.page-garde { position:relative; min-height:720px; display:flex; flex-direction:column; overflow:hidden; background:var(--navy); page-break-after:always; }
.garde-bg { position:absolute; inset:0; overflow:hidden; pointer-events:none; }
.garde-bg::before { content:''; position:absolute; top:-120px; right:-80px; width:560px; height:560px; border:1px solid rgba(201,168,76,0.12); border-radius:50%; }
.garde-bg::after  { content:''; position:absolute; top:-60px; right:-20px; width:380px; height:380px; border:1px solid rgba(201,168,76,0.20); border-radius:50%; }
.garde-dots { position:absolute; top:0; right:0; width:340px; height:100%; background-image:radial-gradient(circle,rgba(201,168,76,0.25) 1px,transparent 1px); background-size:22px 22px; mask-image:linear-gradient(to left,rgba(0,0,0,0.7),transparent); }
.garde-diagonal { position:absolute; inset:0; background:linear-gradient(135deg,transparent 62%,rgba(201,168,76,0.06) 62%); }
.garde-accent-line { position:absolute; left:52px; top:0; bottom:0; width:2px; background:linear-gradient(to bottom,transparent,var(--gold) 20%,var(--gold) 80%,transparent); opacity:0.5; }
.garde-inner { position:relative; z-index:2; display:flex; flex-direction:column; height:100%; }
.garde-header { display:flex; justify-content:space-between; align-items:flex-start; padding:36px 56px 0; }
.garde-logo-wrap img { height:52px; }
.garde-badges { display:flex; flex-direction:column; align-items:flex-end; gap:8px; }
.badge-confidentiel { font-size:7.5pt; font-weight:700; color:var(--rouge); letter-spacing:2.5px; border:1.5px solid var(--rouge); padding:5px 12px; text-transform:uppercase; }
.badge-direction { font-size:7pt; font-weight:600; color:var(--gold); letter-spacing:1.5px; border:1px solid rgba(201,168,76,0.4); padding:4px 10px; text-transform:uppercase; }
.garde-hero { flex:1; display:flex; flex-direction:column; justify-content:center; padding:48px 56px 0 80px; }
.garde-eyebrow { font-size:8pt; font-weight:600; color:var(--gold); letter-spacing:4px; text-transform:uppercase; margin-bottom:20px; }
.garde-titre { font-family:'Playfair Display',serif; font-size:48pt; font-weight:900; color:var(--white); line-height:1.05; margin-bottom:8px; }
.garde-titre em { font-style:italic; color:var(--gold); }
.garde-subtitle { font-family:'Playfair Display',serif; font-size:16pt; font-weight:400; font-style:italic; color:var(--slate-light); margin-bottom:8px; }
.garde-sousdirection { font-size:9pt; font-weight:500; color:var(--gold); letter-spacing:1.5px; margin-bottom:28px; }
.garde-sep { display:flex; align-items:center; gap:16px; margin-bottom:28px; }
.garde-sep-line { height:1px; width:80px; background:var(--gold); opacity:0.6; }
.garde-sep-diamond { width:8px; height:8px; background:var(--gold); transform:rotate(45deg); }
.garde-statut { display:inline-flex; align-items:center; gap:12px; padding:10px 24px; }
.garde-statut-rouge { background:rgba(192,57,43,0.15); border:1px solid rgba(192,57,43,0.5); }
.garde-statut-ambre { background:rgba(230,126,34,0.15); border:1px solid rgba(230,126,34,0.5); }
.garde-statut-vert  { background:rgba(30,132,73,0.15);  border:1px solid rgba(30,132,73,0.5); }
.statut-dot { width:10px; height:10px; border-radius:50%; animation:pulse 2s ease-in-out infinite; }
.statut-dot-rouge { background:var(--rouge); box-shadow:0 0 8px rgba(192,57,43,0.7); }
.statut-dot-ambre { background:var(--orange); box-shadow:0 0 8px rgba(230,126,34,0.7); }
.statut-dot-vert  { background:var(--vert);  box-shadow:0 0 8px rgba(30,132,73,0.7); }
@keyframes pulse { 0%,100%{opacity:1;transform:scale(1)} 50%{opacity:0.7;transform:scale(1.2)} }
.statut-label-rouge { font-size:8pt; font-weight:700; color:var(--rouge); letter-spacing:3px; text-transform:uppercase; }
.statut-label-ambre { font-size:8pt; font-weight:700; color:var(--orange); letter-spacing:3px; text-transform:uppercase; }
.statut-label-vert  { font-size:8pt; font-weight:700; color:var(--vert);  letter-spacing:3px; text-transform:uppercase; }
.garde-footer { border-top:1px solid rgba(201,168,76,0.25); margin:36px 56px 0; padding:24px 0 40px; }
.garde-kpis { display:grid; grid-template-columns:repeat(6,1fr); gap:0; }
.garde-kpi { padding:0 20px 0 0; border-right:1px solid rgba(255,255,255,0.08); }
.garde-kpi:last-child { border-right:none; }
.kpi-label { font-size:6.5pt; font-weight:600; color:var(--slate); text-transform:uppercase; letter-spacing:1.5px; margin-bottom:6px; }
.kpi-value { font-family:'Playfair Display',serif; font-size:11pt; font-weight:700; color:var(--white); line-height:1.2; }
.kpi-value.highlight { color:var(--gold); font-size:13pt; }
.kpi-sub { font-size:7pt; color:var(--slate); margin-top:3px; }
.rapport-body { padding:0; }
.section-header { background:linear-gradient(to right,var(--navy),var(--navy-light)); padding:22px 56px; display:flex; align-items:center; gap:20px; }
.section-num { font-family:'JetBrains Mono',monospace; font-size:9pt; font-weight:500; color:var(--gold); opacity:0.8; min-width:28px; }
.section-titre { font-family:'Playfair Display',serif; font-size:15pt; font-weight:700; color:var(--white); }
.section-body { padding:36px 56px 48px; }
.section-divider { height:1px; background:var(--border); }
.kpi-grid { display:grid; grid-template-columns:repeat(4,1fr); gap:12px; margin-bottom:32px; }
.kpi-card { background:var(--navy); border-radius:6px; padding:20px 18px; position:relative; overflow:hidden; }
.kpi-card::before { content:''; position:absolute; top:0; left:0; right:0; height:3px; background:var(--gold); }
.kpi-card-rouge::before { background:var(--rouge); }
.kpi-card-ambre::before { background:var(--orange); }
.kpi-card-vert::before  { background:var(--vert); }
.kpi-card-label { font-size:7pt; font-weight:600; color:var(--slate); text-transform:uppercase; letter-spacing:1.5px; margin-bottom:8px; }
.kpi-card-value { font-family:'Playfair Display',serif; font-size:16pt; font-weight:700; color:var(--white); line-height:1.1; margin-bottom:6px; }
.kpi-card-value-rouge { color:var(--rouge) !important; }
.kpi-card-value-ambre { color:var(--orange) !important; }
.kpi-card-value-vert  { color:#2ECC71 !important; }
.kpi-card-sub { font-size:7.5pt; color:var(--gold); }
.table-section-title { font-family:'Playfair Display',serif; font-size:10pt; font-weight:600; color:var(--gold); margin:28px 0 12px; padding-bottom:6px; border-bottom:1px solid var(--gold-pale); }
table.premium { width:100%; border-collapse:collapse; font-size:9pt; }
table.premium thead tr { background:var(--navy); }
table.premium thead th { padding:11px 14px; color:var(--white); font-weight:600; font-size:8pt; text-align:left; }
table.premium thead th.right { text-align:right; }
table.premium tbody tr:nth-child(even) { background:#F8FAFB; }
table.premium tbody tr.highlight-gold { background:var(--gold-pale); font-weight:700; }
table.premium tbody td { padding:9px 14px; border-bottom:1px solid var(--border); color:var(--text); }
table.premium tbody td.right { text-align:right; font-family:'JetBrains Mono',monospace; font-size:8.5pt; }
table.premium tbody td.label { font-weight:600; color:var(--navy); }
.badge { display:inline-flex; align-items:center; gap:4px; font-size:7.5pt; font-weight:700; padding:2px 8px; border-radius:3px; }
.badge-vert  { background:rgba(30,132,73,0.12);  color:var(--vert); }
.badge-ambre { background:rgba(230,126,34,0.12); color:var(--orange); }
.badge-rouge { background:rgba(192,57,43,0.12);  color:var(--rouge); }
.tricolonne { display:grid; grid-template-columns:repeat(3,1fr); gap:16px; margin:24px 0; }
.tricolonne-card { border-radius:8px; overflow:hidden; border:1px solid var(--border); }
.tricolonne-header { padding:14px 20px; background:var(--navy); }
.tricolonne-ref { font-size:7pt; font-weight:600; color:var(--gold); letter-spacing:1.5px; text-transform:uppercase; margin-bottom:4px; }
.tricolonne-label { font-size:9pt; font-weight:700; color:var(--white); }
.tricolonne-body { padding:18px 20px; background:#FAFBFC; }
.tricolonne-value { font-family:'Playfair Display',serif; font-size:18pt; font-weight:700; color:var(--navy); margin-bottom:6px; }
.tricolonne-delta { font-size:8pt; color:var(--slate); }
.tricolonne-delta strong { color:var(--orange); }
.scr-grid { display:grid; grid-template-columns:repeat(4,1fr); gap:10px; margin:20px 0; }
.scr-card { padding:14px 16px; border-radius:6px; background:var(--navy); position:relative; }
.scr-card::before { content:''; position:absolute; top:0; left:0; right:0; height:3px; background:var(--gold); border-radius:3px 3px 0 0; }
.scr-card-dominant::before { background:var(--rouge); }
.scr-card-label { font-size:7pt; color:var(--slate); text-transform:uppercase; letter-spacing:1px; margin-bottom:6px; }
.scr-card-value { font-family:'Playfair Display',serif; font-size:13pt; font-weight:700; color:var(--white); }
.scr-card-pct { font-size:8pt; color:var(--gold); margin-top:4px; }
.scr-card-art { font-size:7pt; color:var(--slate); margin-top:2px; font-style:italic; }
.commentaire-wrap { border:1px solid var(--border); border-radius:6px; overflow:hidden; margin-top:8px; }
.commentaire-header { background:var(--navy); padding:14px 20px; display:flex; align-items:center; justify-content:space-between; }
.commentaire-header-title { font-size:8pt; font-weight:600; color:var(--gold); letter-spacing:2px; text-transform:uppercase; }
.commentaire-ai-badge { font-size:7pt; color:var(--slate); font-family:'JetBrains Mono',monospace; }
.commentaire-body { padding:28px 32px; background:#FAFBFC; }
.comm-section-title { font-family:'Playfair Display',serif; font-size:11pt; font-weight:700; color:var(--navy); border-bottom:1.5px solid var(--gold); padding-bottom:6px; margin:24px 0 12px; }
.comm-section-title:first-child { margin-top:0; }
.comm-h4 { font-size:9.5pt; font-weight:700; color:var(--navy); margin:16px 0 8px; }
.comm-p { font-size:9.5pt; line-height:1.85; color:var(--text-mid); margin-bottom:12px; }
.comm-p strong { color:var(--navy); font-weight:700; }
.comm-ul { margin:10px 0 12px 20px; list-style:none; }
.comm-ul li { position:relative; padding-left:16px; font-size:9.5pt; line-height:1.7; color:var(--text-mid); margin-bottom:6px; }
.comm-ul li::before { content:'—'; position:absolute; left:0; color:var(--gold); font-weight:700; }
.comm-ul li strong { color:var(--navy); }
.comm-divider { border:none; border-top:1px solid var(--border); margin:20px 0; }
.comm-footer { font-size:7.5pt; color:var(--slate); font-style:italic; text-align:right; border-top:1px solid var(--border); padding-top:10px; margin-top:8px; }
.avis-box { padding:24px 28px; border-radius:8px; margin:24px 0; display:flex; align-items:flex-start; gap:20px; }
.avis-favorable   { background:rgba(30,132,73,0.08);  border:2px solid rgba(30,132,73,0.4); }
.avis-reserve     { background:rgba(230,126,34,0.08); border:2px solid rgba(230,126,34,0.4); }
.avis-defavorable { background:rgba(192,57,43,0.08);  border:2px solid rgba(192,57,43,0.5); }
.avis-icon { font-size:24pt; flex-shrink:0; }
.avis-label { font-size:7pt; font-weight:700; letter-spacing:2px; text-transform:uppercase; margin-bottom:6px; }
.avis-value { font-family:'Playfair Display',serif; font-size:16pt; font-weight:700; margin-bottom:8px; }
.avis-value-favorable   { color:var(--vert); }
.avis-value-reserve     { color:var(--orange); }
.avis-value-defavorable { color:var(--rouge); }
.avis-text { font-size:9pt; color:var(--text-mid); line-height:1.6; }
.graph-wrap { border:1px solid var(--border); border-radius:6px; overflow:hidden; margin:20px 0; }
.graph-titre { background:var(--navy-light); padding:10px 18px; font-size:8pt; font-weight:600; color:var(--gold); letter-spacing:1px; text-transform:uppercase; }
.graph-body { padding:16px; background:#FAFBFC; }
.pied-de-page { background:var(--navy); padding:20px 56px; display:flex; justify-content:space-between; align-items:center; border-top:2px solid var(--gold); }
.pied-logo img { height:26px; }
.pied-meta { text-align:right; font-size:7.5pt; color:var(--slate); line-height:1.7; }
.pied-meta .confidentiel-footer { color:var(--rouge); font-weight:700; letter-spacing:1px; }
@media print { body{background:white} .rapport-container{box-shadow:none;max-width:none;margin:0} .page-garde{page-break-after:always;min-height:100vh} @page{margin:0} }
</style>
'''


def _md_to_html(texte: str) -> str:
    """Convertit la narration markdown actuarielle en HTML premium."""
    if not texte:
        return ""
    html = []
    for l in texte.split("\n"):
        l = l.rstrip()
        m = re.match(r"^§\s*(\d+)\s*[-—–]\s*(.+)$", l)
        if m:
            html.append(f'<div class="comm-section-title">§{m.group(1)} — {m.group(2).strip()}</div>')
            continue
        if l.startswith("###"):
            html.append(f'<div class="comm-h4">{l.lstrip("#").strip()}</div>')
            continue
        if not l.strip():
            html.append('<hr class="comm-divider">')
            continue
        if l.strip().startswith("- "):
            item = l.strip()[2:]
            item = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", item)
            html.append(f'<ul class="comm-ul"><li>{item}</li></ul>')
            continue
        line = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", l)
        html.append(f'<p class="comm-p">{line}</p>')
    return "\n".join(html)


def _construire_contexte_epre(
    result_ep1, result_ep2, result_ep3, result_ep4,
    result_ep6, result_ep7,
    arrete,
) -> str:
    """Contexte structuré EP-RE transmis à Claude."""
    ep1=result_ep1 or {}; ep2=result_ep2 or {}; ep3=result_ep3 or {}
    ep4=result_ep4 or {}; ep6=result_ep6 or {}; ep7=result_ep7 or {}

    ias    = ep1.get("ias19", {}) or {}
    dbo    = float(ias.get("dbo_total", 0) or 0)
    sc     = float(ias.get("service_cost", 0) or 0)
    ic     = float(ias.get("interest_cost", 0) or 0)
    taux_a = float(ias.get("taux_actu", 0) or 0)
    sens_t = float(ias.get("sensibilite_dbo_taux_pct", 0) or 0)
    prov   = ep3.get("provisions", {}) or {}
    pm_ep  = float(prov.get("pm_encours", 0) or 0)
    ppb_ep = float(prov.get("ppb", prov.get("ppb_total", 0)) or 0)
    rcapi  = float(prov.get("reserve_capi", 0) or 0)
    taux_c = float(ep3.get("taux_couverture_pct", 0) or 0)
    ratio_b= float(ep4.get("ratio_base", 0) or 0)
    ratio_l= float(ep4.get("ratio_longevite", 0) or 0)
    ratio_c= float(ep4.get("ratio_combine", 0) or 0)
    tpb_opt= float(ep7.get("taux_pb_optimal", 0) or 0)
    ppb_fin= float(ep7.get("ppb_finale", 0) or 0)
    scr_pb = float(ep7.get("ratio_scr_post_pb", 0) or 0)
    ecarts = ep6.get("ecarts", {}) or {}
    imp_dbo= float(ep6.get("impact_dbo_total", 0) or 0)
    tar    = ep2.get("tarification", {}) or {}
    cot    = float(tar.get("cotisation_brute_annuelle", 0) or 0)
    rente  = float(tar.get("rente_mensuelle", 0) or 0)
    t_remp = float(tar.get("taux_remplacement_pct", 0) or 0)

    lignes = [
        f"RAPPORT CONSOLIDE EP-RE -- Arrete {arrete}",
        "(Identite de l'organisme non transmise : ne jamais la nommer ni l'inventer.)",
        f"Date generation : {datetime.now().strftime('%d/%m/%Y')}",
        "",
        "=== ENGAGEMENTS IAS 19 ===",
        f"DBO totale (IAS 19 PUC §67)     : {_f(dbo)}",
        f"Service Cost (droit acquis)      : {_f(sc)}",
        f"Interest Cost (desactualisation) : {_f(ic)}",
        f"Taux actualisation iBoxx AA §83  : {_pct(taux_a)}",
        f"Sensibilite DBO a -100bp         : {_pct(sens_t)} (D_mod x DBO, §83)",
        "",
        "=== TARIFICATION ET PRESTATIONS ===",
        f"Cotisation annuelle brute        : {_f(cot)}",
        f"Rente mensuelle estimee          : {_f(rente, dec=2)}",
        f"Taux de remplacement             : {_pct(t_remp)}",
        "",
        "=== PROVISIONS ET COUVERTURE ===",
        f"PM epargne-retraite (ä_x x rente): {_f(pm_ep)}",
        f"PPB EP-RE                        : {_f(ppb_ep)} (plafond C2023-10 : 10% PM = {_f(pm_ep*0.1)})",
        f"Ratio PPB/PM                     : {_pct(ppb_ep/pm_ep*100 if pm_ep else 0)}",
        f"Reserve de capitalisation        : {_f(rcapi)}",
        f"Taux de couverture actifs/PM     : {_pct(taux_c)} (seuil 100%)",
        "",
        "=== STRESS ET SENSIBILITES ===",
        f"Ratio de couverture base         : {_pct(ratio_b)}",
        f"Ratio post-stress longevite      : {_pct(ratio_l)} (+5 ans)",
        f"Ratio post-stress combine        : {_pct(ratio_c)} (pire cas)",
        f"Statut stress                    : {ep4.get('statut_rag', 'non disponible')}",
        "",
        "=== OPTIMISATION PB (ART. L132-29, C2023-10) ===",
        f"Taux PB optimal retenu           : {_pct(tpb_opt)}",
        f"PB minimale reglementaire        : Art. L132-29 (85% prod. fin. / 90% ben. tech.)",
        f"PPB finale apres optimisation    : {_f(ppb_fin)}",
        f"Ratio SCR post-distribution      : {_pct(scr_pb)} (seuil 100%)",
        f"Statut optimisation              : {ep7.get('statut_rag', 'non disponible')}",
        "",
        "=== BACK-TESTING (IAS 19.145) ===",
        f"Ecart mortalite observé          : {ecarts.get('mortalite_bps', 'non disponible')} bp",
        f"Ecart rotation observe           : {ecarts.get('rotation_bps', 'non disponible')} bp",
        f"Ecart revalorisation observe     : {ecarts.get('revalorisation_bps', 'non disponible')} bp",
        f"Ecart taux observe               : {ecarts.get('taux_bps', 'non disponible')} bp",
        f"Impact DBO total                 : {_f(imp_dbo)}",
        f"Statut backtesting               : {ep6.get('statut_rag', 'non disponible')}",
        "",
        "=== STATUTS GLOBAUX ===",
        f"IAS 19 (EP1)                     : {ep1.get('statut_rag', 'non disponible')}",
        f"Tarification (EP2)               : {ep2.get('statut_rag', 'non disponible')}",
        f"Provisionnement (EP3)            : {ep3.get('statut_rag', 'non disponible')}",
        f"Stress (EP4)                     : {ep4.get('statut_rag', 'non disponible')}",
        f"Backtesting (EP6)                : {ep6.get('statut_rag', 'non disponible')}",
        f"Optimisation PB (EP7)            : {ep7.get('statut_rag', 'non disponible')}",
        "",
        "Redige le commentaire actuariel complet en 8 sections (§1 a §8) selon les regles du prompt.",
    ]
    return "\n".join(lignes)


def _narration_claude_api_epre(contexte: str) -> str:
    resp = frontiere_llm.appeler(
        modele=frontiere_llm.MODELE_ETABLI,
        max_tokens=12000,
        systeme=SYSTEM_PROMPT_EPRE,
        messages=[{"role": "user", "content": contexte}],
        cle=frontiere_llm.cle_api_ou_secrets(),
    )
    return frontiere_llm.texte_du_premier_bloc(resp)


def _generer_narration_epre(contexte: str, commentaire: str) -> Tuple[str, str]:
    try:
        txt = _narration_claude_api_epre(contexte)
        if txt:
            return txt, "claude_api"
    except Exception as e:
        logger.warning(f"Claude API : {e}")
    if commentaire and commentaire.strip():
        return _clean(commentaire), "manuel"
    return "", "aucune"


def _kpi_card(label, value, sub="", cls="") -> str:
    cls_val  = f" kpi-card-value-{cls}" if cls else ""
    cls_card = f" kpi-card-{cls}" if cls else ""
    return (
        f'<div class="kpi-card{cls_card}">'''
        f'<div class="kpi-card-label">{label}</div>'''
        f'<div class="kpi-card-value{cls_val}">{value}</div>'''
        f'<div class="kpi-card-sub">{sub}</div></div>'
    )

def _scr_card(label, value, pct, art, dominant=False) -> str:
    cls = " scr-card-dominant" if dominant else ""
    return (
        f'<div class="scr-card{cls}">'''
        f'<div class="scr-card-label">{label}</div>'''
        f'<div class="scr-card-value">{value}</div>'''
        f'<div class="scr-card-pct">{pct}</div>'''
        f'<div class="scr-card-art">{art}</div></div>'
    )

def _tricolonne_card(ref, label, value, delta="") -> str:
    return (
        f'<div class="tricolonne-card">'''
        f'<div class="tricolonne-header">'''
        f'<div class="tricolonne-ref">{ref}</div>'''
        f'<div class="tricolonne-label">{label}</div></div>'''
        f'<div class="tricolonne-body">'''
        f'<div class="tricolonne-value">{value}</div>'''
        f'<div class="tricolonne-delta">{delta}</div>'''
        f'</div></div>'
    )

def _wrap_graph(html_g: str, titre: str) -> str:
    if not html_g:
        return ""
    return (
        f'<div class="graph-wrap">'''
        f'<div class="graph-titre">{titre}</div>'''
        f'<div class="graph-body">{html_g}</div></div>'
    )

def _avis_box(avis: str, texte: str) -> str:
    a = avis.upper()
    if "DEFAV" in a or ("FAVOR" not in a and "DEF" in a):
        cls, icon, vcls = "avis-defavorable", "❌", "avis-value-defavorable"
    elif "RESERVE" in a or "RESERV" in a:
        cls, icon, vcls = "avis-reserve",    "⚠️", "avis-value-reserve"
    else:
        cls, icon, vcls = "avis-favorable",  "✅", "avis-value-favorable"
    return (
        f'<div class="avis-box {cls}">'''
        f'<div class="avis-icon">{icon}</div>'''
        f'<div>'''
        f'<div class="avis-label">Avis de la Fonction Actuarielle — Art. 48 §g Directive S2</div>'''
        f'<div class="avis-value {vcls}">{avis}</div>'''
        f'<div class="avis-text">{texte}</div>'''
        f'</div></div>'
    )


def export_html(
    result_ep1=None, result_ep2=None, result_ep3=None, result_ep4=None,
    result_ep6=None, result_ep7=None,
    commentaire="", ref_client="", arrete="",
    audit_id="", graphiques=None,
    actuaire_nom="", actuaire_numero_ia="",
) -> str:
    """Genere le rapport HTML consolide EP-RE — niveau premium A7."""
    try:
        ep1=result_ep1 or {}; ep2=result_ep2 or {}; ep3=result_ep3 or {}
        ep4=result_ep4 or {}; ep6=result_ep6 or {}; ep7=result_ep7 or {}

        dt  = datetime.now().strftime("%d/%m/%Y %H:%M")
        arr = arrete or datetime.now().strftime("%d/%m/%Y")
        cli = ref_client or "A renseigner"

        ias   = ep1.get("ias19", {}) or {}
        dbo   = float(ias.get("dbo_total", 0) or 0)
        prov  = ep3.get("provisions", {}) or {}
        pm_ep = float(prov.get("pm_encours", 0) or 0)
        ppb_ep= float(prov.get("ppb", prov.get("ppb_total", 0)) or 0)
        taux_c= float(ep3.get("taux_couverture_pct", 0) or 0)
        ratio_b= float(ep4.get("ratio_base", 0) or 0)
        ratio_l= float(ep4.get("ratio_longevite", 0) or 0)
        ratio_c= float(ep4.get("ratio_combine", 0) or 0)
        tpb   = float(ep7.get("taux_pb_optimal", 0) or 0)
        scr_pb= float(ep7.get("ratio_scr_post_pb", 0) or 0)
        avis  = ep4.get("avis_pa", ep3.get("avis_pa", "FAVORABLE AVEC RESERVE"))

        statuts = [ep1.get("statut_rag","AMBRE"), ep3.get("statut_rag","AMBRE"),
                   ep4.get("statut_rag","AMBRE"), ep6.get("statut_rag","AMBRE")]
        sg = "ROUGE" if "ROUGE" in statuts else "AMBRE" if "AMBRE" in statuts else "VERT"
        sc = _statut_cls(sg)
        dot_cls    = f"statut-dot-{sc}"
        statut_cls_str = f"garde-statut-{sc}"
        label_cls  = f"statut-label-{sc}"
        label_txt  = {"rouge":"Vigilance requise","ambre":"Reserves signalees","vert":"Favorable"}[sc]

        # Graphiques
        ghtml = {}
        if graphiques and PLOTLY_OK:
            for nom, fig_or_html in graphiques.items():
                try:
                    ghtml[nom] = fig_or_html if isinstance(fig_or_html, str) else pio.to_html(
                        fig_or_html, full_html=False, include_plotlyjs=False,
                        config={"displayModeBar": False})
                except Exception:
                    pass

        # Narration
        contexte = _construire_contexte_epre(
            ep1,ep2,ep3,ep4,ep6,ep7,arr)
        narration, source = _generer_narration_epre(contexte, commentaire)

        # KPI page de garde
        garde_kpis = (
            f'<div class="garde-kpi"><div class="kpi-label">DBO IAS 19</div>'''
            f'<div class="kpi-value highlight">{_f(dbo)}</div>'''
            f'<div class="kpi-sub">§67 PUC IAS 19</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">PM Épargne-Retraite</div>'''
            f'<div class="kpi-value">{_f(pm_ep)}</div>'''
            f'<div class="kpi-sub">Art. R342-14</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">PPB EP-RE</div>'''
            f'<div class="kpi-value">{_f(ppb_ep)}</div>'''
            f'<div class="kpi-sub">Plafond 10% PM (C2023-10)</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Couverture actifs</div>'''
            f'<div class="kpi-value">{_pct(taux_c)}</div>'''
            f'<div class="kpi-sub">Seuil ≥ 100%</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Ratio SCR post-PB</div>'''
            f'<div class="kpi-value">{_pct(scr_pb)}</div>'''
            f'<div class="kpi-sub">Art. L132-29</div></div>'''
            f'<div class="garde-kpi"><div class="kpi-label">Stress combine</div>'''
            f'<div class="kpi-value">{_pct(ratio_c)}</div>'''
            f'<div class="kpi-sub">Pire scénario stress</div></div>'
        )

        # KPI grid synthese
        kpi_grid = '<div class="kpi-grid">'
        kpi_grid += _kpi_card("DBO IAS 19", _f(dbo), "§67 PUC IAS 19")
        kpi_grid += _kpi_card("PM EP-RE", _f(pm_ep), "Art. R342-14")
        kpi_grid += _kpi_card("PPB EP-RE", _f(ppb_ep),
                               f"Ratio {_pct(ppb_ep/pm_ep*100 if pm_ep else 0)} · Plafond 10%")
        kpi_grid += _kpi_card("Couverture actifs", _pct(taux_c), "Seuil ≥ 100%",
                               cls="rouge" if taux_c < 100 else "vert")
        kpi_grid += _kpi_card("Ratio SCR post-PB", _pct(scr_pb), "Seuil ≥ 100%",
                               cls="rouge" if scr_pb < 100 else "vert")
        kpi_grid += _kpi_card("Stress combine", _pct(ratio_c), "Pire scénario",
                               cls="rouge" if ratio_c < 100 else "ambre" if ratio_c < 110 else "vert")
        kpi_grid += _kpi_card("Taux PB optimal", _pct(tpb), "Art. L132-29 (85%/90%)")
        kpi_grid += _kpi_card("Ratio base", _pct(ratio_b), "Scénario central")
        kpi_grid += '</div>'

        avis_html = _avis_box(avis, ep4.get("conseil_global",
            "Rapport produit par ActuarIA Direction EP-RE."))

        narration_html = _md_to_html(narration) if narration else (
            '<p class="comm-p" style="color:var(--slate);font-style:italic;">'''
            "Narration non disponible — alimenter EP1-EP7 et configurer ANTHROPIC_API_KEY.</p>"
        )
        ai_badge = ("❆ Narration générée par ActuarIA Intelligence"
                    if source == "claude_api" else "❆ Commentaire actuariel manuel")

        H = []
        H.append('<!DOCTYPE html>')
        H.append('<html lang="fr"><head><meta charset="UTF-8">')
        H.append(f'<title>Rapport EP-RE — {cli} — {arr}</title>')
        H.append('<script src="https://cdn.jsdelivr.net/npm/plotly.js-dist@2.26.0/plotly.min.js"></script>')
        H.append(_css())
        H.append('</head><body><div class="rapport-container">')

        # Page de garde
        H.append('<div class="page-garde">')
        H.append('<div class="garde-bg"><div class="garde-dots"></div>')
        H.append('<div class="garde-diagonal"></div><div class="garde-accent-line"></div></div>')
        H.append('<div class="garde-inner">')
        H.append(f'<div class="garde-header"><div class="garde-logo-wrap"><img src="{LOGO_URI}" alt="ActuarIA"/></div>')
        H.append('<div class="garde-badges"><span class="badge-confidentiel">⬛ Confidentiel</span>')
        H.append('<span class="badge-direction">Épargne-Retraite EP-RE</span></div></div>')
        H.append('<div class="garde-hero">')
        H.append('<div class="garde-eyebrow">Rapport Consolide Actuariel</div>')
        H.append('<div class="garde-titre">Épargne &amp; <em>Retraite</em></div>')
        H.append(f'<div class="garde-subtitle">Arrêté au {arr}</div>')
        H.append(f'<div class="garde-sousdirection">Art. 39 / Art. 83 / PER · IAS 19 · Solvabilité 2 · {cli}</div>')
        H.append('<div class="garde-sep"><div class="garde-sep-line"></div>')
        H.append('<div class="garde-sep-diamond"></div><div class="garde-sep-line"></div></div>')
        H.append(f'<div class="garde-statut {statut_cls_str}"><div class="statut-dot {dot_cls}"></div>')
        H.append(f'<div class="{label_cls}">{label_txt}</div></div>')
        H.append('</div>')
        H.append(f'<div class="garde-footer"><div class="garde-kpis">{garde_kpis}</div></div>')
        H.append('</div></div>')

        H.append('<div class="rapport-body">')

        # §1
        H.append('<div class="section-header"><span class="section-num">01</span>')
        H.append('<span class="section-titre">Synthèse exécutive</span></div>')
        H.append(f'<div class="section-body">{kpi_grid}')
        H.append(_wrap_graph(ghtml.get("scorecard_ep5",""), "Scorecard Rapport EP5 — Validation C1/C2/C3"))
        H.append(avis_html)
        H.append('</div><div class="section-divider"></div>')

        # §2
        H.append('<div class="section-header"><span class="section-num">02</span>')
        H.append('<span class="section-titre">Engagements IAS 19 — DBO et Décomposition</span></div>')
        H.append(f'<div class="section-body"><div class="kpi-grid">')
        H.append(_kpi_card("DBO totale", _f(dbo), "§67 PUC IAS 19"))
        H.append(_kpi_card("Service Cost", _f(float(ias.get("service_cost",0) or 0)), "Droit acquis exercice"))
        H.append(_kpi_card("Interest Cost", _f(float(ias.get("interest_cost",0) or 0)), "Désactualisation"))
        H.append(_kpi_card("Taux iBoxx AA", _pct(float(ias.get("taux_actu",0) or 0)), "§83 IAS 19"))
        H.append('</div>')
        H.append(_wrap_graph(ghtml.get("graphiques_ias19",""), "IAS 19 — DBO, Service Cost, Sensibilités"))
        H.append('</div><div class="section-divider"></div>')

        # §3
        H.append('<div class="section-header"><span class="section-num">03</span>')
        H.append('<span class="section-titre">Provisions et Couverture du Régime</span></div>')
        H.append(f'<div class="section-body"><div class="kpi-grid">')
        H.append(_kpi_card("PM EP-RE", _f(pm_ep), "Art. R342-14"))
        H.append(_kpi_card("PPB EP-RE", _f(ppb_ep),
                            f"Plafond 10% PM · C2023-10"))
        H.append(_kpi_card("Couverture actifs", _pct(taux_c), "Seuil ≥ 100%",
                            cls="rouge" if taux_c < 100 else "vert"))
        H.append(_kpi_card("Réserve capitalisation", _f(float(prov.get("reserve_capi",0) or 0)), ""))
        H.append('</div>')
        H.append(_wrap_graph(ghtml.get("graphiques_ep3",""), "Provisions EP-RE — PM, PPB, Couverture"))
        H.append('</div><div class="section-divider"></div>')

        # §4
        H.append('<div class="section-header"><span class="section-num">04</span>')
        H.append('<span class="section-titre">Analyse SCR et Exigences Solvabilité 2</span></div>')
        H.append(f'<div class="section-body">')
        H.append(_wrap_graph(ghtml.get("graphiques_scr",""), "SCR Longevité — Art. 138 S2"))
        H.append('</div><div class="section-divider"></div>')

        # §5
        H.append('<div class="section-header"><span class="section-num">05</span>')
        H.append('<span class="section-titre">Optimisation de la Participation aux Bénéfices</span></div>')
        H.append(f'<div class="section-body"><div class="kpi-grid">')
        H.append(_kpi_card("Taux PB optimal", _pct(tpb), "Art. L132-29 (85%/90%)"))
        H.append(_kpi_card("PPB finale", _f(float(ep7.get("ppb_finale",0) or 0)),
                            f"Plafond C2023-10 : 10% PM"))
        H.append(_kpi_card("SCR post-PB", _pct(scr_pb), "Seuil ≥ 100%",
                            cls="rouge" if scr_pb < 100 else "vert"))
        H.append(_kpi_card("PB min. régl.",
                            _f(float(ep7.get("pb_minimale_reglementaire",0) or 0)),
                            "85% prod. fin."))
        H.append('</div>')
        H.append(_wrap_graph(ghtml.get("graphiques_ep7",""), "Optimisation PB — Projection PPB pluriannuelle"))
        H.append('</div><div class="section-divider"></div>')

        # §6
        H.append('<div class="section-header"><span class="section-num">06</span>')
        H.append('<span class="section-titre">Analyse de Stress et Sensibilités</span></div>')
        H.append(f'<div class="section-body"><div class="kpi-grid">')
        H.append(_kpi_card("Ratio base", _pct(ratio_b), "Scénario central"))
        H.append(_kpi_card("Post-longévité", _pct(ratio_l), "+5 ans",
                            cls="rouge" if ratio_l < 100 else "ambre" if ratio_l < 110 else "vert"))
        H.append(_kpi_card("Stress combiné", _pct(ratio_c), "Pire cas",
                            cls="rouge" if ratio_c < 100 else "ambre" if ratio_c < 110 else "vert"))
        H.append(_kpi_card("Sens. DBO -100bp", _pct(float(ias.get("sensibilite_dbo_taux_pct",0) or 0)), "§83 IAS 19"))
        H.append('</div>')
        H.append(_wrap_graph(ghtml.get("graphiques_stress",""), "Stress EP-RE — Scénarios longévité, inflation, taux"))
        H.append('</div><div class="section-divider"></div>')

        # §7
        H.append('<div class="section-header"><span class="section-num">07</span>')
        H.append('<span class="section-titre">Back-Testing des Hypothèses Actuarielles (IAS 19.145)</span></div>')
        H.append(f'<div class="section-body">')
        H.append(_wrap_graph(ghtml.get("graphiques_bt",""), "Écarts d'expérience — Impact DBO par paramètre"))
        H.append('</div><div class="section-divider"></div>')

        # §8
        H.append('<div class="section-header"><span class="section-num">08</span>')
        H.append('<span class="section-titre">Commentaire Actuariel — Avis de la Fonction Actuarielle</span></div>')
        H.append('<div class="section-body">')
        H.append('<div class="commentaire-wrap">')
        H.append('<div class="commentaire-header">')
        H.append('<div class="commentaire-header-title">Analyse actuarielle consolidée EP-RE</div>')
        H.append(f'<div class="commentaire-ai-badge">{ai_badge}</div></div>')
        H.append('<div class="commentaire-body">')
        H.append(narration_html)
        H.append(f'<div class="comm-footer">Rédigé le {dt} · Actuaire : {actuaire_nom or "A renseigner"} · N° IA : {actuaire_numero_ia or "—"} · CONFIDENTIEL</div>')
        H.append('</div></div></div>')

        H.append('</div>')  # fin rapport-body

        H.append('<div class="pied-de-page">')
        H.append(f'<div class="pied-logo"><img src="{LOGO_URI}" alt="ActuarIA"/></div>')
        H.append(f'<div class="pied-meta"><div>{cli} · Arrêté {arr}</div>')
        H.append(f'<div>Audit ID : {audit_id or "—"} · Généré le {dt}</div>')
        H.append('<div class="confidentiel-footer">⬛ CONFIDENTIEL — USAGE INTERNE EXCLUSIF</div>')
        H.append('</div></div>')
        H.append('</div></body></html>')

        html = "\n".join(H)
        logger.info(f"HTML EP-RE genere : {len(html):,} chars, source={source}")
        return html

    except Exception as e:
        logger.error(f"export_html EP-RE : {e}", exc_info=True)
        return f"<html><body><h1>Erreur</h1><p>{e}</p></body></html>"


def export_pdf(**kwargs) -> bytes:
    """Genere le PDF du rapport Vie via WeasyPrint."""
    try:
        if not WEASYPRINT_OK:
            raise ImportError("weasyprint requis")
        html_str = export_html(**kwargs)
        pdf = WP_HTML(string=html_str).write_pdf()
        logger.info(f"PDF Vie : {len(pdf):,} bytes")
        return pdf
    except Exception as e:
        logger.error(f"export_pdf Vie : {e}", exc_info=True)
        return b""


def export_word(
    result_ep1=None, result_ep2=None, result_ep3=None, result_ep4=None,
    result_ep6=None, result_ep7=None,
    commentaire="", ref_client="", arrete="",
    audit_id="", actuaire_nom="", actuaire_numero_ia="",
) -> bytes:
    """Genere le rapport Word professionnel EP-RE — 8 chapitres."""
    if not DOCX_OK:
        logger.error("python-docx absent"); return b""
    try:
        def rgb(h):
            h = h.lstrip("#")
            return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:],16))
        NR=rgb(NAVY); GR=rgb(GOLD); BR=rgb(WHITE)
        SR=rgb(SLATE); RgR=rgb(ROUGE); VR=rgb(VERT); AR=rgb(ORANGE)

        dt  = datetime.now().strftime("%d/%m/%Y")
        arr = arrete or dt
        cli = ref_client or "A renseigner"

        ep1=result_ep1 or {}; ep2=result_ep2 or {}; ep3=result_ep3 or {}
        ep4=result_ep4 or {}; ep6=result_ep6 or {}; ep7=result_ep7 or {}

        ias   = ep1.get("ias19", {}) or {}
        dbo   = float(ias.get("dbo_total", 0) or 0)
        prov  = ep3.get("provisions", {}) or {}
        pm_ep = float(prov.get("pm_encours", 0) or 0)
        ppb_ep= float(prov.get("ppb", prov.get("ppb_total", 0)) or 0)
        taux_c= float(ep3.get("taux_couverture_pct", 0) or 0)
        ratio_b= float(ep4.get("ratio_base", 0) or 0)
        ratio_l= float(ep4.get("ratio_longevite", 0) or 0)
        ratio_c= float(ep4.get("ratio_combine", 0) or 0)
        scr_pb = float(ep7.get("ratio_scr_post_pb", 0) or 0)
        avis  = ep4.get("avis_pa", ep3.get("avis_pa", "FAVORABLE AVEC RESERVE"))
        ecarts= ep6.get("ecarts", {}) or {}

        contexte = _construire_contexte_epre(ep1,ep2,ep3,ep4,ep6,ep7,arr)
        narration, source = _generer_narration_epre(contexte, commentaire)

        doc = Document()
        for s in doc.sections:
            s.top_margin=Cm(2.2); s.bottom_margin=Cm(2.2)
            s.left_margin=Cm(2.8); s.right_margin=Cm(2.8)

        def _bg(cell, hex6):
            tc=cell._tc; tcp=tc.get_or_add_tcPr()
            sd=OxmlElement("w:shd")
            sd.set(qn("w:fill"), hex6.lstrip("#"))
            sd.set(qn("w:color"),"auto"); sd.set(qn("w:val"),"clear")
            tcp.append(sd)

        def _run(p, txt, bold=False, italic=False, sz=10, col=None):
            r=p.add_run(str(txt)); r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
            if col: r.font.color.rgb=col
            return r

        def _h(txt, lv=1, col=None):
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(10 if lv==1 else 6)
            p.paragraph_format.space_after=Pt(4)
            sz={1:16,2:12,3:10}.get(lv,10)
            c=col or (NR if lv==1 else GR)
            _run(p, txt, bold=True, sz=sz, col=c)

        def _sep():
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(3); p.paragraph_format.space_after=Pt(3)
            pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement("w:pBdr")
            bo=OxmlElement("w:bottom")
            bo.set(qn("w:val"),"single"); bo.set(qn("w:sz"),"6")
            bo.set(qn("w:space"),"1"); bo.set(qn("w:color"),"C9A84C")
            pBdr.append(bo); pPr.append(pBdr)

        def _tbl(heads, rows, ws=None):
            from docx.enum.table import WD_ALIGN_VERTICAL
            t=doc.add_table(rows=1+len(rows), cols=len(heads)); t.style="Table Grid"
            for i,hd in enumerate(heads):
                c=t.rows[0].cells[i]; _bg(c,"0B1E3D")
                pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=pp.add_run(str(hd)); r.bold=True; r.font.size=Pt(9); r.font.color.rgb=BR
            for ri,row in enumerate(rows):
                for ci,v in enumerate(row):
                    c=t.rows[ri+1].cells[ci]
                    if ri%2==1: _bg(c,"EEF2F7")
                    pp=c.paragraphs[0]; pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
                    r=pp.add_run(str(v) if v is not None else "—"); r.font.size=Pt(9)
            if ws:
                for i,w in enumerate(ws):
                    for row in t.rows: row.cells[i].width=Cm(w)
            doc.add_paragraph().paragraph_format.space_after=Pt(2)

        # Page de garde
        p=doc.add_paragraph()
        _run(p,"Actuar",bold=True,sz=28,col=NR); _run(p,"IA",bold=True,sz=28,col=GR)
        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p,"RAPPORT CONSOLIDE EP-RE\n",bold=True,sz=22,col=NR)
        _run(p,"Epargne-Retraite — Art. 39 / Art. 83 / PER",bold=True,sz=14,col=GR)
        doc.add_paragraph()
        avis_col = VR if "FAVORABLE" in avis.upper() and "RESERVE" not in avis.upper() else (
                   AR if "RESERVE" in avis.upper() else RgR)
        p=doc.add_paragraph()
        _run(p,"Avis Fonction Actuarielle : ",sz=11,col=NR)
        _run(p,avis,bold=True,sz=11,col=avis_col)
        doc.add_paragraph()
        _tbl(["Client","Regime","Arrete","Actuaire"],
             [[cli,"Art. 39 / Art. 83 / PER",arr,actuaire_nom or "A renseigner"]],
             ws=[3.5,4.0,3.0,5.5])
        doc.add_page_break()

        # §1
        _h("§1 — Contexte, Perimetre et Gouvernance du Regime"); _sep()
        tar = ep2.get("tarification",{}) or {}
        _tbl(["Indicateur","Valeur","Reference"],
             [["Regime","Art. 39 / Art. 83 / PER","Code assurances"],
              ["Cotisation annuelle brute",_f(float(tar.get("cotisation_brute_annuelle",0) or 0)),"EP2"],
              ["Rente mensuelle estimee",_f(float(tar.get("rente_mensuelle",0) or 0),dec=2),"EP2"],
              ["Taux de remplacement",_pct(float(tar.get("taux_remplacement_pct",0) or 0)),"EP2"],
              ["Date arrete",arr,"—"]],
             ws=[5.5,4.0,6.5])
        doc.add_page_break()

        # §2
        _h("§2 — Engagements IAS 19 — DBO et Decomposition"); _sep()
        _tbl(["Indicateur","Valeur","Reference","Seuil / Note"],
             [["DBO totale",_f(dbo),"§67 PUC IAS 19","—"],
              ["Service Cost (droit acquis)",_f(float(ias.get("service_cost",0) or 0)),"§67","—"],
              ["Interest Cost (desactualisation)",_f(float(ias.get("interest_cost",0) or 0)),"§67","—"],
              ["Taux actualisation iBoxx AA",_pct(float(ias.get("taux_actu",0) or 0)),"§83","±25bp tolerance"],
              ["Sensibilite DBO a -100bp",_pct(float(ias.get("sensibilite_dbo_taux_pct",0) or 0)),"§83 D_mod","Alerter si > 15%"],
              ["Conformite iBoxx",ep1.get("iboxx_audit",{}).get("conformite_iboxx","Non verifiable"),"§83","Obligatoire"]],
             ws=[5.0,3.5,3.5,4.0])
        doc.add_page_break()

        # §3
        _h("§3 — Provisions et Couverture du Regime"); _sep()
        _tbl(["Indicateur","Valeur","Reference","Seuil"],
             [["PM epargne-retraite",_f(pm_ep),"Art. R342-14 / ä_x x rente","—"],
              ["PPB EP-RE",_f(ppb_ep),"Art. R342-14","Plafond 10% PM (C2023-10)"],
              ["Ratio PPB/PM",_pct(ppb_ep/pm_ep*100 if pm_ep else 0),"C2023-10","≤ 10%"],
              ["Reserve capitalisation",_f(float(prov.get("reserve_capi",0) or 0)),"Art. R342-14","—"],
              ["Taux couverture actifs/PM",_pct(taux_c),"Art. R342-14","≥ 100%"]],
             ws=[5.0,3.5,4.0,3.5])
        doc.add_page_break()

        # §4
        _h("§4 — Analyse SCR et Exigences Solvabilite 2"); _sep()
        p=doc.add_paragraph()
        _run(p,"SCR longevite (Art. 138 S2 : -20% qx) : module dominant pour les regimes de retraite a prestations definies.",
             sz=9,col=NR)
        doc.add_page_break()

        # §5
        _h("§5 — Optimisation de la Participation aux Benefices"); _sep()
        tpb_opt = float(ep7.get("taux_pb_optimal",0) or 0)
        ppb_fin = float(ep7.get("ppb_finale",0) or 0)
        pb_min  = float(ep7.get("pb_minimale_reglementaire",0) or 0)
        _tbl(["Indicateur PB","Valeur","Reference"],
             [["Taux PB optimal retenu",_pct(tpb_opt),"Art. L132-29"],
              ["PB minimale reglementaire",_f(pb_min),"85% prod. fin. / 90% ben. tech."],
              ["PPB finale apres optimisation",_f(ppb_fin),"Plafond 10% PM (C2023-10)"],
              ["Ratio SCR post-distribution",_pct(scr_pb),"Seuil >= 100%"],
              ["Statut optimisation",ep7.get("statut_rag","—"),"EP7"]],
             ws=[5.5,3.5,7.0])
        doc.add_page_break()

        # §6
        _h("§6 — Analyse de Stress et Sensibilites"); _sep()
        _tbl(["Scenario","Ratio couverture","Delta vs base","Reference"],
             [["Scenario base",_pct(ratio_b),"—","EP4"],
              ["Stress longevite (+5 ans)",_pct(float(ep4.get("ratio_longevite",0) or 0)),
               _pct(float(ep4.get("ratio_longevite",0) or 0)-ratio_b),"Art. 138 S2"],
              ["Stress taux (-100bp)","— (DBO +"+_pct(float(ias.get("sensibilite_dbo_taux_pct",0) or 0))+")",
               "—","§83 IAS 19"],
              ["Scenario combine (pire cas)",_pct(ratio_c),
               _pct(ratio_c-ratio_b),"EP4"]],
             ws=[5.0,4.0,3.5,3.5])
        doc.add_page_break()

        # §7
        _h("§7 — Back-Testing des Hypotheses Actuarielles (IAS 19.145)"); _sep()
        rows_bt = [
            [param,str(ecarts.get(f"{param}_bps","—"))+" bp",
             _f(float(ep6.get(f"impact_{param}",0) or 0)),
             ep6.get(f"statut_{param}","—")]
            for param in ["mortalite","rotation","revalorisation","taux"]
        ]
        rows_bt.append(["Impact total DBO","—",
                        _f(float(ep6.get("impact_dbo_total",0) or 0)),
                        ep6.get("statut_rag","—")])
        _tbl(["Parametre","Ecart observe","Impact DBO","Statut"],
             rows_bt,ws=[4.0,3.0,3.5,5.5])
        doc.add_page_break()

        # §8
        _h("§8 — Commentaire Actuariel — Avis de la Fonction Actuarielle"); _sep()
        if narration:
            secs = re.split(r"(?=§\d+\s*[-—–])", _clean(narration))
            for sec in secs:
                sec = sec.strip()
                if not sec: continue
                ls = sec.split("\n",1)
                if ls[0]: _h(ls[0],lv=2)
                if len(ls) > 1:
                    for ln in ls[1].split("\n"):
                        ln = ln.strip()
                        if ln:
                            p=doc.add_paragraph()
                            p.paragraph_format.space_after=Pt(3)
                            p.paragraph_format.left_indent=Cm(0.4)
                            _run(p,ln,sz=9,col=NR)
            if source=="claude_api":
                p=doc.add_paragraph()
                _run(p,"❆ Narration generee par ActuarIA Intelligence",sz=7,italic=True,col=SR)
        else:
            p=doc.add_paragraph()
            _run(p,"Narration non disponible — configurer ANTHROPIC_API_KEY.",sz=9,italic=True,col=SR)

        doc.add_paragraph()
        p=doc.add_paragraph()
        _run(p,"Avis Fonction Actuarielle (Art. 48 §g Directive 2009/138/CE) : ",sz=10,col=NR)
        _run(p,avis,bold=True,sz=12,col=avis_col)

        _sep()
        p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        _run(p,f"ActuarIA · {cli} · EP-RE · Arrêté {arr} · {dt} · CONFIDENTIEL",
             sz=7,italic=True,col=SR)

        buf=io.BytesIO(); doc.save(buf); buf.seek(0)
        wb=buf.read()
        logger.info(f"Word EP-RE : {len(wb):,} bytes")
        return wb

    except Exception as e:
        logger.error(f"export_word EP-RE : {e}", exc_info=True)
        return b""
